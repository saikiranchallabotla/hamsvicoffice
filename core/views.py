# core/views.py

import json
import os
import re
import logging
from copy import copy

import inflect
from docx import Document
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, Border, Side, PatternFill
from django.utils import timezone
from django.urls import reverse
from django.core.files.uploadedfile import InMemoryUploadedFile
from django.contrib.auth.decorators import login_required

from django.conf import settings
from django.http import HttpResponse, JsonResponse, HttpResponseNotAllowed
from django.shortcuts import render, redirect, get_object_or_404
from django.template.loader import render_to_string
from django.views.decorators.http import require_POST
from django.utils.crypto import get_random_string

import io
from io import BytesIO
from difflib import SequenceMatcher
# (moved to top-level imports)

from .models import Project, SelfFormattedTemplate, Estimate, Organization, Membership, Upload, Job, OutputFile
from .decorators import org_required, role_required

logger = logging.getLogger(__name__)
from .tasks import process_excel_upload, generate_bill_pdf, generate_workslip_pdf, generate_bill_document_task
from .utils_excel import load_backend, copy_block_with_styles_and_formulas, build_temp_day_rates

p_engine = inflect.engine()

# Define BILL_TEMPLATES_DIR for template operations
BILL_TEMPLATES_DIR = os.path.join(settings.BASE_DIR, "core", "templates", "core", "bill_templates")

_inflect_engine = inflect.engine()

# ============================================================================
# PHASE 3B: HELPER FUNCTIONS FOR ORGANIZATION & ASYNC PROCESSING
# ============================================================================

def get_org_from_request(request):
    """
    Safely extract organization from request.
    
    Returns:
        Organization object (auto-creates if not available)
        
    For single-tenant/development mode, automatically creates
    a default organization and membership for the user.
    """
    if hasattr(request, 'organization') and request.organization:
        return request.organization
    
    # Auto-create organization for logged-in users (single-tenant mode)
    if request.user.is_authenticated:
        from core.models import Organization, Membership
        from django.utils.text import slugify
        
        # Try to find existing membership
        membership = Membership.objects.filter(user=request.user).select_related('organization').first()
        
        if membership:
            request.organization = membership.organization
            return membership.organization
        
        # Create default organization for user
        org_name = f"{request.user.username}'s Organization"
        org_slug = slugify(org_name)[:255]
        
        # Make slug unique if needed
        base_slug = org_slug
        counter = 1
        while Organization.objects.filter(slug=org_slug).exists():
            org_slug = f"{base_slug}-{counter}"
            counter += 1
        
        org, created = Organization.objects.get_or_create(
            name=org_name,
            defaults={
                'slug': org_slug,
                'owner': request.user,
                'is_active': True
            }
        )
        
        # Create membership
        Membership.objects.get_or_create(
            user=request.user,
            organization=org,
            defaults={'role': 'owner'}
        )
        
        request.organization = org
        return org
    
    # Not authenticated - raise error
    from django.http import Http404
    raise Http404("Please login to continue.")


def check_org_access(request, obj):
    """
    Verify that an object belongs to the user's organization.
    
    Args:
        request: Django request object (has organization from middleware)
        obj: Model instance to check (must have organization FK)
        
    Returns:
        True if object belongs to user's org
        
    Raises:
        Http404 if object doesn't belong to user's org
    """
    org = get_org_from_request(request)
    if hasattr(obj, 'organization') and obj.organization != org:
        from django.http import Http404
        raise Http404("You don't have permission to access this object.")
    return True


def enqueue_excel_task(job_id, task_name='excel_parse', **kwargs):
    """
    Enqueue an Excel processing task.
    Runs synchronously if Celery is not available.
    
    Args:
        job_id: ID of Job object to update
        task_name: 'excel_parse', 'generate_bill', 'generate_workslip'
        **kwargs: Additional arguments for the task (must include upload_id for excel_parse)
        
    Returns:
        Task result or mock object with id
    """
    from django.conf import settings
    
    # Get upload_id from kwargs or from job
    upload_id = kwargs.get('upload_id')
    if not upload_id:
        try:
            job = Job.objects.get(id=job_id)
            if job.upload:
                upload_id = job.upload.id
        except Job.DoesNotExist:
            pass
    
    # Try Celery first, fall back to sync if connection fails
    try:
        if task_name == 'generate_bill':
            task = generate_bill_pdf.delay(job_id, kwargs.get('project_id'))
        elif task_name == 'generate_workslip':
            task = generate_workslip_pdf.delay(job_id, kwargs.get('project_id'))
        elif task_name == 'generate_bill_document':
            task = generate_bill_document_task.delay(job_id)
        elif upload_id:
            task = process_excel_upload.delay(upload_id)
        else:
            raise ValueError("upload_id required for excel_parse task")
        return task
    except Exception as e:
        # Celery not available, run synchronously
        logger.warning(f"Celery not available ({e}), running task synchronously")
        
        # Create a mock task result
        class MockTask:
            def __init__(self):
                import uuid
                self.id = str(uuid.uuid4())
        
        # Run task synchronously
        if task_name == 'generate_bill':
            generate_bill_pdf(job_id, kwargs.get('project_id'))
        elif task_name == 'generate_workslip':
            generate_workslip_pdf(job_id, kwargs.get('project_id'))
        elif task_name == 'generate_bill_document':
            generate_bill_document_task(job_id)
        elif upload_id:
            process_excel_upload(upload_id)
        
        return MockTask()


def create_job_for_excel(request, upload=None, job_type='excel_parse', metadata=None):
    """
    Create a Job object and enqueue task for Excel processing.
    
    Args:
        request: Django request (has organization)
        upload: Upload object (optional, will be created if not provided)
        job_type: Type of job ('excel_parse', 'generate_bill', 'generate_workslip')
        metadata: Dict with additional metadata
        
    Returns:
        Tuple: (job, celery_task)
        
    Example:
        job, task = create_job_for_excel(request, upload, 'generate_bill', {'project_id': 123})
        return JsonResponse({'job_id': job.id, 'status_url': reverse('job_status', args=[job.id])})
    """
    org = get_org_from_request(request)
    
    # Create upload if not provided
    if not upload:
        upload = Upload.objects.create(
            filename=metadata.get('filename', 'export.xlsx') if metadata else 'export.xlsx',
            file_size=0,
            status='processing'
        )
    
    # Create job
    job = Job.objects.create(
        organization=org,
        user=request.user,
        upload=upload,
        job_type=job_type,
        status='queued',
        result={'metadata': metadata or {}}  # Store metadata in result field
    )
    
    # Enqueue task
    task_kwargs = {}
    if metadata:
        task_kwargs.update({k: v for k, v in metadata.items() if k not in ('filename', 'job_id')})
    
    task = enqueue_excel_task(job.id, job_type, **task_kwargs)
    
    # Store task ID in job
    job.celery_task_id = task.id
    job.save()
    
    return job, task


# ============================================================================
# END PHASE 3B HELPERS
# ============================================================================


def _number_to_words_rupees(n):
    """
    Convert number to words in Indian numbering system:
      12345.67 -> 'Twelve thousand three hundred and forty-five rupees only'
      1234567 -> 'Twelve lakh thirty-four thousand five hundred and sixty-seven rupees only'
    """
    try:
        integer_part = int(round(float(n)))
    except Exception:
        integer_part = 0

    if integer_part == 0:
        return "Zero rupees only"

    # Indian numbering system
    crores = integer_part // 10000000
    integer_part %= 10000000
    lakhs = integer_part // 100000
    integer_part %= 100000
    thousands = integer_part // 1000
    integer_part %= 1000
    hundreds = integer_part

    parts = []
    if crores > 0:
        parts.append(f"{_inflect_engine.number_to_words(crores)} crore" + ("s" if crores > 1 else ""))
    if lakhs > 0:
        parts.append(f"{_inflect_engine.number_to_words(lakhs)} lakh" + ("s" if lakhs > 1 else ""))
    if thousands > 0:
        parts.append(f"{_inflect_engine.number_to_words(thousands)} thousand")
    if hundreds > 0:
        parts.append(_inflect_engine.number_to_words(hundreds))

    words = " ".join(parts)
    words = words.replace("-", " ")
    return f"{words} rupees only".capitalize()


def _format_indian_number(num):
    """
    Format a number in Indian numbering system with commas.
    Example: 300000 -> "3,00,000" (lakhs), 10000000 -> "1,00,00,000" (crores)
    
    Indian system: 1,00,00,00,000 (crores, lakhs, thousands, hundreds)
    Pattern: last 3 digits, then groups of 2
    """
    try:
        # Handle string input
        if isinstance(num, str):
            num = float(num.replace(',', '').replace(' ', ''))
        
        num = float(num)
        
        # Handle negative numbers
        is_negative = num < 0
        num = abs(num)
        
        # Check if it has decimals
        if num == int(num):
            # No decimals - format as integer
            num_str = str(int(num))
        else:
            # Has decimals - format with 2 decimal places
            num_str = f"{num:.2f}"
            integer_part, decimal_part = num_str.split('.')
            # Format the integer part only
            num_str = integer_part
        
        # Apply Indian comma formatting
        if len(num_str) <= 3:
            formatted = num_str
        else:
            # Last 3 digits
            last_three = num_str[-3:]
            remaining = num_str[:-3]
            
            # Group remaining by 2
            groups = []
            while remaining:
                groups.append(remaining[-2:] if len(remaining) >= 2 else remaining)
                remaining = remaining[:-2]
            
            groups.reverse()
            formatted = ','.join(groups) + ',' + last_three
        
        # Add decimal part back if it exists
        if num != int(num):
            formatted += f".{f'{num:.2f}'.split('.')[1]}"
        
        result = f"-{formatted}" if is_negative else formatted
        return result
        
    except (ValueError, TypeError):
        return str(num)


def _get_current_financial_year():
    """
    Get the current financial year in format "2025-26".
    Financial year runs from April 1 to March 31.
    """
    today = timezone.now().date()
    
    if today.month >= 4:  # April onwards
        fy_start = today.year
        fy_end = (today.year + 1) % 100
    else:  # January to March
        fy_start = today.year - 1
        fy_end = today.year % 100
    
    return f"{fy_start}-{fy_end:02d}"


def _get_current_date_formatted():
    """
    Get current date in format "DD-MM-YYYY"
    """
    today = timezone.now().date()
    return today.strftime("%d-%m-%Y")


def home(request):
    """Home page - shows login/register for guests, modules for authenticated users"""
    if not request.user.is_authenticated:
        # Show landing page with login/register options
        return render(request, "core/landing.html")
    try:
        return render(request, "core/home.html")
    except Exception:
        return HttpResponse("Home page temporarily unavailable.")


@login_required(login_url='login')
def workslip_home(request):
    """
    Landing page for Workslip module:
    - Step 1: Select work type (New Estimate / AMC / Temporary Works)
    - Step 2: Select work mode (Original / Repair)
    - Step 3: Select category (Electrical / Civil)
    Then redirect to workslip main page with parameters
    """
    return render(request, "core/workslip_home.html")


@login_required(login_url='login')
def workslip(request):
    from core.utils_excel import get_available_backends_for_module
    
    # ---- Clear session data on fresh page load (GET without preserve param) ----
    if request.method == "GET" and not request.GET.get("preserve") and not request.GET.get("group"):
        # Clear all workslip session data for a fresh start
        request.session["ws_estimate_rows"] = []
        request.session["ws_exec_map"] = {}
        request.session["ws_tp_percent"] = 0.0
        request.session["ws_tp_type"] = "Excess"
        request.session["ws_supp_items"] = []
        request.session["ws_estimate_grand_total"] = 0.0
        request.session["ws_work_name"] = ""
        request.session["ws_current_phase"] = 1
        request.session["ws_previous_phases"] = []
        request.session["ws_previous_supp_items"] = []
        request.session["ws_previous_ae_data"] = []
        request.session["ws_metadata"] = {}  # Clear workslip metadata
        request.session["ws_deduct_old_material"] = 0.0
        request.session["ws_lc_percent"] = 0.0
        request.session["ws_qc_percent"] = 0.0
        request.session["ws_nac_percent"] = 0.0
        request.session["ws_selected_backend_id"] = None  # Clear backend selection
        request.session["ws_work_type"] = None  # Clear work type selection
        request.session["ws_work_mode"] = None  # Clear work mode selection
        request.session["ws_category"] = None  # Clear category selection
    
    # Handle work type, work mode and category from URL parameters (from workslip_home)
    url_work_type = request.GET.get("work_type")
    url_work_mode = request.GET.get("work_mode")
    url_category = request.GET.get("category")
    
    if url_work_type:
        request.session["ws_work_type"] = url_work_type
    if url_work_mode:
        request.session["ws_work_mode"] = url_work_mode
    if url_category:
        request.session["ws_category"] = url_category
    
    # Get work type, work mode and category from session
    ws_work_type = request.session.get("ws_work_type", "new_estimate") or "new_estimate"
    ws_work_mode = request.session.get("ws_work_mode", "original") or "original"
    ws_category = request.session.get("ws_category", "electrical") or "electrical"
    
    # Handle backend switching via URL parameter
    url_backend_id = request.GET.get("backend_id")
    if url_backend_id:
        try:
            request.session["ws_selected_backend_id"] = int(url_backend_id)
        except (ValueError, TypeError):
            pass
    
    ws_selected_backend_id = request.session.get("ws_selected_backend_id")
    
    # ---- session state ----
    ws_estimate_rows = request.session.get("ws_estimate_rows", []) or []
    ws_exec_map = request.session.get("ws_exec_map", {}) or {}
    ws_tp_percent = request.session.get("ws_tp_percent", 0.0)
    ws_tp_type = request.session.get("ws_tp_type", "Excess")
    ws_supp_items = request.session.get("ws_supp_items", []) or []
    ws_estimate_grand_total = request.session.get("ws_estimate_grand_total", 0.0)
    ws_work_name = request.session.get("ws_work_name", "") or ""
    ws_deduct_old_material = request.session.get("ws_deduct_old_material", 0.0)
    
    # Log session values for debugging
    logger.info(f"[WORKSLIP] Session: work_type={ws_work_type}, category={ws_category}, tp_percent={ws_tp_percent}, tp_type={ws_tp_type}, deduct={ws_deduct_old_material}")
    
    # Multi-phase workslip tracking
    ws_current_phase = request.session.get("ws_current_phase", 1)  # Current phase number (1, 2, 3, etc.)
    ws_previous_phases = request.session.get("ws_previous_phases", [])  # List of previous phase exec_maps
    ws_previous_supp_items = request.session.get("ws_previous_supp_items", [])  # Supplemental items from previous phases

    # Determine module code and categories based on work type
    # - backend_db_category: category stored in ModuleBackend table (electrical/civil)
    # - load_category: category prefix used by load_backend for file paths (amc_electrical, temp_electrical, etc.)
    if ws_work_type == 'amc':
        module_code = 'amc'
        backend_db_category = ws_category  # stored as 'electrical' or 'civil' in DB
        load_category = f'amc_{ws_category}'  # used for file loading
    elif ws_work_type == 'tempworks':
        module_code = 'temp_works'
        backend_db_category = ws_category  # stored as 'electrical' or 'civil' in DB
        load_category = f'temp_{ws_category}'  # used for file loading
    else:  # new_estimate (default)
        module_code = 'new_estimate'
        backend_db_category = ws_category
        load_category = ws_category
    
    category = load_category  # Use load_category for loading backend files

    # Get available backends for dropdown (use DB category format)
    available_backends = get_available_backends_for_module(module_code, backend_db_category)

    # -------- 0) Load backend (groups/items + Master Datas) ----------
    desc_to_item = {}
    try:
        items_list, groups_map, units_map, ws_data, filepath = load_backend(
            category, settings.BASE_DIR,
            backend_id=ws_selected_backend_id,
            module_code=module_code,
            user=request.user
        )
        if ws_data is not None:
            for info in items_list:
                item_name = info["name"]
                start_row = info["start_row"]
                desc_cell = ws_data.cell(row=start_row + 2, column=4).value
                desc_text = str(desc_cell or "").strip()
                if desc_text:
                    desc_to_item.setdefault(desc_text, item_name)
    except Exception:
        items_list, groups_map, ws_data, filepath = [], {}, None, ""

    groups = sorted(groups_map.keys(), key=lambda s: s.lower()) if groups_map else []
    current_group = request.GET.get("group") or (groups[0] if groups else "")

    group_items = groups_map.get(current_group, []) if current_group else []
    detected_names = {i["name"] for i in items_list}
    items_in_group = [name for name in group_items if name in detected_names]

    # helper: group â†’ units
    item_to_group = {}
    for grp_name, item_list_grp in groups_map.items():
        for nm in item_list_grp:
            item_to_group.setdefault(nm, grp_name)

    def units_for(name):
        grp = item_to_group.get(name, "")
        if grp in ("Piping", "Wiring & Cables"):
            return ("Mtrs", "Mtr")
        elif grp == "Points":
            return ("Pts", "Pt")
        else:
            return ("Nos", "No")

    # ------------- 1) POST actions -------------
    if request.method == "POST":
        action = request.POST.get("action")

        # A) Upload Estimate file and parse items
        if action == "upload_estimate":
            file = request.FILES.get("estimate_file")
            if not file:
                return render(request, "core/workslip.html", {
                    "error": "Please upload an Estimate Excel file.",
                    "category": category,
                    "groups": groups,
                    "current_group": current_group,
                    "items_in_group": items_in_group,
                    "ws_estimate_rows": ws_estimate_rows,
                    "preview_rows": [],
                    "tp_percent": ws_tp_percent if ws_tp_percent else "",
                    "tp_type": ws_tp_type,
                    "supp_items_selected": ws_supp_items,
                    "work_name": ws_work_name,
                })

            try:
                # read once, create two workbooks: formulas + values
                excel_bytes = file.read()
                wb_est = load_workbook(BytesIO(excel_bytes), data_only=False)
                wb_est_vals = load_workbook(BytesIO(excel_bytes), data_only=True)
            except Exception as e:
                return render(request, "core/workslip.html", {
                    "error": f"Couldn't read uploaded Estimate file: {e}",
                    "category": category,
                    "groups": groups,
                    "current_group": current_group,
                    "items_in_group": items_in_group,
                    "ws_estimate_rows": ws_estimate_rows,
                    "preview_rows": [],
                    "tp_percent": ws_tp_percent if ws_tp_percent else "",
                    "tp_type": ws_tp_type,
                    "supp_items_selected": ws_supp_items,
                    "work_name": ws_work_name,
                })

            # ---- detect our Estimate-format sheet (ignore sheet name) ----
            def looks_like_our_estimate_sheet(sh):
                """
                Heuristic: row 3 headers should look like our Estimate:
                A: Sl.No, B: Quantity (Unit), D: Item Description, E: Rate, H: Amount
                """
                row = 3
                a = str(sh.cell(row=row, column=1).value or "").strip().lower()
                b = str(sh.cell(row=row, column=2).value or "").strip().lower()
                d = str(sh.cell(row=row, column=4).value or "").strip().lower()
                e = str(sh.cell(row=row, column=5).value or "").strip().lower()
                h = str(sh.cell(row=row, column=8).value or "").strip().lower()

                score = 0
                if "sl" in a and "no" in a:
                    score += 1
                if "quantity" in b:
                    score += 1
                if "item" in d and "description" in d:
                    score += 1
                if "rate" in e:
                    score += 1
                if "amount" in h:
                    score += 1
                return score >= 3  # tolerant

            estimate_sheets = [sh for sh in wb_est.worksheets if looks_like_our_estimate_sheet(sh)]

            if not estimate_sheets:
                return render(request, "core/workslip.html", {
                    "error": "No sheet in the uploaded workbook matches the Estimate format.",
                    "category": category,
                    "groups": groups,
                    "current_group": current_group,
                    "items_in_group": items_in_group,
                    "ws_estimate_rows": ws_estimate_rows,
                    "preview_rows": [],
                    "tp_percent": ws_tp_percent if ws_tp_percent else "",
                    "tp_type": ws_tp_type,
                    "supp_items_selected": ws_supp_items,
                    "work_name": ws_work_name,
                })

            # âœ… Only use the FIRST matching estimate sheet (no multi-sheet aggregation)
            ws_est_sheet = estimate_sheets[0]
            ws_est_vals_sheet = wb_est_vals[ws_est_sheet.title]

            # ---------- Find a sheet with yellow+red item headers (Item Blocks) ---------- #
            def get_heading_name(sheet, row_index: int):
                """
                Returns the heading text in this row (yellow fill + red font in A..J),
                or None if not a heading row.
                """
                for col in range(1, 10 + 1):  # A..J
                    cell = sheet.cell(row=row_index, column=col)
                    fill = getattr(cell, "fill", None)
                    font = getattr(cell, "font", None)

                    pattern = getattr(fill, "patternType", None)
                    fg = getattr(getattr(fill, "fgColor", None), "rgb", "")
                    color = getattr(getattr(font, "color", None), "rgb", "")

                    is_yellow = (
                        fill
                        and pattern
                        and str(pattern).lower() == "solid"
                        and fg
                        and str(fg).upper().endswith("FFFF00")
                    )
                    is_red = (
                        font
                        and getattr(font, "color", None)
                        and color
                        and str(color).upper().endswith("FF0000")
                    )

                    if is_yellow and is_red and str(cell.value or "").strip():
                        return str(cell.value).strip()
                return None

            # Try to find any sheet (except the estimate sheet) that contains such yellow headers
            blocks_sheet = None
            for sh in wb_est.worksheets:
                if sh.title == ws_est_sheet.title:
                    continue
                found = False
                for r in range(1, min(sh.max_row, 200) + 1):
                    if get_heading_name(sh, r):
                        found = True
                        break
                if found:
                    blocks_sheet = sh
                    break

            # Ordered list of item NAMES from the blocks sheet (yellow headers)
            heading_names = []
            if blocks_sheet is not None:
                for r in range(1, blocks_sheet.max_row + 1):
                    nm = get_heading_name(blocks_sheet, r)
                    if nm:
                        heading_names.append(nm)
            # If not found, we'll just fall back in parsing

            # prepare backend rate lookup (Master Datas  -  correct numeric rate)
            item_to_info = {it["name"]: it for it in items_list}
            wb_backend_vals = None
            ws_backend_vals = None
            if filepath and os.path.exists(filepath):
                try:
                    wb_backend_vals = load_workbook(filepath, data_only=True)
                    ws_backend_vals = wb_backend_vals["Master Datas"]
                except Exception:
                    wb_backend_vals = None
                    ws_backend_vals = None

            def backend_rate_for_item(name):
                info = item_to_info.get(name)
                if not info or ws_backend_vals is None:
                    return 0.0
                start_row = info["start_row"]
                end_row = info["end_row"]
                for r in range(end_row, start_row - 1, -1):
                    v = ws_backend_vals.cell(row=r, column=10).value  # col J
                    if v not in (None, ""):
                        try:
                            return float(v)
                        except Exception:
                            return 0.0
                return 0.0

            def to_number(v):
                try:
                    return float(v)
                except Exception:
                    return 0.0

            # ---- parse this Estimate sheet ----
            parsed_rows = []
            max_row = ws_est_sheet.max_row
            r = 4
            grand_total_val = 0.0
            
            # DEBUG: Log sheet info
            logger.info(f"[WORKSLIP DEBUG] Parsing sheet: {ws_est_sheet.title}, max_row={max_row}")

            # "Name of the work" from row 2
            work_name_local = ""
            name_cell = ws_est_sheet.cell(row=2, column=1).value
            if name_cell:
                text = str(name_cell)
                parts = text.split(":", 1)
                if len(parts) > 1:
                    work_name_local = parts[1].strip()
                else:
                    work_name_local = text.strip()

            heading_idx = 0  # to walk through heading_names in order

            while r <= max_row:
                desc = ws_est_sheet.cell(row=r, column=4).value  # D
                desc_str = str(desc or "").strip()
                desc_upper = desc_str.upper()

                # Rate may be formula; get value from data_only sheet
                rate_formula = ws_est_sheet.cell(row=r, column=5).value   # E (formula or value)
                rate_value = ws_est_vals_sheet.cell(row=r, column=5).value  # E (cached value)
                rate_is_empty = (rate_formula is None or str(rate_formula).strip() == "")

                # Quantity may also be formula â†’ use data_only workbook first
                qty_formula = ws_est_sheet.cell(row=r, column=2).value   # B
                qty_value = ws_est_vals_sheet.cell(row=r, column=2).value  # B value
                qty_is_empty = (qty_formula is None or str(qty_formula).strip() == "")
                
                # DEBUG: Log each row's data
                if r <= 10:  # Only first 10 rows to avoid flooding logs
                    logger.info(f"[WORKSLIP DEBUG] Row {r}: desc='{desc_str[:30] if desc_str else 'None'}', "
                               f"rate_formula={rate_formula}, rate_value={rate_value}, "
                               f"qty_formula={qty_formula}, qty_value={qty_value}")

                # If we see any totals keywords â†’ end of items
                if desc_str:
                    if any(
                        kw in desc_upper
                        for kw in ("SUB TOTAL", "SUBTOTAL", "TOTAL", "ECV")
                    ):
                        break

                # completely blank line
                if desc is None and ws_est_sheet.cell(row=r, column=1).value is None:
                    r += 1
                    continue

                # If row has NO rate AND NO quantity, it's a heading/section label â†’ skip it
                if rate_is_empty and qty_is_empty:
                    r += 1
                    continue

                # Process only rows that have BOTH description AND (rate or quantity)
                if desc_str != "" and not (rate_is_empty and qty_is_empty):
                    # Quantity conversion
                    if isinstance(qty_value, (int, float)):
                        qty_num = float(qty_value)
                    else:
                        qty_num = to_number(qty_formula)

                    unit = ws_est_sheet.cell(row=r, column=3).value  # C

                    # backend item name from desc (for rate lookup, etc.)
                    backend_item_name = desc_to_item.get(desc_str, desc_str)

                    # display name from yellow header list (for UI only)
                    if heading_idx < len(heading_names):
                        display_name = heading_names[heading_idx]
                    else:
                        display_name = backend_item_name or desc_str
                    heading_idx += 1

                    # rate: prefer cached numeric from data_only workbook,
                    # fallback to direct numeric or backend
                    if isinstance(rate_value, (int, float)):
                        rate_num = float(rate_value)
                    else:
                        rate_num = to_number(rate_formula)
                        if (rate_formula is None or rate_num == 0.0) and backend_item_name in item_to_info:
                            rate_num = backend_rate_for_item(backend_item_name)

                    parsed_rows.append({
                        "key": f"{ws_est_sheet.title}_row{r}",
                        "excel_row": r,
                        "item_name": backend_item_name,      # backend / mapping name
                        "display_name": display_name,        # yellow header for UI
                        "desc": desc_str,                    # full description from Estimate
                        "qty_est": qty_num,
                        "unit": str(unit or "").strip(),
                        "rate": rate_num,
                    })
                    # DEBUG: Log each parsed row
                    logger.info(f"[WORKSLIP DEBUG] Parsed row {r}: desc={desc_str[:50]}, qty={qty_num}, rate={rate_num}")

                r += 1

            # ---- find GRAND TOTAL *below* items block if present ----
            for rr in range(r, max_row + 1):
                d2 = str(ws_est_sheet.cell(row=rr, column=4).value or "").strip().upper()
                if "GRAND TOTAL" in d2:
                    grand_total_val = to_number(ws_est_vals_sheet.cell(row=rr, column=8).value)
                    break

            # store in session
            ws_estimate_rows = parsed_rows
            ws_exec_map = {}
            ws_supp_items = []
            ws_work_name = work_name_local
            
            # DEBUG: Log total parsed rows before saving to session
            logger.info(f"[WORKSLIP DEBUG] Total parsed rows: {len(parsed_rows)}, grand_total={grand_total_val}")

            # Parse additional metadata from rows 2-7 of the Estimate sheet
            base_metadata = {
                "work_name": work_name_local,
                "estimate_amount": "",
                "admin_sanction": "",
                "tech_sanction": "",
                "agreement": "",
                "agency_name": "",
                "grand_total": grand_total_val,
            }
            for meta_row in range(2, 8):
                cell_val = str(ws_est_sheet.cell(row=meta_row, column=1).value or "").strip()
                cell_lower = cell_val.lower()
                
                if ":" in cell_val:
                    parts = cell_val.split(":", 1)
                    extracted_value = parts[1].strip() if len(parts) > 1 else ""
                    
                    # If no value after colon, check other columns
                    if not extracted_value:
                        for c in range(2, 11):
                            val = ws_est_sheet.cell(row=meta_row, column=c).value
                            if val and str(val).strip():
                                extracted_value = str(val).strip()
                                break
                    
                    if "estimate amount" in cell_lower:
                        base_metadata["estimate_amount"] = extracted_value
                    elif "administrative" in cell_lower or "admin" in cell_lower:
                        base_metadata["admin_sanction"] = extracted_value
                    elif "technical" in cell_lower or "tech" in cell_lower:
                        base_metadata["tech_sanction"] = extracted_value
                    elif "agreement" in cell_lower:
                        base_metadata["agreement"] = extracted_value
                    elif "agency" in cell_lower:
                        base_metadata["agency_name"] = extracted_value

            request.session["ws_estimate_rows"] = ws_estimate_rows
            request.session["ws_exec_map"] = ws_exec_map
            request.session["ws_supp_items"] = ws_supp_items
            request.session["ws_estimate_grand_total"] = grand_total_val
            request.session["ws_work_name"] = ws_work_name
            request.session["ws_metadata"] = base_metadata  # Store metadata from base estimate
            # Reset phase tracking for new estimate upload
            request.session["ws_current_phase"] = 1
            request.session["ws_previous_phases"] = []
            
            # DEBUG: Verify session save
            logger.info(f"[WORKSLIP DEBUG] Session saved. ws_estimate_rows length: {len(request.session.get('ws_estimate_rows', []))}")

            return redirect(reverse('workslip_main') + '?preserve=1')

        # B) (optional) update preview
        elif action == "update_preview":
            exec_str = request.POST.get("exec_map", "")
            tp_percent_str = request.POST.get("tp_percent", "")
            tp_type = request.POST.get("tp_type", "Excess")

            new_exec_map = {}
            if exec_str:
                try:
                    raw = json.loads(exec_str)
                    if isinstance(raw, dict):
                        for k, v in raw.items():
                            try:
                                new_exec_map[str(k)] = float(v)
                            except Exception:
                                continue
                except Exception:
                    pass

            ws_exec_map.update(new_exec_map)
            try:
                ws_tp_percent = float(tp_percent_str) if tp_percent_str != "" else 0.0
            except Exception:
                ws_tp_percent = 0.0
            ws_tp_type = tp_type if tp_type in ("Less", "Excess") else "Excess"

            request.session["ws_exec_map"] = ws_exec_map
            request.session["ws_tp_percent"] = ws_tp_percent
            request.session["ws_tp_type"] = ws_tp_type

            return redirect(reverse('workslip_main') + '?preserve=1')

        # C) Add Supplemental items
        elif action == "add_supplemental":
            selected = request.POST.getlist("supp_items")
            new_list = ws_supp_items[:]
            for nm in selected:
                if nm not in new_list:
                    new_list.append(nm)
            ws_supp_items = new_list

            # merge exec_map & TP coming from hidden fields
            exec_str = request.POST.get("exec_map", "")
            tp_percent_str = request.POST.get("tp_percent", "")
            tp_type = request.POST.get("tp_type", "Excess")

            ws_exec_map_session = request.session.get("ws_exec_map", {}) or {}
            new_exec_map = {}
            if exec_str:
                try:
                    raw = json.loads(exec_str)
                    if isinstance(raw, dict):
                        for k, v in raw.items():
                            try:
                                new_exec_map[str(k)] = float(v)
                            except Exception:
                                continue
                except Exception:
                    pass

            ws_exec_map = ws_exec_map_session.copy()
            ws_exec_map.update(new_exec_map)

            try:
                ws_tp_percent = float(tp_percent_str) if tp_percent_str != "" else 0.0
            except Exception:
                ws_tp_percent = 0.0
            ws_tp_type = tp_type if tp_type in ("Less", "Excess") else "Excess"

            # Filter exec_map: keep all base items, and only supplemental items that are still selected
            filtered_exec_map = {}
            for k, v in ws_exec_map.items():
                if not k.startswith("supp:") or k[5:] in ws_supp_items:
                    filtered_exec_map[k] = v

            request.session["ws_supp_items"] = ws_supp_items
            request.session["ws_exec_map"] = filtered_exec_map
            request.session["ws_tp_percent"] = ws_tp_percent
            request.session["ws_tp_type"] = ws_tp_type

            return redirect(reverse('workslip_main') + '?preserve=1')

        # D-COMBINED) Upload Both Estimate + Previous Workslips Together (Dynamic Multi-Workslip)
        elif action == "upload_combined":
            estimate_file = request.FILES.get("estimate_file")
            
            # Get target workslip from form
            target_workslip_str = request.POST.get("target_workslip", "1")
            try:
                target_workslip = int(target_workslip_str)
                if target_workslip < 1:
                    target_workslip = 1
                if target_workslip > 10:
                    target_workslip = 10
            except (ValueError, TypeError):
                target_workslip = 1
            
            logger.info(f"[MULTI-WORKSLIP] Target workslip: {target_workslip}")
            
            # Collect all previous workslip files based on target
            workslip_files = []
            for i in range(1, target_workslip):
                ws_file = request.FILES.get(f"workslip_file_{i}")
                if ws_file:
                    workslip_files.append((i, ws_file))
                    logger.info(f"[MULTI-WORKSLIP] Found workslip file {i}: {ws_file.name}")
            
            # Legacy support: also check for old "previous_workslip_file" field
            legacy_ws_file = request.FILES.get("previous_workslip_file")
            if legacy_ws_file and not workslip_files:
                workslip_files.append((1, legacy_ws_file))
                logger.info(f"[MULTI-WORKSLIP] Using legacy workslip file: {legacy_ws_file.name}")
            
            # Helper to get heading names from yellow+red cells
            def get_heading_name_from_sheet(sheet, row_index: int):
                for col in range(1, 10 + 1):
                    cell = sheet.cell(row=row_index, column=col)
                    fill = getattr(cell, "fill", None)
                    font = getattr(cell, "font", None)
                    pattern = getattr(fill, "patternType", None)
                    fg = getattr(getattr(fill, "fgColor", None), "rgb", "")
                    color = getattr(getattr(font, "color", None), "rgb", "")
                    is_yellow = (fill and pattern and str(pattern).lower() == "solid" and fg and str(fg).upper().endswith("FFFF00"))
                    is_red = (font and getattr(font, "color", None) and color and str(color).upper().endswith("FF0000"))
                    if is_yellow and is_red and str(cell.value or "").strip():
                        return str(cell.value).strip()
                return None
            
            # Validation based on target workslip
            if target_workslip == 1:
                # For Workslip-1, only estimate is required
                if not estimate_file:
                    return render(request, "core/workslip.html", {
                        "error": "Please upload an Estimate file for Workslip-1.",
                        "category": category, "groups": groups, "current_group": current_group,
                        "items_in_group": items_in_group, "ws_estimate_rows": ws_estimate_rows,
                        "preview_rows": [], "tp_percent": ws_tp_percent if ws_tp_percent else "",
                        "tp_type": ws_tp_type, "supp_items_selected": ws_supp_items,
                        "work_name": ws_work_name, "current_phase": ws_current_phase,
                        "previous_phases": ws_previous_phases, "target_workslip": target_workslip,
                    })
            else:
                # For Workslip-2 and above: Estimate is OPTIONAL if previous workslips are provided
                # Item descriptions can be extracted from previous workslip files
                if not estimate_file and not workslip_files:
                    return render(request, "core/workslip.html", {
                        "error": f"Please upload either an Estimate file OR previous Workslip files for Workslip-{target_workslip}.",
                        "category": category, "groups": groups, "current_group": current_group,
                        "items_in_group": items_in_group, "ws_estimate_rows": ws_estimate_rows,
                        "preview_rows": [], "tp_percent": ws_tp_percent if ws_tp_percent else "",
                        "tp_type": ws_tp_type, "supp_items_selected": ws_supp_items,
                        "work_name": ws_work_name, "current_phase": ws_current_phase,
                        "previous_phases": ws_previous_phases, "target_workslip": target_workslip,
                    })
                
                # Check if at least the most recent previous workslip is uploaded
                # (e.g., for Workslip-3, at least Workslip-2 should be uploaded)
                if workslip_files:
                    uploaded_nums = {num for num, _ in workslip_files}
                    most_recent_required = target_workslip - 1
                    if most_recent_required not in uploaded_nums:
                        return render(request, "core/workslip.html", {
                            "error": f"Please upload at least Workslip-{most_recent_required} to generate Workslip-{target_workslip}.",
                            "category": category, "groups": groups, "current_group": current_group,
                            "items_in_group": items_in_group, "ws_estimate_rows": ws_estimate_rows,
                            "preview_rows": [], "tp_percent": ws_tp_percent if ws_tp_percent else "",
                            "tp_type": ws_tp_type, "supp_items_selected": ws_supp_items,
                            "work_name": ws_work_name, "current_phase": ws_current_phase,
                            "previous_phases": ws_previous_phases, "target_workslip": target_workslip,
                        })
            
            # Parse Estimate file for display names (yellow+red rows from Items Blocks sheet)
            estimate_display_names = []
            est_bytes = None  # Store bytes for reuse later
            if estimate_file:
                try:
                    est_bytes = estimate_file.read()
                    wb_est = load_workbook(BytesIO(est_bytes), data_only=False)
                    
                    # Find Items Blocks sheet (sheet with yellow+red headers)
                    blocks_sheet = None
                    for sh in wb_est.worksheets:
                        found = False
                        for r in range(1, min(sh.max_row, 200) + 1):
                            if get_heading_name_from_sheet(sh, r):
                                found = True
                                break
                        if found:
                            blocks_sheet = sh
                            break
                    
                    if blocks_sheet:
                        for r in range(1, blocks_sheet.max_row + 1):
                            nm = get_heading_name_from_sheet(blocks_sheet, r)
                            if nm:
                                estimate_display_names.append(nm)
                        logger.info(f"[COMBINED UPLOAD] Found {len(estimate_display_names)} item names from estimate Items Blocks sheet")
                except Exception as e:
                    logger.warning(f"[COMBINED UPLOAD] Could not parse estimate for display names: {e}")
            
            # If no estimate file, extract item names from the most recent workslip file
            # This allows generating Workslip-3 using only Workslip-2 (without original estimate)
            if not estimate_file and workslip_files:
                # Use the most recent workslip file for item names
                most_recent_ws = max(workslip_files, key=lambda x: x[0])
                ws_file_num, workslip_file_for_names = most_recent_ws
                logger.info(f"[COMBINED UPLOAD] No estimate file - extracting item names from Workslip-{ws_file_num}")
                
                try:
                    workslip_file_for_names.seek(0)  # Reset file pointer
                    ws_bytes_for_names = workslip_file_for_names.read()
                    workslip_file_for_names.seek(0)  # Reset again for later use
                    wb_ws_for_names = load_workbook(BytesIO(ws_bytes_for_names), data_only=False)
                    
                    # Find Items Blocks sheet in workslip (same structure as estimate)
                    ws_blocks_sheet = None
                    for sh in wb_ws_for_names.worksheets:
                        found = False
                        for r in range(1, min(sh.max_row or 0, 200) + 1):
                            if get_heading_name_from_sheet(sh, r):
                                found = True
                                break
                        if found:
                            ws_blocks_sheet = sh
                            break
                    
                    if ws_blocks_sheet:
                        for r in range(1, (ws_blocks_sheet.max_row or 0) + 1):
                            nm = get_heading_name_from_sheet(ws_blocks_sheet, r)
                            if nm and nm not in estimate_display_names:
                                estimate_display_names.append(nm)
                        logger.info(f"[COMBINED UPLOAD] Extracted {len(estimate_display_names)} item names from Workslip-{ws_file_num}")
                    else:
                        logger.warning(f"[COMBINED UPLOAD] No Items Blocks sheet found in Workslip-{ws_file_num}")
                except Exception as e:
                    logger.warning(f"[COMBINED UPLOAD] Could not extract item names from workslip: {e}")
            
            # Process multiple workslip files for multi-workslip upload
            if workslip_files:
                # Sort workslip files by phase number to process in order
                workslip_files.sort(key=lambda x: x[0])
                
                # Initialize combined data structures for all phases
                all_phase_exec_maps = []  # List of exec_maps, one per phase
                all_phase_ae_data = []    # List of ae_data, one per phase
                all_previous_supp_items = []  # Combined supplemental items from all phases
                parsed_items = None  # Will hold the base items (from the last/most recent workslip)
                ws_metadata = None  # Will hold metadata from the most recent workslip
                
                # Helper function to detect phase and column structure
                def detect_workslip_phase_and_columns(ws):
                    header_row = 8
                    for r in range(1, 15):
                        cell_val = str(ws.cell(row=r, column=1).value or "").strip().lower()
                        if "sl" in cell_val:
                            header_row = r
                            break
                    
                    phase_count = 0
                    col_map = {"header_row": header_row}
                    
                    for c in range(1, 30):
                        header = str(ws.cell(row=header_row, column=c).value or "").strip().lower()
                        if ("execution" in header or "exec" in header or "workslip" in header) and ("qty" in header or "quantity" in header):
                            phase_count += 1
                            col_map[f"exec_qty_phase_{phase_count}"] = c
                        elif ("execution" in header or "exec" in header or "workslip" in header) and ("amount" in header or "amt" in header):
                            col_map[f"exec_amt_phase_{phase_count}"] = c
                        elif "qty" in header and ("est" in header or "estimate" in header):
                            col_map["est_qty"] = c
                        elif "description" in header or "desc" in header:
                            col_map["desc"] = c
                        elif "rate" in header and ("est" in header or "estimate" in header):
                            col_map["est_rate"] = c
                    
                    if phase_count == 0:
                        phase_count = 1
                        col_map["exec_qty_phase_1"] = 7
                        col_map["exec_amt_phase_1"] = 9
                    
                    return phase_count, col_map
                
                # Process each workslip file
                for ws_file_num, workslip_file in workslip_files:
                    logger.info(f"[MULTI-WORKSLIP] Processing Workslip-{ws_file_num}")
                    try:
                        ws_bytes = workslip_file.read()
                        wb_ws = load_workbook(BytesIO(ws_bytes), data_only=True)
                        wb_ws_formulas = load_workbook(BytesIO(ws_bytes), data_only=False)
                    except Exception as e:
                        return render(request, "core/workslip.html", {
                            "error": f"Couldn't read Workslip-{ws_file_num} file: {e}",
                            "category": category, "groups": groups, "current_group": current_group,
                            "items_in_group": items_in_group, "ws_estimate_rows": ws_estimate_rows,
                            "preview_rows": [], "tp_percent": ws_tp_percent if ws_tp_percent else "",
                            "tp_type": ws_tp_type, "supp_items_selected": ws_supp_items,
                            "work_name": ws_work_name, "current_phase": ws_current_phase,
                            "previous_phases": ws_previous_phases, "target_workslip": target_workslip,
                        })
                    
                    # Find WorkSlip sheet (for quantities)
                    ws_sheet = None
                    for sh in wb_ws.worksheets:
                        if "workslip" in sh.title.lower() or "working estimate" in str(sh.cell(row=1, column=1).value or "").lower():
                            ws_sheet = sh
                            logger.info(f"[MULTI-WORKSLIP] Workslip-{ws_file_num}: Found workslip sheet: '{sh.title}'")
                            break
                    if not ws_sheet:
                        ws_sheet = wb_ws.active
                        logger.info(f"[MULTI-WORKSLIP] Workslip-{ws_file_num}: Using active sheet: '{ws_sheet.title}'")
                    
                    # Find Workslip Items Blocks sheet (for supplemental item names - yellow+red rows)
                    ws_blocks_sheet = None
                    for sh in wb_ws_formulas.worksheets:
                        if "items" in sh.title.lower() and "block" in sh.title.lower():
                            found = False
                            for r in range(1, min(sh.max_row, 200) + 1):
                                if get_heading_name_from_sheet(sh, r):
                                    found = True
                                    break
                            if found:
                                ws_blocks_sheet = sh
                                break
                    
                    # Parse Workslip Items Blocks for supplemental item names
                    ws_supp_item_names = []  # Supplemental items from this workslip
                    if ws_blocks_sheet:
                        in_supplemental_section = False
                        for r in range(1, ws_blocks_sheet.max_row + 1):
                            nm = get_heading_name_from_sheet(ws_blocks_sheet, r)
                            if nm:
                                nm_upper = nm.upper()
                                if "SUPPLEMENTAL" in nm_upper:
                                    in_supplemental_section = True
                                    continue
                                if in_supplemental_section:
                                    ws_supp_item_names.append(nm)
                        logger.info(f"[MULTI-WORKSLIP] Workslip-{ws_file_num}: Found {len(ws_supp_item_names)} supplemental items")
                    
                    phase_count, col_map = detect_workslip_phase_and_columns(ws_sheet)
                    header_row = col_map.get("header_row", 8)
                    
                    # Extract metadata from workslip header rows (rows 2-7) - only from last file
                    file_metadata = {
                        "work_name": "",
                        "estimate_amount": "",
                        "admin_sanction": "",
                        "tech_sanction": "",
                        "agreement": "",
                        "agency_name": "",
                        "tp_percent": 0.0,
                        "tp_type": "Excess",
                        "grand_total": 0.0,
                        "deduct_old_material": 0.0,
                        "lc_percent": 0.0,
                        "qc_percent": 0.0,
                        "nac_percent": 0.0,
                    }
                    
                    for r in range(2, 8):
                        cell_val = str(ws_sheet.cell(row=r, column=1).value or "").strip()
                        cell_lower = cell_val.lower()
                    
                        # Check if this is a label row and extract value
                        extracted_value = ""
                        if ":" in cell_val:
                            parts = cell_val.split(":", 1)
                            extracted_value = parts[1].strip() if len(parts) > 1 else ""
                        
                            # If no value after colon, check columns 2-10 for value
                            if not extracted_value:
                                for c in range(2, 11):
                                    val = ws_sheet.cell(row=r, column=c).value
                                    if val and str(val).strip():
                                        extracted_value = str(val).strip()
                                        break
                            
                            if "name of the work" in cell_lower or "work name" in cell_lower:
                                file_metadata["work_name"] = extracted_value
                            elif "estimate amount" in cell_lower:
                                file_metadata["estimate_amount"] = extracted_value
                                import re
                                num_match = re.search(r'[\d,]+\.?\d*', extracted_value.replace(',', ''))
                                if num_match:
                                    try:
                                        file_metadata["grand_total"] = float(num_match.group().replace(',', ''))
                                    except:
                                        pass
                            elif "administrative" in cell_lower or "admin" in cell_lower:
                                file_metadata["admin_sanction"] = extracted_value
                            elif "technical" in cell_lower or "tech" in cell_lower:
                                file_metadata["tech_sanction"] = extracted_value
                            elif "agreement" in cell_lower:
                                file_metadata["agreement"] = extracted_value
                            elif "agency" in cell_lower:
                                file_metadata["agency_name"] = extracted_value
                    
                    # Scan footer section for T.P, Grand Total, Deduct, LC, QC, NAC values
                    desc_col = col_map.get("desc", 2)
                    for r in range(header_row + 1, ws_sheet.max_row + 1):
                        desc_text = str(ws_sheet.cell(row=r, column=desc_col).value or "").strip()
                        desc_upper = desc_text.upper()
                        
                        if ("ADD" in desc_upper or "DEDUCT" in desc_upper) and ("T.P" in desc_upper or "T P" in desc_upper) and "UNUSED" not in desc_upper:
                            import re
                            tp_match = re.search(r'@\s*([\d.]+)\s*%', desc_text)
                            if tp_match:
                                try:
                                    file_metadata["tp_percent"] = float(tp_match.group(1))
                                except:
                                    pass
                            if "EXCESS" in desc_upper:
                                file_metadata["tp_type"] = "Excess"
                            elif "LESS" in desc_upper:
                                file_metadata["tp_type"] = "Less"
                        elif "GRAND TOTAL" in desc_upper:
                            for c in range(ws_sheet.max_column, 0, -1):
                                val = ws_sheet.cell(row=r, column=c).value
                                if val:
                                    try:
                                        file_metadata["grand_total"] = float(val)
                                        break
                                    except:
                                        continue
                        elif "DEDUCT" in desc_upper and "OLD" in desc_upper:
                            for c in range(ws_sheet.max_column, 0, -1):
                                val = ws_sheet.cell(row=r, column=c).value
                                if val:
                                    try:
                                        file_metadata["deduct_old_material"] = abs(float(val))
                                        break
                                    except:
                                        continue
                        elif "L.C" in desc_upper or "LC @" in desc_upper or "ADD LC" in desc_upper:
                            import re
                            lc_match = re.search(r'@\s*([\d.]+)\s*%', desc_text)
                            if lc_match:
                                try:
                                    file_metadata["lc_percent"] = float(lc_match.group(1))
                                except:
                                    pass
                        elif "Q.C" in desc_upper or "QC @" in desc_upper or "ADD QC" in desc_upper:
                            import re
                            qc_match = re.search(r'@\s*([\d.]+)\s*%', desc_text)
                            if qc_match:
                                try:
                                    file_metadata["qc_percent"] = float(qc_match.group(1))
                                except:
                                    pass
                        elif "NAC" in desc_upper:
                            import re
                            nac_match = re.search(r'@\s*([\d.]+)\s*%', desc_text)
                            if nac_match:
                                try:
                                    file_metadata["nac_percent"] = float(nac_match.group(1))
                                except:
                                    pass
                    
                    # Parse main items from this workslip file
                    file_parsed_items = []
                    file_phase_exec_maps = [{} for _ in range(phase_count)]
                    file_phase_ae_data = [{} for _ in range(phase_count)]
                    last_base_key = None
                    display_name_idx = 0
                    
                    for r in range(header_row + 1, ws_sheet.max_row + 1):
                        desc = str(ws_sheet.cell(row=r, column=desc_col).value or "").strip()
                        
                        if not desc:
                            continue
                        
                        if any(kw in desc.upper() for kw in ("SUB TOTAL", "SUBTOTAL", "TOTAL", "GRAND TOTAL", "DEDUCT", "SUPPLEMENTAL")):
                            break
                        
                        desc_upper = desc.upper().strip()
                        is_ae_row = (desc_upper.startswith("AE") and len(desc_upper) >= 2 and 
                                     (len(desc_upper) == 2 or desc_upper[2:].isdigit() or desc_upper[2] == ' '))
                        
                        est_qty = ws_sheet.cell(row=r, column=col_map.get("est_qty", 4)).value or 0
                        est_rate = ws_sheet.cell(row=r, column=col_map.get("est_rate", 5)).value or 0
                        unit = ws_sheet.cell(row=r, column=3).value or ""
                        
                        try:
                            est_qty = float(est_qty)
                        except:
                            est_qty = 0.0
                        try:
                            est_rate = float(est_rate)
                        except:
                            est_rate = 0.0
                        
                        if is_ae_row and last_base_key:
                            for p in range(1, phase_count + 1):
                                exec_qty_col = col_map.get(f"exec_qty_phase_{p}")
                                if exec_qty_col:
                                    exec_qty = ws_sheet.cell(row=r, column=exec_qty_col).value or 0
                                    try:
                                        exec_qty = float(exec_qty)
                                    except:
                                        exec_qty = 0.0
                                    if exec_qty > 0:
                                        if last_base_key in file_phase_exec_maps[p-1]:
                                            file_phase_exec_maps[p-1][last_base_key] += exec_qty
                                        else:
                                            file_phase_exec_maps[p-1][last_base_key] = exec_qty
                                        ae_key = f"{last_base_key}:ae:{desc}"
                                        file_phase_ae_data[p-1][ae_key] = exec_qty
                        else:
                            row_key = f"ws{ws_file_num}_row_{r}"
                            last_base_key = row_key
                            
                            for p in range(1, phase_count + 1):
                                exec_qty_col = col_map.get(f"exec_qty_phase_{p}")
                                if exec_qty_col:
                                    exec_qty = ws_sheet.cell(row=r, column=exec_qty_col).value or 0
                                    try:
                                        exec_qty = float(exec_qty)
                                    except:
                                        exec_qty = 0.0
                                    if exec_qty > 0:
                                        file_phase_exec_maps[p-1][row_key] = exec_qty
                            
                            if estimate_display_names and display_name_idx < len(estimate_display_names):
                                display_name = estimate_display_names[display_name_idx]
                            else:
                                display_name = desc
                            display_name_idx += 1
                            
                            file_parsed_items.append({
                                "key": row_key,
                                "excel_row": r,
                                "item_name": desc,
                                "display_name": display_name,
                                "desc": desc,
                                "qty_est": est_qty,
                                "unit": str(unit).strip(),
                                "rate": est_rate,
                            })
                    
                    # Parse supplemental items from this workslip
                    file_supp_items = []
                    in_supp_section = False
                    current_supp_section = 0  # Track which supplemental section we're in
                    supp_item_idx = 0
                    for r in range(header_row + 1, ws_sheet.max_row + 1):
                        desc = str(ws_sheet.cell(row=r, column=desc_col).value or "").strip()
                        
                        # Check if this is a supplemental section header (e.g., "Supplemental Items-1", "Supplemental Items-2")
                        if "SUPPLEMENTAL" in desc.upper():
                            in_supp_section = True
                            # Extract the section number from header like "Supplemental Items-1" or "Supplemental Items-2"
                            import re
                            supp_match = re.search(r'(\d+)\s*$', desc)
                            if supp_match:
                                current_supp_section = int(supp_match.group(1))
                            else:
                                current_supp_section += 1  # Increment for unnamed sections
                            logger.info(f"[MULTI-WORKSLIP] Found supplemental section {current_supp_section}: {desc}")
                            continue
                        
                        if in_supp_section and desc:
                            if any(kw in desc.upper() for kw in ("SUB TOTAL", "SUBTOTAL", "TOTAL", "GRAND TOTAL", "DEDUCT")):
                                break
                            
                            if supp_item_idx < len(ws_supp_item_names):
                                supp_display_name = ws_supp_item_names[supp_item_idx]
                            else:
                                supp_display_name = desc
                            supp_item_idx += 1
                            
                            # Get unit for supplemental item
                            supp_unit = str(ws_sheet.cell(row=r, column=3).value or "").strip()
                            
                            # Try to get rate from est_rate column first
                            supp_rate_col = col_map.get("est_rate", 5)
                            supp_rate = ws_sheet.cell(row=r, column=supp_rate_col).value or 0
                            try:
                                supp_rate = float(supp_rate)
                            except:
                                supp_rate = 0.0
                            
                            for p in range(1, phase_count + 1):
                                exec_qty_col = col_map.get(f"exec_qty_phase_{p}")
                                exec_amt_col = col_map.get(f"exec_amt_phase_{p}")
                                if exec_qty_col:
                                    exec_qty = ws_sheet.cell(row=r, column=exec_qty_col).value or 0
                                    try:
                                        exec_qty = float(exec_qty)
                                    except:
                                        exec_qty = 0.0
                                    
                                    # If rate is 0, try to calculate from exec amount / qty
                                    if supp_rate == 0 and exec_amt_col and exec_qty > 0:
                                        exec_amt = ws_sheet.cell(row=r, column=exec_amt_col).value or 0
                                        try:
                                            exec_amt = float(exec_amt)
                                            if exec_amt > 0:
                                                supp_rate = exec_amt / exec_qty
                                        except:
                                            pass
                                    
                                    if exec_qty > 0:
                                        supp_amount = exec_qty * supp_rate
                                        file_supp_items.append({
                                            "name": supp_display_name,
                                            "qty": exec_qty,
                                            "phase": p,  # Use the workslip phase number (1, 2, etc.) from the column
                                            "supp_section": current_supp_section,  # Which supplemental section (1 or 2)
                                            "desc": desc,
                                            "unit": supp_unit,
                                            "rate": supp_rate,
                                            "amount": supp_amount,
                                        })
                    
                    # Accumulate data from this workslip file into the combined structures
                    # For base items, use the most recent workslip's items
                    if file_parsed_items:
                        parsed_items = file_parsed_items
                    
                    # Accumulate execution maps from all phases in this file
                    for exec_map in file_phase_exec_maps:
                        if exec_map:  # Only add non-empty maps
                            all_phase_exec_maps.append(exec_map)
                    
                    for ae_data in file_phase_ae_data:
                        if ae_data:
                            all_phase_ae_data.append(ae_data)
                    
                    # Accumulate supplemental items
                    all_previous_supp_items.extend(file_supp_items)
                    
                    # Keep metadata from the most recent workslip
                    ws_metadata = file_metadata
                    
                    logger.info(f"[MULTI-WORKSLIP] Workslip-{ws_file_num}: Parsed {len(file_parsed_items)} items, {len(file_supp_items)} supp items, {phase_count} phases")
                
                # End of workslip files loop - save accumulated data to session
                if parsed_items:
                    ws_estimate_rows = parsed_items
                    request.session["ws_estimate_rows"] = ws_estimate_rows
                    request.session["ws_previous_phases"] = all_phase_exec_maps
                    request.session["ws_previous_ae_data"] = all_phase_ae_data
                    request.session["ws_previous_supp_items"] = all_previous_supp_items
                    request.session["ws_current_phase"] = target_workslip
                    request.session["ws_exec_map"] = {}
                    request.session["ws_target_workslip"] = target_workslip
                    
                    if ws_metadata:
                        request.session["ws_metadata"] = ws_metadata
                        if ws_metadata.get("work_name"):
                            request.session["ws_work_name"] = ws_metadata["work_name"]
                        request.session["ws_tp_percent"] = ws_metadata.get("tp_percent", 0.0)
                        request.session["ws_tp_type"] = ws_metadata.get("tp_type", "Excess")
                        request.session["ws_deduct_old_material"] = ws_metadata.get("deduct_old_material", 0.0)
                        request.session["ws_lc_percent"] = ws_metadata.get("lc_percent", 0.0)
                        request.session["ws_qc_percent"] = ws_metadata.get("qc_percent", 0.0)
                        request.session["ws_nac_percent"] = ws_metadata.get("nac_percent", 0.0)
                        if ws_metadata.get("grand_total", 0) > 0:
                            request.session["ws_estimate_grand_total"] = ws_metadata.get("grand_total", 0.0)
                    
                    logger.info(f"[MULTI-WORKSLIP] Complete: Loaded {len(all_phase_exec_maps)} phases, {len(parsed_items)} items, {len(all_previous_supp_items)} supp items for Workslip-{target_workslip}")
                
                return redirect(reverse('workslip_main') + '?preserve=1')
            
            # Only estimate file was uploaded (no previous workslip) - process as new Workslip-1
            elif estimate_file and est_bytes:
                try:
                    # Use the already-read estimate bytes (from display names extraction)
                    wb_est_parse = load_workbook(BytesIO(est_bytes), data_only=False)
                    wb_est_vals_parse = load_workbook(BytesIO(est_bytes), data_only=True)
                    
                    # Find estimate sheet
                    def looks_like_our_estimate_sheet_local(sh):
                        row = 3
                        a = str(sh.cell(row=row, column=1).value or "").strip().lower()
                        b = str(sh.cell(row=row, column=2).value or "").strip().lower()
                        d = str(sh.cell(row=row, column=4).value or "").strip().lower()
                        e = str(sh.cell(row=row, column=5).value or "").strip().lower()
                        h = str(sh.cell(row=row, column=8).value or "").strip().lower()
                        score = 0
                        if "sl" in a and "no" in a: score += 1
                        if "quantity" in b: score += 1
                        if "item" in d and "description" in d: score += 1
                        if "rate" in e: score += 1
                        if "amount" in h: score += 1
                        return score >= 3
                    
                    estimate_sheets = [sh for sh in wb_est_parse.worksheets if looks_like_our_estimate_sheet_local(sh)]
                    if not estimate_sheets:
                        return render(request, "core/workslip.html", {
                            "error": "No sheet in the uploaded workbook matches the Estimate format.",
                            "category": category, "groups": groups, "current_group": current_group,
                            "items_in_group": items_in_group, "ws_estimate_rows": ws_estimate_rows,
                            "preview_rows": [], "tp_percent": ws_tp_percent if ws_tp_percent else "",
                            "tp_type": ws_tp_type, "supp_items_selected": ws_supp_items,
                            "work_name": ws_work_name, "current_phase": ws_current_phase,
                            "previous_phases": ws_previous_phases,
                        })
                    
                    ws_est_sheet = estimate_sheets[0]
                    ws_est_vals_sheet = wb_est_vals_parse[ws_est_sheet.title]
                    
                    # Use display names from estimate if we found them earlier
                    heading_names_local = estimate_display_names if estimate_display_names else []
                    
                    def to_number_local(v):
                        try:
                            return float(v)
                        except:
                            return 0.0
                    
                    # Parse rows
                    parsed_rows = []
                    max_row = ws_est_sheet.max_row
                    r = 4
                    heading_idx = 0
                    
                    work_name_local = ""
                    name_cell = ws_est_sheet.cell(row=2, column=1).value
                    if name_cell:
                        text = str(name_cell)
                        parts = text.split(":", 1)
                        work_name_local = parts[1].strip() if len(parts) > 1 else text.strip()
                    
                    while r <= max_row:
                        desc = ws_est_sheet.cell(row=r, column=4).value
                        desc_str = str(desc or "").strip()
                        desc_upper = desc_str.upper()
                        
                        rate_formula = ws_est_sheet.cell(row=r, column=5).value
                        rate_value = ws_est_vals_sheet.cell(row=r, column=5).value
                        rate_is_empty = (rate_formula is None or str(rate_formula).strip() == "")
                        
                        qty_formula = ws_est_sheet.cell(row=r, column=2).value
                        qty_value = ws_est_vals_sheet.cell(row=r, column=2).value
                        qty_is_empty = (qty_formula is None or str(qty_formula).strip() == "")
                        
                        if desc_str and any(kw in desc_upper for kw in ("SUB TOTAL", "SUBTOTAL", "TOTAL", "ECV")):
                            break
                        if desc is None and ws_est_sheet.cell(row=r, column=1).value is None:
                            r += 1
                            continue
                        if rate_is_empty and qty_is_empty:
                            r += 1
                            continue
                        
                        if desc_str != "" and not (rate_is_empty and qty_is_empty):
                            qty_num = float(qty_value) if isinstance(qty_value, (int, float)) else to_number_local(qty_formula)
                            unit = ws_est_sheet.cell(row=r, column=3).value
                            backend_item_name = desc_to_item.get(desc_str, desc_str)
                            
                            if heading_idx < len(heading_names_local):
                                display_name = heading_names_local[heading_idx]
                            else:
                                display_name = backend_item_name or desc_str
                            heading_idx += 1
                            
                            rate_num = float(rate_value) if isinstance(rate_value, (int, float)) else to_number_local(rate_formula)
                            
                            parsed_rows.append({
                                "key": f"{ws_est_sheet.title}_row{r}",
                                "excel_row": r,
                                "item_name": backend_item_name,
                                "display_name": display_name,
                                "desc": desc_str,
                                "qty_est": qty_num,
                                "unit": str(unit or "").strip(),
                                "rate": rate_num,
                            })
                        r += 1
                    
                    # Store in session
                    request.session["ws_estimate_rows"] = parsed_rows
                    request.session["ws_exec_map"] = {}
                    request.session["ws_supp_items"] = []
                    request.session["ws_work_name"] = work_name_local
                    request.session["ws_current_phase"] = target_workslip  # Use target workslip number
                    request.session["ws_target_workslip"] = target_workslip
                    request.session["ws_previous_phases"] = []
                    request.session["ws_previous_supp_items"] = []
                    request.session["ws_previous_ae_data"] = []
                    logger.info(f"[COMBINED UPLOAD - ESTIMATE ONLY] Parsed {len(parsed_rows)} items from estimate for Workslip-{target_workslip}")
                    
                except Exception as e:
                    logger.warning(f"[COMBINED UPLOAD] Error parsing estimate only: {e}")
                    return render(request, "core/workslip.html", {
                        "error": f"Error parsing estimate file: {e}",
                        "category": category, "groups": groups, "current_group": current_group,
                        "items_in_group": items_in_group, "ws_estimate_rows": ws_estimate_rows,
                        "preview_rows": [], "tp_percent": ws_tp_percent if ws_tp_percent else "",
                        "tp_type": ws_tp_type, "supp_items_selected": ws_supp_items,
                        "work_name": ws_work_name, "current_phase": ws_current_phase,
                        "previous_phases": ws_previous_phases,
                    })
                
                return redirect(reverse('workslip_main') + '?preserve=1')
            
            return redirect(reverse('workslip_main') + '?preserve=1')

        # D) Clear Everything
        elif action == "clear_all":
            ws_estimate_rows = []
            ws_exec_map = {}
            ws_tp_percent = 0.0
            ws_tp_type = "Excess"
            ws_supp_items = []
            ws_work_name = ""
            request.session["ws_estimate_rows"] = []
            request.session["ws_exec_map"] = {}
            request.session["ws_tp_percent"] = 0.0
            request.session["ws_tp_type"] = "Excess"
            request.session["ws_supp_items"] = []
            request.session["ws_estimate_grand_total"] = 0.0
            request.session["ws_work_name"] = ""
            request.session["ws_current_phase"] = 1
            request.session["ws_target_workslip"] = 1
            request.session["ws_previous_phases"] = []
            request.session["ws_previous_supp_items"] = []
            request.session["ws_previous_ae_data"] = []
            return redirect("workslip_main")

        # D2) Upload Previous Workslip for Next Phase
        elif action == "upload_previous_workslip":
            file = request.FILES.get("previous_workslip_file")
            if not file:
                return render(request, "core/workslip.html", {
                    "error": "Please upload a previous Workslip Excel file.",
                    "category": category,
                    "groups": groups,
                    "current_group": current_group,
                    "items_in_group": items_in_group,
                    "ws_estimate_rows": ws_estimate_rows,
                    "preview_rows": [],
                    "tp_percent": ws_tp_percent if ws_tp_percent else "",
                    "tp_type": ws_tp_type,
                    "supp_items_selected": ws_supp_items,
                    "work_name": ws_work_name,
                    "current_phase": ws_current_phase,
                    "previous_phases": ws_previous_phases,
                })

            try:
                excel_bytes = file.read()
                wb_ws = load_workbook(BytesIO(excel_bytes), data_only=True)
            except Exception as e:
                return render(request, "core/workslip.html", {
                    "error": f"Couldn't read uploaded Workslip file: {e}",
                    "category": category,
                    "groups": groups,
                    "current_group": current_group,
                    "items_in_group": items_in_group,
                    "ws_estimate_rows": ws_estimate_rows,
                    "preview_rows": [],
                    "tp_percent": ws_tp_percent if ws_tp_percent else "",
                    "tp_type": ws_tp_type,
                    "supp_items_selected": ws_supp_items,
                    "work_name": ws_work_name,
                    "current_phase": ws_current_phase,
                    "previous_phases": ws_previous_phases,
                })

            # Find WorkSlip sheet
            ws_sheet = None
            for sh in wb_ws.worksheets:
                if "workslip" in sh.title.lower() or "working estimate" in str(sh.cell(row=1, column=1).value or "").lower():
                    ws_sheet = sh
                    break
            
            if not ws_sheet:
                ws_sheet = wb_ws.active

            # Detect phase number from workslip (look for "Phase X" in title or just count columns)
            # Parse the workslip to extract: item descriptions, estimated qty, and all phase execution data
            
            def detect_workslip_phase_and_columns(ws):
                """
                Detect the phase number and column structure of a workslip.
                Returns: (phase_number, column_map) where column_map has exec_qty columns per phase
                """
                # Find header row (usually row 8)
                header_row = 8
                for r in range(1, 15):
                    cell_val = str(ws.cell(row=r, column=1).value or "").strip().lower()
                    if "sl" in cell_val:
                        header_row = r
                        break
                
                # Count execution columns to determine phase
                # Standard: Sl, Desc, Unit, Est Qty, Est Rate, Est Amt, Exec Qty, Exec Rate, Exec Amt, More, Less, Remarks
                # Multi-phase: ... Phase1 Qty, Phase1 Amt, Phase2 Qty, Phase2 Amt, ...
                
                phase_count = 0
                col_map = {"header_row": header_row}
                
                for c in range(1, 30):
                    header = str(ws.cell(row=header_row, column=c).value or "").strip().lower()
                    if "execution" in header or "exec" in header:
                        if "qty" in header or "quantity" in header:
                            phase_count += 1
                            col_map[f"exec_qty_phase_{phase_count}"] = c
                        elif "amount" in header or "amt" in header:
                            col_map[f"exec_amt_phase_{phase_count}"] = c
                    elif "qty" in header and "est" in header:
                        col_map["est_qty"] = c
                    elif "description" in header or "desc" in header:
                        col_map["desc"] = c
                    elif "rate" in header and "est" in header:
                        col_map["est_rate"] = c
                
                # If no explicit phase columns found, assume standard single-phase layout
                if phase_count == 0:
                    phase_count = 1
                    col_map["exec_qty_phase_1"] = 7  # Column G
                    col_map["exec_amt_phase_1"] = 9  # Column I
                
                return phase_count, col_map
            
            phase_count, col_map = detect_workslip_phase_and_columns(ws_sheet)
            header_row = col_map.get("header_row", 8)
            
            # Parse rows from the workslip
            parsed_items = []
            new_estimate_rows = []
            phase_exec_maps = [{} for _ in range(phase_count)]  # One exec_map per phase
            phase_ae_data = [{} for _ in range(phase_count)]  # Track AE data separately per phase
            
            last_base_key = None  # Track the last base item for merging AE quantities
            
            for r in range(header_row + 1, ws_sheet.max_row + 1):
                desc_col = col_map.get("desc", 2)
                desc = str(ws_sheet.cell(row=r, column=desc_col).value or "").strip()
                
                if not desc:
                    continue
                    
                # Skip total/subtotal rows
                if any(kw in desc.upper() for kw in ("SUB TOTAL", "SUBTOTAL", "TOTAL", "GRAND TOTAL", "DEDUCT", "SUPPLEMENTAL")):
                    break
                
                # Check if this is an AE row (AE1, AE2, etc.)
                desc_upper = desc.upper().strip()
                is_ae_row = (desc_upper.startswith("AE") and len(desc_upper) >= 2 and 
                             (len(desc_upper) == 2 or desc_upper[2:].isdigit() or desc_upper[2] == ' '))
                
                # Get estimate qty and rate
                est_qty = ws_sheet.cell(row=r, column=col_map.get("est_qty", 4)).value or 0
                est_rate = ws_sheet.cell(row=r, column=col_map.get("est_rate", 5)).value or 0
                unit = ws_sheet.cell(row=r, column=3).value or ""
                
                try:
                    est_qty = float(est_qty)
                except:
                    est_qty = 0.0
                try:
                    est_rate = float(est_rate)
                except:
                    est_rate = 0.0
                
                if is_ae_row and last_base_key:
                    # This is an AE row - add its execution quantities to the previous base item
                    for p in range(1, phase_count + 1):
                        exec_qty_col = col_map.get(f"exec_qty_phase_{p}")
                        if exec_qty_col:
                            exec_qty = ws_sheet.cell(row=r, column=exec_qty_col).value or 0
                            try:
                                exec_qty = float(exec_qty)
                            except:
                                exec_qty = 0.0
                            if exec_qty > 0:
                                # Add AE qty to base item's exec_map
                                if last_base_key in phase_exec_maps[p-1]:
                                    phase_exec_maps[p-1][last_base_key] += exec_qty
                                else:
                                    phase_exec_maps[p-1][last_base_key] = exec_qty
                                # Also store the AE quantity separately for Excel output
                                ae_key = f"{last_base_key}:ae:{desc}"
                                phase_ae_data[p-1][ae_key] = exec_qty
                else:
                    # This is a base item row
                    row_key = f"phase_row_{r}"
                    last_base_key = row_key
                    
                    # Collect execution data from all phases
                    for p in range(1, phase_count + 1):
                        exec_qty_col = col_map.get(f"exec_qty_phase_{p}")
                        if exec_qty_col:
                            exec_qty = ws_sheet.cell(row=r, column=exec_qty_col).value or 0
                            try:
                                exec_qty = float(exec_qty)
                            except:
                                exec_qty = 0.0
                            if exec_qty > 0:
                                phase_exec_maps[p-1][row_key] = exec_qty
                    
                    new_estimate_rows.append({
                        "key": row_key,
                        "excel_row": r,
                        "item_name": desc,
                        "display_name": desc,
                        "desc": desc,
                        "qty_est": est_qty,
                        "unit": str(unit).strip(),
                        "rate": est_rate,
                    })
            
            if new_estimate_rows:
                # Store the estimate rows
                ws_estimate_rows = new_estimate_rows
                request.session["ws_estimate_rows"] = ws_estimate_rows
                
                # Store previous phases' execution data (with AE merged into base items)
                ws_previous_phases = phase_exec_maps
                request.session["ws_previous_phases"] = ws_previous_phases
                
                # Store previous phases' AE data separately for Excel output
                request.session["ws_previous_ae_data"] = phase_ae_data
                
                # Set current phase to next phase
                ws_current_phase = phase_count + 1
                request.session["ws_current_phase"] = ws_current_phase
                
                # Clear current exec_map for new phase entry
                request.session["ws_exec_map"] = {}
                
                logger.info(f"[WORKSLIP PHASE] Parsed workslip with {phase_count} phases, {len(new_estimate_rows)} items. Now phase {ws_current_phase}")
            
            return redirect(reverse('workslip_main') + '?preserve=1')

        # E) Download Workslip
        elif action == "download_workslip":
            exec_str = request.POST.get("exec_map", "")
            tp_percent_str = request.POST.get("tp_percent", "")
            tp_type = request.POST.get("tp_type", "Excess")
            deduct_old_material_str = request.POST.get("deduct_old_material", "")
            
            # Get metadata from form
            ws_work_name_form = request.POST.get("ws_work_name", "")
            ws_estimate_amount_form = request.POST.get("ws_estimate_amount", "")
            ws_agency_name_form = request.POST.get("ws_agency_name", "")
            ws_admin_sanction_form = request.POST.get("ws_admin_sanction", "")
            ws_tech_sanction_form = request.POST.get("ws_tech_sanction", "")
            ws_agreement_form = request.POST.get("ws_agreement", "")

            # merge UI exec_map into session map
            ws_exec_map_session = request.session.get("ws_exec_map", {}) or {}
            new_exec_map = {}
            if exec_str:
                try:
                    raw = json.loads(exec_str)
                    if isinstance(raw, dict):
                        for k, v in raw.items():
                            try:
                                new_exec_map[str(k)] = float(v)
                            except Exception:
                                continue
                except Exception:
                    pass
            ws_exec_map = ws_exec_map_session.copy()
            ws_exec_map.update(new_exec_map)

            # Use form values if provided, otherwise fall back to session values (from uploaded workslip)
            try:
                ws_tp_percent = float(tp_percent_str) if tp_percent_str != "" else request.session.get("ws_tp_percent", 0.0)
            except Exception:
                ws_tp_percent = request.session.get("ws_tp_percent", 0.0)
            ws_tp_type = tp_type if tp_type in ("Less", "Excess") else request.session.get("ws_tp_type", "Excess")
            
            # Parse Deduct Old Material Cost - use form value or fall back to session
            try:
                ws_deduct_old_material = float(deduct_old_material_str) if deduct_old_material_str != "" else request.session.get("ws_deduct_old_material", 0.0)
            except Exception:
                ws_deduct_old_material = request.session.get("ws_deduct_old_material", 0.0)

            request.session["ws_exec_map"] = ws_exec_map
            request.session["ws_tp_percent"] = ws_tp_percent
            request.session["ws_tp_type"] = ws_tp_type
            request.session["ws_deduct_old_material"] = ws_deduct_old_material
            
            # Update metadata in session from form values
            ws_metadata_session = request.session.get("ws_metadata", {}) or {}
            if ws_work_name_form:
                ws_metadata_session["work_name"] = ws_work_name_form
            if ws_estimate_amount_form:
                ws_metadata_session["estimate_amount"] = ws_estimate_amount_form
            if ws_agency_name_form:
                ws_metadata_session["agency_name"] = ws_agency_name_form
            if ws_admin_sanction_form:
                ws_metadata_session["admin_sanction"] = ws_admin_sanction_form
            if ws_tech_sanction_form:
                ws_metadata_session["tech_sanction"] = ws_tech_sanction_form
            if ws_agreement_form:
                ws_metadata_session["agreement"] = ws_agreement_form
            # Also update TP values in metadata
            ws_metadata_session["tp_percent"] = ws_tp_percent
            ws_metadata_session["tp_type"] = ws_tp_type
            request.session["ws_metadata"] = ws_metadata_session
            request.session.modified = True

            # helper to safely fetch execution quantity for base estimate rows
            def get_exec_qty_for_base(row_key, item_name, desc):
                candidates = [
                    f"base:{row_key}",
                    row_key,
                    item_name or "",
                    desc or "",
                ]
                for k in candidates:
                    k = str(k).strip()
                    if not k:
                        continue
                    if k in ws_exec_map:
                        try:
                            return float(ws_exec_map[k])
                        except Exception:
                            return 0.0
                return 0.0

            # ---------- build supplemental description+rate from backend ----------
            item_to_info = {it["name"]: it for it in items_list}
            wb_backend_vals = None
            ws_backend_vals = None
            if filepath and os.path.exists(filepath):
                try:
                    wb_backend_vals = load_workbook(filepath, data_only=True)
                    ws_backend_vals = wb_backend_vals["Master Datas"]
                except Exception:
                    wb_backend_vals = None
                    ws_backend_vals = None

            supp_desc_map = {}
            supp_rate_map = {}
            for name in ws_supp_items:
                info = item_to_info.get(name)
                if not info or ws_backend_vals is None or ws_data is None:
                    continue
                start_row = info["start_row"]
                end_row = info["end_row"]
                # Description: 2nd row below yellow header in col D
                desc_cell = ws_data.cell(row=start_row + 2, column=4).value
                supp_desc_map[name] = str(desc_cell or "").strip()
                # Rate from Master Datas col J
                rate_val = 0.0
                for r in range(end_row, start_row - 1, -1):
                    v = ws_backend_vals.cell(row=r, column=10).value
                    if v not in (None, ""):
                        try:
                            rate_val = float(v)
                        except Exception:
                            rate_val = 0.0
                        break
                supp_rate_map[name] = rate_val

            # ---------- create workbook ----------
            wb_out = Workbook()

            # Sheet 1: ItemBlocks (only if supplemental items exist)
            if ws_supp_items:
                ws_blocks = wb_out.active
                ws_blocks.title = "ItemBlocks"
                current_row = 1
                if ws_data is not None:
                    for name in ws_supp_items:
                        info = item_to_info.get(name)
                        if not info:
                            continue
                        start_row = info["start_row"]
                        end_row = info["end_row"]
                        copy_block_with_styles_and_formulas(
                            ws_src=ws_data,
                            ws_dst=ws_blocks,
                            src_min_row=start_row,
                            src_max_row=end_row,
                            dst_start_row=current_row,
                            col_start=1,
                            col_end=10,
                        )
                        current_row += (end_row - start_row + 1) + 1  # 1 blank row between blocks
            else:
                # no supplemental items â†’ remove default sheet so only WorkSlip will exist
                default_sheet = wb_out.active
                wb_out.remove(default_sheet)

            # Sheet 2 (or 1 if no ItemBlocks): WorkSlip
            ws_ws = wb_out.create_sheet("WorkSlip")

            thin = Side(border_style="thin", color="000000")
            border_all = Border(left=thin, right=thin, top=thin, bottom=thin)
            header_fill = PatternFill("solid", fgColor="FFC8C8C8")
            subtotal_fill = PatternFill("solid", fgColor="FFE6E6E6")
            supp_fill = PatternFill("solid", fgColor="FFF5E1")
            phase_fill = PatternFill("solid", fgColor="FFFEF3C7")  # Amber for previous phases
            current_phase_fill = PatternFill("solid", fgColor="FFDBEAFE")  # Blue for current phase
            
            # Get phase data
            ws_previous_phases = request.session.get("ws_previous_phases", []) or []
            ws_current_phase = request.session.get("ws_current_phase", 1)
            num_previous_phases = len(ws_previous_phases)
            
            # Calculate total columns: Base 11 + 2 per previous phase (Qty + Amount)
            # Columns: Sl, Desc, Unit, Est Qty, Est Rate, Est Amt, [Phase1 Qty, Phase1 Amt, ...], Curr Qty, Curr Amt, More, Less, Remarks
            extra_phase_cols = num_previous_phases * 2
            total_cols = 11 + extra_phase_cols
            
            # Build column letter for last column
            def col_letter(n):
                result = ""
                while n > 0:
                    n, remainder = divmod(n - 1, 26)
                    result = chr(65 + remainder) + result
                return result
            
            last_col_letter = col_letter(total_cols)

            # Top main heading - merge across all columns
            ws_ws.merge_cells(f"A1:{last_col_letter}1")
            c = ws_ws["A1"]
            phase_title = f"WORKING ESTIMATE-{ws_current_phase}" if ws_current_phase > 1 else "WORKING ESTIMATE"
            c.value = phase_title
            c.font = Font(bold=True, size=14)
            c.alignment = Alignment(horizontal="center", vertical="center")

            # Get stored metadata from previous workslip
            ws_metadata = request.session.get("ws_metadata", {})
            
            # 6 merged rows below heading - use values from uploaded workslip if available
            work_name_val = ws_metadata.get("work_name", "") or ws_work_name or ""
            estimate_amount_val = ws_metadata.get("estimate_amount", "") or ""
            admin_sanction_val = ws_metadata.get("admin_sanction", "") or ""
            tech_sanction_val = ws_metadata.get("tech_sanction", "") or ""
            agreement_val = ws_metadata.get("agreement", "") or ""
            agency_name_val = ws_metadata.get("agency_name", "") or ""
            
            headings = [
                f"Name of the work : {work_name_val}" if work_name_val else "Name of the work :",
                f"Estimate Amount : {estimate_amount_val}" if estimate_amount_val else "Estimate Amount :",
                f"Ref. to Administrative sanction : {admin_sanction_val}" if admin_sanction_val else "Ref. to Administrative sanction :",
                f"Ref. to Technical sanction : {tech_sanction_val}" if tech_sanction_val else "Ref. to Technical sanction :",
                f"Ref. to Agreement : {agreement_val}" if agreement_val else "Ref. to Agreement :",
                f"Name of the Agency : {agency_name_val}" if agency_name_val else "Name of the Agency :",
            ]
            for i, text in enumerate(headings, start=2):
                merge_range = f"A{i}:{last_col_letter}{i}"
                ws_ws.merge_cells(merge_range)
                cell = ws_ws[f"A{i}"]
                cell.value = text
                cell.font = Font(bold=True)
                cell.alignment = Alignment(horizontal="left", vertical="center")

            # Apply borders to heading area
            for row_idx in range(1, 8):
                for col_idx in range(1, total_cols + 1):
                    cell = ws_ws.cell(row=row_idx, column=col_idx)
                    cell.border = border_all

            # Table header - dynamic based on phases
            header_row = 8
            # Base columns (first 6)
            base_cols = ["Sl.No", "Description of Item", "Unit", "Qty (Estimate)", "Rate (Estimate)", "Amount (Estimate)"]
            
            # Previous phase columns
            phase_cols = []
            for p in range(1, num_previous_phases + 1):
                phase_cols.extend([f"Workslip-{p} Qty", f"Workslip-{p} Amt"])
            
            # Current phase columns (last 5 - no Rate since it's same as Estimate Rate)
            if ws_current_phase > 1:
                current_cols = [f"Workslip-{ws_current_phase} Qty", f"Workslip-{ws_current_phase} Amt", "More", "Less", "Remarks"]
            else:
                current_cols = ["Qty (Execution)", "Amount (Execution)", "More", "Less", "Remarks"]
            
            all_header_cols = base_cols + phase_cols + current_cols
            
            for col_idx, text in enumerate(all_header_cols, start=1):
                cell = ws_ws.cell(row=header_row, column=col_idx, value=text)
                cell.font = Font(bold=True)
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                cell.border = border_all
                cell.fill = header_fill
                # Color phase columns differently
                if col_idx > 6 and col_idx <= 6 + extra_phase_cols:
                    cell.fill = phase_fill
                elif col_idx > 6 + extra_phase_cols and col_idx <= 6 + extra_phase_cols + 2:
                    if ws_current_phase > 1:
                        cell.fill = current_phase_fill

            # Column widths
            ws_ws.column_dimensions["A"].width = 6
            ws_ws.column_dimensions["B"].width = 70
            ws_ws.column_dimensions["C"].width = 10
            ws_ws.column_dimensions["D"].width = 14
            ws_ws.column_dimensions["E"].width = 12
            ws_ws.column_dimensions["F"].width = 14
            
            # Previous phase columns
            for i in range(extra_phase_cols):
                col_l = col_letter(7 + i)
                ws_ws.column_dimensions[col_l].width = 12
            
            # Current execution columns (no Rate - uses Estimate Rate)
            exec_start_col = 7 + extra_phase_cols
            ws_ws.column_dimensions[col_letter(exec_start_col)].width = 14      # Qty
            ws_ws.column_dimensions[col_letter(exec_start_col + 1)].width = 14  # Amount
            ws_ws.column_dimensions[col_letter(exec_start_col + 2)].width = 10  # More
            ws_ws.column_dimensions[col_letter(exec_start_col + 3)].width = 10  # Less
            ws_ws.column_dimensions[col_letter(exec_start_col + 4)].width = 25  # Remarks

            # remark helper
            def remark_for_item(q_est, q_exec, is_supp=False, has_ae_split=False):
                if is_supp:
                    return "Proposed as per site condition"
                if q_exec == 0:
                    return "Deleted"
                # If we're splitting into base + AE rows, base row should not say "Excess"
                if has_ae_split:
                    return ""  # Base row gets no remark when there's an AE row
                if q_exec > q_est:
                    return "Excess as per estimated"
                if 0 < q_exec < q_est:
                    return "Less as per estimated"
                return ""

            data_start = header_row + 1
            out_row = data_start
            sl_counter = 1
            ae_counter = 1
            
            # Column indices for data (no COL_CURR_RATE - uses Estimate Rate)
            COL_SL = 1
            COL_DESC = 2
            COL_UNIT = 3
            COL_EST_QTY = 4
            COL_EST_RATE = 5
            COL_EST_AMT = 6
            COL_PHASE_START = 7  # Previous phases start here
            COL_CURR_QTY = 7 + extra_phase_cols
            COL_CURR_AMT = 8 + extra_phase_cols
            COL_MORE = 9 + extra_phase_cols
            COL_LESS = 10 + extra_phase_cols
            COL_REMARKS = 11 + extra_phase_cols

            # Get previous phases' AE data and supplemental items
            ws_previous_ae_data = request.session.get("ws_previous_ae_data", [])
            ws_previous_supp_items = request.session.get("ws_previous_supp_items", [])
            
            # ---- Base Estimate items with row-splitting ----
            for row in ws_estimate_rows:
                row_key = row["key"]
                qty_est = float(row.get("qty_est", 0) or 0)

                qty_exec = get_exec_qty_for_base(
                    row_key=row_key,
                    item_name=row.get("item_name"),
                    desc=row.get("desc"),
                )

                unit = row.get("unit") or ""
                rate = float(row.get("rate", 0) or 0)
                desc_est = row.get("desc") or row.get("item_name") or ""
                
                # Get previous phases' execution quantities for this row (AE already merged)
                prev_phase_qtys = []
                for phase_map in ws_previous_phases:
                    prev_qty = phase_map.get(row_key, 0)
                    try:
                        prev_qty = float(prev_qty)
                    except:
                        prev_qty = 0.0
                    prev_phase_qtys.append(prev_qty)
                
                # Calculate excess for each previous phase
                prev_phase_excess = []
                for p_qty in prev_phase_qtys:
                    excess = max(0, p_qty - qty_est) if qty_est > 0 else 0
                    prev_phase_excess.append(excess)
                
                # Calculate base qty for each previous phase (capped at estimate)
                prev_base_qtys = []
                for p_qty in prev_phase_qtys:
                    base_qty = min(p_qty, qty_est) if qty_est > 0 else p_qty
                    prev_base_qtys.append(base_qty)
                
                # Calculate current phase excess
                current_excess = max(0, qty_exec - qty_est) if qty_est > 0 else 0
                current_base_qty = min(qty_exec, qty_est) if qty_est > 0 else qty_exec
                
                # Check if any phase (previous or current) has excess - if so, we need ONE AE row
                has_any_excess = any(e > 0 for e in prev_phase_excess) or current_excess > 0

                # FIRST: Always write the base row
                ws_ws.cell(out_row, COL_SL, sl_counter)
                ws_ws.cell(out_row, COL_DESC, desc_est)
                ws_ws.cell(out_row, COL_UNIT, unit)
                ws_ws.cell(out_row, COL_EST_QTY, qty_est)
                ws_ws.cell(out_row, COL_EST_RATE, rate)
                ws_ws.cell(out_row, COL_EST_AMT, f"={col_letter(COL_EST_QTY)}{out_row}*{col_letter(COL_EST_RATE)}{out_row}")
                
                # Previous phases' data - show base qty (capped at estimate)
                for p_idx, p_qty in enumerate(prev_base_qtys):
                    phase_qty_col = COL_PHASE_START + (p_idx * 2)
                    phase_amt_col = phase_qty_col + 1
                    ws_ws.cell(out_row, phase_qty_col, p_qty)
                    ws_ws.cell(out_row, phase_amt_col, f"={col_letter(phase_qty_col)}{out_row}*{col_letter(COL_EST_RATE)}{out_row}")
                    ws_ws.cell(out_row, phase_qty_col).fill = phase_fill
                    ws_ws.cell(out_row, phase_amt_col).fill = phase_fill
                
                # Current execution (base qty capped at estimate if there's excess)
                ws_ws.cell(out_row, COL_CURR_QTY, current_base_qty)
                ws_ws.cell(out_row, COL_CURR_AMT, f"={col_letter(COL_CURR_QTY)}{out_row}*{col_letter(COL_EST_RATE)}{out_row}")
                ws_ws.cell(out_row, COL_MORE, f"=IF({col_letter(COL_CURR_AMT)}{out_row}>{col_letter(COL_EST_AMT)}{out_row},{col_letter(COL_CURR_AMT)}{out_row}-{col_letter(COL_EST_AMT)}{out_row},\"\")")
                ws_ws.cell(out_row, COL_LESS, f"=IF({col_letter(COL_EST_AMT)}{out_row}>{col_letter(COL_CURR_AMT)}{out_row},{col_letter(COL_EST_AMT)}{out_row}-{col_letter(COL_CURR_AMT)}{out_row},\"\")")
                ws_ws.cell(out_row, COL_REMARKS, remark_for_item(qty_est, qty_exec, is_supp=False, has_ae_split=has_any_excess))

                for cidx in range(1, total_cols + 1):
                    cell = ws_ws.cell(out_row, cidx)
                    cell.border = border_all
                    if cidx in (COL_DESC, COL_REMARKS):
                        cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
                    else:
                        cell.alignment = Alignment(horizontal="center", vertical="center")

                out_row += 1
                sl_counter += 1
                base_row_for_rate = out_row - 1  # Save for rate reference
                
                # SECOND: If any phase has excess, write ONE AE row with all phase excess values
                if has_any_excess:
                    ae_label = f"AE{ae_counter}"
                    ae_counter += 1

                    ws_ws.cell(out_row, COL_SL, "")          # no serial number
                    ws_ws.cell(out_row, COL_DESC, ae_label)
                    ws_ws.cell(out_row, COL_UNIT, unit)
                    ws_ws.cell(out_row, COL_EST_QTY, None)
                    ws_ws.cell(out_row, COL_EST_RATE, None)
                    ws_ws.cell(out_row, COL_EST_AMT, None)
                    
                    # Previous phases - show excess for each phase in respective columns
                    for pi, excess_qty in enumerate(prev_phase_excess):
                        phase_qty_col = COL_PHASE_START + (pi * 2)
                        phase_amt_col = phase_qty_col + 1
                        if excess_qty > 0:
                            ws_ws.cell(out_row, phase_qty_col, excess_qty)
                            ws_ws.cell(out_row, phase_amt_col, f"={col_letter(phase_qty_col)}{out_row}*{col_letter(COL_EST_RATE)}{base_row_for_rate}")
                        else:
                            ws_ws.cell(out_row, phase_qty_col, None)
                            ws_ws.cell(out_row, phase_amt_col, None)
                        ws_ws.cell(out_row, phase_qty_col).fill = phase_fill
                        ws_ws.cell(out_row, phase_amt_col).fill = phase_fill
                    
                    # Current phase excess
                    if current_excess > 0:
                        ws_ws.cell(out_row, COL_CURR_QTY, current_excess)
                        ws_ws.cell(out_row, COL_CURR_AMT, f"={col_letter(COL_CURR_QTY)}{out_row}*{col_letter(COL_EST_RATE)}{base_row_for_rate}")
                    else:
                        ws_ws.cell(out_row, COL_CURR_QTY, None)
                        ws_ws.cell(out_row, COL_CURR_AMT, None)
                    
                    ws_ws.cell(out_row, COL_MORE, f"=IF({col_letter(COL_CURR_AMT)}{out_row}>{col_letter(COL_EST_AMT)}{out_row},{col_letter(COL_CURR_AMT)}{out_row}-{col_letter(COL_EST_AMT)}{out_row},\"\")")
                    ws_ws.cell(out_row, COL_LESS, "")
                    ws_ws.cell(out_row, COL_REMARKS, "Excess as per estimated")

                    for cidx in range(1, total_cols + 1):
                        cell = ws_ws.cell(out_row, cidx)
                        cell.border = border_all
                        if cidx in (COL_DESC, COL_REMARKS):
                            cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
                        else:
                            cell.alignment = Alignment(horizontal="center", vertical="center")

                    out_row += 1

            # ---- Previous Workslip Supplemental items (per phase) ----
            if ws_previous_supp_items:
                # Group previous supplemental items by phase
                supp_by_phase = {}
                for supp in ws_previous_supp_items:
                    phase_num = supp.get("phase", 1)
                    if phase_num not in supp_by_phase:
                        supp_by_phase[phase_num] = []
                    supp_by_phase[phase_num].append(supp)
                
                # Output each phase's supplemental items
                for phase_num in sorted(supp_by_phase.keys()):
                    phase_supps = supp_by_phase[phase_num]
                    
                    # Heading row for this phase's supplemental items
                    supp_phase_header = f"Supplemental Items-{phase_num}"
                    supp_cell = ws_ws.cell(out_row, COL_DESC, supp_phase_header)
                    supp_cell.font = Font(bold=True, color="FF0000")  # Red text
                    for col in range(1, total_cols + 1):
                        cell = ws_ws.cell(out_row, col)
                        cell.border = border_all
                        cell.fill = supp_fill
                        if col == COL_DESC:
                            cell.alignment = Alignment(horizontal="left", vertical="center")
                        else:
                            cell.alignment = Alignment(horizontal="center", vertical="center")
                    out_row += 1
                    
                    # Output each supplemental item from this phase
                    for supp in phase_supps:
                        supp_name = supp.get("name", "")
                        supp_qty = float(supp.get("qty", 0) or 0)
                        supp_desc = supp.get("desc", supp_name)
                        supp_unit = supp.get("unit", "-") or "-"
                        supp_rate = float(supp.get("rate", 0) or 0)
                        supp_amount = supp.get("amount", supp_qty * supp_rate)
                        
                        ws_ws.cell(out_row, COL_SL, sl_counter)
                        ws_ws.cell(out_row, COL_DESC, supp_name)
                        ws_ws.cell(out_row, COL_UNIT, supp_unit)
                        ws_ws.cell(out_row, COL_EST_QTY, None)
                        ws_ws.cell(out_row, COL_EST_RATE, supp_rate if supp_rate > 0 else None)
                        ws_ws.cell(out_row, COL_EST_AMT, None)
                        
                        # Put quantity in the correct phase column
                        for p_idx in range(num_previous_phases):
                            phase_qty_col = COL_PHASE_START + (p_idx * 2)
                            phase_amt_col = phase_qty_col + 1
                            if (p_idx + 1) == phase_num:
                                ws_ws.cell(out_row, phase_qty_col, supp_qty)
                                # Amount calculated from rate
                                ws_ws.cell(out_row, phase_amt_col, supp_amount if supp_rate > 0 else None)
                            else:
                                ws_ws.cell(out_row, phase_qty_col, None)
                                ws_ws.cell(out_row, phase_amt_col, None)
                            ws_ws.cell(out_row, phase_qty_col).fill = phase_fill
                            ws_ws.cell(out_row, phase_amt_col).fill = phase_fill
                        
                        # Check if user entered current workslip quantity for this previous supp item
                        prev_supp_key = f"prev_supp:{phase_num}:{supp_name}"
                        prev_supp_curr_qty = float(ws_exec_map.get(prev_supp_key, 0) or 0)
                        
                        if prev_supp_curr_qty > 0:
                            prev_supp_curr_amt = prev_supp_curr_qty * supp_rate
                            ws_ws.cell(out_row, COL_CURR_QTY, prev_supp_curr_qty)
                            ws_ws.cell(out_row, COL_CURR_AMT, prev_supp_curr_amt)
                            ws_ws.cell(out_row, COL_MORE, prev_supp_curr_amt)  # All extra work for prev supp items goes to More
                            ws_ws.cell(out_row, COL_LESS, None)
                            ws_ws.cell(out_row, COL_REMARKS, "Proposed as per site condition")
                        else:
                            ws_ws.cell(out_row, COL_CURR_QTY, None)
                            ws_ws.cell(out_row, COL_CURR_AMT, None)
                            ws_ws.cell(out_row, COL_MORE, None)
                            ws_ws.cell(out_row, COL_LESS, None)
                            ws_ws.cell(out_row, COL_REMARKS, f"From Workslip-{phase_num}")
                        
                        for cidx in range(1, total_cols + 1):
                            cell = ws_ws.cell(out_row, cidx)
                            cell.border = border_all
                            if cidx in (COL_DESC, COL_REMARKS):
                                cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
                            else:
                                cell.alignment = Alignment(horizontal="center", vertical="center")
                        
                        out_row += 1
                        sl_counter += 1
            
            # ---- Current Phase Supplemental items block ----
            if ws_supp_items:
                # heading row with phase-specific name
                supp_header_text = f"Supplemental Items-{ws_current_phase}" if ws_current_phase > 1 else "Supplemental Items"
                supp_cell = ws_ws.cell(out_row, COL_DESC, supp_header_text)
                supp_cell.font = Font(bold=True)
                for col in range(1, total_cols + 1):
                    cell = ws_ws.cell(out_row, col)
                    cell.border = border_all
                    cell.fill = supp_fill
                    if col == COL_DESC:
                        cell.alignment = Alignment(horizontal="left", vertical="center")
                    else:
                        cell.alignment = Alignment(horizontal="center", vertical="center")
                out_row += 1

                # actual supplemental rows
                for name in ws_supp_items:
                    desc_supp = supp_desc_map.get(name, name)
                    unit_pl, _ = units_for(name)
                    rate = float(supp_rate_map.get(name, 0.0) or 0.0)
                    key = f"supp:{name}"
                    qty_exec = float(ws_exec_map.get(key, 0) or 0)

                    ws_ws.cell(out_row, COL_SL, sl_counter)
                    ws_ws.cell(out_row, COL_DESC, desc_supp)
                    ws_ws.cell(out_row, COL_UNIT, unit_pl)
                    ws_ws.cell(out_row, COL_EST_QTY, None)
                    ws_ws.cell(out_row, COL_EST_RATE, rate)  # Put supp rate in Est Rate column for amount calculation
                    ws_ws.cell(out_row, COL_EST_AMT, None)
                    
                    # Previous phases empty for current supplemental
                    for p_idx in range(num_previous_phases):
                        phase_qty_col = COL_PHASE_START + (p_idx * 2)
                        phase_amt_col = phase_qty_col + 1
                        ws_ws.cell(out_row, phase_qty_col, None)
                        ws_ws.cell(out_row, phase_amt_col, None)
                        ws_ws.cell(out_row, phase_qty_col).fill = phase_fill
                        ws_ws.cell(out_row, phase_amt_col).fill = phase_fill
                    
                    ws_ws.cell(out_row, COL_CURR_QTY, qty_exec)
                    ws_ws.cell(out_row, COL_CURR_AMT, f"={col_letter(COL_CURR_QTY)}{out_row}*{col_letter(COL_EST_RATE)}{out_row}")
                    ws_ws.cell(out_row, COL_MORE, f"=IF({col_letter(COL_CURR_AMT)}{out_row}>{col_letter(COL_EST_AMT)}{out_row},{col_letter(COL_CURR_AMT)}{out_row}-{col_letter(COL_EST_AMT)}{out_row},\"\")")
                    ws_ws.cell(out_row, COL_LESS, f"=IF({col_letter(COL_EST_AMT)}{out_row}>{col_letter(COL_CURR_AMT)}{out_row},{col_letter(COL_EST_AMT)}{out_row}-{col_letter(COL_CURR_AMT)}{out_row},\"\")")
                    ws_ws.cell(out_row, COL_REMARKS, remark_for_item(0, qty_exec, is_supp=True))

                    for cidx in range(1, total_cols + 1):
                        cell = ws_ws.cell(out_row, cidx)
                        cell.border = border_all
                        if cidx in (COL_DESC, COL_REMARKS):
                            cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
                        else:
                            cell.alignment = Alignment(horizontal="center", vertical="center")

                    out_row += 1
                    sl_counter += 1

            # ---- Sub Total row (over all items) ----
            sub_row = out_row
            ws_ws.cell(sub_row, COL_DESC, "Sub Total Amount")
            ws_ws.cell(sub_row, COL_EST_AMT, f"=SUM({col_letter(COL_EST_AMT)}{data_start}:{col_letter(COL_EST_AMT)}{sub_row-1})")
            
            # Previous phases' subtotals
            for p_idx in range(num_previous_phases):
                phase_amt_col = COL_PHASE_START + (p_idx * 2) + 1  # Amount column for this phase
                phase_amt_letter = col_letter(phase_amt_col)
                ws_ws.cell(sub_row, phase_amt_col, f"=SUM({phase_amt_letter}{data_start}:{phase_amt_letter}{sub_row-1})")
                ws_ws.cell(sub_row, phase_amt_col).fill = phase_fill
            
            ws_ws.cell(sub_row, COL_CURR_AMT, f"=SUM({col_letter(COL_CURR_AMT)}{data_start}:{col_letter(COL_CURR_AMT)}{sub_row-1})")
            # More / Less for Sub Total row
            ws_ws.cell(sub_row, COL_MORE, f"=SUM({col_letter(COL_MORE)}{data_start}:{col_letter(COL_MORE)}{sub_row-1})")
            ws_ws.cell(sub_row, COL_LESS, f"=SUM({col_letter(COL_LESS)}{data_start}:{col_letter(COL_LESS)}{sub_row-1})")

            for col in range(1, total_cols + 1):
                cell = ws_ws.cell(sub_row, col)
                cell.font = Font(bold=True)
                cell.border = border_all
                cell.fill = subtotal_fill
                if col == COL_DESC:
                    cell.alignment = Alignment(horizontal="left", vertical="center")
                else:
                    cell.alignment = Alignment(horizontal="center", vertical="center")

            # ---- Rows below Sub Total ----
            # Determine if we need a deduct row
            current_row = sub_row
            deduct_row = None
            if ws_deduct_old_material > 0:
                current_row += 1
                deduct_row = current_row
            
            tp_row    = current_row + 1
            sub1_row  = current_row + 2
            lc_row    = current_row + 3
            qc_row    = current_row + 4
            nac_row   = current_row + 5
            sub2_row  = current_row + 6
            gst_row   = current_row + 7
            unused_row= current_row + 8
            ls_row    = current_row + 9
            grand_row = current_row + 10
            
            # Column letters for dynamic formulas
            EST_AMT_COL = col_letter(COL_EST_AMT)
            CURR_AMT_COL = col_letter(COL_CURR_AMT)
            MORE_COL = col_letter(COL_MORE)
            LESS_COL = col_letter(COL_LESS)
            
            # Helper to get phase amount column letter
            def phase_amt_letter(p_idx):
                return col_letter(COL_PHASE_START + (p_idx * 2) + 1)
            
            # Add Deduct Old Material Cost row (if applicable)
            if deduct_row:
                ws_ws.cell(deduct_row, COL_DESC, "Deduct Old Material Cost")
                ws_ws.cell(deduct_row, COL_EST_AMT, -ws_deduct_old_material)  # Estimate - negative
                # Previous phases - same deduction
                for p_idx in range(num_previous_phases):
                    phase_amt_col = COL_PHASE_START + (p_idx * 2) + 1
                    ws_ws.cell(deduct_row, phase_amt_col, -ws_deduct_old_material)
                    ws_ws.cell(deduct_row, phase_amt_col).fill = phase_fill
                ws_ws.cell(deduct_row, COL_CURR_AMT, -ws_deduct_old_material)  # Execution - negative
                ws_ws.cell(deduct_row, COL_MORE, "")  # More
                ws_ws.cell(deduct_row, COL_LESS, "")  # Less

            # i) Add / Deduct T.P
            tp_label_prefix = "Add" if ws_tp_type == "Excess" else "Deduct"
            ws_ws.cell(tp_row, COL_DESC, f"{tp_label_prefix} T.P @ {ws_tp_percent} % {ws_tp_type}")
            ws_ws.cell(tp_row, COL_EST_AMT, None)  # Estimate MUST be empty

            # Previous phases TP
            for p_idx in range(num_previous_phases):
                p_amt_col = COL_PHASE_START + (p_idx * 2) + 1
                p_amt_letter = phase_amt_letter(p_idx)
                if deduct_row:
                    if ws_tp_type == "Excess":
                        ws_ws.cell(tp_row, p_amt_col, f"=({p_amt_letter}{sub_row}+{p_amt_letter}{deduct_row})*{ws_tp_percent}/100")
                    else:
                        ws_ws.cell(tp_row, p_amt_col, f"=-({p_amt_letter}{sub_row}+{p_amt_letter}{deduct_row})*{ws_tp_percent}/100")
                else:
                    if ws_tp_type == "Excess":
                        ws_ws.cell(tp_row, p_amt_col, f"={p_amt_letter}{sub_row}*{ws_tp_percent}/100")
                    else:
                        ws_ws.cell(tp_row, p_amt_col, f"=-{p_amt_letter}{sub_row}*{ws_tp_percent}/100")
                ws_ws.cell(tp_row, p_amt_col).fill = phase_fill

            # Current Amount: positive if Excess, negative if Less
            if deduct_row:
                if ws_tp_type == "Excess":
                    ws_ws.cell(tp_row, COL_CURR_AMT, f"=({CURR_AMT_COL}{sub_row}+{CURR_AMT_COL}{deduct_row})*{ws_tp_percent}/100")
                else:
                    ws_ws.cell(tp_row, COL_CURR_AMT, f"=-({CURR_AMT_COL}{sub_row}+{CURR_AMT_COL}{deduct_row})*{ws_tp_percent}/100")
            else:
                if ws_tp_type == "Excess":
                    ws_ws.cell(tp_row, COL_CURR_AMT, f"={CURR_AMT_COL}{sub_row}*{ws_tp_percent}/100")
                else:
                    ws_ws.cell(tp_row, COL_CURR_AMT, f"=-{CURR_AMT_COL}{sub_row}*{ws_tp_percent}/100")

            # More / Less for TP row
            ws_ws.cell(tp_row, COL_MORE, f"=IF({CURR_AMT_COL}{tp_row}>{EST_AMT_COL}{tp_row},{CURR_AMT_COL}{tp_row}-{EST_AMT_COL}{tp_row},\"\")")
            ws_ws.cell(tp_row, COL_LESS, f"=IF({EST_AMT_COL}{tp_row}>{CURR_AMT_COL}{tp_row},{EST_AMT_COL}{tp_row}-{CURR_AMT_COL}{tp_row},\"\")")

            # ii) Sub Total 1 - includes deduction if present
            ws_ws.cell(sub1_row, COL_DESC, "Sub Total 1")
            if deduct_row:
                ws_ws.cell(sub1_row, COL_EST_AMT, f"={EST_AMT_COL}{sub_row}+{EST_AMT_COL}{deduct_row}")
                # Previous phases Sub Total 1
                for p_idx in range(num_previous_phases):
                    p_amt_col = COL_PHASE_START + (p_idx * 2) + 1
                    p_amt_letter = phase_amt_letter(p_idx)
                    ws_ws.cell(sub1_row, p_amt_col, f"={p_amt_letter}{sub_row}+{p_amt_letter}{deduct_row}+{p_amt_letter}{tp_row}")
                    ws_ws.cell(sub1_row, p_amt_col).fill = phase_fill
                ws_ws.cell(sub1_row, COL_CURR_AMT, f"={CURR_AMT_COL}{sub_row}+{CURR_AMT_COL}{deduct_row}+{CURR_AMT_COL}{tp_row}")
            else:
                ws_ws.cell(sub1_row, COL_EST_AMT, f"={EST_AMT_COL}{sub_row}")
                # Previous phases Sub Total 1
                for p_idx in range(num_previous_phases):
                    p_amt_col = COL_PHASE_START + (p_idx * 2) + 1
                    p_amt_letter = phase_amt_letter(p_idx)
                    ws_ws.cell(sub1_row, p_amt_col, f"={p_amt_letter}{sub_row}+{p_amt_letter}{tp_row}")
                    ws_ws.cell(sub1_row, p_amt_col).fill = phase_fill
                ws_ws.cell(sub1_row, COL_CURR_AMT, f"={CURR_AMT_COL}{sub_row}+{CURR_AMT_COL}{tp_row}")

            # iii) Add LC @ 1%
            ws_ws.cell(lc_row, COL_DESC, "Add LC @ 1%")
            ws_ws.cell(lc_row, COL_EST_AMT, f"={EST_AMT_COL}{sub1_row}*0.01")
            for p_idx in range(num_previous_phases):
                p_amt_col = COL_PHASE_START + (p_idx * 2) + 1
                p_amt_letter = phase_amt_letter(p_idx)
                ws_ws.cell(lc_row, p_amt_col, f"={p_amt_letter}{sub1_row}*0.01")
                ws_ws.cell(lc_row, p_amt_col).fill = phase_fill
            ws_ws.cell(lc_row, COL_CURR_AMT, f"={CURR_AMT_COL}{sub1_row}*0.01")
            ws_ws.cell(lc_row, COL_MORE, f"=IF({CURR_AMT_COL}{lc_row}>{EST_AMT_COL}{lc_row},{CURR_AMT_COL}{lc_row}-{EST_AMT_COL}{lc_row},\"\")")
            ws_ws.cell(lc_row, COL_LESS, f"=IF({EST_AMT_COL}{lc_row}>{CURR_AMT_COL}{lc_row},{EST_AMT_COL}{lc_row}-{CURR_AMT_COL}{lc_row},\"\")")

            # iv) Add QC @ 1%
            ws_ws.cell(qc_row, COL_DESC, "Add QC @ 1%")
            ws_ws.cell(qc_row, COL_EST_AMT, f"={EST_AMT_COL}{sub1_row}*0.01")
            for p_idx in range(num_previous_phases):
                p_amt_col = COL_PHASE_START + (p_idx * 2) + 1
                p_amt_letter = phase_amt_letter(p_idx)
                ws_ws.cell(qc_row, p_amt_col, f"={p_amt_letter}{sub1_row}*0.01")
                ws_ws.cell(qc_row, p_amt_col).fill = phase_fill
            ws_ws.cell(qc_row, COL_CURR_AMT, f"={CURR_AMT_COL}{sub1_row}*0.01")
            ws_ws.cell(qc_row, COL_MORE, f"=IF({CURR_AMT_COL}{qc_row}>{EST_AMT_COL}{qc_row},{CURR_AMT_COL}{qc_row}-{EST_AMT_COL}{qc_row},\"\")")
            ws_ws.cell(qc_row, COL_LESS, f"=IF({EST_AMT_COL}{qc_row}>{CURR_AMT_COL}{qc_row},{EST_AMT_COL}{qc_row}-{CURR_AMT_COL}{qc_row},\"\")")

            # v) Add NAC chargers @ 0.1%
            ws_ws.cell(nac_row, COL_DESC, "Add NAC chargers @ 0.1 %")
            ws_ws.cell(nac_row, COL_EST_AMT, f"={EST_AMT_COL}{sub1_row}*0.001")
            for p_idx in range(num_previous_phases):
                p_amt_col = COL_PHASE_START + (p_idx * 2) + 1
                p_amt_letter = phase_amt_letter(p_idx)
                ws_ws.cell(nac_row, p_amt_col, f"={p_amt_letter}{sub1_row}*0.001")
                ws_ws.cell(nac_row, p_amt_col).fill = phase_fill
            ws_ws.cell(nac_row, COL_CURR_AMT, f"={CURR_AMT_COL}{sub1_row}*0.001")
            ws_ws.cell(nac_row, COL_MORE, f"=IF({CURR_AMT_COL}{nac_row}>{EST_AMT_COL}{nac_row},{CURR_AMT_COL}{nac_row}-{EST_AMT_COL}{nac_row},\"\")")
            ws_ws.cell(nac_row, COL_LESS, f"=IF({EST_AMT_COL}{nac_row}>{CURR_AMT_COL}{nac_row},{EST_AMT_COL}{nac_row}-{CURR_AMT_COL}{nac_row},\"\")")

            # vi) Sub Total 2
            ws_ws.cell(sub2_row, COL_DESC, "Sub Total 2")
            ws_ws.cell(sub2_row, COL_EST_AMT, f"={EST_AMT_COL}{sub1_row}+{EST_AMT_COL}{lc_row}+{EST_AMT_COL}{qc_row}+{EST_AMT_COL}{nac_row}")
            for p_idx in range(num_previous_phases):
                p_amt_col = COL_PHASE_START + (p_idx * 2) + 1
                p_amt_letter = phase_amt_letter(p_idx)
                ws_ws.cell(sub2_row, p_amt_col, f"={p_amt_letter}{sub1_row}+{p_amt_letter}{lc_row}+{p_amt_letter}{qc_row}+{p_amt_letter}{nac_row}")
                ws_ws.cell(sub2_row, p_amt_col).fill = phase_fill
            ws_ws.cell(sub2_row, COL_CURR_AMT, f"={CURR_AMT_COL}{sub1_row}+{CURR_AMT_COL}{lc_row}+{CURR_AMT_COL}{qc_row}+{CURR_AMT_COL}{nac_row}")
            # (NO More/Less formulas in Sub Total 2 as per requirement)

            # vii) Add GST @ 18%
            ws_ws.cell(gst_row, COL_DESC, "Add GST @ 18%")
            ws_ws.cell(gst_row, COL_EST_AMT, f"={EST_AMT_COL}{sub2_row}*0.18")
            for p_idx in range(num_previous_phases):
                p_amt_col = COL_PHASE_START + (p_idx * 2) + 1
                p_amt_letter = phase_amt_letter(p_idx)
                ws_ws.cell(gst_row, p_amt_col, f"={p_amt_letter}{sub2_row}*0.18")
                ws_ws.cell(gst_row, p_amt_col).fill = phase_fill
            ws_ws.cell(gst_row, COL_CURR_AMT, f"={CURR_AMT_COL}{sub2_row}*0.18")
            ws_ws.cell(gst_row, COL_MORE, f"=IF({CURR_AMT_COL}{gst_row}>{EST_AMT_COL}{gst_row},{CURR_AMT_COL}{gst_row}-{EST_AMT_COL}{gst_row},\"\")")
            ws_ws.cell(gst_row, COL_LESS, f"=IF({EST_AMT_COL}{gst_row}>{CURR_AMT_COL}{gst_row},{EST_AMT_COL}{gst_row}-{CURR_AMT_COL}{gst_row},\"\")")

            # viii) Unused T.P @ % on ECV (Estimate empty, Execution uses Estimate of Sub Total row)
            ws_ws.cell(unused_row, COL_DESC, f"Unused T.P @ {ws_tp_percent} % on ECV")
            ws_ws.cell(unused_row, COL_EST_AMT, None)   # Estimate MUST be empty
            for p_idx in range(num_previous_phases):
                p_amt_col = COL_PHASE_START + (p_idx * 2) + 1
                ws_ws.cell(unused_row, p_amt_col, f"={EST_AMT_COL}{sub_row}*{ws_tp_percent}/100")
                ws_ws.cell(unused_row, p_amt_col).fill = phase_fill
            ws_ws.cell(unused_row, COL_CURR_AMT, f"={EST_AMT_COL}{sub_row}*{ws_tp_percent}/100")
            ws_ws.cell(unused_row, COL_MORE, f"=IF({CURR_AMT_COL}{unused_row}>{EST_AMT_COL}{unused_row},{CURR_AMT_COL}{unused_row}-{EST_AMT_COL}{unused_row},\"\")")
            ws_ws.cell(unused_row, COL_LESS, f"=IF({EST_AMT_COL}{unused_row}>{CURR_AMT_COL}{unused_row},{EST_AMT_COL}{unused_row}-{CURR_AMT_COL}{unused_row},\"\")")

            # ix) L.S. provision row
            ws_ws.cell(ls_row, COL_DESC, "L.S provision towards unforeseen items")
            ws_ws.cell(ls_row, COL_EST_AMT, f"={EST_AMT_COL}{grand_row}-{EST_AMT_COL}{unused_row}-{EST_AMT_COL}{gst_row}-{EST_AMT_COL}{sub2_row}")
            for p_idx in range(num_previous_phases):
                p_amt_col = COL_PHASE_START + (p_idx * 2) + 1
                p_amt_letter = phase_amt_letter(p_idx)
                ws_ws.cell(ls_row, p_amt_col, f"={p_amt_letter}{grand_row}-{p_amt_letter}{unused_row}-{p_amt_letter}{gst_row}-{p_amt_letter}{sub2_row}")
                ws_ws.cell(ls_row, p_amt_col).fill = phase_fill
            ws_ws.cell(ls_row, COL_CURR_AMT, f"={CURR_AMT_COL}{grand_row}-{CURR_AMT_COL}{unused_row}-{CURR_AMT_COL}{gst_row}-{CURR_AMT_COL}{sub2_row}")
            ws_ws.cell(ls_row, COL_MORE, f"=IF({CURR_AMT_COL}{ls_row}>{EST_AMT_COL}{ls_row},{CURR_AMT_COL}{ls_row}-{EST_AMT_COL}{ls_row},\"\")")
            ws_ws.cell(ls_row, COL_LESS, f"=IF({EST_AMT_COL}{ls_row}>{CURR_AMT_COL}{ls_row},{EST_AMT_COL}{ls_row}-{CURR_AMT_COL}{ls_row},\"\")")

            # x) Grand Total = Grand Total of uploaded Estimate (both Estimate & Execution same)
            grand_total_val = float(request.session.get("ws_estimate_grand_total", 0.0) or 0.0)
            ws_ws.cell(grand_row, COL_DESC, "Grand Total")
            ws_ws.cell(grand_row, COL_EST_AMT, grand_total_val)
            for p_idx in range(num_previous_phases):
                p_amt_col = COL_PHASE_START + (p_idx * 2) + 1
                ws_ws.cell(grand_row, p_amt_col, grand_total_val)
                ws_ws.cell(grand_row, p_amt_col).fill = phase_fill
            ws_ws.cell(grand_row, COL_CURR_AMT, grand_total_val)
            # More / Less in Grand Total row = sum of Sub Total â†’ LS rows
            ws_ws.cell(grand_row, COL_MORE, f"=SUM({MORE_COL}{sub_row}:{MORE_COL}{ls_row})")
            ws_ws.cell(grand_row, COL_LESS, f"=SUM({LESS_COL}{sub_row}:{LESS_COL}{ls_row})")

            # style all total rows
            rows_to_style = [tp_row, sub1_row, lc_row, qc_row, nac_row,
                        sub2_row, gst_row, unused_row, ls_row, grand_row]
            if deduct_row:
                rows_to_style.insert(0, deduct_row)  # Add deduct row at the beginning
            for r_i in rows_to_style:
                for col in range(1, total_cols + 1):
                    cell = ws_ws.cell(r_i, col)
                    cell.font = Font(bold=True)
                    cell.border = border_all
                    cell.fill = subtotal_fill
                    if col == COL_DESC:
                        cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
                    else:
                        cell.alignment = Alignment(horizontal="center", vertical="center")

            # reset row heights auto
            for r in range(1, ws_ws.max_row + 1):
                ws_ws.row_dimensions[r].height = None

            # Reorder sheets: WorkSlip first, then ItemBlocks
            if "WorkSlip" in wb_out.sheetnames:
                ws_idx = wb_out.sheetnames.index("WorkSlip")
                if ws_idx > 0:
                    wb_out.move_sheet("WorkSlip", offset=-ws_idx)

            # return file
            resp = HttpResponse(
                content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
            resp["Content-Disposition"] = 'attachment; filename="WorkSlip.xlsx"'
            wb_out.save(resp)
            return resp

    # ------------- 2) Build preview_rows for UI (GET or redirect) -------------
    preview_rows = []
    
    # DEBUG: Log session state on GET
    logger.info(f"[WORKSLIP DEBUG] GET - ws_estimate_rows from session: {len(ws_estimate_rows)} rows")
    if ws_estimate_rows and len(ws_estimate_rows) > 0:
        first_row = ws_estimate_rows[0]
        logger.info(f"[WORKSLIP DEBUG] First row structure: {first_row}")
        logger.info(f"[WORKSLIP DEBUG] First row qty_est={first_row.get('qty_est')}, rate={first_row.get('rate')}")

    # base items
    for idx, row in enumerate(ws_estimate_rows, start=1):
        row_key = row["key"]
        # try same candidate keys as in download
        candidates = [
            f"base:{row_key}",
            row_key,
            row.get("item_name") or "",
            row.get("desc") or "",
        ]
        qty_preview = ""
        for k in candidates:
            k = str(k).strip()
            if k and k in ws_exec_map:
                qty_preview = ws_exec_map[k]
                break

        preview_rows.append({
            "row_type": "base",
            "key": row_key,
            "sl": idx,
            # ðŸ‘‡ UI shows item NAME (yellow header) if present
            "desc": row.get("display_name") or row.get("item_name") or row.get("desc"),
            "qty_est": row.get("qty_est", 0),
            "unit": row.get("unit", ""),
            "rate": row.get("rate", 0),
            "qty_exec": qty_preview,
        })

    # supplemental preview with rates from backend
    item_to_info = {it["name"]: it for it in items_list}
    wb_vals = None
    ws_vals = None
    if filepath and os.path.exists(filepath):
        try:
            wb_vals = load_workbook(filepath, data_only=True)
            ws_vals = wb_vals["Master Datas"]
        except Exception:
            wb_vals = None
            ws_vals = None

    supp_details = []
    if ws_vals is not None:
        for name in ws_supp_items:
            info = item_to_info.get(name)
            if not info:
                continue
            start_row = info["start_row"]
            end_row = info["end_row"]
            rate_val = 0.0
            for r in range(end_row, start_row - 1, -1):
                v = ws_vals.cell(row=r, column=10).value
                if v not in (None, ""):
                    try:
                        rate_val = float(v)
                    except Exception:
                        rate_val = 0.0
                    break
            unit_pl, _ = units_for(name)
            key = f"supp:{name}"
            supp_details.append({
                "name": name,
                "unit": unit_pl,
                "rate": rate_val,
                "qty_exec": ws_exec_map.get(key, ""),
            })

    # Add Previous Workslip Supplemental Items BEFORE current supplemental items
    if ws_previous_supp_items:
        # Group items by supplemental section AND consolidate quantities by workslip phase
        # Structure: { supp_section: { item_name: { phase1_qty, phase2_qty, ... } } }
        supp_by_section = {}
        
        for supp in ws_previous_supp_items:
            section_num = supp.get("supp_section", supp.get("phase", 1))  # Fallback to phase for backward compat
            item_name = supp.get("name", "")
            phase_num = supp.get("phase", 1)
            
            if section_num not in supp_by_section:
                supp_by_section[section_num] = {}
            
            if item_name not in supp_by_section[section_num]:
                supp_by_section[section_num][item_name] = {
                    "name": item_name,
                    "desc": supp.get("desc", item_name),
                    "unit": supp.get("unit", "-") or "-",
                    "rate": supp.get("rate", 0) or 0,
                    "phase_quantities": {},  # { phase_num: qty }
                    "phase_amounts": {},     # { phase_num: amount }
                }
            
            # Add this phase's quantity to the item
            supp_by_section[section_num][item_name]["phase_quantities"][phase_num] = supp.get("qty", 0)
            supp_by_section[section_num][item_name]["phase_amounts"][phase_num] = supp.get("amount", 0)
            # Update rate if not set
            if supp_by_section[section_num][item_name]["rate"] == 0:
                supp_by_section[section_num][item_name]["rate"] = supp.get("rate", 0) or 0
        
        base_count = len(preview_rows)
        preview_rows.append({
            "row_type": "heading",
            "label": "Previous Workslip Supplemental Items",
        })
        
        for section_num in sorted(supp_by_section.keys()):
            section_items = supp_by_section[section_num]
            preview_rows.append({
                "row_type": "heading",
                "label": f"Supplemental Items-{section_num}",
            })
            for i, (item_name, item_data) in enumerate(section_items.items(), start=1):
                supp_key = f"prev_supp:{section_num}:{item_name}"
                
                # Build previous phases execution data for display
                # Must include ALL phases to match the table columns
                num_prev_phases = len(ws_previous_phases) if ws_previous_phases else 0
                prev_phases_exec = []
                for phase_idx in range(1, num_prev_phases + 1):
                    prev_phases_exec.append({
                        "phase": phase_idx,
                        "qty": item_data["phase_quantities"].get(phase_idx, 0),
                        "amount": item_data["phase_amounts"].get(phase_idx, 0),
                    })
                
                prev_supp_row = {
                    "row_type": "prev_supp",
                    "key": supp_key,
                    "name": item_data["name"],
                    "sl": base_count + i,
                    "desc": item_data["name"],
                    "qty_est": "-",
                    "unit": item_data["unit"],
                    "rate": item_data["rate"],
                    "qty_exec": ws_exec_map.get(supp_key, ""),
                    "supp_section": section_num,
                    "previous_phases_exec": prev_phases_exec,  # List of {phase, qty, amount}
                }
                preview_rows.append(prev_supp_row)
                base_count += 1

    # Add Current Phase Supplemental Items
    if supp_details:
        base_count = len(preview_rows)
        preview_rows.append({
            "row_type": "heading",
            "label": f"Supplemental Items{'-' + str(ws_current_phase) if ws_current_phase > 1 else ''}",
        })
        for i, s in enumerate(supp_details, start=1):
            key = f"supp:{s['name']}"
            preview_rows.append({
                "row_type": "supp",
                "key": key,
                "name": s["name"],
                "sl": base_count + i,
                "desc": s["name"],      # UI shows supplemental item name
                "qty_est": 0,
                "unit": s["unit"],
                "rate": s["rate"],
                "qty_exec": ws_exec_map.get(key, ""),
            })

    tp_percent_str = "" if not ws_tp_percent else str(ws_tp_percent)
    ws_deduct_old_material = request.session.get("ws_deduct_old_material", 0.0) or 0.0
    ws_lc_percent = request.session.get("ws_lc_percent", 0.0) or 0.0
    ws_qc_percent = request.session.get("ws_qc_percent", 0.0) or 0.0
    ws_nac_percent = request.session.get("ws_nac_percent", 0.0) or 0.0
    ws_metadata = request.session.get("ws_metadata", {}) or {}
    
    # Add phase data to preview_rows for display (skip prev_supp rows which already have their data)
    for row in preview_rows:
        if row.get("row_type") not in ("heading", "prev_supp"):
            row_key = row.get("key", "")
            # Add previous phases' execution data
            row["previous_phases_exec"] = []
            for phase_idx, phase_map in enumerate(ws_previous_phases):
                phase_exec = phase_map.get(row_key, 0)
                row["previous_phases_exec"].append({
                    "phase": phase_idx + 1,
                    "qty": phase_exec,
                    "amount": phase_exec * row.get("rate", 0) if phase_exec else 0
                })

    return render(request, "core/workslip.html", {
        "category": category,
        "work_type": ws_work_type,
        "work_type_display": {
            'new_estimate': 'New Estimate',
            'amc': 'AMC',
            'tempworks': 'Temporary Works'
        }.get(ws_work_type, 'New Estimate'),
        "work_mode": ws_work_mode,
        "work_mode_display": {
            'original': 'Original',
            'repair': 'Repair'
        }.get(ws_work_mode, 'Original'),
        "category_display": ws_category.title() if ws_category else 'Electrical',
        "module_code": module_code,
        "groups": groups,
        "current_group": current_group,
        "items_in_group": items_in_group,
        "ws_estimate_rows": ws_estimate_rows,
        "preview_rows": preview_rows,
        "tp_percent": tp_percent_str,
        "tp_type": ws_tp_type,
        "supp_items_selected": ws_supp_items,
        "work_name": ws_work_name,
        "deduct_old_material": ws_deduct_old_material if ws_deduct_old_material > 0 else "",
        "lc_percent": ws_lc_percent if ws_lc_percent > 0 else "",
        "qc_percent": ws_qc_percent if ws_qc_percent > 0 else "",
        "nac_percent": ws_nac_percent if ws_nac_percent > 0 else "",
        "current_phase": ws_current_phase,
        "target_workslip": request.session.get("ws_target_workslip", 1) or 1,
        "previous_phases": ws_previous_phases,
        "total_phases": len(ws_previous_phases) if ws_previous_phases else 0,
        "previous_supp_items": ws_previous_supp_items,
        "ws_metadata": ws_metadata,
        "available_backends": available_backends,
        "selected_backend_id": ws_selected_backend_id,
    })


# -----------------------
# WORKSLIP AJAX TOGGLE SUPPLEMENTAL ITEM
# -----------------------
@login_required(login_url='login')
def workslip_ajax_toggle_supp(request):
    """
    AJAX endpoint to toggle a supplemental item in workslip without page reload.
    POST with JSON: { "item": "item_name", "action": "add" or "remove" }
    Returns JSON: { "status": "ok", "supp_items": [...], "action_taken": "added" or "removed", "item_info": {...} }
    """
    if request.method != "POST":
        return JsonResponse({"status": "error", "message": "POST required"}, status=405)
    
    try:
        if request.content_type and 'application/json' in request.content_type:
            data = json.loads(request.body.decode("utf-8") or "{}")
        else:
            data = request.POST
        
        item = data.get("item", "").strip()
        action = data.get("action", "toggle")  # "add", "remove", or "toggle"
        
        if not item:
            return JsonResponse({"status": "error", "message": "No item specified"}, status=400)
        
        ws_supp_items = request.session.get("ws_supp_items", []) or []
        action_taken = None
        
        if action == "add":
            if item not in ws_supp_items:
                ws_supp_items.append(item)
                action_taken = "added"
            else:
                action_taken = "already_exists"
        elif action == "remove":
            if item in ws_supp_items:
                ws_supp_items.remove(item)
                action_taken = "removed"
                # Also remove from exec_map
                ws_exec_map = request.session.get("ws_exec_map", {}) or {}
                key = f"supp:{item}"
                if key in ws_exec_map:
                    del ws_exec_map[key]
                request.session["ws_exec_map"] = ws_exec_map
            else:
                action_taken = "not_found"
        else:  # toggle
            if item in ws_supp_items:
                ws_supp_items.remove(item)
                action_taken = "removed"
                # Also remove from exec_map
                ws_exec_map = request.session.get("ws_exec_map", {}) or {}
                key = f"supp:{item}"
                if key in ws_exec_map:
                    del ws_exec_map[key]
                request.session["ws_exec_map"] = ws_exec_map
            else:
                ws_supp_items.append(item)
                action_taken = "added"
        
        request.session["ws_supp_items"] = ws_supp_items
        
        # Get item info (rate, unit) for newly added items
        item_info = None
        if action_taken == "added":
            try:
                # Load backend for the default category
                category = "original"
                items_list, groups_map, units_map, ws_data, filepath = load_backend(category, settings.BASE_DIR)
                
                # Get rate
                wb_vals = load_workbook(filepath, data_only=True)
                ws_vals = wb_vals["Master Datas"]
                
                item_rate = None
                for info in items_list:
                    if info["name"] == item:
                        start_row = info["start_row"]
                        end_row = info["end_row"]
                        for r in range(end_row, start_row - 1, -1):
                            val = ws_vals.cell(row=r, column=10).value  # column J
                            if val not in (None, ""):
                                item_rate = val
                                break
                        break
                
                # Get unit from backend units_map (Column D of Groups sheet)
                unit = units_map.get(item, "Nos")
                
                item_info = {
                    "name": item,
                    "rate": item_rate,
                    "unit": unit
                }
                
                wb_vals.close()
            except Exception as e:
                # If we can't get item info, just return without it
                item_info = {"name": item, "rate": None, "unit": "Nos"}
        
        return JsonResponse({
            "status": "ok",
            "supp_items": ws_supp_items,
            "action_taken": action_taken,
            "item": item,
            "item_info": item_info
        })
        
    except Exception as e:
        return JsonResponse({"status": "error", "message": str(e)}, status=400)


# Provide placeholder views for any names expected by URLConf but not yet
# implemented during the ongoing refactor. This uses module-level
# __getattr__ (PEP 562) so imports like `from core import views` succeed.
_placeholder_views = {}

def _make_placeholder(name):
    def view(request, *args, **kwargs):
        return HttpResponse(f"Placeholder view: {name}")
    view.__name__ = name
    return view

def __getattr__(name: str):
    # Common view names referenced by estimate_site/urls.py â€” if any are
    # missing, return a lightweight placeholder so the server can start.
    if name.startswith("_"):
        raise AttributeError(name)
    if name not in _placeholder_views:
        _placeholder_views[name] = _make_placeholder(name)
    return _placeholder_views[name]


def _safe_float(v):
    try:
        return float(v)
    except Exception:
        return 0.0


def _extract_header_data_fuzzy_from_wb(wb):
    """
    Scan all sheets, top ~40 rows, trying to find:
      - Name of work
      - Estimate amount / ECV (optional, not strictly needed here)
      - Admin sanction (ignored for docs)
      - Technical sanction (ignored for docs)
      - Agreement
      - Agency

    We are NOT strict about exact label text: things like
      'Name of work', 'Name of the work', 'Work name', etc
      'Agreement', 'Agt', 'Agrmt', ...
      'Agency', 'Contractor', ...
    are all accepted.

    We return values *after* the colon if present.

    Returns dict:
      {
        "name_of_work": "...",
        "agreement": "...",
        "agency": "...",
      }
    (others may be empty)
    """
    header = {
        "name_of_work": "",
        "estimate_amount": "",
        "admin_sanction": "",
        "tech_sanction": "",
        "agreement": "",
        "agency": "",
    }

    def clean_value_for_header(val):
        if not val:
            return ""
        s = str(val).strip()
        # if there's a colon, return RHS only
        if ":" in s:
            s = s.split(":", 1)[1].strip()
        return s

    for ws in wb.worksheets:
        max_row = min(ws.max_row, 40)
        max_col = min(ws.max_column, 20)
        for r in range(1, max_row + 1):
            for c in range(1, max_col + 1):
                raw = ws.cell(row=r, column=c).value
                if raw is None:
                    continue
                s_full = str(raw).strip()
                low = s_full.lower()
                # skip obviously numeric things
                if not low:
                    continue

                # tokens set
                tokens = set(low.replace(":", " ").replace(".", " ").split())

                # ---- Name of work ----
                if not header["name_of_work"]:
                    # look for 'name' + 'work' anywhere
                    if "name" in tokens and "work" in tokens:
                        header["name_of_work"] = clean_value_for_header(s_full)
                        continue

                # ---- Estimate Amount / ECV (optional) ----
                if not header["estimate_amount"]:
                    if "estimate" in tokens and ("amount" in tokens or "ecv" in tokens):
                        header["estimate_amount"] = clean_value_for_header(s_full)
                        continue

                # ---- Admin sanction (optional, not used in docs) ----
                if not header["admin_sanction"]:
                    if "admin" in low and "sanction" in low:
                        header["admin_sanction"] = clean_value_for_header(s_full)
                        continue

                # ---- Technical sanction (optional, not used in docs) ----
                if not header["tech_sanction"]:
                    if ("tech" in low or "technical" in low) and "sanction" in low:
                        header["tech_sanction"] = clean_value_for_header(s_full)
                        continue

                # ---- Agreement (Agt / Agrmt etc.) ----
                if not header["agreement"]:
                    if (
                        "agreement" in low
                        or "agrmt" in low
                        or "agt" in tokens
                        or "agt." in tokens
                    ):
                        header["agreement"] = clean_value_for_header(s_full)
                        continue

                # ---- Agency (Contractor / Firm etc.) ----
                if not header["agency"]:
                    if "agency" in tokens or "contractor" in tokens or "firm" in tokens:
                        header["agency"] = clean_value_for_header(s_full)
                        continue

    return header


def _extract_header_data_from_sheet(ws):
    """
    Extract header data (name of work, agreement, agency, etc.) from a specific sheet.
    
    This is a per-sheet version of _extract_header_data_fuzzy_from_wb that reads
    data from the first ~40 rows of the given worksheet.
    
    Returns dict:
      {
        "name_of_work": "...",
        "estimate_amount": "...",
        "admin_sanction": "...",
        "tech_sanction": "...",
        "agreement": "...",
        "agency": "...",
      }
    """
    header = {
        "name_of_work": "",
        "estimate_amount": "",
        "admin_sanction": "",
        "tech_sanction": "",
        "agreement": "",
        "agency": "",
    }

    def clean_value_for_header(val):
        if not val:
            return ""
        s = str(val).strip()
        # if there's a colon, return RHS only
        if ":" in s:
            s = s.split(":", 1)[1].strip()
        return s

    max_row = min(ws.max_row, 40)
    max_col = min(ws.max_column, 20)
    for r in range(1, max_row + 1):
        for c in range(1, max_col + 1):
            raw = ws.cell(row=r, column=c).value
            if raw is None:
                continue
            s_full = str(raw).strip()
            low = s_full.lower()
            # skip obviously numeric things
            if not low:
                continue

            # tokens set
            tokens = set(low.replace(":", " ").replace(".", " ").split())

            # ---- Name of work ----
            if not header["name_of_work"]:
                # look for 'name' + 'work' anywhere
                if "name" in tokens and "work" in tokens:
                    header["name_of_work"] = clean_value_for_header(s_full)
                    continue

            # ---- Estimate Amount / ECV (optional) ----
            if not header["estimate_amount"]:
                if "estimate" in tokens and ("amount" in tokens or "ecv" in tokens):
                    header["estimate_amount"] = clean_value_for_header(s_full)
                    continue

            # ---- Admin sanction (optional, not used in docs) ----
            if not header["admin_sanction"]:
                if "admin" in low and "sanction" in low:
                    header["admin_sanction"] = clean_value_for_header(s_full)
                    continue

            # ---- Technical sanction (optional, not used in docs) ----
            if not header["tech_sanction"]:
                if ("tech" in low or "technical" in low) and "sanction" in low:
                    header["tech_sanction"] = clean_value_for_header(s_full)
                    continue

            # ---- Agreement (Agt / Agrmt etc.) ----
            if not header["agreement"]:
                if (
                    "agreement" in low
                    or "agrmt" in low
                    or "agt" in tokens
                    or "agt." in tokens
                ):
                    header["agreement"] = clean_value_for_header(s_full)
                    continue

            # ---- Agency (Contractor / Firm etc.) ----
            if not header["agency"]:
                if "agency" in tokens or "contractor" in tokens or "firm" in tokens:
                    header["agency"] = clean_value_for_header(s_full)
                    continue

    return header


def _detect_bill_format(ws):
    """
    Try to detect if this sheet looks like:
      - 'First Bill' 8-column format (S.No, Quantity, Unit, Item, Rate, Per, Unit, Amount)
      - 'Nth Bill' 10-column format with 'Quantity Till Date' headers

    Returns one of:
      "first_bill_8col", "nth_bill_10col", or None
    """
    max_scan_row = min(ws.max_row, 20)
    for r in range(1, max_scan_row + 1):
        a = str(ws.cell(row=r, column=1).value or "").strip().lower()
        b = str(ws.cell(row=r, column=2).value or "").strip().lower()
        c = str(ws.cell(row=r, column=3).value or "").strip().lower()
        d = str(ws.cell(row=r, column=4).value or "").strip().lower()
        # fifth column value not needed for detection here
        
        # First Bill style: "sl" in A, "quantity" in B, "item"/"description" in D
        if "sl" in a and "quantity" in b and ("item" in d or "description" in d):
            # assume 8-col bill
            return "first_bill_8col"

        # Nth Bill style: "sl" in A and "quantity till date" in C
        if "sl" in a and "quantity till date" in c:
            return "nth_bill_10col"

    return None


def _extract_total_amount_from_bill_wb(wb):
    """
    Try to find the current bill amount (the 'Total' row) in a Bill workbook.

    - For 'First Bill' 8-col format -> read from column H in 'Total' row.
    - For 'Nth Bill' 10-col format -> read from column I (Since Last Amount) in 'Total' row.
    - Fallback: last numeric cell in the row containing 'Total'.

    Returns float.
    """
    if not wb.worksheets:
        return 0.0

    ws = wb.worksheets[0]  # assume bill is in first sheet
    return _extract_total_amount_from_single_sheet(ws)


def _extract_total_amount_from_single_sheet(ws):
    """
    Extract total amount from a single bill sheet (worksheet).
    
    - For 'First Bill' 8-col format -> read from column H in 'Total' row.
    - For 'Nth Bill' 10-col format -> read from column I (Since Last Amount) in 'Total' row.
    - Fallback: last numeric cell in the row containing 'Total'.

    Returns float.
    """
    if not ws:
        return 0.0
    
    bill_format = _detect_bill_format(ws)

    max_row = ws.max_row
    total_row = None

    for r in range(1, max_row + 1):
        found_total_label = False
        for c in range(1, ws.max_column + 1):
            val = ws.cell(row=r, column=c).value
            if isinstance(val, str):
                low = val.strip().lower()
                if low.startswith("total"):  # ignore 'sub total' rows by checking exact 'total'
                    # but allow 'total' and 'total amount' etc.
                    if not low.startswith("sub total") and not low.startswith("subtotal"):
                        found_total_label = True
                        break
        if found_total_label:
            total_row = r
            break

    if total_row is None:
        return 0.0

    # Format-specific read
    if bill_format == "first_bill_8col":
        # H = col 8
        amount_val = ws.cell(row=total_row, column=8).value
        if amount_val is not None and str(amount_val).strip() != "":
            return _safe_float(amount_val)

    if bill_format == "nth_bill_10col":
        # 'Since Last Amount' total is in column I (9)
        amount_val = ws.cell(row=total_row, column=9).value
        if amount_val is not None and str(amount_val).strip() != "":
            return _safe_float(amount_val)

    # Fallback: last numeric cell in the Total row
    last_numeric = 0.0
    for c in range(1, ws.max_column + 1):
        val = ws.cell(row=total_row, column=c).value
        if isinstance(val, (int, float)):
            last_numeric = float(val)
    return last_numeric


# Duplicate helper definitions removed here; canonical `_build_mb_details_string`
# and `_resolve_cc_header` are defined earlier in the file and will be used.




# -----------------------
# BILL FROM WORKSLIP
# -----------------------
# -----------------------
# BILL  -  from Estimate / from WorkSlip
# -----------------------
# imports consolidated at top of file


# ---------- Module-level helpers for `bill()` (extracted from nested defs) ----------
def to_number(v):
    try:
        return float(v)
    except Exception:
        return 0.0


def is_merged_cell(ws, row, col):
    for merged in ws.merged_cells.ranges:
        if merged.min_row <= row <= merged.max_row and merged.min_col <= col <= merged.max_col:
            return True
    return False


def find_estimate_sheet_and_header_row(wb):
    for ws in wb.worksheets:
        for r in range(1, 26):
            a = str(ws.cell(row=r, column=1).value or "").strip().lower()
            b = str(ws.cell(row=r, column=2).value or "").strip().lower()
            d = str(ws.cell(row=r, column=4).value or "").strip().lower()
            if "sl" in a and "quantity" in b and ("item" in d or "description" in d):
                return ws, r
    return wb.worksheets[0], 3


def looks_like_header(ws, r):
    a = str(ws.cell(row=r, column=1).value or "").strip().lower()
    b = str(ws.cell(row=r, column=2).value or "").strip().lower()
    c = str(ws.cell(row=r, column=3).value or "").strip().lower()
    d = str(ws.cell(row=r, column=4).value or "").strip().lower()
    e = str(ws.cell(row=r, column=5).value or "").strip().lower()

    has_sl = ("sl" in a) or ("s.no" in a) or ("serial" in a)
    has_qty = ("qty" in b or "quantity" in b or "qty" in c or "quantity" in c)
    has_desc = (
        "item" in c or "item" in d or "item" in e or
        "description" in c or "description" in d or "description" in e
    )
    return has_sl and has_qty and has_desc


def find_all_estimate_sheets_and_header_rows(wb):
    results = []

    for ws in wb.worksheets:
        header_row = None
        max_scan = min(ws.max_row, 60)
        for r in range(1, max_scan + 1):
            if looks_like_header(ws, r):
                header_row = r
                break
        if header_row:
            results.append((ws, header_row))

    if not results:
        for ws in wb.worksheets:
            hr = None
            max_scan = min(ws.max_row or 0, 30)
            for r in range(1, max_scan + 1):
                if looks_like_header(ws, r):
                    hr = r
                    break
            if not hr:
                hr = 3
            results.append((ws, hr))

    return results


def parse_estimate_items(ws, header_row):
    items = []
    max_row = min(ws.max_row, 5000)

    for r in range(header_row + 1, max_row + 1):
        desc_raw = ws.cell(row=r, column=4).value
        desc = str(desc_raw or "").strip()
        rate_raw = ws.cell(row=r, column=5).value
        amt_raw = ws.cell(row=r, column=8).value

        desc_low = desc.lower()

        if desc_low.startswith(("ecv", "sub total", "subtotal", "total")):
            break

        rate_str = "" if rate_raw is None else str(rate_raw).strip()
        is_rate_empty = (rate_str == "")

        if desc and (is_rate_empty or is_merged_cell(ws, r, 4)):
            continue

        qty_raw = ws.cell(row=r, column=2).value
        unit_raw = ws.cell(row=r, column=3).value

        qty_str = "" if qty_raw is None else str(qty_raw).strip()
        unit_str = "" if unit_raw is None else str(unit_raw).strip()
        amt_str = "" if amt_raw is None else str(amt_raw).strip()

        all_blank = (
            desc == "" and
            is_rate_empty and
            qty_str == "" and
            unit_str == "" and
            amt_str == ""
        )

        if all_blank:
            continue

        if desc == "" and is_rate_empty:
            continue

        qty = to_number(qty_raw)
        unit = unit_str
        rate = to_number(rate_raw)
        amt = to_number(amt_raw)

        if amt != 0:
            if qty == 0 and rate != 0:
                qty = amt / rate
            elif rate == 0 and qty != 0:
                rate = amt / qty

        items.append({
            "qty": qty,
            "unit": unit,
            "desc": desc,
            "rate": rate,
        })

    return items


def find_workslip_sheet(wb):
    for ws in wb.worksheets:
        for r in range(1, 40):
            b = str(ws.cell(row=r, column=2).value or "").strip().lower()
            g = str(ws.cell(row=r, column=7).value or "").strip().lower()
            if "description" in b and ("qty" in g or "quantity" in g):
                return ws
    return wb.worksheets[0]


def find_all_workslip_sheets(wb):
    results = []
    for ws in wb.worksheets:
        max_scan = min(ws.max_row, 60)
        for r in range(1, max_scan + 1):
            b = str(ws.cell(row=r, column=2).value or "").strip().lower()
            c = str(ws.cell(row=r, column=3).value or "").strip().lower()

            has_desc = (
                "description of item" in b or
                "description of item" in c or
                ("description" in b and "item" in b) or
                ("description" in c and "item" in c)
            )

            has_qty = False
            for col in range(5, 12):
                t = str(ws.cell(row=r, column=col).value or "").strip().lower()
                if "qty" in t or "quantity" in t:
                    has_qty = True
                    break

            if has_desc and has_qty:
                results.append(ws)
                break

    if not results:
        results.append(find_workslip_sheet(wb))

    return results


def extract_tp_from_workslip(ws):
    """Extract T.P percentage and type (Less/Excess) from a workslip sheet."""
    max_scan = min(ws.max_row, 400)
    last_sub_total = None
    import re
    
    def has_tp_keyword(text):
        """Check if text contains T.P or tender premium keyword."""
        return "tp" in text or "tender premium" in text or "tenderpremium" in text

    for r in range(1, max_scan + 1):
        # Build row text for easy matching
        row_text_parts = []
        for col in range(1, 15):
            row_text_parts.append(str(ws.cell(row=r, column=col).value or ""))
        row_text = " ".join(row_text_parts).strip().lower()
        # Normalize "t.p" "t.p." "t. p" etc. to "tp" for matching
        row_text_norm = re.sub(r't\s*\.\s*p\.?', 'tp', row_text)

        # Track the latest Sub Total value (amount columns to the right)
        if "sub total" in row_text and "total" in row_text:
            # This may match both, but keep as potential
            pass
        if "sub total" in row_text and not row_text.startswith("total"):
            for amt_col in [8, 9, 10, 11, 12]:
                sub_val = to_number(ws.cell(row=r, column=amt_col).value)
                if sub_val is not None:
                    last_sub_total = sub_val
                    break

        # Direct percent extraction if present (check both original and normalized)
        if (has_tp_keyword(row_text) or has_tp_keyword(row_text_norm)) and ("%" in row_text or "percent" in row_text):
            percent_match = re.search(r"(\d+\.?\d*)\s*%", row_text)
            if percent_match:
                tp_percent = float(percent_match.group(1))
                # Check for Less/Deduct keywords vs Excess/Add
                is_less = "less" in row_text or "deduct" in row_text
                tp_type = "Less" if is_less else "Excess"
                print(f"DEBUG extract_tp_from_workslip: Found TP row with percent. row_text='{row_text[:100]}' percent={tp_percent} type={tp_type}")
                return tp_percent, tp_type

        # If no percent text, but TP row has an amount, derive percent from last_sub_total
        if has_tp_keyword(row_text) or has_tp_keyword(row_text_norm):
            for amt_col in [8, 9, 10, 11, 12]:
                tp_val = to_number(ws.cell(row=r, column=amt_col).value)
                if tp_val is None:
                    continue
                if last_sub_total and last_sub_total != 0:
                    tp_percent = abs(tp_val) * 100.0 / abs(last_sub_total)
                    # Check for Less/Deduct keywords vs Excess/Add; also use sign of amount
                    is_less = "less" in row_text or "deduct" in row_text or tp_val < 0
                    tp_type = "Less" if is_less else "Excess"
                    print(f"DEBUG extract_tp_from_workslip: Found TP row with amount. row_text='{row_text[:100]}' tp_val={tp_val} derived percent={tp_percent} type={tp_type}")
                    return tp_percent, tp_type

    # Default if not found
    print(f"DEBUG extract_tp_from_workslip: No TP found, using default 0.0 Excess")
    return 0.0, "Excess"


def parse_workslip_items(ws):
    """
    Parse workslip items, automatically detecting the last workslip phase columns.
    
    For multi-phase workslips (Workslip-1, Workslip-2, etc.), this function will
    find and use the LAST workslip's Qty/Amt columns for bill generation.
    
    The Rate is ALWAYS taken from the Estimate Rate column (column 5), as
    multi-phase workslips don't have separate Rate columns per phase.
    """
    items = []
    max_row = min(ws.max_row, 5000)
    
    # Step 1: Find header row and detect column structure
    header_row = 8  # Default
    for r in range(1, 15):
        cell_val = str(ws.cell(row=r, column=1).value or "").strip().lower()
        if "sl" in cell_val:
            header_row = r
            break
    
    # Step 2: Scan header row to find all workslip/execution qty columns
    # We need to find the LAST Qty column - this is the latest workslip phase
    qty_columns = []  # List of column indices for execution qty columns
    amt_columns = []  # List of column indices for execution amt columns
    estimate_rate_col = 5  # Default - Rate (Estimate) column
    
    for c in range(1, 40):  # Scan more columns for multi-phase workslips
        header = str(ws.cell(row=header_row, column=c).value or "").strip().lower()
        
        # Check for workslip/execution columns
        is_exec_col = ("execution" in header or "exec" in header or "workslip" in header)
        is_qty = ("qty" in header or "quantity" in header)
        is_amt = ("amount" in header or "amt" in header)
        is_estimate = ("est" in header or "estimate" in header)
        is_rate = ("rate" in header)
        
        # Find estimate rate column
        if is_estimate and is_rate:
            estimate_rate_col = c
        
        # Skip estimate columns for qty/amt detection
        if is_estimate:
            continue
        
        # Skip More/Less/Remarks columns
        if "more" in header or "less" in header or "remark" in header:
            continue
            
        if is_exec_col and is_qty:
            qty_columns.append(c)
        elif is_exec_col and is_amt:
            amt_columns.append(c)
    
    # Step 3: Determine the columns to use
    # Use the LAST Qty column (latest phase) and its corresponding Amount column
    # Rate is always from Estimate Rate column
    if qty_columns:
        qty_col = qty_columns[-1]  # Last qty column = latest workslip phase
        # The Amount column should be immediately after Qty in multi-phase format
        amt_col = amt_columns[-1] if amt_columns else qty_col + 1
        rate_col = estimate_rate_col  # Always use Estimate Rate
    else:
        # Fallback to default columns (old single-phase format: col 7=Qty, 8=Rate, 9=Amt)
        qty_col = 7
        rate_col = 8  # Old format had Rate at col 8
        amt_col = 9
    
    print(f"DEBUG parse_workslip_items: Using columns - qty={qty_col}, rate={rate_col}, amt={amt_col} (detected from header row {header_row})")

    for r in range(1, max_row + 1):
        desc_raw = ws.cell(row=r, column=2).value
        desc = str(desc_raw or "").strip()
        if desc == "":
            continue

        low = desc.lower()

        if low == "description of item":
            continue
        if low.startswith("sub total") or low.startswith("sub-total") or low.startswith("ecv"):
            break
        if "supplemental items" in low:
            continue

        if is_merged_cell(ws, r, 2):
            continue

        is_ae = low.startswith("ae")

        qty_raw = ws.cell(row=r, column=qty_col).value
        rate_raw = ws.cell(row=r, column=rate_col).value
        amt_raw = ws.cell(row=r, column=amt_col).value

        qty_exec = to_number(qty_raw)
        rate_exec = to_number(rate_raw)
        amt_exec = to_number(amt_raw)

        if amt_exec != 0:
            if qty_exec == 0 and rate_exec != 0:
                qty_exec = amt_exec / rate_exec
            elif rate_exec == 0 and qty_exec != 0:
                rate_exec = amt_exec / qty_exec

        if qty_exec == 0 and rate_exec == 0:
            continue

        unit = str(ws.cell(row=r, column=3).value or "").strip()

        items.append({
            "qty": qty_exec,
            "unit": unit,
            "desc": desc,
            "rate": rate_exec,
            "is_ae": is_ae,
        })

    return items


def find_nth_bill_sheet_and_header_row(wb):
    for ws in wb.worksheets:
        for r in range(1, 40):
            a = str(ws.cell(row=r, column=1).value or "").strip().lower()
            c = str(ws.cell(row=r, column=3).value or "").strip().lower()
            if "sl" in a and "quantity till date" in c:
                return ws, r
    return wb.worksheets[0], 10


def parse_first_bill_for_nth(ws, header_row):
    """Parse items from a First Bill (8-column format).
    
    header_row is the row with column headers (S.No, Quantity, Unit, Item, Rate, Per, Unit, Amount).
    We skip that header row and start parsing from header_row + 1.
    Only treat rows as items if they have a valid amount in column H.
    """
    items = []
    max_row = min(ws.max_row, 5000)

    for r in range(header_row + 1, max_row + 1):
        desc_raw = ws.cell(row=r, column=4).value
        desc = str(desc_raw or "").strip()
        rate_raw = ws.cell(row=r, column=5).value
        amt_cell_raw = ws.cell(row=r, column=8).value

        desc_low = desc.lower()

        if desc_low.startswith(("ecv", "sub total", "subtotal", "total")):
            break

        # Require amount to be non-empty and non-zero to be considered a valid item
        amt_val = to_number(amt_cell_raw)
        if amt_val == 0 or amt_val is None:
            # No amount = not a data item
            continue

        rate_str = "" if rate_raw is None else str(rate_raw).strip()
        is_rate_empty = (rate_str == "")

        if (
            desc
            and is_rate_empty
            and not desc_low.startswith(("ecv", "sub total", "subtotal", "total"))
        ) or (
            desc
            and is_merged_cell(ws, r, 4)
        ):
            continue

        prev_qty_raw = ws.cell(row=r, column=2).value
        unit_raw = ws.cell(row=r, column=3).value

        qty_str = "" if prev_qty_raw is None else str(prev_qty_raw).strip()
        unit_str = "" if unit_raw is None else str(unit_raw).strip()

        if desc == "" and is_rate_empty:
            continue

        unit = unit_str
        qty_val = to_number(prev_qty_raw)
        rate_val = to_number(rate_raw)
        prev_amount_val = amt_val

        if (amt_cell_raw is None or str(amt_cell_raw).strip() == "") and (qty_val != 0 or rate_val != 0):
            prev_amount_val = qty_val * rate_val

        items.append({
            "desc": desc,
            "unit": unit,
            "rate": rate_val,
            "prev_qty": qty_val,
            "prev_amount": prev_amount_val,
        })

    return items


def parse_nth_bill_for_next(ws, header_row):
    """Parse items from an Nth bill (11-column format with Unit column).
    
    Format: A: S.No, B: Item, C: Quantity, D: Unit, E: Rate, F: Total Value,
            G-H: Deduct Previous (Qty, Amt), I-J: Since Last (Qty, Amt), K: Remarks
    
    We look at the header row to determine if there's a Unit column.
    """
    items = []
    max_row = min(ws.max_row, 5000)
    start_row = header_row + 2
    
    # Check if this is an 11-column format with Unit column
    # by examining the header row for 'Unit' in column D or similar patterns
    has_unit_column = False
    for col in range(3, 6):  # Check columns C, D, E
        hdr_val = str(ws.cell(row=header_row, column=col).value or "").strip().lower()
        if hdr_val == "unit":
            has_unit_column = True
            break
    
    for r in range(start_row, max_row + 1):
        desc_raw = ws.cell(row=r, column=2).value
        desc = str(desc_raw or "").strip()
        if desc == "":
            continue

        low = desc.lower()
        if low.startswith("sub total") or low.startswith("subtotal"):
            break

        if has_unit_column:
            # 11-column format: A: S.No, B: Item, C: Qty, D: Unit, E: Rate, F: Total Value
            unit_raw = ws.cell(row=r, column=4).value
            rate_raw = ws.cell(row=r, column=5).value
            prev_qty_raw = ws.cell(row=r, column=3).value
            amt_cell_raw = ws.cell(row=r, column=6).value
        else:
            # 10-column format: A: S.No, B: Item, C: Qty, D: Rate, E: Total Value
            unit_raw = None
            rate_raw = ws.cell(row=r, column=4).value
            prev_qty_raw = ws.cell(row=r, column=3).value
            amt_cell_raw = ws.cell(row=r, column=5).value

        unit_str = str(unit_raw or "").strip() if unit_raw else ""
        qty_val = to_number(prev_qty_raw)
        rate_val = to_number(rate_raw)
        prev_amount_val = to_number(amt_cell_raw)

        if (amt_cell_raw is None or str(amt_cell_raw).strip() == "") and (qty_val != 0 or rate_val != 0):
            prev_amount_val = qty_val * rate_val

        items.append({
            "desc": desc,
            "unit": unit_str,
            "rate": rate_val,
            "prev_qty": qty_val,
            "prev_amount": prev_amount_val,
        })

    return items


def read_tp_from_sheet(ws):
    tp_percent = None
    tp_type = None
    max_row = ws.max_row
    for r in range(1, max_row + 1):
        for col in (2, 4):
            val = ws.cell(row=r, column=col).value
            if not val:
                continue
            s = str(val).strip()
            low = s.lower()
            if "t.p" in low:
                if low.startswith("add"):
                    tp_type = "Excess"
                elif low.startswith("deduct"):
                    tp_type = "Less"
                m = re.search(r"(\d+(\.\d+)?)", low)
                if m:
                    tp_percent = float(m.group(1))
                return tp_percent, tp_type
    return None, None


def singular_unit(plural):
    p = str(plural or "").strip()
    if p.lower().endswith("s") and len(p) > 1:
        return p[:-1]
    return p


def ordinal_word(n):
    mapping = {
        1: "First",
        2: "Second",
        3: "Third",
        4: "Fourth",
        5: "Fifth",
        6: "Sixth",
        7: "Seventh",
        8: "Eighth",
        9: "Ninth",
        10: "Tenth",
    }
    if n in mapping:
        return mapping[n]
    if 10 < n % 100 < 14:
        suffix = "th"
    else:
        suffix = {1: "st", 2: "nd", 3: "rd"}.get(n % 10, "th")
        if suffix is None:
            suffix = "th"
    return f"{n}{suffix}"


def create_first_bill_sheet(
    wb_out,
    sheet_name,
    items,
    header_data,
    title_text,
    tp_percent,
    tp_type,
    mb_measure_no, mb_measure_p_from, mb_measure_p_to,
    mb_abs_no, mb_abs_p_from, mb_abs_p_to,
    doi, doc, domr, dobr
):
    if wb_out.worksheets and wb_out.worksheets[0].cell(row=1, column=1).value is None and len(wb_out.worksheets) == 1:
        ws_bill = wb_out.active
        ws_bill.title = sheet_name
    else:
        ws_bill = wb_out.create_sheet(title=sheet_name)

    thin = Side(border_style="thin", color="000000")
    border_all = Border(left=thin, right=thin, top=thin, bottom=thin)
    header_fill = PatternFill("solid", fgColor="FFC8C8C8")
    subtotal_fill = PatternFill("solid", fgColor="FFE6E6E6")

    ws_bill.merge_cells("A1:H1")
    c = ws_bill["A1"]
    c.value = title_text
    c.font = Font(bold=True, size=14)
    c.alignment = Alignment(horizontal="center", vertical="center")

    work_val = header_data.get("name_of_work", "").strip()
    est_val = header_data.get("estimate_amount", "").strip()
    adm_val = header_data.get("admin_sanction", "").strip()
    tech_val = header_data.get("tech_sanction", "").strip()
    agt_val = header_data.get("agreement", "").strip()
    agency_val = header_data.get("agency", "").strip()

    ws_bill.merge_cells("A2:H2")
    c2 = ws_bill["A2"]
    c2.value = f"Name of the work : {work_val}" if work_val else "Name of the work :"
    c2.font = Font(bold=True)
    c2.alignment = Alignment(horizontal="left", vertical="center")

    ws_bill.merge_cells("A3:H3")
    c3 = ws_bill["A3"]
    c3.value = f"Estimate Amount : {est_val}" if est_val else "Estimate Amount :"
    c3.font = Font(bold=True)
    c3.alignment = Alignment(horizontal="left", vertical="center")

    ws_bill.merge_cells("A4:H4")
    c4 = ws_bill["A4"]
    c4.value = (
        f"Ref. to Administrative sanction : {adm_val}"
        if adm_val else "Ref. to Administrative sanction :"
    )
    c4.font = Font(bold=True)
    c4.alignment = Alignment(horizontal="left", vertical="center")

    ws_bill.merge_cells("A5:H5")
    c5 = ws_bill["A5"]
    c5.value = (
        f"Ref. to Technical sanction : {tech_val}"
        if tech_val else "Ref. to Technical sanction :"
    )
    c5.font = Font(bold=True)
    c5.alignment = Alignment(horizontal="left", vertical="center")

    ws_bill.merge_cells("A6:H6")
    c6 = ws_bill["A6"]
    c6.value = (
        f"Ref. to Agreement : {agt_val}"
        if agt_val else "Ref. to Agreement :"
    )
    c6.font = Font(bold=True)
    c6.alignment = Alignment(horizontal="left", vertical="center")

    ws_bill.merge_cells("A7:H7")
    c7 = ws_bill["A7"]
    c7.value = (
        f"Name of the Agency : {agency_val}"
        if agency_val else "Name of the Agency :"
    )
    c7.font = Font(bold=True)
    c7.alignment = Alignment(horizontal="left", vertical="center")

    ws_bill.merge_cells("A8:H8")
    c8 = ws_bill["A8"]
    c8.value = (
        f"M.B.No Details: MB.No. {mb_measure_no} P.No. {mb_measure_p_from} to {mb_measure_p_to} (Measurements)   "
        f"&   MB.No. {mb_abs_no} P.No. {mb_abs_p_from} to {mb_abs_p_to} (Abstract)"
    )
    c8.font = Font(bold=True)
    c8.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)

    ws_bill.merge_cells("A9:H9")
    c9 = ws_bill["A9"]
    c9.value = f"DOI : {doi}    DOC : {doc}    DOMR : {domr}    DOBR : {dobr}"
    c9.font = Font(bold=True)
    c9.alignment = Alignment(horizontal="left", vertical="center")

    for r in range(1, 10):
        for c_idx in range(1, 9):
            cell = ws_bill.cell(row=r, column=c_idx)
            cell.border = border_all

    header_row = 10
    headers = ["S.No", "Quantity", "Unit", "Item", "Rate", "Per", "Unit", "Amount"]
    for col_idx, text in enumerate(headers, start=1):
        cell = ws_bill.cell(row=header_row, column=col_idx, value=text)
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = border_all
        cell.fill = header_fill

    ws_bill.column_dimensions["A"].width = 6
    ws_bill.column_dimensions["B"].width = 10
    ws_bill.column_dimensions["C"].width = 10
    ws_bill.column_dimensions["D"].width = 45
    ws_bill.column_dimensions["E"].width = 10
    ws_bill.column_dimensions["F"].width = 6
    ws_bill.column_dimensions["G"].width = 10
    ws_bill.column_dimensions["H"].width = 15

    data_start = header_row + 1
    row_idx = data_start
    slno = 1

    for it in items:
        qty = it.get("qty", 0)
        unit_pl = str(it.get("unit") or "").strip()
        desc = it.get("desc") or ""
        rate = it.get("rate", 0.0)
        is_ae = bool(it.get("is_ae", False))

        if is_ae:
            ws_bill.cell(row=row_idx, column=1, value=None)
        else:
            ws_bill.cell(row=row_idx, column=1, value=slno)

        ws_bill.cell(row=row_idx, column=2, value=qty)
        ws_bill.cell(row=row_idx, column=3, value=unit_pl)
        ws_bill.cell(row=row_idx, column=4, value=desc)
        ws_bill.cell(row=row_idx, column=5, value=rate)
        ws_bill.cell(row=row_idx, column=6, value=1)
        ws_bill.cell(row=row_idx, column=7, value=singular_unit(unit_pl))
        ws_bill.cell(row=row_idx, column=8, value=f"=B{row_idx}*E{row_idx}")

        for c_idx in range(1, 9):
            cell = ws_bill.cell(row=row_idx, column=c_idx)
            cell.border = border_all
            if c_idx == 4:
                cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
            else:
                cell.alignment = Alignment(horizontal="center", vertical="center")

        if not is_ae:
            slno += 1

        row_idx += 1

    last_item_row = row_idx - 1

    sub_row = row_idx
    ws_bill.cell(row=sub_row, column=4, value="Sub Total Amount")
    ws_bill.cell(row=sub_row, column=8, value=f"=SUM(H{data_start}:H{last_item_row})")

    for c_idx in range(1, 9):
        cell = ws_bill.cell(row=sub_row, column=c_idx)
        cell.font = Font(bold=True)
        cell.border = border_all
        cell.fill = subtotal_fill
        if c_idx == 4:
            cell.alignment = Alignment(horizontal="left", vertical="center")
        else:
            cell.alignment = Alignment(horizontal="center", vertical="center")

    tp_row = sub_row + 1
    tp_percent = float(tp_percent or 0.0)
    tp_type = tp_type if tp_type in ("Less", "Excess") else "Excess"

    # Show Add or Deduct explicitly based on tp_type
    label_prefix = "Deduct" if tp_type == "Less" else "Add"
    label_tp = f"{label_prefix} T.P @ {tp_percent} % {tp_type}"
    ws_bill.cell(row=tp_row, column=4, value=label_tp)
    ws_bill.cell(row=tp_row, column=8, value=f"=H{sub_row}*{abs(tp_percent)}/100")

    for c_idx in range(1, 9):
        cell = ws_bill.cell(row=tp_row, column=c_idx)
        cell.font = Font(bold=True)
        cell.border = border_all
        cell.fill = subtotal_fill
        if c_idx == 4:
            cell.alignment = Alignment(horizontal="left", vertical="center")
        else:
            cell.alignment = Alignment(horizontal="center", vertical="center")

    total_row = tp_row + 1
    ws_bill.cell(row=total_row, column=4, value="Total")

    if tp_type == "Less":
        ws_bill.cell(row=total_row, column=8, value=f"=H{sub_row}-H{tp_row}")
    else:
        ws_bill.cell(row=total_row, column=8, value=f"=H{sub_row}+H{tp_row}")

    for c_idx in range(1, 9):
        cell = ws_bill.cell(row=total_row, column=c_idx)
        cell.font = Font(bold=True)
        cell.border = border_all
        cell.fill = subtotal_fill
        if c_idx == 4:
            cell.alignment = Alignment(horizontal="left", vertical="center")
        else:
            cell.alignment = Alignment(horizontal="center", vertical="center")

    for r in range(1, ws_bill.max_row + 1):
        ws_bill.row_dimensions[r].height = None


def build_first_bill_wb(items, header_data, title_text,
                        tp_percent, tp_type,
                        mb_measure_no, mb_measure_p_from, mb_measure_p_to,
                        mb_abs_no, mb_abs_p_from, mb_abs_p_to,
                        doi, doc, domr, dobr):
    wb_out = Workbook()
    create_first_bill_sheet(
        wb_out,
        sheet_name="Bill",
        items=items,
        header_data=header_data,
        title_text=title_text,
        tp_percent=tp_percent,
        tp_type=tp_type,
        mb_measure_no=mb_measure_no,
        mb_measure_p_from=mb_measure_p_from,
        mb_measure_p_to=mb_measure_p_to,
        mb_abs_no=mb_abs_no,
        mb_abs_p_from=mb_abs_p_from,
        mb_abs_p_to=mb_abs_p_to,
        doi=doi,
        doc=doc,
        domr=domr,
        dobr=dobr,
    )
    return wb_out


def build_nth_bill_wb(items, header_data, title_text,
                      tp_percent, tp_type,
                      mb_measure_no, mb_measure_p_from, mb_measure_p_to,
                      mb_abs_no, mb_abs_p_from, mb_abs_p_to,
                      doi, doc, domr, dobr):
    """
    Build a complete Nth bill workbook with 11-column format:
    Sl.No, Item, Quantity Till Date, Unit, Rate per Unit, Total Value Till Date,
    Deduct Previous (Qty, Amount), Since Last (Qty, Amount), Remarks
    """
    wb_out = Workbook()
    ws = wb_out.active
    ws.title = "Bill"

    thin = Side(border_style="thin", color="000000")
    border_all = Border(left=thin, right=thin, top=thin, bottom=thin)
    header_fill = PatternFill("solid", fgColor="FFC8C8C8")
    subtotal_fill = PatternFill("solid", fgColor="FFE6E6E6")

    ws.merge_cells("A1:K1")
    c1 = ws["A1"]
    c1.value = title_text
    c1.font = Font(bold=True, size=14)
    c1.alignment = Alignment(horizontal="center", vertical="center")

    work_val = header_data.get("name_of_work", "").strip()
    est_val = header_data.get("estimate_amount", "").strip()
    adm_val = header_data.get("admin_sanction", "").strip()
    tech_val = header_data.get("tech_sanction", "").strip()
    agt_val = header_data.get("agreement", "").strip()
    agency_val = header_data.get("agency", "").strip()

    ws.merge_cells("A2:K2")
    c2 = ws["A2"]
    c2.value = f"Name of the work : {work_val}" if work_val else "Name of the work :"
    c2.font = Font(bold=True)
    c2.alignment = Alignment(horizontal="left", vertical="center")

    ws.merge_cells("A3:K3")
    c3 = ws["A3"]
    c3.value = f"Estimate Amount : {est_val}" if est_val else "Estimate Amount :"
    c3.font = Font(bold=True)
    c3.alignment = Alignment(horizontal="left", vertical="center")

    ws.merge_cells("A4:K4")
    c4 = ws["A4"]
    c4.value = (
        f"Ref. to Administrative sanction : {adm_val}"
        if adm_val else "Ref. to Administrative sanction :"
    )
    c4.font = Font(bold=True)
    c4.alignment = Alignment(horizontal="left", vertical="center")

    ws.merge_cells("A5:K5")
    c5 = ws["A5"]
    c5.value = (
        f"Ref. to Technical sanction : {tech_val}"
        if tech_val else "Ref. to Technical sanction :"
    )
    c5.font = Font(bold=True)
    c5.alignment = Alignment(horizontal="left", vertical="center")

    ws.merge_cells("A6:K6")
    c6 = ws["A6"]
    c6.value = (
        f"Ref. to Agreement : {agt_val}"
        if agt_val else "Ref. to Agreement :"
    )
    c6.font = Font(bold=True)
    c6.alignment = Alignment(horizontal="left", vertical="center")

    ws.merge_cells("A7:K7")
    c7 = ws["A7"]
    c7.value = (
        f"Name of the Agency : {agency_val}"
        if agency_val else "Name of the Agency :"
    )
    c7.font = Font(bold=True)
    c7.alignment = Alignment(horizontal="left", vertical="center")

    # -------- ROW 8: Label + MB details (NTH BILL) IN ONE MERGED CELL --------
    ws.merge_cells("A8:K8")
    c8 = ws["A8"]
    c8.value = (
        f"M.B.No Details: MB.No. {mb_measure_no} P.No. {mb_measure_p_from} to {mb_measure_p_to} (Measurements)   "
        f"&   MB.No. {mb_abs_no} P.No. {mb_abs_p_from} to {mb_abs_p_to} (Abstract)"
    )
    c8.font = Font(bold=True)
    c8.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)

    ws.merge_cells("A9:K9")
    c9 = ws["A9"]
    c9.value = f"DOI : {doi}    DOC : {doc}    DOMR : {domr}    DOBR : {dobr}"
    c9.font = Font(bold=True)
    c9.alignment = Alignment(horizontal="left", vertical="center")

    for r in range(1, 10):
        for col in range(1, 12):
            cell = ws.cell(row=r, column=col)
            cell.border = border_all

    # Merge header cells: S.No, Item, Qty Till Date, Unit, Rate, Total Value, Remarks span 2 rows
    for col in [1, 2, 3, 4, 5, 6, 11]:
        ws.merge_cells(start_row=10, start_column=col, end_row=11, end_column=col)

    ws.merge_cells("G10:H10")  # Deduct Previous
    ws.merge_cells("I10:J10")  # Since Last

    ws.cell(row=10, column=1, value="S.No")
    ws.cell(row=10, column=2, value="Item")
    ws.cell(row=10, column=3, value="Quantity Till Date")
    ws.cell(row=10, column=4, value="Unit")
    ws.cell(row=10, column=5, value="Rate per Unit")
    ws.cell(row=10, column=6, value="Total Value till date")
    ws.cell(row=10, column=7, value="Deduct Previous Measurements")
    ws.cell(row=10, column=9, value="Since Last Measurements")
    ws.cell(row=10, column=11, value="Remarks")

    ws.cell(row=11, column=7, value="Quantity")
    ws.cell(row=11, column=8, value="Amount")
    ws.cell(row=11, column=9, value="Quantity")
    ws.cell(row=11, column=10, value="Amount")

    for r in (10, 11):
        for col in range(1, 12):
            cell = ws.cell(row=r, column=col)
            cell.font = Font(bold=True)
            cell.border = border_all
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    ws.column_dimensions["A"].width = 6
    ws.column_dimensions["B"].width = 45
    ws.column_dimensions["C"].width = 14
    ws.column_dimensions["D"].width = 8
    ws.column_dimensions["E"].width = 12
    ws.column_dimensions["F"].width = 16
    ws.column_dimensions["G"].width = 14
    ws.column_dimensions["H"].width = 16
    ws.column_dimensions["I"].width = 14
    ws.column_dimensions["J"].width = 16
    ws.column_dimensions["K"].width = 20

    data_start = 12
    r = data_start
    sl = 1

    for it in items:
        desc = it.get("desc") or ""
        unit = it.get("unit") or ""
        rate = it.get("rate", 0.0)
        prev_qty = it.get("prev_qty", 0.0)
        prev_amount = it.get("prev_amount", 0.0)

        ws.cell(row=r, column=1, value=sl)
        ws.cell(row=r, column=2, value=desc)
        ws.cell(row=r, column=3, value=None)  # Quantity Till Date (to be filled)
        ws.cell(row=r, column=4, value=unit)
        ws.cell(row=r, column=5, value=rate)
        ws.cell(row=r, column=6, value=f"=C{r}*E{r}")
        ws.cell(row=r, column=7, value=prev_qty)
        ws.cell(row=r, column=8, value=prev_amount)
        ws.cell(row=r, column=9, value=f"=C{r}-G{r}")
        ws.cell(row=r, column=10, value=f"=F{r}-H{r}")
        ws.cell(row=r, column=11, value="")

        for col in range(1, 12):
            cell = ws.cell(row=r, column=col)
            cell.border = border_all
            if col == 2:
                cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
            else:
                cell.alignment = Alignment(horizontal="center", vertical="center")

        r += 1
        sl += 1

    last_item_row = r - 1

    sub_row = r
    ws.cell(row=sub_row, column=2, value="Sub Total")
    ws.cell(row=sub_row, column=6, value=f"=SUM(F{data_start}:F{last_item_row})")
    ws.cell(row=sub_row, column=8, value=f"=SUM(H{data_start}:H{last_item_row})")
    ws.cell(row=sub_row, column=10, value=f"=SUM(J{data_start}:J{last_item_row})")

    tp_row = sub_row + 1
    tp_percent = float(tp_percent or 0.0)
    tp_type = tp_type if tp_type in ("Less", "Excess") else "Excess"

    # Show Add or Deduct explicitly based on tp_type
    label_prefix = "Deduct" if tp_type == "Less" else "Add"
    ws.cell(row=tp_row, column=2, value=f"{label_prefix} T.P @ {tp_percent} % {tp_type}")
    ws.cell(row=tp_row, column=6, value=f"=F{sub_row}*{abs(tp_percent)}/100")
    ws.cell(row=tp_row, column=8, value=f"=H{sub_row}*{abs(tp_percent)}/100")
    ws.cell(row=tp_row, column=10, value=f"=J{sub_row}*{abs(tp_percent)}/100")

    total_row = tp_row + 1
    ws.cell(row=total_row, column=2, value="Total")

    if tp_type == "Less":
        ws.cell(row=total_row, column=6, value=f"=F{sub_row}-F{tp_row}")
        ws.cell(row=total_row, column=8, value=f"=H{sub_row}-H{tp_row}")
        ws.cell(row=total_row, column=10, value=f"=J{sub_row}-J{tp_row}")
    else:
        ws.cell(row=total_row, column=6, value=f"=F{sub_row}+F{tp_row}")
        ws.cell(row=total_row, column=8, value=f"=H{sub_row}+H{tp_row}")
        ws.cell(row=total_row, column=10, value=f"=J{sub_row}+J{tp_row}")

    for rr in [sub_row, tp_row, total_row]:
        for col in range(1, 12):
            cell = ws.cell(row=rr, column=col)
            cell.font = Font(bold=True)
            cell.border = border_all
            cell.fill = subtotal_fill
            if col == 2:
                cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            else:
                cell.alignment = Alignment(horizontal="center", vertical="center")

    for rr in range(1, ws.max_row + 1):
        ws.row_dimensions[rr].height = None

    return wb_out


def _populate_nth_bill_sheet(ws, items, header_data, title_text,
                             tp_percent, tp_type,
                             mb_measure_no, mb_measure_p_from, mb_measure_p_to,
                             mb_abs_no, mb_abs_p_from, mb_abs_p_to,
                             doi, doc, domr, dobr):
    """Populate an existing worksheet with Nth bill data.
    
    This is similar to build_nth_bill_wb but works on an existing sheet.
    Uses 11-column format with Unit column.
    """
    thin = Side(border_style="thin", color="000000")
    border_all = Border(left=thin, right=thin, top=thin, bottom=thin)
    header_fill = PatternFill("solid", fgColor="FFC8C8C8")
    subtotal_fill = PatternFill("solid", fgColor="FFE6E6E6")

    ws.merge_cells("A1:K1")
    c1 = ws["A1"]
    c1.value = title_text
    c1.font = Font(bold=True, size=14)
    c1.alignment = Alignment(horizontal="center", vertical="center")

    work_val = header_data.get("name_of_work", "").strip()
    est_val = header_data.get("estimate_amount", "").strip()
    adm_val = header_data.get("admin_sanction", "").strip()
    tech_val = header_data.get("tech_sanction", "").strip()
    agt_val = header_data.get("agreement", "").strip()
    agency_val = header_data.get("agency", "").strip()

    ws.merge_cells("A2:K2")
    c2 = ws["A2"]
    c2.value = f"Name of the work : {work_val}" if work_val else "Name of the work :"
    c2.font = Font(bold=True)
    c2.alignment = Alignment(horizontal="left", vertical="center")

    ws.merge_cells("A3:K3")
    c3 = ws["A3"]
    c3.value = f"Estimate Amount : {est_val}" if est_val else "Estimate Amount :"
    c3.font = Font(bold=True)
    c3.alignment = Alignment(horizontal="left", vertical="center")

    ws.merge_cells("A4:K4")
    c4 = ws["A4"]
    c4.value = (
        f"Ref. to Administrative sanction : {adm_val}"
        if adm_val else "Ref. to Administrative sanction :"
    )
    c4.font = Font(bold=True)
    c4.alignment = Alignment(horizontal="left", vertical="center")

    ws.merge_cells("A5:K5")
    c5 = ws["A5"]
    c5.value = (
        f"Ref. to Technical sanction : {tech_val}"
        if tech_val else "Ref. to Technical sanction :"
    )
    c5.font = Font(bold=True)
    c5.alignment = Alignment(horizontal="left", vertical="center")

    ws.merge_cells("A6:K6")
    c6 = ws["A6"]
    c6.value = (
        f"Ref. to Agreement : {agt_val}"
        if agt_val else "Ref. to Agreement :"
    )
    c6.font = Font(bold=True)
    c6.alignment = Alignment(horizontal="left", vertical="center")

    ws.merge_cells("A7:K7")
    c7 = ws["A7"]
    c7.value = (
        f"Name of the Agency : {agency_val}"
        if agency_val else "Name of the Agency :"
    )
    c7.font = Font(bold=True)
    c7.alignment = Alignment(horizontal="left", vertical="center")

    ws.merge_cells("A8:K8")
    c8 = ws["A8"]
    c8.value = (
        f"M.B.No Details: MB.No. {mb_measure_no} P.No. {mb_measure_p_from} to {mb_measure_p_to} (Measurements)   "
        f"&   MB.No. {mb_abs_no} P.No. {mb_abs_p_from} to {mb_abs_p_to} (Abstract)"
    )
    c8.font = Font(bold=True)
    c8.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)

    ws.merge_cells("A9:K9")
    c9 = ws["A9"]
    c9.value = f"DOI : {doi}    DOC : {doc}    DOMR : {domr}    DOBR : {dobr}"
    c9.font = Font(bold=True)
    c9.alignment = Alignment(horizontal="left", vertical="center")

    for r in range(1, 10):
        for col in range(1, 12):
            cell = ws.cell(row=r, column=col)
            cell.border = border_all

    # Merge header cells: S.No, Item, Qty Till Date, Unit, Rate, Total Value, Remarks span 2 rows
    for col in [1, 2, 3, 4, 5, 6, 11]:
        ws.merge_cells(start_row=10, start_column=col, end_row=11, end_column=col)

    ws.merge_cells("G10:H10")  # Deduct Previous
    ws.merge_cells("I10:J10")  # Since Last

    ws.cell(row=10, column=1, value="S.No")
    ws.cell(row=10, column=2, value="Item")
    ws.cell(row=10, column=3, value="Quantity Till Date")
    ws.cell(row=10, column=4, value="Unit")
    ws.cell(row=10, column=5, value="Rate per Unit")
    ws.cell(row=10, column=6, value="Total Value till date")
    ws.cell(row=10, column=7, value="Deduct Previous Measurements")
    ws.cell(row=10, column=9, value="Since Last Measurements")
    ws.cell(row=10, column=11, value="Remarks")

    ws.cell(row=11, column=7, value="Quantity")
    ws.cell(row=11, column=8, value="Amount")
    ws.cell(row=11, column=9, value="Quantity")
    ws.cell(row=11, column=10, value="Amount")

    for r in (10, 11):
        for col in range(1, 12):
            cell = ws.cell(row=r, column=col)
            cell.font = Font(bold=True)
            cell.border = border_all
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    ws.column_dimensions["A"].width = 6
    ws.column_dimensions["B"].width = 45
    ws.column_dimensions["C"].width = 14
    ws.column_dimensions["D"].width = 8
    ws.column_dimensions["E"].width = 12
    ws.column_dimensions["F"].width = 16
    ws.column_dimensions["G"].width = 14
    ws.column_dimensions["H"].width = 16
    ws.column_dimensions["I"].width = 14
    ws.column_dimensions["J"].width = 16
    ws.column_dimensions["K"].width = 20

    data_start = 12
    r = data_start
    sl = 1

    for it in items:
        desc = it.get("desc") or ""
        unit = it.get("unit") or ""
        rate = it.get("rate", 0.0)
        prev_qty = it.get("prev_qty", 0.0)
        prev_amount = it.get("prev_amount", 0.0)

        ws.cell(row=r, column=1, value=sl)
        ws.cell(row=r, column=2, value=desc)
        ws.cell(row=r, column=3, value=None)  # Quantity Till Date (to be filled)
        ws.cell(row=r, column=4, value=unit)
        ws.cell(row=r, column=5, value=rate)
        ws.cell(row=r, column=6, value=f"=C{r}*E{r}")
        ws.cell(row=r, column=7, value=prev_qty)
        ws.cell(row=r, column=8, value=prev_amount)
        ws.cell(row=r, column=9, value=f"=C{r}-G{r}")
        ws.cell(row=r, column=10, value=f"=F{r}-H{r}")
        ws.cell(row=r, column=11, value="")

        for col in range(1, 12):
            cell = ws.cell(row=r, column=col)
            cell.border = border_all
            if col == 2:
                cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
            else:
                cell.alignment = Alignment(horizontal="center", vertical="center")

        r += 1
        sl += 1

    last_item_row = r - 1

    sub_row = r
    ws.cell(row=sub_row, column=2, value="Sub Total")
    ws.cell(row=sub_row, column=6, value=f"=SUM(F{data_start}:F{last_item_row})")
    ws.cell(row=sub_row, column=8, value=f"=SUM(H{data_start}:H{last_item_row})")
    ws.cell(row=sub_row, column=10, value=f"=SUM(J{data_start}:J{last_item_row})")

    tp_row = sub_row + 1
    tp_percent = float(tp_percent or 0.0)
    tp_type = tp_type if tp_type in ("Less", "Excess") else "Excess"

    label_prefix = "Deduct" if tp_type == "Less" else "Add"
    ws.cell(row=tp_row, column=2, value=f"{label_prefix} T.P @ {tp_percent} % {tp_type}")
    ws.cell(row=tp_row, column=6, value=f"=F{sub_row}*{abs(tp_percent)}/100")
    ws.cell(row=tp_row, column=8, value=f"=H{sub_row}*{abs(tp_percent)}/100")
    ws.cell(row=tp_row, column=10, value=f"=J{sub_row}*{abs(tp_percent)}/100")

    total_row = tp_row + 1
    ws.cell(row=total_row, column=2, value="Total")

    if tp_type == "Less":
        ws.cell(row=total_row, column=6, value=f"=F{sub_row}-F{tp_row}")
        ws.cell(row=total_row, column=8, value=f"=H{sub_row}-H{tp_row}")
        ws.cell(row=total_row, column=10, value=f"=J{sub_row}-J{tp_row}")
    else:
        ws.cell(row=total_row, column=6, value=f"=F{sub_row}+F{tp_row}")
        ws.cell(row=total_row, column=8, value=f"=H{sub_row}+H{tp_row}")
        ws.cell(row=total_row, column=10, value=f"=J{sub_row}+J{tp_row}")

    for rr in [sub_row, tp_row, total_row]:
        for col in range(1, 12):
            cell = ws.cell(row=rr, column=col)
            cell.font = Font(bold=True)
            cell.border = border_all
            cell.fill = subtotal_fill
            if col == 2:
                cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            else:
                cell.alignment = Alignment(horizontal="center", vertical="center")

    for rr in range(1, ws.max_row + 1):
        ws.row_dimensions[rr].height = None


@login_required(login_url='login')
def bill(request):
    """
    Bill generator with four sections:

    SECTION 1: Bill from Estimate
        action = "estimate_first_part"   â†’ CC First & Part Bill
        action = "estimate_first_final"  â†’ CC First & Final Bill

    SECTION 2: Bill from WorkSlip
        action = "workslip_first_part"   â†’ CC First & Part Bill
        action = "workslip_first_final"  â†’ CC First & Final Bill

    SECTION 3: Nth & Part / 2nd & Final from First & Part Bill
        action = "firstpart_nth_part"    â†’ CC Nth & Part Bill
        action = "firstpart_2nd_final"   â†’ CC Second & Final Bill
        - Prev Qty (F) = Column B of uploaded First Bill
        - Prev Amt (G) = Column H of uploaded First Bill
                         (if H is formula with no value â†’ recompute as B * E)

    SECTION 4: Nth & Part / Nth & Final from Nth & Part Bill
        action = "nth_nth_part"          â†’ CC Nth & Part Bill
        action = "nth_nth_final"         â†’ CC Nth & Final Bill
        - Prev Qty (F) = Column C (Quantity Till Date)
        - Prev Amt (G) = Column E (Total Value till date)
                         (if E has no value â†’ recompute as C * D)
    """

    print("DEBUG: Entered bill view; method=", getattr(request, 'method', None))
    # If GET, render the page. For POST we currently return a safe stub
    # (temporary) so the view never returns None while refactoring continues.
    method = getattr(request, 'method', '').upper()
    if method == 'GET':
        # Get user's document templates for display
        from core.template_views import get_user_template
        covering_letter_template = get_user_template(request.user, 'covering_letter')
        movement_slip_template = get_user_template(request.user, 'movement_slip')
        return render(request, "core/bill.html", {
            'covering_letter_template': covering_letter_template,
            'movement_slip_template': movement_slip_template,
        })

    if method == 'POST':
        action = str(request.POST.get('action') or '').strip()
        bill_type = str(request.POST.get('bill_type') or '').strip()
        uploaded = request.FILES.get('bill_file') or request.FILES.get('file')
        if not uploaded:
            return JsonResponse({"error": "no uploaded file"}, status=400)

        # Extract MB details from POST request
        mb_measure_no = str(request.POST.get('mb_measure_no') or '').strip()
        mb_measure_p_from = str(request.POST.get('mb_measure_p_from') or '').strip()
        mb_measure_p_to = str(request.POST.get('mb_measure_p_to') or '').strip()
        mb_abs_no = str(request.POST.get('mb_abstract_no') or '').strip()
        mb_abs_p_from = str(request.POST.get('mb_abstract_p_from') or '').strip()
        mb_abs_p_to = str(request.POST.get('mb_abstract_p_to') or '').strip()
        
        # Extract dates from POST request
        doi = str(request.POST.get('doi') or '').strip()
        doc = str(request.POST.get('doc') or '').strip()
        domr = str(request.POST.get('domr') or '').strip()
        dobr = str(request.POST.get('dobr') or '').strip()

        try:
            wb = load_workbook(uploaded, data_only=True)
        except Exception as e:
            return JsonResponse({"error": f"failed to read uploaded workbook: {e}"}, status=400)

        # Helper function to generate bill title (used by both estimate and workslip)
        def get_bill_title(action, idx, total):
            """Generate appropriate bill title based on action (button clicked)."""
            if not action:
                # No action, just use "Bill" without numbers
                return "CC Bill"
            
            # Map action to bill type descriptions with CC prefix
            type_map = {
                "estimate_first_part": "CC First & Part Bill",
                "estimate_first_final": "CC First & Final Bill",
                "workslip_first_part": "CC First & Part Bill",
                "workslip_first_final": "CC First & Final Bill",
                "firstpart_nth_part": "CC Nth & Part Bill",
                "firstpart_2nd_final": "CC 2nd & Final Bill",
                "nth_nth_part": "CC Nth & Part Bill",
                "nth_nth_final": "CC Nth & Final Bill",
            }
            
            # Get the bill type name from the action
            bill_type_name = type_map.get(action, "CC Bill")
            
            # If only one bill, return type name without numbering
            if total <= 1:
                return bill_type_name
            
            # Otherwise return type name (no numbers added, just the type)
            return bill_type_name

        # Support estimate -> first bill (minimal)
        if action.startswith('estimate_'):
            # Try to detect obvious estimate sheets first
            pairs = find_all_estimate_sheets_and_header_rows(wb)

            # If detection failed, fallback to trying every worksheet with a
            # conservative header guess. We'll keep only sheets that parse
            # into at least one item.
            if not pairs:
                candidates = []
                for ws in wb.worksheets:
                    items_try = parse_estimate_items(ws, 3)
                    if items_try:
                        candidates.append((ws, 3))
                if not candidates:
                    # final fallback: try header_row=2 and 4 as last resort
                    for ws in wb.worksheets:
                        for hr in (2, 4):
                            items_try = parse_estimate_items(ws, hr)
                            if items_try:
                                candidates.append((ws, hr))
                                break
                pairs = candidates

            if not pairs:
                return JsonResponse({"error": "no estimate-like sheets found"}, status=400)

            # Build output workbook with one Bill sheet per parsed estimate sheet.
            wb_out = Workbook()
            created = 0
            for idx, (ws, header_row) in enumerate(pairs, start=1):
                items = parse_estimate_items(ws, header_row)
                if not items:
                    continue
                created += 1
                
                # Extract header data from this specific sheet
                sheet_header_data = _extract_header_data_from_sheet(ws)
                
                # Extract T.P percentage and type from workslip
                tp_percent, tp_type = extract_tp_from_workslip(ws)
                
                # Generate appropriate title
                title_text = get_bill_title(action, idx, len(pairs))
                
                # Only add sheet numbering if there are multiple sheets
                if len(pairs) > 1:
                    sheet_name = f"Bill-{created}"
                else:
                    sheet_name = "Bill"
                
                create_first_bill_sheet(
                    wb_out,
                    sheet_name=sheet_name,
                    items=items,
                    header_data=sheet_header_data,
                    title_text=title_text,
                    tp_percent=None,
                    tp_type=None,
                    mb_measure_no=mb_measure_no, 
                    mb_measure_p_from=mb_measure_p_from, 
                    mb_measure_p_to=mb_measure_p_to,
                    mb_abs_no=mb_abs_no, 
                    mb_abs_p_from=mb_abs_p_from, 
                    mb_abs_p_to=mb_abs_p_to,
                    doi=doi, 
                    doc=doc, 
                    domr=domr, 
                    dobr=dobr
                )

            if created == 0:
                return JsonResponse({"error": "no items parsed from upload"}, status=400)

            resp = HttpResponse(
                content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
            resp["Content-Disposition"] = 'attachment; filename="Bills.xlsx"'
            wb_out.save(resp)
            return resp

        # Support workslip -> first bill
        if action.startswith('workslip_'):
            # Try to detect workslip sheets
            workslip_sheets = find_all_workslip_sheets(wb)
            # Defaults in case extraction fails
            tp_percent = 0.0
            tp_type = "Excess"
            
            # Build output workbook with one Bill sheet per workslip sheet
            wb_out = Workbook()
            created = 0
            for idx, ws in enumerate(workslip_sheets, start=1):
                items = parse_workslip_items(ws)
                if not items:
                    continue
                created += 1
                
                # Extract header data from this specific sheet
                sheet_header_data = _extract_header_data_from_sheet(ws)
                
                # Extract T.P percentage and type from workslip
                try:
                    tp_percent, tp_type = extract_tp_from_workslip(ws)
                    print(f"DEBUG: Workslip '{ws.title}' TP extracted percent={tp_percent} type={tp_type}")
                except Exception as e:
                    print(f"DEBUG: Failed to extract TP from workslip '{ws.title}': {e}")
                    tp_percent, tp_type = 0.0, "Excess"
                
                # Generate appropriate title
                title_text = get_bill_title(action, idx, len(workslip_sheets))
                
                # Only add sheet numbering if there are multiple sheets
                if len(workslip_sheets) > 1:
                    sheet_name = f"Bill-{created}"
                else:
                    sheet_name = "Bill"
                
                create_first_bill_sheet(
                    wb_out,
                    sheet_name=sheet_name,
                    items=items,
                    header_data=sheet_header_data,
                    title_text=title_text,
                    tp_percent=tp_percent,
                    tp_type=tp_type,
                    mb_measure_no=mb_measure_no, 
                    mb_measure_p_from=mb_measure_p_from, 
                    mb_measure_p_to=mb_measure_p_to,
                    mb_abs_no=mb_abs_no, 
                    mb_abs_p_from=mb_abs_p_from, 
                    mb_abs_p_to=mb_abs_p_to,
                    doi=doi, 
                    doc=doc, 
                    domr=domr, 
                    dobr=dobr
                )

            if created == 0:
                return JsonResponse({"error": "no workslip items parsed from upload"}, status=400)

            resp = HttpResponse(
                content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
            resp["Content-Disposition"] = 'attachment; filename="Bills.xlsx"'
            wb_out.save(resp)
            return resp

        # Support Nth/2nd bills from First Bill (temporary shim)
        # Support Nth/2nd bills from First Bill (multi-sheet support)
        if action in ("firstpart_nth_part", "firstpart_2nd_final"):
            # Find all bill sheets (from multiple estimates)
            bill_sheets = [ws for ws in wb.worksheets if ws.title.startswith("Bill")]
            print(f"DEBUG: Found {len(bill_sheets)} sheets starting with 'Bill': {[ws.title for ws in bill_sheets]}")
            
            if not bill_sheets:
                # Fallback: find all sheets that look like estimate/bill sheets
                bill_sheets = []
                for ws in wb.worksheets:
                    print(f"DEBUG: Checking sheet '{ws.title}'...")
                    for r in range(1, 30):
                        a = str(ws.cell(row=r, column=1).value or "").strip().lower()
                        b = str(ws.cell(row=r, column=2).value or "").strip().lower()
                        d = str(ws.cell(row=r, column=4).value or "").strip().lower()
                        if "sl" in a and "quantity" in b and ("item" in d or "description" in d):
                            if ws not in bill_sheets:
                                bill_sheets.append(ws)
                            print(f"DEBUG: Found bill-like sheet '{ws.title}' with header at row {r}")
                            break
                
                print(f"DEBUG: Fallback found {len(bill_sheets)} bill sheets: {[ws.title for ws in bill_sheets]}")
            
            if not bill_sheets:
                # Last fallback: try all non-empty sheets
                bill_sheets = [ws for ws in wb.worksheets if ws.max_row > 1]
                print(f"DEBUG: No bill sheets found by fallback, using all non-empty sheets: {[ws.title for ws in bill_sheets]}")
            
            if not bill_sheets:
                bill_sheets = [wb.worksheets[0]]  # fallback to first sheet
                print(f"DEBUG: No bill sheets found, using first sheet '{bill_sheets[0].title}'")
            
            print(f"DEBUG: Processing {len(bill_sheets)} bill sheets total: {[ws.title for ws in bill_sheets]}")
            
            # Extract Nth number
            if action == "firstpart_nth_part":
                nth_str = request.POST.get("nth_number", "").strip()
                try:
                    nth_val = int(nth_str)
                except Exception:
                    nth_val = 2
                if nth_val < 2:
                    nth_val = 2
                ord_word = ordinal_word(nth_val)
                # Prefer bill_type (dropdown) if present, else action
                which = bill_type if bill_type in ("firstpart_nth_part", "firstpart_2nd_final") else action
                if which == "firstpart_nth_part":
                    title_text = f"CC {ord_word} & Part Bill"
                else:
                    title_text = f"CC {ord_word} & Final Bill"
            else:
                ord_word = ordinal_word(2)
                title_text = f"CC {ord_word} & Final Bill"

            # Build output workbook with one Nth bill sheet per input bill sheet
            wb_out = Workbook()
            created = 0
            
            for idx, ws_first in enumerate(bill_sheets, start=1):
                print(f"DEBUG: Processing sheet {idx}/{len(bill_sheets)}: '{ws_first.title}'")
                
                # Find header row in this sheet
                header_row = None
                for r in range(1, 30):
                    a = str(ws_first.cell(row=r, column=1).value or "").strip().lower()
                    b = str(ws_first.cell(row=r, column=2).value or "").strip().lower()
                    d = str(ws_first.cell(row=r, column=4).value or "").strip().lower()
                    if "sl" in a and "quantity" in b and ("item" in d or "description" in d):
                        header_row = r
                        print(f"DEBUG: Found header row {r} in sheet '{ws_first.title}'")
                        break
                
                if header_row is None:
                    header_row = 10  # default fallback
                    print(f"DEBUG: No header found, using default row 10")
                
                header_data = _extract_header_data_from_sheet(ws_first)
                items = parse_first_bill_for_nth(ws_first, header_row)
                print(f"DEBUG: Parsed {len(items)} items from sheet '{ws_first.title}'")

                if not items:
                    print(f"DEBUG: No items found, skipping sheet '{ws_first.title}'")
                    continue
                created += 1

                tp_percent, tp_type = read_tp_from_sheet(ws_first)
                if tp_percent is None:
                    tp_percent = 0.0
                if tp_type is None:
                    tp_type = "Excess"

                # Determine sheet name
                if len(bill_sheets) > 1:
                    sheet_name = f"Bill-{created}"
                else:
                    sheet_name = "Bill"

                # Keep a single title text even when multiple sheets are generated
                sheet_title = title_text

                # Create Nth bill sheet
                if created == 1:
                    ws_nth = wb_out.active
                    ws_nth.title = sheet_name
                    print(f"DEBUG: Using active sheet for output, naming it '{sheet_name}'")
                else:
                    ws_nth = wb_out.create_sheet(title=sheet_name)
                    print(f"DEBUG: Created new sheet '{sheet_name}'")

                # Populate this Nth bill sheet
                print(f"DEBUG: Populating sheet '{sheet_name}' with {len(items)} items")
                _populate_nth_bill_sheet(
                    ws_nth,
                    items=items,
                    header_data=header_data,
                    title_text=sheet_title,
                    tp_percent=tp_percent,
                    tp_type=tp_type,
                    mb_measure_no=mb_measure_no,
                    mb_measure_p_from=mb_measure_p_from,
                    mb_measure_p_to=mb_measure_p_to,
                    mb_abs_no=mb_abs_no,
                    mb_abs_p_from=mb_abs_p_from,
                    mb_abs_p_to=mb_abs_p_to,
                    doi=doi,
                    doc=doc,
                    domr=domr,
                    dobr=dobr,
                )
                print(f"DEBUG: Populated sheet '{sheet_name}' successfully")

            print(f"DEBUG: Created {created} output sheets total")

            if created == 0:
                return JsonResponse({"error": "no items parsed from any First Bill sheets"}, status=400)

            resp = HttpResponse(
                content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
            filename = "Nth_Bill_from_FirstPart.xlsx" if action == "firstpart_nth_part" else "Second_Final_from_FirstPart.xlsx"
            resp["Content-Disposition"] = f'attachment; filename="{filename}"'
            wb_out.save(resp)
            return resp

        # Support Nth from Nth bill (temporary shim)
        # Support Nth from Nth bill (multi-sheet support)
        if action in ("nth_nth_part", "nth_nth_final"):
            # Find all Nth bill sheets
            bill_sheets = [ws for ws in wb.worksheets if ws.title.startswith("Bill")]
            print(f"DEBUG (Nthâ†’Nth): Found {len(bill_sheets)} sheets starting with 'Bill': {[ws.title for ws in bill_sheets]}")
            
            if not bill_sheets:
                # Fallback: find sheets with "Quantity Till Date" header
                bill_sheets = []
                for ws in wb.worksheets:
                    print(f"DEBUG (Nthâ†’Nth): Checking sheet '{ws.title}'...")
                    for r in range(1, 30):
                        c = str(ws.cell(row=r, column=3).value or "").strip().lower()
                        if "quantity till date" in c:
                            if ws not in bill_sheets:
                                bill_sheets.append(ws)
                            print(f"DEBUG (Nthâ†’Nth): Found Nth-bill-like sheet '{ws.title}' with header at row {r}")
                            break
                
                print(f"DEBUG (Nthâ†’Nth): Fallback found {len(bill_sheets)} bill sheets: {[ws.title for ws in bill_sheets]}")
            
            if not bill_sheets:
                # Last fallback: try all non-empty sheets
                bill_sheets = [ws for ws in wb.worksheets if ws.max_row > 1]
                print(f"DEBUG (Nthâ†’Nth): No bill sheets found by fallback, using all non-empty sheets: {[ws.title for ws in bill_sheets]}")
            
            if not bill_sheets:
                bill_sheets = [wb.worksheets[0]]  # fallback to first sheet
                print(f"DEBUG (Nthâ†’Nth): No bill sheets found, using first sheet '{bill_sheets[0].title}'")
            
            print(f"DEBUG (Nthâ†’Nth): Processing {len(bill_sheets)} bill sheets total: {[ws.title for ws in bill_sheets]}")
            
            # Extract Nth number
            nth_str = request.POST.get("nth_number", "").strip()
            try:
                nth_val = int(nth_str)
            except Exception:
                nth_val = 2
            if nth_val < 2:
                nth_val = 2
            ord_word = ordinal_word(nth_val)

            # Prefer bill_type (dropdown) if present, else action
            which = bill_type if bill_type in ("nth_nth_part", "nth_nth_final") else action
            if which == "nth_nth_part":
                title_text = f"CC {ord_word} & Part Bill"
            else:
                title_text = f"CC {ord_word} & Final Bill"

            # Build output workbook with one Nth bill sheet per input sheet
            wb_out = Workbook()
            created = 0
            
            print(f"DEBUG (Nthâ†’Nth): Processing {len(bill_sheets)} input bill sheets")
            
            for idx, ws_nth in enumerate(bill_sheets, start=1):
                print(f"DEBUG (Nthâ†’Nth): Processing sheet {idx}/{len(bill_sheets)}: '{ws_nth.title}'")
                
                # Find header row in this sheet
                header_row = None
                for r in range(1, 30):
                    a = str(ws_nth.cell(row=r, column=1).value or "").strip().lower()
                    c = str(ws_nth.cell(row=r, column=3).value or "").strip().lower()
                    if "sl" in a and "quantity till date" in c:
                        header_row = r
                        print(f"DEBUG (Nthâ†’Nth): Found header row {r} in sheet '{ws_nth.title}'")
                        break
                
                if header_row is None:
                    header_row = 10  # default fallback
                    print(f"DEBUG (Nthâ†’Nth): No header found, using default row 10")
                
                header_data = _extract_header_data_from_sheet(ws_nth)
                items = parse_nth_bill_for_next(ws_nth, header_row)
                print(f"DEBUG (Nthâ†’Nth): Parsed {len(items)} items from sheet '{ws_nth.title}'")

                if not items:
                    print(f"DEBUG (Nthâ†’Nth): No items found, skipping sheet '{ws_nth.title}'")
                    continue
                created += 1

                tp_percent, tp_type = read_tp_from_sheet(ws_nth)
                if tp_percent is None:
                    tp_percent = 0.0
                if tp_type is None:
                    tp_type = "Excess"

                # Determine sheet name
                if len(bill_sheets) > 1:
                    sheet_name = f"Bill-{created}"
                else:
                    sheet_name = "Bill"

                # Keep a single title text even when multiple sheets are generated
                sheet_title = title_text

                # Create Nth bill sheet
                if created == 1:
                    ws_out = wb_out.active
                    ws_out.title = sheet_name
                    print(f"DEBUG (Nthâ†’Nth): Using active sheet for output, naming it '{sheet_name}'")
                else:
                    ws_out = wb_out.create_sheet(title=sheet_name)
                    print(f"DEBUG (Nthâ†’Nth): Created new sheet '{sheet_name}'")

                # Populate this Nth bill sheet
                print(f"DEBUG (Nthâ†’Nth): Populating sheet '{sheet_name}' with {len(items)} items")
                _populate_nth_bill_sheet(
                    ws_out,
                    items=items,
                    header_data=header_data,
                    title_text=sheet_title,
                    tp_percent=tp_percent,
                    tp_type=tp_type,
                    mb_measure_no=mb_measure_no,
                    mb_measure_p_from=mb_measure_p_from,
                    mb_measure_p_to=mb_measure_p_to,
                    mb_abs_no=mb_abs_no,
                    mb_abs_p_from=mb_abs_p_from,
                    mb_abs_p_to=mb_abs_p_to,
                    doi=doi,
                    doc=doc,
                    domr=domr,
                    dobr=dobr,
                )
                print(f"DEBUG (Nthâ†’Nth): Populated sheet '{sheet_name}' successfully")

            print(f"DEBUG (Nthâ†’Nth): Created {created} output sheets total")
            
            if created == 0:
                return JsonResponse({"error": "no items parsed from any Nth Bill sheets"}, status=400)

            resp = HttpResponse(
                content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
            filename = "Nth_Bill_from_NthPart.xlsx" if action == "nth_nth_part" else "Nth_Final_from_NthPart.xlsx"
            resp["Content-Disposition"] = f'attachment; filename="{filename}"'
            wb_out.save(resp)
            return resp

        return JsonResponse({"error": "action not supported in temporary handler"}, status=501)

    return HttpResponseNotAllowed(['GET', 'POST'])
def build_estimate_wb(ws_src, blocks):
    """Minimal helper to build an estimate workbook from source and blocks.

    This is a conservative stub that returns an empty workbook with a single
    'Estimate' sheet so callers expecting a Workbook can proceed. Replace
    with real implementation when available.
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "Estimate"
    return wb
    print('DEBUG: defined build_estimate_wb')

    def parse_estimate_items(ws, header_row):
        """
        Estimate / Bill-like format:
          A: Sl.No, B: Quantity, C: Unit, D: Item, E: Rate, F: Per, G: Unit, H: Amount

        Rules:
          - HEADING ROW (skip, but do NOT stop):
              * D has text AND (E is empty OR D is merged)
          - Repeat this skipping for all such heading rows
            until we hit:
              * 'ECV', 'Sub Total', 'Subtotal', 'Total' â†’ STOP.
          - Completely blank rows are just skipped (never used as stop condition).
        """
        items = []
        max_row = min(ws.max_row, 5000)

        for r in range(header_row + 1, max_row + 1):
            desc_raw = ws.cell(row=r, column=4).value  # D
            desc = str(desc_raw or "").strip()
            rate_raw = ws.cell(row=r, column=5).value  # E
            amt_raw = ws.cell(row=r, column=8).value   # H

            desc_low = desc.lower()

            # ---- HARD STOP on totals / ECV rows ----
            if desc_low.startswith(("ecv", "sub total", "subtotal", "total")):
                break

            rate_str = "" if rate_raw is None else str(rate_raw).strip()
            is_rate_empty = (rate_str == "")

            # ---------- HEADING ROW ----------
            if desc and (is_rate_empty or is_merged_cell(ws, r, 4)):
                # pure heading â†’ skip, but DO NOT stop
                continue

            qty_raw = ws.cell(row=r, column=2).value  # B
            unit_raw = ws.cell(row=r, column=3).value  # C

            qty_str = "" if qty_raw is None else str(qty_raw).strip()
            unit_str = "" if unit_raw is None else str(unit_raw).strip()
            amt_str = "" if amt_raw is None else str(amt_raw).strip()

            all_blank = (
                desc == "" and
                is_rate_empty and
                qty_str == "" and
                unit_str == "" and
                amt_str == ""
            )

            # Completely blank row â†’ just skip
            if all_blank:
                continue

            if desc == "" and is_rate_empty:
                # still junk, skip
                continue

            # ---------- REAL ITEM ROW ----------
            qty = to_number(qty_raw)
            unit = unit_str
            rate = to_number(rate_raw)
            amt = to_number(amt_raw)

            if amt != 0:
                if qty == 0 and rate != 0:
                    qty = amt / rate
                elif rate == 0 and qty != 0:
                    rate = amt / qty

            items.append({
                "qty": qty,
                "unit": unit,
                "desc": desc,
                "rate": rate,
            })

        return items

    def find_workslip_sheet(wb):
        """
        Single-sheet heuristic for backward compatibility.
        """
        for ws in wb.worksheets:
            for r in range(1, 40):
                b = str(ws.cell(row=r, column=2).value or "").strip().lower()
                g = str(ws.cell(row=r, column=7).value or "").strip().lower()
                if "description" in b and ("qty" in g or "quantity" in g):
                    return ws
        return wb.worksheets[0]

    def find_all_workslip_sheets(wb):
        """
        NEW: multi-sheet heuristic for WorkSlips.
        """
        results = []
        for ws in wb.worksheets:
            max_scan = min(ws.max_row, 60)
            for r in range(1, max_scan + 1):
                b = str(ws.cell(row=r, column=2).value or "").strip().lower()
                c = str(ws.cell(row=r, column=3).value or "").strip().lower()

                has_desc = (
                    "description of item" in b or
                    "description of item" in c or
                    ("description" in b and "item" in b) or
                    ("description" in c and "item" in c)
                )

                has_qty = False
                for col in range(5, 12):
                    t = str(ws.cell(row=r, column=col).value or "").strip().lower()
                    if "qty" in t or "quantity" in t:
                        has_qty = True
                        break

                if has_desc and has_qty:
                    results.append(ws)
                    break

        if not results:
            results.append(find_workslip_sheet(wb))

        return results

    def parse_workslip_items(ws):
        """
        Parse workslip items, automatically detecting the last workslip phase columns.
        
        For multi-phase workslips (Workslip-1, Workslip-2, etc.), this function will
        find and use the LAST workslip's Qty/Amt columns for bill generation.
        Rate is always from Estimate Rate column for multi-phase workslips.
        """
        items = []
        max_row = min(ws.max_row, 5000)
        
        # Step 1: Find header row and detect column structure
        header_row = 8  # Default
        for r in range(1, 15):
            cell_val = str(ws.cell(row=r, column=1).value or "").strip().lower()
            if "sl" in cell_val:
                header_row = r
                break
        
        # Step 2: Scan header row to find all workslip/execution qty columns
        qty_columns = []
        amt_columns = []
        estimate_rate_col = 5  # Default - Rate (Estimate) column
        
        for c in range(1, 40):  # Scan more columns for multi-phase workslips
            header = str(ws.cell(row=header_row, column=c).value or "").strip().lower()
            
            # Check for workslip/execution qty columns
            is_exec_col = ("execution" in header or "exec" in header or "workslip" in header)
            is_qty = ("qty" in header or "quantity" in header)
            is_amt = ("amount" in header or "amt" in header)
            is_estimate = ("est" in header or "estimate" in header)
            is_rate = ("rate" in header)
            
            # Find estimate rate column
            if is_estimate and is_rate:
                estimate_rate_col = c
            
            # Skip estimate columns - we only want execution/workslip columns
            if is_estimate:
                continue
            
            # Skip More/Less/Remarks columns
            if "more" in header or "less" in header or "remark" in header:
                continue
                
            if is_exec_col and is_qty:
                qty_columns.append(c)
            elif is_exec_col and is_amt:
                amt_columns.append(c)
        
        # Step 3: Determine the columns to use
        # Use the LAST Qty column (latest phase) and its corresponding Amount column
        if qty_columns:
            qty_col = qty_columns[-1]  # Last qty column = latest workslip phase
            amt_col = amt_columns[-1] if amt_columns else qty_col + 1
            rate_col = estimate_rate_col  # Always use Estimate Rate
        else:
            # Fallback to old format
            qty_col = 7
            rate_col = 8
            amt_col = 9
        
        print(f"DEBUG parse_workslip_items (local): Using columns - qty={qty_col}, rate={rate_col}, amt={amt_col}")

        for r in range(1, max_row + 1):
            desc_raw = ws.cell(row=r, column=2).value  # B
            desc = str(desc_raw or "").strip()
            if desc == "":
                continue

            low = desc.lower()

            if low == "description of item":
                continue
            if low.startswith("sub total") or low.startswith("sub-total") or low.startswith("ecv"):
                break
            if "supplemental items" in low:
                continue

            # Heading row in merged cell -> skip
            if is_merged_cell(ws, r, 2):
                continue

            is_ae = low.startswith("ae")

            qty_raw = ws.cell(row=r, column=qty_col).value
            rate_raw = ws.cell(row=r, column=rate_col).value
            amt_raw = ws.cell(row=r, column=amt_col).value

            qty_exec = to_number(qty_raw)
            rate_exec = to_number(rate_raw)
            amt_exec = to_number(amt_raw)

            if amt_exec != 0:
                if qty_exec == 0 and rate_exec != 0:
                    qty_exec = amt_exec / rate_exec
                elif rate_exec == 0 and qty_exec != 0:
                    rate_exec = amt_exec / qty_exec

            # Pure heading: has text but no qty & no rate
            if qty_exec == 0 and rate_exec == 0:
                continue

            unit = str(ws.cell(row=r, column=3).value or "").strip()  # C

            items.append({
                "qty": qty_exec,
                "unit": unit,
                "desc": desc,
                "rate": rate_exec,
                "is_ae": is_ae,
            })

        return items

    def find_nth_bill_sheet_and_header_row(wb):
        for ws in wb.worksheets:
            for r in range(1, 40):
                a = str(ws.cell(row=r, column=1).value or "").strip().lower()
                c = str(ws.cell(row=r, column=3).value or "").strip().lower()
                if "sl" in a and "quantity till date" in c:
                    return ws, r
        return wb.worksheets[0], 10

    def parse_first_bill_for_nth(ws, header_row):
        """
        First & Part / First & Final Bill (8-column format):
          A: Sl.No, B: Quantity, C: Unit, D: Item, E: Rate, F: Per, G: Unit, H: Amount

        Prev qty  = Column B
        Prev amt  = Column H; if H has no usable value (formula without cache),
                    fallback to B * E.

        Heading rows:
          - Description cell (D) is merged, OR
          - D has text but E is empty (and not 'Total' / 'Sub Total' / 'ECV').

        Stop only on totals rows, never on blanks.
        """
        items = []
        max_row = min(ws.max_row, 5000)

        for r in range(header_row + 1, max_row + 1):
            desc_raw = ws.cell(row=r, column=4).value  # D
            desc = str(desc_raw or "").strip()
            rate_raw = ws.cell(row=r, column=5).value  # E

            desc_low = desc.lower()

            if desc_low.startswith(("ecv", "sub total", "subtotal", "total")):
                break

            rate_str = "" if rate_raw is None else str(rate_raw).strip()
            is_rate_empty = (rate_str == "")

            # Heading row â†’ skip only
            if (
                desc
                and is_rate_empty
                and not desc_low.startswith(("ecv", "sub total", "subtotal", "total"))
            ) or (
                desc
                and is_merged_cell(ws, r, 4)
            ):
                continue

            prev_qty_raw = ws.cell(row=r, column=2).value  # B
            amt_cell_raw = ws.cell(row=r, column=8).value  # H
            unit_raw = ws.cell(row=r, column=3).value      # C

            qty_str = "" if prev_qty_raw is None else str(prev_qty_raw).strip()
            unit_str = "" if unit_raw is None else str(unit_raw).strip()
            amt_str = "" if amt_cell_raw is None else str(amt_cell_raw).strip()

            all_blank = (
                desc == "" and is_rate_empty and qty_str == "" and unit_str == "" and amt_str == ""
            )
            if all_blank:
                continue

            if desc == "" and is_rate_empty:
                continue

            unit = unit_str
            qty_val = to_number(prev_qty_raw)
            rate_val = to_number(rate_raw)
            prev_amount_val = to_number(amt_cell_raw)

            if (amt_cell_raw is None or str(amt_cell_raw).strip() == "") and (qty_val != 0 or rate_val != 0):
                prev_amount_val = qty_val * rate_val

            items.append({
                "desc": desc,
                "unit": unit,
                "rate": rate_val,
                "prev_qty": qty_val,
                "prev_amount": prev_amount_val,
            })

        return items

    def parse_nth_bill_for_next(ws, header_row):
        """
        Nth bill format (supports both 10-column and 11-column with Unit):

        10-column: A: S.No, B: Item, C: Qty, D: Rate, E: Total Value
        11-column: A: S.No, B: Item, C: Qty, D: Unit, E: Rate, F: Total Value
        
        Prev qty  = column C (Quantity Till Date)
        Prev amt  = column E or F (Total Value till date)
        """
        items = []
        max_row = min(ws.max_row, 5000)
        start_row = header_row + 2  # skip row 10 & 11 headings
        
        # Check if this is an 11-column format with Unit column
        has_unit_column = False
        for col in range(3, 6):  # Check columns C, D, E
            hdr_val = str(ws.cell(row=header_row, column=col).value or "").strip().lower()
            if hdr_val == "unit":
                has_unit_column = True
                break

        for r in range(start_row, max_row + 1):
            desc_raw = ws.cell(row=r, column=2).value  # B
            desc = str(desc_raw or "").strip()
            if desc == "":
                continue

            low = desc.lower()
            if low.startswith("sub total") or low.startswith("subtotal"):
                break

            if has_unit_column:
                # 11-column format: A: S.No, B: Item, C: Qty, D: Unit, E: Rate, F: Total Value
                unit_raw = ws.cell(row=r, column=4).value  # D
                rate_raw = ws.cell(row=r, column=5).value  # E
                prev_qty_raw = ws.cell(row=r, column=3).value  # C
                amt_cell_raw = ws.cell(row=r, column=6).value  # F
            else:
                # 10-column format: A: S.No, B: Item, C: Qty, D: Rate, E: Total Value
                unit_raw = None
                rate_raw = ws.cell(row=r, column=4).value  # D
                prev_qty_raw = ws.cell(row=r, column=3).value  # C
                amt_cell_raw = ws.cell(row=r, column=5).value  # E

            unit_str = str(unit_raw or "").strip() if unit_raw else ""
            qty_val = to_number(prev_qty_raw)
            rate_val = to_number(rate_raw)
            prev_amount_val = to_number(amt_cell_raw)

            if (amt_cell_raw is None or str(amt_cell_raw).strip() == "") and (qty_val != 0 or rate_val != 0):
                prev_amount_val = qty_val * rate_val

            items.append({
                "desc": desc,
                "unit": unit_str,
                "rate": rate_val,
                "prev_qty": qty_val,
                "prev_amount": prev_amount_val,
            })

        return items

    def read_tp_from_sheet(ws):
        """
        Generic TP reader: look for a cell (B or D) containing 'T.P'
        starting with 'Add' or 'Deduct'.
        """
        tp_percent = None
        tp_type = None
        max_row = ws.max_row
        for r in range(1, max_row + 1):
            for col in (2, 4):
                val = ws.cell(row=r, column=col).value
                if not val:
                    continue
                s = str(val).strip()
                low = s.lower()
                if "t.p" in low:
                    if low.startswith("add"):
                        tp_type = "Excess"
                    elif low.startswith("deduct"):
                        tp_type = "Less"
                    m = re.search(r"(\d+(\.\d+)?)", low)
                    if m:
                        tp_percent = float(m.group(1))
                    return tp_percent, tp_type
        return None, None

    def singular_unit(plural):
        p = str(plural or "").strip()
        if p.lower().endswith("s") and len(p) > 1:
            return p[:-1]
        return p

    def ordinal_word(n):
        mapping = {
            1: "First",
            2: "Second",
            3: "Third",
            4: "Fourth",
            5: "Fifth",
            6: "Sixth",
            7: "Seventh",
            8: "Eighth",
            9: "Ninth",
            10: "Tenth",
        }
        if n in mapping:
            return mapping[n]
        if 10 < n % 100 < 14:
            suffix = "th"
        else:
            suffix = {1: "st", 2: "nd", 3: "rd"}.get(n % 10, "th")
            if suffix is None:
                suffix = "th"
        return f"{n}{suffix}"

    def create_first_bill_sheet(
        wb_out,
        sheet_name,
        items,
        header_data,
        title_text,
        tp_percent,
        tp_type,
        mb_measure_no, mb_measure_p_from, mb_measure_p_to,
        mb_abs_no, mb_abs_p_from, mb_abs_p_to,
        doi, doc, domr, dobr
    ):
        """
        Create ONE 8-column bill sheet in an existing workbook.
        Same layout as your original First Bill.
        """
        # first sheet: reuse active; others: create new
        if wb_out.worksheets and wb_out.worksheets[0].cell(row=1, column=1).value is None and len(wb_out.worksheets) == 1:
            ws_bill = wb_out.active
            ws_bill.title = sheet_name
        else:
            ws_bill = wb_out.create_sheet(title=sheet_name)

        thin = Side(border_style="thin", color="000000")
        border_all = Border(left=thin, right=thin, top=thin, bottom=thin)
        header_fill = PatternFill("solid", fgColor="FFC8C8C8")
        subtotal_fill = PatternFill("solid", fgColor="FFE6E6E6")

        ws_bill.merge_cells("A1:H1")
        c = ws_bill["A1"]
        c.value = title_text
        c.font = Font(bold=True, size=14)
        c.alignment = Alignment(horizontal="center", vertical="center")

        work_val = header_data.get("name_of_work", "").strip()
        est_val = header_data.get("estimate_amount", "").strip()
        adm_val = header_data.get("admin_sanction", "").strip()
        tech_val = header_data.get("tech_sanction", "").strip()
        agt_val = header_data.get("agreement", "").strip()
        agency_val = header_data.get("agency", "").strip()

        ws_bill.merge_cells("A2:H2")
        c2 = ws_bill["A2"]
        c2.value = f"Name of the work : {work_val}" if work_val else "Name of the work :"
        c2.font = Font(bold=True)
        c2.alignment = Alignment(horizontal="left", vertical="center")

        ws_bill.merge_cells("A3:H3")
        c3 = ws_bill["A3"]
        c3.value = f"Estimate Amount : {est_val}" if est_val else "Estimate Amount :"
        c3.font = Font(bold=True)
        c3.alignment = Alignment(horizontal="left", vertical="center")

        ws_bill.merge_cells("A4:H4")
        c4 = ws_bill["A4"]
        c4.value = (
            f"Ref. to Administrative sanction : {adm_val}"
            if adm_val else "Ref. to Administrative sanction :"
        )
        c4.font = Font(bold=True)
        c4.alignment = Alignment(horizontal="left", vertical="center")

        ws_bill.merge_cells("A5:H5")
        c5 = ws_bill["A5"]
        c5.value = (
            f"Ref. to Technical sanction : {tech_val}"
            if tech_val else "Ref. to Technical sanction :"
        )
        c5.font = Font(bold=True)
        c5.alignment = Alignment(horizontal="left", vertical="center")

        ws_bill.merge_cells("A6:H6")
        c6 = ws_bill["A6"]
        c6.value = (
            f"Ref. to Agreement : {agt_val}"
            if agt_val else "Ref. to Agreement :"
        )
        c6.font = Font(bold=True)
        c6.alignment = Alignment(horizontal="left", vertical="center")

        ws_bill.merge_cells("A7:H7")
        c7 = ws_bill["A7"]
        c7.value = (
            f"Name of the Agency : {agency_val}"
            if agency_val else "Name of the Agency :"
        )
        c7.font = Font(bold=True)
        c7.alignment = Alignment(horizontal="left", vertical="center")

        # -------- ROW 8: Label + MB details (FIRST BILL) IN ONE MERGED CELL --------
        ws_bill.merge_cells("A8:H8")
        c8 = ws_bill["A8"]
        c8.value = (
            f"M.B.No Details: MB.No. {mb_measure_no} P.No. {mb_measure_p_from} to {mb_measure_p_to} (Measurements)   "
            f"&   MB.No. {mb_abs_no} P.No. {mb_abs_p_from} to {mb_abs_p_to} (Abstract)"
        )
        c8.font = Font(bold=True)
        c8.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)

        ws_bill.merge_cells("A9:H9")
        c9 = ws_bill["A9"]
        c9.value = f"DOI : {doi}    DOC : {doc}    DOMR : {domr}    DOBR : {dobr}"
        c9.font = Font(bold=True)
        c9.alignment = Alignment(horizontal="left", vertical="center")

        for r in range(1, 10):
            for c_idx in range(1, 9):
                cell = ws_bill.cell(row=r, column=c_idx)
                cell.border = border_all

        header_row = 10
        headers = ["S.No", "Quantity", "Unit", "Item", "Rate", "Per", "Unit", "Amount"]
        for col_idx, text in enumerate(headers, start=1):
            cell = ws_bill.cell(row=header_row, column=col_idx, value=text)
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = border_all
            cell.fill = header_fill

        ws_bill.column_dimensions["A"].width = 6
        ws_bill.column_dimensions["B"].width = 10
        ws_bill.column_dimensions["C"].width = 10
        ws_bill.column_dimensions["D"].width = 45
        ws_bill.column_dimensions["E"].width = 10
        ws_bill.column_dimensions["F"].width = 6
        ws_bill.column_dimensions["G"].width = 10
        ws_bill.column_dimensions["H"].width = 15

        data_start = header_row + 1
        row_idx = data_start
        slno = 1

        for it in items:
            qty = it.get("qty", 0)
            unit_pl = str(it.get("unit") or "").strip()
            desc = it.get("desc") or ""
            rate = it.get("rate", 0.0)
            is_ae = bool(it.get("is_ae", False))

            if is_ae:
                ws_bill.cell(row=row_idx, column=1, value=None)
            else:
                ws_bill.cell(row=row_idx, column=1, value=slno)

            ws_bill.cell(row=row_idx, column=2, value=qty)
            ws_bill.cell(row=row_idx, column=3, value=unit_pl)
            ws_bill.cell(row=row_idx, column=4, value=desc)
            ws_bill.cell(row=row_idx, column=5, value=rate)
            ws_bill.cell(row=row_idx, column=6, value=1)
            ws_bill.cell(row=row_idx, column=7, value=singular_unit(unit_pl))
            ws_bill.cell(row=row_idx, column=8, value=f"=B{row_idx}*E{row_idx}")

            for c_idx in range(1, 9):
                cell = ws_bill.cell(row=row_idx, column=c_idx)
                cell.border = border_all
                if c_idx == 4:
                    cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
                else:
                    cell.alignment = Alignment(horizontal="center", vertical="center")

            if not is_ae:
                slno += 1

            row_idx += 1

        last_item_row = row_idx - 1

        sub_row = row_idx
        ws_bill.cell(row=sub_row, column=4, value="Sub Total Amount")
        ws_bill.cell(row=sub_row, column=8, value=f"=SUM(H{data_start}:H{last_item_row})")

        for c_idx in range(1, 9):
            cell = ws_bill.cell(row=sub_row, column=c_idx)
            cell.font = Font(bold=True)
            cell.border = border_all
            cell.fill = subtotal_fill
            if c_idx == 4:
                cell.alignment = Alignment(horizontal="left", vertical="center")
            else:
                cell.alignment = Alignment(horizontal="center", vertical="center")

        tp_row = sub_row + 1
        tp_percent = float(tp_percent or 0.0)
        tp_type = tp_type if tp_type in ("Less", "Excess") else "Excess"

        label_tp = f"Add/ Deduct T.P @ {tp_percent} % {tp_type}"
        ws_bill.cell(row=tp_row, column=4, value=label_tp)
        ws_bill.cell(row=tp_row, column=8, value=f"=H{sub_row}*{abs(tp_percent)}/100")

        for c_idx in range(1, 9):
            cell = ws_bill.cell(row=tp_row, column=c_idx)
            cell.font = Font(bold=True)
            cell.border = border_all
            cell.fill = subtotal_fill
            if c_idx == 4:
                cell.alignment = Alignment(horizontal="left", vertical="center")
            else:
                cell.alignment = Alignment(horizontal="center", vertical="center")

        total_row = tp_row + 1
        ws_bill.cell(row=total_row, column=4, value="Total")

        if tp_type == "Less":
            ws_bill.cell(row=total_row, column=8, value=f"=H{sub_row}-H{tp_row}")
        else:
            ws_bill.cell(row=total_row, column=8, value=f"=H{sub_row}+H{tp_row}")

        for c_idx in range(1, 9):
            cell = ws_bill.cell(row=total_row, column=c_idx)
            cell.font = Font(bold=True)
            cell.border = border_all
            cell.fill = subtotal_fill
            if c_idx == 4:
                cell.alignment = Alignment(horizontal="left", vertical="center")
            else:
                cell.alignment = Alignment(horizontal="center", vertical="center")

        # let Excel auto-fit default heights
        for r in range(1, ws_bill.max_row + 1):
            ws_bill.row_dimensions[r].height = None

    def build_first_bill_wb(items, header_data, title_text,
                            tp_percent, tp_type,
                            mb_measure_no, mb_measure_p_from, mb_measure_p_to,
                            mb_abs_no, mb_abs_p_from, mb_abs_p_to,
                            doi, doc, domr, dobr):
        """
        Backward-compatible wrapper: build a single-sheet bill workbook.
        """
        wb_out = Workbook()
        create_first_bill_sheet(
            wb_out,
            sheet_name="Bill",
            items=items,
            header_data=header_data,
            title_text=title_text,
            tp_percent=tp_percent,
            tp_type=tp_type,
            mb_measure_no=mb_measure_no,
            mb_measure_p_from=mb_measure_p_from,
            mb_measure_p_to=mb_measure_p_to,
            mb_abs_no=mb_abs_no,
            mb_abs_p_from=mb_abs_p_from,
            mb_abs_p_to=mb_abs_p_to,
            doi=doi,
            doc=doc,
            domr=domr,
            dobr=dobr,
        )
        return wb_out

    def build_nth_bill_wb(items, header_data, title_text,
                          tp_percent, tp_type,
                          mb_measure_no, mb_measure_p_from, mb_measure_p_to,
                          mb_abs_no, mb_abs_p_from, mb_abs_p_to,
                          doi, doc, domr, dobr):
        """
        Same as your earlier Nth bill format, single sheet only.
        Uses 11-column format with Unit column.
        """
        wb_out = Workbook()
        ws = wb_out.active
        ws.title = "Bill"

        thin = Side(border_style="thin", color="000000")
        border_all = Border(left=thin, right=thin, top=thin, bottom=thin)
        header_fill = PatternFill("solid", fgColor="FFC8C8C8")
        subtotal_fill = PatternFill("solid", fgColor="FFE6E6E6")

        ws.merge_cells("A1:K1")
        c1 = ws["A1"]
        c1.value = title_text
        c1.font = Font(bold=True, size=14)
        c1.alignment = Alignment(horizontal="center", vertical="center")

        work_val = header_data.get("name_of_work", "").strip()
        est_val = header_data.get("estimate_amount", "").strip()
        adm_val = header_data.get("admin_sanction", "").strip()
        tech_val = header_data.get("tech_sanction", "").strip()
        agt_val = header_data.get("agreement", "").strip()
        agency_val = header_data.get("agency", "").strip()

        ws.merge_cells("A2:K2")
        c2 = ws["A2"]
        c2.value = f"Name of the work : {work_val}" if work_val else "Name of the work :"
        c2.font = Font(bold=True)
        c2.alignment = Alignment(horizontal="left", vertical="center")

        ws.merge_cells("A3:K3")
        c3 = ws["A3"]
        c3.value = f"Estimate Amount : {est_val}" if est_val else "Estimate Amount :"
        c3.font = Font(bold=True)
        c3.alignment = Alignment(horizontal="left", vertical="center")

        ws.merge_cells("A4:K4")
        c4 = ws["A4"]
        c4.value = (
            f"Ref. to Administrative sanction : {adm_val}"
            if adm_val else "Ref. to Administrative sanction :"
        )
        c4.font = Font(bold=True)
        c4.alignment = Alignment(horizontal="left", vertical="center")

        ws.merge_cells("A5:K5")
        c5 = ws["A5"]
        c5.value = (
            f"Ref. to Technical sanction : {tech_val}"
            if tech_val else "Ref. to Technical sanction :"
        )
        c5.font = Font(bold=True)
        c5.alignment = Alignment(horizontal="left", vertical="center")

        ws.merge_cells("A6:K6")
        c6 = ws["A6"]
        c6.value = (
            f"Ref. to Agreement : {agt_val}"
            if agt_val else "Ref. to Agreement :"
        )
        c6.font = Font(bold=True)
        c6.alignment = Alignment(horizontal="left", vertical="center")

        ws.merge_cells("A7:K7")
        c7 = ws["A7"]
        c7.value = (
            f"Name of the Agency : {agency_val}"
            if agency_val else "Name of the Agency :"
        )
        c7.font = Font(bold=True)
        c7.alignment = Alignment(horizontal="left", vertical="center")

        # -------- ROW 8: Label + MB details (NTH BILL) IN ONE MERGED CELL --------
        ws.merge_cells("A8:K8")
        c8 = ws["A8"]
        c8.value = (
            f"M.B.No Details: MB.No. {mb_measure_no} P.No. {mb_measure_p_from} to {mb_measure_p_to} (Measurements)   "
            f"&   MB.No. {mb_abs_no} P.No. {mb_abs_p_from} to {mb_abs_p_to} (Abstract)"
        )
        c8.font = Font(bold=True)
        c8.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)

        ws.merge_cells("A9:K9")
        c9 = ws["A9"]
        c9.value = f"DOI : {doi}    DOC : {doc}    DOMR : {domr}    DOBR : {dobr}"
        c9.font = Font(bold=True)
        c9.alignment = Alignment(horizontal="left", vertical="center")

        for r in range(1, 10):
            for col in range(1, 12):
                cell = ws.cell(row=r, column=col)
                cell.border = border_all

        # Merge header cells: S.No, Item, Qty Till Date, Unit, Rate, Total Value, Remarks span 2 rows
        for col in [1, 2, 3, 4, 5, 6, 11]:
            ws.merge_cells(start_row=10, start_column=col, end_row=11, end_column=col)

        ws.merge_cells("G10:H10")  # Deduct Previous
        ws.merge_cells("I10:J10")  # Since Last

        ws.cell(row=10, column=1, value="S.No")
        ws.cell(row=10, column=2, value="Item")
        ws.cell(row=10, column=3, value="Quantity Till Date")
        ws.cell(row=10, column=4, value="Unit")
        ws.cell(row=10, column=5, value="Rate per Unit")
        ws.cell(row=10, column=6, value="Total Value till date")
        ws.cell(row=10, column=7, value="Deduct Previous Measurements")
        ws.cell(row=10, column=9, value="Since Last Measurements")
        ws.cell(row=10, column=11, value="Remarks")

        ws.cell(row=11, column=7, value="Quantity")
        ws.cell(row=11, column=8, value="Amount")
        ws.cell(row=11, column=9, value="Quantity")
        ws.cell(row=11, column=10, value="Amount")

        for r in (10, 11):
            for col in range(1, 12):
                cell = ws.cell(row=r, column=col)
                cell.font = Font(bold=True)
                cell.border = border_all
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

        ws.column_dimensions["A"].width = 6
        ws.column_dimensions["B"].width = 45
        ws.column_dimensions["C"].width = 14
        ws.column_dimensions["D"].width = 8
        ws.column_dimensions["E"].width = 12
        ws.column_dimensions["F"].width = 16
        ws.column_dimensions["G"].width = 14
        ws.column_dimensions["H"].width = 16
        ws.column_dimensions["I"].width = 14
        ws.column_dimensions["J"].width = 16
        ws.column_dimensions["K"].width = 20

        data_start = 12
        r = data_start
        sl = 1

        for it in items:
            desc = it.get("desc") or ""
            unit = it.get("unit") or ""
            rate = it.get("rate", 0.0)
            prev_qty = it.get("prev_qty", 0.0)
            prev_amount = it.get("prev_amount", 0.0)

            ws.cell(row=r, column=1, value=sl)
            ws.cell(row=r, column=2, value=desc)
            ws.cell(row=r, column=3, value=None)  # Quantity Till Date (to be filled)
            ws.cell(row=r, column=4, value=unit)
            ws.cell(row=r, column=5, value=rate)
            ws.cell(row=r, column=6, value=f"=C{r}*E{r}")
            ws.cell(row=r, column=7, value=prev_qty)
            ws.cell(row=r, column=8, value=prev_amount)
            ws.cell(row=r, column=9, value=f"=C{r}-G{r}")
            ws.cell(row=r, column=10, value=f"=F{r}-H{r}")
            ws.cell(row=r, column=11, value="")

            for col in range(1, 12):
                cell = ws.cell(row=r, column=col)
                cell.border = border_all
                if col == 2:
                    cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
                else:
                    cell.alignment = Alignment(horizontal="center", vertical="center")

            r += 1
            sl += 1

        last_item_row = r - 1

        sub_row = r
        ws.cell(row=sub_row, column=2, value="Sub Total")
        ws.cell(row=sub_row, column=6, value=f"=SUM(F{data_start}:F{last_item_row})")
        ws.cell(row=sub_row, column=8, value=f"=SUM(H{data_start}:H{last_item_row})")
        ws.cell(row=sub_row, column=10, value=f"=SUM(J{data_start}:J{last_item_row})")

        tp_row = sub_row + 1
        tp_percent = float(tp_percent or 0.0)
        tp_type = tp_type if tp_type in ("Less", "Excess") else "Excess"

        label_prefix = "Deduct" if tp_type == "Less" else "Add"
        ws.cell(row=tp_row, column=2, value=f"{label_prefix} T.P @ {tp_percent} % {tp_type}")
        ws.cell(row=tp_row, column=6, value=f"=F{sub_row}*{abs(tp_percent)}/100")
        ws.cell(row=tp_row, column=8, value=f"=H{sub_row}*{abs(tp_percent)}/100")
        ws.cell(row=tp_row, column=10, value=f"=J{sub_row}*{abs(tp_percent)}/100")

        total_row = tp_row + 1
        ws.cell(row=total_row, column=2, value="Total")

        if tp_type == "Less":
            ws.cell(row=total_row, column=6, value=f"=F{sub_row}-F{tp_row}")
            ws.cell(row=total_row, column=8, value=f"=H{sub_row}-H{tp_row}")
            ws.cell(row=total_row, column=10, value=f"=J{sub_row}-J{tp_row}")
        else:
            ws.cell(row=total_row, column=6, value=f"=F{sub_row}+F{tp_row}")
            ws.cell(row=total_row, column=8, value=f"=H{sub_row}+H{tp_row}")
            ws.cell(row=total_row, column=10, value=f"=J{sub_row}+J{tp_row}")

        for rr in [sub_row, tp_row, total_row]:
            for col in range(1, 12):
                cell = ws.cell(row=rr, column=col)
                cell.font = Font(bold=True)
                cell.border = border_all
                cell.fill = subtotal_fill
                if col == 2:
                    cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
                else:
                    cell.alignment = Alignment(horizontal="center", vertical="center")

        for rr in range(1, ws.max_row + 1):
            ws.row_dimensions[rr].height = None

        return wb_out

    # ---------- main bill() logic ----------
    print('DEBUG: entering main bill logic')

    if request.method == "GET":
        print("DEBUG: bill handling GET -> rendering bill.html")
        resp = render(request, "core/bill.html")
        print("DEBUG: render returned type", type(resp))
        return resp

    action = request.POST.get("action", "").strip()
    uploaded = request.FILES.get("bill_file")

    if not uploaded:
        return render(request, "core/bill.html", {
            "error": "Please upload an Excel file.",
        })

    mb_measure_no = request.POST.get("mb_measure_no", "").strip()
    mb_measure_p_from = request.POST.get("mb_measure_p_from", "").strip()
    mb_measure_p_to = request.POST.get("mb_measure_p_to", "").strip()
    mb_abs_no = request.POST.get("mb_abstract_no", "").strip()
    mb_abs_p_from = request.POST.get("mb_abstract_p_from", "").strip()
    mb_abs_p_to = request.POST.get("mb_abstract_p_to", "").strip()
    doi = request.POST.get("doi", "").strip()
    doc = request.POST.get("doc", "").strip()
    domr = request.POST.get("domr", "").strip()
    dobr = request.POST.get("dobr", "").strip()

    tp_percent_session = to_number(request.session.get("ws_tp_percent", 0.0))
    tp_type_session = request.session.get("ws_tp_type", "Excess")

    try:
        wb_in = load_workbook(uploaded, data_only=True)
        print('DEBUG: loaded workbook, sheets=', [s.title for s in wb_in.worksheets])
    except Exception as e:
        return render(request, "core/bill.html", {
            "error": f"Error reading uploaded Excel: {e}",
        })

    print('DEBUG: action=', action)
    # SECTION 1  -  Bill from Estimate (multi-sheet)
    if action in ("estimate_first_part", "estimate_first_final"):
        sheet_info_list = find_all_estimate_sheets_and_header_rows(wb_in)

        # Build one output workbook with one Bill sheet per estimate sheet
        wb_out = Workbook()

        any_items = False
        for idx, (ws_est, header_row) in enumerate(sheet_info_list, start=1):
            header_data = extract_header_data(ws_est)
            items = parse_estimate_items(ws_est, header_row)
            if not items:
                continue

            any_items = True
            title_text = "CC First & Part Bill" if action == "estimate_first_part" else "CC First & Final Bill"
            tp_percent = tp_percent_session
            tp_type = tp_type_session

            if len(sheet_info_list) == 1:
                sheet_name = "Bill"
            else:
                # Name by sheet: 'Bill_<SheetName>' (trim to Excel limit)
                base = ws_est.title.strip() or f"Sheet{idx}"
                sheet_name = f"Bill_{base}"
                if len(sheet_name) > 31:
                    sheet_name = sheet_name[:31]

            create_first_bill_sheet(
                wb_out,
                sheet_name=sheet_name,
                items=items,
                header_data=header_data,
                title_text=title_text,
                tp_percent=tp_percent,
                tp_type=tp_type,
                mb_measure_no=mb_measure_no,
                mb_measure_p_from=mb_measure_p_from,
                mb_measure_p_to=mb_measure_p_to,
                mb_abs_no=mb_abs_no,
                mb_abs_p_from=mb_abs_p_from,
                mb_abs_p_to=mb_abs_p_to,
                doi=doi,
                doc=doc,
                domr=domr,
                dobr=dobr,
            )

        if not any_items:
            return render(request, "core/bill.html", {
                "error": "Could not detect any items in the uploaded Estimate (all sheets).",
            })

        resp = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        resp["Content-Disposition"] = 'attachment; filename="Bill_from_Estimate.xlsx"'
        wb_out.save(resp)
        return resp

    # SECTION 2  -  Bill from WorkSlip (multi-sheet)
    if action in ("workslip_first_part", "workslip_first_final"):
        ws_list = find_all_workslip_sheets(wb_in)

        wb_out = Workbook()
        any_items = False

        for idx, ws_ws in enumerate(ws_list, start=1):
            header_data = extract_header_data(ws_ws)
            items = parse_workslip_items(ws_ws)
            if not items:
                continue

            any_items = True
            title_text = "CC First & Part Bill" if action == "workslip_first_part" else "CC First & Final Bill"

            tp_percent, tp_type = read_tp_from_sheet(ws_ws)
            if tp_percent is None:
                tp_percent = tp_percent_session
            if tp_type is None:
                tp_type = tp_type_session

            if len(ws_list) == 1:
                sheet_name = "Bill"
            else:
                base = ws_ws.title.strip() or f"WS{idx}"
                sheet_name = f"Bill_{base}"
                if len(sheet_name) > 31:
                    sheet_name = sheet_name[:31]

            create_first_bill_sheet(
                wb_out,
                sheet_name=sheet_name,
                items=items,
                header_data=header_data,
                title_text=title_text,
                tp_percent=tp_percent,
                tp_type=tp_type,
                mb_measure_no=mb_measure_no,
                mb_measure_p_from=mb_measure_p_from,
                mb_measure_p_to=mb_measure_p_to,
                mb_abs_no=mb_abs_no,
                mb_abs_p_from=mb_abs_p_from,
                mb_abs_p_to=mb_abs_p_to,
                doi=doi,
                doc=doc,
                domr=domr,
                dobr=dobr,
            )

        if not any_items:
            return render(request, "core/bill.html", {
                "error": "Could not detect any executed items in the uploaded WorkSlip (all sheets).",
            })

        resp = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        resp["Content-Disposition"] = 'attachment; filename="Bill_from_WorkSlip.xlsx"'
        wb_out.save(resp)
        return resp

    # SECTION 3  -  Nth & Part / 2nd & Final from First & Part Bill
    if action in ("firstpart_nth_part", "firstpart_2nd_final"):
        ws_first, header_row = find_estimate_sheet_and_header_row(wb_in)
        header_data = extract_header_data(ws_first)
        items = parse_first_bill_for_nth(ws_first, header_row)

        if not items:
            return render(request, "core/bill.html", {
                "error": "Could not detect any items in the uploaded First & Part / First & Final Bill.",
            })

        if action == "firstpart_nth_part":
            nth_str = request.POST.get("nth_number", "").strip()
            try:
                nth_val = int(nth_str)
            except Exception:
                nth_val = 2
            if nth_val < 2:
                nth_val = 2
            ord_word = ordinal_word(nth_val)
            title_text = f"CC {ord_word} & Part Bill"
        else:
            ord_word = ordinal_word(2)
            title_text = f"CC {ord_word} & Final Bill"

        tp_percent, tp_type = read_tp_from_sheet(ws_first)
        if tp_percent is None:
            tp_percent = tp_percent_session
        if tp_type is None:
            tp_type = tp_type_session

        wb_out = build_nth_bill_wb(
            items=items,
            header_data=header_data,
            title_text=title_text,
            tp_percent=tp_percent,
            tp_type=tp_type,
            mb_measure_no=mb_measure_no,
            mb_measure_p_from=mb_measure_p_from,
            mb_measure_p_to=mb_measure_p_to,
            mb_abs_no=mb_abs_no,
            mb_abs_p_from=mb_abs_p_from,
            mb_abs_p_to=mb_abs_p_to,
            doi=doi,
            doc=doc,
            domr=domr,
            dobr=dobr,
        )

        resp = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        filename = "Nth_Bill_from_FirstPart.xlsx" if action == "firstpart_nth_part" else "Second_Final_from_FirstPart.xlsx"
        resp["Content-Disposition"] = f'attachment; filename="{filename}"'
        wb_out.save(resp)
        return resp

    # SECTION 4  -  Nth & Part / Nth & Final from Nth & Part Bill
    if action in ("nth_nth_part", "nth_nth_final"):
        ws_nth, header_row = find_nth_bill_sheet_and_header_row(wb_in)
        header_data = extract_header_data(ws_nth)
        items = parse_nth_bill_for_next(ws_nth, header_row)

        if not items:
            return render(request, "core/bill.html", {
                "error": "Could not detect any items in the uploaded Nth & Part Bill.",
            })

        nth_str = request.POST.get("nth_number", "").strip()
        try:
            nth_val = int(nth_str)
        except Exception:
            nth_val = 2
        if nth_val < 2:
            nth_val = 2
        ord_word = ordinal_word(nth_val)

        if action == "nth_nth_part":
            title_text = f"CC {ord_word} & Part Bill"
        else:
            title_text = f"CC {ord_word} & Final Bill"

        tp_percent, tp_type = read_tp_from_sheet(ws_nth)
        if tp_percent is None:
            tp_percent = tp_percent_session
        if tp_type is None:
            tp_type = tp_type_session

        wb_out = build_nth_bill_wb(
            items=items,
            header_data=header_data,
            title_text=title_text,
            tp_percent=tp_percent,
            tp_type=tp_type,
            mb_measure_no=mb_measure_no,
            mb_measure_p_from=mb_measure_p_from,
            mb_measure_p_to=mb_measure_p_to,
            mb_abs_no=mb_abs_no,
            mb_abs_p_from=mb_abs_p_from,
            mb_abs_p_to=mb_abs_p_to,
            doi=doi,
            doc=doc,
            domr=domr,
            dobr=dobr,
        )

        resp = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        filename = "Nth_Bill_from_NthPart.xlsx" if action == "nth_nth_part" else "Nth_Final_from_NthPart.xlsx"
        resp["Content-Disposition"] = f'attachment; filename="{filename}"'
        wb_out.save(resp)
        return resp

    return render(request, "core/bill.html", {
        "error": "This action is not implemented or is invalid.",
    })





# ---------- Helpers only for amount calculation for Covering/Movement ----------

def _to_number_amt(v):
    try:
        return float(v)
    except Exception:
        return 0.0


def _is_merged_cell_amt(ws, row, col):
    """
    Local helper for amount extraction:
    True if (row, col) lies in any merged range on this sheet.
    """
    for merged in ws.merged_cells.ranges:
        if merged.min_row <= row <= merged.max_row and merged.min_col <= col <= merged.max_col:
            return True
    return False


def _find_estimate_sheet_header_for_amount(wb):
    """
    Same header logic as your find_estimate_sheet_and_header_row:
      A: 'Sl', B: 'Quantity', D: 'Item'/'Description'
    """
    for ws in wb.worksheets:
        for r in range(1, 26):
            a = str(ws.cell(row=r, column=1).value or "").strip().lower()
            b = str(ws.cell(row=r, column=2).value or "").strip().lower()
            d = str(ws.cell(row=r, column=4).value or "").strip().lower()
            if "sl" in a and "quantity" in b and ("item" in d or "description" in d):
                return ws, r
    return wb.worksheets[0], 3


def _parse_estimate_items_for_amount(ws, header_row):
    """
    Mirror of your parse_estimate_items(), but only returns qty & rate for
    computing subtotal in Python.

    Format:
      A: Sl.No, B: Quantity, C: Unit, D: Item, E: Rate, F: Per, G: Unit, H: Amount
    """
    items = []
    max_row = min(ws.max_row, 5000)

    for r in range(header_row + 1, max_row + 1):
        desc_raw = ws.cell(row=r, column=4).value  # D
        desc = str(desc_raw or "").strip()
        rate_raw = ws.cell(row=r, column=5).value  # E
        amt_raw = ws.cell(row=r, column=8).value   # H

        desc_low = desc.lower()

        # STOP on ECV / Sub Total / Total
        if desc_low.startswith(("ecv", "sub total", "subtotal", "total")):
            break

        rate_str = "" if rate_raw is None else str(rate_raw).strip()
        is_rate_empty = (rate_str == "")

        # Heading row: description present, but rate empty OR description cell merged
        if desc and (is_rate_empty or _is_merged_cell_amt(ws, r, 4)):
            continue

        qty_raw = ws.cell(row=r, column=2).value  # B
        unit_raw = ws.cell(row=r, column=3).value  # C

        qty_str = "" if qty_raw is None else str(qty_raw).strip()
        unit_str = "" if unit_raw is None else str(unit_raw).strip()
        amt_str = "" if amt_raw is None else str(amt_raw).strip()

        all_blank = (
            desc == "" and
            is_rate_empty and
            qty_str == "" and
            unit_str == "" and
            amt_str == ""
        )
        if all_blank:
            continue

        if desc == "" and is_rate_empty:
            continue

        qty = _to_number_amt(qty_raw)
        rate = _to_number_amt(rate_raw)
        amt = _to_number_amt(amt_raw)

        # if only Amount given, but qty/rate missing, use that
        if amt != 0:
            if qty == 0 and rate != 0:
                qty = amt / rate
            elif rate == 0 and qty != 0:
                rate = amt / qty

        items.append({
            "qty": qty,
            "rate": rate,
        })

    return items


def _find_workslip_sheet_for_amount(wb):
    """
    Same idea as your find_workslip_sheet.
    """
    for ws in wb.worksheets:
        for r in range(1, 40):
            b = str(ws.cell(row=r, column=2).value or "").strip().lower()
            g = str(ws.cell(row=r, column=7).value or "").strip().lower()
            if "description" in b and ("qty" in g or "quantity" in g):
                return ws
    return wb.worksheets[0]


def _parse_workslip_items_for_amount(ws):
    """
    Mirror of parse_workslip_items(), but only returns qty & rate for subtotal.
    Automatically detects and uses the LAST workslip phase columns.
    Rate is always from Estimate Rate column for multi-phase workslips.
    """
    items = []
    max_row = min(ws.max_row, 5000)
    
    # Step 1: Find header row
    header_row = 8
    for r in range(1, 15):
        cell_val = str(ws.cell(row=r, column=1).value or "").strip().lower()
        if "sl" in cell_val:
            header_row = r
            break
    
    # Step 2: Scan header row to find all workslip/execution qty columns
    qty_columns = []
    amt_columns = []
    estimate_rate_col = 5  # Default - Rate (Estimate) column
    
    for c in range(1, 40):  # Scan more columns for multi-phase workslips
        header = str(ws.cell(row=header_row, column=c).value or "").strip().lower()
        
        is_exec_col = ("execution" in header or "exec" in header or "workslip" in header)
        is_qty = ("qty" in header or "quantity" in header)
        is_amt = ("amount" in header or "amt" in header)
        is_estimate = ("est" in header or "estimate" in header)
        is_rate = ("rate" in header)
        
        # Find estimate rate column
        if is_estimate and is_rate:
            estimate_rate_col = c
        
        if is_estimate:
            continue
        
        # Skip More/Less/Remarks columns
        if "more" in header or "less" in header or "remark" in header:
            continue
            
        if is_exec_col and is_qty:
            qty_columns.append(c)
        elif is_exec_col and is_amt:
            amt_columns.append(c)
    
    # Step 3: Use the LAST workslip columns (latest phase)
    if qty_columns:
        qty_col = qty_columns[-1]  # Last qty column = latest workslip phase
        amt_col = amt_columns[-1] if amt_columns else qty_col + 1
        rate_col = estimate_rate_col  # Always use Estimate Rate
    else:
        # Fallback to old format
        qty_col = 7
        rate_col = 8
        amt_col = 9

    for r in range(1, max_row + 1):
        desc_raw = ws.cell(row=r, column=2).value  # B
        desc = str(desc_raw or "").strip()
        if desc == "":
            continue

        low = desc.lower()

        if low == "description of item":
            continue
        if low.startswith("sub total") or low.startswith("sub-total") or low.startswith("ecv"):
            break
        if "supplemental items" in low:
            continue

        if _is_merged_cell_amt(ws, r, 2):
            continue

        qty_raw = ws.cell(row=r, column=qty_col).value
        rate_raw = ws.cell(row=r, column=rate_col).value
        amt_raw = ws.cell(row=r, column=amt_col).value

        qty_exec = _to_number_amt(qty_raw)
        rate_exec = _to_number_amt(rate_raw)
        amt_exec = _to_number_amt(amt_raw)

        if amt_exec != 0:
            if qty_exec == 0 and rate_exec != 0:
                qty_exec = amt_exec / rate_exec
            elif rate_exec == 0 and qty_exec != 0:
                rate_exec = amt_exec / qty_exec

        # Pure heading: text but no qty, no rate, no amount
        if qty_exec == 0 and rate_exec == 0 and amt_exec == 0:
            continue

        items.append({
            "qty": qty_exec,
            "rate": rate_exec,
        })

    return items


def _read_tp_from_sheet_for_amount(ws):
    """
    Same T.P reader as your read_tp_from_sheet().
    Looks in B or D for 'T.P', starting with 'Add' or 'Deduct'.
    """
    tp_percent = None
    tp_type = None
    max_row = ws.max_row
    for r in range(1, max_row + 1):
        for col in (2, 4):
            val = ws.cell(row=r, column=col).value
            if not val:
                continue
            s = str(val).strip()
            low = s.lower()
            if "t.p" in low:
                if low.startswith("add"):
                    tp_type = "Excess"
                elif low.startswith("deduct"):
                    tp_type = "Less"
                m = re.search(r"(\d+(\.\d+)?)", low)
                if m:
                    tp_percent = float(m.group(1))
                return tp_percent, tp_type
    return None, None


def _extract_total_amount_for_action(wb, action: str, request=None) -> float:
    """
    Compute the TOTAL amount for the bill that will be generated, using the
    SAME logic as your bill() function.

    - For estimate_first_part / estimate_first_final:
        * Treat uploaded file as Estimate
        * Find ALL estimate-like sheets (like bill() does)
        * For each sheet:
              subtotal_sheet = sum(Qty * Rate) over parsed items
              TP_sheet = subtotal_sheet * abs(tp_percent)/100
              total_sheet = subtotal_sheet +/- TP_sheet
        * Return sum of total_sheet over all such sheets.
          (This matches multi-sheet Bill_from_Estimate.xlsx.)

    - For workslip_first_part / workslip_first_final:
        * Treat uploaded file as WorkSlip
        * Find ALL workslip-like sheets (like bill() does)
        * For each sheet:
              subtotal_sheet = sum(Qty(Exec) * Rate(Exec)) over parsed items
              TP_sheet from sheet T.P row or session fallback
              total_sheet = subtotal_sheet +/- TP_sheet
        * Return sum of total_sheet over all such sheets.
          (Matches multi-sheet Bill_from_WorkSlip.xlsx.)

    - For all other actions (Nth & Part / Nth & Final):
        * We assume the uploaded Excel is already a Bill
        * Use _extract_total_amount_from_bill_wb(wb)
          which:
              - for 8-col bills reads Column H (Total row)
              - for 10-col Nth bills reads Column I (Total row)
    """
    action = (action or "").strip()

    # ---------- Small local helpers ----------

    def to_num(v):
        try:
            return float(v)
        except Exception:
            return 0.0

    def is_merged(ws, row, col):
        for merged in ws.merged_cells.ranges:
            if merged.min_row <= row <= merged.max_row and merged.min_col <= col <= merged.max_col:
                return True
        return False

    # --- ESTIMATE HELPERS (copied from your bill() logic, but only for amount) ---

    def looks_like_estimate_header(ws, r):
        a = str(ws.cell(row=r, column=1).value or "").strip().lower()
        b = str(ws.cell(row=r, column=2).value or "").strip().lower()
        c = str(ws.cell(row=r, column=3).value or "").strip().lower()
        d = str(ws.cell(row=r, column=4).value or "").strip().lower()
        e = str(ws.cell(row=r, column=5).value or "").strip().lower()

        has_sl = ("sl" in a) or ("s.no" in a) or ("serial" in a)
        has_qty = ("qty" in b or "quantity" in b or "qty" in c or "quantity" in c)
        has_desc = (
            "item" in c or "item" in d or "item" in e or
            "description" in c or "description" in d or "description" in e
        )
        return has_sl and has_qty and has_desc

    def find_all_estimate_sheets_amt(wb_inner):
        results = []
        for ws in wb_inner.worksheets:
            header_row = None
            max_scan = min(ws.max_row, 60)
            for r in range(1, max_scan + 1):
                if looks_like_estimate_header(ws, r):
                    header_row = r
                    break
            if header_row:
                results.append((ws, header_row))

        if not results:
            # fallback similar to find_estimate_sheet_and_header_row
            for ws in wb_inner.worksheets:
                for r in range(1, 26):
                    a = str(ws.cell(row=r, column=1).value or "").strip().lower()
                    b = str(ws.cell(row=r, column=2).value or "").strip().lower()
                    d = str(ws.cell(row=r, column=4).value or "").strip().lower()
                    if "sl" in a and "quantity" in b and ("item" in d or "description" in d):
                        return [(ws, r)]
            # ultimate fallback
            return [(wb_inner.worksheets[0], 3)]
        return results

    def parse_estimate_items_amt(ws, header_row):
        items = []
        max_row = min(ws.max_row, 5000)

        for r in range(header_row + 1, max_row + 1):
            desc_raw = ws.cell(row=r, column=4).value  # D
            desc = str(desc_raw or "").strip()
            rate_raw = ws.cell(row=r, column=5).value  # E
            amt_raw = ws.cell(row=r, column=8).value   # H

            desc_low = desc.lower()

            # STOP on ECV / Sub Total / Total
            if desc_low.startswith(("ecv", "sub total", "subtotal", "total")):
                break

            rate_str = "" if rate_raw is None else str(rate_raw).strip()
            is_rate_empty = (rate_str == "")

            # Heading row: description present, but rate empty OR description cell merged
            if desc and (is_rate_empty or is_merged(ws, r, 4)):
                continue

            qty_raw = ws.cell(row=r, column=2).value  # B
            unit_raw = ws.cell(row=r, column=3).value  # C

            qty_str = "" if qty_raw is None else str(qty_raw).strip()
            unit_str = "" if unit_raw is None else str(unit_raw).strip()
            amt_str = "" if amt_raw is None else str(amt_raw).strip()

            all_blank = (
                desc == "" and
                is_rate_empty and
                qty_str == "" and
                unit_str == "" and
                amt_str == ""
            )
            if all_blank:
                continue

            if desc == "" and is_rate_empty:
                continue

            qty = to_num(qty_raw)
            rate = to_num(rate_raw)
            amt = to_num(amt_raw)

            # If only Amount given, infer qty or rate
            if amt != 0:
                if qty == 0 and rate != 0:
                    qty = amt / rate
                elif rate == 0 and qty != 0:
                    rate = amt / qty

            items.append({"qty": qty, "rate": rate})

        return items

    # --- WORKSLIP HELPERS (copied from your bill() logic, but only for amount) ---

    def find_all_workslip_sheets_amt(wb_inner):
        results = []
        for ws in wb_inner.worksheets:
            max_scan = min(ws.max_row, 60)
            for r in range(1, max_scan + 1):
                b = str(ws.cell(row=r, column=2).value or "").strip().lower()
                c = str(ws.cell(row=r, column=3).value or "").strip().lower()

                has_desc = (
                    "description of item" in b or
                    "description of item" in c or
                    ("description" in b and "item" in b) or
                    ("description" in c and "item" in c)
                )

                has_qty = False
                for col in range(5, 12):
                    t = str(ws.cell(row=r, column=col).value or "").strip().lower()
                    if "qty" in t or "quantity" in t:
                        has_qty = True
                        break

                if has_desc and has_qty:
                    results.append(ws)
                    break

        if not results:
            # fallback similar to your find_workslip_sheet
            for ws in wb_inner.worksheets:
                for r in range(1, 40):
                    b = str(ws.cell(row=r, column=2).value or "").strip().lower()
                    g = str(ws.cell(row=r, column=7).value or "").strip().lower()
                    if "description" in b and ("qty" in g or "quantity" in g):
                        return [ws]
            return [wb_inner.worksheets[0]]
        return results

    def parse_workslip_items_amt(ws):
        """
        Parse workslip items for amount calculation, detecting the last workslip phase columns.
        Rate is always from Estimate Rate column for multi-phase workslips.
        """
        items = []
        max_row = min(ws.max_row, 5000)
        
        # Step 1: Find header row
        header_row = 8
        for r in range(1, 15):
            cell_val = str(ws.cell(row=r, column=1).value or "").strip().lower()
            if "sl" in cell_val:
                header_row = r
                break
        
        # Step 2: Scan header row to find all workslip/execution qty columns
        qty_columns = []
        amt_columns = []
        estimate_rate_col = 5  # Default - Rate (Estimate) column
        
        for c in range(1, 40):  # Scan more columns for multi-phase workslips
            header = str(ws.cell(row=header_row, column=c).value or "").strip().lower()
            
            is_exec_col = ("execution" in header or "exec" in header or "workslip" in header)
            is_qty = ("qty" in header or "quantity" in header)
            is_amt = ("amount" in header or "amt" in header)
            is_estimate = ("est" in header or "estimate" in header)
            is_rate = ("rate" in header)
            
            # Find estimate rate column
            if is_estimate and is_rate:
                estimate_rate_col = c
            
            if is_estimate:
                continue
            
            # Skip More/Less/Remarks columns
            if "more" in header or "less" in header or "remark" in header:
                continue
                
            if is_exec_col and is_qty:
                qty_columns.append(c)
            elif is_exec_col and is_amt:
                amt_columns.append(c)
        
        # Step 3: Use the LAST workslip columns (latest phase)
        if qty_columns:
            qty_col = qty_columns[-1]
            amt_col = amt_columns[-1] if amt_columns else qty_col + 1
            rate_col = estimate_rate_col  # Always use Estimate Rate
        else:
            qty_col = 7
            rate_col = 8
            amt_col = 9

        for r in range(1, max_row + 1):
            desc_raw = ws.cell(row=r, column=2).value  # B
            desc = str(desc_raw or "").strip()
            if desc == "":
                continue

            low = desc.lower()

            if low == "description of item":
                continue
            if low.startswith("sub total") or low.startswith("sub-total") or low.startswith("ecv"):
                break
            if "supplemental items" in low:
                continue

            if is_merged(ws, r, 2):
                continue

            qty_raw = ws.cell(row=r, column=qty_col).value
            rate_raw = ws.cell(row=r, column=rate_col).value
            amt_raw = ws.cell(row=r, column=amt_col).value

            qty_exec = to_num(qty_raw)
            rate_exec = to_num(rate_raw)
            amt_exec = to_num(amt_raw)

            if amt_exec != 0:
                if qty_exec == 0 and rate_exec != 0:
                    qty_exec = amt_exec / rate_exec
                elif rate_exec == 0 and qty_exec != 0:
                    rate_exec = amt_exec / qty_exec

            # Pure heading: text but no qty, no rate, no amount
            if qty_exec == 0 and rate_exec == 0 and amt_exec == 0:
                continue

            items.append({"qty": qty_exec, "rate": rate_exec})

        return items

    def read_tp_from_sheet_amt(ws):
        tp_percent = None
        tp_type = None
        max_row = ws.max_row
        for r in range(1, max_row + 1):
            for col in (2, 4):
                val = ws.cell(row=r, column=col).value
                if not val:
                    continue
                s = str(val).strip()
                low = s.lower()
                if "t.p" in low:
                    if low.startswith("add"):
                        tp_type = "Excess"
                    elif low.startswith("deduct"):
                        tp_type = "Less"
                    m = re.search(r"(\d+(\.\d+)?)", low)
                    if m:
                        tp_percent = float(m.group(1))
                    return tp_percent, tp_type
        return None, None

    # --------------- CASE 1: Estimate â†’ First / First & Final ---------------

    if action in ("estimate_first_part", "estimate_first_final"):
        sheets_info = find_all_estimate_sheets_amt(wb)
        if not sheets_info:
            return 0.0

        total_all_sheets = 0.0

        # T.P from session (same as bill())
        tp_percent_session = 0.0
        tp_type_session = "Excess"
        if request is not None:
            tp_percent_session = to_num(request.session.get("ws_tp_percent", 0.0))
            tp_type_session = (request.session.get("ws_tp_type") or "Excess")

        for (ws_est, header_row) in sheets_info:
            items = parse_estimate_items_amt(ws_est, header_row)
            if not items:
                continue

            subtotal = 0.0
            for it in items:
                subtotal += it["qty"] * it["rate"]

            tp_percent = tp_percent_session
            tp_type = tp_type_session

            tp_abs = subtotal * abs(tp_percent) / 100.0
            if tp_type == "Less":
                total_sheet = subtotal - tp_abs
            else:
                total_sheet = subtotal + tp_abs

            total_all_sheets += total_sheet

        return total_all_sheets

    # --------------- CASE 2: WorkSlip â†’ First / First & Final ---------------

    if action in ("workslip_first_part", "workslip_first_final"):
        ws_list = find_all_workslip_sheets_amt(wb)
        if not ws_list:
            return _extract_total_amount_from_bill_wb(wb)

        total_all_sheets = 0.0

        for ws_ws in ws_list:
            items = parse_workslip_items_amt(ws_ws)
            if not items:
                continue

            subtotal = 0.0
            for it in items:
                subtotal += it["qty"] * it["rate"]

            # T.P from sheet; if not found, use session  -  same as bill()
            tp_percent, tp_type = read_tp_from_sheet_amt(ws_ws)

            if request is not None:
                if tp_percent is None:
                    tp_percent = to_num(request.session.get("ws_tp_percent", 0.0))
                if not tp_type:
                    tp_type = (request.session.get("ws_tp_type") or "Excess")
            else:
                if tp_percent is None:
                    tp_percent = 0.0
                if not tp_type:
                    tp_type = "Excess"

            tp_abs = subtotal * abs(tp_percent) / 100.0
            if tp_type == "Less":
                total_sheet = subtotal - tp_abs
            else:
                total_sheet = subtotal + tp_abs

            total_all_sheets += total_sheet

        return total_all_sheets

    # --------------- CASE 3: Nth & Part / Nth & Final (N â‰¥ 2) ---------------

    # For these, the uploaded Excel is a Bill already.
    # _extract_total_amount_from_bill_wb:
    #  - detects 8-col vs 10-col
    #  - reads:
    #       H (col 8) of 'Total' row for 1st / First & Final
    #       I (col 9) of 'Total' row for Nth bills.
    return _extract_total_amount_from_bill_wb(wb)



def _build_mb_details_string(mb_measure_no, mb_measure_from, mb_measure_to,
                             mb_abs_no, mb_abs_from, mb_abs_to):
    return (
        f"MB.No. {mb_measure_no} P.No. {mb_measure_from} to {mb_measure_to} (Measurements) "
        f"& MB.No. {mb_abs_no} P.No. {mb_abs_from} to {mb_abs_to} (Abstract)"
    )


def _resolve_cc_header(action, nth_number_str=None):
    """
    Rebuild the CC header line based on 'action' (same actions used in bill view):

      - estimate_first_part      -> 'CC First & Part Bill'
      - estimate_first_final     -> 'CC First & Final Bill'
      - workslip_first_part      -> 'CC First & Part Bill'
      - workslip_first_final     -> 'CC First & Final Bill'
      - firstpart_nth_part, nth_nth_part   -> 'CC <Nth> & Part Bill'
      - firstpart_2nd_final, nth_nth_final -> 'CC <Nth> & Final Bill'
    """
    def ordinal_word(n):
        mapping = {
            1: "First",
            2: "Second",
            3: "Third",
            4: "Fourth",
            5: "Fifth",
            6: "Sixth",
            7: "Seventh",
            8: "Eighth",
            9: "Ninth",
            10: "Tenth",
        }
        if n in mapping:
            return mapping[n]
        if 10 < n % 100 < 14:
            suffix = "th"
        else:
            suffix = {1: "st", 2: "nd", 3: "rd"}.get(n % 10, "th")
        return f"{n}{suffix}"

    base_action = action or ""

    if base_action in ("estimate_first_part", "workslip_first_part"):
        return "CC First & Part Bill"
    if base_action in ("estimate_first_final", "workslip_first_final"):
        return "CC First & Final Bill"

    # Nth-type actions
    try:
        n_val = int(nth_number_str or "2")
    except Exception:
        n_val = 2
    if n_val < 2:
        n_val = 2
    ord_word = ordinal_word(n_val)

    if base_action in ("firstpart_nth_part", "nth_nth_part"):
        return f"CC {ord_word} & Part Bill"
    if base_action in ("firstpart_2nd_final", "nth_nth_final"):
        return f"CC {ord_word} & Final Bill"

    # fallback
    return "CC Bill"


def _fill_excel_template(template_filename, context_dict):
    """
    Open Excel template from BILL_TEMPLATES_DIR, replace {{KEY}} placeholders
    with values from context_dict, return openpyxl Workbook.
    """
    template_path = os.path.join(BILL_TEMPLATES_DIR, template_filename)
    wb = load_workbook(template_path)

    # placeholders: {{NAME_OF_WORK}}, etc.
    placeholders = list(context_dict.keys())

    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str):
                    text = cell.value
                    # replace {{KEY}} style placeholders
                    for key in placeholders:
                        placeholder = "{{" + key + "}}"
                        if placeholder in text:
                            text = text.replace(placeholder, str(context_dict[key]))
                    cell.value = text
    return wb


def _fill_docx_template(template_filename, context_dict):
    """
    Open Word template from BILL_TEMPLATES_DIR, replace {{KEY}} placeholders
    in paragraphs and tables, return python-docx Document.
    """
    template_path = os.path.join(BILL_TEMPLATES_DIR, template_filename)
    doc = Document(template_path)
    placeholders = list(context_dict.keys())

    # Helper to replace in a python-docx paragraph (handling runs)
    def replace_in_paragraph(paragraph):
        if not paragraph.text:
            return
        # Join runs to single string
        text = paragraph.text
        replaced = text
        for key in placeholders:
            ph = "{{" + key + "}}"
            if ph in replaced:
                replaced = replaced.replace(ph, str(context_dict[key]))
        if replaced != text:
            # Clear runs, set one run with replaced text (keeps style of first run)
            for _ in range(len(paragraph.runs)):
                paragraph.runs[0].text = ""
                paragraph.runs.pop(0)
            paragraph.add_run(replaced)

    for p in doc.paragraphs:
        replace_in_paragraph(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)

    return doc

# -----------------------
# BILL FROM WORKSLIP
# -----------------------
# -----------------------
# BILL  -  from Estimate / from WorkSlip
# -----------------------
# imports for bill/templating consolidated at top of file

# BILL_TEMPLATES_DIR, _number_to_words_rupees, _extract_header_data_fuzzy_from_wb,
# _extract_total_amount_from_bill_wb, _build_mb_details_string,
# _resolve_cc_header, _fill_excel_template, _fill_docx_template
# are assumed to be defined ABOVE this function (as in your file).


@org_required
def bill_document(request):
    """
    Generate:
      - LS Form (Part / Final) -> Excel (multiple sheets if multiple bills)
      - Covering Letter (Word) -> Multiple pages if multiple bills
      - Movement Slip (Word) -> Multiple pages if multiple bills

    directly from an uploaded Bill Excel + MB details.
    Returns the file directly for download (synchronous).

    POST fields:
      - bill_file     : uploaded Excel (Bill / Estimate / Nth bill)
      - doc_kind      : 'ls_part', 'ls_final', 'covering', 'movement'
      - action        : same as in bill() (estimate_first_part, firstpart_nth_part, ...)
      - nth_number    : Nth number (for N>=2 bills)
      - mb_measure_no, mb_measure_p_from, mb_measure_p_to
      - mb_abstract_no, mb_abstract_p_from, mb_abstract_p_to
    """
    if request.method != "POST":
        return HttpResponseNotAllowed(["POST"], "Use POST to generate documents.")

    # 1) Get uploaded Bill Excel
    uploaded = request.FILES.get("bill_file")
    if not uploaded:
        return HttpResponse(
            "Please upload the Bill Excel file for document generation.",
            status=400,
        )

    # 2) Which document?
    doc_kind = (request.POST.get("doc_kind") or "").strip()
    if not doc_kind:
        return HttpResponse("Missing 'doc_kind' in request.", status=400)

    # Action + Nth number (for CC header)
    action = (request.POST.get("action") or "").strip()
    nth_number_str = (request.POST.get("nth_number") or "").strip()

    # 3) MB details from form
    mb_measure_no = (request.POST.get("mb_measure_no") or "").strip()
    mb_measure_p_from = (request.POST.get("mb_measure_p_from") or "").strip()
    mb_measure_p_to = (request.POST.get("mb_measure_p_to") or "").strip()
    mb_abs_no = (request.POST.get("mb_abstract_no") or "").strip()
    mb_abs_p_from = (request.POST.get("mb_abstract_p_from") or "").strip()
    mb_abs_p_to = (request.POST.get("mb_abstract_p_to") or "").strip()

    # 4) Open uploaded Excel
    try:
        wb_in = load_workbook(uploaded, data_only=True)
    except Exception as e:
        return HttpResponse(f"Error reading uploaded Bill Excel: {e}", status=400)
    
    # Also open with formulas for fallback calculation
    try:
        uploaded.seek(0)  # Reset file pointer
        wb_formulas = load_workbook(uploaded, data_only=False)
    except:
        wb_formulas = None

    # Find all bill sheets (sheets that look like bills/estimates)
    bill_sheets = [ws for ws in wb_in.worksheets if ws.title.startswith("Bill")]
    if not bill_sheets:
        # Fallback: try to find sheets with estimate/bill-like structure
        bill_sheets = []
        for ws in wb_in.worksheets:
            # Check if it has typical bill structure
            if ws.max_row > 5:
                bill_sheets.append(ws)
        if not bill_sheets:
            bill_sheets = wb_in.worksheets  # use all sheets
    
    has_multiple_bills = len(bill_sheets) > 1

    # MB details string & CC header (same for all sheets)
    mb_details_str = _build_mb_details_string(
        mb_measure_no,
        mb_measure_p_from,
        mb_measure_p_to,
        mb_abs_no,
        mb_abs_p_from,
        mb_abs_p_to,
    )
    cc_header = _resolve_cc_header(action, nth_number_str=nth_number_str)
    
    # Current month + year
    now = timezone.now()
    mm_yyyy = f"{now.month:02d}.{now.year}"

    # Helper function to extract total from a single sheet
    def _extract_total_from_sheet(ws, ws_formulas=None):
        total = 0.0
        max_scan = min(ws.max_row, 200)
        
        for r in range(1, max_scan + 1):
            for check_col in [3, 4, 5]:
                cell_val = str(ws.cell(row=r, column=check_col).value or "").strip().lower()
                
                if cell_val == "total":
                    for amt_col in [8, 9, 10]:
                        amt_val = ws.cell(row=r, column=amt_col).value
                        
                        if (amt_val is None or amt_val == 0 or amt_val == '') and ws_formulas:
                            try:
                                formula_cell = ws_formulas.cell(row=r, column=amt_col)
                                formula_val = formula_cell.value
                                
                                if isinstance(formula_val, str) and formula_val.startswith('='):
                                    match = re.match(r'=([A-Z]+)(\d+)([\+\-])([A-Z]+)(\d+)', formula_val)
                                    if match:
                                        col1, row1, op, col2, row2 = match.groups()
                                        def col_to_num(col):
                                            return ord(col) - ord('A') + 1
                                        val1 = ws.cell(row=int(row1), column=col_to_num(col1)).value
                                        val2 = ws.cell(row=int(row2), column=col_to_num(col2)).value
                                        if val1 and val2:
                                            if op == '+':
                                                amt_val = float(val1) + float(val2)
                                            elif op == '-':
                                                amt_val = float(val1) - float(val2)
                            except Exception:
                                pass
                        
                        try:
                            num_val = float(amt_val) if amt_val else 0
                            if num_val != 0:
                                return num_val
                        except (TypeError, ValueError):
                            continue
                    break
        return total

    # 10) LS FORMS (EXCEL) - Multiple sheets support
    if doc_kind in ("ls_part", "ls_final"):
        if doc_kind == "ls_part":
            template_name = "LS_Form_Part.xlsx"
            download_name = "LS_Form_Part.xlsx"
        else:
            template_name = "LS_Form_Final.xlsx"
            download_name = "LS_Form_Final.xlsx"

        template_path = os.path.join(BILL_TEMPLATES_DIR, template_name)
        if not os.path.exists(template_path):
            return HttpResponse(f"LS template not found: {template_name}", status=404)

        # If multiple bill sheets, create one LS sheet per bill
        if has_multiple_bills:
            wb_out = Workbook()
            # Remove default sheet
            if wb_out.worksheets:
                wb_out.remove(wb_out.worksheets[0])
            
            for sheet_idx, bill_ws in enumerate(bill_sheets, start=1):
                # Extract per-sheet header data
                sheet_header = _extract_header_data_from_sheet(bill_ws)
                sheet_name_of_work = sheet_header.get("name_of_work", "") or ""
                sheet_agreement_ref = sheet_header.get("agreement", "") or ""
                sheet_agency_name = sheet_header.get("agency", "") or ""
                
                # Get formulas sheet if available
                ws_formulas = None
                if wb_formulas:
                    try:
                        ws_formulas = wb_formulas[bill_ws.title]
                    except:
                        pass
                
                # Calculate total for this sheet
                sheet_total = _extract_total_from_sheet(bill_ws, ws_formulas)
                try:
                    sheet_total_str = f"{float(sheet_total):,.2f}"
                except:
                    sheet_total_str = str(sheet_total)
                sheet_amount_words = _number_to_words_rupees(sheet_total)
                
                # Build context for this sheet
                sheet_ctx = {
                    "NAME_OF_WORK": sheet_name_of_work,
                    "NAME_OF_AGENCY": sheet_agency_name,
                    "AGENCY_NAME": sheet_agency_name,
                    "REF_OF_AGREEMENT": sheet_agreement_ref,
                    "AGREEMENT_REF": sheet_agreement_ref,
                    "MB_DETAILS": mb_details_str,
                    "CC_HEADER": cc_header,
                    "AMOUNT": sheet_total_str,
                    "TOTAL_AMOUNT": sheet_total_str,
                    "AMOUNT_IN_WORDS": sheet_amount_words,
                }
                
                # Load fresh template for this sheet
                try:
                    template_wb = load_workbook(template_path)
                except Exception as e:
                    return HttpResponse(f"Error loading template: {e}", status=500)
                
                # Fill placeholders in template
                for ws in template_wb.worksheets:
                    for row in ws.iter_rows():
                        for cell in row:
                            if isinstance(cell.value, str):
                                text = cell.value
                                for key, val in sheet_ctx.items():
                                    placeholder = "{{" + key + "}}"
                                    if placeholder in text:
                                        text = text.replace(placeholder, str(val or ""))
                                cell.value = text
                    
                    # Copy this sheet to output workbook
                    new_sheet_name = f"{bill_ws.title}_LS" if len(bill_ws.title) <= 25 else f"LS_{sheet_idx}"
                    new_ws = wb_out.create_sheet(title=new_sheet_name)
                    
                    # Copy cell values, styles, and merged cells
                    for row in ws.iter_rows():
                        for cell in row:
                            new_cell = new_ws[cell.coordinate]
                            new_cell.value = cell.value
                            if cell.has_style:
                                new_cell.font = copy(cell.font)
                                new_cell.border = copy(cell.border)
                                new_cell.fill = copy(cell.fill)
                                new_cell.number_format = copy(cell.number_format)
                                new_cell.protection = copy(cell.protection)
                                new_cell.alignment = copy(cell.alignment)
                    
                    # Copy merged cells
                    for merged_range in ws.merged_cells.ranges:
                        new_ws.merge_cells(str(merged_range))
                    
                    # Copy column widths
                    for col_letter, col_dim in ws.column_dimensions.items():
                        new_ws.column_dimensions[col_letter].width = col_dim.width
                    
                    # Copy row heights
                    for row_num, row_dim in ws.row_dimensions.items():
                        new_ws.row_dimensions[row_num].height = row_dim.height
            
            download_name = f"LS_Forms_{'Part' if doc_kind == 'ls_part' else 'Final'}.xlsx"
        else:
            # Single sheet - use original logic
            header = _extract_header_data_fuzzy_from_wb(wb_in)
            total_amount = _extract_total_from_sheet(bill_sheets[0], wb_formulas[bill_sheets[0].title] if wb_formulas else None)
            try:
                total_amount_str = f"{float(total_amount):,.2f}"
            except:
                total_amount_str = str(total_amount)
            
            ctx = {
                "NAME_OF_WORK": header.get("name_of_work", ""),
                "NAME_OF_AGENCY": header.get("agency", ""),
                "AGENCY_NAME": header.get("agency", ""),
                "REF_OF_AGREEMENT": header.get("agreement", ""),
                "AGREEMENT_REF": header.get("agreement", ""),
                "MB_DETAILS": mb_details_str,
                "CC_HEADER": cc_header,
                "AMOUNT": total_amount_str,
                "TOTAL_AMOUNT": total_amount_str,
                "AMOUNT_IN_WORDS": _number_to_words_rupees(total_amount),
            }
            
            try:
                wb_out = _fill_excel_template(template_name, ctx)
            except FileNotFoundError:
                return HttpResponse(f"LS template not found: {template_name}", status=404)
            except Exception as e:
                return HttpResponse(f"Error filling LS template: {e}", status=500)

        resp = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        resp["Content-Disposition"] = f'attachment; filename="{download_name}"'
        wb_out.save(resp)
        return resp

    # 11) COVERING LETTER / MOVEMENT SLIP (WORD) - Multiple pages support
    if doc_kind in ("covering", "movement"):
        from core.template_views import get_user_template
        
        if doc_kind == "covering":
            template_type = "covering_letter"
            base_download_name = "Covering_Letter"
        else:
            template_type = "movement_slip"
            base_download_name = "Movement_Slip"
        
        # Try to get user's custom template first
        user_template = get_user_template(request.user, template_type)
        
        if user_template:
            # User has uploaded their own template
            template_path = user_template.file.path
        else:
            # No user template - show error with link to upload
            template_type_display = "Covering Letter" if doc_kind == "covering" else "Movement Slip"
            error_html = f"""
            <html>
            <head><title>Template Required</title></head>
            <body style="font-family: Arial, sans-serif; padding: 40px; text-align: center;">
                <h2>âš ï¸ {template_type_display} Template Not Found</h2>
                <p>You haven't uploaded a {template_type_display} template yet.</p>
                <p>Please upload your own template with your officer names and formatting.</p>
                <a href="/templates/" style="display: inline-block; margin-top: 20px; padding: 12px 24px; background: #007bff; color: white; text-decoration: none; border-radius: 5px;">
                    Upload Template
                </a>
                <br><br>
                <a href="/bill/" style="color: #666;">â† Back to Bill Generator</a>
            </body>
            </html>
            """
            return HttpResponse(error_html, status=404)
        
        if not os.path.exists(template_path):
            return HttpResponse(f"Template file not found. Please re-upload your template.", status=404)

        def replace_in_paragraphs(paragraphs, placeholder_map, mm_yyyy_val):
            for p in paragraphs:
                if not p.runs:
                    continue
                for run in p.runs:
                    if not run.text:
                        continue
                    text = run.text
                    changed = False
                    for ph, val in placeholder_map.items():
                        if ph in text:
                            text = text.replace(ph, val or "")
                            changed = True
                    if "dd.mm.yyyy" in text:
                        text = text.replace("dd.mm.yyyy", f"dd.{mm_yyyy_val}")
                        changed = True
                    if changed:
                        run.text = text

        # If multiple bill sheets, create combined document with page breaks
        if has_multiple_bills:
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            
            # Start with the template and clear its content
            combined_doc = Document(template_path)
            
            # Clear the template content first
            for element in list(combined_doc.element.body):
                combined_doc.element.body.remove(element)
            
            for sheet_idx, bill_ws in enumerate(bill_sheets, start=1):
                # Extract per-sheet header data
                sheet_header = _extract_header_data_from_sheet(bill_ws)
                sheet_name_of_work = sheet_header.get("name_of_work", "") or ""
                sheet_agreement_ref = sheet_header.get("agreement", "") or ""
                sheet_agency_name = sheet_header.get("agency", "") or ""
                
                # Get formulas sheet if available
                ws_formulas = None
                if wb_formulas:
                    try:
                        ws_formulas = wb_formulas[bill_ws.title]
                    except:
                        pass
                
                # Calculate total for this sheet
                sheet_total = _extract_total_from_sheet(bill_ws, ws_formulas)
                try:
                    sheet_total_str = f"{float(sheet_total):,.2f}"
                except:
                    sheet_total_str = str(sheet_total)
                sheet_amount_words = _number_to_words_rupees(sheet_total)
                
                # Build placeholder map for this sheet
                sheet_placeholder_map = {
                    "{{NAME_OF_WORK}}": sheet_name_of_work,
                    "{{AGENCY_NAME}}": sheet_agency_name,
                    "{{NAME_OF_AGENCY}}": sheet_agency_name,
                    "{{AGREEMENT_REF}}": sheet_agreement_ref,
                    "{{REF_OF_AGREEMENT}}": sheet_agreement_ref,
                    "{{CC_HEADER}}": cc_header,
                    "{{MB_DETAILS}}": mb_details_str,
                    "{{AMOUNT}}": sheet_total_str,
                    "{{TOTAL_AMOUNT}}": sheet_total_str,
                    "{{AMOUNT_IN_WORDS}}": sheet_amount_words,
                }
                
                # Load fresh template for this sheet
                sheet_doc = Document(template_path)
                
                # Replace placeholders in body
                replace_in_paragraphs(sheet_doc.paragraphs, sheet_placeholder_map, mm_yyyy)
                
                # Replace in tables
                for table in sheet_doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            replace_in_paragraphs(cell.paragraphs, sheet_placeholder_map, mm_yyyy)
                
                # Add page break at the END of this document (except for the last one)
                # This ensures each estimate starts on a new page without blank pages
                if sheet_idx < len(bill_sheets):
                    sheet_doc.add_page_break()
                
                # Append this document's content to combined document
                for element in sheet_doc.element.body:
                    combined_doc.element.body.append(element)
            
            buf = io.BytesIO()
            combined_doc.save(buf)
            buf.seek(0)
            
            download_name = f"{base_download_name}s.docx"
            response = HttpResponse(
                buf.getvalue(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            response["Content-Disposition"] = f'attachment; filename="{download_name}"'
            return response
        
        else:
            # Single sheet - original logic
            header = _extract_header_data_fuzzy_from_wb(wb_in)
            total_amount = _extract_total_from_sheet(bill_sheets[0], wb_formulas[bill_sheets[0].title] if wb_formulas else None)
            try:
                total_amount_str = f"{float(total_amount):,.2f}"
            except:
                total_amount_str = str(total_amount)
            
            placeholder_map = {
                "{{NAME_OF_WORK}}": header.get("name_of_work", ""),
                "{{AGENCY_NAME}}": header.get("agency", ""),
                "{{NAME_OF_AGENCY}}": header.get("agency", ""),
                "{{AGREEMENT_REF}}": header.get("agreement", ""),
                "{{REF_OF_AGREEMENT}}": header.get("agreement", ""),
                "{{CC_HEADER}}": cc_header,
                "{{MB_DETAILS}}": mb_details_str,
                "{{AMOUNT}}": total_amount_str,
                "{{TOTAL_AMOUNT}}": total_amount_str,
                "{{AMOUNT_IN_WORDS}}": _number_to_words_rupees(total_amount),
            }

            doc = Document(template_path)
            replace_in_paragraphs(doc.paragraphs, placeholder_map, mm_yyyy)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        replace_in_paragraphs(cell.paragraphs, placeholder_map, mm_yyyy)

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)

            download_name = f"{base_download_name}.docx"
            response = HttpResponse(
                buf.getvalue(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            response["Content-Disposition"] = f'attachment; filename="{download_name}"'
            return response

    return HttpResponse(f"Unknown doc_kind: {doc_kind}", status=400)


@org_required
def self_formatted_document(request):
    """
    Generate self-formatted documents asynchronously.
    
    Now uses async job processing instead of in-request generation.
    
    User uploads:
      - bill_file      : Excel (your Bill / Estimate / Nth Bill)
      - template_file  : DOCX or XLSX (user's own format with {{PLACEHOLDERS}})
      - action         : same as bill() (estimate_first_part, workslip_first_final, ...)
      - nth_number     : N value for Nth bills (optional)
      - mb_measure_no, mb_measure_p_from, mb_measure_p_to
      - mb_abstract_no, mb_abstract_p_from, mb_abstract_p_to
    """
    if request.method != "POST":
        return HttpResponseNotAllowed(["POST"], "Use POST to generate self-formatted documents.")

    org = get_org_from_request(request)

    # ---------- Required uploads ----------
    bill_file = request.FILES.get("bill_file")
    template_file = request.FILES.get("template_file")

    if not bill_file:
        return JsonResponse({"error": "Please upload the Bill Excel file (bill_file)."}, status=400)
    if not template_file:
        return JsonResponse({"error": "Please upload the template file (template_file)."}, status=400)

    # ---------- Meta info ----------
    action = (request.POST.get("action") or "").strip()
    nth_number_str = (request.POST.get("nth_number") or "").strip()

    # MB details
    mb_measure_no = (request.POST.get("mb_measure_no") or "").strip()
    mb_measure_p_from = (request.POST.get("mb_measure_p_from") or "").strip()
    mb_measure_p_to = (request.POST.get("mb_measure_p_to") or "").strip()
    mb_abs_no = (request.POST.get("mb_abstract_no") or "").strip()
    mb_abs_p_from = (request.POST.get("mb_abstract_p_from") or "").strip()
    mb_abs_p_to = (request.POST.get("mb_abstract_p_to") or "").strip()

    # Create Upload for bill file
    try:
        upload = Upload.objects.create(
            organization=org,
            user=request.user,
            filename=bill_file.name,
            file_size=bill_file.size,
            status='processing'
        )
    except Exception as e:
        return JsonResponse({"error": f"Failed to save upload: {e}"}, status=400)

    # Enqueue async job
    metadata = {
        'action': action,
        'nth_number': nth_number_str,
        'mb_measure_no': mb_measure_no,
        'mb_measure_p_from': mb_measure_p_from,
        'mb_measure_p_to': mb_measure_p_to,
        'mb_abs_no': mb_abs_no,
        'mb_abs_p_from': mb_abs_p_from,
        'mb_abs_p_to': mb_abs_p_to,
        'template_filename': template_file.name,
        'upload_id': upload.id,
    }

    try:
        job, task = create_job_for_excel(
            request,
            upload=upload,
            job_type='generate_self_formatted_document',
            metadata=metadata
        )
        
        return JsonResponse({
            'job_id': job.id,
            'status_url': reverse('job_status', args=[job.id]),
            'message': 'Generating self-formatted document. You will be notified when ready.'
        })
    except Exception as e:
        return JsonResponse({"error": f"Failed to enqueue task: {e}"}, status=500)


# -----------------------
# SIMPLE PAGES & PROJECTS
# -----------------------
@login_required(login_url='login')
def my_subscription(request):
    return render(request, "core/my_subscription.html")

@org_required
def my_projects(request):
    org = get_org_from_request(request)
    projects = Project.objects.for_org(org)
    return render(request, "core/my_projects.html", {"projects": projects})

@org_required
def create_project(request):
    org = get_org_from_request(request)
    if request.method == "POST":
        name = request.POST.get("project_name")
        if name:
            Project.objects.get_or_create(organization=org, name=name)
    return redirect("my_projects")

@login_required(login_url='login')
def datas(request):
    """
    Landing page for 'New Estimate'.

    - Clears current selection (as you had before)
    - Reads ?work_type=original/repair from URL and stores in session
    - Defaults to 'original' if nothing selected
    """
    # Always start fresh when entering Datas
    request.session["fetched_items"] = []
    request.session["current_project_name"] = None
    request.session["qty_map"] = {}
    request.session["unit_map"] = {}
    request.session["work_name"] = ""
    request.session["grand_total"] = ""
    request.session["selected_backend_id"] = None  # Clear any previous backend selection

    mode = request.GET.get("work_type")

    if mode in ("original", "repair"):
        request.session["work_type"] = mode

    # If nothing in session yet, default to original
    if "work_type" not in request.session:
        request.session["work_type"] = "original"

    return render(
        request,
        "core/datas.html",
        {"work_type": request.session["work_type"]},
    )



# -----------------------
# STEP 1: SELECT PROJECT
# -----------------------
@login_required(login_url='login')
def select_project(request):
    org = get_org_from_request(request)
    projects = Project.objects.for_org(org)

    use_id = request.GET.get("use")
    if use_id:
        request.session["selected_project_id"] = use_id
        request.session["fetched_items"] = []
        return redirect("choose_category")

    if request.method == "POST":
        project_name = request.POST.get("project_name")
        if project_name:
            project, created = Project.objects.get_or_create(organization=org, name=project_name)
            request.session["selected_project_id"] = project.id
            request.session["fetched_items"] = []
            return redirect("choose_category")

    return render(request, "core/select_project.html", {"projects": projects})


# -----------------------
# STEP 2: SELECT CATEGORY
# -----------------------
@login_required(login_url='login')
def choose_category(request):
    return render(request, "core/choose_category.html")


# -----------------------
# STEP 3: GROUPS PAGE (redirects to items)
# (imports consolidated at top)
# -----------------------
@login_required(login_url='login')
def datas_groups(request, category):
    # NEW: remember work_type in session if passed in URL
    work_type = (request.GET.get("work_type") or "").lower()
    if work_type in ("original", "repair"):
        request.session["work_type"] = work_type

    # NEW: Store selected backend_id in session for consistent use throughout the flow
    backend_id = request.GET.get("backend_id")
    if backend_id:
        try:
            request.session["selected_backend_id"] = int(backend_id)
        except (ValueError, TypeError):
            request.session["selected_backend_id"] = None
    
    # Get backend_id from session for loading
    selected_backend_id = request.session.get("selected_backend_id")
    
    try:
        items_list, groups_map, _, ws_data, filepath = load_backend(
            category, settings.BASE_DIR, 
            backend_id=selected_backend_id,
            module_code='new_estimate',
            user=request.user if request.user.is_authenticated else None
        )
    except FileNotFoundError as e:
        logger.info(f"No backend data available for category {category} - showing Coming Soon")
        # Check if the other category has backends available
        other_category = 'electrical' if category == 'civil' else 'civil'
        other_available = False
        try:
            from core.utils_excel import get_available_backends_for_module
            other_backends = get_available_backends_for_module('new_estimate', other_category)
            other_available = len(other_backends) > 0
        except:
            pass
        return render(request, "core/coming_soon.html", {
            "category": category,
            "module_name": "New Estimate",
            "other_category": other_category,
            "other_category_available": other_available,
        })
    except Exception as e:
        logger.error(f"Error loading backend for {category}: {e}")
        return render(request, "core/error.html", {
            "error_title": "Loading Error",
            "error_message": f"Could not load backend data: {str(e)}",
        })
    
    groups = sorted(groups_map.keys(), key=lambda s: s.lower())

    if not groups:
        return render(request, "core/groups.html", {
            "category": category,
            "groups": [],
            "error": "No groups found in backend Excel.",
        })

    default_group = request.GET.get("group") or groups[0]
    return redirect("datas_items", category=category, group=default_group)



# -----------------------
# STEP 4: ITEMS IN GROUP (3-panel UI)
# -----------------------
@login_required(login_url='login')
def datas_items(request, category, group):
    from core.utils_excel import get_available_backends_for_module
    
    # Check if backend_id is passed in URL (for switching backends on this page)
    url_backend_id = request.GET.get("backend_id")
    if url_backend_id:
        try:
            request.session["selected_backend_id"] = int(url_backend_id)
        except (ValueError, TypeError):
            pass
    
    # Get backend_id from session for consistent loading throughout the flow
    selected_backend_id = request.session.get("selected_backend_id")
    
    # Get available backends for the dropdown
    try:
        available_backends = get_available_backends_for_module('new_estimate', category)
    except Exception as e:
        logger.error(f"Error getting available backends: {e}")
        available_backends = []
    
    try:
        items_list, groups_map, backend_units_map, ws_data, filepath = load_backend(
            category, settings.BASE_DIR,
            backend_id=selected_backend_id,
            module_code='new_estimate',
            user=request.user if request.user.is_authenticated else None
        )
    except FileNotFoundError as e:
        logger.info(f"No backend data available for category {category} - showing Coming Soon")
        # Check if the other category has backends available
        other_category = 'electrical' if category == 'civil' else 'civil'
        other_available = False
        try:
            other_backends = get_available_backends_for_module('new_estimate', other_category)
            other_available = len(other_backends) > 0
        except:
            pass
        return render(request, "core/coming_soon.html", {
            "category": category,
            "module_name": "New Estimate",
            "other_category": other_category,
            "other_category_available": other_available,
        })
    except Exception as e:
        logger.error(f"Error loading backend for {category}: {e}")
        return render(request, "core/error.html", {
            "error_title": "Loading Error",
            "error_message": f"Could not load backend data: {str(e)}",
        })

    groups = sorted(groups_map.keys(), key=lambda s: s.lower())

    group_items = groups_map.get(group, [])
    detected_names = {i["name"] for i in items_list}
    display_items = [name for name in group_items if name in detected_names]

    wb_vals = load_workbook(filepath, data_only=True)
    ws_vals = wb_vals["Master Datas"]

    item_rates = {}
    for info in items_list:
        name = info["name"]
        start_row = info["start_row"]
        end_row = info["end_row"]
        rate = None
        for r in range(end_row, start_row - 1, -1):
            val = ws_vals.cell(row=r, column=10).value  # column J
            if val not in (None, ""):
                rate = val
                break
        item_rates[name] = rate

    item_to_group = {}
    for grp_name, item_list_in_grp in groups_map.items():
        for nm in item_list_in_grp:
            item_to_group.setdefault(nm, grp_name)

    def units_for(name):
        # First check backend_units_map (Column D from backend Excel)
        backend_unit = backend_units_map.get(name, "")
        if backend_unit:
            return (backend_unit, backend_unit)  # Use same for both plural and singular display
        # Fall back to group-based defaults
        grp_name = item_to_group.get(name, "")
        if grp_name in ("Piping", "Wiring & Cables"):
            return ("Mtrs", "Mtr")
        elif grp_name == "Points":
            return ("Pts", "Pt")
        else:
            return ("Nos", "No")

    # Build item subtypes map: items with ":" are subtypes
    # Group subtypes by their parent name (part before ":")
    item_subtypes = {}  # parent_name -> [list of full subtype names]
    parent_items = set()  # items that have subtypes
    
    for name in display_items:
        if " : " in name:
            # This is a subtype - extract parent name
            parent_name = name.split(" : ")[0].strip()
            if parent_name not in item_subtypes:
                item_subtypes[parent_name] = []
            item_subtypes[parent_name].append(name)
            parent_items.add(parent_name)
    
    items_info = []
    seen_parents = set()
    for name in display_items:
        if " : " in name:
            # This is a subtype - check if we already added the parent
            parent_name = name.split(" : ")[0].strip()
            if parent_name not in seen_parents:
                # Add the parent item with subtypes info - serialize subtypes as JSON string
                subtypes_list = item_subtypes.get(parent_name, [])
                items_info.append({
                    "name": parent_name,
                    "rate": None,  # Parent doesn't have its own rate
                    "has_subtypes": True,
                    "subtypes": json.dumps(subtypes_list),
                    "subtypes_count": len(subtypes_list),
                })
                seen_parents.add(parent_name)
        else:
            # Regular item without subtypes
            items_info.append({
                "name": name,
                "rate": item_rates.get(name),
                "has_subtypes": False,
                "subtypes": "[]",
                "subtypes_count": 0,
            })


    fetched = request.session.get("fetched_items", [])

    qty_map = request.session.get("qty_map", {}) or {}
    unit_map = request.session.get("unit_map", {}) or {}
    work_name = request.session.get("work_name", "") or ""
    grand_total = request.session.get("grand_total", "") or ""

    estimate_rows = []
    for idx, name in enumerate(fetched, start=1):
        default_plural, singular = units_for(name)
        # Priority: 1) user-entered unit from UI, 2) backend_units_map default
        custom_unit = unit_map.get(name, "")
        display_unit = custom_unit if custom_unit else default_plural
        estimate_rows.append({
            "sl": idx,
            "name": name,
            "rate": item_rates.get(name),
            "unit": display_unit,
            "default_unit": default_plural,
            "qty": qty_map.get(name, ""),
        })

    work_type = request.session.get("work_type", "original") or "original"
    excess_tp_percent = request.session.get("excess_tp_percent", "") or ""
    ls_special_name = request.session.get("ls_special_name", "") or ""
    ls_special_amount = request.session.get("ls_special_amount", "") or ""
    deduct_old_material = request.session.get("deduct_old_material", "") or ""
    
    return render(request, "core/items.html", {
        "category": category,
        "group": group,
        "groups": groups,
        "items_info": items_info,
        "fetched": fetched,
        "estimate_rows": estimate_rows,
        "work_name": work_name,
        "grand_total": grand_total,
        "work_type": work_type,
        "excess_tp_percent": excess_tp_percent,
        "ls_special_name": ls_special_name,
        "ls_special_amount": ls_special_amount,
        "deduct_old_material": deduct_old_material,
        "fetched_json": json.dumps(fetched),
        "available_backends": available_backends,
        "selected_backend_id": selected_backend_id,
    })


# -----------------------
# STEP 5: FETCH / UN-FETCH ITEM (toggle)
# -----------------------
@login_required(login_url='login')
def fetch_item(request, category, group, item):
    fetched = request.session.get("fetched_items", []) or []

    if item in fetched:
        fetched.remove(item)
    else:
        fetched.append(item)

    request.session["fetched_items"] = fetched

    work_name = request.GET.get("work_name")
    if work_name is not None:
        request.session["work_name"] = work_name

    return redirect("datas_items", category=category, group=group)


# -----------------------
# AJAX TOGGLE ITEM (no page reload)
# -----------------------
@login_required(login_url='login')
def ajax_toggle_item(request, category):
    """
    AJAX endpoint to toggle an item in fetched list without page reload.
    POST with JSON: { "item": "item_name", "action": "add" or "remove", "work_name": "..." }
    Returns JSON: { "status": "ok", "fetched": [...], "action_taken": "added" or "removed", "item_info": {...} }
    """
    if request.method != "POST":
        return JsonResponse({"status": "error", "message": "POST required"}, status=405)
    
    try:
        if request.content_type and 'application/json' in request.content_type:
            data = json.loads(request.body.decode("utf-8") or "{}")
        else:
            data = request.POST
        
        item = data.get("item", "").strip()
        action = data.get("action", "toggle")  # "add", "remove", or "toggle"
        work_name = data.get("work_name")
        
        if not item:
            return JsonResponse({"status": "error", "message": "No item specified"}, status=400)
        
        fetched = request.session.get("fetched_items", []) or []
        action_taken = None
        
        if action == "add":
            if item not in fetched:
                fetched.append(item)
                action_taken = "added"
            else:
                action_taken = "already_exists"
        elif action == "remove":
            if item in fetched:
                fetched.remove(item)
                action_taken = "removed"
            else:
                action_taken = "not_found"
        else:  # toggle
            if item in fetched:
                fetched.remove(item)
                action_taken = "removed"
            else:
                fetched.append(item)
                action_taken = "added"
        
        request.session["fetched_items"] = fetched
        
        if work_name is not None:
            request.session["work_name"] = work_name
        
        # Get item info (rate, unit) for newly added items
        item_info = None
        if action_taken == "added":
            try:
                # Use same backend as the items page - get from session
                selected_backend_id = request.session.get("selected_backend_id")
                items_list, groups_map, units_map, ws_data, filepath = load_backend(
                    category, settings.BASE_DIR,
                    backend_id=selected_backend_id,
                    module_code='new_estimate',
                    user=request.user
                )
                
                # Get rate
                wb_vals = load_workbook(filepath, data_only=True)
                ws_vals = wb_vals["Master Datas"]
                
                item_rate = None
                for info in items_list:
                    if info["name"] == item:
                        start_row = info["start_row"]
                        end_row = info["end_row"]
                        for r in range(end_row, start_row - 1, -1):
                            val = ws_vals.cell(row=r, column=10).value  # column J
                            if val not in (None, ""):
                                item_rate = val
                                break
                        break
                
                # Get unit with smart fallback (same logic as datas_items)
                # Priority: 1) units_map from backend, 2) group-based defaults
                unit = units_map.get(item, "")
                if not unit:
                    # Find item's group for fallback
                    item_group = ""
                    for grp_name, grp_items in groups_map.items():
                        if item in grp_items:
                            item_group = grp_name
                            break
                    # Group-based defaults
                    if item_group in ("Piping", "Wiring & Cables", "Run of Mains", "Sheathed Cables", "U.G Cabling"):
                        unit = "Mtrs"
                    elif item_group == "Points":
                        unit = "Pts"
                    else:
                        unit = "Nos"
                
                item_info = {
                    "name": item,
                    "rate": item_rate,
                    "unit": unit
                }
                
                wb_vals.close()
            except Exception as e:
                # If we can't get item info, just return without it
                item_info = {"name": item, "rate": None, "unit": "Nos"}
        
        return JsonResponse({
            "status": "ok",
            "fetched": fetched,
            "action_taken": action_taken,
            "item": item,
            "item_info": item_info
        })
        
    except Exception as e:
        return JsonResponse({"status": "error", "message": str(e)}, status=400)


# -----------------------
# AJAX REORDER ITEMS
# -----------------------
@login_required(login_url='login')
def ajax_reorder_items(request, category):
    """
    AJAX endpoint to reorder fetched items list.
    POST with JSON: { "items": ["item1", "item2", ...] }
    Returns JSON: { "status": "ok", "fetched": [...] }
    """
    if request.method != "POST":
        return JsonResponse({"status": "error", "message": "POST required"}, status=405)
    
    try:
        if request.content_type and 'application/json' in request.content_type:
            data = json.loads(request.body.decode("utf-8") or "{}")
        else:
            return JsonResponse({"status": "error", "message": "JSON required"}, status=400)
        
        new_order = data.get("items", [])
        
        if not isinstance(new_order, list):
            return JsonResponse({"status": "error", "message": "items must be a list"}, status=400)
        
        # Validate: new_order should contain the same items as current fetched
        current_fetched = set(request.session.get("fetched_items", []) or [])
        new_order_set = set(new_order)
        
        # Only reorder if sets match (no items added/removed via this endpoint)
        if current_fetched == new_order_set:
            request.session["fetched_items"] = new_order
        else:
            # Allow partial reorder - use the intersection
            valid_items = [item for item in new_order if item in current_fetched]
            request.session["fetched_items"] = valid_items
        
        return JsonResponse({
            "status": "ok",
            "fetched": request.session["fetched_items"]
        })
        
    except Exception as e:
        return JsonResponse({"status": "error", "message": str(e)}, status=400)


# -----------------------
# STEP 6: OUTPUT PANEL
# -----------------------
@login_required(login_url='login')
def output_panel(request, category):
    fetched = request.session.get("fetched_items", [])
    return render(request, "core/output.html", {
        "category": category,
        "items": fetched
    })


# local imports consolidated at top


@login_required(login_url='login')
def download_output(request, category):
    """
    Async Excel generation endpoint.
    
    POST with:
      - qty_map: JSON string of quantities
      - unit_map: JSON string of custom units per item
      - work_name: Name of work
      - work_type: "original" or "repair"
    
    Returns JSON with job_id and status_url for polling.
    Actual Excel generation happens asynchronously via Celery task.
    """
    fetched = request.session.get("fetched_items", [])
    if not fetched:
        return JsonResponse({"error": "No items selected"}, status=400)

    # Parse input
    item_qtys = {}
    item_units = {}
    work_name = ""
    grand_total = None
    
    # Initialize optional parameters with defaults (before POST check)
    excess_tp_enabled = False
    excess_tp_percent = None
    ls_special_enabled = False
    ls_special_name = None
    ls_special_amount = None
    deduct_old_material = None

    if request.method == "POST":
        qty_map_str = request.POST.get("qty_map", "")
        if qty_map_str:
            try:
                raw = json.loads(qty_map_str)
                if isinstance(raw, dict):
                    for k, v in raw.items():
                        try:
                            item_qtys[str(k)] = float(v)
                        except Exception:
                            continue
            except Exception:
                pass

        # Parse unit_map from POST
        unit_map_str = request.POST.get("unit_map", "")
        if unit_map_str:
            try:
                raw_units = json.loads(unit_map_str)
                if isinstance(raw_units, dict):
                    for k, v in raw_units.items():
                        item_units[str(k)] = str(v).strip()
            except Exception:
                pass

        work_name = (request.POST.get("work_name") or "").strip()
        
        # Parse grand_total from POST
        grand_total_str = request.POST.get("grand_total", "").strip()
        if grand_total_str:
            try:
                grand_total = float(grand_total_str)
            except ValueError:
                grand_total = None
        
        # Parse additional options
        excess_tp_enabled = request.POST.get("excess_tp_enabled", "").strip().lower() == 'true'
        if excess_tp_enabled:
            excess_tp_str = request.POST.get("excess_tp_percent", "").strip()
            if excess_tp_str:
                try:
                    excess_tp_percent = float(excess_tp_str)
                except ValueError:
                    excess_tp_percent = None
        
        ls_special_enabled = request.POST.get("ls_special_enabled", "").strip().lower() == 'true'
        if ls_special_enabled:
            ls_special_name = request.POST.get("ls_special_name", "").strip() or None
            ls_special_amount_str = request.POST.get("ls_special_amount", "").strip()
            if ls_special_amount_str:
                try:
                    ls_special_amount = float(ls_special_amount_str)
                except ValueError:
                    ls_special_amount = None
        
        # Parse Deduct Old Material Cost (for repair work)
        deduct_old_material_str = request.POST.get("deduct_old_material", "").strip()
        if deduct_old_material_str:
            try:
                deduct_old_material = float(deduct_old_material_str)
            except ValueError:
                deduct_old_material = None

    work_type = (request.POST.get("work_type")
                 or request.session.get("work_type")
                 or "original").lower()
    request.session["work_type"] = work_type
    
    # Get selected backend ID from session (for multi-backend support)
    selected_backend_id = request.session.get("selected_backend_id")

    # For development: Run synchronously without Celery
    # This bypasses the async task queue when Redis/Celery isn't available
    from django.conf import settings
    
    if getattr(settings, 'CELERY_TASK_ALWAYS_EAGER', True):
        # Synchronous mode - call task directly without queue
        try:
            org = get_org_from_request(request)
            
            job = Job.objects.create(
                organization=org,
                user=request.user,
                job_type='generate_output_excel',
                status='queued',
                current_step="Processing...",
            )
            
            # Store inputs in job result
            job.result = {
                'fetched_items': fetched,
                'qty_map': item_qtys,
                'unit_map': item_units,
                'work_name': work_name,
                'work_type': work_type,
                'grand_total': grand_total,
                'excess_tp_enabled': excess_tp_enabled,
                'excess_tp_percent': excess_tp_percent,
                'ls_special_enabled': ls_special_enabled,
                'ls_special_name': ls_special_name,
                'ls_special_amount': ls_special_amount,
                'deduct_old_material': deduct_old_material,
                'backend_id': selected_backend_id,
            }
            job.save()
            
            # Call task function directly (synchronous, no Celery)
            from core.tasks import generate_output_excel
            result = generate_output_excel.apply(args=(
                job.id,
                category,
                json.dumps(item_qtys),
                json.dumps(item_units),
                work_name,
                work_type,
                grand_total,
                excess_tp_percent,
                ls_special_name,
                ls_special_amount,
                deduct_old_material,
                selected_backend_id,
            )).get()
            
            # Refresh job to get updated result
            job.refresh_from_db()
            
            if job.status == 'completed' and job.result.get('output_file_id'):
                # Redirect to file download
                from core.models import OutputFile
                try:
                    output_file = OutputFile.objects.get(id=job.result['output_file_id'])
                    from django.http import FileResponse
                    import os
                    
                    if output_file.file and os.path.exists(output_file.file.path):
                        response = FileResponse(
                            open(output_file.file.path, 'rb'),
                            as_attachment=True,
                            filename=output_file.filename or f"{category}_output.xlsx",
                        )
                        return response
                except OutputFile.DoesNotExist:
                    pass
            
            # Fallback - return JSON with job status
            return JsonResponse({
                'job_id': job.id,
                'status': job.status,
                'message': job.current_step or 'Processing complete',
                'error': job.error_message if job.status == 'failed' else None,
            })
                
        except Exception as e:
            logger.error(f"Failed to generate output Excel: {e}")
            return JsonResponse({"error": str(e)}, status=500)
    
    # Async mode with Celery (production)
    try:
        org = get_org_from_request(request)
        
        job = Job.objects.create(
            organization=org,
            user=request.user,
            job_type='generate_output_excel',
            status='queued',
            current_step="Queued for processing",
        )
        
        # Store inputs in job result temporarily
        job.result = {
            'fetched_items': fetched,
            'qty_map': item_qtys,
            'unit_map': item_units,
            'work_name': work_name,
            'work_type': work_type,
            'grand_total': grand_total,
            'excess_tp_enabled': excess_tp_enabled,
            'excess_tp_percent': excess_tp_percent,
            'ls_special_enabled': ls_special_enabled,
            'ls_special_name': ls_special_name,
            'ls_special_amount': ls_special_amount,
            'deduct_old_material': deduct_old_material,
            'backend_id': selected_backend_id,
        }
        job.save()
        
        # Enqueue async task
        from core.tasks import generate_output_excel
        task = generate_output_excel.delay(
            job.id,
            category,
            json.dumps(item_qtys),
            json.dumps(item_units),
            work_name,
            work_type,
            grand_total,
            excess_tp_percent,
            ls_special_name,
            ls_special_amount,
            deduct_old_material,
            selected_backend_id,
        )
        
        job.celery_task_id = task.id
        job.save()
        
        return JsonResponse({
            'job_id': job.id,
            'status_url': reverse('job_status', args=[job.id]),
            'message': f'Generating {category} output estimate. Please wait...'
        })
        
    except Exception as e:
        logger.error(f"Failed to enqueue output Excel task: {e}")
        return JsonResponse({"error": str(e)}, status=500)






# -----------------------
# CLEAR OUTPUT
# -----------------------
@login_required(login_url='login')
def clear_output(request, category):
    request.session["fetched_items"] = []
    request.session["qty_map"] = {}
    request.session["unit_map"] = {}
    request.session["work_name"] = ""
    request.session["grand_total"] = ""

    group = request.GET.get("group")
    if group:
        return redirect("datas_items", category=category, group=group)

    return redirect("datas_groups", category=category)


# -----------------------
# SAVE QTY MAP (AJAX endpoint for preserving quantities during navigation)
# -----------------------
@login_required(login_url='login')
def save_qty_map(request, category):
    """
    AJAX endpoint to save quantity map, unit map, and grand total to session.
    Called before navigation to preserve entered quantities and units.
    """
    if request.method == "POST":
        try:
            qty_map_str = request.POST.get("qty_map", "")
            unit_map_str = request.POST.get("unit_map", "")
            grand_total_str = request.POST.get("grand_total", "")
            work_name = request.POST.get("work_name", "")
            excess_tp_percent = request.POST.get("excess_tp_percent", "")
            ls_special_name = request.POST.get("ls_special_name", "")
            ls_special_amount = request.POST.get("ls_special_amount", "")
            deduct_old_material = request.POST.get("deduct_old_material", "")
            work_type = request.POST.get("work_type", "")
            
            if qty_map_str:
                try:
                    qty_map = json.loads(qty_map_str)
                    if isinstance(qty_map, dict):
                        request.session["qty_map"] = qty_map
                except json.JSONDecodeError:
                    pass
            
            if unit_map_str:
                try:
                    unit_map = json.loads(unit_map_str)
                    if isinstance(unit_map, dict):
                        request.session["unit_map"] = unit_map
                except json.JSONDecodeError:
                    pass
            
            if grand_total_str:
                request.session["grand_total"] = grand_total_str
            
            if work_name:
                request.session["work_name"] = work_name
            
            if excess_tp_percent:
                request.session["excess_tp_percent"] = excess_tp_percent
            
            if ls_special_name:
                request.session["ls_special_name"] = ls_special_name
            
            if ls_special_amount:
                request.session["ls_special_amount"] = ls_special_amount
            
            if deduct_old_material:
                request.session["deduct_old_material"] = deduct_old_material
            
            if work_type:
                request.session["work_type"] = work_type
            
            return JsonResponse({"status": "ok"})
        except Exception as e:
            return JsonResponse({"status": "error", "message": str(e)}, status=400)
    
    return JsonResponse({"status": "error", "message": "POST required"}, status=405)


@login_required(login_url='login')
@org_required
def save_project(request, category):
    org = get_org_from_request(request)
    fetched = request.session.get("fetched_items", []) or []
    if not fetched:
        return redirect("datas_groups", category=category)

    if request.method == "POST":
        qty_map = {}
        qty_map_str = request.POST.get("qty_map", "")
        if qty_map_str:
            try:
                raw = json.loads(qty_map_str)
                if isinstance(raw, dict):
                    for k, v in raw.items():
                        try:
                            num = float(v)
                        except (TypeError, ValueError):
                            continue
                        qty_map[str(k)] = num
            except json.JSONDecodeError:
                qty_map = {}

        work_name = (request.POST.get("work_name_hidden") or "").strip()

        request.session["qty_map"] = qty_map
        request.session["work_name"] = work_name

        project_name = request.POST.get("project_name") or request.session.get("current_project_name")

        if project_name:
            project, created = Project.objects.get_or_create(organization=org, name=project_name)
            project.category = category

            data = {
                "items": fetched,
                "qty_map": qty_map,
                "work_name": work_name,
            }
            project.set_items(data)
            project.save()

            request.session["current_project_name"] = project.name

            return redirect("my_projects")

    return redirect("datas_groups", category=category)


@org_required
def load_project(request, project_id):
    org = get_org_from_request(request)
    project = get_object_or_404(Project, id=project_id, organization=org)
    stored = project.get_items()

    if isinstance(stored, dict):
        fetched = stored.get("items", []) or []
        qty_map = stored.get("qty_map", {}) or {}
        work_name = stored.get("work_name", "") or ""
    else:
        fetched = stored or []
        qty_map = {}
        work_name = ""

    request.session["fetched_items"] = fetched
    request.session["qty_map"] = qty_map
    request.session["work_name"] = work_name
    request.session["current_project_name"] = project.name

    return redirect("datas_groups", category=project.category)


@login_required(login_url='login')
def download_estimate(request, category):
    """
    Async estimate Excel generation endpoint.
    
    Uses session['fetched_items'] to generate estimate-only workbook asynchronously.
    Returns JSON with job_id and status_url for polling.
    """
    fetched = request.session.get("fetched_items", [])
    if not fetched:
        return JsonResponse({"error": "No items selected"}, status=400)
    
    # Get selected backend ID from session (for multi-backend support)
    selected_backend_id = request.session.get("selected_backend_id")

    try:
        org = get_org_from_request(request)
        
        job = Job.objects.create(
            organization=org,
            user=request.user,
            job_type='generate_estimate_excel',
            status='queued',
            current_step="Queued for processing",
        )
        
        # Store fetched items and backend_id in result temporarily
        job.result = {
            'fetched_items': fetched,
            'backend_id': selected_backend_id,
        }
        job.save()
        
        # Enqueue async task
        from core.tasks import generate_estimate_excel
        task = generate_estimate_excel.delay(
            job.id,
            category,
            json.dumps(fetched),
            selected_backend_id,
        )
        
        job.celery_task_id = task.id
        job.save()
        
        return JsonResponse({
            'job_id': job.id,
            'status_url': reverse('job_status', args=[job.id]),
            'message': f'Generating {category} estimate. Please wait...'
        })
        
    except Exception as e:
        logger.error(f"Failed to enqueue estimate Excel task: {e}")
        return JsonResponse({"error": str(e)}, status=500)



def build_fetched_details(category, base_dir, fetched_names):
    """
    Build a list of dicts for the UI side panel:
    [
      {"slno": 1, "name": "Item 1", "rate": 123.45, "rate_display": "123.45"},
      ...
    ]
    Rates are read from the backend Excel (last non-empty J cell of each block).
    """
    if not fetched_names:
        return []

    try:
        items_list, groups_map, _, ws_data, filepath = load_backend(category, base_dir)
    except Exception:
        return [
            {"slno": idx, "name": name, "rate": "", "rate_display": ""}
            for idx, name in enumerate(fetched_names, start=1)
        ]

    block_map = {it["name"]: (it["start_row"], it["end_row"]) for it in items_list}

    try:
        wb_vals = load_workbook(filepath, data_only=True)
        ws_vals = wb_vals["Master Datas"]
    except Exception:
        wb_vals = None
        ws_vals = None

    details = []
    for idx, name in enumerate(fetched_names, start=1):
        rate_value = ""
        rate_display = ""
        block = block_map.get(name)

        if block and ws_vals is not None:
            start_row, end_row = block
            for r in range(end_row, start_row - 1, -1):
                val = ws_vals.cell(row=r, column=10).value
                if val not in (None, ""):
                    rate_display = val
                    try:
                        rate_value = float(val)
                    except Exception:
                        rate_value = ""
                    break

        details.append(
            {
                "slno": idx,
                "name": name,
                "rate": rate_value,
                "rate_display": rate_display,
            }
        )

    return details


@login_required(login_url='login')
@require_POST
def toggle_item(request, category, group):
    """
    AJAX endpoint: add/remove an item from session['fetched_items']
    based on checkbox state, and return updated side panel HTML.
    """
    item = request.POST.get("item", "").strip()
    checked = request.POST.get("checked") == "true"

    fetched = request.session.get("fetched_items", [])

    if checked:
        if item and item not in fetched:
            fetched.append(item)
    else:
        fetched = [x for x in fetched if x != item]

    request.session["fetched_items"] = fetched

    html = render_to_string(
        "core/sidebar_output.html",
        {"category": category, "fetched": fetched},
        request=request,
    )

    return JsonResponse({"html": html})


@login_required(login_url='login')
def new_project(request):
    """
    Clear current selection and start from scratch.
    """
    request.session["fetched_items"] = []
    request.session["qty_map"] = {}
    request.session["work_name"] = ""
    request.session["current_project_name"] = None
    return redirect("datas")


@org_required
def delete_project(request, project_id):
    org = get_org_from_request(request)
    project = get_object_or_404(Project, id=project_id, organization=org)
    project.delete()
    return redirect("my_projects")

def _extract_value_part_from_line(s: str) -> str:
    """
    For a line like 'Name of the work : Construction of XYZ',
    returns 'Construction of XYZ'.
    Also handles OCR artifacts like leading numbers/asterisks.
    """
    s = (s or "").strip()
    # Remove leading line numbers like "1 ", "4 ", etc.
    s = re.sub(r'^\d+\s+', '', s)
    
    if ":" in s:
        val = s.split(":", 1)[1].strip()
        # Clean OCR artifacts
        val = val.lstrip('*').strip()
        return val
    if " - " in s:
        return s.split(" - ", 1)[1].strip()
    return ""


def _collect_multiline_value(lines, start_idx, max_lines=4):
    """
    Collect value that spans multiple lines (common in OCR output).
    Stops when hitting a new numbered line or label pattern.
    """
    collected = []
    for i in range(start_idx, min(start_idx + max_lines, len(lines))):
        line = str(lines[i]).strip()
        if not line:
            continue
        # Stop if we hit a new numbered item (like "2 Ref to...")
        if re.match(r'^\d+\s+[A-Z]', line):
            break
        # Stop if we hit another label with colon
        if re.search(r'^[A-Za-z\s]+\s*:\s*$', line):
            break
        collected.append(line)
    return " ".join(collected).strip()


def _extract_labels_from_lines(lines):
    """
    Fuzzy-read important labels from list of text lines.
    Handles OCR output from scanned documents like Forwarding Slips, Estimates, Bills.
    """
    labels = {
        "name_of_work": "",
        "agreement": "",
        "admin_sanction": "",
        "admin_sanction_amount": "",
        "tech_sanction": "",
        "tech_sanction_amount": "",
        "agency": "",
        "contractor": "",
        "contractor_address": "",
        "mb_details": "",
        "tp_details": "",
        "cc_header": "",
        "amount": "",
        "estimate_amount": "",
        "bond_no": "",
        "nit_no": "",
        "tender_premium": "",
        "period_of_completion": "",
        "earnest_money": "",
        "security_deposit": "",
        "work_order_no": "",
        "work_order_date": "",
        "date_of_commencement": "",
        "date_of_completion": "",
    }

    for idx, raw in enumerate(lines):
        s = str(raw or "").strip()
        if not s:
            continue
        low = s.lower()

        # i. Name of the work (various formats)
        if not labels["name_of_work"]:
            # Standard patterns - "Name of Work:" or "Name of the Work:"
            if "name of work" in low or "name of the work" in low:
                val = _extract_value_part_from_line(s)
                # If value is empty or too short, collect from next lines
                if not val or len(val) < 5:
                    val = _collect_multiline_value(lines, idx + 1, max_lines=5)
                if val:
                    # Clean up OCR artifacts like "ast :"
                    val = re.sub(r'^[a-z]+\s*:\s*', '', val, flags=re.I)
                    labels["name_of_work"] = val
                    continue
            
            # Look for "for the work of" pattern (common in Forwarding Slips)
            if "for the work of" in low:
                # Extract text after "for the work of"
                match = re.search(r'for the work of\s*["\']?\s*(.+?)(?:["\']?\s*duly|$)', s, re.I)
                if match:
                    val = match.group(1).strip().strip('"\'')
                    if val and len(val) > 5:
                        labels["name_of_work"] = val
                        continue
                # Value might be empty here (""), collect from next lines
                if idx + 1 < len(lines):
                    val = _collect_multiline_value(lines, idx + 1, max_lines=3)
                    if val and len(val) > 10:
                        labels["name_of_work"] = val
                        continue
            
            # Look for work description keywords (maintenance, servicing, generator, etc.)
            work_keywords = ['maintenance', 'servicing', 'generator', 'diesel', 'kva', 'annual']
            if any(kw in low for kw in work_keywords) and len(s) > 30:
                # Skip if line starts with fragment like "rk:" or is too short
                if not re.match(r'^[a-z]{1,3}:', low) and not low.startswith('rk'):
                    # Collect this and following lines
                    work_parts = [s]
                    for j in range(idx + 1, min(idx + 5, len(lines))):
                        next_line = str(lines[j]).strip()
                        if not next_line or len(next_line) < 3:
                            continue
                        next_low = next_line.lower()
                        if any(x in next_low for x in ['sanction', 'contractor', 'nit', 'accord', 'permit', 'memo']):
                            break
                        if len(next_line) > 5:
                            work_parts.append(next_line)
                    full_work = " ".join(work_parts)
                    # Add "Providing A" prefix if work starts with "nual" (fragment of "Annual")
                    if full_work.lower().startswith('nual'):
                        full_work = "Providing A" + full_work
                    labels["name_of_work"] = full_work
                    continue
            
            # Look for "Providing" or "Supply" at start of line (common work name patterns)
            # Also handle OCR fragments like "viding" (from "Providing")
            if (low.startswith("providing") or low.startswith("supply") or low.startswith("construction") or
                low.startswith("viding") or "viding annual" in low or "viding " in low):
                # Skip fragments like "rk: Providing"
                if re.match(r'^[a-z]{1,3}:', low):
                    continue
                # This line and following lines might be the work name - collect them
                if len(s) > 10:
                    # Collect multiple lines for fragmented OCR
                    work_parts = [s]
                    for j in range(idx + 1, min(idx + 8, len(lines))):
                        next_line = str(lines[j]).strip()
                        if not next_line or len(next_line) < 3:
                            continue
                        next_low = next_line.lower()
                        # Stop if we hit a new section
                        if any(x in next_low for x in ['sanction', 'contractor', 'nit', 'tender', 'amount', 'period']):
                            break
                        # Stop if line looks like a label
                        if ':' in next_line and len(next_line.split(':')[0]) < 25:
                            break
                        work_parts.append(next_line)
                    # Join and clean up the work name
                    full_work = " ".join(work_parts)
                    # Add "Pro" prefix if it starts with "viding"
                    if full_work.lower().startswith("viding"):
                        full_work = "Pro" + full_work
                    labels["name_of_work"] = full_work
                    continue
            
            # Check if previous line was just "Name of Work" label (table format)
            if idx > 0:
                prev_line = str(lines[idx - 1]).strip().lower()
                if prev_line in ("name of work", "name of the work", "name of work:"):
                    if len(s) > 10 and not any(x in low for x in ['sanction', 'amount', 'contractor', 'nit']):
                        labels["name_of_work"] = s
                        continue
                    continue

        # Bond No / Agreement (from header like "Suppl.Agreement Bond No. 30/2024-2025")
        if not labels["agreement"]:
            if "bond no" in low or "agreement bond" in low or "agreement" in low:
                # Try to extract the number pattern like "30/2024-2025"
                match = re.search(r'(\d+\s*[\/\-]\s*\d{4}\s*[\/\-]?\s*\d*)', s)
                if match:
                    labels["agreement"] = match.group(1).strip()
                    continue
                # Also check for Agreement.No pattern
                match = re.search(r'Agreement\.?No\.?\s*[:\.]?\s*([^\s,]+)', s, re.I)
                if match:
                    labels["agreement"] = match.group(1).strip()
                    continue

        # Admin Sanction / Sanctioned Estimate (TA.No pattern or Ref. to Administrative sanction)
        if not labels["admin_sanction"]:
            if "administrative sanction" in low or "admin sanction" in low or ("a)" in low and "sanction" in low):
                # Try Memo.No pattern first (capture full reference with date)
                memo_match = re.search(r'(Memo\.?\s*No\.?\s*[A-Za-z0-9\/\-\.\s]+(?:dt[:\.\s]*[\d\.\-\/]+)?)', s, re.I)
                if memo_match:
                    labels["admin_sanction"] = memo_match.group(1).strip()
                else:
                    # Try to extract value after colon
                    val = _extract_value_part_from_line(s)
                    if val:
                        labels["admin_sanction"] = val
                        # If value seems incomplete, collect from next lines
                        if len(val) < 20 and idx + 1 < len(lines):
                            next_val = _collect_multiline_value(lines, idx + 1, max_lines=2)
                            if next_val:
                                labels["admin_sanction"] = val + " " + next_val
                # Also try to get amount if present on this or next line
                amount_match = re.findall(r'Rs\.?\s*([\d,]+(?:\.\d+)?)', s, re.I)
                if amount_match and not labels["admin_sanction_amount"]:
                    labels["admin_sanction_amount"] = "Rs." + amount_match[-1] + "/-"
                continue
        
        # Check for "b)Administrative sanction Amount" or "Technical Sanction Amount" patterns
        if not labels["admin_sanction_amount"]:
            if ("admin" in low or "a)" in low) and "sanction" in low and "amount" in low:
                amount_match = re.findall(r'Rs\.?\s*([\d,]+(?:\.\d+)?)', s, re.I)
                if amount_match:
                    labels["admin_sanction_amount"] = "Rs." + amount_match[-1] + "/-"
                else:
                    # Value might be on next line
                    if idx + 1 < len(lines):
                        next_line = str(lines[idx + 1]).strip()
                        amt_match = re.findall(r'Rs\.?\s*([\d,]+(?:\.\d+)?)', next_line, re.I)
                        if amt_match:
                            labels["admin_sanction_amount"] = "Rs." + amt_match[-1] + "/-"
                continue
        
        # Tech Sanction Amount specifically
        if not labels["tech_sanction_amount"]:
            if ("tech" in low or "b)" in low) and "sanction" in low and "amount" in low:
                amount_match = re.findall(r'Rs\.?\s*([\d,]+(?:\.\d+)?)', s, re.I)
                if amount_match:
                    labels["tech_sanction_amount"] = "Rs." + amount_match[-1] + "/-"
                else:
                    # Value might be on next line
                    if idx + 1 < len(lines):
                        next_line = str(lines[idx + 1]).strip()
                        amt_match = re.findall(r'Rs\.?\s*([\d,]+(?:\.\d+)?)', next_line, re.I)
                        if amt_match:
                            labels["tech_sanction_amount"] = "Rs." + amt_match[-1] + "/-"
                continue

        # Tech Sanction - handle DR.NO pattern and value on next line
        if not labels["tech_sanction"]:
            if ("tech" in low and "sanc" in low) or "b)technical" in low or "technical sanction" in low:
                # Try DR.NO pattern first
                dr_match = re.search(r'(DR\.?\s*NO\.?\s*[\d\/\-]+(?:\s*,?\s*(?:Dt|dt)[:\.]?\s*[\d\.\/\-]+)?)', s, re.I)
                if dr_match:
                    labels["tech_sanction"] = dr_match.group(1).strip()
                else:
                    val = _extract_value_part_from_line(s)
                    if val:
                        labels["tech_sanction"] = val
                    elif idx + 1 < len(lines):
                        # Value might be on next line
                        next_line = str(lines[idx + 1]).strip()
                        dr_match = re.search(r'(DR\.?\s*NO\.?\s*[\d\/\-]+(?:\s*,?\s*(?:Dt|dt)[:\.]?\s*[\d\.\/\-]+)?)', next_line, re.I)
                        if dr_match:
                            labels["tech_sanction"] = dr_match.group(1).strip()
                        elif next_line and len(next_line) > 3:
                            labels["tech_sanction"] = next_line
                # Get amount if present
                amount_match = re.findall(r'Rs\.?\s*([\d,]+(?:\.\d+)?)', s, re.I)
                if amount_match and not labels["tech_sanction_amount"]:
                    labels["tech_sanction_amount"] = "Rs." + amount_match[-1] + "/-"
                continue

        # Contractor / Agency - treat as same thing
        # "Name of the Contractor" = "Name of the Agency"
        if not labels["agency"]:
            if "contractor" in low or "agency" in low or "name of the contractor" in low:
                # Look for M/s. pattern first (common for contractor names) - case insensitive
                match = re.search(r'M/s\.?\s*([A-Za-z][A-Za-z0-9\s\.\&\-]+)', s, re.I)
                if match:
                    labels["agency"] = "M/s." + match.group(1).strip()
                    continue
                # Otherwise use value after colon
                val = _extract_value_part_from_line(s)
                if val and len(val) > 2 and not val.lower().startswith('yes') and not val.lower().startswith('no'):
                    # Skip if value looks like an address (starts with number pattern like 1-6-620)
                    if re.match(r'^\d+[\-\/]\d+', val):
                        # This is an address, not contractor name
                        labels["contractor_address"] = val
                        continue
                    # Check if val itself contains M/s pattern
                    ms_match = re.search(r'(?:M/s|Mis|M\.s)\.?\s*(.+)', val, re.I)
                    if ms_match:
                        labels["agency"] = "M/s." + ms_match.group(1).strip()
                    else:
                        labels["agency"] = val
                    continue
                # Value might be on next line
                if idx + 1 < len(lines):
                    next_line = str(lines[idx + 1]).strip()
                    # Skip if next line looks like address
                    if re.match(r'^\d+[\-\/]\d+', next_line):
                        labels["contractor_address"] = next_line
                        continue
                    match = re.search(r'(?:M/s|Mis|M\.s)\.?\s*([A-Za-z][A-Za-z0-9\s\.\&\-]+)', next_line, re.I)
                    if match:
                        labels["agency"] = "M/s." + match.group(1).strip()
                    if match:
                        labels["agency"] = "M/s." + match.group(1).strip()
                        continue
                    elif next_line and len(next_line) > 5:
                        labels["agency"] = next_line
                        continue
        
        # Also detect standalone M/s. lines (contractor name without label)
        if not labels["agency"]:
            # Match M/s patterns including OCR variations like "M/s.", "Mis.", "M/S", etc.
            if low.startswith("m/s") or low.startswith("m/s.") or low.startswith("mis.") or low.startswith("m.s"):
                # This line is likely a contractor name
                match = re.search(r'(?:M/s|Mis|M\.s)\.?\s*(.+)', s, re.I)
                if match:
                    labels["agency"] = "M/s." + match.group(1).strip()
                    continue

        # Also check specifically for "Name of the Contractor" with value on same line
        if not labels["agency"]:
            if "name of the contractor" in low or "name of contractor" in low or "contractor name" in low:
                val = _extract_value_part_from_line(s)
                if val:
                    ms_match = re.search(r'(?:M/s|Mis|M\.s)\.?\s*(.+)', val, re.I)
                    if ms_match:
                        labels["agency"] = "M/s." + ms_match.group(1).strip()
                    else:
                        labels["agency"] = val
                    continue
        
        # Contractor Address - separate field (look for address patterns)
        if not labels["contractor_address"]:
            # Detect address-like patterns: "1-6-620/1 Near..." or contains locality names
            if "contractor" in low and "address" in low:
                val = _extract_value_part_from_line(s)
                if val:
                    labels["contractor_address"] = val
                    continue
            # Standalone address line pattern (e.g., "1-6-620/1 Near Ramalayam...")
            address_match = re.match(r'^(\d+[\-\/]\d+[\-\/]?\d*\s+(?:Near|Opp|Behind|At|H\.?No)?\s*.+)', s, re.I)
            if address_match and not labels["contractor_address"]:
                # This looks like an address - but only capture if we already have contractor
                if labels["agency"]:
                    labels["contractor_address"] = s
                    continue
                val = _extract_value_part_from_line(s)
                if val:
                    labels["contractor_address"] = val
                    # Address often spans multiple lines - collect more
                    if idx + 1 < len(lines):
                        addr_parts = [val]
                        for j in range(idx + 1, min(idx + 4, len(lines))):
                            addr_line = str(lines[j]).strip()
                            # Stop if we hit a new label/section
                            if ':' in addr_line[:20] or re.match(r'^\\d+[\\)\\.]', addr_line):
                                break
                            if any(x in addr_line.lower() for x in ['estimate', 'period', 'sanction', 'nit']):
                                break
                            if addr_line and len(addr_line) > 3:
                                addr_parts.append(addr_line)
                        labels["contractor_address"] = " ".join(addr_parts)
                    continue
                # Value might be on next line
                if idx + 1 < len(lines):
                    next_line = str(lines[idx + 1]).strip()
                    if next_line and len(next_line) > 10:
                        labels["contractor_address"] = next_line
                        # Collect more address lines
                        addr_parts = [next_line]
                        for j in range(idx + 2, min(idx + 5, len(lines))):
                            addr_line = str(lines[j]).strip()
                            if ':' in addr_line[:20] or re.match(r'^\\d+[\\)\\.]', addr_line):
                                break
                            if any(x in addr_line.lower() for x in ['estimate', 'period', 'sanction', 'nit']):
                                break
                            if addr_line and len(addr_line) > 3:
                                addr_parts.append(addr_line)
                        labels["contractor_address"] = " ".join(addr_parts)
                        continue

        # Amount of Estimate / Estimate Amount (value often on next line like "*2,04,798/-" or "Rs. 102820.48")
        if not labels["estimate_amount"]:
            if "amount of estimate" in low or "estimate amount" in low or ("estimate" in low and "amount" in low):
                val = _extract_value_part_from_line(s)
                # Check for Rs. pattern in the value
                if val:
                    rs_match = re.search(r'Rs\\.?\\s*([\\d,]+(?:\\.\\d+)?)', val, re.I)
                    if rs_match:
                        labels["estimate_amount"] = "Rs. " + rs_match.group(1)
                    else:
                        labels["estimate_amount"] = val.lstrip('*').rstrip('/-').strip()
                    continue
                if not val and idx + 1 < len(lines):
                    # Value is on next line - extract number
                    next_line = str(lines[idx + 1]).strip()
                    # Match patterns like "Rs. 102820.48" or "*2,04,798/-" or "2,04,798"
                    rs_match = re.search(r'Rs\\.?\\s*([\\d,]+(?:\\.\\d+)?)', next_line, re.I)
                    if rs_match:
                        labels["estimate_amount"] = "Rs. " + rs_match.group(1)
                        continue
                    match = re.search(r'[\\*]?([\\d,]+(?:\\.\\d+)?)', next_line)
                    if match:
                        val = match.group(1)
                        labels["estimate_amount"] = "Rs. " + val.lstrip('*').rstrip('/-').strip()
                        continue

        # MB Details - Measurement Book details
        if not labels["mb_details"]:
            # Various MB patterns
            if "mb" in low and ("no" in low or "details" in low or "page" in low or "sl" in low):
                val = _extract_value_part_from_line(s)
                if val:
                    labels["mb_details"] = val
                    continue
            # Look for patterns like "MB No. 123, Page 45-67"
            if "mb" in low:
                match = re.search(r'MB\s*(?:No\.?|No)?\s*:?\s*(\d+)[,\s]*(?:Page|Pg\.?|P)?\s*(?:No\.?)?\s*:?\s*(\d+(?:\s*[-to]+\s*\d+)?)', s, re.I)
                if match:
                    mb_no = match.group(1)
                    pages = match.group(2) if match.group(2) else ""
                    labels["mb_details"] = f"MB No. {mb_no}, Page {pages}" if pages else f"MB No. {mb_no}"
                    continue
            # Look for "Sl. No." or "Serial No." patterns (often have MB reference)
            if ("sl." in low or "serial" in low) and "no" in low:
                val = _extract_value_part_from_line(s)
                if val:
                    labels["mb_details"] = val
                    continue

        # T.P Details (like "T.P @7.86 % Less")
        if not labels["tp_details"]:
            if "t.p" in low or "tp@" in low or "tp %" in low:
                # Look for percentage pattern
                match = re.search(r'(T\.?P\.?\s*@?\s*[\d\.]+\s*%?\s*(?:less|more)?)', s, re.I)
                if match:
                    labels["tp_details"] = match.group(1).strip()
                else:
                    val = _extract_value_part_from_line(s)
                    if val:
                        labels["tp_details"] = val
                continue

        # CC Header - Bill Type (L.S. Bill, 1st Part Bill, Final Bill, etc.)
        if not labels["cc_header"]:
            # Look for "Bill of L.S." or "L.S. Bill" pattern
            if "bill of l.s" in low or "l.s. bill" in low or "l.s bill" in low:
                labels["cc_header"] = "L.S. Bill"
                continue
            # Look for Part Bills
            part_match = re.search(r'(\d+(?:st|nd|rd|th)?\s*part\s*bill)', low, re.I)
            if part_match:
                labels["cc_header"] = part_match.group(1).title()
                continue
            # Look for Final Bill
            if "final bill" in low or "final part" in low:
                labels["cc_header"] = "Final Bill"
                continue
            # Look for generic CC Bill pattern
            if "cc" in low and "bill" in low:
                labels["cc_header"] = s.strip()
                continue

        # NIT No. / Tender No.
        if not labels["nit_no"]:
            if "nit" in low or ("tender" in low and "no" in low):
                # Look for NIT number pattern - capture full reference including date
                # Pattern like: NIT.No14/EE/GI/DB/HD/TA1/2025-26, Dt:04.04.2025
                match = re.search(r'(NIT\.?\s*No\.?\s*[A-Za-z0-9\/\-]+(?:\s*,?\s*Dt[:\.]?\s*[\d\.\/\-]+)?)', s, re.I)
                if match:
                    labels["nit_no"] = match.group(1).strip()
                    continue
                val = _extract_value_part_from_line(s)
                if val:
                    labels["nit_no"] = val
                    # Value might continue on next line (date part)
                    if idx + 1 < len(lines):
                        next_line = str(lines[idx + 1]).strip()
                        if next_line.lower().startswith('dt') or re.match(r'^[\d\.\-\/]+', next_line):
                            labels["nit_no"] = val + ", " + next_line
                    continue

        # Tender Premium
        if not labels["tender_premium"]:
            if "tender premium" in low or ("premium" in low and ("%" in s or "er" in low)):
                # Look for percentage pattern
                match = re.search(r'([\d\.]+\s*%?\s*(?:ER|less|more|above|below)?)', s, re.I)
                if match:
                    labels["tender_premium"] = match.group(1).strip()
                    continue
                val = _extract_value_part_from_line(s)
                if val:
                    labels["tender_premium"] = val
                    continue

        # Period of Completion
        if not labels["period_of_completion"]:
            if "period of completion" in low or "completion period" in low or "stipulated period" in low or "period of compl" in low:
                # Look for months/days pattern on this line
                match = re.search(r'(\d+\s*(?:months?|days?|years?|weeks?))', s, re.I)
                if match:
                    labels["period_of_completion"] = match.group(1).strip()
                    continue
                val = _extract_value_part_from_line(s)
                if val:
                    labels["period_of_completion"] = val
                    continue
                # Value might be on next line
                if idx + 1 < len(lines):
                    next_line = str(lines[idx + 1]).strip()
                    match = re.search(r'(\d+\s*(?:months?|days?|years?|weeks?))', next_line, re.I)
                    if match:
                        labels["period_of_completion"] = match.group(1).strip()
                        continue

        # Earnest Money Deposit
        if not labels["earnest_money"]:
            if "earnest money" in low or "emd" in low:
                nums = re.findall(r'Rs\.?\s*([\d,]+(?:\.\d+)?)', s, re.I)
                if nums:
                    labels["earnest_money"] = nums[0]
                    continue
                val = _extract_value_part_from_line(s)
                if val:
                    labels["earnest_money"] = val
                    continue

        # Security Deposit
        if not labels["security_deposit"]:
            if "security deposit" in low or "s.d" in low:
                nums = re.findall(r'Rs\.?\s*([\d,]+(?:\.\d+)?)', s, re.I)
                if nums:
                    labels["security_deposit"] = nums[0]
                    continue
                val = _extract_value_part_from_line(s)
                if val:
                    labels["security_deposit"] = val
                    continue

        # Work Order No. / Date
        if not labels["work_order_no"]:
            if "work order" in low:
                match = re.search(r'No\.?\s*[:\.]?\s*([\w\/\-]+)', s, re.I)
                if match:
                    labels["work_order_no"] = match.group(1).strip()
                # Also look for date
                date_match = re.search(r'(?:Dt|Date|dated)[:\.]?\s*([\d\.\/\-]+)', s, re.I)
                if date_match:
                    labels["work_order_date"] = date_match.group(1).strip()
                continue

        # Date of Commencement
        if not labels["date_of_commencement"]:
            if "commencement" in low or "start date" in low:
                date_match = re.search(r'([\d]{1,2}[\.\-\/][\d]{1,2}[\.\-\/][\d]{2,4})', s)
                if date_match:
                    labels["date_of_commencement"] = date_match.group(1)
                    continue
                val = _extract_value_part_from_line(s)
                if val:
                    labels["date_of_commencement"] = val
                    continue

        # Date of Completion
        if not labels["date_of_completion"]:
            if "date of completion" in low or "completion date" in low:
                date_match = re.search(r'([\d]{1,2}[\.\-\/][\d]{1,2}[\.\-\/][\d]{2,4})', s)
                if date_match:
                    labels["date_of_completion"] = date_match.group(1)
                    continue
                val = _extract_value_part_from_line(s)
                if val:
                    labels["date_of_completion"] = val
                    continue

        # Total / Amount (with Rs. prefix) - must be a significant amount
        # Prefer bill amounts (1,10,000) over estimate amounts (102820)
        if not labels["amount"]:
            # Look for bill amount, total amount, sanction amount patterns
            if ("total" in low or "bill amount" in low or "amount of bill" in low or 
                "sanction amount" in low or "approximate value" in low):
                nums = re.findall(r'Rs\.?\s*([\d,]+(?:\.\d+)?)', s, re.I)
                if nums:
                    # Filter out small amounts like security deposits, stamps etc
                    for num in reversed(nums):
                        clean_num = num.replace(',', '')
                        try:
                            val = float(clean_num)
                            # Prefer amounts like 1,10,000 (>100000) over amounts like 102820.48
                            if val >= 100000:
                                labels["amount"] = num
                                break
                        except:
                            labels["amount"] = num
                            break
                    if labels["amount"]:
                        continue
            
            # Look for standalone Rs. amounts on sanction amount lines
            if "sanction" in low and "amount" in low:
                nums = re.findall(r'Rs\.?\s*([\d,]+(?:\.\d+)?)/?-?', s, re.I)
                for num in nums:
                    clean_num = num.replace(',', '')
                    try:
                        if float(clean_num) >= 100000:
                            labels["amount"] = num
                            break
                    except:
                        pass

    # Use estimate_amount as amount if no amount found
    if not labels["amount"] and labels["estimate_amount"]:
        labels["amount"] = labels["estimate_amount"]
    
    # Use contractor as agency if agency not found
    if not labels["agency"] and labels["contractor"]:
        labels["agency"] = labels["contractor"]
    
    # FIX COMMON OCR SPELLING ERRORS
    ocr_corrections = {
        # Maintenance
        "mala tenance": "maintenance",
        "malatenance": "maintenance", 
        "maintenace": "maintenance",
        "maintainance": "maintenance",
        "maintanance": "maintenance",
        # Monthly
        "mouthly": "monthly",
        "monthely": "monthly",
        "montly": "monthly",
        # Servicing
        "servlcing": "servicing",
        "serviclng": "servicing",
        # Annual
        "annuai": "annual",
        "annuat": "annual",
        # Generator
        "generater": "generator",
        "genertor": "generator",
        # Providing
        "provlding": "providing",
        "providlng": "providing",
        # Located
        "Iocated": "located",
        # Guest house
        "bow se": "house",
        "bowse": "house",
        # Complex
        "complec": "complex",
        "compex": "complex",
        # Begumpet
        "Bite pad": "Begumpet",
        "Bitépad": "Begumpet",
        # Hyderabad
        "Hydera bad": "Hyderabad",
        # Common OCR errors
        "sanctlon": "sanction",
        "sanchon": "sanction",
        "Techical": "Technical",
        "Addinistrative": "Administrative",
        "ot cated": "located",
        "st ate": "state",
    }
    
    # Apply corrections to name_of_work
    if labels.get("name_of_work"):
        work = labels["name_of_work"]
        for wrong, correct in ocr_corrections.items():
            work = re.sub(re.escape(wrong), correct, work, flags=re.I)
        labels["name_of_work"] = work
    
    # SECOND PASS: Handle table-format where labels and values are on separate lines
    # This handles OCR from table documents where "Name of Work" is one row and value is next row
    label_to_key = {
        "name of work": "name_of_work",
        "name of the work": "name_of_work",
        "administrative sanction": "admin_sanction",
        "admin sanction": "admin_sanction",
        "administrative sanction amount": "admin_sanction_amount",
        "admin sanction amount": "admin_sanction_amount",
        "admin sanction ref": "admin_sanction",
        "technical sanction": "tech_sanction",
        "tech sanction": "tech_sanction",
        "technical sanction amount": "tech_sanction_amount",
        "tech sanction amount": "tech_sanction_amount",
        "nit no": "nit_no",
        "nit no. & date": "nit_no",
        "nit no & date": "nit_no",
        "contractor name": "agency",
        "name of contractor": "agency",
        "name of the contractor": "agency",
        "contractor address": "contractor_address",
        "estimate amount": "estimate_amount",
        "amount of estimate": "estimate_amount",
        "period of completion": "period_of_completion",
        "period of compl": "period_of_completion",
        "tender premium": "tender_premium",
    }
    
    for idx, raw in enumerate(lines):
        s = str(raw or "").strip()
        if not s:
            continue
        low = s.lower().rstrip(':').strip()
        
        # Check if this line is just a label (no value after it)
        if low in label_to_key:
            key = label_to_key[low]
            # Only fill if not already set
            if not labels.get(key) and idx + 1 < len(lines):
                next_val = str(lines[idx + 1]).strip()
                # Make sure next line isn't another label
                if next_val and next_val.lower().rstrip(':') not in label_to_key:
                    # Clean up Rs. format for amount fields
                    if key in ("admin_sanction_amount", "tech_sanction_amount", "estimate_amount"):
                        rs_match = re.search(r'Rs\.?\s*([\d,\.]+)', next_val, re.I)
                        if rs_match:
                            labels[key] = "Rs." + rs_match.group(1) + "/-"
                        else:
                            labels[key] = next_val
                    # Handle contractor/agency with M/s. pattern
                    elif key == "agency":
                        ms_match = re.search(r'M/s\.?\s*(.+)', next_val, re.I)
                        if ms_match:
                            labels[key] = "M/s." + ms_match.group(1).strip()
                        else:
                            labels[key] = next_val
                    # For multiline values like name_of_work, collect more lines
                    elif key == "name_of_work" and len(next_val) < 100:
                        full_val = _collect_multiline_value(lines, idx + 1, max_lines=5)
                        labels[key] = full_val if full_val else next_val
                    else:
                        labels[key] = next_val

    return labels


def _preprocess_image_for_ocr(img):
    """
    Preprocess image for better OCR accuracy, especially for blurred/low-quality images.
    Uses various image enhancement techniques.
    """
    try:
        from PIL import Image, ImageEnhance, ImageFilter
        import numpy as np
    except ImportError:
        return img  # Return original if dependencies not available
    
    try:
        # Convert to RGB if necessary
        if img.mode != 'RGB':
            img = img.convert('RGB')
        
        # 1. Resize if image is too small (upscale for better OCR)
        width, height = img.size
        if width < 1000 or height < 1000:
            scale_factor = max(1500 / width, 1500 / height)
            new_width = int(width * scale_factor)
            new_height = int(height * scale_factor)
            img = img.resize((new_width, new_height), Image.LANCZOS)
        
        # 2. Convert to grayscale
        img = img.convert('L')
        
        # 3. Increase contrast
        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(2.0)
        
        # 4. Increase sharpness (helps with blurred images)
        img = img.convert('RGB')  # Convert back for sharpness
        enhancer = ImageEnhance.Sharpness(img)
        img = enhancer.enhance(2.5)
        
        # 5. Apply unsharp mask for deblurring effect
        img = img.filter(ImageFilter.UnsharpMask(radius=2, percent=150, threshold=3))
        
        # 6. Convert back to grayscale for OCR
        img = img.convert('L')
        
        # 7. Apply adaptive thresholding using numpy
        try:
            img_array = np.array(img)
            
            # Apply Gaussian blur to reduce noise
            from PIL import ImageFilter
            img_blurred = Image.fromarray(img_array).filter(ImageFilter.GaussianBlur(radius=1))
            img_array = np.array(img_blurred)
            
            # Simple adaptive thresholding
            mean_val = np.mean(img_array)
            threshold = mean_val * 0.85
            img_array = np.where(img_array > threshold, 255, 0).astype(np.uint8)
            img = Image.fromarray(img_array)
        except Exception:
            pass  # Continue with enhanced image if thresholding fails
        
        # 8. Apply slight median filter to remove noise
        img = img.filter(ImageFilter.MedianFilter(size=3))
        
        return img
    except Exception as e:
        logger.warning(f"Image preprocessing failed: {e}")
        return img  # Return original on error


def _ocr_with_multiple_configs(img, lang='eng'):
    """
    Fast OCR extraction - uses single optimal config for speed.
    """
    import logging
    logger = logging.getLogger(__name__)
    
    try:
        import pytesseract
    except ImportError:
        return ""
    
    # Use single best config for speed (PSM 6 is optimal for documents)
    try:
        result = pytesseract.image_to_string(img, lang=lang, config='--oem 3 --psm 6')
        logger.debug(f"OCR extracted {len(result)} chars")
        return result
    except Exception as e:
        logger.debug(f"OCR failed: {e}")
        return ""


def _extract_labels_from_source_file(uploaded_file):
    """
    Read uploaded file (Excel / Word / PDF / Image / text)
    and return:
      - labels: dict from _extract_labels_from_lines
      - lines:  flattened list of text lines (for custom placeholders)
    
    Supports: .xlsx, .xlsm, .docx, .pdf, .jpg, .jpeg, .png, .bmp, .tiff, .gif, .txt, .csv
    Uses advanced image preprocessing for blurred/low-quality scanned documents.
    """
    filename = uploaded_file.name or ""
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    lines = []

    if ext in ("xlsx", "xlsm"):
        wb = load_workbook(uploaded_file, data_only=True)
        for ws in wb.worksheets:
            max_r = min(ws.max_row or 0, 150)
            max_c = min(ws.max_column or 0, 20)
            for r in range(1, max_r + 1):
                vals = []
                for c in range(1, max_c + 1):
                    v = ws.cell(row=r, column=c).value
                    if v is not None:
                        vals.append(str(v).strip())
                if vals:
                    # For 2-column table format (Label | Value), create "Label: Value" format
                    if len(vals) == 2:
                        label = vals[0].lower().strip()
                        # Check if first column looks like a label
                        if any(x in label for x in ['name', 'work', 'sanction', 'amount', 'contractor', 
                                                      'nit', 'estimate', 'period', 'address', 'premium']):
                            lines.append(f"{vals[0]}: {vals[1]}")
                        else:
                            lines.append(" ".join(vals))
                    # For single cell or multi-column, just join
                    else:
                        lines.append(" ".join(vals))

    elif ext == "docx":
        doc = Document(uploaded_file)
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                lines.append(t)
        for table in doc.tables:
            for row in table.rows:
                vals = [(cell.text or "").strip() for cell in row.cells]
                vals = [v for v in vals if v]
                if vals:
                    # For 2-column table format (Label | Value), create "Label: Value" format
                    if len(vals) == 2:
                        label = vals[0].lower().strip()
                        # Check if first column looks like a label
                        if any(x in label for x in ['name', 'work', 'sanction', 'amount', 'contractor', 
                                                      'nit', 'estimate', 'period', 'address', 'premium']):
                            lines.append(f"{vals[0]}: {vals[1]}")
                        else:
                            lines.append(" ".join(vals))
                    else:
                        lines.append(" ".join(vals))

    elif ext == "pdf":
        import logging
        logger = logging.getLogger(__name__)
        
        # First try PyPDF2 for text-based PDFs
        try:
            import PyPDF2
            uploaded_file.seek(0)  # Reset file pointer
            reader = PyPDF2.PdfReader(uploaded_file)
            for page in reader.pages:
                txt = page.extract_text() or ""
                # PDF text often has inconsistent spacing/newlines
                for ln in txt.splitlines():
                    ln = ln.strip()
                    if ln:
                        # Handle cases where multiple fields are on one line
                        if ':' in ln and len(ln) > 50:
                            parts = re.split(r'(?<=[a-zA-Z])\s*:\s*(?=[A-Z])', ln)
                            if len(parts) > 1:
                                for p in parts:
                                    if p.strip():
                                        lines.append(p.strip())
                                continue
                        lines.append(ln)
        except Exception as e:
            logger.warning(f"PyPDF2 extraction failed: {e}")
        
        # If no text extracted or very little text, try OCR for scanned/blurred PDFs
        if len(lines) < 3:
            try:
                from pdf2image import convert_from_bytes
                import pytesseract
                from PIL import Image
                
                logger.info("Attempting OCR for scanned PDF...")
                uploaded_file.seek(0)
                pdf_bytes = uploaded_file.read()
                
                # Convert PDF pages to images at lower DPI for speed
                images = convert_from_bytes(pdf_bytes, dpi=200)
                
                ocr_lines = []
                
                for idx, img in enumerate(images):
                    logger.info(f"Processing page {idx + 1} with OCR...")
                    
                    # Single OCR pass for speed
                    txt = _ocr_with_multiple_configs(img, lang='eng')
                    
                    for ln in txt.splitlines():
                        ln = ln.strip()
                        if ln and len(ln) > 1:
                            ocr_lines.append(ln)
                
                # Use OCR results if they have more content
                if len(ocr_lines) > len(lines):
                    lines = ocr_lines
                    logger.info(f"OCR extracted {len(lines)} lines from scanned PDF")
                    
            except ImportError as e:
                logger.warning(f"OCR libraries not available: {e}. Install pdf2image, pytesseract, and Pillow.")
            except Exception as e:
                logger.warning(f"OCR extraction failed: {e}")

    elif ext in ("jpg", "jpeg", "png", "bmp", "tiff", "tif", "gif", "webp"):
        # Direct image file support
        import logging
        logger = logging.getLogger(__name__)
        
        try:
            from PIL import Image
            import pytesseract
            
            logger.info(f"Processing image file: {filename}")
            uploaded_file.seek(0)
            img = Image.open(uploaded_file)
            
            # Single OCR pass for speed
            txt = _ocr_with_multiple_configs(img, lang='eng')
            
            for ln in txt.splitlines():
                ln = ln.strip()
                if ln and len(ln) > 1:  # Skip single characters
                    lines.append(ln)
            
            logger.info(f"OCR extracted {len(lines)} lines from image")
            
        except ImportError as e:
            logger.warning(f"Image OCR libraries not available: {e}. Install Pillow and pytesseract.")
        except Exception as e:
            logger.warning(f"Image OCR extraction failed: {e}")

    else:
        # plain text / csv / unknown  -  treat as text
        try:
            content = uploaded_file.read()
        except Exception:
            content = b""
        try:
            text = content.decode("utf-8", errors="ignore")
        except Exception:
            text = str(content)
        for ln in text.splitlines():
            ln = ln.strip()
            if ln:
                lines.append(ln)

    # Debug logging - show what was extracted from source file
    import logging
    logger = logging.getLogger(__name__)
    logger.info(f"=== SOURCE FILE LINES ({len(lines)} total) ===")
    for i, line in enumerate(lines[:30]):  # First 30 lines
        logger.info(f"  [{i}] {line[:100]}...")  # First 100 chars
    logger.info("=== END SOURCE FILE LINES ===")

    labels = _extract_labels_from_lines(lines)
    
    # Debug logging - show extracted labels
    logger.info("=== EXTRACTED LABELS ===")
    for key, val in labels.items():
        if val:
            logger.info(f"  {key}: {val[:80]}..." if len(str(val)) > 80 else f"  {key}: {val}")
    logger.info("=== END EXTRACTED LABELS ===")
    
    return labels, lines


# -------------------------------------------
#  SELF FORMATTED MODULE  -  HELPERS
# -------------------------------------------

# imports consolidated at top of file
# from .utils import _number_to_words_rupees, _extract_labels_from_source_file
# or similar  -  just make sure they are imported somewhere above.


def _fuzzy_find_from_lines(lines, label_hint: str, threshold: float = 0.55) -> str:
    """
    Given a list of text lines (strings) and a label hint (e.g. "Name of the work"),
    find the line that best matches the label text using fuzzy similarity.

    Used by _build_placeholder_map() to map a custom KEY to the nearest
    line in the source document.
    """
    label = (label_hint or "").strip()
    if not label:
        return ""

    if not lines:
        return ""

    candidates = []

    # Normalize lines into plain strings
    for ln in lines:
        # If it's a list/tuple from some extractor, join elements
        if isinstance(ln, (list, tuple)):
            text = " ".join(str(x) for x in ln if x is not None)
        else:
            text = str(ln)
        text = text.strip()
        if not text:
            continue
        candidates.append(text)

    if not candidates:
        return ""

    best_line = ""
    best_score = 0.0
    label_low = label.lower()

    for text in candidates:
        score = SequenceMatcher(None, label_low, text.lower()).ratio()
        if score > best_score:
            best_score = score
            best_line = text

    # If nothing is similar enough, return empty string
    if best_score < threshold:
        return ""

    return best_line


def _extract_last_number(text: str) -> str:
    """
    From a string like 'Grand Total          105000.00',
    return '105000.00'. Returns "" if no number.
    """
    if not text:
        return ""
    cleaned = text.replace(",", "")
    nums = re.findall(r"[-+]?\d+(?:\.\d+)?", cleaned)
    return nums[-1] if nums else ""


def _build_placeholder_map(labels, lines, custom_text: str):
    """
    labels: from _extract_labels_from_source_file()
    lines : flattened text lines
    custom_text: textarea content "KEY = label text"

    Returns: dict { "{{PLACEHOLDER}}": "value" }
    """

    # -------- AMOUNT + INDIAN WORDS HANDLING --------
    # 'amount' is expected to be in labels (e.g. "56,000.00")
    raw_amount = (labels.get("amount") or "").strip()
    # remove common currency prefixes like 'Rs.', 'INR', or the rupee symbol
    raw_amount = re.sub(r'^(?:rs\.?|inr|â‚¹)\s*[:\- - â€”]?\s*', '', raw_amount, flags=re.I)
    raw_amount = raw_amount.replace(",", "").strip()
    total_amount = 0.0
    amount_in_words = ""

    if raw_amount:
        try:
            total_amount = float(raw_amount)
        except ValueError:
            total_amount = 0.0

    if total_amount:
        try:
            # Indian-style: Crores, Lakhs, Thousands, Rupees, Paise
            amount_in_words = _number_to_words_rupees(total_amount)
        except Exception:
            amount_in_words = ""

    # -------- ESTIMATE GRAND TOTAL (special case) --------
    est_grand_total = labels.get("est_grand_total", "")

    # If extractor did not supply it, try to guess from lines
    if not est_grand_total:
        grand_line = _fuzzy_find_from_lines(lines, "Grand Total")
        if grand_line:
            est_grand_total = _extract_last_number(grand_line)

    # As a fallback, use generic "amount"
    if not est_grand_total:
        est_grand_total = labels.get("amount", "")

    # prepare formatted amount strings
    formatted_amount = ""
    amount_raw = (labels.get("amount") or "").strip()
    if total_amount:
        try:
            if "." in amount_raw:
                formatted_amount = f"{total_amount:,.2f}"
            else:
                formatted_amount = f"{int(total_amount):,}"
        except Exception:
            formatted_amount = str(total_amount)

    # -------- BUILT-IN PLACEHOLDERS --------
    placeholder_map = {
        "{{NAME_OF_WORK}}":     labels.get("name_of_work", ""),
        "{{REF_OF_AGREEMENT}}": labels.get("agreement", ""),
        "{{AGREEMENT_REF}}":    labels.get("agreement", ""),
        "{{BOND_NO}}":          labels.get("agreement", ""),  # Alias for agreement/bond
        "{{ADMIN_SANCTION}}":   labels.get("admin_sanction", ""),
        "{{TECH_SANCTION}}":    labels.get("tech_sanction", ""),
        "{{NAME_OF_AGENCY}}":   labels.get("agency", ""),
        "{{AGENCY_NAME}}":      labels.get("agency", ""),
        "{{CONTRACTOR}}":       labels.get("agency", ""),  # Alias for agency/contractor
        "{{NAME_OF_CONTRACTOR}}": labels.get("agency", ""),

        # MB details content only (label part stripped in _extract_labels_from_source_file)
        "{{MB_DETAILS}}":       labels.get("mb_details", ""),

        "{{TP_DETAILS}}":       labels.get("tp_details", ""),
        "{{TENDER_PREMIUM}}":   labels.get("tender_premium", ""),
        "{{CC_HEADER}}":        labels.get("cc_header", ""),

        # Estimate amount (from "Amount of Estimate" line)
        "{{ESTIMATE_AMOUNT}}":  labels.get("estimate_amount", ""),
        
        # Admin/Tech Sanction amounts
        "{{ADMIN_SANCTION_AMOUNT}}": labels.get("admin_sanction_amount", ""),
        "{{TECH_SANCTION_AMOUNT}}": labels.get("tech_sanction_amount", ""),
        
        # NIT / Tender details
        "{{NIT_NO}}":           labels.get("nit_no", ""),
        "{{TENDER_NO}}":        labels.get("nit_no", ""),
        
        # Contractor address
        "{{CONTRACTOR_ADDRESS}}": labels.get("contractor_address", ""),
        
        # Period and Dates
        "{{PERIOD_OF_COMPLETION}}": labels.get("period_of_completion", ""),
        "{{STIPULATED_PERIOD}}": labels.get("period_of_completion", ""),
        "{{DATE_OF_COMMENCEMENT}}": labels.get("date_of_commencement", ""),
        "{{DATE_OF_COMPLETION}}": labels.get("date_of_completion", ""),
        "{{WORK_ORDER_NO}}":    labels.get("work_order_no", ""),
        "{{WORK_ORDER_DATE}}":  labels.get("work_order_date", ""),
        
        # Money related
        "{{EARNEST_MONEY}}":    labels.get("earnest_money", ""),
        "{{EMD}}":              labels.get("earnest_money", ""),
        "{{SECURITY_DEPOSIT}}": labels.get("security_deposit", ""),
        "{{SD}}":               labels.get("security_deposit", ""),

        # Generic amount from source (Bill / Estimate / etc.)
        # {{AMOUNT}}: formatted with commas (e.g. 1,23,456) when possible
        "{{AMOUNT}}":           formatted_amount or labels.get("amount", ""),
        "{{TOTAL_AMOUNT}}":     formatted_amount or labels.get("amount", ""),
        # raw numeric value without commas/currency
        "{{AMOUNT_RAW}}":       amount_raw.replace(",", ""),
        # convenience: amount with Rs. prefix
        "{{AMOUNT_WITH_RS}}":   ("Rs. " + (formatted_amount or labels.get("amount", ""))).strip(),

        # For your estimate covering letter  -  maps to Grand Total
        "{{EST_GRAND_TOTAL}}":  est_grand_total,

        "{{AMOUNT_IN_WORDS}}":  amount_in_words,
    }

    # -------- CUSTOM PLACEHOLDERS (from textarea) --------
    # Syntax: MY_PLACEHOLDER = label text
    custom_raw = (custom_text or "").strip()
    if custom_raw:
        for line in custom_raw.splitlines():
            line = line.strip()
            if not line or line.startswith("#"):
                continue
            if "=" not in line:
                continue

            key, label_hint = line.split("=", 1)
            key = key.strip()
            label_hint = label_hint.strip()
            if not key or not label_hint:
                continue

            # normalise to {{KEY}} format
            if not key.startswith("{{"):
                key = "{{" + key
            if not key.endswith("}}"):
                key = key + "}}"

            match_line = _fuzzy_find_from_lines(lines, label_hint)
            if not match_line:
                continue

            # Try to be smart: if there is a number on that line, use the number;
            # otherwise use the full line.
            value_num = _extract_last_number(match_line)
            value = value_num or match_line

            placeholder_map[key] = value

    return placeholder_map


def _fill_template_file(template_file, placeholder_map):
    """
    Apply placeholders to template_file and return a HttpResponse with the
    filled file.
    """
    template_name = template_file.name or "template"
    ext = template_name.lower().rsplit(".", 1)[-1] if "." in template_name else ""

    # -------- DOCX --------
    if ext == "docx":
        doc = Document(template_file)

        def replace_in_paragraphs(paragraphs):
            for p in paragraphs:
                # join all runs to a single string to handle placeholders split across runs
                full = ''.join([r.text or '' for r in p.runs])
                new_full = full
                for ph, val in placeholder_map.items():
                    new_full = new_full.replace(ph, str(val) if val is not None else "")
                if new_full != full:
                    # write back into runs: put entire text in first run, clear others
                    if p.runs:
                        p.runs[0].text = new_full
                        for rr in p.runs[1:]:
                            rr.text = ''
                    else:
                        p.add_run(new_full)

        replace_in_paragraphs(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_in_paragraphs(cell.paragraphs)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        resp = HttpResponse(
            buf.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        resp["Content-Disposition"] = f'attachment; filename="Filled_{os.path.basename(template_name)}"'
        return resp

    # -------- Excel (XLSX / XLSM) --------
    if ext in ("xlsx", "xlsm"):
        wb = load_workbook(template_file)
        for ws in wb.worksheets:
            max_r = ws.max_row or 0
            max_c = ws.max_column or 0
            for r in range(1, max_r + 1):
                for c in range(1, max_c + 1):
                    cell = ws.cell(row=r, column=c)
                    if isinstance(cell.value, str):
                        txt = cell.value
                        changed = False
                        for ph, val in placeholder_map.items():
                            if ph in txt:
                                txt = txt.replace(ph, str(val) if val is not None else "")
                                changed = True
                        if changed:
                            cell.value = txt

        resp = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        resp["Content-Disposition"] = f'attachment; filename="Filled_{os.path.basename(template_name)}"'
        wb.save(resp)
        return resp

    # -------- Text / CSV --------
    if ext in ("txt", "csv"):
        try:
            content = template_file.read()
        except Exception:
            content = b""
        try:
            text = content.decode("utf-8", errors="ignore")
        except Exception:
            text = str(content)

        for ph, val in placeholder_map.items():
            text = text.replace(ph, str(val) if val is not None else "")

        resp = HttpResponse(text, content_type="text/plain; charset=utf-8")
        resp["Content-Disposition"] = f'attachment; filename="Filled_{os.path.basename(template_name)}"'
        return resp

    return HttpResponse(
        f"Unsupported template type .{ext}. Use DOCX / XLSX / XLSM / TXT / CSV.",
        status=400,
    )


# ============================================
#  SELF-FORMATTED FORMS  -  VIEWS
# ============================================

@login_required(login_url='login')
def self_formatted_form_page(request):
    """
    Shows:
      - Quick one-time generation form
      - Create reusable format form
      - List of saved formats
    Optimized: Limited query with only necessary fields for faster load.
    """
    # Only fetch the fields needed for display, limit to recent 20 formats
    saved_formats = SelfFormattedTemplate.objects.only(
        'id', 'name', 'created_at'
    ).order_by("-created_at")[:20]
    error_message = request.GET.get("error")  # optional error via redirect

    return render(request, "core/self_formatted.html", {
        "saved_formats": saved_formats,
        "error_message": error_message,
    })


@login_required(login_url='login')
def self_formatted_generate(request):
    """
    Quick one-time generation: user uploads source + template, optional
    custom placeholders text. Does not save anything in DB.
    Optimized: No database queries during file processing for speed.
    """
    if request.method != "POST":
        return HttpResponseNotAllowed(["POST"])

    source_file = request.FILES.get("source_file")
    template_file = request.FILES.get("template_file")
    custom_text = request.POST.get("custom_placeholders", "")

    # On error, redirect with error message instead of querying DB
    if not source_file or not template_file:
        return redirect(f"{reverse('self_formatted_form_page')}?error=Please+upload+both+source+file+and+template+file")

    labels, lines = _extract_labels_from_source_file(source_file)
    
    # Check if no text was extracted (likely scanned PDF)
    if not lines:
        filename = source_file.name or ""
        if filename.lower().endswith('.pdf'):
            error_msg = "PDF+appears+to+be+scanned.+Use+Excel+or+Word+file+instead."
        else:
            error_msg = "No+text+could+be+extracted+from+the+source+file."
        return redirect(f"{reverse('self_formatted_form_page')}?error={error_msg}")
    
    placeholder_map = _build_placeholder_map(labels, lines, custom_text)

    return _fill_template_file(template_file, placeholder_map)


@login_required(login_url='login')
def self_formatted_preview(request):
    """AJAX endpoint: compute placeholder_map for a given source file + custom text and return JSON.
    Used by the UI to preview mappings before generation.
    Uses faster/lighter OCR settings for quicker preview response.
    """
    from django.http import JsonResponse
    import logging
    logger = logging.getLogger(__name__)

    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)

    source_file = request.FILES.get("source_file")
    custom_text = request.POST.get("custom_placeholders", "")

    if not source_file:
        return JsonResponse({"error": "source_file required"}, status=400)

    try:
        logger.info(f"Preview: Processing {source_file.name}")
        labels, lines = _extract_labels_from_source_file(source_file)
        
        # Check if no text was extracted
        if not lines:
            filename = source_file.name or ""
            if filename.lower().endswith('.pdf'):
                return JsonResponse({
                    "error": "No text extracted. This appears to be a scanned PDF. Use Excel/Word file or install Tesseract OCR."
                }, status=400)
            return JsonResponse({"error": "No text could be extracted from the file."}, status=400)
        
        logger.info(f"Preview: Extracted {len(lines)} lines, building placeholders...")
        placeholder_map = _build_placeholder_map(labels, lines, custom_text)
        logger.info(f"Preview: Found {len(placeholder_map)} placeholders")
    except Exception as e:
        logger.error(f"Preview error: {e}")
        return JsonResponse({"error": str(e)}, status=500)

    return JsonResponse({"placeholders": placeholder_map})


@org_required
def self_formatted_save_format(request):
    """
    Save a reusable format (name + description + template file + custom placeholders).
    Optimized: Redirect with error instead of querying DB.
    """
    if request.method != "POST":
        return redirect("self_formatted_form_page")

    format_name = request.POST.get("format_name", "").strip()
    format_description = request.POST.get("format_description", "").strip()
    template_file = request.FILES.get("format_template_file")
    raw_custom = request.POST.get("format_custom_placeholders", "").strip()

    if not format_name or not template_file:
        return redirect(f"{reverse('self_formatted_form_page')}?error=Format+name+and+template+file+are+required")

    fmt = SelfFormattedTemplate(
        name=format_name,
        description=format_description,
        template_file=template_file,
        custom_placeholders=raw_custom,
    )
    fmt.save()

    return redirect("self_formatted_form_page")


@org_required
def self_formatted_use_format(request, pk):
    """
    Use a saved format:
      GET  -> show page asking only for source_file upload.
      POST -> generate document using saved template + placeholders.
    """
    fmt = get_object_or_404(SelfFormattedTemplate, pk=pk)

    if request.method == "GET":
        return render(request, "core/self_formatted_use.html", {
            "format": fmt,
        })

    if request.method == "POST":
        source_file = request.FILES.get("source_file")
        if not source_file:
            return HttpResponse("Please upload a source file.", status=400)

        labels, lines = _extract_labels_from_source_file(source_file)
        placeholder_source = fmt.custom_placeholders or ""
        placeholder_map = _build_placeholder_map(labels, lines, placeholder_source)

        # Reopen template file from disk
        try:
            with fmt.template_file.open("rb") as f:
                data = f.read()
        except FileNotFoundError:
            # Template file was moved/deleted from media folder
            return redirect(
                f"{reverse('self_formatted_form_page')}?error="
                "Template file not found on server. "
                "Delete this saved format and create it again."
            )

        mem = io.BytesIO(data)
        uploaded = InMemoryUploadedFile(
            mem,
            field_name="template_file",
            name=os.path.basename(fmt.template_file.name),
            content_type="application/octet-stream",
            size=len(data),
            charset=None,
        )
        return _fill_template_file(uploaded, placeholder_map)

    return HttpResponseNotAllowed(["GET", "POST"])


@org_required
def self_formatted_delete_format(request, pk):
    """
    Delete a saved format (and its underlying template file).
    """
    if request.method != "POST":
        return HttpResponseNotAllowed(["POST"])

    fmt = get_object_or_404(SelfFormattedTemplate, pk=pk)

    template = fmt.template_file
    storage = template.storage if template else None
    name = template.name if template else None

    fmt.delete()

    if storage and name and storage.exists(name):
        storage.delete(name)

    return redirect("self_formatted_form_page")


@org_required
def self_formatted_edit_format(request, pk):
    """Edit an existing SelfFormattedTemplate.
    GET: show edit form
    POST: apply updates (name, description, optional new template file, custom placeholders)
    """
    fmt = get_object_or_404(SelfFormattedTemplate, pk=pk)

    if request.method == "GET":
        preview_text = None
        template_url = None
        try:
            if fmt.template_file and fmt.template_file.name:
                template_url = fmt.template_file.url
                # attempt lightweight preview depending on extension
                name = fmt.template_file.name.lower()
                with fmt.template_file.open('rb') as f:
                    data = f.read()

                if name.endswith('.txt') or name.endswith('.csv'):
                    preview_text = data.decode('utf-8', errors='replace')[:4000]
                elif name.endswith('.docx'):
                    try:
                        from docx import Document
                        mem = io.BytesIO(data)
                        doc = Document(mem)
                        paras = [p.text for p in doc.paragraphs if p.text]
                        preview_text = '\n'.join(paras)[:4000]
                    except Exception:
                        preview_text = None
                elif name.endswith('.pdf'):
                    try:
                        from PyPDF2 import PdfReader
                        mem = io.BytesIO(data)
                        reader = PdfReader(mem)
                        text_parts = []
                        if reader.pages:
                            text_parts.append(reader.pages[0].extract_text() or '')
                        preview_text = '\n'.join(text_parts)[:4000]
                    except Exception:
                        preview_text = None
                elif name.endswith('.xlsx') or name.endswith('.xlsm'):
                    try:
                        from openpyxl import load_workbook
                        mem = io.BytesIO(data)
                        wb = load_workbook(mem, read_only=True, data_only=True)
                        sheet = wb.active
                        rows = []
                        for r in sheet.iter_rows(min_row=1, max_row=8, max_col=6, values_only=True):
                            rows.append('\t'.join([str(c) if c is not None else '' for c in r]))
                        preview_text = '\n'.join(rows)
                    except Exception:
                        preview_text = None
        except Exception:
            preview_text = None

        return render(request, "core/self_formatted_edit.html", {
            "format": fmt,
            "preview_text": preview_text,
            "template_url": template_url,
        })

    # POST
    name = request.POST.get("format_name", "").strip()
    description = request.POST.get("format_description", "").strip()
    raw_custom = request.POST.get("format_custom_placeholders", "").strip()
    new_template = request.FILES.get("format_template_file")

    if not name:
        return render(request, "core/self_formatted_edit.html", {
            "format": fmt,
            "error_message": "Name is required.",
        }, status=400)

    # Replace template file if new file provided
    old_name = None
    storage = None
    if new_template:
        if fmt.template_file and fmt.template_file.name:
            old_name = fmt.template_file.name
            storage = fmt.template_file.storage
        fmt.template_file = new_template

    fmt.name = name
    fmt.description = description
    fmt.custom_placeholders = raw_custom
    fmt.save()

    # delete old template file from storage after saving new one
    try:
        if old_name and storage and storage.exists(old_name):
            storage.delete(old_name)
    except Exception:
        pass

    return redirect("self_formatted_form_page")

# ==========================
#  TEMPORARY WORKS MODULE
#  (completely separate from New Estimate)
# ==========================


@login_required(login_url='login')
def tempworks_home(request):
    """
    Landing page for Temporary Works:
    - Just shows buttons to choose category (temp_electrical, temp_civil, etc.)
    - Clears only the temp session keys.
    """
    request.session["temp_entries"] = []  # list of {"id":..., "name":...}
    request.session["temp_work_name"] = ""
    request.session["temp_selected_backend_id"] = None  # Clear backend selection
    return render(request, "core/tempworks_home.html")


@login_required(login_url='login')
def temp_groups(request, category):
    """
    Step 1 in Temporary Works: show groups for given temp category.
    Example categories: 'temp_electrical', 'temp_civil'
    """
    from core.utils_excel import get_available_backends_for_module
    
    # Handle backend switching via URL parameter
    url_backend_id = request.GET.get("backend_id")
    if url_backend_id:
        try:
            request.session["temp_selected_backend_id"] = int(url_backend_id)
        except (ValueError, TypeError):
            pass
    
    temp_selected_backend_id = request.session.get("temp_selected_backend_id")
    
    # Map temp category to base category for backend lookup
    base_category = category.replace('temp_', '')  # temp_electrical -> electrical
    
    try:
        items_list, groups_map, _, ws_data, filepath = load_backend(
            category, settings.BASE_DIR,
            backend_id=temp_selected_backend_id,
            module_code='temp_works',  # Use temp_works module's own backends
            user=request.user
        )
    except FileNotFoundError as e:
        logger.info(f"No backend data available for temp category {category} - showing Coming Soon")
        # Check if the other category has backends available
        other_base = 'electrical' if base_category == 'civil' else 'civil'
        other_category = f'temp_{other_base}'
        other_available = False
        try:
            other_backends = get_available_backends_for_module('temp_works', other_base)
            other_available = len(other_backends) > 0
        except:
            pass
        return render(request, "core/coming_soon.html", {
            "category": category,
            "module_name": "Temporary Works",
            "other_category": other_category,
            "other_category_available": other_available,
        })
    except Exception as e:
        logger.error(f"Error loading temp backend for {category}: {e}")
        return render(request, "core/error.html", {
            "error_title": "Loading Error",
            "error_message": f"Could not load temporary works data: {str(e)}",
        })
        
    groups = sorted(groups_map.keys(), key=lambda s: s.lower())

    if not groups:
        return render(
            request,
            "core/temp_items.html",
            {
                "category": category,
                "group": "",
                "groups": [],
                "items_info": [],
                "entries": [],
                "day_rates_json": "{}",
                "is_temporary": True,
                "error": "No groups found in Temporary Works backend.",
            },
        )

    default_group = request.GET.get("group") or groups[0]
    return redirect("temp_items", category=category, group=default_group)


@login_required(login_url='login')
def temp_items(request, category, group):
    """
    Temporary Works UI:
    - rates must be computed from temp workbook (Column C day, Column J rate)
    - Column J has formulas (ROUND etc), so we evaluate if cached value is missing
    """
    from core.utils_excel import get_available_backends_for_module
    
    # Handle backend switching via URL parameter
    url_backend_id = request.GET.get("backend_id")
    if url_backend_id:
        try:
            request.session["temp_selected_backend_id"] = int(url_backend_id)
        except (ValueError, TypeError):
            pass
    
    temp_selected_backend_id = request.session.get("temp_selected_backend_id")
    
    # Map temp category to base category for backend lookup
    base_category = category.replace('temp_', '')  # temp_electrical -> electrical
    
    # Get available backends for dropdown (temp_works has its own backends)
    available_backends = get_available_backends_for_module('temp_works', base_category)
    
    try:
        items_list, groups_map, _, ws_src, filepath = load_backend(
            category, settings.BASE_DIR,
            backend_id=temp_selected_backend_id,
            module_code='temp_works',  # Use temp_works module's own backends
            user=request.user
        )
    except FileNotFoundError as e:
        logger.info(f"No backend data available for temp category {category} - showing Coming Soon")
        other_base = 'electrical' if base_category == 'civil' else 'civil'
        other_category = f'temp_{other_base}'
        other_available = False
        try:
            other_backends = get_available_backends_for_module('temp_works', other_base)
            other_available = len(other_backends) > 0
        except:
            pass
        return render(request, "core/coming_soon.html", {
            "category": category,
            "module_name": "Temporary Works",
            "other_category": other_category,
            "other_category_available": other_available,
        })
    except Exception as e:
        logger.error(f"Error loading temp backend for {category}: {e}")
        return render(request, "core/error.html", {
            "error_title": "Loading Error",
            "error_message": f"Could not load temporary works data: {str(e)}",
        })

    groups = sorted(groups_map.keys(), key=lambda s: s.lower())
    group_items = groups_map.get(group, [])

    detected_names = {i["name"] for i in items_list}
    display_items = [name for name in group_items if name in detected_names]
    items_info = [{"name": name} for name in display_items]

    # units mapping (same as your code)
    item_to_group = {}
    for grp_name, item_list_in_grp in groups_map.items():
        for nm in item_list_in_grp:
            item_to_group.setdefault(nm, grp_name)

    def units_for(name):
        grp_name = (item_to_group.get(name, "") or "").lower()
        if grp_name in ("piping", "wiring & cables", "wiring and cables"):
            return ("Mtrs", "Mtr")
        elif grp_name == "points":
            return ("Pts", "Pt")
        else:
            return ("Nos", "No")

    temp_entries = request.session.get("temp_entries", []) or []
    display_entries = []
    for idx, ent in enumerate(temp_entries, start=1):
        plural, _singular = units_for(ent["name"])
        display_entries.append(
            {
                "id": ent["id"],
                "sl": idx,
                "name": ent["name"],
                "unit": plural,
                "qty": ent.get("qty", ""),
                "days": ent.get("days", 1),
            }
        )

    # âœ… IMPORTANT: build day rates using filepath (so we can load wb twice)
    day_rates = build_temp_day_rates(filepath, items_list)
    day_rates_json = json.dumps(day_rates)

    work_name = request.session.get("temp_work_name", "") or ""
    grand_total = request.session.get("temp_grand_total", "") or ""

    context = {
        "category": category,
        "group": group,
        "groups": groups,
        "items_info": items_info,
        "entries": display_entries,
        "day_rates_json": day_rates_json,
        "is_temporary": True,
        "work_name": work_name,
        "grand_total": grand_total,
        "available_backends": available_backends,
        "selected_backend_id": temp_selected_backend_id,
    }
    return render(request, "core/temp_items.html", context)


@login_required(login_url='login')
def temp_day_rates_debug(request, category):
    """Debug endpoint: return JSON of computed day rates for a category.
    Useful to inspect what the view passes to the template.
    """
    from django.http import JsonResponse

    try:
        items_list, groups_map, _, ws_src, filepath = load_backend(category, settings.BASE_DIR)
        day_rates = build_temp_day_rates(filepath, items_list)
    except Exception as e:
        return JsonResponse({"error": str(e)}, status=500)

    return JsonResponse(day_rates, safe=False)

@login_required(login_url='login')
def temp_add_item(request, category, group, item):
    """
    Add one more row for given item (duplicates allowed).
    """
    temp_entries = request.session.get("temp_entries", []) or []
    entry = {
        "id": get_random_string(8),
        "name": item,
        # we don't store qty/days here; they live in the form & are sent on POST
    }
    temp_entries.append(entry)
    request.session["temp_entries"] = temp_entries

    # Preserve 'work_name' typed in the box
    work_name = request.GET.get("work_name")
    if work_name is not None:
        request.session["temp_work_name"] = work_name

    return redirect("temp_items", category=category, group=group)


# -----------------------
# TEMP AJAX ADD ITEM (no page reload)
# -----------------------
@login_required(login_url='login')
def temp_ajax_add_item(request, category):
    """
    AJAX endpoint to add an item to temp entries without page reload.
    POST with JSON: { "item": "item_name", "work_name": "..." }
    Returns JSON: { "status": "ok", "entries": [...], "entry_id": "...", "item_info": {...} }
    """
    if request.method != "POST":
        return JsonResponse({"status": "error", "message": "POST required"}, status=405)
    
    try:
        if request.content_type and 'application/json' in request.content_type:
            data = json.loads(request.body.decode("utf-8") or "{}")
        else:
            data = request.POST
        
        item = data.get("item", "").strip()
        work_name = data.get("work_name")
        
        if not item:
            return JsonResponse({"status": "error", "message": "No item specified"}, status=400)
        
        temp_entries = request.session.get("temp_entries", []) or []
        entry = {
            "id": get_random_string(8),
            "name": item,
        }
        temp_entries.append(entry)
        request.session["temp_entries"] = temp_entries
        
        if work_name is not None:
            request.session["temp_work_name"] = work_name
        
        # Get unit from backend units_map (Column D of Groups sheet)
        item_info = {"name": item, "unit": "Nos"}
        try:
            items_list, groups_map, units_map, ws_data, filepath = load_backend(category, settings.BASE_DIR)
            
            # Get unit from backend units_map
            unit = units_map.get(item, "Nos")
            
            item_info["unit"] = unit
        except Exception as e:
            # If we can't determine unit, default to Nos
            pass
        
        return JsonResponse({
            "status": "ok",
            "entries": temp_entries,
            "entry_id": entry["id"],
            "item": item,
            "item_info": item_info
        })
        
    except Exception as e:
        return JsonResponse({"status": "error", "message": str(e)}, status=400)


# -----------------------
# TEMP AJAX REORDER ITEMS
# -----------------------
@login_required(login_url='login')
def temp_ajax_reorder_items(request, category):
    """
    AJAX endpoint to reorder temp entries list.
    POST with JSON: { "entries": [{"id": "...", "name": "..."}, ...] }
    Returns JSON: { "status": "ok", "entries": [...] }
    """
    if request.method != "POST":
        return JsonResponse({"status": "error", "message": "POST required"}, status=405)
    
    try:
        if request.content_type and 'application/json' in request.content_type:
            data = json.loads(request.body.decode("utf-8") or "{}")
        else:
            return JsonResponse({"status": "error", "message": "JSON required"}, status=400)
        
        new_order = data.get("entries", [])
        
        if not isinstance(new_order, list):
            return JsonResponse({"status": "error", "message": "entries must be a list"}, status=400)
        
        # Validate entries have required fields
        current_entries = request.session.get("temp_entries", []) or []
        current_ids = {e["id"] for e in current_entries}
        
        # Reorder based on provided IDs
        valid_entries = []
        for entry in new_order:
            if isinstance(entry, dict) and entry.get("id") in current_ids:
                # Find the original entry to preserve all data
                original = next((e for e in current_entries if e["id"] == entry["id"]), None)
                if original:
                    valid_entries.append(original)
        
        request.session["temp_entries"] = valid_entries
        
        return JsonResponse({
            "status": "ok",
            "entries": valid_entries
        })
        
    except Exception as e:
        return JsonResponse({"status": "error", "message": str(e)}, status=400)


# -----------------------
# TEMP CLEAR ITEMS (stay on same page)
# -----------------------
@login_required(login_url='login')
def temp_clear_items(request, category, group):
    """
    Clear all temp entries but stay on the same page.
    """
    request.session["temp_entries"] = []
    request.session["temp_work_name"] = ""
    request.session["temp_grand_total"] = ""
    return redirect("temp_items", category=category, group=group)


@login_required(login_url='login')
def temp_save_state(request, category):
    """Save current temporary entries (from client) into session.
    Expects JSON body: { entries: [{id, name, qty, days}, ...], work_name: "...", grand_total: "..." }
    """
    from django.http import JsonResponse

    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)

    try:
        payload = json.loads(request.body.decode("utf-8") or "{}")
    except Exception:
        return JsonResponse({"error": "invalid json"}, status=400)

    entries = payload.get("entries", [])
    work_name = payload.get("work_name", "")
    grand_total = payload.get("grand_total", "")

    # Basic validation: ensure list of dicts
    if not isinstance(entries, list):
        entries = []

    # Save into session
    request.session["temp_entries"] = entries
    request.session["temp_work_name"] = work_name or ""
    request.session["temp_grand_total"] = grand_total or ""

    return JsonResponse({"ok": True})


@login_required(login_url='login')
def temp_remove_item(request, category, group, entry_id):
    """
    Remove one selected temp row.
    """
    temp_entries = request.session.get("temp_entries", []) or []
    temp_entries = [e for e in temp_entries if e.get("id") != entry_id]
    request.session["temp_entries"] = temp_entries
    return redirect("temp_items", category=category, group=group)


# -----------------------
# TEMP AJAX REMOVE ITEM
# -----------------------
@login_required(login_url='login')
def temp_ajax_remove_item(request, category):
    """
    AJAX endpoint to remove an item from temp entries without page reload.
    POST with JSON: { "entry_id": "..." }
    Returns JSON: { "status": "ok", "entries": [...] }
    """
    if request.method != "POST":
        return JsonResponse({"status": "error", "message": "POST required"}, status=405)
    
    try:
        if request.content_type and 'application/json' in request.content_type:
            data = json.loads(request.body.decode("utf-8") or "{}")
        else:
            data = request.POST
        
        entry_id = data.get("entry_id", "").strip()
        
        if not entry_id:
            return JsonResponse({"status": "error", "message": "No entry_id specified"}, status=400)
        
        temp_entries = request.session.get("temp_entries", []) or []
        temp_entries = [e for e in temp_entries if e.get("id") != entry_id]
        request.session["temp_entries"] = temp_entries
        
        return JsonResponse({
            "status": "ok",
            "entries": temp_entries
        })
        
    except Exception as e:
        return JsonResponse({"status": "error", "message": str(e)}, status=400)


@login_required(login_url='login')
def temp_download_output(request, category):
    """
    Build Output + Estimate for Temporary Works ONLY.
    Uses:
      - entries_json (from hidden field): list of {id, name, qty, days}
      - work_name
    For each entry we:
      - copy the item block up to the matching 'Hire charges per X day(s)' row
      - take rate from that row (column J)
      - add suffix 'for X day(s)' in Estimate description
    New Estimate module is NOT touched by this view.
    """
    if request.method != "POST":
        return redirect("temp_groups", category=category)

    entries_json = request.POST.get("entries_json", "[]")
    try:
        entries = json.loads(entries_json)
        if not isinstance(entries, list):
            entries = []
    except Exception:
        entries = []

    if not entries:
        return redirect("temp_groups", category=category)

    work_name = (request.POST.get("work_name") or "").strip()
    
    # Get selected backend ID from session (for multi-backend support)
    temp_selected_backend_id = request.session.get("temp_selected_backend_id")

    # ----- load backend -----
    try:
        items_list, groups_map, _, ws_src, filepath = load_backend(
            category, settings.BASE_DIR,
            backend_id=temp_selected_backend_id,
            module_code='temp_works'  # Use temp_works module's own backends
        )
    except FileNotFoundError as e:
        logger.error(f"Backend not found for temp download: {category} - {e}")
        return redirect("temp_groups", category=category)
    except Exception as e:
        logger.error(f"Error loading temp backend for download: {category} - {e}")
        return redirect("temp_groups", category=category)
    
    name_to_info = {it["name"]: it for it in items_list}

    # map item -> group for units
    item_to_group = {}
    for grp_name, item_list_in_grp in groups_map.items():
        for nm in item_list_in_grp:
            item_to_group.setdefault(nm, grp_name)

    def units_for(name):
        grp_name = (item_to_group.get(name, "") or "").lower()
        if grp_name in ("piping", "wiring & cables", "wiring and cables"):
            return ("Mtrs", "Mtr")
        elif grp_name == "points":
            return ("Pts", "Pt")
        else:
            return ("Nos", "No")

    # ----- workbook & styles -----
    wb = Workbook()
    ws_out = wb.active
    ws_out.title = "Output"

    thin = Side(border_style="thin", color="000000")
    border_all = Border(top=thin, left=thin, right=thin, bottom=thin)

    cursor = 1  # current row in Output
    rate_rows = []  # each element: (entry_index, row_in_output)

    # =====================================================
    # 1) OUTPUT SHEET: one block per entry (supports dupes)
    # =====================================================
    for idx, entry in enumerate(entries, start=1):
        name = entry.get("name")
        days = int(entry.get("days") or 1)

        info = name_to_info.get(name)
        if not info:
            continue

        src_min = info["start_row"]
        src_max = info["end_row"]

        # find row that matches "Hire charges per X day(s)" nearest this days
        target_row = None
        search1 = f"per {days} day"
        search2 = f"per {days} days"

        for r in range(src_min, src_max + 1):
            txt = str(ws_src.cell(row=r, column=4).value or "").lower()
            if "hire charges per" in txt and (search1 in txt or search2 in txt):
                target_row = r
                break

        if target_row is None:
            # fallback: use entire block and rate = last non-empty J
            effective_end = src_max
            rate_src_row = None
            for r in range(src_max, src_min - 1, -1):
                v = ws_src.cell(row=r, column=10).value
                if v not in (None, ""):
                    rate_src_row = r
                    break
        else:
            effective_end = target_row
            rate_src_row = target_row

        dst_start = cursor

        copy_block_with_styles_and_formulas(
            ws_src=ws_src,
            ws_dst=ws_out,
            src_min_row=src_min,
            src_max_row=effective_end,
            dst_start_row=dst_start,
            col_start=1,
            col_end=10,
        )

        # Label first row as Data block
        ws_out.cell(row=dst_start, column=1).value = f"Data {idx}"

        if rate_src_row:
            rate_rows.append((idx - 1, dst_start + (rate_src_row - src_min)))
        else:
            rate_rows.append((idx - 1, None))

        cursor += (effective_end - src_min + 1)

    # =====================================================
    # 2) ESTIMATE SHEET
    # =====================================================
    ws_est = wb.create_sheet("Estimate")

    # Title row
    ws_est.merge_cells("A1:H1")
    c1 = ws_est["A1"]
    c1.value = "ESTIMATE"
    c1.font = Font(bold=True, size=14)
    c1.alignment = Alignment(horizontal="center", vertical="center")

    ws_est.merge_cells("A2:H2")
    c2 = ws_est["A2"]
    if work_name:
        c2.value = f"Name of the work : {work_name}"
    else:
        c2.value = "Name of the work : "
    c2.font = Font(bold=True, size=11)
    c2.alignment = Alignment(horizontal="left", vertical="center")

    for row in (1, 2):
        for col in range(1, 9):
            ws_est.cell(row=row, column=col).border = border_all

    # Header row
    headers = ["Sl.No", "Quantity (Unit)", "", "Item Description",
               "Rate", "Per Unit", "", "Amount"]
    for col, text in enumerate(headers, start=1):
        cell = ws_est.cell(row=3, column=col, value=text)
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = border_all

    ws_est.merge_cells("B3:C3")
    ws_est.merge_cells("F3:G3")

    ws_est.column_dimensions["A"].width = 6
    ws_est.column_dimensions["B"].width = 10
    ws_est.column_dimensions["C"].width = 10
    ws_est.column_dimensions["D"].width = 45
    ws_est.column_dimensions["E"].width = 10
    ws_est.column_dimensions["F"].width = 8
    ws_est.column_dimensions["G"].width = 10
    ws_est.column_dimensions["H"].width = 15

    # ---- Fill estimate rows ----
    row_est = 4
    slno = 1

    for idx, entry in enumerate(entries, start=1):
        name = entry.get("name")
        qty_val = float(entry.get("qty") or 0) or None
        days = int(entry.get("days") or 1)

        info = name_to_info.get(name)
        if not info:
            continue

        # base description: 2 rows below heading
        start_row = info["start_row"]
        base_desc = ws_src.cell(row=start_row + 2, column=4).value or ""
        base_desc_str = str(base_desc).strip()

        # suffix: "for X day(s)"
        if days == 1:
            suffix = f"for {days} day"
        else:
            suffix = f"for {days} days"

        desc = f"{base_desc_str}, {suffix}" if base_desc_str else suffix

        # rate formula references Output!J<row>
        _, rr = rate_rows[idx - 1]
        rate_formula = f"=Output!J{rr}" if rr else ""

        plural, singular = units_for(name)

        a = ws_est.cell(row=row_est, column=1, value=slno)
        a.alignment = Alignment(horizontal="center", vertical="center")
        a.border = border_all

        b = ws_est.cell(row=row_est, column=2, value=qty_val)
        b.alignment = Alignment(horizontal="center", vertical="center")
        b.border = border_all

        c = ws_est.cell(row=row_est, column=3, value=plural)
        c.alignment = Alignment(horizontal="center", vertical="center")
        c.border = border_all

        d_cell = ws_est.cell(row=row_est, column=4, value=desc)
        d_cell.alignment = Alignment(horizontal="justify", vertical="top", wrap_text=True)
        d_cell.border = border_all

        e = ws_est.cell(row=row_est, column=5, value=rate_formula)
        e.alignment = Alignment(horizontal="center", vertical="center")
        e.border = border_all

        f_cell = ws_est.cell(row=row_est, column=6, value=1)
        f_cell.alignment = Alignment(horizontal="center", vertical="center")
        f_cell.border = border_all

        g = ws_est.cell(row=row_est, column=7, value=singular)
        g.alignment = Alignment(horizontal="center", vertical="center")
        g.border = border_all

        h = ws_est.cell(row=row_est, column=8, value=f"=B{row_est}*E{row_est}")
        h.alignment = Alignment(horizontal="center", vertical="center")
        h.border = border_all

        row_est += 1
        slno += 1

    # ---- Totals (same style as your main estimate) ----
    ecv_row = row_est
    ws_est.cell(row=ecv_row, column=4, value="ECV")
    ws_est.cell(row=ecv_row, column=8, value=f"=SUM(H4:H{ecv_row-1})")

    lc_row = ecv_row + 1
    qc_row = ecv_row + 2
    nac_row = ecv_row + 3
    sub_row = ecv_row + 4
    gst_row = ecv_row + 5
    ls_row = ecv_row + 6
    gt_row = ecv_row + 7

    ws_est.cell(row=lc_row, column=4, value="Add LC @ 1 %")
    ws_est.cell(row=lc_row, column=8, value=f"=H{ecv_row}*0.01")

    ws_est.cell(row=qc_row, column=4, value="Add QC @ 1 %")
    ws_est.cell(row=qc_row, column=8, value=f"=H{ecv_row}*0.01")

    ws_est.cell(row=nac_row, column=4, value="Add NAC @ 0.1 %")
    ws_est.cell(row=nac_row, column=8, value=f"=H{ecv_row}*0.001")

    ws_est.cell(row=sub_row, column=4, value="Sub Total")
    ws_est.cell(row=sub_row, column=8, value=f"=H{ecv_row}+H{lc_row}+H{qc_row}+H{nac_row}")

    ws_est.cell(row=gst_row, column=4, value="Add GST@18 %")
    ws_est.cell(row=gst_row, column=8, value=f"=H{sub_row}*0.18")

    ws_est.cell(row=ls_row, column=4, value="L.S Provision towards unforeseen items")
    ws_est.cell(row=ls_row, column=8, value=f"=H{gt_row}-H{gst_row}-H{sub_row}")

    ws_est.cell(row=gt_row, column=4, value="Grand Total")
    # Set Grand Total value if provided by user, otherwise leave empty
    grand_total_str = request.POST.get("grand_total", "").strip()
    if grand_total_str:
        try:
            grand_total_val = float(grand_total_str)
            if grand_total_val > 0:
                ws_est.cell(row=gt_row, column=8, value=grand_total_val)
        except ValueError:
            pass

    for r in range(ecv_row, gt_row + 1):
        for c in range(1, 9):
            cell = ws_est.cell(row=r, column=c)
            cell.border = border_all
            if c == 4:
                cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            else:
                cell.alignment = Alignment(horizontal="center", vertical="center")
        ws_est.cell(row=r, column=4).font = Font(bold=True)
        ws_est.cell(row=r, column=8).font = Font(bold=True)

    # Reorder sheets: Estimate first, then Output
    if "Estimate" in wb.sheetnames:
        est_idx = wb.sheetnames.index("Estimate")
        if est_idx > 0:
            wb.move_sheet("Estimate", offset=-est_idx)

    # ----- return workbook -----
    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    response["Content-Disposition"] = f'attachment; filename="{category}_temp_output_estimate.xlsx"'
    wb.save(response)
    return response


@login_required(login_url='login')
def estimate(request):
    """
    Estimate module: Upload item blocks sheet (in backend format)
    and generate estimate in the standard format with Output and Estimate sheets.
    
    The uploaded file should contain item blocks in the same format as:
    - core/data/electrical_backend.xlsx
    - core/data/civil_backend.xlsx
    
    Items are detected by yellow background + red text cells in the header row.
    """
    if request.method == 'GET':
        return render(request, 'core/estimate.html', {})
    
    if request.method == 'POST':
        output_file = request.FILES.get('output_file')
        
        if not output_file:
            return render(request, 'core/estimate.html', {
                'error': 'Please upload an Excel file with item blocks.'
            })
        
        # Verify file size and content
        if output_file.size == 0:
            return render(request, 'core/estimate.html', {
                'error': 'Uploaded file is empty. Please select a valid Excel file.'
            })
        
        try:
            # Load the uploaded workbook
            # Reset file position before loading to fix first-attempt upload issue
            try:
                output_file.seek(0)
            except Exception as seek_error:
                pass  # Some file-like objects don't support seek
            
            try:
                wb_upload = load_workbook(output_file, data_only=False)
            except Exception as load_error:
                return render(request, 'core/estimate.html', {
                    'error': f'Failed to read Excel file: {str(load_error)}'
                })
            
            try:
                output_file.seek(0)  # Reset again for second load
            except Exception as seek_error:
                pass
            
            try:
                wb_upload_vals = load_workbook(output_file, data_only=True)
            except Exception as load_error:
                return render(request, 'core/estimate.html', {
                    'error': f'Failed to process Excel file: {str(load_error)}'
                })
            
            # ---- Helper: Check if cell is yellow with red text ----
            def cell_is_yellow(cell):
                fill = cell.fill
                if not fill or not fill.patternType or fill.patternType.lower() != "solid":
                    return False
                rgb = getattr(fill.fgColor, "rgb", None)
                if rgb and str(rgb).upper().endswith("FFFF00"):
                    return True
                if getattr(fill.fgColor, "type", None) == "theme":
                    if getattr(fill.fgColor, "theme", None) in (4, 5, 6):
                        return True
                if getattr(fill.fgColor, "indexed", None) == 6:
                    return True
                return False
            
            def cell_is_red_text(cell):
                font = cell.font
                if not font or not font.color:
                    return False
                rgb = getattr(font.color, "rgb", None)
                if rgb and str(rgb).upper().endswith("FF0000"):
                    return True
                if getattr(font.color, "type", None) == "theme":
                    return True
                if getattr(font.color, "indexed", None) == 3:
                    return True
                return False
            
            def is_yellow_and_red(cell):
                return cell_is_yellow(cell) and cell_is_red_text(cell)
            
            def is_valid_item_block(ws_src, start_row, end_row):
                """
                Check if this block looks like a valid item block (has rate data in column J).
                This helps distinguish real item blocks from headings or signature sections.
                """
                # A valid item block should have at least one non-empty value in column J (rate column)
                for r in range(start_row, min(end_row + 1, start_row + 50)):  # Check up to 50 rows
                    val = ws_src.cell(row=r, column=10).value
                    if val not in (None, "") and str(val).strip():
                        return True
                return False
            
            def find_item_block_end(ws_src, start_row, max_row):
                """
                Find the true end of an item block by looking for the rate row in column J.
                The block ends at the last row that has meaningful data before the next heading
                or signature section.
                """
                # First, find where the next yellow+red heading is
                next_heading_row = max_row + 1
                for rr in range(start_row + 1, max_row + 1):
                    for c in range(1, 11):
                        cell = ws_src.cell(row=rr, column=c)
                        if is_yellow_and_red(cell) and str(cell.value or "").strip():
                            next_heading_row = rr
                            break
                    if next_heading_row <= max_row:
                        break
                
                # The block should end before the next heading
                potential_end = next_heading_row - 1
                
                # Find the last row with rate data (column J) - this is the true end of item block
                last_rate_row = start_row
                for r in range(start_row, potential_end + 1):
                    val = ws_src.cell(row=r, column=10).value
                    if val not in (None, "") and str(val).strip():
                        last_rate_row = r
                
                # The block ends at the last rate row (don't include signature/footer content)
                return last_rate_row, next_heading_row
            
            def extract_items_from_sheet(ws_src):
                """Extract all item blocks from a single sheet."""
                fetched_items = []
                item_blocks = {}  # name -> (start_row, end_row)
                
                max_row = ws_src.max_row
                r = 1
                first_item_found = False
                
                while r <= max_row:
                    heading_name = None
                    heading_col = None
                    # Check columns A..J for yellow+red cell
                    for c in range(1, 11):
                        cell = ws_src.cell(row=r, column=c)
                        if is_yellow_and_red(cell) and str(cell.value or "").strip():
                            heading_name = str(cell.value).strip()
                            heading_col = c
                            break
                    
                    if heading_name:
                        start_row = r
                        
                        # Find the proper end of this block
                        end_row, next_heading_row = find_item_block_end(ws_src, start_row, max_row)
                        
                        # Validate this is a real item block (not a sheet heading or signature section)
                        if is_valid_item_block(ws_src, start_row, end_row):
                            fetched_items.append(heading_name)
                            item_blocks[heading_name] = (start_row, end_row)
                            first_item_found = True
                        
                        # Move to next heading position
                        r = next_heading_row if next_heading_row <= max_row else end_row + 1
                    else:
                        r += 1
                
                return fetched_items, item_blocks
            
            def create_output_and_estimate_sheets(wb_out, ws_src, fetched_items, item_blocks, 
                                                   output_sheet_name, estimate_sheet_name):
                """Create Output and Estimate sheets for a single source sheet."""
                thin = Side(border_style="thin", color="000000")
                border_all = Border(top=thin, left=thin, right=thin, bottom=thin)
                
                # Create Output sheet
                ws_out = wb_out.create_sheet(output_sheet_name)
                
                # Build Output sheet by copying item blocks
                cursor = 1
                rate_pos = {}
                data_serial = 1
                
                for item_name in fetched_items:
                    src_min, src_max = item_blocks[item_name]
                    
                    # Find rate row (non-empty in column J)
                    rate_src_row = None
                    for r in range(src_max, src_min - 1, -1):
                        v = ws_src.cell(row=r, column=10).value
                        if v not in (None, ""):
                            rate_src_row = r
                            break
                    
                    dst_start = cursor
                    
                    # Copy block with styles
                    copy_block_with_styles_and_formulas(
                        ws_src=ws_src,
                        ws_dst=ws_out,
                        src_min_row=src_min,
                        src_max_row=src_max,
                        dst_start_row=dst_start,
                        col_start=1,
                        col_end=10
                    )
                    
                    ws_out.cell(row=dst_start, column=1).value = f"Data {data_serial}"
                    data_serial += 1
                    
                    # Store rate row position
                    if rate_src_row:
                        rate_pos[item_name] = dst_start + (rate_src_row - src_min)
                    
                    cursor += (src_max - src_min + 1)
                
                # Create Estimate sheet
                ws_est = wb_out.create_sheet(estimate_sheet_name)
                
                # Title row
                ws_est.merge_cells("A1:H1")
                c1 = ws_est["A1"]
                c1.value = "ESTIMATE"
                c1.font = Font(bold=True, size=14)
                c1.alignment = Alignment(horizontal="center", vertical="center")
                
                # Title row 2
                ws_est.merge_cells("A2:H2")
                c2 = ws_est["A2"]
                c2.value = "Name of the work : "
                c2.font = Font(bold=True, size=11)
                c2.alignment = Alignment(horizontal="left", vertical="center")
                
                for row in (1, 2):
                    for col in range(1, 9):
                        ws_est.cell(row=row, column=col).border = border_all
                
                # Header row
                headers = ["Sl.No", "Quantity (Unit)", "", "Item Description",
                           "Rate", "Per Unit", "", "Amount"]
                
                for col, text in enumerate(headers, start=1):
                    cell = ws_est.cell(row=3, column=col, value=text)
                    cell.font = Font(bold=True)
                    cell.alignment = Alignment(horizontal="center", vertical="center")
                    cell.border = border_all
                
                ws_est.merge_cells("B3:C3")
                ws_est.merge_cells("F3:G3")
                
                ws_est.column_dimensions["A"].width = 6
                ws_est.column_dimensions["B"].width = 10
                ws_est.column_dimensions["C"].width = 10
                ws_est.column_dimensions["D"].width = 38
                ws_est.column_dimensions["E"].width = 10
                ws_est.column_dimensions["F"].width = 8
                ws_est.column_dimensions["G"].width = 10
                ws_est.column_dimensions["H"].width = 15
                
                # Fill estimate rows
                row_est = 4
                slno = 1
                
                def to_plural(unit):
                    """Convert unit to plural form."""
                    unit_lower = unit.lower()
                    if unit_lower == "no":
                        return "Nos"
                    elif unit_lower == "nos":
                        return "Nos"
                    elif unit_lower == "mtr":
                        return "Mtrs"
                    elif unit_lower == "mtrs":
                        return "Mtrs"
                    elif unit_lower == "pts":
                        return "Pts"
                    elif unit_lower == "pt":
                        return "Pts"
                    elif unit_lower == "cum":
                        return "Cum"
                    elif unit_lower == "kg":
                        return "Kg"
                    elif unit_lower == "l":
                        return "L"
                    elif unit_lower == "kg":
                        return "Kg"
                    else:
                        return unit + "s" if unit else "Nos"
                
                def to_singular(unit):
                    """Convert unit to singular form."""
                    unit_lower = unit.lower()
                    if unit_lower == "nos":
                        return "No"
                    elif unit_lower == "no":
                        return "No"
                    elif unit_lower == "mtrs":
                        return "Mtr"
                    elif unit_lower == "mtr":
                        return "Mtr"
                    elif unit_lower == "pts":
                        return "Pt"
                    elif unit_lower == "pt":
                        return "Pt"
                    elif unit_lower == "cum":
                        return "Cum"
                    elif unit_lower == "kg":
                        return "Kg"
                    elif unit_lower == "l":
                        return "L"
                    else:
                        return unit
                
                def determine_unit_from_heading(heading_name):
                    """
                    Intelligently determine unit based on heading name meaning.
                    """
                    heading_lower = heading_name.lower()
                    
                    # Light Point or Fan Point â†’ Pts
                    if "light point" in heading_lower or "fan point" in heading_lower:
                        return "Pts"
                    
                    # Light or Fan (fixtures/bulbs) â†’ Nos (check BEFORE pipe keywords to avoid "tube light" being "Mtr")
                    light_fan_keywords = ["light", "fan", "bulb", "fixture", "downlight", "spotlight", "batten"]
                    for keyword in light_fan_keywords:
                        if keyword in heading_lower:
                            return "Nos"
                    
                    # Pipes, wires, cables â†’ Mtr (meters)
                    pipe_keywords = ["pipe", "wire", "cable", "conduit", "duct", "channel", "rod", "bar", "rail", "tube"]
                    for keyword in pipe_keywords:
                        if keyword in heading_lower:
                            return "Mtr"
                    
                    # Points â†’ Pts
                    if "point" in heading_lower or "pts" in heading_lower:
                        return "Pts"
                    
                    # Default to Nos
                    return "Nos"
                
                for item_name in fetched_items:
                    src_min, src_max = item_blocks[item_name]
                    
                    # Get the description from the second row (src_min + 2)
                    base_desc = ws_src.cell(row=src_min + 2, column=4).value or ""
                    base_desc_str = str(base_desc).strip()
                    desc = base_desc_str

                    # Determine unit intelligently from heading name
                    best_unit = determine_unit_from_heading(item_name)
                    unit_plural = to_plural(best_unit)
                    unit_singular = to_singular(best_unit)

                    # Rate from Output sheet (reference the correct output sheet)
                    rr = rate_pos.get(item_name)
                    # Excel sheet names with spaces need quotes
                    safe_output_name = f"'{output_sheet_name}'" if ' ' in output_sheet_name else output_sheet_name
                    rate_formula = f"={safe_output_name}!J{rr}" if rr else ""

                    # Write row
                    a = ws_est.cell(row=row_est, column=1, value=slno)
                    a.alignment = Alignment(horizontal="center", vertical="center")
                    a.border = border_all

                    b = ws_est.cell(row=row_est, column=2, value="")
                    b.alignment = Alignment(horizontal="center", vertical="center")
                    b.border = border_all

                    c = ws_est.cell(row=row_est, column=3, value=unit_plural)
                    c.alignment = Alignment(horizontal="center", vertical="center")
                    c.border = border_all

                    d = ws_est.cell(row=row_est, column=4, value=desc)
                    d.alignment = Alignment(horizontal="justify", vertical="top", wrap_text=True)
                    d.border = border_all

                    e = ws_est.cell(row=row_est, column=5, value=rate_formula)
                    e.alignment = Alignment(horizontal="center", vertical="center")
                    e.border = border_all

                    f = ws_est.cell(row=row_est, column=6, value=1)
                    f.alignment = Alignment(horizontal="center", vertical="center")
                    f.border = border_all

                    g = ws_est.cell(row=row_est, column=7, value=unit_singular)
                    g.alignment = Alignment(horizontal="center", vertical="center")
                    g.border = border_all

                    h = ws_est.cell(row=row_est, column=8, value=f"=B{row_est}*E{row_est}")
                    h.alignment = Alignment(horizontal="center", vertical="center")
                    h.border = border_all

                    row_est += 1
                    slno += 1
                
                # ---- Add totals rows ----
                ecv_row = row_est
                ws_est.cell(row=ecv_row, column=4, value="ECV")
                ws_est.cell(row=ecv_row, column=8, value=f"=SUM(H4:H{ecv_row-1})")
                
                lc_row = ecv_row + 1
                qc_row = ecv_row + 2
                nac_row = ecv_row + 3
                sub_row = ecv_row + 4
                gst_row = ecv_row + 5
                ls_row = ecv_row + 6
                gt_row = ecv_row + 7
                
                ws_est.cell(row=lc_row, column=4, value="Add LC @ 1 %")
                ws_est.cell(row=lc_row, column=8, value=f"=H{ecv_row}*0.01")
                
                ws_est.cell(row=qc_row, column=4, value="Add QC @ 1 %")
                ws_est.cell(row=qc_row, column=8, value=f"=H{ecv_row}*0.01")
                
                ws_est.cell(row=nac_row, column=4, value="Add NAC @ 0.1 %")
                ws_est.cell(row=nac_row, column=8, value=f"=H{ecv_row}*0.001")
                
                ws_est.cell(row=sub_row, column=4, value="Sub Total")
                ws_est.cell(row=sub_row, column=8, value=f"=H{ecv_row}+H{lc_row}+H{qc_row}+H{nac_row}")
                
                ws_est.cell(row=gst_row, column=4, value="Add GST@18 %")
                ws_est.cell(row=gst_row, column=8, value=f"=H{sub_row}*0.18")
                
                ws_est.cell(row=ls_row, column=4, value="L.S Provision towards unforeseen items")
                ws_est.cell(row=ls_row, column=8, value=f"=H{gt_row}-H{gst_row}-H{sub_row}")
                
                ws_est.cell(row=gt_row, column=4, value="Grand Total")
                
                # Apply borders to totals
                for r in range(ecv_row, gt_row + 1):
                    for c in range(1, 9):
                        cell = ws_est.cell(row=r, column=c)
                        cell.border = border_all
                        if c == 4:
                            cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
                        else:
                            cell.alignment = Alignment(horizontal="center", vertical="center")
                    ws_est.cell(row=r, column=4).font = Font(bold=True)
                    ws_est.cell(row=r, column=8).font = Font(bold=True)
                
                return ws_out, ws_est
            
            # ---- Multi-sheet processing ----
            # Find all sheets with item blocks
            sheets_with_items = []
            for sheet_name in wb_upload.sheetnames:
                ws_src = wb_upload[sheet_name]
                fetched_items, item_blocks = extract_items_from_sheet(ws_src)
                if fetched_items:
                    sheets_with_items.append({
                        'name': sheet_name,
                        'ws_src': ws_src,
                        'fetched_items': fetched_items,
                        'item_blocks': item_blocks
                    })
            
            if not sheets_with_items:
                return render(request, 'core/estimate.html', {
                    'error': 'No item blocks found in any sheet. Make sure item headers have yellow background and red text.'
                })
            
            # Create output workbook
            wb_out = Workbook()
            # Remove the default sheet, we'll create our own
            default_sheet = wb_out.active
            
            # Generate Output and Estimate sheets for each source sheet with items
            total_sheets = len(sheets_with_items)
            
            for idx, sheet_info in enumerate(sheets_with_items):
                src_sheet_name = sheet_info['name']
                ws_src = sheet_info['ws_src']
                fetched_items = sheet_info['fetched_items']
                item_blocks = sheet_info['item_blocks']
                
                # Determine sheet names
                if total_sheets == 1:
                    # Single sheet: use simple names
                    output_sheet_name = "Datas"
                    estimate_sheet_name = "Estimate"
                else:
                    # Multiple sheets: append source sheet name
                    # Truncate to fit Excel's 31 character limit
                    base_name = src_sheet_name[:20] if len(src_sheet_name) > 20 else src_sheet_name
                    output_sheet_name = f"Datas_{base_name}"[:31]
                    estimate_sheet_name = f"Estimate_{base_name}"[:31]
                
                create_output_and_estimate_sheets(
                    wb_out=wb_out,
                    ws_src=ws_src,
                    fetched_items=fetched_items,
                    item_blocks=item_blocks,
                    output_sheet_name=output_sheet_name,
                    estimate_sheet_name=estimate_sheet_name
                )
            
            # Remove the default empty sheet if we created our own
            if default_sheet.title in wb_out.sheetnames and len(wb_out.sheetnames) > 1:
                wb_out.remove(default_sheet)
            
            # Reorder sheets: Estimate followed by its corresponding Datas sheet
            # Pattern: Estimate_1, Datas_1, Estimate_2, Datas_2, etc.
            estimate_sheets = [s for s in wb_out.sheetnames if s.startswith("Estimate") or s == "Estimate"]
            output_sheets = [s for s in wb_out.sheetnames if s.startswith("Datas") or s == "Datas"]
            
            # Build pairs based on suffix matching
            ordered_sheets = []
            for est_name in estimate_sheets:
                ordered_sheets.append(est_name)
                # Find matching Datas sheet
                if est_name == "Estimate":
                    # Single sheet case
                    if "Datas" in output_sheets:
                        ordered_sheets.append("Datas")
                else:
                    # Multi-sheet case: Estimate_XYZ -> Datas_XYZ
                    suffix = est_name[8:]  # Remove "Estimate" prefix (8 chars)
                    output_name = f"Datas{suffix}"
                    if output_name in output_sheets:
                        ordered_sheets.append(output_name)
            
            # Add any remaining output sheets that weren't paired
            for out_name in output_sheets:
                if out_name not in ordered_sheets:
                    ordered_sheets.append(out_name)
            
            # Reorder sheets according to the new order
            for i, sheet_name in enumerate(ordered_sheets):
                if sheet_name in wb_out.sheetnames:
                    current_idx = wb_out.sheetnames.index(sheet_name)
                    if current_idx != i:
                        wb_out.move_sheet(sheet_name, offset=(i - current_idx))
            
            # Return the estimate workbook
            response = HttpResponse(
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = 'attachment; filename="estimate.xlsx"'
            wb_out.save(response)
            return response
            
        except Exception as e:
            import traceback
            return render(request, 'core/estimate.html', {
                'error': f'Error processing file: {str(e)}'
            })


@org_required
@login_required
def download_specification_report(request, estimate_id):
    """Generate and download specification report as Word document"""
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    org = request.organization
    estimate = get_object_or_404(Estimate, id=estimate_id, organization=org)
    
    try:
        # Create Word document
        doc = Document()
        
        # Title
        title = doc.add_heading('Specification report accompanying the estimate :-', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in title.runs:
            run.font.size = Pt(12)
            run.font.bold = True
        
        # Introduction paragraph with work name
        estimate_data = estimate.estimate_data or {}
        work_name = estimate_data.get('name_of_work', 'the work')
        
        intro_text = f'The estimate is prepared for the work {work_name}'
        intro_para = doc.add_paragraph(intro_text)
        intro_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in intro_para.runs:
            run.font.size = Pt(11)
        
        # Estimate amount
        total_amount = estimate.total_amount or 0
        amount_para = doc.add_paragraph()
        amount_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        amount_run = amount_para.add_run(f'Est.Amount: Rs. {total_amount:,.2f}')
        amount_run.font.size = Pt(11)
        amount_run.font.bold = True
        amount_run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
        
        doc.add_paragraph()  # Blank line
        
        # Body of letter (manually entered, placeholder for user to edit)
        body_label = doc.add_paragraph('{{BODY_OF_LETTER}}')
        body_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in body_label.runs:
            run.font.size = Pt(11)
            run.font.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)  # Gray placeholder
        
        doc.add_paragraph()  # Blank line
        
        doc.add_paragraph('Hence, this estimate has been prepared accordingly.')
        
        doc.add_paragraph()  # Blank line
        
        # Extract items with quantities and units as bullet points
        # Support both 'items' key and 'ws_estimate_rows' key for different estimate formats
        estimate_items = estimate_data.get('ws_estimate_rows', estimate_data.get('items', []))
        
        for item in estimate_items:
            # Extract values from item structure
            # Support multiple key names: desc/description, qty_est/qty/quantity
            item_description = item.get('desc', item.get('description', item.get('display_name', '')))
            quantity = item.get('qty_est', item.get('qty', item.get('quantity', '')))
            unit = item.get('unit', '')
            
            # Format: "Description  -  Quantity Unit"
            if quantity and unit:
                # Convert quantity to clean format
                qty_str = str(quantity).strip()
                if '.' in qty_str and qty_str.endswith('.0'):
                    qty_str = qty_str.replace('.0', '')
                bullet_text = f'{item_description}  -  {qty_str} {unit}'
            elif quantity:
                qty_str = str(quantity).strip()
                if '.' in qty_str and qty_str.endswith('.0'):
                    qty_str = qty_str.replace('.0', '')
                bullet_text = f'{item_description}  -  {qty_str}'
            else:
                bullet_text = item_description
            
            bullet_para = doc.add_paragraph(bullet_text, style='List Bullet')
            for run in bullet_para.runs:
                run.font.size = Pt(11)
        
        doc.add_paragraph()  # Blank line
        
        # Calculate financial year (April to March)
        from datetime import datetime
        today = datetime.now()
        if today.month >= 4:  # April onwards
            fy_start = today.year
            fy_end = (today.year + 1) % 100
        else:  # Jan to March
            fy_start = today.year - 1
            fy_end = today.year % 100
        financial_year = f"{fy_start}-{fy_end:02d}"
        
        # Footer text about rates and provisions
        footer_text = (f'The rates proposed in the estimate are as per SQR {financial_year} and Approved rates. L.S. Provision is made in the '
                      'estimate towards GST at 18%, QC amount at 1%, Labour Cess at 1% and NAC amount at 0.1% as per actual and LS Provision Towards, unforeseen items & rounding off also proposed in the estimate.')
        footer_para = doc.add_paragraph(footer_text)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in footer_para.runs:
            run.font.size = Pt(10)
        
        doc.add_paragraph()  # Blank line
        
        # Funds section
        funds_text = ('FUNDS: The estimate requires Administrative sanction and also fixes up the agency with provision of funds '
                     'under relevant head of account for taking up the work from the Government. Telangana State Hyderabad')
        funds_para = doc.add_paragraph(funds_text)
        funds_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in funds_para.runs:
            run.font.size = Pt(10)
            run.font.bold = True
        
        # Generate filename
        safe_work_name = work_name.replace(" ", "_").replace("/", "_")[:30]
        filename = f'Specification_Report_{safe_work_name}_{timezone.now().strftime("%Y%m%d_%H%M%S")}.docx'
        
        # Save to response
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        doc.save(response)
        return response
        
    except Exception as e:
        logger.error(f'Error generating specification report: {str(e)}', exc_info=True)
        from django.contrib import messages
        messages.error(request, f'Error generating report: {str(e)}')
        return redirect('view_estimate', estimate_id=estimate_id)


@login_required(login_url='login')
def generate_specification_report_from_file(request):
    """
    Generate specification report from uploaded Excel workbook.
    Extracts item headings (red text + yellow background) from Item Blocks sheet
    and matches quantities/units from the Estimate sheet.
    """
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    if request.method != 'POST':
        return render(request, 'core/estimate.html', {
            'error': 'Invalid request method'
        })
    
    output_file = request.FILES.get('output_file')
    if not output_file:
        return render(request, 'core/estimate.html', {
            'error': 'Please upload an Excel file.'
        })
    
    try:
        # Reset file position
        try:
            output_file.seek(0)
        except:
            pass
        
        wb_upload = load_workbook(output_file, data_only=False)
        
        # Also load with data_only=True for values
        output_file.seek(0)
        wb_values = load_workbook(output_file, data_only=True)
        
        # Helper functions for cell detection
        def cell_is_yellow(cell):
            fill = cell.fill
            if not fill or not fill.fgColor:
                return False
            fg = fill.fgColor
            # Check RGB values
            rgb = getattr(fg, "rgb", None)
            if rgb:
                rgb_upper = str(rgb).upper()
                if rgb_upper in ('FFFFFF00', 'FFFF00', 'FFFFE000', 'FFFFC000', 'FFFFCC00'):
                    return True
                # Check if it ends with yellow-ish color
                if rgb_upper.endswith('FFFF00') or (rgb_upper.endswith('FF00') and 'FF' in rgb_upper[:4]):
                    return True
            # Check indexed colors (Excel default palette)
            if getattr(fg, "indexed", None) in (6, 13):
                return True
            # Check theme colors
            if getattr(fg, "type", None) == "theme":
                return True
            return False
        
        def cell_is_red_text(cell):
            font = cell.font
            if not font or not font.color:
                return False
            c = font.color
            rgb = getattr(c, "rgb", None)
            if rgb:
                rgb_upper = str(rgb).upper()
                if rgb_upper in ('FFFF0000', 'FF0000', 'FFC00000', 'FFFF0000'):
                    return True
                if rgb_upper.endswith('FF0000'):
                    return True
            # Check indexed red
            if getattr(c, "indexed", None) == 3:
                return True
            return False
        
        def is_yellow_and_red(cell):
            return cell_is_yellow(cell) and cell_is_red_text(cell)
        
        # Find sheets - look for "Estimate" and "Item Blocks" sheets
        estimate_sheet = None
        estimate_sheet_values = None
        item_blocks_sheet = None
        
        for sheet_name in wb_upload.sheetnames:
            sheet_lower = sheet_name.lower().strip()
            if 'estimate' in sheet_lower and 'item' not in sheet_lower and 'datas' not in sheet_lower:
                estimate_sheet = wb_upload[sheet_name]
                estimate_sheet_values = wb_values[sheet_name]
            elif 'item' in sheet_lower and 'block' in sheet_lower:
                item_blocks_sheet = wb_upload[sheet_name]
            elif 'datas' in sheet_lower:
                # "Datas" sheet - treat as item blocks sheet
                item_blocks_sheet = wb_upload[sheet_name]
        
        # Fallback: If no specific sheets found, look for sheet with yellow+red items
        if not item_blocks_sheet:
            # Try to find any sheet with yellow+red items
            for sheet_name in wb_upload.sheetnames:
                ws = wb_upload[sheet_name]
                for row in range(1, min(50, ws.max_row + 1)):
                    for col in range(1, 10):
                        cell = ws.cell(row=row, column=col)
                        if is_yellow_and_red(cell):
                            item_blocks_sheet = ws
                            break
                    if item_blocks_sheet:
                        break
                if item_blocks_sheet:
                    break
        
        if not item_blocks_sheet:
            return render(request, 'core/estimate.html', {
                'error': 'Could not find Datas sheet. Please ensure your workbook has a sheet containing items with red text and yellow background.'
            })
        
        # Step 1: Extract red+yellow items from Item Blocks sheet
        item_headings = []
        for row in range(1, item_blocks_sheet.max_row + 1):
            for col in range(1, 15):  # Check columns A through O
                cell = item_blocks_sheet.cell(row=row, column=col)
                if is_yellow_and_red(cell):
                    heading = str(cell.value or "").strip()
                    if not heading or len(heading) < 3:
                        continue
                    
                    # Skip if it looks like a label/header
                    heading_lower = heading.lower()
                    if heading_lower in ('sl', 'sl.', 'sl.no', 'description', 'unit', 'qty', 'rate', 'amount', 'total'):
                        continue
                    
                    # Avoid duplicates
                    if heading not in item_headings:
                        item_headings.append(heading)
                    break  # Move to next row after finding a heading
        
        if not item_headings:
            return render(request, 'core/estimate.html', {
                'error': 'No item headings found with red text and yellow background in the workbook.'
            })
        
        # Step 2: Build quantity/unit map from Estimate sheet
        qty_unit_map = {}  # heading -> {'qty': ..., 'unit': ...}
        work_name = ""
        total_amount = ""
        
        if estimate_sheet:
            # Find work name - look for "Name of the work" or "Name of work" 
            # The work name might be in the same cell after a colon, or in adjacent cells
            for row in range(1, min(30, estimate_sheet.max_row + 1)):
                if work_name:
                    break
                for col in range(1, 10):
                    cell_val = str(estimate_sheet.cell(row=row, column=col).value or "").strip()
                    cell_lower = cell_val.lower()
                    
                    # Check if this cell contains "name of the work" or "name of work"
                    if "name of the work" in cell_lower or "name of work" in cell_lower:
                        # Check if work name is in the same cell after colon
                        if ':' in cell_val:
                            parts = cell_val.split(':', 1)
                            if len(parts) > 1 and parts[1].strip():
                                work_name = parts[1].strip()
                                break
                        
                        # Try next columns in same row
                        if not work_name:
                            for next_col in range(col + 1, col + 8):
                                next_cell = estimate_sheet.cell(row=row, column=next_col).value
                                if next_cell and str(next_cell).strip():
                                    work_name = str(next_cell).strip()
                                    break
                        
                        # Try next row
                        if not work_name:
                            next_row_val = estimate_sheet.cell(row=row + 1, column=col).value
                            if next_row_val and str(next_row_val).strip():
                                work_name = str(next_row_val).strip()
                        break
            
            # Find Grand Total / Estimate Amount
            # Search for "Grand Total" row and get the amount from adjacent column
            for row in range(1, estimate_sheet.max_row + 1):
                for col in range(1, 10):
                    cell_val = str(estimate_sheet.cell(row=row, column=col).value or "").strip().lower()
                    if 'grand total' in cell_val:
                        # Look for amount in columns to the right (E, F, G, H)
                        for amt_col in range(col + 1, col + 6):
                            if estimate_sheet_values:
                                amt_cell = estimate_sheet_values.cell(row=row, column=amt_col)
                            else:
                                amt_cell = estimate_sheet.cell(row=row, column=amt_col)
                            if amt_cell.value is not None:
                                try:
                                    amt_val = float(amt_cell.value)
                                    if amt_val > 1000:  # Grand total should be a significant amount
                                        # Format as currency
                                        total_amount = f"{amt_val:,.2f}"
                                        break
                                except (ValueError, TypeError):
                                    pass
                        if total_amount:
                            break
                if total_amount:
                    break
            
            # Search Estimate sheet for each item heading to get qty and unit
            # Structure: Column A=Sl.No, Column B=Quantity, Column C=Unit, Column D=Description
            
            # Helper function to check if item heading matches description using keywords
            def heading_matches_description(heading, description):
                heading_lower = heading.lower().strip()
                desc_lower = description.lower().strip()
                
                # Skip if description is too short (likely a header row without real data)
                if len(desc_lower) < 15:
                    return False
                
                # Exact match
                if heading_lower == desc_lower:
                    return True
                
                # Heading appears exactly in description
                if heading_lower in desc_lower:
                    return True
                
                # Normalize description - remove extra spaces
                desc_normalized = ' '.join(desc_lower.split())
                
                # Helper to check if all words appear in description
                def all_words_in_desc(words):
                    return all(w in desc_normalized for w in words)
                
                # Helper to check if any word appears in description
                def any_word_in_desc(words):
                    return any(w in desc_normalized for w in words)
                
                # Specific matching rules for electrical items
                
                # PVC Pipe matching - must distinguish concealed vs surface
                if 'concealed' in heading_lower and 'pvc' in heading_lower:
                    return 'concealed' in desc_normalized and 'pvc' in desc_normalized and 'surface' not in desc_normalized
                
                if 'surface' in heading_lower and 'pvc' in heading_lower:
                    return 'surface' in desc_normalized and 'pvc' in desc_normalized
                
                # Light & Bell Points - row 8 type items
                if ('light' in heading_lower and 'bell' in heading_lower) or ('light' in heading_lower and 'point' in heading_lower):
                    # Match descriptions with "light and bell point" or "for light" 
                    if 'light and bell' in desc_normalized:
                        return True
                    if 'for light' in desc_normalized and 'bell' in desc_normalized and 'point' in desc_normalized:
                        return True
                    if 'light point' in desc_normalized:
                        return True
                    return False
                
                # Fan and Exhaust Fan Points - row 9 type items  
                if 'fan' in heading_lower and 'exhaust' in heading_lower and 'point' in heading_lower:
                    return ('exhaust' in desc_normalized and 'fan' in desc_normalized and 'point' in desc_normalized and 
                            'light and bell' not in desc_normalized)
                
                # Common Switch Board - row 11
                if 'switch' in heading_lower and 'board' in heading_lower:
                    # Check for "switch board" or "switch  board" with any spacing, or switchboard
                    if 'switch' in desc_normalized and 'board' in desc_normalized:
                        return True
                    if 'switchboard' in desc_normalized:
                        return True
                    return False
                
                # Power Plug - row 13, 14
                if 'power' in heading_lower and 'plug' in heading_lower:
                    # Power Plug items - look for 16A socket/plug descriptions
                    # "Two Nos" in heading is part of item name (2-gang type), not quantity
                    if 'two' in heading_lower:
                        # Power Plug Two Nos - has "2 Nos" or "2 Module" or "2 way" in description
                        return ('16a' in desc_normalized or '16 a' in desc_normalized) and \
                               ('2 nos' in desc_normalized or '2 no.s' in desc_normalized or 
                                '2 module' in desc_normalized or '2 way' in desc_normalized or 
                                'two nos' in desc_normalized or 'twin' in desc_normalized)
                    else:
                        # Regular single power plug - should NOT have "2 Nos" pattern
                        has_16a = '16a' in desc_normalized or '16 a' in desc_normalized
                        has_plug = 'socket' in desc_normalized or 'plug' in desc_normalized
                        is_two_type = '2 nos' in desc_normalized or '2 no.s' in desc_normalized or 'two nos' in desc_normalized
                        return has_16a and has_plug and not is_two_type
                
                # Module (sockets/switches) - rows 16, 17, 18
                # "Two Nos", "Three Nos", "Four Nos" in heading = type of module (number of socket+switch combos)
                # Descriptions mention: "2 no.s", "3 Nos", "4 Nos" and modular box sizes: 6, 8, 12 Modular
                if 'module' in heading_lower:
                    # Must have "modular" in description
                    if 'modular' not in desc_normalized:
                        return False
                    
                    # 6Module = Two Nos = 2 socket+switch combos, uses common switch board
                    if '6module' in heading_lower.replace(' ', '') or ('two nos' in heading_lower):
                        # Look for "2 nos" or "2 no.s" with common switch board or 6 modular
                        has_two = '2 nos' in desc_normalized or '2 no.s' in desc_normalized or '2nos' in desc_normalized
                        has_common_board = 'common switch board' in desc_normalized or 'common switch  board' in desc_normalized
                        # Exclude if it's 8 or 12 modular box
                        is_larger_box = '8 modular' in desc_normalized or '12 modular' in desc_normalized
                        return has_two and has_common_board and not is_larger_box
                    
                    # 8Module = Three Nos = 3 socket+switch combos, uses 8 Modular box
                    elif '8module' in heading_lower.replace(' ', '') or ('three nos' in heading_lower):
                        has_three = '3 nos' in desc_normalized or '3 no.s' in desc_normalized or '3nos' in desc_normalized
                        has_8_box = '8 modular' in desc_normalized
                        return has_three and has_8_box
                    
                    # 12Module = Four Nos = 4 socket+switch combos, uses 12 Modular box
                    elif '12module' in heading_lower.replace(' ', '') or ('four nos' in heading_lower):
                        has_four = '4 nos' in desc_normalized or '4 no.s' in desc_normalized or '4nos' in desc_normalized
                        has_12_box = '12 modular' in desc_normalized
                        return has_four and has_12_box
                    
                    return 'modular' in desc_normalized and 'socket' in desc_normalized
                
                # Ding Dong Bell
                if 'ding' in heading_lower and 'dong' in heading_lower:
                    return 'ding dong' in desc_normalized or ('calling bell' in desc_normalized)
                if 'bell' in heading_lower and 'ding' not in heading_lower and 'point' not in heading_lower:
                    return 'calling bell' in desc_normalized or 'door bell' in desc_normalized
                
                # LED light
                if 'led' in heading_lower:
                    if '1200' in heading_lower or 'length' in heading_lower:
                        return 'led' in desc_normalized and ('1200' in desc_normalized or 'tube' in desc_normalized or 'batten' in desc_normalized)
                    return 'led' in desc_normalized
                
                # Ceiling Fans
                if 'ceiling' in heading_lower and 'fan' in heading_lower:
                    return 'ceiling fan' in desc_normalized or ('ceiling' in desc_normalized and 'fan' in desc_normalized)
                
                # Stepped Electronic Regulator
                if 'regulator' in heading_lower:
                    return 'regulator' in desc_normalized
                
                # Exhaust fans (not points)
                if 'exhaust' in heading_lower and 'fan' in heading_lower and 'point' not in heading_lower:
                    if 'kitchen' in heading_lower or 'metallic' in heading_lower or '12' in heading_lower:
                        # 12" Metallic Exhaust Fan for Kitchen
                        return ('exhaust' in desc_normalized and 'fan' in desc_normalized) and \
                               ('kitchen' in desc_normalized or 'metallic' in desc_normalized or '12' in desc_normalized or '300mm' in desc_normalized)
                    elif 'bathroom' in heading_lower or 'shutter' in heading_lower or '6' in heading_lower:
                        # 6" Shutter type Exhaust Fan for Bathrooms
                        return ('exhaust' in desc_normalized and 'fan' in desc_normalized) and \
                               ('bathroom' in desc_normalized or 'shutter' in desc_normalized or '6' in desc_normalized or '150mm' in desc_normalized)
                    return 'exhaust' in desc_normalized and 'fan' in desc_normalized
                
                # Copper cable runs
                if 'copper' in heading_lower and 'cable' in heading_lower:
                    if '2.5' in heading_lower:
                        return 'copper' in desc_normalized and '2.5' in desc_normalized
                    elif '4.0' in heading_lower or '4 sq' in heading_lower:
                        return 'copper' in desc_normalized and ('4.0' in desc_normalized or '4 sq' in desc_normalized or '4sq' in desc_normalized)
                    return 'copper' in desc_normalized and 'cable' in desc_normalized
                
                # WPTC
                if 'wptc' in heading_lower:
                    return 'wptc' in desc_normalized
                
                # TPN DB
                if 'tpn' in heading_lower:
                    return 'tpn' in desc_normalized
                if 'db' in heading_lower and 'way' in heading_lower:
                    return 'db' in desc_normalized and 'way' in desc_normalized
                
                # Water heater / Geyser
                if 'geyser' in heading_lower or 'water heater' in heading_lower:
                    return 'geyser' in desc_normalized or 'water heater' in desc_normalized or 'gyser' in desc_normalized
                
                # Generic fallback - require ALL significant words to be present
                stop_words = {'and', 'or', 'the', 'a', 'an', 'in', 'on', 'of', 'for', 'with', 'to', 'nos', 'no', 'type', 'rb', 'r.b', 'r.b.', 'n.r.b', 'n.r.b.'}
                heading_words = [w.strip() for w in heading_lower.replace('&', ' ').replace('-', ' ').replace('/', ' ').replace('(', ' ').replace(')', ' ').split() 
                                if len(w.strip()) > 2 and w.strip() not in stop_words]
                
                if heading_words and len(heading_words) >= 2:
                    # ALL significant words must appear in description
                    if all(word in desc_normalized for word in heading_words):
                        return True
                
                return False
            
            for row in range(1, estimate_sheet.max_row + 1):
                # Get the description from column D (or search nearby columns)
                desc_cell_val = None
                for desc_col in [4, 5, 3]:  # D, E, C - prioritize column D
                    cell_val = estimate_sheet.cell(row=row, column=desc_col).value
                    if cell_val and str(cell_val).strip():
                        desc_cell_val = str(cell_val).strip()
                        break
                
                if not desc_cell_val:
                    continue
                
                # Check if this row matches any of our item headings
                for heading in item_headings:
                    if heading in qty_unit_map:
                        continue  # Already found
                    
                    if heading_matches_description(heading, desc_cell_val):
                        # Get quantity from column B and unit from column C
                        qty = ""
                        unit = ""
                        
                        # Try to get quantity from column B first, then other columns
                        if estimate_sheet_values:
                            for qty_col in [2, 6, 3]:  # B, F, C
                                qty_cell = estimate_sheet_values.cell(row=row, column=qty_col)
                                if qty_cell.value is not None:
                                    try:
                                        qty_val = float(qty_cell.value)
                                        if qty_val > 0:
                                            qty = str(qty_val)
                                            break
                                    except (ValueError, TypeError):
                                        pass
                        
                        # Get unit from column C first, then other columns
                        for unit_col in [3, 7, 5]:  # C, G, E
                            unit_cell = estimate_sheet.cell(row=row, column=unit_col)
                            if unit_cell.value is not None:
                                unit_val = str(unit_cell.value).strip()
                                if unit_val and unit_val.lower() in ('nos', 'no', 'mtr', 'mtrs', 'pts', 'pt', 'sqm', 'cum', 'kg', 'l', 'rm', 'each', 'set', 'job', 'ls', 'rmt', 'sqmtr', 'metre', 'meters'):
                                    unit = unit_val
                                    break
                        
                        qty_unit_map[heading] = {'qty': qty, 'unit': unit}
        
        # Build final items list
        all_items = []
        for heading in item_headings:
            item_data = qty_unit_map.get(heading, {'qty': '', 'unit': ''})
            all_items.append({
                'desc': heading,
                'qty': item_data['qty'],
                'unit': item_data['unit']
            })
        
        # Create Word document
        doc = Document()
        
        # Title
        title = doc.add_heading('Specification report accompanying the estimate :-', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in title.runs:
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.underline = True
        
        # Introduction paragraph
        if not work_name:
            work_name = "{{NAME_OF_WORK}}"
        
        intro_para = doc.add_paragraph()
        intro_para.add_run(f'The estimate is prepared for the work {work_name}')
        
        doc.add_paragraph()
        
        # Estimate amount
        amount_para = doc.add_paragraph()
        amount_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if total_amount:
            amount_run = amount_para.add_run(f'Est.Amount: Rs. {total_amount}')
        else:
            amount_run = amount_para.add_run('Est.Amount: {{EST_AMOUNT}}')
        amount_run.font.bold = True
        amount_run.font.underline = True
        
        doc.add_paragraph()
        
        # Body of letter placeholder
        body_para = doc.add_paragraph('{{BODY_OF_LETTER}}')
        for run in body_para.runs:
            run.font.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph()
        
        # "Hence" statement
        doc.add_paragraph('Hence, this estimate has been prepared accordingly.')
        
        doc.add_paragraph()
        
        # Item bullet points
        for item in all_items:
            desc = item.get('desc', '')
            qty = item.get('qty', '')
            unit = item.get('unit', '')
            
            # Format quantity
            if qty:
                try:
                    qty_float = float(qty)
                    if qty_float == int(qty_float):
                        qty = str(int(qty_float))
                    else:
                        qty = str(qty_float)
                except:
                    pass
            
            # Build bullet text
            if qty and unit:
                bullet_text = f'{desc}  -  {qty} {unit}'
            elif qty:
                bullet_text = f'{desc}  -  {qty}'
            else:
                bullet_text = desc
            
            bullet_para = doc.add_paragraph(bullet_text, style='List Bullet')
            for run in bullet_para.runs:
                run.font.size = Pt(11)
                run.font.bold = True
        
        doc.add_paragraph()
        
        # Calculate financial year (April to March)
        from datetime import datetime
        today = datetime.now()
        if today.month >= 4:  # April onwards
            fy_start = today.year
            fy_end = (today.year + 1) % 100
        else:  # Jan to March
            fy_start = today.year - 1
            fy_end = today.year % 100
        financial_year = f"{fy_start}-{fy_end:02d}"
        
        # Footer about rates
        footer_text = (f'The rates proposed in the estimate are as per SQR {financial_year} and Approved rates. L.S. Provision is made in the '
                      'estimate towards GST at 18%, QC amount at 1%, Labour Cess at 1% and NAC amount at 0.1% as per actual '
                      'and LS Provision Towards, unforeseen items & rounding off also proposed in the estimate.')
        footer_para = doc.add_paragraph(footer_text)
        for run in footer_para.runs:
            run.font.size = Pt(10)
        
        doc.add_paragraph()
        
        # Funds section
        funds_para = doc.add_paragraph()
        funds_run = funds_para.add_run('FUNDS: ')
        funds_run.font.bold = True
        funds_run.font.underline = True
        funds_para.add_run('The estimate requires Administrative sanction and also fixes up the agency with provision of funds '
                          'under relevant head of account for taking up the work from the Government, Telangana State Hyderabad')
        
        # Generate filename
        safe_name = work_name.replace(" ", "_").replace("/", "_").replace("{{", "").replace("}}", "")[:25]
        filename = f'Specification_Report_{safe_name}_{timezone.now().strftime("%Y%m%d_%H%M%S")}.docx'
        
        # Return as download
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        doc.save(response)
        return response
        
    except Exception as e:
        logger.error(f'Error generating specification report: {str(e)}', exc_info=True)
        return render(request, 'core/estimate.html', {
            'error': f'Error generating specification report: {str(e)}'
        })


@login_required(login_url='login')
def download_specification_report_live(request, category):
    """
    Generate specification report from live estimate items (New Estimate module).
    Receives items as JSON from the frontend.
    """
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    if request.method != 'POST':
        return redirect('datas_groups', category=category)
    
    try:
        # Get data from POST
        items_json = request.POST.get('items', '[]')
        work_name = request.POST.get('work_name', '{{NAME_OF_WORK}}')
        total_amount = request.POST.get('total_amount', '0.00')
        
        items = json.loads(items_json)
        
        if not items:
            from django.contrib import messages
            messages.error(request, 'No items with quantities to generate specification report')
            return redirect('datas_groups', category=category)
        
        # Create Word document
        doc = Document()
        
        # Title
        title = doc.add_heading('Specification report accompanying the estimate :-', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in title.runs:
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.underline = True
        
        # Introduction paragraph
        intro_para = doc.add_paragraph()
        intro_para.add_run(f'The estimate is prepared for the work {work_name}')
        
        doc.add_paragraph()
        
        # Estimate amount
        amount_para = doc.add_paragraph()
        amount_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        amount_run = amount_para.add_run(f'Est.Amount: Rs. {total_amount}')
        amount_run.font.bold = True
        amount_run.font.underline = True
        
        doc.add_paragraph()
        
        # Body of letter placeholder
        body_para = doc.add_paragraph('{{BODY_OF_LETTER}}')
        for run in body_para.runs:
            run.font.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph()
        
        # "Hence" statement
        doc.add_paragraph('Hence, this estimate has been prepared accordingly.')
        
        doc.add_paragraph()
        
        # Item bullet points
        for item in items:
            desc = item.get('desc', '')
            qty = item.get('qty', '')
            unit = item.get('unit', '')
            
            # Format quantity - remove trailing .0 if whole number
            if qty:
                try:
                    qty_float = float(qty)
                    if qty_float == int(qty_float):
                        qty = str(int(qty_float))
                    else:
                        qty = str(qty_float)
                except:
                    pass
            
            # Build bullet text
            if qty and unit:
                bullet_text = f'{desc}  -  {qty} {unit}'
            elif qty:
                bullet_text = f'{desc}  -  {qty}'
            else:
                bullet_text = desc
            
            bullet_para = doc.add_paragraph(bullet_text, style='List Bullet')
            for run in bullet_para.runs:
                run.font.size = Pt(11)
                run.font.bold = True
        
        doc.add_paragraph()
        
        # Calculate financial year (April to March)
        from datetime import datetime
        today = datetime.now()
        if today.month >= 4:  # April onwards
            fy_start = today.year
            fy_end = (today.year + 1) % 100
        else:  # Jan to March
            fy_start = today.year - 1
            fy_end = today.year % 100
        financial_year = f"{fy_start}-{fy_end:02d}"
        
        # Footer about rates
        footer_text = (f'The rates proposed in the estimate are as per SQR {financial_year} and Approved rates. L.S. Provision is made in the '
                      'estimate towards GST at 18%, QC amount at 1%, Labour Cess at 1% and NAC amount at 0.1% as per actual '
                      'and LS Provision Towards, unforeseen items & rounding off also proposed in the estimate.')
        footer_para = doc.add_paragraph(footer_text)
        for run in footer_para.runs:
            run.font.size = Pt(10)
        
        doc.add_paragraph()
        
        # Funds section
        funds_para = doc.add_paragraph()
        funds_run = funds_para.add_run('FUNDS: ')
        funds_run.font.bold = True
        funds_run.font.underline = True
        funds_para.add_run('The estimate requires Administrative sanction and also fixes up the agency with provision of funds '
                          'under relevant head of account for taking up the work from the Government, Telangana State Hyderabad')
        
        # Generate filename
        safe_name = work_name.replace(" ", "_").replace("/", "_").replace("{{", "").replace("}}", "")[:25]
        filename = f'Specification_Report_{safe_name}_{timezone.now().strftime("%Y%m%d_%H%M%S")}.docx'
        
        # Return as download
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        doc.save(response)
        return response
        
    except Exception as e:
        logger.error(f'Error generating specification report: {str(e)}', exc_info=True)
        from django.contrib import messages
        messages.error(request, f'Error generating specification report: {str(e)}')
        return redirect('datas_groups', category=category)


@login_required(login_url='login')
def generate_estimate_forwarding_letter(request):
    """
    Generate a forwarding letter for multi-sheet estimates.
    
    Extracts name of work and grand total from each Estimate sheet in the workbook
    and generates a formal forwarding letter in Word format with:
    - Serial numbered table of works and amounts
    - Indian number formatting for amounts
    - Dynamic financial year and date
    - Generic officer designations
    """
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    if request.method != 'POST':
        return render(request, 'core/estimate.html', {
            'error': 'Invalid request method'
        })
    
    output_file = request.FILES.get('output_file')
    if not output_file:
        return render(request, 'core/estimate.html', {
            'error': 'Please upload an Excel file.'
        })
    
    try:
        # Reset file position
        try:
            output_file.seek(0)
        except:
            pass
        
        wb_upload = load_workbook(output_file, data_only=False)
        
        # Also load with data_only=True for values
        output_file.seek(0)
        wb_values = load_workbook(output_file, data_only=True)
        
        # Find all Estimate sheets and extract work names and grand totals
        estimates_data = []
        
        for sheet_name in wb_upload.sheetnames:
            sheet_lower = sheet_name.lower().strip()
            
            # Look for sheets that contain "estimate" in their name
            if 'estimate' in sheet_lower:
                ws = wb_upload[sheet_name]
                ws_values = wb_values[sheet_name]
                
                work_name = ""
                grand_total = 0.0
                
                # Find work name - look for "Name of the work" or "Name of work"
                for row in range(1, min(30, ws.max_row + 1)):
                    if work_name:
                        break
                    for col in range(1, 10):
                        cell_val = str(ws.cell(row=row, column=col).value or "").strip()
                        cell_lower = cell_val.lower()
                        
                        if "name of the work" in cell_lower or "name of work" in cell_lower:
                            # Check if work name is in the same cell after colon
                            if ':' in cell_val:
                                parts = cell_val.split(':', 1)
                                if len(parts) > 1 and parts[1].strip():
                                    work_name = parts[1].strip()
                                    break
                            
                            # Try next columns in same row
                            if not work_name:
                                for next_col in range(col + 1, col + 8):
                                    if next_col <= ws.max_column:
                                        next_cell = ws.cell(row=row, column=next_col).value
                                        if next_cell and str(next_cell).strip():
                                            work_name = str(next_cell).strip()
                                            break
                            
                            # Try next row
                            if not work_name:
                                if row + 1 <= ws.max_row:
                                    next_row_val = ws.cell(row=row + 1, column=col).value
                                    if next_row_val and str(next_row_val).strip():
                                        work_name = str(next_row_val).strip()
                            break
                
                # Find Grand Total
                for row in range(1, ws.max_row + 1):
                    for col in range(1, 10):
                        cell_val = str(ws.cell(row=row, column=col).value or "").strip().lower()
                        if 'grand total' in cell_val:
                            # Look for amount in columns to the right (especially column H)
                            for amt_col in range(col + 1, col + 8):
                                if amt_col <= ws.max_column:
                                    amt_cell = ws_values.cell(row=row, column=amt_col)
                                    if amt_cell.value is not None:
                                        try:
                                            amt_val = float(amt_cell.value)
                                            if amt_val > 100:  # Should be a significant amount
                                                grand_total = amt_val
                                                break
                                        except (ValueError, TypeError):
                                            pass
                            
                            if grand_total > 0:
                                break
                    if grand_total > 0:
                        break
                
                # Only add if we found meaningful data
                if work_name or grand_total > 0:
                    estimates_data.append({
                        'sheet_name': sheet_name,
                        'work_name': work_name or f"Work from {sheet_name}",
                        'grand_total': grand_total
                    })
        
        if not estimates_data:
            return render(request, 'core/estimate.html', {
                'error': 'No Estimate sheets found with work names or grand totals. Make sure your workbook has sheets with "Estimate" in the name.'
            })
        
        # Get current date and financial year
        current_date = _get_current_date_formatted()
        financial_year = _get_current_financial_year()
        
        # Create Word document
        doc = Document()
        
        # Light gray color for placeholders
        placeholder_color = RGBColor(169, 169, 169)  # Light gray
        
        # Set page margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Header - Department name (placeholder)
        header1 = doc.add_paragraph()
        header1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = header1.add_run('[GOVERNMENT / ORGANIZATION NAME]')
        run1.font.bold = True
        run1.font.size = Pt(14)
        run1.font.color.rgb = placeholder_color
        run1.font.italic = True
        
        header2 = doc.add_paragraph()
        header2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = header2.add_run('[DEPARTMENT NAME]')
        run2.font.bold = True
        run2.font.size = Pt(13)
        run2.font.color.rgb = placeholder_color
        run2.font.italic = True
        
        doc.add_paragraph()  # Blank line
        
        # From/To section in a table
        from_to_table = doc.add_table(rows=1, cols=2)
        from_to_table.autofit = True
        
        from_cell = from_to_table.cell(0, 0)
        from_para = from_cell.paragraphs[0]
        from_para.add_run('From: -\n')
        # Placeholder for sender details
        from_run1 = from_para.add_run('[Officer Name, Qualification],\n')
        from_run1.font.color.rgb = placeholder_color
        from_run1.font.italic = True
        from_run2 = from_para.add_run('[Designation],\n')
        from_run2.font.color.rgb = placeholder_color
        from_run2.font.italic = True
        from_run3 = from_para.add_run('[Sub Division, Office Address].')
        from_run3.font.color.rgb = placeholder_color
        from_run3.font.italic = True
        
        to_cell = from_to_table.cell(0, 1)
        to_para = to_cell.paragraphs[0]
        to_para.add_run('To,\n')
        # Placeholder for recipient details
        to_run1 = to_para.add_run('[Officer Designation],\n')
        to_run1.font.color.rgb = placeholder_color
        to_run1.font.italic = True
        to_run2 = to_para.add_run('[Division Name],\n')
        to_run2.font.color.rgb = placeholder_color
        to_run2.font.italic = True
        to_run3 = to_para.add_run('[Address].')
        to_run3.font.color.rgb = placeholder_color
        to_run3.font.italic = True
        
        doc.add_paragraph()
        
        # Letter number and date on same line
        today = timezone.now().date()
        month_num = today.month
        year_short = today.year % 100
        
        lr_para = doc.add_paragraph()
        lr_para.add_run('Lr No. ')
        lr_placeholder = lr_para.add_run('[Office Code]')
        lr_placeholder.font.color.rgb = placeholder_color
        lr_placeholder.font.italic = True
        lr_placeholder.font.underline = True
        lr_para.add_run(f'/{financial_year}/          ')
        lr_para.add_run(f'\t\t\t\t\tDate:-    - {today.strftime("%m")} - {today.year}.')
        
        doc.add_paragraph()
        
        # Sir,
        sir_para = doc.add_paragraph()
        sir_para.add_run('Sir,')
        
        doc.add_paragraph()
        
        # Subject
        subject_para = doc.add_paragraph()
        subj_run = subject_para.add_run('Sub:-')
        subj_run.font.underline = True
        subject_para.add_run('\t')
        subj_placeholder = subject_para.add_run(f'[Subject of the letter] ')
        subj_placeholder.font.color.rgb = placeholder_color
        subj_placeholder.font.italic = True
        subject_para.add_run(f'for the year {financial_year}.  -  Submission  -  Request for obtaining administrative sanction  -  Regarding.')
        
        doc.add_paragraph()
        
        # Reference
        ref_para = doc.add_paragraph()
        ref_run = ref_para.add_run('Ref:-')
        ref_run.font.underline = True
        ref_para.add_run('\tMemo No.')
        ref_placeholder = ref_para.add_run('[Reference Number]')
        ref_placeholder.font.color.rgb = placeholder_color
        ref_placeholder.font.italic = True
        ref_placeholder.font.underline = True
        ref_para.add_run(f'/{financial_year} Dt.')
        ref_date_placeholder = ref_para.add_run('[DD.MM.YYYY]')
        ref_date_placeholder.font.color.rgb = placeholder_color
        ref_date_placeholder.font.italic = True
        ref_date_placeholder.font.underline = True
        
        doc.add_paragraph()
        
        # Stars separator
        stars_para = doc.add_paragraph()
        stars_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        stars_para.add_run('**.**')
        
        doc.add_paragraph()
        
        # Main body
        body_para = doc.add_paragraph()
        body_para.add_run('With reference to the subject cited, I submit here ')
        with_run = body_para.add_run(f'with  {len(estimates_data)}')
        with_run.font.underline = True
        body_para.add_run(' Nos. estimates for the following works for the amounts specifies against each work.')
        
        doc.add_paragraph()
        
        # Create table for estimates
        table = doc.add_table(rows=len(estimates_data) + 1, cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set column widths
        for cell in table.columns[0].cells:
            cell.width = Inches(0.5)
        for cell in table.columns[1].cells:
            cell.width = Inches(4.5)
        for cell in table.columns[2].cells:
            cell.width = Inches(1.5)
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Sl.\nNo'
        header_cells[1].text = 'Name of work'
        header_cells[2].text = 'Amount'
        
        # Center align and bold header
        for cell in header_cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.bold = True
        
        # Data rows
        for idx, est_data in enumerate(estimates_data, start=1):
            row_cells = table.rows[idx].cells
            
            # Serial number
            row_cells[0].text = str(idx)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Name of work
            row_cells[1].text = est_data['work_name']
            
            # Amount in Indian format
            formatted_amount = _format_indian_number(est_data['grand_total'])
            row_cells[2].text = f"Rs.{formatted_amount}"
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()
        
        # Specification statement
        spec_para = doc.add_paragraph()
        spec_para.add_run("Specification report accompanying each estimate explains the necessity and provisions made they're in detail.")
        
        doc.add_paragraph()
        
        # Request paragraph
        request_para = doc.add_paragraph()
        request_para.add_run('I request the ')
        req_placeholder = request_para.add_run('[Superior Officer Designation]')
        req_placeholder.font.color.rgb = placeholder_color
        req_placeholder.font.italic = True
        request_para.add_run(' to kindly arrange to obtain administrative sanction the above estimates and arrange to finalize the agencies at the earliest for taking up the works.')
        
        doc.add_paragraph()
        
        # Enclosure
        enc_para = doc.add_paragraph()
        enc_para.add_run('Enclosure: -')
        doc.add_paragraph(f'Estimates  - {len(estimates_data)} No\'s,')
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Signature section
        sign_para = doc.add_paragraph()
        sign_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sign_para.add_run('Yours faithfully,')
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_title = title_para.add_run('[Officer Designation]\n')
        run_title.font.bold = True
        run_title.font.color.rgb = placeholder_color
        run_title.font.italic = True
        sub_div_run = title_para.add_run('[Sub Division Name],\n')
        sub_div_run.font.color.rgb = placeholder_color
        sub_div_run.font.italic = True
        addr_run = title_para.add_run('[Office Address].')
        addr_run.font.color.rgb = placeholder_color
        addr_run.font.italic = True
        
        doc.add_paragraph()
        
        # Copy to
        copy_para = doc.add_paragraph()
        copy_para.add_run('Copy to the ')
        copy_placeholder = copy_para.add_run('[Officer Designation, Section Name]')
        copy_placeholder.font.color.rgb = placeholder_color
        copy_placeholder.font.italic = True
        copy_para.add_run(' for information.')
        
        # Generate filename
        filename = f'Forwarding_Letter_{financial_year}_{timezone.now().strftime("%Y%m%d_%H%M%S")}.docx'
        
        # Return as download
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        doc.save(response)
        return response
        
    except Exception as e:
        logger.error(f'Error generating forwarding letter: {str(e)}', exc_info=True)
        return render(request, 'core/estimate.html', {
            'error': f'Error generating forwarding letter: {str(e)}'
        })


# ==============================================================================
# AMC MODULE VIEWS
# ==============================================================================
# AMC (Annual Maintenance Contract) module works similar to New Estimate module
# but uses a custom backend sheet uploaded via Admin Panel

@login_required(login_url='login')
def amc_home(request):
    """
    Landing page for AMC Module.
    Similar to datas() view for New Estimate.
    - Clears current selection
    - Reads ?work_type=original/repair from URL and stores in session
    """
    # Always start fresh when entering AMC
    request.session["amc_fetched_items"] = []
    request.session["amc_current_project_name"] = None
    request.session["amc_qty_map"] = {}
    request.session["amc_work_name"] = ""
    request.session["amc_grand_total"] = ""
    request.session["amc_selected_backend_id"] = None  # Clear backend selection

    mode = request.GET.get("work_type")
    if mode in ("original", "repair"):
        request.session["amc_work_type"] = mode

    # If nothing in session yet, default to original
    if "amc_work_type" not in request.session:
        request.session["amc_work_type"] = "original"

    return render(
        request,
        "core/amc/amc_home.html",
        {"work_type": request.session["amc_work_type"]},
    )


@login_required(login_url='login')
def amc_groups(request, category):
    """
    AMC Groups page - shows available groups for the selected backend sheet category.
    Category here refers to the backend_sheet_name set in the Module.
    """
    from core.utils_excel import get_available_backends_for_module
    
    # Remember work_type in session if passed in URL
    work_type = (request.GET.get("work_type") or "").lower()
    if work_type in ("original", "repair"):
        request.session["amc_work_type"] = work_type

    # Handle backend switching via URL parameter
    url_backend_id = request.GET.get("backend_id")
    if url_backend_id:
        try:
            request.session["amc_selected_backend_id"] = int(url_backend_id)
        except (ValueError, TypeError):
            pass
    
    amc_selected_backend_id = request.session.get("amc_selected_backend_id")
    
    # Map amc category to base category for backend lookup
    base_category = category.replace('amc_', '')  # amc_electrical -> electrical

    try:
        items_list, groups_map, _, ws_data, filepath = load_backend(
            category, settings.BASE_DIR,
            backend_id=amc_selected_backend_id,
            module_code='amc',  # Use amc module's own backends
            user=request.user
        )
    except FileNotFoundError as e:
        logger.info(f"No backend data available for AMC category {category} - showing Coming Soon")
        other_base = 'electrical' if base_category == 'civil' else 'civil'
        other_category = f'amc_{other_base}'
        other_available = False
        try:
            other_backends = get_available_backends_for_module('amc', other_base)
            other_available = len(other_backends) > 0
        except:
            pass
        return render(request, "core/coming_soon.html", {
            "category": category,
            "module_name": "AMC",
            "other_category": other_category,
            "other_category_available": other_available,
        })
    except ValueError as e:
        return render(request, "core/amc/amc_groups.html", {
            "category": category,
            "groups": [],
            "error": str(e),
        })
    
    groups = sorted(groups_map.keys(), key=lambda s: s.lower())

    if not groups:
        return render(request, "core/amc/amc_groups.html", {
            "category": category,
            "groups": [],
            "error": "No groups found in backend Excel.",
        })

    default_group = request.GET.get("group") or groups[0]
    return redirect("amc_items", category=category, group=default_group)


@login_required(login_url='login')
def amc_items(request, category, group):
    """
    AMC Items page - 3-panel UI for selecting items.
    Similar to datas_items() view.
    """
    from core.utils_excel import get_available_backends_for_module
    
    # Handle backend switching via URL parameter
    url_backend_id = request.GET.get("backend_id")
    if url_backend_id:
        try:
            request.session["amc_selected_backend_id"] = int(url_backend_id)
        except (ValueError, TypeError):
            pass
    
    amc_selected_backend_id = request.session.get("amc_selected_backend_id")
    
    # Map amc category to base category for backend lookup
    base_category = category.replace('amc_', '')  # amc_electrical -> electrical
    
    # Get available backends for dropdown (amc has its own backends)
    available_backends = get_available_backends_for_module('amc', base_category)
    
    try:
        items_list, groups_map, _, ws_data, filepath = load_backend(
            category, settings.BASE_DIR,
            backend_id=amc_selected_backend_id,
            module_code='amc',  # Use amc module's own backends
            user=request.user
        )
    except FileNotFoundError as e:
        logger.info(f"No backend data available for AMC category {category} - showing Coming Soon")
        other_base = 'electrical' if base_category == 'civil' else 'civil'
        other_category = f'amc_{other_base}'
        other_available = False
        try:
            other_backends = get_available_backends_for_module('amc', other_base)
            other_available = len(other_backends) > 0
        except:
            pass
        return render(request, "core/coming_soon.html", {
            "category": category,
            "module_name": "AMC",
            "other_category": other_category,
            "other_category_available": other_available,
        })
    except ValueError as e:
        return render(request, "core/amc/amc_items.html", {
            "category": category,
            "group": group,
            "groups": [],
            "error": str(e),
        })

    groups = sorted(groups_map.keys(), key=lambda s: s.lower())

    group_items = groups_map.get(group, [])
    detected_names = {i["name"] for i in items_list}
    display_items = [name for name in group_items if name in detected_names]

    wb_vals = load_workbook(filepath, data_only=True)
    ws_vals = wb_vals["Master Datas"]

    item_rates = {}
    for info in items_list:
        name = info["name"]
        start_row = info["start_row"]
        end_row = info["end_row"]
        rate = None
        for r in range(end_row, start_row - 1, -1):
            val = ws_vals.cell(row=r, column=10).value  # column J
            if val not in (None, ""):
                rate = val
                break
        item_rates[name] = rate

    item_to_group = {}
    for grp_name, item_list_in_grp in groups_map.items():
        for nm in item_list_in_grp:
            item_to_group.setdefault(nm, grp_name)

    def units_for(name):
        grp_name = item_to_group.get(name, "")
        if grp_name in ("Piping", "Wiring & Cables"):
            return ("Mtrs", "Mtr")
        elif grp_name == "Points":
            return ("Pts", "Pt")
        else:
            return ("Nos", "No")

    items_info = []
    for name in display_items:
        items_info.append({
            "name": name,
            "rate": item_rates.get(name),
        })

    fetched = request.session.get("amc_fetched_items", [])

    qty_map = request.session.get("amc_qty_map", {}) or {}
    work_name = request.session.get("amc_work_name", "") or ""
    grand_total = request.session.get("amc_grand_total", "") or ""

    estimate_rows = []
    for idx, name in enumerate(fetched, start=1):
        plural, singular = units_for(name)
        estimate_rows.append({
            "sl": idx,
            "name": name,
            "rate": item_rates.get(name),
            "unit": plural,
            "qty": qty_map.get(name, ""),
        })

    return render(request, "core/amc/amc_items.html", {
        "category": category,
        "group": group,
        "groups": groups,
        "items_info": items_info,
        "fetched": fetched,
        "estimate_rows": estimate_rows,
        "work_name": work_name,
        "grand_total": grand_total,
        "available_backends": available_backends,
        "selected_backend_id": amc_selected_backend_id,
    })


@login_required(login_url='login')
def amc_fetch_item(request, category, group, item):
    """
    Toggle fetched items for AMC module.
    """
    fetched = request.session.get("amc_fetched_items", []) or []

    if item in fetched:
        fetched.remove(item)
    else:
        fetched.append(item)

    request.session["amc_fetched_items"] = fetched

    work_name = request.GET.get("work_name")
    if work_name is not None:
        request.session["amc_work_name"] = work_name

    return redirect("amc_items", category=category, group=group)


# -----------------------
# AMC AJAX TOGGLE ITEM (no page reload)
# -----------------------
@login_required(login_url='login')
def amc_ajax_toggle_item(request, category):
    """
    AJAX endpoint to toggle an item in AMC fetched list without page reload.
    POST with JSON: { "item": "item_name", "action": "add" or "remove", "work_name": "..." }
    Returns JSON: { "status": "ok", "fetched": [...], "action_taken": "added" or "removed", "item_info": {...} }
    """
    if request.method != "POST":
        return JsonResponse({"status": "error", "message": "POST required"}, status=405)
    
    try:
        if request.content_type and 'application/json' in request.content_type:
            data = json.loads(request.body.decode("utf-8") or "{}")
        else:
            data = request.POST
        
        item = data.get("item", "").strip()
        action = data.get("action", "toggle")  # "add", "remove", or "toggle"
        work_name = data.get("work_name")
        
        if not item:
            return JsonResponse({"status": "error", "message": "No item specified"}, status=400)
        
        fetched = request.session.get("amc_fetched_items", []) or []
        action_taken = None
        
        if action == "add":
            if item not in fetched:
                fetched.append(item)
                action_taken = "added"
            else:
                action_taken = "already_exists"
        elif action == "remove":
            if item in fetched:
                fetched.remove(item)
                action_taken = "removed"
            else:
                action_taken = "not_found"
        else:  # toggle
            if item in fetched:
                fetched.remove(item)
                action_taken = "removed"
            else:
                fetched.append(item)
                action_taken = "added"
        
        request.session["amc_fetched_items"] = fetched
        
        if work_name is not None:
            request.session["amc_work_name"] = work_name
        
        # Get item info (rate, unit) for newly added items
        item_info = None
        if action_taken == "added":
            try:
                # Use same backend as the AMC items page - get from session
                amc_selected_backend_id = request.session.get("amc_selected_backend_id")
                items_list, groups_map, units_map, ws_data, filepath = load_backend(
                    category, settings.BASE_DIR,
                    backend_id=amc_selected_backend_id,
                    module_code='amc',
                    user=request.user
                )
                
                # Get rate
                wb_vals = load_workbook(filepath, data_only=True)
                ws_vals = wb_vals["Master Datas"]
                
                item_rate = None
                for info in items_list:
                    if info["name"] == item:
                        start_row = info["start_row"]
                        end_row = info["end_row"]
                        for r in range(end_row, start_row - 1, -1):
                            val = ws_vals.cell(row=r, column=10).value  # column J
                            if val not in (None, ""):
                                item_rate = val
                                break
                        break
                
                # Get unit with smart fallback (same logic as datas_items)
                # Priority: 1) units_map from backend, 2) group-based defaults
                unit = units_map.get(item, "")
                if not unit:
                    # Find item's group for fallback
                    item_group = ""
                    for grp_name, grp_items in groups_map.items():
                        if item in grp_items:
                            item_group = grp_name
                            break
                    # Group-based defaults
                    if item_group in ("Piping", "Wiring & Cables", "Run of Mains", "Sheathed Cables", "U.G Cabling"):
                        unit = "Mtrs"
                    elif item_group == "Points":
                        unit = "Pts"
                    else:
                        unit = "Nos"
                
                item_info = {
                    "name": item,
                    "rate": item_rate,
                    "unit": unit
                }
                
                wb_vals.close()
            except Exception as e:
                # If we can't get item info, just return without it
                item_info = {"name": item, "rate": None, "unit": "Nos"}
        
        return JsonResponse({
            "status": "ok",
            "fetched": fetched,
            "action_taken": action_taken,
            "item": item,
            "item_info": item_info
        })
        
    except Exception as e:
        return JsonResponse({"status": "error", "message": str(e)}, status=400)


# -----------------------
# AMC AJAX REORDER ITEMS
# -----------------------
@login_required(login_url='login')
def amc_ajax_reorder_items(request, category):
    """
    AJAX endpoint to reorder AMC fetched items list.
    POST with JSON: { "items": ["item1", "item2", ...] }
    Returns JSON: { "status": "ok", "fetched": [...] }
    """
    if request.method != "POST":
        return JsonResponse({"status": "error", "message": "POST required"}, status=405)
    
    try:
        if request.content_type and 'application/json' in request.content_type:
            data = json.loads(request.body.decode("utf-8") or "{}")
        else:
            return JsonResponse({"status": "error", "message": "JSON required"}, status=400)
        
        new_order = data.get("items", [])
        
        if not isinstance(new_order, list):
            return JsonResponse({"status": "error", "message": "items must be a list"}, status=400)
        
        # Validate: new_order should contain the same items as current fetched
        current_fetched = set(request.session.get("amc_fetched_items", []) or [])
        new_order_set = set(new_order)
        
        # Only reorder if sets match (no items added/removed via this endpoint)
        if current_fetched == new_order_set:
            request.session["amc_fetched_items"] = new_order
        else:
            # Allow partial reorder - use the intersection
            valid_items = [item for item in new_order if item in current_fetched]
            request.session["amc_fetched_items"] = valid_items
        
        return JsonResponse({
            "status": "ok",
            "fetched": request.session["amc_fetched_items"]
        })
        
    except Exception as e:
        return JsonResponse({"status": "error", "message": str(e)}, status=400)


@login_required(login_url='login')
def amc_clear_output(request, category):
    """
    Clear all AMC session data.
    """
    request.session["amc_fetched_items"] = []
    request.session["amc_qty_map"] = {}
    request.session["amc_work_name"] = ""
    request.session["amc_grand_total"] = ""

    group = request.GET.get("group")
    if group:
        return redirect("amc_items", category=category, group=group)

    return redirect("amc_groups", category=category)


@login_required(login_url='login')
def amc_save_qty_map(request, category):
    """
    AJAX endpoint to save AMC quantity map and grand total to session.
    """
    if request.method == "POST":
        try:
            qty_map_str = request.POST.get("qty_map", "")
            grand_total_str = request.POST.get("grand_total", "")
            work_name = request.POST.get("work_name", "")
            
            if qty_map_str:
                try:
                    qty_map = json.loads(qty_map_str)
                    if isinstance(qty_map, dict):
                        request.session["amc_qty_map"] = qty_map
                except json.JSONDecodeError:
                    pass
            
            if grand_total_str:
                request.session["amc_grand_total"] = grand_total_str
            
            if work_name:
                request.session["amc_work_name"] = work_name
            
            return JsonResponse({"status": "ok"})
        except Exception as e:
            return JsonResponse({"status": "error", "message": str(e)}, status=400)
    
    return JsonResponse({"status": "error", "message": "POST required"}, status=405)


@login_required(login_url='login')
def amc_download_output(request, category):
    """
    Generate AMC output Excel file - similar to download_output for New Estimate.
    """
    fetched = request.session.get("amc_fetched_items", [])
    if not fetched:
        return JsonResponse({"error": "No items selected"}, status=400)

    # Parse input
    item_qtys = {}
    work_name = ""
    grand_total = None

    if request.method == "POST":
        qty_map_str = request.POST.get("qty_map", "")
        if qty_map_str:
            try:
                raw = json.loads(qty_map_str)
                if isinstance(raw, dict):
                    for k, v in raw.items():
                        try:
                            item_qtys[str(k)] = float(v)
                        except Exception:
                            continue
            except Exception:
                pass

        work_name = (request.POST.get("work_name") or "").strip()
        
        grand_total_str = request.POST.get("grand_total", "").strip()
        if grand_total_str:
            try:
                grand_total = float(grand_total_str)
            except ValueError:
                grand_total = None

    work_type = (request.POST.get("work_type")
                 or request.session.get("amc_work_type")
                 or "original").lower()
    request.session["amc_work_type"] = work_type
    
    # Get selected backend ID from session (for multi-backend support)
    amc_selected_backend_id = request.session.get("amc_selected_backend_id")

    from django.conf import settings as django_settings
    
    if getattr(django_settings, 'CELERY_TASK_ALWAYS_EAGER', True):
        # Synchronous mode
        try:
            org = get_org_from_request(request)
            
            job = Job.objects.create(
                organization=org,
                user=request.user,
                job_type='generate_output_excel',
                status='queued',
                current_step="Processing AMC...",
            )
            
            job.result = {
                'fetched_items': fetched,
                'qty_map': item_qtys,
                'work_name': work_name,
                'work_type': work_type,
                'grand_total': grand_total,
                'module': 'amc',
                'backend_id': amc_selected_backend_id,
            }
            job.save()
            
            # Call task function directly (synchronous)
            from core.tasks import generate_output_excel
            result = generate_output_excel.apply(args=(
                job.id,
                category,
                json.dumps(item_qtys),
                json.dumps({}),  # unit_map
                work_name,
                work_type,
                grand_total,
                None,  # excess_tp_percent
                None,  # ls_special_name
                None,  # ls_special_amount
                None,  # deduct_old_material
                amc_selected_backend_id,
            )).get()
            
            job.refresh_from_db()
            
            if job.status == 'completed' and job.result.get('output_file_id'):
                from core.models import OutputFile
                try:
                    output_file = OutputFile.objects.get(id=job.result['output_file_id'])
                    from django.http import FileResponse
                    import os
                    
                    if output_file.file and os.path.exists(output_file.file.path):
                        response = FileResponse(
                            open(output_file.file.path, 'rb'),
                            as_attachment=True,
                            filename=output_file.filename or f"amc_{category}_output.xlsx",
                        )
                        return response
                except OutputFile.DoesNotExist:
                    pass
            
            return JsonResponse({
                'job_id': job.id,
                'status': job.status,
                'message': job.current_step or 'Processing complete',
                'error': job.error_message if job.status == 'failed' else None,
            })
                
        except Exception as e:
            logger.error(f"Failed to generate AMC output Excel: {e}")
            return JsonResponse({"error": str(e)}, status=500)
    
    # Async mode with Celery
    try:
        org = get_org_from_request(request)
        
        job = Job.objects.create(
            organization=org,
            user=request.user,
            job_type='generate_output_excel',
            status='queued',
            current_step="Queued for AMC processing",
        )
        
        job.result = {
            'fetched_items': fetched,
            'qty_map': item_qtys,
            'work_name': work_name,
            'work_type': work_type,
            'grand_total': grand_total,
            'module': 'amc',
            'backend_id': amc_selected_backend_id,
        }
        job.save()
        
        from core.tasks import generate_output_excel
        task = generate_output_excel.delay(
            job.id,
            category,
            json.dumps(item_qtys),
            json.dumps({}),  # unit_map
            work_name,
            work_type,
            grand_total,
            None,  # excess_tp_percent
            None,  # ls_special_name
            None,  # ls_special_amount
            None,  # deduct_old_material
            amc_selected_backend_id,
        )
        
        job.celery_task_id = task.id
        job.save()
        
        return JsonResponse({
            'job_id': job.id,
            'status_url': reverse('job_status', args=[job.id]),
            'message': f'Generating AMC {category} output. Please wait...'
        })
        
    except Exception as e:
        logger.error(f"Failed to enqueue AMC output Excel task: {e}")
        return JsonResponse({"error": str(e)}, status=500)
