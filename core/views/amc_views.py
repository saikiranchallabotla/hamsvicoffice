# Auto-generated from core/views.py split
import json
import os
import re
import logging
from copy import copy

import inflect
from docx import Document
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, Border, Side, PatternFill
from django.utils import timezone
from django.urls import reverse
from django.core.files.uploadedfile import InMemoryUploadedFile
from django.contrib.auth.decorators import login_required

from django.conf import settings
from django.http import HttpResponse, JsonResponse, HttpResponseNotAllowed
from django.shortcuts import render, redirect, get_object_or_404
from django.template.loader import render_to_string
from django.views.decorators.http import require_POST
from django.utils.crypto import get_random_string

import io
from io import BytesIO
from difflib import SequenceMatcher

from ..models import Project, SelfFormattedTemplate, Estimate, Organization, Membership, Upload, Job, OutputFile, LetterSettings
from ..decorators import org_required, role_required

logger = logging.getLogger(__name__)
from ..tasks import process_excel_upload, generate_bill_pdf, generate_workslip_pdf, generate_bill_document_task
from ..utils_excel import load_backend, copy_block_with_styles_and_formulas, build_temp_day_rates

p_engine = inflect.engine()
BILL_TEMPLATES_DIR = os.path.join(settings.BASE_DIR, "core", "templates", "core", "bill_templates")
_inflect_engine = inflect.engine()

from .utils import (get_org_from_request, _apply_print_settings,
    _format_indian_number, _number_to_words_rupees,
    _get_current_financial_year, _get_letter_settings)

@login_required(login_url='login')
def amc_home(request):
    """
    Landing page for AMC Module.
    Similar to datas() view for New Estimate.
    - Clears current selection
    - Reads ?work_type=original/repair from URL and stores in session
    """
    # Always start fresh when entering AMC
    request.session["amc_fetched_items"] = []
    request.session["amc_current_project_name"] = None
    request.session["amc_qty_map"] = {}
    request.session["amc_work_name"] = ""
    request.session["amc_grand_total"] = ""
    request.session["amc_selected_backend_id"] = None  # Clear backend selection
    request.session["current_saved_work_id"] = None  # Clear so save modal shows "Save" not "Update"

    mode = request.GET.get("work_type")
    if mode in ("original", "repair"):
        request.session["amc_work_type"] = mode

    # If nothing in session yet, default to original
    if "amc_work_type" not in request.session:
        request.session["amc_work_type"] = "original"

    return render(
        request,
        "core/amc/amc_home.html",
        {"work_type": request.session["amc_work_type"]},
    )


@login_required(login_url='login')
def amc_groups(request, category):
    """
    AMC Groups page - shows available groups for the selected backend sheet category.
    Category here refers to the backend_sheet_name set in the Module.
    """
    from core.utils_excel import get_available_backends_for_module
    
    # Remember work_type in session if passed in URL
    work_type = (request.GET.get("work_type") or "").lower()
    if work_type in ("original", "repair"):
        request.session["amc_work_type"] = work_type

    # Handle backend switching via URL parameter
    url_backend_id = request.GET.get("backend_id")
    if url_backend_id:
        try:
            request.session["amc_selected_backend_id"] = int(url_backend_id)
        except (ValueError, TypeError):
            pass
    
    amc_selected_backend_id = request.session.get("amc_selected_backend_id")
    
    # Map amc category to base category for backend lookup
    base_category = category.replace('amc_', '')  # amc_electrical -> electrical

    try:
        items_list, groups_map, _, ws_data, filepath = load_backend(
            category, settings.BASE_DIR,
            backend_id=amc_selected_backend_id,
            module_code='amc',  # Use amc module's own backends
            user=request.user
        )
    except FileNotFoundError as e:
        logger.info(f"No backend data available for AMC category {category} - showing Coming Soon")
        other_base = 'electrical' if base_category == 'civil' else 'civil'
        other_category = f'amc_{other_base}'
        other_available = False
        try:
            other_backends = get_available_backends_for_module('amc', other_base)
            other_available = len(other_backends) > 0
        except:
            pass
        return render(request, "core/coming_soon.html", {
            "category": category,
            "module_name": "AMC",
            "other_category": other_category,
            "other_category_available": other_available,
        })
    except ValueError as e:
        return render(request, "core/amc/amc_groups.html", {
            "category": category,
            "groups": [],
            "error": str(e),
        })
    
    groups = sorted(groups_map.keys(), key=lambda s: s.lower())

    if not groups:
        return render(request, "core/amc/amc_groups.html", {
            "category": category,
            "groups": [],
            "error": "No groups found in backend Excel.",
        })

    default_group = request.GET.get("group") or groups[0]
    return redirect("amc_items", category=category, group=default_group)


@login_required(login_url='login')
def amc_items(request, category, group):
    """
    AMC Items page - 3-panel UI for selecting items.
    Similar to datas_items() view.
    """
    from core.utils_excel import get_available_backends_for_module
    
    # Handle backend switching via URL parameter
    url_backend_id = request.GET.get("backend_id")
    if url_backend_id:
        try:
            request.session["amc_selected_backend_id"] = int(url_backend_id)
        except (ValueError, TypeError):
            pass
    
    amc_selected_backend_id = request.session.get("amc_selected_backend_id")
    
    # Map amc category to base category for backend lookup
    base_category = category.replace('amc_', '')  # amc_electrical -> electrical
    
    # Get available backends for dropdown (amc has its own backends)
    available_backends = get_available_backends_for_module('amc', base_category)
    
    try:
        items_list, groups_map, units_map, ws_data, filepath = load_backend(
            category, settings.BASE_DIR,
            backend_id=amc_selected_backend_id,
            module_code='amc',  # Use amc module's own backends
            user=request.user
        )
    except FileNotFoundError as e:
        logger.info(f"No backend data available for AMC category {category} - showing Coming Soon")
        other_base = 'electrical' if base_category == 'civil' else 'civil'
        other_category = f'amc_{other_base}'
        other_available = False
        try:
            other_backends = get_available_backends_for_module('amc', other_base)
            other_available = len(other_backends) > 0
        except:
            pass
        return render(request, "core/coming_soon.html", {
            "category": category,
            "module_name": "AMC",
            "other_category": other_category,
            "other_category_available": other_available,
        })
    except ValueError as e:
        return render(request, "core/amc/amc_items.html", {
            "category": category,
            "group": group,
            "groups": [],
            "error": str(e),
        })

    groups = sorted(groups_map.keys(), key=lambda s: s.lower())

    group_items = groups_map.get(group, [])
    detected_names = {i["name"] for i in items_list}
    display_items = [name for name in group_items if name in detected_names]

    wb_vals = load_workbook(filepath, data_only=True)
    ws_vals = wb_vals["Master Datas"]

    item_rates = {}
    for info in items_list:
        name = info["name"]
        start_row = info["start_row"]
        end_row = info["end_row"]
        rate = None
        for r in range(end_row, start_row - 1, -1):
            val = ws_vals.cell(row=r, column=10).value  # column J
            if val not in (None, ""):
                rate = val
                break
        item_rates[name] = rate

    item_to_group = {}
    for grp_name, item_list_in_grp in groups_map.items():
        for nm in item_list_in_grp:
            item_to_group.setdefault(nm, grp_name)

    def units_for(name):
        # Priority 1: Use unit from Groups sheet (Column D) via units_map
        backend_unit = units_map.get(name, "") if units_map else ""
        if backend_unit:
            bu = backend_unit.lower()
            if bu in ("mtrs", "mtr", "metre", "meters"):
                return ("Mtrs", "Mtr")
            elif bu in ("pts", "pt", "point", "points"):
                return ("Pts", "Pt")
            elif bu in ("nos", "no"):
                return ("Nos", "No")
            else:
                return (backend_unit, backend_unit)
        # Fallback: group-based
        grp_name = item_to_group.get(name, "")
        if grp_name in ("Piping", "Wiring & Cables"):
            return ("Mtrs", "Mtr")
        elif grp_name == "Points":
            return ("Pts", "Pt")
        else:
            return ("Nos", "No")

    # Build item subtypes map: items with ":" are subtypes
    # Group subtypes by their parent name (part before ":")
    # Support both " : " (with spaces) and ":" (without spaces)
    import re as _re
    _colon_re = _re.compile(r'\s*:\s*')

    def _has_colon(name):
        """Check if item name contains a colon separator (with or without spaces)."""
        return bool(_colon_re.search(name))

    def _split_parent(name):
        """Extract parent name from a colon-separated item name."""
        return _colon_re.split(name, 1)[0].strip()

    item_subtypes = {}  # parent_name -> [list of full subtype names]
    parent_items = set()  # items that have subtypes

    for name in display_items:
        if _has_colon(name):
            parent_name = _split_parent(name)
            if parent_name not in item_subtypes:
                item_subtypes[parent_name] = []
            item_subtypes[parent_name].append(name)
            parent_items.add(parent_name)

    items_info = []
    seen_parents = set()
    for name in display_items:
        if _has_colon(name):
            parent_name = _split_parent(name)
            if parent_name not in seen_parents:
                subtypes_list = item_subtypes.get(parent_name, [])
                items_info.append({
                    "name": parent_name,
                    "rate": None,
                    "has_subtypes": True,
                    "subtypes": json.dumps(subtypes_list),
                    "subtypes_count": len(subtypes_list),
                })
                seen_parents.add(parent_name)
        else:
            items_info.append({
                "name": name,
                "rate": item_rates.get(name),
                "has_subtypes": False,
                "subtypes": "[]",
                "subtypes_count": 0,
            })

    fetched = request.session.get("amc_fetched_items", [])

    qty_map = request.session.get("amc_qty_map", {}) or {}
    work_name = request.session.get("amc_work_name", "") or ""
    grand_total = request.session.get("amc_grand_total", "") or ""

    estimate_rows = []
    for idx, name in enumerate(fetched, start=1):
        plural, singular = units_for(name)
        estimate_rows.append({
            "sl": idx,
            "name": name,
            "rate": item_rates.get(name),
            "unit": plural,
            "qty": qty_map.get(name, ""),
        })

    return render(request, "core/amc/amc_items.html", {
        "category": category,
        "group": group,
        "groups": groups,
        "items_info": items_info,
        "fetched": fetched,
        "estimate_rows": estimate_rows,
        "work_name": work_name,
        "grand_total": grand_total,
        "available_backends": available_backends,
        "selected_backend_id": amc_selected_backend_id,
    })


@login_required(login_url='login')
def amc_fetch_item(request, category, group, item):
    """
    Toggle fetched items for AMC module.
    """
    fetched = request.session.get("amc_fetched_items", []) or []

    if item in fetched:
        fetched.remove(item)
    else:
        fetched.append(item)

    request.session["amc_fetched_items"] = fetched

    work_name = request.GET.get("work_name")
    if work_name is not None:
        request.session["amc_work_name"] = work_name

    return redirect("amc_items", category=category, group=group)


# -----------------------
# AMC AJAX TOGGLE ITEM (no page reload)
# -----------------------
@login_required(login_url='login')
def amc_ajax_toggle_item(request, category):
    """
    AJAX endpoint to toggle an item in AMC fetched list without page reload.
    POST with JSON: { "item": "item_name", "action": "add" or "remove", "work_name": "..." }
    Returns JSON: { "status": "ok", "fetched": [...], "action_taken": "added" or "removed", "item_info": {...} }
    """
    if request.method != "POST":
        return JsonResponse({"status": "error", "message": "POST required"}, status=405)
    
    try:
        if request.content_type and 'application/json' in request.content_type:
            data = json.loads(request.body.decode("utf-8") or "{}")
        else:
            data = request.POST
        
        item = data.get("item", "").strip()
        action = data.get("action", "toggle")  # "add", "remove", or "toggle"
        work_name = data.get("work_name")
        
        if not item:
            return JsonResponse({"status": "error", "message": "No item specified"}, status=400)
        
        fetched = request.session.get("amc_fetched_items", []) or []
        action_taken = None
        
        if action == "add":
            if item not in fetched:
                fetched.append(item)
                action_taken = "added"
            else:
                action_taken = "already_exists"
        elif action == "remove":
            if item in fetched:
                fetched.remove(item)
                action_taken = "removed"
            else:
                action_taken = "not_found"
        else:  # toggle
            if item in fetched:
                fetched.remove(item)
                action_taken = "removed"
            else:
                fetched.append(item)
                action_taken = "added"
        
        request.session["amc_fetched_items"] = fetched
        
        if work_name is not None:
            request.session["amc_work_name"] = work_name
        
        # Get item info (rate, unit) for newly added items
        item_info = None
        if action_taken == "added":
            try:
                # Use same backend as the AMC items page - get from session
                amc_selected_backend_id = request.session.get("amc_selected_backend_id")
                items_list, groups_map, units_map, ws_data, filepath = load_backend(
                    category, settings.BASE_DIR,
                    backend_id=amc_selected_backend_id,
                    module_code='amc',
                    user=request.user
                )
                
                # Get rate
                wb_vals = load_workbook(filepath, data_only=True)
                ws_vals = wb_vals["Master Datas"]
                
                item_rate = None
                for info in items_list:
                    if info["name"] == item:
                        start_row = info["start_row"]
                        end_row = info["end_row"]
                        for r in range(end_row, start_row - 1, -1):
                            val = ws_vals.cell(row=r, column=10).value  # column J
                            if val not in (None, ""):
                                item_rate = val
                                break
                        break
                
                # Get unit with smart fallback (same logic as datas_items)
                # Priority: 1) units_map from backend, 2) group-based defaults
                unit = units_map.get(item, "")
                if not unit:
                    # Find item's group for fallback
                    item_group = ""
                    for grp_name, grp_items in groups_map.items():
                        if item in grp_items:
                            item_group = grp_name
                            break
                    # Group-based defaults
                    if item_group in ("Piping", "Wiring & Cables", "Run of Mains", "Sheathed Cables", "U.G Cabling"):
                        unit = "Mtrs"
                    elif item_group == "Points":
                        unit = "Pts"
                    else:
                        unit = "Nos"
                
                item_info = {
                    "name": item,
                    "rate": item_rate,
                    "unit": unit
                }
                
                wb_vals.close()
            except Exception as e:
                # If we can't get item info, just return without it
                item_info = {"name": item, "rate": None, "unit": "Nos"}
        
        return JsonResponse({
            "status": "ok",
            "fetched": fetched,
            "action_taken": action_taken,
            "item": item,
            "item_info": item_info
        })
        
    except Exception as e:
        return JsonResponse({"status": "error", "message": str(e)}, status=400)


# -----------------------
# AMC AJAX REORDER ITEMS
# -----------------------
@login_required(login_url='login')
def amc_ajax_reorder_items(request, category):
    """
    AJAX endpoint to reorder AMC fetched items list.
    POST with JSON: { "items": ["item1", "item2", ...] }
    Returns JSON: { "status": "ok", "fetched": [...] }
    """
    if request.method != "POST":
        return JsonResponse({"status": "error", "message": "POST required"}, status=405)
    
    try:
        if request.content_type and 'application/json' in request.content_type:
            data = json.loads(request.body.decode("utf-8") or "{}")
        else:
            return JsonResponse({"status": "error", "message": "JSON required"}, status=400)
        
        new_order = data.get("items", [])
        
        if not isinstance(new_order, list):
            return JsonResponse({"status": "error", "message": "items must be a list"}, status=400)
        
        # Validate: new_order should contain the same items as current fetched
        current_fetched = set(request.session.get("amc_fetched_items", []) or [])
        new_order_set = set(new_order)
        
        # Only reorder if sets match (no items added/removed via this endpoint)
        if current_fetched == new_order_set:
            request.session["amc_fetched_items"] = new_order
        else:
            # Allow partial reorder - use the intersection
            valid_items = [item for item in new_order if item in current_fetched]
            request.session["amc_fetched_items"] = valid_items
        
        return JsonResponse({
            "status": "ok",
            "fetched": request.session["amc_fetched_items"]
        })
        
    except Exception as e:
        return JsonResponse({"status": "error", "message": str(e)}, status=400)


@login_required(login_url='login')
def amc_download_specification_report(request, category):
    """
    Generate specification report from live AMC items.
    """
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    if request.method != 'POST':
        return redirect('amc_groups', category=category)
    
    try:
        items_json = request.POST.get('items', '[]')
        work_name = request.POST.get('work_name', '{{NAME_OF_WORK}}')
        total_amount = request.POST.get('total_amount', '0.00')
        
        items = json.loads(items_json)
        
        if not items:
            from django.contrib import messages
            messages.error(request, 'No items with quantities to generate specification report')
            return redirect('amc_groups', category=category)
        
        doc = Document()
        
        title = doc.add_heading('Specification report accompanying the estimate :-', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in title.runs:
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.underline = True
        
        intro_para = doc.add_paragraph()
        intro_para.add_run(f'The estimate is prepared for the work {work_name}')
        
        doc.add_paragraph()
        
        amount_para = doc.add_paragraph()
        amount_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        amount_run = amount_para.add_run(f'Est.Amount: Rs. {total_amount}')
        amount_run.font.bold = True
        amount_run.font.underline = True
        
        doc.add_paragraph()
        
        body_para = doc.add_paragraph('{{BODY_OF_LETTER}}')
        for run in body_para.runs:
            run.font.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph()
        doc.add_paragraph('Hence, this estimate has been prepared accordingly.')
        doc.add_paragraph()
        
        for item in items:
            desc = item.get('desc', '')
            qty = item.get('qty', '')
            unit = item.get('unit', 'Nos')
            
            if qty:
                try:
                    qty_float = float(qty)
                    if qty_float == int(qty_float):
                        qty = str(int(qty_float))
                    else:
                        qty = str(qty_float)
                except:
                    pass
            
            if qty and unit:
                bullet_text = f'{desc}  -  {qty} {unit}'
            elif qty:
                bullet_text = f'{desc}  -  {qty}'
            else:
                bullet_text = desc
            
            bullet_para = doc.add_paragraph(bullet_text, style='List Bullet')
            for run in bullet_para.runs:
                run.font.size = Pt(11)
                run.font.bold = True
        
        doc.add_paragraph()
        
        from datetime import datetime
        today = datetime.now()
        if today.month >= 4:
            fy_start = today.year
            fy_end = (today.year + 1) % 100
        else:
            fy_start = today.year - 1
            fy_end = today.year % 100
        financial_year = f"{fy_start}-{fy_end:02d}"
        
        footer_text = (f'The rates proposed in the estimate are as per SQR {financial_year} and Approved rates. L.S. Provision is made in the '
                      'estimate towards GST at 18%, QC amount at 1%, Labour Cess at 1% and NAC amount at 0.1% as per actual '
                      'and LS Provision Towards, unforeseen items & rounding off also proposed in the estimate.')
        footer_para = doc.add_paragraph(footer_text)
        for run in footer_para.runs:
            run.font.size = Pt(10)
        
        doc.add_paragraph()
        
        funds_para = doc.add_paragraph()
        funds_run = funds_para.add_run('FUNDS: ')
        funds_run.font.bold = True
        funds_run.font.underline = True
        funds_para.add_run('The estimate requires Administrative sanction and also fixes up the agency with provision of funds '
                          'under relevant head of account for taking up the work from the Government, Telangana State Hyderabad')
        
        safe_name = work_name.replace(" ", "_").replace("/", "_").replace("{{", "").replace("}}", "")[:25]
        filename = f'AMC_Specification_Report_{safe_name}_{timezone.now().strftime("%Y%m%d_%H%M%S")}.docx'
        
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        doc.save(response)
        return response
        
    except Exception as e:
        logger.error(f'Error generating AMC specification report: {str(e)}', exc_info=True)
        from django.contrib import messages
        messages.error(request, f'Error generating specification report: {str(e)}')
        return redirect('amc_groups', category=category)


@login_required(login_url='login')
def amc_download_forwarding_letter(request, category):
    """
    Generate forwarding letter from live AMC items.
    """
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    
    if request.method != 'POST':
        return redirect('amc_groups', category=category)
    
    try:
        items_json = request.POST.get('items', '[]')
        work_name = request.POST.get('work_name', '{{NAME_OF_WORK}}')
        total_amount = request.POST.get('total_amount', '0.00')
        
        items = json.loads(items_json)
        
        if not items:
            from django.contrib import messages
            messages.error(request, 'No items with quantities to generate forwarding letter')
            return redirect('amc_groups', category=category)
        
        try:
            grand_total = float(total_amount.replace(',', '').replace('Rs.', '').replace('₹', '').strip())
        except:
            grand_total = 0.0
        
        financial_year = _get_current_financial_year()
        today = timezone.now().date()
        
        # Get user's letter settings
        letter_settings = _get_letter_settings(request.user)
        
        doc = Document()
        placeholder_color = RGBColor(169, 169, 169)
        
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Header - Government/Organization name
        header1 = doc.add_paragraph()
        header1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if letter_settings and letter_settings.government_name:
            run1 = header1.add_run(letter_settings.government_name.upper())
            run1.font.bold = True
            run1.font.size = Pt(14)
        else:
            run1 = header1.add_run('[GOVERNMENT / ORGANIZATION NAME]')
            run1.font.bold = True
            run1.font.size = Pt(14)
            run1.font.color.rgb = placeholder_color
            run1.font.italic = True
        
        # Header - Department name
        header2 = doc.add_paragraph()
        header2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if letter_settings and letter_settings.department_name:
            run2 = header2.add_run(letter_settings.department_name.upper())
            run2.font.bold = True
            run2.font.size = Pt(13)
        else:
            run2 = header2.add_run('[DEPARTMENT NAME]')
            run2.font.bold = True
            run2.font.size = Pt(13)
            run2.font.color.rgb = placeholder_color
            run2.font.italic = True
        
        doc.add_paragraph()
        
        from_to_table = doc.add_table(rows=1, cols=2)
        from_to_table.autofit = True
        
        from_cell = from_to_table.cell(0, 0)
        from_para = from_cell.paragraphs[0]
        from_para.add_run('From: -\n')
        
        # From section - Officer details
        if letter_settings and letter_settings.officer_name:
            name_qual = letter_settings.officer_name
            if letter_settings.officer_qualification:
                name_qual += f", {letter_settings.officer_qualification}"
            from_run1 = from_para.add_run(f'{name_qual},\n')
        else:
            from_run1 = from_para.add_run('[Officer Name, Qualification],\n')
            from_run1.font.color.rgb = placeholder_color
            from_run1.font.italic = True
        
        if letter_settings and letter_settings.officer_designation:
            from_run2 = from_para.add_run(f'{letter_settings.officer_designation},\n')
        else:
            from_run2 = from_para.add_run('[Designation],\n')
            from_run2.font.color.rgb = placeholder_color
            from_run2.font.italic = True
        
        if letter_settings and (letter_settings.sub_division or letter_settings.office_address):
            sub_addr = letter_settings.sub_division
            if letter_settings.office_address:
                sub_addr += f", {letter_settings.office_address}" if sub_addr else letter_settings.office_address
            from_run3 = from_para.add_run(f'{sub_addr}.')
        else:
            from_run3 = from_para.add_run('[Sub Division, Office Address].')
            from_run3.font.color.rgb = placeholder_color
            from_run3.font.italic = True
        
        to_cell = from_to_table.cell(0, 1)
        to_para = to_cell.paragraphs[0]
        to_para.add_run('To,\n')
        
        if letter_settings and letter_settings.recipient_designation:
            to_run1 = to_para.add_run(f'{letter_settings.recipient_designation},\n')
        else:
            to_run1 = to_para.add_run('[Officer Designation],\n')
            to_run1.font.color.rgb = placeholder_color
            to_run1.font.italic = True
        
        if letter_settings and letter_settings.recipient_division:
            to_run2 = to_para.add_run(f'{letter_settings.recipient_division},\n')
        else:
            to_run2 = to_para.add_run('[Division Name],\n')
            to_run2.font.color.rgb = placeholder_color
            to_run2.font.italic = True
        
        if letter_settings and letter_settings.recipient_address:
            to_run3 = to_para.add_run(f'{letter_settings.recipient_address}.')
        else:
            to_run3 = to_para.add_run('[Address].')
            to_run3.font.color.rgb = placeholder_color
            to_run3.font.italic = True
        
        doc.add_paragraph()
        
        lr_para = doc.add_paragraph()
        lr_para.add_run('Lr No. ')
        if letter_settings and letter_settings.office_code:
            lr_code = lr_para.add_run(letter_settings.office_code)
            lr_code.font.underline = True
        else:
            lr_placeholder = lr_para.add_run('[Office Code]')
            lr_placeholder.font.color.rgb = placeholder_color
            lr_placeholder.font.italic = True
            lr_placeholder.font.underline = True
        lr_para.add_run(f'/{financial_year}/          ')
        lr_para.add_run(f'\t\t\t\t\tDate:-    - {today.strftime("%m")} - {today.year}.')
        
        doc.add_paragraph()
        
        sir_para = doc.add_paragraph()
        sir_para.add_run('Sir,')
        
        doc.add_paragraph()
        
        subject_para = doc.add_paragraph()
        subj_run = subject_para.add_run('Sub:-')
        subj_run.font.underline = True
        subject_para.add_run('\t')
        subj_work = subject_para.add_run(f'{work_name} ')
        subj_work.font.bold = True
        subject_para.add_run(f'for the year {financial_year}.  -  Submission  -  Request for obtaining administrative sanction  -  Regarding.')
        
        doc.add_paragraph()
        
        ref_para = doc.add_paragraph()
        ref_run = ref_para.add_run('Ref:-')
        ref_run.font.underline = True
        ref_para.add_run('\tMemo No.')
        ref_placeholder = ref_para.add_run('[Reference Number]')
        ref_placeholder.font.color.rgb = placeholder_color
        ref_placeholder.font.italic = True
        ref_placeholder.font.underline = True
        ref_para.add_run(f'/{financial_year} Dt.')
        ref_date_placeholder = ref_para.add_run('[DD.MM.YYYY]')
        ref_date_placeholder.font.color.rgb = placeholder_color
        ref_date_placeholder.font.italic = True
        ref_date_placeholder.font.underline = True
        
        doc.add_paragraph()
        
        stars_para = doc.add_paragraph()
        stars_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        stars_para.add_run('**.**')
        
        doc.add_paragraph()
        
        body_para = doc.add_paragraph()
        body_para.add_run('With reference to the subject cited, I submit here ')
        with_run = body_para.add_run('with  1')
        with_run.font.underline = True
        body_para.add_run(' No. estimate for the following work for the amount specified.')
        
        doc.add_paragraph()
        
        table = doc.add_table(rows=2, cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for cell in table.columns[0].cells:
            cell.width = Inches(0.5)
        for cell in table.columns[1].cells:
            cell.width = Inches(4.5)
        for cell in table.columns[2].cells:
            cell.width = Inches(1.5)
        
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Sl.\nNo'
        header_cells[1].text = 'Name of work'
        header_cells[2].text = 'Amount'
        
        for cell in header_cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.bold = True
        
        row_cells = table.rows[1].cells
        row_cells[0].text = '1'
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[1].text = work_name
        
        formatted_amount = _format_indian_number(grand_total)
        row_cells[2].text = f"Rs.{formatted_amount}"
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()
        
        spec_para = doc.add_paragraph()
        spec_para.add_run("Specification report accompanying the estimate explains the necessity and provisions made therein in detail.")
        
        doc.add_paragraph()
        
        request_para = doc.add_paragraph()
        request_para.add_run('I request the ')
        if letter_settings and letter_settings.superior_designation:
            req_run = request_para.add_run(letter_settings.superior_designation)
        else:
            req_placeholder = request_para.add_run('[Superior Officer Designation]')
            req_placeholder.font.color.rgb = placeholder_color
            req_placeholder.font.italic = True
        request_para.add_run(' to kindly arrange to obtain administrative sanction for the above estimate and arrange to finalize the agency at the earliest for taking up the work.')
        
        doc.add_paragraph()
        
        enc_para = doc.add_paragraph()
        enc_para.add_run('Enclosure: -')
        doc.add_paragraph('Estimate  - 1 No.')
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        sign_para = doc.add_paragraph()
        sign_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sign_para.add_run('Yours faithfully,')
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        if letter_settings and letter_settings.officer_designation:
            run_title = title_para.add_run(f'{letter_settings.officer_designation}\n')
            run_title.font.bold = True
        else:
            run_title = title_para.add_run('[Officer Designation]\n')
            run_title.font.bold = True
            run_title.font.color.rgb = placeholder_color
            run_title.font.italic = True
        
        if letter_settings and letter_settings.sub_division:
            sub_div_run = title_para.add_run(f'{letter_settings.sub_division},\n')
        else:
            sub_div_run = title_para.add_run('[Sub Division Name],\n')
            sub_div_run.font.color.rgb = placeholder_color
            sub_div_run.font.italic = True
        
        if letter_settings and letter_settings.office_address:
            addr_run = title_para.add_run(f'{letter_settings.office_address}.')
        else:
            addr_run = title_para.add_run('[Office Address].')
            addr_run.font.color.rgb = placeholder_color
            addr_run.font.italic = True
        
        doc.add_paragraph()
        
        copy_para = doc.add_paragraph()
        copy_para.add_run('Copy to the ')
        if letter_settings and (letter_settings.copy_to_designation or letter_settings.copy_to_section):
            copy_text = letter_settings.copy_to_designation or ''
            if letter_settings.copy_to_section:
                copy_text += f", {letter_settings.copy_to_section}" if copy_text else letter_settings.copy_to_section
            copy_run = copy_para.add_run(copy_text)
        else:
            copy_placeholder = copy_para.add_run('[Officer Designation, Section Name]')
            copy_placeholder.font.color.rgb = placeholder_color
            copy_placeholder.font.italic = True
        copy_para.add_run(' for information.')
        
        safe_name = work_name.replace(" ", "_").replace("/", "_").replace("{{", "").replace("}}", "")[:25]
        filename = f'AMC_Forwarding_Letter_{safe_name}_{timezone.now().strftime("%Y%m%d_%H%M%S")}.docx'
        
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        doc.save(response)
        return response
        
    except Exception as e:
        logger.error(f'Error generating AMC forwarding letter: {str(e)}', exc_info=True)
        from django.contrib import messages
        messages.error(request, f'Error generating forwarding letter: {str(e)}')
        return redirect('amc_groups', category=category)


@login_required(login_url='login')
def amc_clear_output(request, category):
    """
    Clear all AMC session data.
    """
    request.session["amc_fetched_items"] = []
    request.session["amc_qty_map"] = {}
    request.session["amc_work_name"] = ""
    request.session["amc_grand_total"] = ""

    group = request.GET.get("group")
    if group:
        return redirect("amc_items", category=category, group=group)

    return redirect("amc_groups", category=category)


@login_required(login_url='login')
def amc_save_qty_map(request, category):
    """
    AJAX endpoint to save AMC quantity map and grand total to session.
    """
    if request.method == "POST":
        try:
            qty_map_str = request.POST.get("qty_map", "")
            grand_total_str = request.POST.get("grand_total", "")
            work_name = request.POST.get("work_name", "")
            
            if qty_map_str:
                try:
                    qty_map = json.loads(qty_map_str)
                    if isinstance(qty_map, dict):
                        request.session["amc_qty_map"] = qty_map
                except json.JSONDecodeError:
                    pass
            
            if grand_total_str:
                request.session["amc_grand_total"] = grand_total_str
            
            if work_name:
                request.session["amc_work_name"] = work_name
            
            return JsonResponse({"status": "ok"})
        except Exception as e:
            return JsonResponse({"status": "error", "message": str(e)}, status=400)
    
    return JsonResponse({"status": "error", "message": "POST required"}, status=405)


@login_required(login_url='login')
def amc_download_output(request, category):
    """
    Generate AMC output Excel file - similar to download_output for New Estimate.
    """
    fetched = request.session.get("amc_fetched_items", [])
    if not fetched:
        return JsonResponse({"error": "No items selected"}, status=400)

    # Parse input
    item_qtys = {}
    work_name = ""
    grand_total = None

    if request.method == "POST":
        qty_map_str = request.POST.get("qty_map", "")
        if qty_map_str:
            try:
                raw = json.loads(qty_map_str)
                if isinstance(raw, dict):
                    for k, v in raw.items():
                        try:
                            item_qtys[str(k)] = float(v)
                        except Exception:
                            continue
            except Exception:
                pass

        work_name = (request.POST.get("work_name") or "").strip()
        
        grand_total_str = request.POST.get("grand_total", "").strip()
        if grand_total_str:
            try:
                grand_total = float(grand_total_str)
            except ValueError:
                grand_total = None

    work_type = (request.POST.get("work_type")
                 or request.session.get("amc_work_type")
                 or "original").lower()
    request.session["amc_work_type"] = work_type
    
    # Get selected backend ID from session (for multi-backend support)
    amc_selected_backend_id = request.session.get("amc_selected_backend_id")

    from django.conf import settings as django_settings
    
    if getattr(django_settings, 'CELERY_TASK_ALWAYS_EAGER', True):
        # Synchronous mode
        try:
            org = get_org_from_request(request)
            
            job = Job.objects.create(
                organization=org,
                user=request.user,
                job_type='generate_output_excel',
                status='queued',
                current_step="Processing AMC...",
            )
            
            job.result = {
                'fetched_items': fetched,
                'qty_map': item_qtys,
                'work_name': work_name,
                'work_type': work_type,
                'grand_total': grand_total,
                'module': 'amc',
                'backend_id': amc_selected_backend_id,
            }
            job.save()
            
            # Call task function directly (synchronous)
            from core.tasks import generate_output_excel
            result = generate_output_excel.apply(args=(
                job.id,
                category,
                json.dumps(item_qtys),
                json.dumps({}),  # unit_map
                work_name,
                work_type,
                grand_total,
                None,  # excess_tp_percent
                None,  # ls_special_name
                None,  # ls_special_amount
                None,  # deduct_old_material
                amc_selected_backend_id,
            )).get()
            
            job.refresh_from_db()
            
            if job.status == 'completed' and job.result.get('output_file_id'):
                from core.models import OutputFile
                try:
                    output_file = OutputFile.objects.get(id=job.result['output_file_id'])
                    from django.http import FileResponse
                    import os
                    
                    if output_file.file and os.path.exists(output_file.file.path):
                        response = FileResponse(
                            open(output_file.file.path, 'rb'),
                            as_attachment=True,
                            filename=output_file.filename or f"amc_{category}_output.xlsx",
                        )
                        return response
                except OutputFile.DoesNotExist:
                    pass
            
            return JsonResponse({
                'job_id': job.id,
                'status': job.status,
                'message': job.current_step or 'Processing complete',
                'error': job.error_message if job.status == 'failed' else None,
            })
                
        except Exception as e:
            logger.error(f"Failed to generate AMC output Excel: {e}")
            return JsonResponse({"error": str(e)}, status=500)
    
    # Async mode with Celery
    try:
        org = get_org_from_request(request)
        
        job = Job.objects.create(
            organization=org,
            user=request.user,
            job_type='generate_output_excel',
            status='queued',
            current_step="Queued for AMC processing",
        )
        
        job.result = {
            'fetched_items': fetched,
            'qty_map': item_qtys,
            'work_name': work_name,
            'work_type': work_type,
            'grand_total': grand_total,
            'module': 'amc',
            'backend_id': amc_selected_backend_id,
        }
        job.save()
        
        from core.tasks import generate_output_excel
        task = generate_output_excel.delay(
            job.id,
            category,
            json.dumps(item_qtys),
            json.dumps({}),  # unit_map
            work_name,
            work_type,
            grand_total,
            None,  # excess_tp_percent
            None,  # ls_special_name
            None,  # ls_special_amount
            None,  # deduct_old_material
            amc_selected_backend_id,
        )
        
        job.celery_task_id = task.id
        job.save()
        
        return JsonResponse({
            'job_id': job.id,
            'status_url': reverse('job_status', args=[job.id]),
            'message': f'Generating AMC {category} output. Please wait...'
        })
        
    except Exception as e:
        logger.error(f"Failed to enqueue AMC output Excel task: {e}")
        return JsonResponse({"error": str(e)}, status=500)
