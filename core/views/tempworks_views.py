# Auto-generated from core/views.py split
import json
import os
import re
import logging
from copy import copy

import inflect
from docx import Document
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, Border, Side, PatternFill
from django.utils import timezone
from django.urls import reverse
from django.core.files.uploadedfile import InMemoryUploadedFile
from django.contrib.auth.decorators import login_required

from django.conf import settings
from django.http import HttpResponse, JsonResponse, HttpResponseNotAllowed
from django.shortcuts import render, redirect, get_object_or_404
from django.template.loader import render_to_string
from django.views.decorators.http import require_POST
from django.utils.crypto import get_random_string

from accounts.models import UserCustomBackend
from core.utils_group_order import apply_group_order, save_group_order

import io
from io import BytesIO
from difflib import SequenceMatcher

from ..models import Project, SelfFormattedTemplate, Estimate, Organization, Membership, Upload, Job, OutputFile, LetterSettings
from ..decorators import org_required, role_required

logger = logging.getLogger(__name__)
from ..tasks import process_excel_upload, generate_bill_pdf, generate_workslip_pdf, generate_bill_document_task
from ..utils_excel import load_backend, copy_block_with_styles_and_formulas, build_temp_day_rates

p_engine = inflect.engine()
BILL_TEMPLATES_DIR = os.path.join(settings.BASE_DIR, "core", "templates", "core", "bill_templates")
_inflect_engine = inflect.engine()

from .utils import (get_org_from_request, _apply_print_settings,
    _format_indian_number, _number_to_words_rupees,
    _get_current_financial_year, _get_letter_settings)

@login_required(login_url='login')
def tempworks_home(request):
    """
    Step 1 in Temporary Works flow: ask user to pick Single Event vs Multiple Events.
    Stores the chosen mode in session, then user proceeds to category chooser.
    Clears only the temp session keys.
    """
    request.session["temp_entries"] = []  # list of entries (single shape OR multi shape)
    request.session["temp_work_name"] = ""
    request.session["temp_selected_backend_id"] = None
    request.session["current_saved_work_id"] = None
    # Default mode: keep existing single-event behavior if user lands directly via legacy URL
    if "temp_mode" not in request.session:
        request.session["temp_mode"] = "single"
    return render(request, "core/tempworks_mode_select.html")


@login_required(login_url='login')
def tempworks_set_mode(request, mode):
    """
    Persist the user's Single/Multiple choice.
    - single: go straight to the category chooser
    - multi: go to the events-setup screen first (enter all event names once)
    """
    mode = (mode or "").strip().lower()
    if mode not in ("single", "multi"):
        mode = "single"
    request.session["temp_mode"] = mode
    # Reset entries when switching mode to avoid mixing shapes
    request.session["temp_entries"] = []
    if mode == "multi":
        # Keep any existing events list so user can re-enter the same screen and edit
        request.session.setdefault("temp_events_list", [])
        return redirect("tempworks_events_setup")
    # single mode
    request.session["temp_events_list"] = []
    return render(request, "core/tempworks_category_select.html", {"temp_mode": mode})


@login_required(login_url='login')
def tempworks_events_setup(request):
    """
    Step 1b (multi mode only): user enters the list of event names for the whole work.
    GET  -> render the form with existing names (if any).
    POST -> save the list and forward to the category chooser.
    Each event in session is stored as {"id": "<8 char>", "name": "<event name>"}.
    """
    if request.session.get("temp_mode") != "multi":
        return redirect("tempworks_home")

    if request.method == "POST":
        raw_names = request.POST.getlist("event_name")
        raw_ids = request.POST.getlist("event_id")
        existing = request.session.get("temp_events_list", []) or []
        id_to_name = {e.get("id"): e.get("name") for e in existing if isinstance(e, dict)}

        cleaned = []
        used_ids = set()
        for i, name in enumerate(raw_names):
            n = (name or "").strip()
            if not n:
                continue
            eid = (raw_ids[i] if i < len(raw_ids) else "") or ""
            if not eid or eid in used_ids or eid not in id_to_name:
                eid = get_random_string(8)
                while eid in used_ids:
                    eid = get_random_string(8)
            used_ids.add(eid)
            cleaned.append({"id": eid, "name": n})

        request.session["temp_events_list"] = cleaned
        # Reset entries because events list changed and existing references may be stale
        # Only reset if the list of event IDs actually changed
        old_ids = set(id_to_name.keys())
        new_ids = set(e["id"] for e in cleaned)
        if old_ids != new_ids:
            # Prune dropped events from any existing entries (preserve still-valid ones)
            valid_names = {e["name"] for e in cleaned}
            entries = request.session.get("temp_entries", []) or []
            for ent in entries:
                if (ent or {}).get("mode") == "multi":
                    ent["events"] = [
                        ev for ev in (ent.get("events") or [])
                        if (ev.get("event_name") or "").strip() in valid_names
                    ]
            request.session["temp_entries"] = entries
        return render(request, "core/tempworks_category_select.html", {"temp_mode": "multi"})

    events = request.session.get("temp_events_list", []) or []
    return render(request, "core/tempworks_events_setup.html", {"events": events})


@login_required(login_url='login')
@require_POST
def tempworks_ajax_update_events(request):
    """
    AJAX: replace the work's events list. Used by the 'Edit Events' modal on the 3-panel page.
    POST JSON: {"events": [{"id": "...", "name": "..."}, ...]}
    Returns:  {"ok": True, "events": [...]}
    Also prunes per-item event refs that no longer match any event name in the new list.
    """
    try:
        data = json.loads(request.body.decode("utf-8") or "{}")
    except Exception:
        return JsonResponse({"ok": False, "error": "invalid json"}, status=400)

    raw = data.get("events") or []
    cleaned = []
    used_ids = set()
    for item in raw:
        if not isinstance(item, dict):
            continue
        name = (item.get("name") or "").strip()
        if not name:
            continue
        eid = (item.get("id") or "").strip() or get_random_string(8)
        while eid in used_ids:
            eid = get_random_string(8)
        used_ids.add(eid)
        cleaned.append({"id": eid, "name": name})

    request.session["temp_events_list"] = cleaned

    # Prune orphaned event references in entries (events removed from list)
    valid_names = {e["name"] for e in cleaned}
    entries = request.session.get("temp_entries", []) or []
    for ent in entries:
        if (ent or {}).get("mode") == "multi":
            ent["events"] = [
                ev for ev in (ent.get("events") or [])
                if (ev.get("event_name") or "").strip() in valid_names
            ]
    request.session["temp_entries"] = entries

    return JsonResponse({"ok": True, "events": cleaned})


@login_required(login_url='login')
def temp_groups(request, category):
    """
    Step 1 in Temporary Works: show groups for given temp category.
    Example categories: 'temp_electrical', 'temp_civil'
    """
    from core.utils_excel import get_available_backends_for_module
    
    # Handle backend switching via URL parameter
    url_backend_id = request.GET.get("backend_id")
    if url_backend_id:
        try:
            request.session["temp_selected_backend_id"] = int(url_backend_id)
        except (ValueError, TypeError):
            pass
    
    temp_selected_backend_id = request.session.get("temp_selected_backend_id")
    
    # Map temp category to base category for backend lookup
    base_category = category.replace('temp_', '')  # temp_electrical -> electrical
    
    try:
        items_list, groups_map, _, ws_data, filepath = load_backend(
            category, settings.BASE_DIR,
            backend_id=temp_selected_backend_id,
            module_code='temp_works',  # Use temp_works module's own backends
            user=request.user
        )
    except FileNotFoundError as e:
        logger.info(f"No backend data available for temp category {category} - showing Coming Soon")
        # Check if the other category has backends available
        other_base = 'electrical' if base_category == 'civil' else 'civil'
        other_category = f'temp_{other_base}'
        other_available = False
        try:
            other_backends = get_available_backends_for_module('temp_works', other_base)
            other_available = len(other_backends) > 0
        except:
            pass
        return render(request, "core/coming_soon.html", {
            "category": category,
            "module_name": "Temporary Works",
            "other_category": other_category,
            "other_category_available": other_available,
        })
    except Exception as e:
        logger.error(f"Error loading temp backend for {category}: {e}")
        return render(request, "core/error.html", {
            "error_title": "Loading Error",
            "error_message": f"Could not load temporary works data: {str(e)}",
        })
        
    groups = apply_group_order(request.user, 'temp', category, groups_map.keys())

    if not groups:
        return render(
            request,
            "core/temp_items.html",
            {
                "category": category,
                "group": "",
                "groups": [],
                "items_info": [],
                "entries": [],
                "day_rates_json": "{}",
                "is_temporary": True,
                "error": "No groups found in Temporary Works backend.",
            },
        )

    default_group = request.GET.get("group") or groups[0]
    return redirect("temp_items", category=category, group=default_group)


@login_required(login_url='login')
def temp_items(request, category, group):
    """
    Temporary Works UI:
    - rates must be computed from temp workbook (Column C day, Column J rate)
    - Column J has formulas (ROUND etc), so we evaluate if cached value is missing
    """
    from core.utils_excel import get_available_backends_for_module
    
    # Handle backend switching via URL parameter
    url_backend_id = request.GET.get("backend_id")
    if url_backend_id:
        try:
            request.session["temp_selected_backend_id"] = int(url_backend_id)
        except (ValueError, TypeError):
            pass
    
    temp_selected_backend_id = request.session.get("temp_selected_backend_id")
    
    # Map temp category to base category for backend lookup
    base_category = category.replace('temp_', '')  # temp_electrical -> electrical
    
    # Get available backends for dropdown (temp_works has its own backends)
    available_backends = get_available_backends_for_module('temp_works', base_category)
    
    try:
        items_list, groups_map, units_map, ws_src, filepath = load_backend(
            category, settings.BASE_DIR,
            backend_id=temp_selected_backend_id,
            module_code='temp_works',  # Use temp_works module's own backends
            user=request.user
        )
    except FileNotFoundError as e:
        logger.info(f"No backend data available for temp category {category} - showing Coming Soon")
        other_base = 'electrical' if base_category == 'civil' else 'civil'
        other_category = f'temp_{other_base}'
        other_available = False
        try:
            other_backends = get_available_backends_for_module('temp_works', other_base)
            other_available = len(other_backends) > 0
        except:
            pass
        return render(request, "core/coming_soon.html", {
            "category": category,
            "module_name": "Temporary Works",
            "other_category": other_category,
            "other_category_available": other_available,
        })
    except Exception as e:
        logger.error(f"Error loading temp backend for {category}: {e}")
        return render(request, "core/error.html", {
            "error_title": "Loading Error",
            "error_message": f"Could not load temporary works data: {str(e)}",
        })

    groups = apply_group_order(request.user, 'temp', category, groups_map.keys())
    group_items = groups_map.get(group, [])

    detected_names = {i["name"] for i in items_list}
    display_items = [name for name in group_items if name in detected_names]

    # Build item subtypes map: items with ":" are subtypes
    # Support both " : " (with spaces) and ":" (without spaces)
    import re as _re
    _colon_re = _re.compile(r'\s*:\s*')

    def _has_colon(name):
        return bool(_colon_re.search(name))

    def _split_parent(name):
        return _colon_re.split(name, 1)[0].strip()

    item_subtypes = {}  # parent_name -> [list of full subtype names]
    parent_items = set()

    for name in display_items:
        if _has_colon(name):
            parent_name = _split_parent(name)
            if parent_name not in item_subtypes:
                item_subtypes[parent_name] = []
            item_subtypes[parent_name].append(name)
            parent_items.add(parent_name)

    items_info = []
    seen_parents = set()
    for name in display_items:
        if _has_colon(name):
            parent_name = _split_parent(name)
            if parent_name not in seen_parents:
                subtypes_list = item_subtypes.get(parent_name, [])
                items_info.append({
                    "name": parent_name,
                    "has_subtypes": True,
                    "subtypes": json.dumps(subtypes_list),
                    "subtypes_count": len(subtypes_list),
                })
                seen_parents.add(parent_name)
        else:
            items_info.append({
                "name": name,
                "has_subtypes": False,
                "subtypes": "[]",
                "subtypes_count": 0,
            })

    # units mapping (same as your code)
    item_to_group = {}
    for grp_name, item_list_in_grp in groups_map.items():
        for nm in item_list_in_grp:
            item_to_group.setdefault(nm, grp_name)

    def units_for(name):
        # Priority 1: Use unit from Groups sheet (Column D) via units_map
        backend_unit = units_map.get(name, "") if units_map else ""
        if backend_unit:
            bu = backend_unit.lower()
            if bu in ("mtrs", "mtr", "metre", "meters"):
                return ("Mtrs", "Mtr")
            elif bu in ("pts", "pt", "point", "points"):
                return ("Pts", "Pt")
            elif bu in ("nos", "no"):
                return ("Nos", "No")
            else:
                return (backend_unit, backend_unit)
        # Fallback: group-based
        grp_name = (item_to_group.get(name, "") or "").lower()
        if grp_name in ("piping", "wiring & cables", "wiring and cables"):
            return ("Mtrs", "Mtr")
        elif grp_name == "points":
            return ("Pts", "Pt")
        else:
            return ("Nos", "No")

    temp_entries = request.session.get("temp_entries", []) or []
    temp_mode = request.session.get("temp_mode", "single")
    temp_events_list = request.session.get("temp_events_list", []) or []
    # Map event_name -> {id, name} for quick lookup when rendering checkboxes
    event_name_to_id = {e["name"]: e["id"] for e in temp_events_list if isinstance(e, dict) and e.get("name")}
    display_entries = []
    for idx, ent in enumerate(temp_entries, start=1):
        plural, _singular = units_for(ent["name"])
        entry_mode = ent.get("mode", "single")
        # For multi entries, build a per-event-list rendering: for each defined event,
        # include the entry's stored days/qty if it had a selection, else default.
        per_event_rows = []
        if entry_mode == "multi":
            stored = {(ev.get("event_name") or "").strip(): ev for ev in (ent.get("events") or [])}
            for ev_def in temp_events_list:
                ev_name = ev_def.get("name", "")
                sel = stored.get(ev_name)
                per_event_rows.append({
                    "event_id": ev_def.get("id", ""),
                    "event_name": ev_name,
                    "selected": sel is not None,
                    "days": (sel or {}).get("days", 1),
                    "qty": (sel or {}).get("qty", ""),
                })
        display_entries.append(
            {
                "id": ent["id"],
                "sl": idx,
                "name": ent["name"],
                "unit": plural,
                "mode": entry_mode,
                # Single-mode fields (unchanged):
                "qty": ent.get("qty", ""),
                "days": ent.get("days", 1),
                # Multi-mode: pre-built list of {event_id, event_name, selected, days, qty}
                "event_rows": per_event_rows if entry_mode == "multi" else [],
            }
        )

    # âœ… IMPORTANT: build day rates using filepath (so we can load wb twice)
    day_rates = build_temp_day_rates(filepath, items_list)
    day_rates_json = json.dumps(day_rates)

    work_name = request.session.get("temp_work_name", "") or ""
    grand_total = request.session.get("temp_grand_total", "") or ""

    context = {
        "category": category,
        "group": group,
        "groups": groups,
        "items_info": items_info,
        "entries": display_entries,
        "day_rates_json": day_rates_json,
        "is_temporary": True,
        "work_name": work_name,
        "grand_total": grand_total,
        "available_backends": available_backends,
        "selected_backend_id": temp_selected_backend_id,
        "custom_groups": UserCustomBackend.custom_group_names(request.user, 'temp_works', base_category),
        "temp_mode": temp_mode,
        "temp_events_list": temp_events_list,
    }
    return render(request, "core/temp_items.html", context)


@login_required(login_url='login')
def temp_day_rates_debug(request, category):
    """Debug endpoint: return JSON of computed day rates for a category.
    Useful to inspect what the view passes to the template.
    """
    from django.http import JsonResponse

    try:
        items_list, groups_map, _, ws_src, filepath = load_backend(category, settings.BASE_DIR)
        day_rates = build_temp_day_rates(filepath, items_list)
    except Exception as e:
        return JsonResponse({"error": str(e)}, status=500)

    return JsonResponse(day_rates, safe=False)

@login_required(login_url='login')
def temp_add_item(request, category, group, item):
    """
    Add one more row for given item (duplicates allowed).
    """
    temp_entries = request.session.get("temp_entries", []) or []
    entry = {
        "id": get_random_string(8),
        "name": item,
        # we don't store qty/days here; they live in the form & are sent on POST
    }
    temp_entries.append(entry)
    request.session["temp_entries"] = temp_entries

    # Preserve 'work_name' typed in the box
    work_name = request.GET.get("work_name")
    if work_name is not None:
        request.session["temp_work_name"] = work_name

    return redirect("temp_items", category=category, group=group)


# -----------------------
# TEMP AJAX ADD ITEM (no page reload)
# -----------------------
@login_required(login_url='login')
def temp_ajax_add_item(request, category):
    """
    AJAX endpoint to add an item to temp entries without page reload.
    POST with JSON: { "item": "item_name", "work_name": "..." }
    Returns JSON: { "status": "ok", "entries": [...], "entry_id": "...", "item_info": {...} }
    """
    if request.method != "POST":
        return JsonResponse({"status": "error", "message": "POST required"}, status=405)
    
    try:
        if request.content_type and 'application/json' in request.content_type:
            data = json.loads(request.body.decode("utf-8") or "{}")
        else:
            data = request.POST
        
        item = data.get("item", "").strip()
        work_name = data.get("work_name")
        
        if not item:
            return JsonResponse({"status": "error", "message": "No item specified"}, status=400)
        
        temp_entries = request.session.get("temp_entries", []) or []
        temp_mode = request.session.get("temp_mode", "single")
        entry = {
            "id": get_random_string(8),
            "name": item,
        }
        if temp_mode == "multi":
            entry["mode"] = "multi"
            entry["events"] = []  # filled by client via save_state
        temp_entries.append(entry)
        request.session["temp_entries"] = temp_entries
        
        if work_name is not None:
            request.session["temp_work_name"] = work_name
        
        # Get unit from backend units_map (Column D of Groups sheet)
        item_info = {"name": item, "unit": "Nos"}
        try:
            items_list, groups_map, units_map, ws_data, filepath = load_backend(category, settings.BASE_DIR)
            
            # Get unit from backend units_map
            unit = units_map.get(item, "Nos")
            
            item_info["unit"] = unit
        except Exception as e:
            # If we can't determine unit, default to Nos
            pass
        
        return JsonResponse({
            "status": "ok",
            "entries": temp_entries,
            "entry_id": entry["id"],
            "item": item,
            "item_info": item_info,
            "temp_mode": temp_mode,
        })
        
    except Exception as e:
        return JsonResponse({"status": "error", "message": str(e)}, status=400)


# -----------------------
# TEMP AJAX REORDER ITEMS
# -----------------------
@login_required(login_url='login')
def temp_ajax_reorder_items(request, category):
    """
    AJAX endpoint to reorder temp entries list.
    POST with JSON: { "entries": [{"id": "...", "name": "..."}, ...] }
    Returns JSON: { "status": "ok", "entries": [...] }
    """
    if request.method != "POST":
        return JsonResponse({"status": "error", "message": "POST required"}, status=405)
    
    try:
        if request.content_type and 'application/json' in request.content_type:
            data = json.loads(request.body.decode("utf-8") or "{}")
        else:
            return JsonResponse({"status": "error", "message": "JSON required"}, status=400)
        
        new_order = data.get("entries", [])
        
        if not isinstance(new_order, list):
            return JsonResponse({"status": "error", "message": "entries must be a list"}, status=400)
        
        # Validate entries have required fields
        current_entries = request.session.get("temp_entries", []) or []
        current_ids = {e["id"] for e in current_entries}
        
        # Reorder based on provided IDs
        valid_entries = []
        for entry in new_order:
            if isinstance(entry, dict) and entry.get("id") in current_ids:
                # Find the original entry to preserve all data
                original = next((e for e in current_entries if e["id"] == entry["id"]), None)
                if original:
                    valid_entries.append(original)
        
        request.session["temp_entries"] = valid_entries
        
        return JsonResponse({
            "status": "ok",
            "entries": valid_entries
        })
        
    except Exception as e:
        return JsonResponse({"status": "error", "message": str(e)}, status=400)


# -----------------------
# TEMP CLEAR ITEMS (stay on same page)
# -----------------------
@login_required(login_url='login')
def temp_clear_items(request, category, group):
    """
    Clear all temp entries but stay on the same page.
    """
    request.session["temp_entries"] = []
    request.session["temp_work_name"] = ""
    request.session["temp_grand_total"] = ""
    return redirect("temp_items", category=category, group=group)


@login_required(login_url='login')
def temp_save_state(request, category):
    """Save current temporary entries (from client) into session.
    Expects JSON body: { entries: [{id, name, qty, days}, ...], work_name: "...", grand_total: "..." }
    """
    from django.http import JsonResponse

    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)

    try:
        payload = json.loads(request.body.decode("utf-8") or "{}")
    except Exception:
        return JsonResponse({"error": "invalid json"}, status=400)

    entries = payload.get("entries", [])
    work_name = payload.get("work_name", "")
    grand_total = payload.get("grand_total", "")

    # Basic validation: ensure list of dicts
    if not isinstance(entries, list):
        entries = []

    # Save into session
    request.session["temp_entries"] = entries
    request.session["temp_work_name"] = work_name or ""
    request.session["temp_grand_total"] = grand_total or ""

    return JsonResponse({"ok": True})


@login_required(login_url='login')
def temp_remove_item(request, category, group, entry_id):
    """
    Remove one selected temp row.
    """
    temp_entries = request.session.get("temp_entries", []) or []
    temp_entries = [e for e in temp_entries if e.get("id") != entry_id]
    request.session["temp_entries"] = temp_entries
    return redirect("temp_items", category=category, group=group)


# -----------------------
# TEMP AJAX REMOVE ITEM
# -----------------------
@login_required(login_url='login')
def temp_ajax_remove_item(request, category):
    """
    AJAX endpoint to remove an item from temp entries without page reload.
    POST with JSON: { "entry_id": "..." }
    Returns JSON: { "status": "ok", "entries": [...] }
    """
    if request.method != "POST":
        return JsonResponse({"status": "error", "message": "POST required"}, status=405)
    
    try:
        if request.content_type and 'application/json' in request.content_type:
            data = json.loads(request.body.decode("utf-8") or "{}")
        else:
            data = request.POST
        
        entry_id = data.get("entry_id", "").strip()
        
        if not entry_id:
            return JsonResponse({"status": "error", "message": "No entry_id specified"}, status=400)
        
        temp_entries = request.session.get("temp_entries", []) or []
        temp_entries = [e for e in temp_entries if e.get("id") != entry_id]
        request.session["temp_entries"] = temp_entries
        
        return JsonResponse({
            "status": "ok",
            "entries": temp_entries
        })
        
    except Exception as e:
        return JsonResponse({"status": "error", "message": str(e)}, status=400)


@login_required(login_url='login')
def temp_download_output(request, category):
    """
    Build Output + Estimate for Temporary Works ONLY.
    Uses:
      - entries_json (from hidden field): list of {id, name, qty, days}
      - work_name
    For each entry we:
      - copy the item block up to the matching 'Hire charges per X day(s)' row
      - take rate from that row (column J)
      - add suffix 'for X day(s)' in Estimate description
    New Estimate module is NOT touched by this view.
    """
    if request.method != "POST":
        return redirect("temp_groups", category=category)

    entries_json = request.POST.get("entries_json", "[]")
    try:
        entries = json.loads(entries_json)
        if not isinstance(entries, list):
            entries = []
    except Exception:
        entries = []

    if not entries:
        return redirect("temp_groups", category=category)

    work_name = (request.POST.get("work_name") or "").strip()
    
    # Get selected backend ID from session (for multi-backend support)
    temp_selected_backend_id = request.session.get("temp_selected_backend_id")

    # ----- load backend -----
    try:
        items_list, groups_map, units_map, ws_src, filepath = load_backend(
            category, settings.BASE_DIR,
            backend_id=temp_selected_backend_id,
            module_code='temp_works'  # Use temp_works module's own backends
        )
    except FileNotFoundError as e:
        logger.error(f"Backend not found for temp download: {category} - {e}")
        return redirect("temp_groups", category=category)
    except Exception as e:
        logger.error(f"Error loading temp backend for download: {category} - {e}")
        return redirect("temp_groups", category=category)
    
    # Also load with data_only=True to get actual cached rate values
    try:
        wb_vals = load_workbook(filepath, data_only=True)
        ws_vals = wb_vals["Master Datas"]
    except Exception as e:
        logger.error(f"Error loading backend values: {e}")
        ws_vals = None
    
    name_to_info = {it["name"]: it for it in items_list}

    # Build day rates map like the UI does (for reliable rate lookup)
    day_rates = build_temp_day_rates(filepath, items_list)
    
    def _norm_name(s):
        """Normalize item name for lookup in day_rates - must match _norm_item_name in utils_excel."""
        s = "" if s is None else str(s)
        s = s.replace("\n", " ").replace("\r", " ")
        s = re.sub(r"\s+", " ", s).strip()
        return s

    def _to_roman_lower(n):
        """1 -> 'i', 4 -> 'iv', etc. Returns '' for non-positive integers."""
        try:
            n = int(n)
        except (TypeError, ValueError):
            return ""
        if n <= 0:
            return ""
        vals = [
            (1000, "m"), (900, "cm"), (500, "d"), (400, "cd"),
            (100, "c"), (90, "xc"), (50, "l"), (40, "xl"),
            (10, "x"), (9, "ix"), (5, "v"), (4, "iv"), (1, "i"),
        ]
        out = []
        for v, sym in vals:
            while n >= v:
                out.append(sym)
                n -= v
        return "".join(out)

    # map item -> group for units
    item_to_group = {}
    for grp_name, item_list_in_grp in groups_map.items():
        for nm in item_list_in_grp:
            item_to_group.setdefault(nm, grp_name)

    def units_for(name):
        # Priority 1: Use unit from Groups sheet (Column D) via units_map
        backend_unit = units_map.get(name, "") if units_map else ""
        if backend_unit:
            bu = backend_unit.lower()
            if bu in ("mtrs", "mtr", "metre", "meters"):
                return ("Mtrs", "Mtr")
            elif bu in ("pts", "pt", "point", "points"):
                return ("Pts", "Pt")
            elif bu in ("nos", "no"):
                return ("Nos", "No")
            else:
                return (backend_unit, backend_unit)
        # Fallback: group-based
        grp_name = (item_to_group.get(name, "") or "").lower()
        if grp_name in ("piping", "wiring & cables", "wiring and cables"):
            return ("Mtrs", "Mtr")
        elif grp_name == "points":
            return ("Pts", "Pt")
        else:
            return ("Nos", "No")

    # ----- workbook & styles -----
    wb = Workbook()
    ws_out = wb.active
    ws_out.title = "Output"

    thin = Side(border_style="thin", color="000000")
    border_all = Border(top=thin, left=thin, right=thin, bottom=thin)

    # Add "Name of Work" header at top of Output sheet
    ws_out.merge_cells("A1:J1")
    hdr = ws_out["A1"]
    hdr.value = f"Name of the work : {work_name}" if work_name else "Name of the work : "
    hdr.font = Font(bold=True, size=11)
    hdr.alignment = Alignment(horizontal="left", vertical="center")

    cursor = 3  # start blocks after header + blank row
    rate_rows = {}  # dict mapping entry_index -> row_in_output for rate (single mode only)

    def _entry_mode(e):
        m = (e or {}).get("mode", "single")
        return "multi" if m == "multi" else "single"

    def _entry_max_days(e):
        """For multi: max days across events; for single: just entry.days."""
        if _entry_mode(e) == "multi":
            evs = e.get("events") or []
            ds = []
            for ev in evs:
                try:
                    d = int(ev.get("days") or 0)
                    if d > 0:
                        ds.append(d)
                except (TypeError, ValueError):
                    pass
            return max(ds) if ds else 1
        try:
            return int(e.get("days") or 1)
        except (TypeError, ValueError):
            return 1

    # =====================================================
    # 1) OUTPUT SHEET: one block per entry (supports dupes)
    # =====================================================
    for idx, entry in enumerate(entries, start=1):
        name = entry.get("name")
        # For multi: use MAX days among events; for single: entry.days
        days = _entry_max_days(entry)

        info = name_to_info.get(name)
        if not info:
            continue

        src_min = info["start_row"]
        src_max = info["end_row"]
        _src_ws = info.get('_source_ws') or ws_src

        # find row that matches "Hire charges per X day(s)" nearest this days
        target_row = None
        search1 = f"per {days} day"
        search2 = f"per {days} days"

        for r in range(src_min, src_max + 1):
            txt = str(_src_ws.cell(row=r, column=4).value or "").lower()
            if "hire charges per" in txt and (search1 in txt or search2 in txt):
                target_row = r
                break

        if target_row is None:
            # fallback: use entire block and rate = last non-empty J
            effective_end = src_max
            rate_src_row = None
            for r in range(src_max, src_min - 1, -1):
                v = _src_ws.cell(row=r, column=10).value
                if v not in (None, ""):
                    rate_src_row = r
                    break
        else:
            effective_end = target_row
            rate_src_row = target_row

        dst_start = cursor

        copy_block_with_styles_and_formulas(
            ws_src=_src_ws,
            ws_dst=ws_out,
            src_min_row=src_min,
            src_max_row=effective_end,
            dst_start_row=dst_start,
            col_start=1,
            col_end=10,
            block_max_row=src_max,
        )

        # Fix rate columns (I=9, J=10) where formulas reference cells outside copied block
        # Overlay actual cached values from ws_vals for rows with day numbers in column C
        if ws_vals:
            for src_r in range(src_min, effective_end + 1):
                dst_r = dst_start + (src_r - src_min)
                # Check if this row has a day number in column C
                day_cell = ws_vals.cell(row=src_r, column=3).value
                if day_cell not in (None, ""):
                    try:
                        day_no = int(float(day_cell))
                        if day_no > 0:
                            # This is a rate row - overlay actual values for columns I and J
                            for col in (9, 10):
                                cached_val = ws_vals.cell(row=src_r, column=col).value
                                if cached_val is not None:
                                    ws_out.cell(row=dst_r, column=col).value = cached_val
                    except (ValueError, TypeError):
                        pass

        # Label first row as Data block
        ws_out.cell(row=dst_start, column=1).value = f"Data {idx}"

        if rate_src_row:
            rate_rows[idx] = dst_start + (rate_src_row - src_min)
        else:
            rate_rows[idx] = None

        cursor += (effective_end - src_min + 1)

    # =====================================================
    # 2) ESTIMATE SHEET
    # =====================================================
    ws_est = wb.create_sheet("Estimate")

    # Title row
    ws_est.merge_cells("A1:H1")
    c1 = ws_est["A1"]
    c1.value = "ESTIMATE"
    c1.font = Font(bold=True, size=14)
    c1.alignment = Alignment(horizontal="center", vertical="center")

    ws_est.merge_cells("A2:H2")
    c2 = ws_est["A2"]
    if work_name:
        c2.value = f"Name of the work : {work_name}"
    else:
        c2.value = "Name of the work : "
    c2.font = Font(bold=True, size=11)
    c2.alignment = Alignment(horizontal="left", vertical="center")

    for row in (1, 2):
        for col in range(1, 9):
            ws_est.cell(row=row, column=col).border = border_all

    # Header row
    headers = ["Sl.No", "Quantity (Unit)", "", "Item Description",
               "Rate", "Per Unit", "", "Amount"]
    for col, text in enumerate(headers, start=1):
        cell = ws_est.cell(row=3, column=col, value=text)
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = border_all

    ws_est.merge_cells("B3:C3")
    ws_est.merge_cells("F3:G3")

    ws_est.column_dimensions["A"].width = 6
    ws_est.column_dimensions["B"].width = 10
    ws_est.column_dimensions["C"].width = 10
    ws_est.column_dimensions["D"].width = 45
    ws_est.column_dimensions["E"].width = 10
    ws_est.column_dimensions["F"].width = 8
    ws_est.column_dimensions["G"].width = 10
    ws_est.column_dimensions["H"].width = 15

    # ---- Fill estimate rows ----
    row_est = 4
    slno = 1

    for idx, entry in enumerate(entries, start=1):
        name = entry.get("name")
        mode = _entry_mode(entry)

        info = name_to_info.get(name)
        if not info:
            continue

        # Get base description from row+2 (standard position in backend)
        start_row = info["start_row"]
        end_row = info["end_row"]

        # Use data_only worksheet to get cached values
        desc_ws = ws_vals if ws_vals else ws_src
        base_desc = desc_ws.cell(row=start_row + 2, column=4).value or ""
        base_desc_str = str(base_desc).strip()

        plural, singular = units_for(name)
        norm_name = _norm_name(name)
        item_day_rates = day_rates.get(norm_name, {})

        def _rate_for_days(days):
            r = item_day_rates.get(days, 0)
            if r == 0 and item_day_rates:
                available_days = sorted([int(k) for k in item_day_rates.keys()])
                closest_day = None
                for d in available_days:
                    if d >= days:
                        closest_day = d
                        break
                if closest_day is None and available_days:
                    closest_day = available_days[-1]
                if closest_day:
                    r = item_day_rates.get(closest_day, 0)
            return r

        def _write_row(desc, qty_val, rate_value, kind="normal"):
            """
            kind:
              - 'normal': single-mode row (Sl.No, qty, rate, amount formula). Increments slno.
              - 'header': multi-mode item header row (Sl.No + bold desc; qty/rate/amount blank). Increments slno.
              - 'event':  multi-mode event sub-row (no Sl.No; qty/rate/amount filled). Does NOT increment slno.
            """
            nonlocal row_est, slno
            is_event = (kind == "event")
            is_header = (kind == "header")

            sl_text = "" if is_event else slno
            a = ws_est.cell(row=row_est, column=1, value=sl_text)
            a.alignment = Alignment(horizontal="center", vertical="center")
            a.border = border_all

            qty_cell_val = "" if (is_header or qty_val in (None, "")) else qty_val
            b = ws_est.cell(row=row_est, column=2, value=qty_cell_val)
            b.alignment = Alignment(horizontal="center", vertical="center")
            b.border = border_all

            unit_text = "" if is_header else plural
            c = ws_est.cell(row=row_est, column=3, value=unit_text)
            c.alignment = Alignment(horizontal="center", vertical="center")
            c.border = border_all

            d_cell = ws_est.cell(row=row_est, column=4, value=desc)
            d_cell.alignment = Alignment(horizontal="justify", vertical="top", wrap_text=True)
            d_cell.border = border_all
            if is_header:
                d_cell.font = Font(bold=True)

            rate_cell_val = "" if (is_header or not rate_value) else rate_value
            e = ws_est.cell(row=row_est, column=5, value=rate_cell_val)
            e.alignment = Alignment(horizontal="center", vertical="center")
            e.border = border_all
            e.number_format = '#,##0.00'

            per_unit_val = "" if is_header else 1
            f_cell = ws_est.cell(row=row_est, column=6, value=per_unit_val)
            f_cell.alignment = Alignment(horizontal="center", vertical="center")
            f_cell.border = border_all

            singular_text = "" if is_header else singular
            g = ws_est.cell(row=row_est, column=7, value=singular_text)
            g.alignment = Alignment(horizontal="center", vertical="center")
            g.border = border_all

            amount_val = "" if is_header else f"=ROUND(B{row_est}*E{row_est},2)"
            h = ws_est.cell(row=row_est, column=8, value=amount_val)
            h.alignment = Alignment(horizontal="center", vertical="center")
            h.border = border_all
            h.number_format = '#,##0.00'

            row_est += 1
            if not is_event:
                slno += 1

        if mode == "multi":
            # New output format:
            #   1) One row with the item description (full base desc, no event suffix)
            #   2) Below it, one row per valid event: "i. <event_name> for X day(s)" — NO item desc repeated
            events = entry.get("events") or []
            valid_events = []
            for ev in events:
                ev_name = (ev.get("event_name") or "").strip()
                try:
                    ev_days = int(ev.get("days") or 0)
                except (TypeError, ValueError):
                    ev_days = 0
                try:
                    ev_qty = float(ev.get("qty") or 0)
                except (TypeError, ValueError):
                    ev_qty = 0.0
                if ev_name and ev_days > 0 and ev_qty > 0:
                    valid_events.append({"name": ev_name, "days": ev_days, "qty": ev_qty})

            if not valid_events:
                # Nothing to output for this item
                continue

            # Header row: item description, no qty/rate/amount on the header line itself
            header_desc = base_desc_str or name
            _write_row(header_desc, None, None, kind="header")

            # Event rows with roman-numeral prefix and no item description
            for i, ev in enumerate(valid_events, start=1):
                roman = _to_roman_lower(i)
                day_word = "day" if ev["days"] == 1 else "days"
                desc = f"{roman}. {ev['name']} for {ev['days']} {day_word}"
                rate_value = _rate_for_days(ev["days"])
                _write_row(desc, ev["qty"], rate_value, kind="event")
        else:
            # Single mode — existing behavior preserved
            try:
                qty_val = float(entry.get("qty") or 0) or None
            except (TypeError, ValueError):
                qty_val = None
            days = int(entry.get("days") or 1)
            suffix = f"for {days} day" if days == 1 else f"for {days} days"
            desc = f"{base_desc_str}, {suffix}" if base_desc_str else suffix
            rate_value = _rate_for_days(days)
            _write_row(desc, qty_val, rate_value)

    # ---- Totals (same style as your main estimate) ----
    ecv_row = row_est
    ws_est.cell(row=ecv_row, column=4, value="ECV")
    ws_est.cell(row=ecv_row, column=8, value=f"=ROUND(SUM(H4:H{ecv_row-1}),2)")

    lc_row = ecv_row + 1
    qc_row = ecv_row + 2
    nac_row = ecv_row + 3
    sub_row = ecv_row + 4
    gst_row = ecv_row + 5
    ls_row = ecv_row + 6
    gt_row = ecv_row + 7

    ws_est.cell(row=lc_row, column=4, value="Add LC @ 1 %")
    ws_est.cell(row=lc_row, column=8, value=f"=ROUND(H{ecv_row}*0.01,2)")

    ws_est.cell(row=qc_row, column=4, value="Add QC @ 1 %")
    ws_est.cell(row=qc_row, column=8, value=f"=ROUND(H{ecv_row}*0.01,2)")

    ws_est.cell(row=nac_row, column=4, value="Add NAC @ 0.1 %")
    ws_est.cell(row=nac_row, column=8, value=f"=ROUND(H{ecv_row}*0.001,2)")

    ws_est.cell(row=sub_row, column=4, value="Sub Total")
    ws_est.cell(row=sub_row, column=8, value=f"=ROUND(H{ecv_row}+H{lc_row}+H{qc_row}+H{nac_row},2)")

    ws_est.cell(row=gst_row, column=4, value="Add GST@18 %")
    ws_est.cell(row=gst_row, column=8, value=f"=ROUND(H{sub_row}*0.18,2)")

    ws_est.cell(row=ls_row, column=4, value="L.S Provision towards unforeseen items")
    ws_est.cell(row=ls_row, column=8, value=f"=ROUND(H{gt_row}-H{gst_row}-H{sub_row},2)")

    ws_est.cell(row=gt_row, column=4, value="Grand Total")
    # Set Grand Total value if provided by user, otherwise leave empty
    grand_total_str = request.POST.get("grand_total", "").strip()
    if grand_total_str:
        try:
            grand_total_val = float(grand_total_str)
            if grand_total_val > 0:
                ws_est.cell(row=gt_row, column=8, value=round(grand_total_val, 2))
        except ValueError:
            pass

    for r in range(ecv_row, gt_row + 1):
        for c in range(1, 9):
            cell = ws_est.cell(row=r, column=c)
            cell.border = border_all
            if c == 4:
                cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            else:
                cell.alignment = Alignment(horizontal="center", vertical="center")
        ws_est.cell(row=r, column=4).font = Font(bold=True)
        ws_est.cell(row=r, column=8).font = Font(bold=True)

    # Reorder sheets: Estimate first, then Output
    if "Estimate" in wb.sheetnames:
        est_idx = wb.sheetnames.index("Estimate")
        if est_idx > 0:
            wb.move_sheet("Estimate", offset=-est_idx)

    # ----- return workbook -----
    _apply_print_settings(wb)
    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    response["Content-Disposition"] = 'attachment; filename="Temp_Estimate.xlsx"'
    wb.save(response)
    return response


@login_required(login_url='login')
def temp_download_specification_report(request, category):
    """
    Generate specification report from live Temporary Works items.
    """
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    if request.method != 'POST':
        return redirect('temp_groups', category=category)
    
    try:
        items_json = request.POST.get('items', '[]')
        work_name = request.POST.get('work_name', '{{NAME_OF_WORK}}')
        total_amount = request.POST.get('total_amount', '0.00')
        
        items = json.loads(items_json)
        
        if not items:
            from django.contrib import messages
            messages.error(request, 'No items with quantities to generate specification report')
            return redirect('temp_groups', category=category)
        
        doc = Document()
        
        title = doc.add_heading('Specification report accompanying the estimate :-', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in title.runs:
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.underline = True
        
        intro_para = doc.add_paragraph()
        intro_para.add_run(f'The estimate is prepared for the work {work_name}')
        
        doc.add_paragraph()
        
        amount_para = doc.add_paragraph()
        amount_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        amount_run = amount_para.add_run(f'Est.Amount: Rs. {total_amount}')
        amount_run.font.bold = True
        amount_run.font.underline = True
        
        doc.add_paragraph()
        
        body_para = doc.add_paragraph('{{BODY_OF_LETTER}}')
        for run in body_para.runs:
            run.font.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph()
        doc.add_paragraph('Hence, this estimate has been prepared accordingly.')
        doc.add_paragraph()
        
        for item in items:
            desc = item.get('desc', '')
            qty = item.get('qty', '')
            days = item.get('days', '')
            unit = item.get('unit', 'Nos')
            
            if qty:
                try:
                    qty_float = float(qty)
                    if qty_float == int(qty_float):
                        qty = str(int(qty_float))
                    else:
                        qty = str(qty_float)
                except:
                    pass
            
            if qty and days:
                bullet_text = f'{desc}  -  {qty} {unit} x {days} Days'
            elif qty:
                bullet_text = f'{desc}  -  {qty} {unit}'
            else:
                bullet_text = desc
            
            bullet_para = doc.add_paragraph(bullet_text, style='List Bullet')
            for run in bullet_para.runs:
                run.font.size = Pt(11)
                run.font.bold = True
        
        doc.add_paragraph()
        
        from datetime import datetime
        today = datetime.now()
        if today.month >= 4:
            fy_start = today.year
            fy_end = (today.year + 1) % 100
        else:
            fy_start = today.year - 1
            fy_end = today.year % 100
        financial_year = f"{fy_start}-{fy_end:02d}"
        
        footer_text = (f'The rates proposed in the estimate are as per SOR {financial_year} and Approved rates. L.S. Provision is made in the '
                      'estimate towards GST at 18%, QC amount at 1%, Labour Cess at 1% and NAC amount at 0.1% as per actual '
                      'and LS Provision Towards, unforeseen items & rounding off also proposed in the estimate.')
        footer_para = doc.add_paragraph(footer_text)
        for run in footer_para.runs:
            run.font.size = Pt(10)
        
        doc.add_paragraph()
        
        funds_para = doc.add_paragraph()
        funds_run = funds_para.add_run('FUNDS: ')
        funds_run.font.bold = True
        funds_run.font.underline = True
        funds_para.add_run('The estimate requires Administrative sanction and also fixes up the agency with provision of funds '
                          'under relevant head of account for taking up the work from the Government, Telangana State Hyderabad')
        
        filename = 'Temp_Spec_Report.docx'
        
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        doc.save(response)
        return response
        
    except Exception as e:
        logger.error(f'Error generating temp works specification report: {str(e)}', exc_info=True)
        from django.contrib import messages
        messages.error(request, f'Error generating specification report: {str(e)}')
        return redirect('temp_groups', category=category)


@login_required(login_url='login')
def temp_download_forwarding_letter(request, category):
    """
    Generate forwarding letter from live Temporary Works items.
    """
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    
    if request.method != 'POST':
        return redirect('temp_groups', category=category)
    
    try:
        items_json = request.POST.get('items', '[]')
        work_name = request.POST.get('work_name', '{{NAME_OF_WORK}}')
        total_amount = request.POST.get('total_amount', '0.00')
        
        items = json.loads(items_json)
        
        if not items:
            from django.contrib import messages
            messages.error(request, 'No items with quantities to generate forwarding letter')
            return redirect('temp_groups', category=category)
        
        try:
            grand_total = float(total_amount.replace(',', '').replace('Rs.', '').replace('₹', '').strip())
        except:
            grand_total = 0.0
        
        financial_year = _get_current_financial_year()
        today = timezone.now().date()
        
        # Get user's letter settings
        letter_settings = _get_letter_settings(request.user)
        
        doc = Document()
        placeholder_color = RGBColor(169, 169, 169)
        
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Header - Government/Organization name
        header1 = doc.add_paragraph()
        header1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if letter_settings and letter_settings.government_name:
            run1 = header1.add_run(letter_settings.government_name.upper())
            run1.font.bold = True
            run1.font.size = Pt(14)
        else:
            run1 = header1.add_run('[GOVERNMENT / ORGANIZATION NAME]')
            run1.font.bold = True
            run1.font.size = Pt(14)
            run1.font.color.rgb = placeholder_color
            run1.font.italic = True
        
        # Header - Department name
        header2 = doc.add_paragraph()
        header2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if letter_settings and letter_settings.department_name:
            run2 = header2.add_run(letter_settings.department_name.upper())
            run2.font.bold = True
            run2.font.size = Pt(13)
        else:
            run2 = header2.add_run('[DEPARTMENT NAME]')
            run2.font.bold = True
            run2.font.size = Pt(13)
            run2.font.color.rgb = placeholder_color
            run2.font.italic = True
        
        doc.add_paragraph()
        
        from_to_table = doc.add_table(rows=1, cols=2)
        from_to_table.autofit = True
        
        from_cell = from_to_table.cell(0, 0)
        from_para = from_cell.paragraphs[0]
        from_para.add_run('From: -\n')
        
        # From section - Officer details
        if letter_settings and letter_settings.officer_name:
            name_qual = letter_settings.officer_name
            if letter_settings.officer_qualification:
                name_qual += f", {letter_settings.officer_qualification}"
            from_run1 = from_para.add_run(f'{name_qual},\n')
        else:
            from_run1 = from_para.add_run('[Officer Name, Qualification],\n')
            from_run1.font.color.rgb = placeholder_color
            from_run1.font.italic = True
        
        if letter_settings and letter_settings.officer_designation:
            from_run2 = from_para.add_run(f'{letter_settings.officer_designation},\n')
        else:
            from_run2 = from_para.add_run('[Designation],\n')
            from_run2.font.color.rgb = placeholder_color
            from_run2.font.italic = True
        
        if letter_settings and (letter_settings.sub_division or letter_settings.office_address):
            sub_addr = letter_settings.sub_division
            if letter_settings.office_address:
                sub_addr += f", {letter_settings.office_address}" if sub_addr else letter_settings.office_address
            from_run3 = from_para.add_run(f'{sub_addr}.')
        else:
            from_run3 = from_para.add_run('[Sub Division, Office Address].')
            from_run3.font.color.rgb = placeholder_color
            from_run3.font.italic = True
        
        to_cell = from_to_table.cell(0, 1)
        to_para = to_cell.paragraphs[0]
        to_para.add_run('To,\n')
        
        if letter_settings and letter_settings.recipient_designation:
            to_run1 = to_para.add_run(f'{letter_settings.recipient_designation},\n')
        else:
            to_run1 = to_para.add_run('[Officer Designation],\n')
            to_run1.font.color.rgb = placeholder_color
            to_run1.font.italic = True
        
        if letter_settings and letter_settings.recipient_division:
            to_run2 = to_para.add_run(f'{letter_settings.recipient_division},\n')
        else:
            to_run2 = to_para.add_run('[Division Name],\n')
            to_run2.font.color.rgb = placeholder_color
            to_run2.font.italic = True
        
        if letter_settings and letter_settings.recipient_address:
            to_run3 = to_para.add_run(f'{letter_settings.recipient_address}.')
        else:
            to_run3 = to_para.add_run('[Address].')
            to_run3.font.color.rgb = placeholder_color
            to_run3.font.italic = True
        
        doc.add_paragraph()
        
        lr_para = doc.add_paragraph()
        lr_para.add_run('Lr No. ')
        if letter_settings and letter_settings.office_code:
            lr_code = lr_para.add_run(letter_settings.office_code)
            lr_code.font.underline = True
        else:
            lr_placeholder = lr_para.add_run('[Office Code]')
            lr_placeholder.font.color.rgb = placeholder_color
            lr_placeholder.font.italic = True
            lr_placeholder.font.underline = True
        lr_para.add_run(f'/{financial_year}/          ')
        lr_para.add_run(f'\t\t\t\t\tDate:-    - {today.strftime("%m")} - {today.year}.')
        
        doc.add_paragraph()
        
        sir_para = doc.add_paragraph()
        sir_para.add_run('Sir,')
        
        doc.add_paragraph()
        
        subject_para = doc.add_paragraph()
        subj_run = subject_para.add_run('Sub:-')
        subj_run.font.underline = True
        subject_para.add_run('\t')
        subj_work = subject_para.add_run(f'{work_name} ')
        subj_work.font.bold = True
        subject_para.add_run(f'for the year {financial_year}.  -  Submission  -  Request for obtaining administrative sanction  -  Regarding.')
        
        doc.add_paragraph()
        
        ref_para = doc.add_paragraph()
        ref_run = ref_para.add_run('Ref:-')
        ref_run.font.underline = True
        ref_para.add_run('\tMemo No.')
        ref_placeholder = ref_para.add_run('[Reference Number]')
        ref_placeholder.font.color.rgb = placeholder_color
        ref_placeholder.font.italic = True
        ref_placeholder.font.underline = True
        ref_para.add_run(f'/{financial_year} Dt.')
        ref_date_placeholder = ref_para.add_run('[DD.MM.YYYY]')
        ref_date_placeholder.font.color.rgb = placeholder_color
        ref_date_placeholder.font.italic = True
        ref_date_placeholder.font.underline = True
        
        doc.add_paragraph()
        
        stars_para = doc.add_paragraph()
        stars_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        stars_para.add_run('**.**')
        
        doc.add_paragraph()
        
        body_para = doc.add_paragraph()
        body_para.add_run('With reference to the subject cited, I submit here ')
        with_run = body_para.add_run('with  1')
        with_run.font.underline = True
        body_para.add_run(' No. estimate for the following work for the amount specified.')
        
        doc.add_paragraph()
        
        table = doc.add_table(rows=2, cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for cell in table.columns[0].cells:
            cell.width = Inches(0.5)
        for cell in table.columns[1].cells:
            cell.width = Inches(4.5)
        for cell in table.columns[2].cells:
            cell.width = Inches(1.5)
        
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Sl.\nNo'
        header_cells[1].text = 'Name of work'
        header_cells[2].text = 'Amount'
        
        for cell in header_cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.bold = True
        
        row_cells = table.rows[1].cells
        row_cells[0].text = '1'
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[1].text = work_name
        
        formatted_amount = _format_indian_number(grand_total)
        row_cells[2].text = f"Rs.{formatted_amount}"
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()
        
        spec_para = doc.add_paragraph()
        spec_para.add_run("Specification report accompanying the estimate explains the necessity and provisions made therein in detail.")
        
        doc.add_paragraph()
        
        request_para = doc.add_paragraph()
        request_para.add_run('I request the ')
        if letter_settings and letter_settings.superior_designation:
            req_run = request_para.add_run(letter_settings.superior_designation)
        else:
            req_placeholder = request_para.add_run('[Superior Officer Designation]')
            req_placeholder.font.color.rgb = placeholder_color
            req_placeholder.font.italic = True
        request_para.add_run(' to kindly arrange to obtain administrative sanction for the above estimate and arrange to finalize the agency at the earliest for taking up the work.')
        
        doc.add_paragraph()
        
        enc_para = doc.add_paragraph()
        enc_para.add_run('Enclosure: -')
        doc.add_paragraph('Estimate  - 1 No.')
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        sign_para = doc.add_paragraph()
        sign_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sign_para.add_run('Yours faithfully,')
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        if letter_settings and letter_settings.officer_designation:
            run_title = title_para.add_run(f'{letter_settings.officer_designation}\n')
            run_title.font.bold = True
        else:
            run_title = title_para.add_run('[Officer Designation]\n')
            run_title.font.bold = True
            run_title.font.color.rgb = placeholder_color
            run_title.font.italic = True
        
        if letter_settings and letter_settings.sub_division:
            sub_div_run = title_para.add_run(f'{letter_settings.sub_division},\n')
        else:
            sub_div_run = title_para.add_run('[Sub Division Name],\n')
            sub_div_run.font.color.rgb = placeholder_color
            sub_div_run.font.italic = True
        
        if letter_settings and letter_settings.office_address:
            addr_run = title_para.add_run(f'{letter_settings.office_address}.')
        else:
            addr_run = title_para.add_run('[Office Address].')
            addr_run.font.color.rgb = placeholder_color
            addr_run.font.italic = True
        
        doc.add_paragraph()
        
        copy_para = doc.add_paragraph()
        copy_para.add_run('Copy to the ')
        if letter_settings and (letter_settings.copy_to_designation or letter_settings.copy_to_section):
            copy_text = letter_settings.copy_to_designation or ''
            if letter_settings.copy_to_section:
                copy_text += f", {letter_settings.copy_to_section}" if copy_text else letter_settings.copy_to_section
            copy_run = copy_para.add_run(copy_text)
        else:
            copy_placeholder = copy_para.add_run('[Officer Designation, Section Name]')
            copy_placeholder.font.color.rgb = placeholder_color
            copy_placeholder.font.italic = True
        copy_para.add_run(' for information.')
        
        filename = 'Temp_Fwd_Letter.docx'
        
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        doc.save(response)
        return response
        
    except Exception as e:
        logger.error(f'Error generating temp works forwarding letter: {str(e)}', exc_info=True)
        from django.contrib import messages
        messages.error(request, f'Error generating forwarding letter: {str(e)}')
        return redirect('temp_groups', category=category)




@login_required(login_url='login')
@require_POST
def temp_ajax_save_group_order(request, category):
    """Save the user's custom group order for the temporary works scope."""
    try:
        data = json.loads(request.body.decode('utf-8') or '{}')
    except ValueError:
        return JsonResponse({'ok': False, 'error': 'invalid json'}, status=400)
    order = data.get('order') or []
    if not isinstance(order, list):
        return JsonResponse({'ok': False, 'error': 'order must be a list'}, status=400)
    save_group_order(request.user, 'temp', category, order)
    return JsonResponse({'ok': True})
