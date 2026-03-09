# Auto-generated from core/views.py split
import json
import os
import re
import logging
from copy import copy

import inflect
from docx import Document
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, Border, Side, PatternFill
from django.utils import timezone
from django.urls import reverse
from django.core.files.uploadedfile import InMemoryUploadedFile
from django.contrib.auth.decorators import login_required

from django.conf import settings
from django.http import HttpResponse, JsonResponse, HttpResponseNotAllowed
from django.shortcuts import render, redirect, get_object_or_404
from django.template.loader import render_to_string
from django.views.decorators.http import require_POST
from django.utils.crypto import get_random_string

import io
from io import BytesIO
from difflib import SequenceMatcher

from ..models import Project, SelfFormattedTemplate, Estimate, Organization, Membership, Upload, Job, OutputFile, LetterSettings
from ..decorators import org_required, role_required

logger = logging.getLogger(__name__)
from ..tasks import process_excel_upload, generate_bill_pdf, generate_workslip_pdf, generate_bill_document_task
from ..utils_excel import load_backend, copy_block_with_styles_and_formulas, build_temp_day_rates

p_engine = inflect.engine()
BILL_TEMPLATES_DIR = os.path.join(settings.BASE_DIR, "core", "templates", "core", "bill_templates")
_inflect_engine = inflect.engine()

from .utils import (get_org_from_request, _apply_print_settings,
    _format_indian_number, _number_to_words_rupees,
    _get_current_financial_year, _get_current_date_formatted,
    _get_letter_settings, _format_date_to_ddmmyyyy)

@login_required(login_url='login')
def estimate(request):
    """
    Estimate module: Upload item blocks sheet (in backend format)
    and generate estimate in the standard format with Output and Estimate sheets.
    
    The uploaded file should contain item blocks in the same format as:
    - core/data/electrical_backend.xlsx
    - core/data/civil_backend.xlsx
    
    Items are detected by yellow background + red text cells in the header row.
    """
    if request.method == 'GET':
        return render(request, 'core/estimate.html', {})
    
    if request.method == 'POST':
        output_file = request.FILES.get('output_file')
        
        if not output_file:
            return render(request, 'core/estimate.html', {
                'error': 'Please upload an Excel file with item blocks.'
            })
        
        # Verify file size and content
        if output_file.size == 0:
            return render(request, 'core/estimate.html', {
                'error': 'Uploaded file is empty. Please select a valid Excel file.'
            })
        
        try:
            # Load the uploaded workbook
            # Reset file position before loading to fix first-attempt upload issue
            try:
                output_file.seek(0)
            except Exception as seek_error:
                pass  # Some file-like objects don't support seek
            
            try:
                wb_upload = load_workbook(output_file, data_only=False)
            except Exception as load_error:
                return render(request, 'core/estimate.html', {
                    'error': f'Failed to read Excel file: {str(load_error)}'
                })
            
            try:
                output_file.seek(0)  # Reset again for second load
            except Exception as seek_error:
                pass
            
            try:
                wb_upload_vals = load_workbook(output_file, data_only=True)
            except Exception as load_error:
                return render(request, 'core/estimate.html', {
                    'error': f'Failed to process Excel file: {str(load_error)}'
                })
            
            # ---- Read units from Groups sheet if available ----
            upload_units_map = {}
            try:
                if "Groups" in wb_upload.sheetnames:
                    from core.utils_excel import read_groups
                    _, upload_units_map = read_groups(wb_upload["Groups"])
            except Exception:
                pass  # If Groups sheet is missing or malformed, fall back to heuristic
            
            # ---- Helper: Check if cell is yellow with red text ----
            def cell_is_yellow(cell):
                fill = cell.fill
                if not fill or not fill.patternType or fill.patternType.lower() != "solid":
                    return False
                rgb = getattr(fill.fgColor, "rgb", None)
                if rgb and str(rgb).upper().endswith("FFFF00"):
                    return True
                if getattr(fill.fgColor, "type", None) == "theme":
                    if getattr(fill.fgColor, "theme", None) in (4, 5, 6):
                        return True
                if getattr(fill.fgColor, "indexed", None) == 6:
                    return True
                return False
            
            def cell_is_red_text(cell):
                font = cell.font
                if not font or not font.color:
                    return False
                rgb = getattr(font.color, "rgb", None)
                if rgb and str(rgb).upper().endswith("FF0000"):
                    return True
                if getattr(font.color, "type", None) == "theme":
                    return True
                if getattr(font.color, "indexed", None) == 3:
                    return True
                return False
            
            def is_yellow_and_red(cell):
                return cell_is_yellow(cell) and cell_is_red_text(cell)
            
            def is_valid_item_block(ws_src, start_row, end_row):
                """
                Check if this block looks like a valid item block (has rate data in column J).
                This helps distinguish real item blocks from headings or signature sections.
                """
                # A valid item block should have at least one non-empty value in column J (rate column)
                for r in range(start_row, min(end_row + 1, start_row + 50)):  # Check up to 50 rows
                    val = ws_src.cell(row=r, column=10).value
                    if val not in (None, "") and str(val).strip():
                        return True
                return False
            
            def find_item_block_end(ws_src, start_row, max_row):
                """
                Find the true end of an item block by looking for the rate row in column J.
                The block ends at the last row that has meaningful data before the next heading
                or signature section.
                """
                # First, find where the next yellow+red heading is
                next_heading_row = max_row + 1
                for rr in range(start_row + 1, max_row + 1):
                    for c in range(1, 11):
                        cell = ws_src.cell(row=rr, column=c)
                        if is_yellow_and_red(cell) and str(cell.value or "").strip():
                            next_heading_row = rr
                            break
                    if next_heading_row <= max_row:
                        break
                
                # The block should end before the next heading
                potential_end = next_heading_row - 1
                
                # Find the last row with rate data (column J) - this is the true end of item block
                last_rate_row = start_row
                for r in range(start_row, potential_end + 1):
                    val = ws_src.cell(row=r, column=10).value
                    if val not in (None, "") and str(val).strip():
                        last_rate_row = r
                
                # The block ends at the last rate row (don't include signature/footer content)
                return last_rate_row, next_heading_row
            
            def extract_items_from_sheet(ws_src):
                """Extract all item blocks from a single sheet."""
                fetched_items = []
                item_blocks = {}  # name -> (start_row, end_row)
                
                max_row = ws_src.max_row
                r = 1
                first_item_found = False
                
                while r <= max_row:
                    heading_name = None
                    heading_col = None
                    # Check columns A..J for yellow+red cell
                    for c in range(1, 11):
                        cell = ws_src.cell(row=r, column=c)
                        if is_yellow_and_red(cell) and str(cell.value or "").strip():
                            heading_name = str(cell.value).strip()
                            heading_col = c
                            break
                    
                    if heading_name:
                        start_row = r
                        
                        # Find the proper end of this block
                        end_row, next_heading_row = find_item_block_end(ws_src, start_row, max_row)
                        
                        # Validate this is a real item block (not a sheet heading or signature section)
                        if is_valid_item_block(ws_src, start_row, end_row):
                            fetched_items.append(heading_name)
                            item_blocks[heading_name] = (start_row, end_row)
                            first_item_found = True
                        
                        # Move to next heading position
                        r = next_heading_row if next_heading_row <= max_row else end_row + 1
                    else:
                        r += 1
                
                return fetched_items, item_blocks
            
            def create_output_and_estimate_sheets(wb_out, ws_src, fetched_items, item_blocks, 
                                                   output_sheet_name, estimate_sheet_name):
                """Create Output and Estimate sheets for a single source sheet."""
                thin = Side(border_style="thin", color="000000")
                border_all = Border(top=thin, left=thin, right=thin, bottom=thin)
                
                # Create Output sheet
                ws_out = wb_out.create_sheet(output_sheet_name)
                
                # Build Output sheet by copying item blocks
                cursor = 1
                rate_pos = {}
                data_serial = 1
                
                for item_name in fetched_items:
                    src_min, src_max = item_blocks[item_name]
                    
                    # Find rate row (non-empty in column J)
                    rate_src_row = None
                    for r in range(src_max, src_min - 1, -1):
                        v = ws_src.cell(row=r, column=10).value
                        if v not in (None, ""):
                            rate_src_row = r
                            break
                    
                    dst_start = cursor
                    
                    # Copy block with styles
                    copy_block_with_styles_and_formulas(
                        ws_src=ws_src,
                        ws_dst=ws_out,
                        src_min_row=src_min,
                        src_max_row=src_max,
                        dst_start_row=dst_start,
                        col_start=1,
                        col_end=10
                    )
                    
                    ws_out.cell(row=dst_start, column=1).value = f"Data {data_serial}"
                    data_serial += 1
                    
                    # Store rate row position
                    if rate_src_row:
                        rate_pos[item_name] = dst_start + (rate_src_row - src_min)
                    
                    cursor += (src_max - src_min + 1)
                
                # Create Estimate sheet
                ws_est = wb_out.create_sheet(estimate_sheet_name)
                
                # Title row
                ws_est.merge_cells("A1:H1")
                c1 = ws_est["A1"]
                c1.value = "ESTIMATE"
                c1.font = Font(bold=True, size=14)
                c1.alignment = Alignment(horizontal="center", vertical="center")
                
                # Title row 2
                ws_est.merge_cells("A2:H2")
                c2 = ws_est["A2"]
                c2.value = "Name of the work : "
                c2.font = Font(bold=True, size=11)
                c2.alignment = Alignment(horizontal="left", vertical="center")
                
                for row in (1, 2):
                    for col in range(1, 9):
                        ws_est.cell(row=row, column=col).border = border_all
                
                # Header row
                headers = ["Sl.No", "Quantity (Unit)", "", "Item Description",
                           "Rate", "Per Unit", "", "Amount"]
                
                for col, text in enumerate(headers, start=1):
                    cell = ws_est.cell(row=3, column=col, value=text)
                    cell.font = Font(bold=True)
                    cell.alignment = Alignment(horizontal="center", vertical="center")
                    cell.border = border_all
                
                ws_est.merge_cells("B3:C3")
                ws_est.merge_cells("F3:G3")
                
                ws_est.column_dimensions["A"].width = 6
                ws_est.column_dimensions["B"].width = 10
                ws_est.column_dimensions["C"].width = 10
                ws_est.column_dimensions["D"].width = 38
                ws_est.column_dimensions["E"].width = 10
                ws_est.column_dimensions["F"].width = 8
                ws_est.column_dimensions["G"].width = 10
                ws_est.column_dimensions["H"].width = 15
                
                # Fill estimate rows
                row_est = 4
                slno = 1
                
                def to_plural(unit):
                    """Convert unit to plural form."""
                    unit_lower = unit.lower()
                    if unit_lower == "no":
                        return "Nos"
                    elif unit_lower == "nos":
                        return "Nos"
                    elif unit_lower == "mtr":
                        return "Mtrs"
                    elif unit_lower == "mtrs":
                        return "Mtrs"
                    elif unit_lower == "pts":
                        return "Pts"
                    elif unit_lower == "pt":
                        return "Pts"
                    elif unit_lower == "cum":
                        return "Cum"
                    elif unit_lower == "kg":
                        return "Kg"
                    elif unit_lower == "l":
                        return "L"
                    elif unit_lower == "kg":
                        return "Kg"
                    else:
                        return unit + "s" if unit else "Nos"
                
                def to_singular(unit):
                    """Convert unit to singular form."""
                    unit_lower = unit.lower()
                    if unit_lower == "nos":
                        return "No"
                    elif unit_lower == "no":
                        return "No"
                    elif unit_lower == "mtrs":
                        return "Mtr"
                    elif unit_lower == "mtr":
                        return "Mtr"
                    elif unit_lower == "pts":
                        return "Pt"
                    elif unit_lower == "pt":
                        return "Pt"
                    elif unit_lower == "cum":
                        return "Cum"
                    elif unit_lower == "kg":
                        return "Kg"
                    elif unit_lower == "l":
                        return "L"
                    else:
                        return unit
                
                def determine_unit_from_heading(heading_name):
                    """
                    Determine unit from Groups sheet units_map first,
                    then fall back to heading name heuristic.
                    """
                    # Priority 1: Use unit from Groups sheet (authoritative)
                    if upload_units_map and heading_name in upload_units_map:
                        return upload_units_map[heading_name]
                    
                    heading_lower = heading_name.lower()
                    
                    # Light Point or Fan Point â†’ Pts
                    if "light point" in heading_lower or "fan point" in heading_lower:
                        return "Pts"
                    
                    # Light or Fan (fixtures/bulbs) â†’ Nos (check BEFORE pipe keywords to avoid "tube light" being "Mtr")
                    light_fan_keywords = ["light", "fan", "bulb", "fixture", "downlight", "spotlight", "batten"]
                    for keyword in light_fan_keywords:
                        if keyword in heading_lower:
                            return "Nos"
                    
                    # Pipes, wires, cables â†’ Mtr (meters)
                    pipe_keywords = ["pipe", "wire", "cable", "conduit", "duct", "channel", "rod", "bar", "rail", "tube"]
                    for keyword in pipe_keywords:
                        if keyword in heading_lower:
                            return "Mtr"
                    
                    # Points â†’ Pts
                    if "point" in heading_lower or "pts" in heading_lower:
                        return "Pts"
                    
                    # Default to Nos
                    return "Nos"
                
                for item_name in fetched_items:
                    src_min, src_max = item_blocks[item_name]
                    
                    # Get the description from the second row (src_min + 2)
                    base_desc = ws_src.cell(row=src_min + 2, column=4).value or ""
                    base_desc_str = str(base_desc).strip()
                    desc = base_desc_str

                    # Determine unit intelligently from heading name
                    best_unit = determine_unit_from_heading(item_name)
                    unit_plural = to_plural(best_unit)
                    unit_singular = to_singular(best_unit)

                    # Rate from Output sheet (reference the correct output sheet)
                    rr = rate_pos.get(item_name)
                    # Excel sheet names with spaces need quotes
                    safe_output_name = f"'{output_sheet_name}'" if ' ' in output_sheet_name else output_sheet_name
                    rate_formula = f"={safe_output_name}!J{rr}" if rr else ""

                    # Write row
                    a = ws_est.cell(row=row_est, column=1, value=slno)
                    a.alignment = Alignment(horizontal="center", vertical="center")
                    a.border = border_all

                    b = ws_est.cell(row=row_est, column=2, value="")
                    b.alignment = Alignment(horizontal="center", vertical="center")
                    b.border = border_all

                    c = ws_est.cell(row=row_est, column=3, value=unit_plural)
                    c.alignment = Alignment(horizontal="center", vertical="center")
                    c.border = border_all

                    d = ws_est.cell(row=row_est, column=4, value=desc)
                    d.alignment = Alignment(horizontal="justify", vertical="top", wrap_text=True)
                    d.border = border_all

                    e = ws_est.cell(row=row_est, column=5, value=rate_formula)
                    e.alignment = Alignment(horizontal="center", vertical="center")
                    e.border = border_all

                    f = ws_est.cell(row=row_est, column=6, value=1)
                    f.alignment = Alignment(horizontal="center", vertical="center")
                    f.border = border_all

                    g = ws_est.cell(row=row_est, column=7, value=unit_singular)
                    g.alignment = Alignment(horizontal="center", vertical="center")
                    g.border = border_all

                    h = ws_est.cell(row=row_est, column=8, value=f"=B{row_est}*E{row_est}")
                    h.alignment = Alignment(horizontal="center", vertical="center")
                    h.border = border_all

                    row_est += 1
                    slno += 1
                
                # ---- Add totals rows ----
                ecv_row = row_est
                ws_est.cell(row=ecv_row, column=4, value="ECV")
                ws_est.cell(row=ecv_row, column=8, value=f"=SUM(H4:H{ecv_row-1})")
                
                lc_row = ecv_row + 1
                qc_row = ecv_row + 2
                nac_row = ecv_row + 3
                sub_row = ecv_row + 4
                gst_row = ecv_row + 5
                ls_row = ecv_row + 6
                gt_row = ecv_row + 7
                
                ws_est.cell(row=lc_row, column=4, value="Add LC @ 1 %")
                ws_est.cell(row=lc_row, column=8, value=f"=H{ecv_row}*0.01")
                
                ws_est.cell(row=qc_row, column=4, value="Add QC @ 1 %")
                ws_est.cell(row=qc_row, column=8, value=f"=H{ecv_row}*0.01")
                
                ws_est.cell(row=nac_row, column=4, value="Add NAC @ 0.1 %")
                ws_est.cell(row=nac_row, column=8, value=f"=H{ecv_row}*0.001")
                
                ws_est.cell(row=sub_row, column=4, value="Sub Total")
                ws_est.cell(row=sub_row, column=8, value=f"=H{ecv_row}+H{lc_row}+H{qc_row}+H{nac_row}")
                
                ws_est.cell(row=gst_row, column=4, value="Add GST@18 %")
                ws_est.cell(row=gst_row, column=8, value=f"=H{sub_row}*0.18")
                
                ws_est.cell(row=ls_row, column=4, value="L.S Provision towards unforeseen items")
                ws_est.cell(row=ls_row, column=8, value=f"=H{gt_row}-H{gst_row}-H{sub_row}")
                
                ws_est.cell(row=gt_row, column=4, value="Grand Total")
                
                # Apply borders to totals
                for r in range(ecv_row, gt_row + 1):
                    for c in range(1, 9):
                        cell = ws_est.cell(row=r, column=c)
                        cell.border = border_all
                        if c == 4:
                            cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
                        else:
                            cell.alignment = Alignment(horizontal="center", vertical="center")
                    ws_est.cell(row=r, column=4).font = Font(bold=True)
                    ws_est.cell(row=r, column=8).font = Font(bold=True)
                
                return ws_out, ws_est
            
            # ---- Multi-sheet processing ----
            # Find all sheets with item blocks
            sheets_with_items = []
            for sheet_name in wb_upload.sheetnames:
                ws_src = wb_upload[sheet_name]
                fetched_items, item_blocks = extract_items_from_sheet(ws_src)
                if fetched_items:
                    sheets_with_items.append({
                        'name': sheet_name,
                        'ws_src': ws_src,
                        'fetched_items': fetched_items,
                        'item_blocks': item_blocks
                    })
            
            if not sheets_with_items:
                return render(request, 'core/estimate.html', {
                    'error': 'No item blocks found in any sheet. Make sure item headers have yellow background and red text.'
                })
            
            # Create output workbook
            wb_out = Workbook()
            # Remove the default sheet, we'll create our own
            default_sheet = wb_out.active
            
            # Generate Output and Estimate sheets for each source sheet with items
            total_sheets = len(sheets_with_items)
            
            for idx, sheet_info in enumerate(sheets_with_items):
                src_sheet_name = sheet_info['name']
                ws_src = sheet_info['ws_src']
                fetched_items = sheet_info['fetched_items']
                item_blocks = sheet_info['item_blocks']
                
                # Determine sheet names
                if total_sheets == 1:
                    # Single sheet: use simple names
                    output_sheet_name = "Datas"
                    estimate_sheet_name = "Estimate"
                else:
                    # Multiple sheets: append source sheet name
                    # Truncate to fit Excel's 31 character limit
                    base_name = src_sheet_name[:20] if len(src_sheet_name) > 20 else src_sheet_name
                    output_sheet_name = f"Datas_{base_name}"[:31]
                    estimate_sheet_name = f"Estimate_{base_name}"[:31]
                
                create_output_and_estimate_sheets(
                    wb_out=wb_out,
                    ws_src=ws_src,
                    fetched_items=fetched_items,
                    item_blocks=item_blocks,
                    output_sheet_name=output_sheet_name,
                    estimate_sheet_name=estimate_sheet_name
                )
            
            # Remove the default empty sheet if we created our own
            if default_sheet.title in wb_out.sheetnames and len(wb_out.sheetnames) > 1:
                wb_out.remove(default_sheet)
            
            # Reorder sheets: Estimate followed by its corresponding Datas sheet
            # Pattern: Estimate_1, Datas_1, Estimate_2, Datas_2, etc.
            estimate_sheets = [s for s in wb_out.sheetnames if s.startswith("Estimate") or s == "Estimate"]
            output_sheets = [s for s in wb_out.sheetnames if s.startswith("Datas") or s == "Datas"]
            
            # Build pairs based on suffix matching
            ordered_sheets = []
            for est_name in estimate_sheets:
                ordered_sheets.append(est_name)
                # Find matching Datas sheet
                if est_name == "Estimate":
                    # Single sheet case
                    if "Datas" in output_sheets:
                        ordered_sheets.append("Datas")
                else:
                    # Multi-sheet case: Estimate_XYZ -> Datas_XYZ
                    suffix = est_name[8:]  # Remove "Estimate" prefix (8 chars)
                    output_name = f"Datas{suffix}"
                    if output_name in output_sheets:
                        ordered_sheets.append(output_name)
            
            # Add any remaining output sheets that weren't paired
            for out_name in output_sheets:
                if out_name not in ordered_sheets:
                    ordered_sheets.append(out_name)
            
            # Reorder sheets according to the new order
            for i, sheet_name in enumerate(ordered_sheets):
                if sheet_name in wb_out.sheetnames:
                    current_idx = wb_out.sheetnames.index(sheet_name)
                    if current_idx != i:
                        wb_out.move_sheet(sheet_name, offset=(i - current_idx))
            
            # Return the estimate workbook
            _apply_print_settings(wb_out)
            response = HttpResponse(
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = 'attachment; filename="estimate.xlsx"'
            wb_out.save(response)
            return response
            
        except Exception as e:
            import traceback
            return render(request, 'core/estimate.html', {
                'error': f'Error processing file: {str(e)}'
            })


@org_required
@login_required
def download_specification_report(request, estimate_id):
    """Generate and download specification report as Word document"""
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    org = request.organization
    estimate = get_object_or_404(Estimate, id=estimate_id, organization=org)
    
    try:
        # Create Word document
        doc = Document()
        
        # Title
        title = doc.add_heading('Specification report accompanying the estimate :-', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in title.runs:
            run.font.size = Pt(12)
            run.font.bold = True
        
        # Introduction paragraph with work name
        estimate_data = estimate.estimate_data or {}
        work_name = estimate_data.get('name_of_work', 'the work')
        
        intro_text = f'The estimate is prepared for the work {work_name}'
        intro_para = doc.add_paragraph(intro_text)
        intro_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in intro_para.runs:
            run.font.size = Pt(11)
        
        # Estimate amount
        total_amount = estimate.total_amount or 0
        amount_para = doc.add_paragraph()
        amount_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        amount_run = amount_para.add_run(f'Est.Amount: Rs. {total_amount:,.2f}')
        amount_run.font.size = Pt(11)
        amount_run.font.bold = True
        amount_run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
        
        doc.add_paragraph()  # Blank line
        
        # Body of letter (manually entered, placeholder for user to edit)
        body_label = doc.add_paragraph('{{BODY_OF_LETTER}}')
        body_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in body_label.runs:
            run.font.size = Pt(11)
            run.font.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)  # Gray placeholder
        
        doc.add_paragraph()  # Blank line
        
        doc.add_paragraph('Hence, this estimate has been prepared accordingly.')
        
        doc.add_paragraph()  # Blank line
        
        # Extract items with quantities and units as bullet points
        # Support both 'items' key and 'ws_estimate_rows' key for different estimate formats
        estimate_items = estimate_data.get('ws_estimate_rows', estimate_data.get('items', []))
        
        for item in estimate_items:
            # Extract values from item structure
            # Support multiple key names: desc/description, qty_est/qty/quantity
            item_description = item.get('desc', item.get('description', item.get('display_name', '')))
            quantity = item.get('qty_est', item.get('qty', item.get('quantity', '')))
            unit = item.get('unit', '')
            
            # Format: "Description  -  Quantity Unit"
            if quantity and unit:
                # Convert quantity to clean format
                qty_str = str(quantity).strip()
                if '.' in qty_str and qty_str.endswith('.0'):
                    qty_str = qty_str.replace('.0', '')
                bullet_text = f'{item_description}  -  {qty_str} {unit}'
            elif quantity:
                qty_str = str(quantity).strip()
                if '.' in qty_str and qty_str.endswith('.0'):
                    qty_str = qty_str.replace('.0', '')
                bullet_text = f'{item_description}  -  {qty_str}'
            else:
                bullet_text = item_description
            
            bullet_para = doc.add_paragraph(bullet_text, style='List Bullet')
            for run in bullet_para.runs:
                run.font.size = Pt(11)
        
        doc.add_paragraph()  # Blank line
        
        # Calculate financial year (April to March)
        from datetime import datetime
        today = datetime.now()
        if today.month >= 4:  # April onwards
            fy_start = today.year
            fy_end = (today.year + 1) % 100
        else:  # Jan to March
            fy_start = today.year - 1
            fy_end = today.year % 100
        financial_year = f"{fy_start}-{fy_end:02d}"
        
        # Footer text about rates and provisions
        footer_text = (f'The rates proposed in the estimate are as per SQR {financial_year} and Approved rates. L.S. Provision is made in the '
                      'estimate towards GST at 18%, QC amount at 1%, Labour Cess at 1% and NAC amount at 0.1% as per actual and LS Provision Towards, unforeseen items & rounding off also proposed in the estimate.')
        footer_para = doc.add_paragraph(footer_text)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in footer_para.runs:
            run.font.size = Pt(10)
        
        doc.add_paragraph()  # Blank line
        
        # Funds section
        funds_text = ('FUNDS: The estimate requires Administrative sanction and also fixes up the agency with provision of funds '
                     'under relevant head of account for taking up the work from the Government. Telangana State Hyderabad')
        funds_para = doc.add_paragraph(funds_text)
        funds_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in funds_para.runs:
            run.font.size = Pt(10)
            run.font.bold = True
        
        # Generate filename
        filename = 'Spec_Report.docx'
        
        # Save to response
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        doc.save(response)
        return response
        
    except Exception as e:
        logger.error(f'Error generating specification report: {str(e)}', exc_info=True)
        from django.contrib import messages
        messages.error(request, f'Error generating report: {str(e)}')
        return redirect('view_estimate', estimate_id=estimate_id)


@login_required(login_url='login')
def generate_specification_report_from_file(request):
    """
    Generate specification report from uploaded Excel workbook.
    Extracts item headings (red text + yellow background) from Item Blocks sheet
    and matches quantities/units from the Estimate sheet.
    """
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    if request.method != 'POST':
        return render(request, 'core/estimate.html', {
            'error': 'Invalid request method'
        })
    
    output_file = request.FILES.get('output_file')
    if not output_file:
        return render(request, 'core/estimate.html', {
            'error': 'Please upload an Excel file.'
        })
    
    try:
        # Reset file position
        try:
            output_file.seek(0)
        except:
            pass
        
        wb_upload = load_workbook(output_file, data_only=False)
        
        # Also load with data_only=True for values
        output_file.seek(0)
        wb_values = load_workbook(output_file, data_only=True)
        
        # Helper functions for cell detection
        def cell_is_yellow(cell):
            fill = cell.fill
            if not fill or not fill.fgColor:
                return False
            fg = fill.fgColor
            # Check RGB values
            rgb = getattr(fg, "rgb", None)
            if rgb:
                rgb_upper = str(rgb).upper()
                if rgb_upper in ('FFFFFF00', 'FFFF00', 'FFFFE000', 'FFFFC000', 'FFFFCC00'):
                    return True
                # Check if it ends with yellow-ish color
                if rgb_upper.endswith('FFFF00') or (rgb_upper.endswith('FF00') and 'FF' in rgb_upper[:4]):
                    return True
            # Check indexed colors (Excel default palette)
            if getattr(fg, "indexed", None) in (6, 13):
                return True
            # Check theme colors
            if getattr(fg, "type", None) == "theme":
                return True
            return False
        
        def cell_is_red_text(cell):
            font = cell.font
            if not font or not font.color:
                return False
            c = font.color
            rgb = getattr(c, "rgb", None)
            if rgb:
                rgb_upper = str(rgb).upper()
                if rgb_upper in ('FFFF0000', 'FF0000', 'FFC00000', 'FFFF0000'):
                    return True
                if rgb_upper.endswith('FF0000'):
                    return True
            # Check indexed red
            if getattr(c, "indexed", None) == 3:
                return True
            return False
        
        def is_yellow_and_red(cell):
            return cell_is_yellow(cell) and cell_is_red_text(cell)
        
        # Find sheets - look for "Estimate" and "Item Blocks" sheets
        estimate_sheet = None
        estimate_sheet_values = None
        item_blocks_sheet = None
        
        for sheet_name in wb_upload.sheetnames:
            sheet_lower = sheet_name.lower().strip()
            if 'estimate' in sheet_lower and 'item' not in sheet_lower and 'datas' not in sheet_lower:
                estimate_sheet = wb_upload[sheet_name]
                estimate_sheet_values = wb_values[sheet_name]
            elif 'item' in sheet_lower and 'block' in sheet_lower:
                item_blocks_sheet = wb_upload[sheet_name]
            elif 'datas' in sheet_lower:
                # "Datas" sheet - treat as item blocks sheet
                item_blocks_sheet = wb_upload[sheet_name]
        
        # Fallback: If no specific sheets found, look for sheet with yellow+red items
        if not item_blocks_sheet:
            # Try to find any sheet with yellow+red items
            for sheet_name in wb_upload.sheetnames:
                ws = wb_upload[sheet_name]
                for row in range(1, min(50, ws.max_row + 1)):
                    for col in range(1, 10):
                        cell = ws.cell(row=row, column=col)
                        if is_yellow_and_red(cell):
                            item_blocks_sheet = ws
                            break
                    if item_blocks_sheet:
                        break
                if item_blocks_sheet:
                    break
        
        if not item_blocks_sheet:
            return render(request, 'core/estimate.html', {
                'error': 'Could not find Datas sheet. Please ensure your workbook has a sheet containing items with red text and yellow background.'
            })
        
        # Step 1: Extract red+yellow items from Item Blocks sheet
        item_headings = []
        for row in range(1, item_blocks_sheet.max_row + 1):
            for col in range(1, 15):  # Check columns A through O
                cell = item_blocks_sheet.cell(row=row, column=col)
                if is_yellow_and_red(cell):
                    heading = str(cell.value or "").strip()
                    if not heading or len(heading) < 3:
                        continue
                    
                    # Skip if it looks like a label/header
                    heading_lower = heading.lower()
                    if heading_lower in ('sl', 'sl.', 'sl.no', 'description', 'unit', 'qty', 'rate', 'amount', 'total'):
                        continue
                    
                    # Avoid duplicates
                    if heading not in item_headings:
                        item_headings.append(heading)
                    break  # Move to next row after finding a heading
        
        if not item_headings:
            return render(request, 'core/estimate.html', {
                'error': 'No item headings found with red text and yellow background in the workbook.'
            })
        
        # Step 2: Build quantity/unit map from Estimate sheet
        qty_unit_map = {}  # heading -> {'qty': ..., 'unit': ...}
        work_name = ""
        total_amount = ""
        
        if estimate_sheet:
            # Find work name - look for "Name of the work" or "Name of work" 
            # The work name might be in the same cell after a colon, or in adjacent cells
            for row in range(1, min(30, estimate_sheet.max_row + 1)):
                if work_name:
                    break
                for col in range(1, 10):
                    cell_val = str(estimate_sheet.cell(row=row, column=col).value or "").strip()
                    cell_lower = cell_val.lower()
                    
                    # Check if this cell contains "name of the work" or "name of work"
                    if "name of the work" in cell_lower or "name of work" in cell_lower:
                        # Check if work name is in the same cell after colon
                        if ':' in cell_val:
                            parts = cell_val.split(':', 1)
                            if len(parts) > 1 and parts[1].strip():
                                work_name = parts[1].strip()
                                break
                        
                        # Try next columns in same row
                        if not work_name:
                            for next_col in range(col + 1, col + 8):
                                next_cell = estimate_sheet.cell(row=row, column=next_col).value
                                if next_cell and str(next_cell).strip():
                                    work_name = str(next_cell).strip()
                                    break
                        
                        # Try next row
                        if not work_name:
                            next_row_val = estimate_sheet.cell(row=row + 1, column=col).value
                            if next_row_val and str(next_row_val).strip():
                                work_name = str(next_row_val).strip()
                        break
            
            # Find Grand Total / Estimate Amount
            # Search for "Grand Total" row and get the amount from adjacent column
            for row in range(1, estimate_sheet.max_row + 1):
                for col in range(1, 10):
                    cell_val = str(estimate_sheet.cell(row=row, column=col).value or "").strip().lower()
                    if 'grand total' in cell_val:
                        # Look for amount in columns to the right (E, F, G, H)
                        for amt_col in range(col + 1, col + 6):
                            if estimate_sheet_values:
                                amt_cell = estimate_sheet_values.cell(row=row, column=amt_col)
                            else:
                                amt_cell = estimate_sheet.cell(row=row, column=amt_col)
                            if amt_cell.value is not None:
                                try:
                                    amt_val = float(amt_cell.value)
                                    if amt_val > 1000:  # Grand total should be a significant amount
                                        # Format as currency
                                        total_amount = f"{amt_val:,.2f}"
                                        break
                                except (ValueError, TypeError):
                                    pass
                        if total_amount:
                            break
                if total_amount:
                    break
            
            # Search Estimate sheet for each item heading to get qty and unit
            # Structure: Column A=Sl.No, Column B=Quantity, Column C=Unit, Column D=Description
            
            # Helper function to check if item heading matches description using keywords
            def heading_matches_description(heading, description):
                heading_lower = heading.lower().strip()
                desc_lower = description.lower().strip()
                
                # Skip if description is too short (likely a header row without real data)
                if len(desc_lower) < 15:
                    return False
                
                # Exact match
                if heading_lower == desc_lower:
                    return True
                
                # Heading appears exactly in description
                if heading_lower in desc_lower:
                    return True
                
                # Normalize description - remove extra spaces
                desc_normalized = ' '.join(desc_lower.split())
                
                # Helper to check if all words appear in description
                def all_words_in_desc(words):
                    return all(w in desc_normalized for w in words)
                
                # Helper to check if any word appears in description
                def any_word_in_desc(words):
                    return any(w in desc_normalized for w in words)
                
                # Specific matching rules for electrical items
                
                # PVC Pipe matching - must distinguish concealed vs surface
                if 'concealed' in heading_lower and 'pvc' in heading_lower:
                    return 'concealed' in desc_normalized and 'pvc' in desc_normalized and 'surface' not in desc_normalized
                
                if 'surface' in heading_lower and 'pvc' in heading_lower:
                    return 'surface' in desc_normalized and 'pvc' in desc_normalized
                
                # Light & Bell Points - row 8 type items
                if ('light' in heading_lower and 'bell' in heading_lower) or ('light' in heading_lower and 'point' in heading_lower):
                    # Match descriptions with "light and bell point" or "for light" 
                    if 'light and bell' in desc_normalized:
                        return True
                    if 'for light' in desc_normalized and 'bell' in desc_normalized and 'point' in desc_normalized:
                        return True
                    if 'light point' in desc_normalized:
                        return True
                    return False
                
                # Fan and Exhaust Fan Points - row 9 type items  
                if 'fan' in heading_lower and 'exhaust' in heading_lower and 'point' in heading_lower:
                    return ('exhaust' in desc_normalized and 'fan' in desc_normalized and 'point' in desc_normalized and 
                            'light and bell' not in desc_normalized)
                
                # Common Switch Board - row 11
                if 'switch' in heading_lower and 'board' in heading_lower:
                    # Check for "switch board" or "switch  board" with any spacing, or switchboard
                    if 'switch' in desc_normalized and 'board' in desc_normalized:
                        return True
                    if 'switchboard' in desc_normalized:
                        return True
                    return False
                
                # Power Plug - row 13, 14
                if 'power' in heading_lower and 'plug' in heading_lower:
                    # Power Plug items - look for 16A socket/plug descriptions
                    # "Two Nos" in heading is part of item name (2-gang type), not quantity
                    if 'two' in heading_lower:
                        # Power Plug Two Nos - has "2 Nos" or "2 Module" or "2 way" in description
                        return ('16a' in desc_normalized or '16 a' in desc_normalized) and \
                               ('2 nos' in desc_normalized or '2 no.s' in desc_normalized or 
                                '2 module' in desc_normalized or '2 way' in desc_normalized or 
                                'two nos' in desc_normalized or 'twin' in desc_normalized)
                    else:
                        # Regular single power plug - should NOT have "2 Nos" pattern
                        has_16a = '16a' in desc_normalized or '16 a' in desc_normalized
                        has_plug = 'socket' in desc_normalized or 'plug' in desc_normalized
                        is_two_type = '2 nos' in desc_normalized or '2 no.s' in desc_normalized or 'two nos' in desc_normalized
                        return has_16a and has_plug and not is_two_type
                
                # Module (sockets/switches) - rows 16, 17, 18
                # "Two Nos", "Three Nos", "Four Nos" in heading = type of module (number of socket+switch combos)
                # Descriptions mention: "2 no.s", "3 Nos", "4 Nos" and modular box sizes: 6, 8, 12 Modular
                if 'module' in heading_lower:
                    # Must have "modular" in description
                    if 'modular' not in desc_normalized:
                        return False
                    
                    # 6Module = Two Nos = 2 socket+switch combos, uses common switch board
                    if '6module' in heading_lower.replace(' ', '') or ('two nos' in heading_lower):
                        # Look for "2 nos" or "2 no.s" with common switch board or 6 modular
                        has_two = '2 nos' in desc_normalized or '2 no.s' in desc_normalized or '2nos' in desc_normalized
                        has_common_board = 'common switch board' in desc_normalized or 'common switch  board' in desc_normalized
                        # Exclude if it's 8 or 12 modular box
                        is_larger_box = '8 modular' in desc_normalized or '12 modular' in desc_normalized
                        return has_two and has_common_board and not is_larger_box
                    
                    # 8Module = Three Nos = 3 socket+switch combos, uses 8 Modular box
                    elif '8module' in heading_lower.replace(' ', '') or ('three nos' in heading_lower):
                        has_three = '3 nos' in desc_normalized or '3 no.s' in desc_normalized or '3nos' in desc_normalized
                        has_8_box = '8 modular' in desc_normalized
                        return has_three and has_8_box
                    
                    # 12Module = Four Nos = 4 socket+switch combos, uses 12 Modular box
                    elif '12module' in heading_lower.replace(' ', '') or ('four nos' in heading_lower):
                        has_four = '4 nos' in desc_normalized or '4 no.s' in desc_normalized or '4nos' in desc_normalized
                        has_12_box = '12 modular' in desc_normalized
                        return has_four and has_12_box
                    
                    return 'modular' in desc_normalized and 'socket' in desc_normalized
                
                # Ding Dong Bell
                if 'ding' in heading_lower and 'dong' in heading_lower:
                    return 'ding dong' in desc_normalized or ('calling bell' in desc_normalized)
                if 'bell' in heading_lower and 'ding' not in heading_lower and 'point' not in heading_lower:
                    return 'calling bell' in desc_normalized or 'door bell' in desc_normalized
                
                # LED light
                if 'led' in heading_lower:
                    if '1200' in heading_lower or 'length' in heading_lower:
                        return 'led' in desc_normalized and ('1200' in desc_normalized or 'tube' in desc_normalized or 'batten' in desc_normalized)
                    return 'led' in desc_normalized
                
                # Ceiling Fans
                if 'ceiling' in heading_lower and 'fan' in heading_lower:
                    return 'ceiling fan' in desc_normalized or ('ceiling' in desc_normalized and 'fan' in desc_normalized)
                
                # Stepped Electronic Regulator
                if 'regulator' in heading_lower:
                    return 'regulator' in desc_normalized
                
                # Exhaust fans (not points)
                if 'exhaust' in heading_lower and 'fan' in heading_lower and 'point' not in heading_lower:
                    if 'kitchen' in heading_lower or 'metallic' in heading_lower or '12' in heading_lower:
                        # 12" Metallic Exhaust Fan for Kitchen
                        return ('exhaust' in desc_normalized and 'fan' in desc_normalized) and \
                               ('kitchen' in desc_normalized or 'metallic' in desc_normalized or '12' in desc_normalized or '300mm' in desc_normalized)
                    elif 'bathroom' in heading_lower or 'shutter' in heading_lower or '6' in heading_lower:
                        # 6" Shutter type Exhaust Fan for Bathrooms
                        return ('exhaust' in desc_normalized and 'fan' in desc_normalized) and \
                               ('bathroom' in desc_normalized or 'shutter' in desc_normalized or '6' in desc_normalized or '150mm' in desc_normalized)
                    return 'exhaust' in desc_normalized and 'fan' in desc_normalized
                
                # Copper cable runs
                if 'copper' in heading_lower and 'cable' in heading_lower:
                    if '2.5' in heading_lower:
                        return 'copper' in desc_normalized and '2.5' in desc_normalized
                    elif '4.0' in heading_lower or '4 sq' in heading_lower:
                        return 'copper' in desc_normalized and ('4.0' in desc_normalized or '4 sq' in desc_normalized or '4sq' in desc_normalized)
                    return 'copper' in desc_normalized and 'cable' in desc_normalized
                
                # WPTC
                if 'wptc' in heading_lower:
                    return 'wptc' in desc_normalized
                
                # TPN DB
                if 'tpn' in heading_lower:
                    return 'tpn' in desc_normalized
                if 'db' in heading_lower and 'way' in heading_lower:
                    return 'db' in desc_normalized and 'way' in desc_normalized
                
                # Water heater / Geyser
                if 'geyser' in heading_lower or 'water heater' in heading_lower:
                    return 'geyser' in desc_normalized or 'water heater' in desc_normalized or 'gyser' in desc_normalized
                
                # Generic fallback - require ALL significant words to be present
                stop_words = {'and', 'or', 'the', 'a', 'an', 'in', 'on', 'of', 'for', 'with', 'to', 'nos', 'no', 'type', 'rb', 'r.b', 'r.b.', 'n.r.b', 'n.r.b.'}
                heading_words = [w.strip() for w in heading_lower.replace('&', ' ').replace('-', ' ').replace('/', ' ').replace('(', ' ').replace(')', ' ').split() 
                                if len(w.strip()) > 2 and w.strip() not in stop_words]
                
                if heading_words and len(heading_words) >= 2:
                    # ALL significant words must appear in description
                    if all(word in desc_normalized for word in heading_words):
                        return True
                
                return False
            
            for row in range(1, estimate_sheet.max_row + 1):
                # Get the description from column D (or search nearby columns)
                desc_cell_val = None
                for desc_col in [4, 5, 3]:  # D, E, C - prioritize column D
                    cell_val = estimate_sheet.cell(row=row, column=desc_col).value
                    if cell_val and str(cell_val).strip():
                        desc_cell_val = str(cell_val).strip()
                        break
                
                if not desc_cell_val:
                    continue
                
                # Check if this row matches any of our item headings
                for heading in item_headings:
                    if heading in qty_unit_map:
                        continue  # Already found
                    
                    if heading_matches_description(heading, desc_cell_val):
                        # Get quantity from column B and unit from column C
                        qty = ""
                        unit = ""
                        
                        # Try to get quantity from column B first, then other columns
                        if estimate_sheet_values:
                            for qty_col in [2, 6, 3]:  # B, F, C
                                qty_cell = estimate_sheet_values.cell(row=row, column=qty_col)
                                if qty_cell.value is not None:
                                    try:
                                        qty_val = float(qty_cell.value)
                                        if qty_val > 0:
                                            qty = str(qty_val)
                                            break
                                    except (ValueError, TypeError):
                                        pass
                        
                        # Get unit from column C first, then other columns
                        for unit_col in [3, 7, 5]:  # C, G, E
                            unit_cell = estimate_sheet.cell(row=row, column=unit_col)
                            if unit_cell.value is not None:
                                unit_val = str(unit_cell.value).strip()
                                if unit_val and unit_val.lower() in ('nos', 'no', 'mtr', 'mtrs', 'pts', 'pt', 'sqm', 'cum', 'kg', 'l', 'rm', 'each', 'set', 'job', 'ls', 'rmt', 'sqmtr', 'metre', 'meters'):
                                    unit = unit_val
                                    break
                        
                        qty_unit_map[heading] = {'qty': qty, 'unit': unit}
        
        # Build final items list
        all_items = []
        for heading in item_headings:
            item_data = qty_unit_map.get(heading, {'qty': '', 'unit': ''})
            all_items.append({
                'desc': heading,
                'qty': item_data['qty'],
                'unit': item_data['unit']
            })
        
        # Create Word document
        doc = Document()
        
        # Title
        title = doc.add_heading('Specification report accompanying the estimate :-', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in title.runs:
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.underline = True
        
        # Introduction paragraph
        if not work_name:
            work_name = "{{NAME_OF_WORK}}"
        
        intro_para = doc.add_paragraph()
        intro_para.add_run(f'The estimate is prepared for the work {work_name}')
        
        doc.add_paragraph()
        
        # Estimate amount
        amount_para = doc.add_paragraph()
        amount_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if total_amount:
            amount_run = amount_para.add_run(f'Est.Amount: Rs. {total_amount}')
        else:
            amount_run = amount_para.add_run('Est.Amount: {{EST_AMOUNT}}')
        amount_run.font.bold = True
        amount_run.font.underline = True
        
        doc.add_paragraph()
        
        # Body of letter placeholder
        body_para = doc.add_paragraph('{{BODY_OF_LETTER}}')
        for run in body_para.runs:
            run.font.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph()
        
        # "Hence" statement
        doc.add_paragraph('Hence, this estimate has been prepared accordingly.')
        
        doc.add_paragraph()
        
        # Item bullet points
        for item in all_items:
            desc = item.get('desc', '')
            qty = item.get('qty', '')
            unit = item.get('unit', '')
            
            # Format quantity
            if qty:
                try:
                    qty_float = float(qty)
                    if qty_float == int(qty_float):
                        qty = str(int(qty_float))
                    else:
                        qty = str(qty_float)
                except:
                    pass
            
            # Build bullet text
            if qty and unit:
                bullet_text = f'{desc}  -  {qty} {unit}'
            elif qty:
                bullet_text = f'{desc}  -  {qty}'
            else:
                bullet_text = desc
            
            bullet_para = doc.add_paragraph(bullet_text, style='List Bullet')
            for run in bullet_para.runs:
                run.font.size = Pt(11)
                run.font.bold = True
        
        doc.add_paragraph()
        
        # Calculate financial year (April to March)
        from datetime import datetime
        today = datetime.now()
        if today.month >= 4:  # April onwards
            fy_start = today.year
            fy_end = (today.year + 1) % 100
        else:  # Jan to March
            fy_start = today.year - 1
            fy_end = today.year % 100
        financial_year = f"{fy_start}-{fy_end:02d}"
        
        # Footer about rates
        footer_text = (f'The rates proposed in the estimate are as per SQR {financial_year} and Approved rates. L.S. Provision is made in the '
                      'estimate towards GST at 18%, QC amount at 1%, Labour Cess at 1% and NAC amount at 0.1% as per actual '
                      'and LS Provision Towards, unforeseen items & rounding off also proposed in the estimate.')
        footer_para = doc.add_paragraph(footer_text)
        for run in footer_para.runs:
            run.font.size = Pt(10)
        
        doc.add_paragraph()
        
        # Funds section
        funds_para = doc.add_paragraph()
        funds_run = funds_para.add_run('FUNDS: ')
        funds_run.font.bold = True
        funds_run.font.underline = True
        funds_para.add_run('The estimate requires Administrative sanction and also fixes up the agency with provision of funds '
                          'under relevant head of account for taking up the work from the Government, Telangana State Hyderabad')
        
        # Generate filename
        filename = 'Spec_Report.docx'
        
        # Return as download
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        doc.save(response)
        return response
        
    except Exception as e:
        logger.error(f'Error generating specification report: {str(e)}', exc_info=True)
        return render(request, 'core/estimate.html', {
            'error': f'Error generating specification report: {str(e)}'
        })


@login_required(login_url='login')
def download_specification_report_live(request, category):
    """
    Generate specification report from live estimate items (New Estimate module).
    Receives items as JSON from the frontend.
    """
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    if request.method != 'POST':
        return redirect('datas_groups', category=category)
    
    try:
        # Get data from POST
        items_json = request.POST.get('items', '[]')
        work_name = request.POST.get('work_name', '{{NAME_OF_WORK}}')
        total_amount = request.POST.get('total_amount', '0.00')
        
        items = json.loads(items_json)
        
        if not items:
            from django.contrib import messages
            messages.error(request, 'No items with quantities to generate specification report')
            return redirect('datas_groups', category=category)
        
        # Create Word document
        doc = Document()
        
        # Title
        title = doc.add_heading('Specification report accompanying the estimate :-', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in title.runs:
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.underline = True
        
        # Introduction paragraph
        intro_para = doc.add_paragraph()
        intro_para.add_run(f'The estimate is prepared for the work {work_name}')
        
        doc.add_paragraph()
        
        # Estimate amount
        amount_para = doc.add_paragraph()
        amount_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        amount_run = amount_para.add_run(f'Est.Amount: Rs. {total_amount}')
        amount_run.font.bold = True
        amount_run.font.underline = True
        
        doc.add_paragraph()
        
        # Body of letter placeholder
        body_para = doc.add_paragraph('{{BODY_OF_LETTER}}')
        for run in body_para.runs:
            run.font.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph()
        
        # "Hence" statement
        doc.add_paragraph('Hence, this estimate has been prepared accordingly.')
        
        doc.add_paragraph()
        
        # Item bullet points
        for item in items:
            desc = item.get('desc', '')
            qty = item.get('qty', '')
            unit = item.get('unit', '')
            
            # Format quantity - remove trailing .0 if whole number
            if qty:
                try:
                    qty_float = float(qty)
                    if qty_float == int(qty_float):
                        qty = str(int(qty_float))
                    else:
                        qty = str(qty_float)
                except:
                    pass
            
            # Build bullet text
            if qty and unit:
                bullet_text = f'{desc}  -  {qty} {unit}'
            elif qty:
                bullet_text = f'{desc}  -  {qty}'
            else:
                bullet_text = desc
            
            bullet_para = doc.add_paragraph(bullet_text, style='List Bullet')
            for run in bullet_para.runs:
                run.font.size = Pt(11)
                run.font.bold = True
        
        doc.add_paragraph()
        
        # Calculate financial year (April to March)
        from datetime import datetime
        today = datetime.now()
        if today.month >= 4:  # April onwards
            fy_start = today.year
            fy_end = (today.year + 1) % 100
        else:  # Jan to March
            fy_start = today.year - 1
            fy_end = today.year % 100
        financial_year = f"{fy_start}-{fy_end:02d}"
        
        # Footer about rates
        footer_text = (f'The rates proposed in the estimate are as per SQR {financial_year} and Approved rates. L.S. Provision is made in the '
                      'estimate towards GST at 18%, QC amount at 1%, Labour Cess at 1% and NAC amount at 0.1% as per actual '
                      'and LS Provision Towards, unforeseen items & rounding off also proposed in the estimate.')
        footer_para = doc.add_paragraph(footer_text)
        for run in footer_para.runs:
            run.font.size = Pt(10)
        
        doc.add_paragraph()
        
        # Funds section
        funds_para = doc.add_paragraph()
        funds_run = funds_para.add_run('FUNDS: ')
        funds_run.font.bold = True
        funds_run.font.underline = True
        funds_para.add_run('The estimate requires Administrative sanction and also fixes up the agency with provision of funds '
                          'under relevant head of account for taking up the work from the Government, Telangana State Hyderabad')
        
        # Generate filename
        filename = 'Spec_Report.docx'
        
        # Return as download
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        doc.save(response)
        return response
        
    except Exception as e:
        logger.error(f'Error generating specification report: {str(e)}', exc_info=True)
        from django.contrib import messages
        messages.error(request, f'Error generating specification report: {str(e)}')
        return redirect('datas_groups', category=category)


@login_required(login_url='login')
def download_forwarding_letter_live(request, category):
    """
    Generate forwarding letter from live estimate items (New Estimate module).
    Receives items as JSON from the frontend - similar to specification report.
    """
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    
    if request.method != 'POST':
        return redirect('datas_groups', category=category)
    
    try:
        # Get data from POST
        items_json = request.POST.get('items', '[]')
        work_name = request.POST.get('work_name', '{{NAME_OF_WORK}}')
        total_amount = request.POST.get('total_amount', '0.00')
        
        items = json.loads(items_json)
        
        if not items:
            from django.contrib import messages
            messages.error(request, 'No items with quantities to generate forwarding letter')
            return redirect('datas_groups', category=category)
        
        # Parse total amount
        try:
            grand_total = float(total_amount.replace(',', '').replace('Rs.', '').replace('₹', '').strip())
        except:
            grand_total = 0.0
        
        # Get current date and financial year
        current_date = _get_current_date_formatted()
        financial_year = _get_current_financial_year()
        today = timezone.now().date()
        
        # Get user's letter settings
        letter_settings = _get_letter_settings(request.user)
        
        # Create Word document
        doc = Document()
        
        # Light gray color for placeholders
        placeholder_color = RGBColor(169, 169, 169)
        
        # Set page margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Header - Government/Organization name
        header1 = doc.add_paragraph()
        header1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if letter_settings and letter_settings.government_name:
            run1 = header1.add_run(letter_settings.government_name.upper())
            run1.font.bold = True
            run1.font.size = Pt(14)
        else:
            run1 = header1.add_run('[GOVERNMENT / ORGANIZATION NAME]')
            run1.font.bold = True
            run1.font.size = Pt(14)
            run1.font.color.rgb = placeholder_color
            run1.font.italic = True
        
        # Header - Department name
        header2 = doc.add_paragraph()
        header2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if letter_settings and letter_settings.department_name:
            run2 = header2.add_run(letter_settings.department_name.upper())
            run2.font.bold = True
            run2.font.size = Pt(13)
        else:
            run2 = header2.add_run('[DEPARTMENT NAME]')
            run2.font.bold = True
            run2.font.size = Pt(13)
            run2.font.color.rgb = placeholder_color
            run2.font.italic = True
        
        doc.add_paragraph()
        
        # From/To section in a table
        from_to_table = doc.add_table(rows=1, cols=2)
        from_to_table.autofit = True
        
        from_cell = from_to_table.cell(0, 0)
        from_para = from_cell.paragraphs[0]
        from_para.add_run('From: -\n')
        
        # From section - Officer details
        if letter_settings and letter_settings.officer_name:
            name_qual = letter_settings.officer_name
            if letter_settings.officer_qualification:
                name_qual += f", {letter_settings.officer_qualification}"
            from_run1 = from_para.add_run(f'{name_qual},\n')
        else:
            from_run1 = from_para.add_run('[Officer Name, Qualification],\n')
            from_run1.font.color.rgb = placeholder_color
            from_run1.font.italic = True
        
        if letter_settings and letter_settings.officer_designation:
            from_run2 = from_para.add_run(f'{letter_settings.officer_designation},\n')
        else:
            from_run2 = from_para.add_run('[Designation],\n')
            from_run2.font.color.rgb = placeholder_color
            from_run2.font.italic = True
        
        if letter_settings and (letter_settings.sub_division or letter_settings.office_address):
            sub_addr = letter_settings.sub_division
            if letter_settings.office_address:
                sub_addr += f", {letter_settings.office_address}" if sub_addr else letter_settings.office_address
            from_run3 = from_para.add_run(f'{sub_addr}.')
        else:
            from_run3 = from_para.add_run('[Sub Division, Office Address].')
            from_run3.font.color.rgb = placeholder_color
            from_run3.font.italic = True
        
        # To section - Recipient details
        to_cell = from_to_table.cell(0, 1)
        to_para = to_cell.paragraphs[0]
        to_para.add_run('To,\n')
        
        if letter_settings and letter_settings.recipient_designation:
            to_run1 = to_para.add_run(f'{letter_settings.recipient_designation},\n')
        else:
            to_run1 = to_para.add_run('[Officer Designation],\n')
            to_run1.font.color.rgb = placeholder_color
            to_run1.font.italic = True
        
        if letter_settings and letter_settings.recipient_division:
            to_run2 = to_para.add_run(f'{letter_settings.recipient_division},\n')
        else:
            to_run2 = to_para.add_run('[Division Name],\n')
            to_run2.font.color.rgb = placeholder_color
            to_run2.font.italic = True
        
        if letter_settings and letter_settings.recipient_address:
            to_run3 = to_para.add_run(f'{letter_settings.recipient_address}.')
        else:
            to_run3 = to_para.add_run('[Address].')
            to_run3.font.color.rgb = placeholder_color
            to_run3.font.italic = True
        
        doc.add_paragraph()
        
        # Letter number and date
        lr_para = doc.add_paragraph()
        lr_para.add_run('Lr No. ')
        if letter_settings and letter_settings.office_code:
            lr_code = lr_para.add_run(letter_settings.office_code)
            lr_code.font.underline = True
        else:
            lr_placeholder = lr_para.add_run('[Office Code]')
            lr_placeholder.font.color.rgb = placeholder_color
            lr_placeholder.font.italic = True
            lr_placeholder.font.underline = True
        lr_para.add_run(f'/{financial_year}/          ')
        lr_para.add_run(f'\t\t\t\t\tDate:-    - {today.strftime("%m")} - {today.year}.')
        
        doc.add_paragraph()
        
        # Sir,
        sir_para = doc.add_paragraph()
        sir_para.add_run('Sir,')
        
        doc.add_paragraph()
        
        # Subject
        subject_para = doc.add_paragraph()
        subj_run = subject_para.add_run('Sub:-')
        subj_run.font.underline = True
        subject_para.add_run('\t')
        subj_work = subject_para.add_run(f'{work_name} ')
        subj_work.font.bold = True
        subject_para.add_run(f'for the year {financial_year}.  -  Submission  -  Request for obtaining administrative sanction  -  Regarding.')
        
        doc.add_paragraph()
        
        # Reference
        ref_para = doc.add_paragraph()
        ref_run = ref_para.add_run('Ref:-')
        ref_run.font.underline = True
        ref_para.add_run('\tMemo No.')
        ref_placeholder = ref_para.add_run('[Reference Number]')
        ref_placeholder.font.color.rgb = placeholder_color
        ref_placeholder.font.italic = True
        ref_placeholder.font.underline = True
        ref_para.add_run(f'/{financial_year} Dt.')
        ref_date_placeholder = ref_para.add_run('[DD.MM.YYYY]')
        ref_date_placeholder.font.color.rgb = placeholder_color
        ref_date_placeholder.font.italic = True
        ref_date_placeholder.font.underline = True
        
        doc.add_paragraph()
        
        # Stars separator
        stars_para = doc.add_paragraph()
        stars_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        stars_para.add_run('**.**')
        
        doc.add_paragraph()
        
        # Main body - for single estimate
        body_para = doc.add_paragraph()
        body_para.add_run('With reference to the subject cited, I submit here ')
        with_run = body_para.add_run('with  1')
        with_run.font.underline = True
        body_para.add_run(' No. estimate for the following work for the amount specified.')
        
        doc.add_paragraph()
        
        # Create table for estimate
        table = doc.add_table(rows=2, cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set column widths
        for cell in table.columns[0].cells:
            cell.width = Inches(0.5)
        for cell in table.columns[1].cells:
            cell.width = Inches(4.5)
        for cell in table.columns[2].cells:
            cell.width = Inches(1.5)
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Sl.\nNo'
        header_cells[1].text = 'Name of work'
        header_cells[2].text = 'Amount'
        
        # Center align and bold header
        for cell in header_cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.bold = True
        
        # Data row
        row_cells = table.rows[1].cells
        row_cells[0].text = '1'
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[1].text = work_name
        
        # Amount in Indian format
        formatted_amount = _format_indian_number(grand_total)
        row_cells[2].text = f"Rs.{formatted_amount}"
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()
        
        # Specification statement
        spec_para = doc.add_paragraph()
        spec_para.add_run("Specification report accompanying the estimate explains the necessity and provisions made therein in detail.")
        
        doc.add_paragraph()
        
        # Request paragraph
        request_para = doc.add_paragraph()
        request_para.add_run('I request the ')
        if letter_settings and letter_settings.superior_designation:
            req_run = request_para.add_run(letter_settings.superior_designation)
        else:
            req_placeholder = request_para.add_run('[Superior Officer Designation]')
            req_placeholder.font.color.rgb = placeholder_color
            req_placeholder.font.italic = True
        request_para.add_run(' to kindly arrange to obtain administrative sanction for the above estimate and arrange to finalize the agency at the earliest for taking up the work.')
        
        doc.add_paragraph()
        
        # Enclosure
        enc_para = doc.add_paragraph()
        enc_para.add_run('Enclosure: -')
        doc.add_paragraph('Estimate  - 1 No.')
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Signature section
        sign_para = doc.add_paragraph()
        sign_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sign_para.add_run('Yours faithfully,')
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        if letter_settings and letter_settings.officer_designation:
            run_title = title_para.add_run(f'{letter_settings.officer_designation}\n')
            run_title.font.bold = True
        else:
            run_title = title_para.add_run('[Officer Designation]\n')
            run_title.font.bold = True
            run_title.font.color.rgb = placeholder_color
            run_title.font.italic = True
        
        if letter_settings and letter_settings.sub_division:
            sub_div_run = title_para.add_run(f'{letter_settings.sub_division},\n')
        else:
            sub_div_run = title_para.add_run('[Sub Division Name],\n')
            sub_div_run.font.color.rgb = placeholder_color
            sub_div_run.font.italic = True
        
        if letter_settings and letter_settings.office_address:
            addr_run = title_para.add_run(f'{letter_settings.office_address}.')
        else:
            addr_run = title_para.add_run('[Office Address].')
            addr_run.font.color.rgb = placeholder_color
            addr_run.font.italic = True
        
        doc.add_paragraph()
        
        # Copy to
        copy_para = doc.add_paragraph()
        copy_para.add_run('Copy to the ')
        if letter_settings and (letter_settings.copy_to_designation or letter_settings.copy_to_section):
            copy_text = letter_settings.copy_to_designation or ''
            if letter_settings.copy_to_section:
                copy_text += f", {letter_settings.copy_to_section}" if copy_text else letter_settings.copy_to_section
            copy_run = copy_para.add_run(copy_text)
        else:
            copy_placeholder = copy_para.add_run('[Officer Designation, Section Name]')
            copy_placeholder.font.color.rgb = placeholder_color
            copy_placeholder.font.italic = True
        copy_para.add_run(' for information.')
        
        # Generate filename
        filename = 'Fwd_Letter.docx'
        
        # Return as download
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        doc.save(response)
        return response
        
    except Exception as e:
        logger.error(f'Error generating forwarding letter: {str(e)}', exc_info=True)
        from django.contrib import messages
        messages.error(request, f'Error generating forwarding letter: {str(e)}')
        return redirect('datas_groups', category=category)


@login_required(login_url='login')
def generate_estimate_forwarding_letter(request):
    """
    Generate a forwarding letter for multi-sheet estimates.
    
    Extracts name of work and grand total from each Estimate sheet in the workbook
    and generates a formal forwarding letter in Word format with:
    - Serial numbered table of works and amounts
    - Indian number formatting for amounts
    - Dynamic financial year and date
    - Generic officer designations
    """
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    if request.method != 'POST':
        return render(request, 'core/estimate.html', {
            'error': 'Invalid request method'
        })
    
    output_file = request.FILES.get('output_file')
    if not output_file:
        return render(request, 'core/estimate.html', {
            'error': 'Please upload an Excel file.'
        })
    
    try:
        # Reset file position
        try:
            output_file.seek(0)
        except:
            pass
        
        wb_upload = load_workbook(output_file, data_only=False)
        
        # Also load with data_only=True for values
        output_file.seek(0)
        wb_values = load_workbook(output_file, data_only=True)
        
        # Find all Estimate sheets and extract work names and grand totals
        estimates_data = []
        
        for sheet_name in wb_upload.sheetnames:
            sheet_lower = sheet_name.lower().strip()
            
            # Look for sheets that contain "estimate" in their name
            if 'estimate' in sheet_lower:
                ws = wb_upload[sheet_name]
                ws_values = wb_values[sheet_name]
                
                work_name = ""
                grand_total = 0.0
                
                # Find work name - look for "Name of the work" or "Name of work"
                for row in range(1, min(30, ws.max_row + 1)):
                    if work_name:
                        break
                    for col in range(1, 10):
                        cell_val = str(ws.cell(row=row, column=col).value or "").strip()
                        cell_lower = cell_val.lower()
                        
                        if "name of the work" in cell_lower or "name of work" in cell_lower:
                            # Check if work name is in the same cell after colon
                            if ':' in cell_val:
                                parts = cell_val.split(':', 1)
                                if len(parts) > 1 and parts[1].strip():
                                    work_name = parts[1].strip()
                                    break
                            
                            # Try next columns in same row
                            if not work_name:
                                for next_col in range(col + 1, col + 8):
                                    if next_col <= ws.max_column:
                                        next_cell = ws.cell(row=row, column=next_col).value
                                        if next_cell and str(next_cell).strip():
                                            work_name = str(next_cell).strip()
                                            break
                            
                            # Try next row
                            if not work_name:
                                if row + 1 <= ws.max_row:
                                    next_row_val = ws.cell(row=row + 1, column=col).value
                                    if next_row_val and str(next_row_val).strip():
                                        work_name = str(next_row_val).strip()
                            break
                
                # Find Grand Total
                for row in range(1, ws.max_row + 1):
                    for col in range(1, 10):
                        cell_val = str(ws.cell(row=row, column=col).value or "").strip().lower()
                        if 'grand total' in cell_val:
                            # Look for amount in columns to the right (especially column H)
                            for amt_col in range(col + 1, col + 8):
                                if amt_col <= ws.max_column:
                                    amt_cell = ws_values.cell(row=row, column=amt_col)
                                    if amt_cell.value is not None:
                                        try:
                                            amt_val = float(amt_cell.value)
                                            if amt_val > 100:  # Should be a significant amount
                                                grand_total = amt_val
                                                break
                                        except (ValueError, TypeError):
                                            pass
                            
                            if grand_total > 0:
                                break
                    if grand_total > 0:
                        break
                
                # Only add if we found meaningful data
                if work_name or grand_total > 0:
                    estimates_data.append({
                        'sheet_name': sheet_name,
                        'work_name': work_name or f"Work from {sheet_name}",
                        'grand_total': grand_total
                    })
        
        if not estimates_data:
            return render(request, 'core/estimate.html', {
                'error': 'No Estimate sheets found with work names or grand totals. Make sure your workbook has sheets with "Estimate" in the name.'
            })
        
        # Get current date and financial year
        current_date = _get_current_date_formatted()
        financial_year = _get_current_financial_year()
        
        # Get user's letter settings
        letter_settings = _get_letter_settings(request.user)
        
        # Create Word document
        doc = Document()
        
        # Light gray color for placeholders
        placeholder_color = RGBColor(169, 169, 169)  # Light gray
        
        # Set page margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Header - Government/Organization name
        header1 = doc.add_paragraph()
        header1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if letter_settings and letter_settings.government_name:
            run1 = header1.add_run(letter_settings.government_name.upper())
            run1.font.bold = True
            run1.font.size = Pt(14)
        else:
            run1 = header1.add_run('[GOVERNMENT / ORGANIZATION NAME]')
            run1.font.bold = True
            run1.font.size = Pt(14)
            run1.font.color.rgb = placeholder_color
            run1.font.italic = True
        
        # Header - Department name
        header2 = doc.add_paragraph()
        header2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if letter_settings and letter_settings.department_name:
            run2 = header2.add_run(letter_settings.department_name.upper())
            run2.font.bold = True
            run2.font.size = Pt(13)
        else:
            run2 = header2.add_run('[DEPARTMENT NAME]')
            run2.font.bold = True
            run2.font.size = Pt(13)
            run2.font.color.rgb = placeholder_color
            run2.font.italic = True
        
        doc.add_paragraph()  # Blank line
        
        # From/To section in a table
        from_to_table = doc.add_table(rows=1, cols=2)
        from_to_table.autofit = True
        
        from_cell = from_to_table.cell(0, 0)
        from_para = from_cell.paragraphs[0]
        from_para.add_run('From: -\n')
        
        # From section - Officer details
        if letter_settings and letter_settings.officer_name:
            name_qual = letter_settings.officer_name
            if letter_settings.officer_qualification:
                name_qual += f", {letter_settings.officer_qualification}"
            from_run1 = from_para.add_run(f'{name_qual},\n')
        else:
            from_run1 = from_para.add_run('[Officer Name, Qualification],\n')
            from_run1.font.color.rgb = placeholder_color
            from_run1.font.italic = True
        
        if letter_settings and letter_settings.officer_designation:
            from_run2 = from_para.add_run(f'{letter_settings.officer_designation},\n')
        else:
            from_run2 = from_para.add_run('[Designation],\n')
            from_run2.font.color.rgb = placeholder_color
            from_run2.font.italic = True
        
        if letter_settings and (letter_settings.sub_division or letter_settings.office_address):
            sub_addr = letter_settings.sub_division
            if letter_settings.office_address:
                sub_addr += f", {letter_settings.office_address}" if sub_addr else letter_settings.office_address
            from_run3 = from_para.add_run(f'{sub_addr}.')
        else:
            from_run3 = from_para.add_run('[Sub Division, Office Address].')
            from_run3.font.color.rgb = placeholder_color
            from_run3.font.italic = True
        
        # To section - Recipient details
        to_cell = from_to_table.cell(0, 1)
        to_para = to_cell.paragraphs[0]
        to_para.add_run('To,\n')
        
        if letter_settings and letter_settings.recipient_designation:
            to_run1 = to_para.add_run(f'{letter_settings.recipient_designation},\n')
        else:
            to_run1 = to_para.add_run('[Officer Designation],\n')
            to_run1.font.color.rgb = placeholder_color
            to_run1.font.italic = True
        
        if letter_settings and letter_settings.recipient_division:
            to_run2 = to_para.add_run(f'{letter_settings.recipient_division},\n')
        else:
            to_run2 = to_para.add_run('[Division Name],\n')
            to_run2.font.color.rgb = placeholder_color
            to_run2.font.italic = True
        
        if letter_settings and letter_settings.recipient_address:
            to_run3 = to_para.add_run(f'{letter_settings.recipient_address}.')
        else:
            to_run3 = to_para.add_run('[Address].')
            to_run3.font.color.rgb = placeholder_color
            to_run3.font.italic = True
        
        doc.add_paragraph()
        
        # Letter number and date on same line
        today = timezone.now().date()
        month_num = today.month
        year_short = today.year % 100
        
        lr_para = doc.add_paragraph()
        lr_para.add_run('Lr No. ')
        if letter_settings and letter_settings.office_code:
            lr_code = lr_para.add_run(letter_settings.office_code)
            lr_code.font.underline = True
        else:
            lr_placeholder = lr_para.add_run('[Office Code]')
            lr_placeholder.font.color.rgb = placeholder_color
            lr_placeholder.font.italic = True
            lr_placeholder.font.underline = True
        lr_para.add_run(f'/{financial_year}/          ')
        lr_para.add_run(f'\t\t\t\t\tDate:-    - {today.strftime("%m")} - {today.year}.')
        
        doc.add_paragraph()
        
        # Sir,
        sir_para = doc.add_paragraph()
        sir_para.add_run('Sir,')
        
        doc.add_paragraph()
        
        # Subject - use first work name if single estimate, otherwise placeholder
        subject_para = doc.add_paragraph()
        subj_run = subject_para.add_run('Sub:-')
        subj_run.font.underline = True
        subject_para.add_run('\t')
        if len(estimates_data) == 1:
            subj_work = subject_para.add_run(f'{estimates_data[0]["work_name"]} ')
            subj_work.font.bold = True
        else:
            subj_placeholder = subject_para.add_run(f'[Subject of the letter] ')
            subj_placeholder.font.color.rgb = placeholder_color
            subj_placeholder.font.italic = True
        subject_para.add_run(f'for the year {financial_year}.  -  Submission  -  Request for obtaining administrative sanction  -  Regarding.')
        
        doc.add_paragraph()
        
        # Reference
        ref_para = doc.add_paragraph()
        ref_run = ref_para.add_run('Ref:-')
        ref_run.font.underline = True
        ref_para.add_run('\tMemo No.')
        ref_placeholder = ref_para.add_run('[Reference Number]')
        ref_placeholder.font.color.rgb = placeholder_color
        ref_placeholder.font.italic = True
        ref_placeholder.font.underline = True
        ref_para.add_run(f'/{financial_year} Dt.')
        ref_date_placeholder = ref_para.add_run('[DD.MM.YYYY]')
        ref_date_placeholder.font.color.rgb = placeholder_color
        ref_date_placeholder.font.italic = True
        ref_date_placeholder.font.underline = True
        
        doc.add_paragraph()
        
        # Stars separator
        stars_para = doc.add_paragraph()
        stars_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        stars_para.add_run('**.**')
        
        doc.add_paragraph()
        
        # Main body
        body_para = doc.add_paragraph()
        body_para.add_run('With reference to the subject cited, I submit here ')
        with_run = body_para.add_run(f'with  {len(estimates_data)}')
        with_run.font.underline = True
        body_para.add_run(' Nos. estimates for the following works for the amounts specifies against each work.')
        
        doc.add_paragraph()
        
        # Create table for estimates
        table = doc.add_table(rows=len(estimates_data) + 1, cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set column widths
        for cell in table.columns[0].cells:
            cell.width = Inches(0.5)
        for cell in table.columns[1].cells:
            cell.width = Inches(4.5)
        for cell in table.columns[2].cells:
            cell.width = Inches(1.5)
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Sl.\nNo'
        header_cells[1].text = 'Name of work'
        header_cells[2].text = 'Amount'
        
        # Center align and bold header
        for cell in header_cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.bold = True
        
        # Data rows
        for idx, est_data in enumerate(estimates_data, start=1):
            row_cells = table.rows[idx].cells
            
            # Serial number
            row_cells[0].text = str(idx)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Name of work
            row_cells[1].text = est_data['work_name']
            
            # Amount in Indian format
            formatted_amount = _format_indian_number(est_data['grand_total'])
            row_cells[2].text = f"Rs.{formatted_amount}"
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()
        
        # Specification statement
        spec_para = doc.add_paragraph()
        spec_para.add_run("Specification report accompanying each estimate explains the necessity and provisions made they're in detail.")
        
        doc.add_paragraph()
        
        # Request paragraph
        request_para = doc.add_paragraph()
        request_para.add_run('I request the ')
        if letter_settings and letter_settings.superior_designation:
            req_run = request_para.add_run(letter_settings.superior_designation)
        else:
            req_placeholder = request_para.add_run('[Superior Officer Designation]')
            req_placeholder.font.color.rgb = placeholder_color
            req_placeholder.font.italic = True
        request_para.add_run(' to kindly arrange to obtain administrative sanction the above estimates and arrange to finalize the agencies at the earliest for taking up the works.')
        
        doc.add_paragraph()
        
        # Enclosure
        enc_para = doc.add_paragraph()
        enc_para.add_run('Enclosure: -')
        doc.add_paragraph(f'Estimates  - {len(estimates_data)} No\'s,')
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Signature section
        sign_para = doc.add_paragraph()
        sign_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sign_para.add_run('Yours faithfully,')
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        if letter_settings and letter_settings.officer_designation:
            run_title = title_para.add_run(f'{letter_settings.officer_designation}\n')
            run_title.font.bold = True
        else:
            run_title = title_para.add_run('[Officer Designation]\n')
            run_title.font.bold = True
            run_title.font.color.rgb = placeholder_color
            run_title.font.italic = True
        
        if letter_settings and letter_settings.sub_division:
            sub_div_run = title_para.add_run(f'{letter_settings.sub_division},\n')
        else:
            sub_div_run = title_para.add_run('[Sub Division Name],\n')
            sub_div_run.font.color.rgb = placeholder_color
            sub_div_run.font.italic = True
        
        if letter_settings and letter_settings.office_address:
            addr_run = title_para.add_run(f'{letter_settings.office_address}.')
        else:
            addr_run = title_para.add_run('[Office Address].')
            addr_run.font.color.rgb = placeholder_color
            addr_run.font.italic = True
        
        doc.add_paragraph()
        
        # Copy to
        copy_para = doc.add_paragraph()
        copy_para.add_run('Copy to the ')
        if letter_settings and (letter_settings.copy_to_designation or letter_settings.copy_to_section):
            copy_text = letter_settings.copy_to_designation or ''
            if letter_settings.copy_to_section:
                copy_text += f", {letter_settings.copy_to_section}" if copy_text else letter_settings.copy_to_section
            copy_run = copy_para.add_run(copy_text)
        else:
            copy_placeholder = copy_para.add_run('[Officer Designation, Section Name]')
            copy_placeholder.font.color.rgb = placeholder_color
            copy_placeholder.font.italic = True
        copy_para.add_run(' for information.')
        
        # Generate filename
        filename = 'Fwd_Letter.docx'
        
        # Return as download
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        doc.save(response)
        return response
        
    except Exception as e:
        logger.error(f'Error generating forwarding letter: {str(e)}', exc_info=True)
        return render(request, 'core/estimate.html', {
            'error': f'Error generating forwarding letter: {str(e)}'
        })


# ==============================================================================
# AMC MODULE VIEWS
# ==============================================================================
# AMC (Annual Maintenance Contract) module works similar to New Estimate module
# but uses a custom backend sheet uploaded via Admin Panel

