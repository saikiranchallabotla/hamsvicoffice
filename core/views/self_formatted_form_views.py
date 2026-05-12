# Auto-generated from core/views.py split
import json
import os
import re
import logging
from copy import copy

import inflect
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, Border, Side, PatternFill
from openpyxl.utils import get_column_letter
from django.utils import timezone
from django.urls import reverse
from django.core.files.uploadedfile import InMemoryUploadedFile
from django.contrib.auth.decorators import login_required

from django.conf import settings
from django.http import Http404, HttpResponse, JsonResponse, HttpResponseNotAllowed
from django.shortcuts import render, redirect, get_object_or_404
from django.template.loader import render_to_string
from django.views.decorators.http import require_POST
from django.utils.crypto import get_random_string

import io
from io import BytesIO
from difflib import SequenceMatcher

from ..models import Project, SelfFormattedTemplate, Estimate, Organization, Membership, Upload, Job, OutputFile, LetterSettings
from ..decorators import org_required, role_required

logger = logging.getLogger(__name__)
from ..tasks import process_excel_upload, generate_bill_pdf, generate_workslip_pdf, generate_bill_document_task
from ..utils_excel import load_backend, copy_block_with_styles_and_formulas, build_temp_day_rates

p_engine = inflect.engine()
BILL_TEMPLATES_DIR = os.path.join(settings.BASE_DIR, "core", "templates", "core", "bill_templates")
_inflect_engine = inflect.engine()

from .utils import (get_org_from_request, _get_letter_settings,
    _get_current_financial_year, _format_indian_number)
from .self_formatted_views import (_extract_labels_from_source_file,
    _build_placeholder_map, _fill_template_file,
    _replace_placeholders_in_docx_xml,
    _extract_labels_per_work,
    _extract_labels_from_lines)

# Pre-migration-safe columns for SelfFormattedTemplate queries.
# Fields from migration 0022 (is_locked, template_file_backup, etc.) may not
# exist yet, so use only() with these safe columns to avoid ProgrammingError.
_SAFE_TEMPLATE_FIELDS = (
    'id', 'name', 'description', 'template_file',
    'custom_placeholders', 'is_shared', 'user_id',
    'organization_id', 'created_at', 'updated_at',
)


def _safe_get_template(pk, user):
    """Get a SelfFormattedTemplate by pk/user, safe even if migration 0022 hasn't run."""
    try:
        return SelfFormattedTemplate.objects.get(pk=pk, user=user)
    except SelfFormattedTemplate.DoesNotExist:
        return None
    except Exception:
        # Column missing - fall back to safe fields only
        try:
            return SelfFormattedTemplate.objects.only(*_SAFE_TEMPLATE_FIELDS).get(pk=pk, user=user)
        except SelfFormattedTemplate.DoesNotExist:
            return None

@login_required(login_url='login')
def self_formatted_form_page(request):
    """
    Shows:
      - Quick one-time generation form
      - Create reusable format form
      - List of saved formats (with lock status)
    Optimized: Limited query with only necessary fields for faster load.
    """
    # Only fetch the current user's formats (not other users')
    # Use only() to select known-safe columns, avoiding new fields that may
    # not yet exist in the DB (migration 0022 may not have run).
    # list() forces evaluation inside the try/except so DB errors are caught.
    try:
        saved_formats = list(
            SelfFormattedTemplate.objects.filter(
                user=request.user
            ).only(
                'id', 'name', 'description', 'template_file',
                'custom_placeholders', 'is_shared', 'user_id',
                'organization_id', 'created_at', 'updated_at'
            ).order_by("-created_at")[:20]
        )
    except Exception:
        saved_formats = []
    
    error_message = request.GET.get("error")  # optional error via redirect
    success_message = request.GET.get("success")  # optional success message

    return render(request, "core/self_formatted.html", {
        "saved_formats": saved_formats,
        "error_message": error_message,
        "success_message": success_message,
    })


@login_required(login_url='login')
def self_formatted_generate(request):
    """
    Quick one-time generation: user uploads source + template, optional
    custom placeholders text. Does not save anything in DB.
    Optimized: No database queries during file processing for speed.
    """
    if request.method != "POST":
        return HttpResponseNotAllowed(["POST"])

    source_file = request.FILES.get("source_file")
    template_file = request.FILES.get("template_file")
    custom_text = request.POST.get("custom_placeholders", "")

    # On error, redirect with error message instead of querying DB
    if not source_file or not template_file:
        return redirect(f"{reverse('self_formatted_form_page')}?error=Please+upload+both+source+file+and+template+file")

    try:
        labels, lines = _extract_labels_from_source_file(source_file)
    except Exception as e:
        import traceback
        logger.error(f"Source file extraction failed: {e}\n{traceback.format_exc()}")
        return redirect(f"{reverse('self_formatted_form_page')}?error=Failed+to+read+source+file.+Please+check+the+file+format.")

    # Check if no text was extracted (likely scanned PDF)
    if not lines:
        filename = source_file.name or ""
        if filename.lower().endswith('.pdf'):
            error_msg = "PDF+appears+to+be+scanned.+Use+Excel+or+Word+file+instead."
        else:
            error_msg = "No+text+could+be+extracted+from+the+source+file."
        return redirect(f"{reverse('self_formatted_form_page')}?error={error_msg}")

    placeholder_map = _build_placeholder_map(labels, lines, custom_text)

    try:
        return _fill_template_file(template_file, placeholder_map)
    except Exception as e:
        import traceback
        logger.error(f"Template fill failed: {e}\n{traceback.format_exc()}")
        return redirect(f"{reverse('self_formatted_form_page')}?error=Failed+to+fill+template.+Please+check+the+template+file+format.")


@login_required(login_url='login')
def self_formatted_preview(request):
    """AJAX endpoint: compute placeholder_map for a given source file + custom text and return JSON.
    Used by the UI to preview mappings before generation.
    Uses faster/lighter OCR settings for quicker preview response.
    """
    from django.http import JsonResponse
    import logging
    logger = logging.getLogger(__name__)

    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)

    source_file = request.FILES.get("source_file")
    custom_text = request.POST.get("custom_placeholders", "")

    if not source_file:
        return JsonResponse({"error": "source_file required"}, status=400)

    try:
        logger.info(f"Preview: Processing {source_file.name}")
        labels, lines = _extract_labels_from_source_file(source_file)
        
        # Check if no text was extracted
        if not lines:
            filename = source_file.name or ""
            if filename.lower().endswith('.pdf'):
                return JsonResponse({
                    "error": "No text extracted. This appears to be a scanned PDF. Use Excel/Word file or install Tesseract OCR."
                }, status=400)
            return JsonResponse({"error": "No text could be extracted from the file."}, status=400)
        
        logger.info(f"Preview: Extracted {len(lines)} lines, building placeholders...")
        placeholder_map = _build_placeholder_map(labels, lines, custom_text)
        logger.info(f"Preview: Found {len(placeholder_map)} placeholders")
    except Exception as e:
        logger.error(f"Preview error: {e}")
        return JsonResponse({"error": str(e)}, status=500)

    return JsonResponse({"placeholders": placeholder_map})


@org_required
def self_formatted_save_format(request):
    """
    Save a reusable format (name + description + template file + custom placeholders).
    Automatically creates a database backup of the template for maximum persistence.
    Optimized: Redirect with error instead of querying DB.
    """
    if request.method != "POST":
        return redirect("self_formatted_form_page")

    format_name = request.POST.get("format_name", "").strip()
    format_description = request.POST.get("format_description", "").strip()
    template_file = request.FILES.get("format_template_file")
    raw_custom = request.POST.get("format_custom_placeholders", "").strip()

    if not format_name or not template_file:
        return redirect(f"{reverse('self_formatted_form_page')}?error=Format+name+and+template+file+are+required")

    # Read the template file content for backup before save
    template_file.seek(0)
    template_content = template_file.read()
    template_file.seek(0)  # Reset for FileField save

    fmt = SelfFormattedTemplate(
        name=format_name,
        description=format_description,
        template_file=template_file,
        custom_placeholders=raw_custom,
        user=request.user,
        organization=request.organization,
    )
    
    # Set persistence fields (from migration 0022)
    fmt.is_locked = True
    fmt.template_file_backup = template_content
    fmt.template_file_name = template_file.name
    fmt.template_file_size = len(template_content)
    fmt.save()

    return redirect("self_formatted_form_page")


@org_required
def self_formatted_use_format(request, pk):
    """
    Use a saved format:
      GET  -> show page asking only for source_file upload.
      POST -> generate document using saved template + placeholders.
    """
    fmt = _safe_get_template(pk, request.user)
    if fmt is None:
        raise Http404

    if request.method == "GET":
        return render(request, "core/self_formatted_use.html", {
            "format": fmt,
        })

    if request.method == "POST":
        source_file = request.FILES.get("source_file")
        if not source_file:
            return HttpResponse("Please upload a source file.", status=400)

        try:
            labels, lines = _extract_labels_from_source_file(source_file)
        except Exception as e:
            logger.error(f"Source file extraction failed in use_format: {e}")
            return HttpResponse("Failed to read source file. Please check the file format.", status=400)

        if not lines:
            return HttpResponse("No text could be extracted from the source file.", status=400)

        placeholder_source = fmt.custom_placeholders or ""
        placeholder_map = _build_placeholder_map(labels, lines, placeholder_source)

        # Use the new get_template_content method which falls back to backup
        data = fmt.get_template_content()

        if not data:
            # Template file was not found in file storage OR database backup
            return redirect(
                f"{reverse('self_formatted_form_page')}?error="
                "Template file not found. The template may have been corrupted. "
                "Please re-create this format."
            )

        # Determine filename safely - use backup name or fallback
        template_name = ""
        try:
            template_name = fmt.template_file.name if fmt.template_file else ""
        except Exception:
            pass
        if not template_name:
            template_name = getattr(fmt, 'template_file_name', '') or f"template_{fmt.pk}.docx"
        template_name = os.path.basename(template_name)

        mem = io.BytesIO(data)
        uploaded = InMemoryUploadedFile(
            mem,
            field_name="template_file",
            name=template_name,
            content_type="application/octet-stream",
            size=len(data),
            charset=None,
        )
        try:
            return _fill_template_file(uploaded, placeholder_map)
        except Exception as e:
            logger.error(f"Template fill failed in use_format: {e}")
            return HttpResponse("Failed to fill template. Please check the template file.", status=500)

    return HttpResponseNotAllowed(["GET", "POST"])


@org_required
def self_formatted_delete_format(request, pk):
    """
    Delete a saved format (and its underlying template file).
    Respects is_locked flag - locked templates require explicit unlock first.
    """
    if request.method != "POST":
        return HttpResponseNotAllowed(["POST"])

    fmt = _safe_get_template(pk, request.user)
    if fmt is None:
        raise Http404
    
    # Check if format is locked (backwards compatible - default to True if field missing)
    is_locked = getattr(fmt, 'is_locked', True)
    if is_locked:
        # Check for unlock confirmation in request
        confirm_delete = request.POST.get('confirm_delete', '').lower()
        if confirm_delete != 'yes':
            return redirect(
                f"{reverse('self_formatted_form_page')}?error="
                "Format is locked. Click the lock icon to unlock before deleting."
            )

    template = fmt.template_file
    storage = template.storage if template else None
    name = template.name if template else None

    fmt.delete()

    if storage and name and storage.exists(name):
        storage.delete(name)

    return redirect("self_formatted_form_page")


@org_required
def self_formatted_edit_format(request, pk):
    """Edit an existing SelfFormattedTemplate.
    GET: show edit form
    POST: apply updates (name, description, optional new template file, custom placeholders)
    """
    fmt = _safe_get_template(pk, request.user)
    if fmt is None:
        raise Http404

    if request.method == "GET":
        preview_text = None
        template_url = None
        try:
            if fmt.template_file and fmt.template_file.name:
                template_url = fmt.template_file.url
                # attempt lightweight preview depending on extension
                name = fmt.template_file.name.lower()
                with fmt.template_file.open('rb') as f:
                    data = f.read()

                if name.endswith('.txt') or name.endswith('.csv'):
                    preview_text = data.decode('utf-8', errors='replace')[:4000]
                elif name.endswith('.docx'):
                    try:
                        from docx import Document
                        mem = io.BytesIO(data)
                        doc = Document(mem)
                        paras = [p.text for p in doc.paragraphs if p.text]
                        preview_text = '\n'.join(paras)[:4000]
                    except Exception:
                        preview_text = None
                elif name.endswith('.pdf'):
                    try:
                        from PyPDF2 import PdfReader
                        mem = io.BytesIO(data)
                        reader = PdfReader(mem)
                        text_parts = []
                        if reader.pages:
                            text_parts.append(reader.pages[0].extract_text() or '')
                        preview_text = '\n'.join(text_parts)[:4000]
                    except Exception:
                        preview_text = None
                elif name.endswith('.xlsx') or name.endswith('.xlsm'):
                    try:
                        from openpyxl import load_workbook
                        mem = io.BytesIO(data)
                        wb = load_workbook(mem, read_only=True, data_only=True)
                        sheet = wb.active
                        rows = []
                        for r in sheet.iter_rows(min_row=1, max_row=8, max_col=6, values_only=True):
                            rows.append('\t'.join([str(c) if c is not None else '' for c in r]))
                        preview_text = '\n'.join(rows)
                    except Exception:
                        preview_text = None
        except Exception:
            preview_text = None

        return render(request, "core/self_formatted_edit.html", {
            "format": fmt,
            "preview_text": preview_text,
            "template_url": template_url,
        })

    # POST
    name = request.POST.get("format_name", "").strip()
    description = request.POST.get("format_description", "").strip()
    raw_custom = request.POST.get("format_custom_placeholders", "").strip()
    new_template = request.FILES.get("format_template_file")

    if not name:
        return render(request, "core/self_formatted_edit.html", {
            "format": fmt,
            "error_message": "Name is required.",
        }, status=400)

    # Replace template file if new file provided
    old_name = None
    storage = None
    if new_template:
        if fmt.template_file and fmt.template_file.name:
            old_name = fmt.template_file.name
            storage = fmt.template_file.storage
        fmt.template_file = new_template
        # Update backup with new template content (backwards compatible)
        try:
            new_template.seek(0)
            backup_content = new_template.read()
            fmt.template_file_backup = backup_content
            fmt.template_file_name = new_template.name
            fmt.template_file_size = len(backup_content)
            new_template.seek(0)
        except Exception:
            pass  # Fields don't exist yet in DB schema

    fmt.name = name
    fmt.description = description
    fmt.custom_placeholders = raw_custom
    fmt.save()

    # delete old template file from storage after saving new one
    try:
        if old_name and storage and storage.exists(old_name):
            storage.delete(old_name)
    except Exception:
        pass

    return redirect("self_formatted_form_page")


@org_required
@require_POST
def self_formatted_toggle_lock(request, pk):
    """
    Toggle the lock status of a saved format.
    Locked formats require extra confirmation to delete.
    Returns JSON response for AJAX calls.
    """
    fmt = _safe_get_template(pk, request.user)
    if fmt is None:
        raise Http404
    
    # Check if is_locked field exists (backwards compatible)
    if not hasattr(fmt, 'is_locked'):
        return JsonResponse({
            'success': False,
            'message': 'Lock feature not available. Please run database migrations.'
        }, status=400)
    
    try:
        fmt.is_locked = not fmt.is_locked
        fmt.save(update_fields=['is_locked'])
        
        return JsonResponse({
            'success': True,
            'is_locked': fmt.is_locked,
            'message': f"Format {'locked' if fmt.is_locked else 'unlocked'} successfully"
        })
    except Exception as e:
        return JsonResponse({
            'success': False,
            'message': f'Failed to toggle lock: {str(e)}'
        }, status=500)


@org_required
@require_POST
def self_formatted_restore_backup(request, pk):
    """
    Restore a format's template file from database backup.
    Used when file storage fails but backup exists.
    """
    fmt = _safe_get_template(pk, request.user)
    if fmt is None:
        raise Http404
    
    # Check if backup field exists (backwards compatible)
    template_file_backup = getattr(fmt, 'template_file_backup', None)
    if not template_file_backup:
        return JsonResponse({
            'success': False,
            'message': 'No backup available for this format'
        }, status=400)
    
    try:
        success = fmt.restore_from_backup()
        
        if success:
            return JsonResponse({
                'success': True,
                'message': 'Template restored from backup successfully'
            })
        else:
            return JsonResponse({
                'success': False,
                'message': 'Failed to restore template from backup'
            }, status=500)
    except Exception as e:
        return JsonResponse({
            'success': False,
            'message': f'Restore failed: {str(e)}'
        }, status=500)


def _round_numeric_value(val):
    """Round numeric strings to 2 decimal places. Handles Rs. prefix, commas, /- suffix."""
    if not val or not isinstance(val, str):
        return val
    s = val.strip()
    # Extract number part, preserving prefix/suffix
    prefix = ""
    suffix = ""
    # Check for Rs. prefix
    m = re.match(r'^(Rs\.?\s*)', s, re.I)
    if m:
        prefix = m.group(1)
        s = s[m.end():]
    # Check for /- suffix
    if s.endswith('/-'):
        suffix = '/-'
        s = s[:-2]
    s = s.strip()
    # Remove commas for parsing
    cleaned = s.replace(',', '')
    if not cleaned:
        return val
    try:
        num = float(cleaned)
        # Only round if it has more than 2 decimal places
        rounded = round(num, 2)
        if ',' in s:
            # Preserve comma formatting (Indian style)
            int_part = int(rounded)
            dec_part = rounded - int_part
            # Format with commas
            int_str = f"{int_part:,}"
            if dec_part:
                result = f"{int_str}.{str(round(dec_part, 2))[2:]:0<2}"
            else:
                result = int_str
        else:
            if rounded == int(rounded):
                result = str(int(rounded))
            else:
                result = f"{rounded:.2f}"
        return f"{prefix}{result}{suffix}"
    except (ValueError, TypeError):
        return val


@login_required(login_url='login')
def self_formatted_progress_report(request):
    """
    Progress Report: accept multiple source files, extract labels from each,
    and produce a single Excel workbook with one row per work.
    Excel files with multiple sheets produce one row per sheet.
    """
    if request.method != "POST":
        return HttpResponseNotAllowed(["POST"])

    files = request.FILES.getlist("source_files")
    if not files:
        return redirect(
            f"{reverse('self_formatted_form_page')}?error=Please+upload+at+least+one+source+file"
        )

    from .utils import _apply_print_settings

    COLUMNS = [
        "Sr.No",
        "Name of Work",
        "Name of Agency",
        "Administrative Sanction",
        "Technical Sanction",
        "Agreement Number",
        "Estimate Amount",
        "Amount Paid",
        "M.B No Details",
        "DOI",
        "DOC",
        "DOMR",
        "DOBR",
        "Remarks",
    ]
    LABEL_KEYS = [
        "name_of_work",
        "agency",
        "admin_sanction",
        "tech_sanction",
        "agreement",
        "estimate_amount",
        "amount_paid",
        "mb_details",
        "doi",
        "doc",
        "domr",
        "dobr",
    ]

    wb = Workbook()
    ws = wb.active
    ws.title = "Progress Report"

    # Styles
    thin = Side(border_style="thin", color="000000")
    border_all = Border(left=thin, right=thin, top=thin, bottom=thin)
    header_fill = PatternFill("solid", fgColor="FFC8C8C8")
    header_font = Font(bold=True, size=11)
    cell_font = Font(size=10)
    wrap_align = Alignment(horizontal="left", vertical="top", wrap_text=True)
    center_align = Alignment(horizontal="center", vertical="center")

    # Header row
    for col_idx, col_name in enumerate(COLUMNS, start=1):
        cell = ws.cell(row=1, column=col_idx, value=col_name)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = border_all
        cell.alignment = center_align

    # Data rows — each file may yield multiple works (Excel with multiple sheets)
    row_num = 2
    sr_no = 1
    errors = []
    for f in files:
        try:
            works = _extract_labels_per_work(f)
        except Exception as e:
            logger.error(f"Progress report: failed to extract from {f.name}: {e}")
            errors.append(f.name)
            continue

        for source_name, labels in works:
            # Skip entries with no useful data
            if not any(labels.get(k) for k in LABEL_KEYS):
                continue

            # Sr.No
            ws.cell(row=row_num, column=1, value=sr_no).font = cell_font
            ws.cell(row=row_num, column=1).border = border_all
            ws.cell(row=row_num, column=1).alignment = center_align

            # Data columns (columns 2..13)
            for col_offset, key in enumerate(LABEL_KEYS):
                raw_val = labels.get(key, "")
                # Round numeric values to 2 decimals
                cell_val = _round_numeric_value(raw_val)
                cell = ws.cell(row=row_num, column=col_offset + 2, value=cell_val)
                cell.font = cell_font
                cell.border = border_all
                cell.alignment = wrap_align

            # Remarks column (last column = 14) — based on cc_header
            remarks = ""
            cc = (labels.get("cc_header") or "").strip()
            if cc:
                cc_low = cc.lower()
                if "workslip" in cc_low or "work slip" in cc_low:
                    remarks = "Workslip was submitted"
                elif "final" in cc_low:
                    remarks = "Final Bill submitted"
                else:
                    remarks = f"{cc} submitted"
            remarks_cell = ws.cell(row=row_num, column=len(COLUMNS), value=remarks)
            remarks_cell.font = cell_font
            remarks_cell.border = border_all
            remarks_cell.alignment = wrap_align

            row_num += 1
            sr_no += 1

    # Column widths
    col_widths = [6, 40, 25, 25, 25, 18, 16, 16, 25, 12, 12, 12, 12, 28]
    for i, w in enumerate(col_widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w

    _apply_print_settings(wb, landscape=True)

    # Build response
    resp = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    resp["Content-Disposition"] = 'attachment; filename="Progress_Report.xlsx"'
    wb.save(resp)
    return resp



def _extract_estimate_sheet_for_covering_letter(ws):
    """
    Inspect a single worksheet and return (work_name, amount) if it looks like
    an estimate summary sheet, or None if it should be skipped (Datas/items sheets).

    Sheet-skip rules:
      - Sheet name contains 'datas', 'data', 'groups', 'group', 'master', 'item'.
    Content-skip rule:
      - If no recognisable estimate label (name_of_work) is found.

    Amount logic: rightmost filled column, bottommost row that holds a numeric value.
    """
    sheet_name_lower = (ws.title or '').lower().strip()
    _SKIP_NAMES = ('datas', 'data', 'groups', 'group', 'master', 'items', 'item')
    if any(sheet_name_lower == k or sheet_name_lower.startswith(k) for k in _SKIP_NAMES):
        return None

    max_r = min(ws.max_row or 0, 150)
    max_c = min(ws.max_column or 0, 20)
    if max_r == 0:
        return None

    # Build text lines for label extraction (same logic as _extract_labels_per_work)
    lines = []
    for r in range(1, max_r + 1):
        vals = []
        for c in range(1, max_c + 1):
            v = ws.cell(row=r, column=c).value
            if v is not None:
                vals.append(str(v).strip())
        if vals:
            if len(vals) == 2:
                lbl = vals[0].lower().strip()
                if any(x in lbl for x in ['name', 'work', 'sanction', 'amount', 'contractor',
                                           'nit', 'estimate', 'period', 'address', 'premium']):
                    lines.append(f"{vals[0]}: {vals[1]}")
                else:
                    lines.append(' '.join(vals))
            else:
                lines.append(' '.join(vals))

    labels = _extract_labels_from_lines(lines)
    work_name = labels.get('name_of_work', '').strip()
    if not work_name:
        return None  # No estimate label found — skip this sheet

    # Amount: scan bottom-up, right-to-left for the first positive numeric value
    amount = 0.0
    for r in range(max_r, 0, -1):
        for c in range(max_c, 0, -1):
            v = ws.cell(row=r, column=c).value
            if v is None:
                continue
            cleaned = re.sub(r'^Rs\.?\s*', '', str(v).strip(), flags=re.I)
            cleaned = cleaned.replace(',', '').rstrip('/-').strip()
            try:
                num_val = float(cleaned)
                if num_val > 0:
                    amount = round(num_val, 2)
                    break
            except (ValueError, TypeError):
                continue
        if amount > 0:
            break

    return (work_name, amount)


@login_required(login_url='login')
def self_formatted_bulk_covering_letter(request):
    """
    Bulk Covering Letter: accept multiple source files (or multi-sheet Excel),
    extract work name + amount from each, and produce a single forwarding letter
    Word document with all estimates in the table.
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    if request.method != 'POST':
        return HttpResponseNotAllowed(['POST'])

    files = request.FILES.getlist('source_files')
    if not files:
        return redirect(
            f"{reverse('self_formatted_form_page')}?error=Please+upload+at+least+one+source+file"
        )

    # Extract one (work_name, amount) entry per file / per sheet.
    # For Excel: only process estimate-summary sheets (skip Datas/items sheets).
    # Amount: rightmost filled column, bottommost numeric value.
    works = []
    for f in files:
        try:
            filename = f.name or ''
            ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''
            if ext in ('xlsx', 'xlsm'):
                wb = load_workbook(f, data_only=True)
                for ws in wb.worksheets:
                    result = _extract_estimate_sheet_for_covering_letter(ws)
                    if result is not None:
                        works.append({'name': result[0], 'amount': result[1]})
            else:
                # Non-Excel: use existing label extraction
                labels, _lines = _extract_labels_from_source_file(f)
                name = (labels.get('name_of_work') or labels.get('cc_header') or filename).strip()
                raw = (labels.get('estimate_amount') or labels.get('amount') or
                       labels.get('admin_sanction_amount') or '0')
                try:
                    amt = float(str(raw).replace(',', '').replace('Rs.', '').replace('\u20b9', '').strip())
                except Exception:
                    amt = 0.0
                if name:
                    works.append({'name': name, 'amount': amt})
        except Exception as e:
            logger.error(f"Bulk covering letter: failed to extract from {f.name}: {e}")
            continue

    if not works:
        return redirect(
            f"{reverse('self_formatted_form_page')}?error=Could+not+extract+data+from+uploaded+files"
        )

    financial_year = _get_current_financial_year()
    today = timezone.now().date()
    letter_settings = _get_letter_settings(request.user)
    placeholder_color = RGBColor(169, 169, 169)
    n = len(works)
    grand_total = sum(w['amount'] for w in works)

    doc = Document()

    normal_style = doc.styles['Normal']
    normal_style.paragraph_format.space_after = Pt(4)
    normal_style.paragraph_format.space_before = Pt(0)

    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # --- Header ---
    h1 = doc.add_paragraph()
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h1.add_run((letter_settings.government_name.upper() if letter_settings and letter_settings.government_name else '[GOVERNMENT / ORGANIZATION NAME]'))
    r.font.bold = True
    r.font.size = Pt(14)
    if not (letter_settings and letter_settings.government_name):
        r.font.color.rgb = placeholder_color

    h2 = doc.add_paragraph()
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h2.paragraph_format.space_after = Pt(8)
    r2 = h2.add_run((letter_settings.department_name.upper() if letter_settings and letter_settings.department_name else '[DEPARTMENT NAME]'))
    r2.font.bold = True
    r2.font.size = Pt(13)
    if not (letter_settings and letter_settings.department_name):
        r2.font.color.rgb = placeholder_color

    # --- From / To table ---
    ft = doc.add_table(rows=1, cols=2)
    ft.autofit = True

    from_para = ft.cell(0, 0).paragraphs[0]
    from_para.add_run('From: -\n').font.bold = True
    if letter_settings and letter_settings.officer_name:
        nq = letter_settings.officer_name
        if letter_settings.officer_qualification:
            nq += f', {letter_settings.officer_qualification}'
        from_para.add_run(f'{nq},\n')
    else:
        r_ = from_para.add_run('[Officer Name, Qualification],\n')
        r_.font.color.rgb = placeholder_color
    if letter_settings and letter_settings.officer_designation:
        from_para.add_run(f'{letter_settings.officer_designation},\n')
    else:
        r_ = from_para.add_run('[Designation],\n')
        r_.font.color.rgb = placeholder_color
    if letter_settings and (letter_settings.sub_division or letter_settings.office_address):
        sub_addr = letter_settings.sub_division or ''
        if letter_settings.office_address:
            sub_addr += (f', {letter_settings.office_address}' if sub_addr else letter_settings.office_address)
        from_para.add_run(f'{sub_addr}.')
    else:
        r_ = from_para.add_run('[Sub Division, Office Address].')
        r_.font.color.rgb = placeholder_color

    to_para = ft.cell(0, 1).paragraphs[0]
    to_para.add_run('To,\n').font.bold = True
    if letter_settings and letter_settings.recipient_designation:
        to_para.add_run(f'{letter_settings.recipient_designation},\n')
    else:
        r_ = to_para.add_run('[Officer Designation],\n')
        r_.font.color.rgb = placeholder_color
    if letter_settings and letter_settings.recipient_division:
        to_para.add_run(f'{letter_settings.recipient_division},\n')
    else:
        r_ = to_para.add_run('[Division Name],\n')
        r_.font.color.rgb = placeholder_color
    if letter_settings and letter_settings.recipient_address:
        to_para.add_run(f'{letter_settings.recipient_address}.')
    else:
        r_ = to_para.add_run('[Address].')
        r_.font.color.rgb = placeholder_color

    # --- Lr. No. / Date ---
    lr_para = doc.add_paragraph()
    lr_para.paragraph_format.space_before = Pt(8)
    lr_run = lr_para.add_run('Lr No. ')
    lr_run.font.bold = True
    lr_run.font.underline = True
    if letter_settings and letter_settings.office_code:
        r_ = lr_para.add_run(letter_settings.office_code)
        r_.font.bold = True
        r_.font.underline = True
    else:
        r_ = lr_para.add_run('[Office Code]')
        r_.font.color.rgb = placeholder_color
        r_.font.bold = True
        r_.font.underline = True
    r_ = lr_para.add_run(f'/{financial_year}/          ')
    r_.font.bold = True
    r_.font.underline = True
    r_ = lr_para.add_run(f'\t\t\t\t\tDate:-    - {today.strftime("%m")} - {today.year}.')
    r_.font.bold = True
    r_.font.underline = True

    sir_para = doc.add_paragraph()
    sir_para.paragraph_format.space_before = Pt(6)
    sir_para.add_run('Sir,')

    # --- Subject ---
    subj_para = doc.add_paragraph()
    subj_para.paragraph_format.space_before = Pt(6)
    r_ = subj_para.add_run('Sub:-')
    r_.font.underline = True
    subj_para.add_run('\t')
    subj_para.add_run(f'Submission of Estimates for the year {financial_year}')
    subj_para.add_run('  -  Request for obtaining Administrative Sanction  -  Regarding.')

    # --- Reference ---
    ref_para = doc.add_paragraph()
    r_ = ref_para.add_run('Ref:-')
    r_.font.underline = True
    ref_para.add_run('\tMemo No.')
    r_ = ref_para.add_run('[Reference Number]')
    r_.font.color.rgb = placeholder_color
    r_.font.underline = True
    ref_para.add_run(f'/{financial_year} Dt.')
    r_ = ref_para.add_run('[DD.MM.YYYY]')
    r_.font.color.rgb = placeholder_color
    r_.font.underline = True

    stars = doc.add_paragraph()
    stars.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stars.add_run('**.**')

    # --- Body ---
    body_para = doc.add_paragraph()
    body_para.add_run('With reference to the subject cited, I submit here ')
    r_ = body_para.add_run(f'with  {n}')
    r_.font.underline = True
    body_para.add_run(' No. estimates for the following works for the amounts specified.')

    # --- Estimates table ---
    table = doc.add_table(rows=n + 2, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for cell in table.columns[0].cells:
        cell.width = Inches(0.5)
    for cell in table.columns[1].cells:
        cell.width = Inches(4.5)
    for cell in table.columns[2].cells:
        cell.width = Inches(1.5)

    # Header row
    hdr = table.rows[0].cells
    hdr[0].text = 'Sl.\nNo'
    hdr[1].text = 'Name of work'
    hdr[2].text = 'Amount'
    for cell in hdr:
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True

    # Data rows
    for i, w in enumerate(works, start=1):
        row = table.rows[i].cells
        row[0].text = str(i)
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[1].text = w['name']
        row[2].text = f"Rs.{_format_indian_number(w['amount'])}"
        row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Total row
    total_row = table.rows[n + 1].cells
    total_row[0].merge(total_row[1])
    total_para = total_row[0].paragraphs[0]
    total_run = total_para.add_run('Total')
    total_run.font.bold = True
    total_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    total_row[2].text = f"Rs.{_format_indian_number(grand_total)}"
    for run in total_row[2].paragraphs[0].runs:
        run.font.bold = True
    total_row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # --- Spec paragraph ---
    spec_para = doc.add_paragraph()
    spec_para.paragraph_format.space_before = Pt(6)
    spec_para.add_run("Specification report accompanying the estimates explains the necessity and provisions made therein in detail.")

    # --- Request paragraph ---
    req_para = doc.add_paragraph()
    req_para.add_run('I request the ')
    if letter_settings and letter_settings.superior_designation:
        req_para.add_run(letter_settings.superior_designation)
    else:
        r_ = req_para.add_run('[Superior Officer Designation]')
        r_.font.color.rgb = placeholder_color
    req_para.add_run(' to kindly arrange to obtain administrative sanction for the above estimates and arrange to finalize the agencies at the earliest for taking up the works.')

    # --- Enclosure ---
    enc_para = doc.add_paragraph()
    enc_para.paragraph_format.space_before = Pt(6)
    enc_para.add_run('Enclosure: -')
    doc.add_paragraph(f'Estimates  - {n} Nos.')

    # --- Signature ---
    sign_para = doc.add_paragraph()
    sign_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sign_para.paragraph_format.space_before = Pt(14)
    sign_para.add_run('Yours faithfully,')

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    title_para.paragraph_format.space_before = Pt(18)
    if letter_settings and letter_settings.officer_designation:
        r_ = title_para.add_run(f'{letter_settings.officer_designation}\n')
        r_.font.bold = True
    else:
        r_ = title_para.add_run('[Officer Designation]\n')
        r_.font.bold = True
        r_.font.color.rgb = placeholder_color
    if letter_settings and letter_settings.sub_division:
        title_para.add_run(f'{letter_settings.sub_division},\n')
    else:
        r_ = title_para.add_run('[Sub Division Name],\n')
        r_.font.color.rgb = placeholder_color
    if letter_settings and letter_settings.office_address:
        title_para.add_run(f'{letter_settings.office_address}.')
    else:
        r_ = title_para.add_run('[Office Address].')
        r_.font.color.rgb = placeholder_color

    # --- Copy To ---
    copy_para = doc.add_paragraph()
    copy_para.paragraph_format.space_before = Pt(8)
    copy_para.add_run('Copy to the ')
    if letter_settings and (letter_settings.copy_to_designation or letter_settings.copy_to_section):
        copy_text = letter_settings.copy_to_designation or ''
        if letter_settings.copy_to_section:
            copy_text += (f', {letter_settings.copy_to_section}' if copy_text else letter_settings.copy_to_section)
        copy_para.add_run(copy_text)
    else:
        r_ = copy_para.add_run('[Officer Designation, Section Name]')
        r_.font.color.rgb = placeholder_color
    copy_para.add_run(' for information.')

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename="Bulk_Covering_Letter.docx"'
    doc.save(response)
    return response


# ==========================
#  TEMPORARY WORKS MODULE
#  (completely separate from New Estimate)
# ==========================


