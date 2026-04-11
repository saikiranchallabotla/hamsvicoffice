# Auto-generated from core/views.py split
import json
import os
import re
import logging
from copy import copy

import inflect
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, Border, Side, PatternFill
from django.utils import timezone
from django.urls import reverse
from django.core.files.uploadedfile import InMemoryUploadedFile
from django.contrib.auth.decorators import login_required

from django.conf import settings
from django.http import HttpResponse, JsonResponse, HttpResponseNotAllowed
from django.shortcuts import render, redirect, get_object_or_404
from django.template.loader import render_to_string
from django.views.decorators.http import require_POST
from django.utils.crypto import get_random_string

import io
from io import BytesIO
from difflib import SequenceMatcher

from ..models import Project, SelfFormattedTemplate, Estimate, Organization, Membership, Upload, Job, OutputFile, LetterSettings
from ..decorators import org_required, role_required

logger = logging.getLogger(__name__)
from ..tasks import process_excel_upload, generate_bill_pdf, generate_workslip_pdf, generate_bill_document_task
from ..utils_excel import load_backend, copy_block_with_styles_and_formulas, build_temp_day_rates

p_engine = inflect.engine()
BILL_TEMPLATES_DIR = os.path.join(settings.BASE_DIR, "core", "templates", "core", "bill_templates")
_inflect_engine = inflect.engine()

from .utils import (get_org_from_request, check_org_access,
    _apply_print_settings, _format_indian_number, _number_to_words_rupees,
    _get_current_financial_year, _get_current_date_formatted,
    _get_letter_settings)

@login_required(login_url='login')
def my_subscription(request):
    return render(request, "core/my_subscription.html")

@org_required
def my_projects(request):
    org = get_org_from_request(request)
    projects = Project.objects.for_org(org)
    return render(request, "core/my_projects.html", {"projects": projects})

@org_required
def create_project(request):
    org = get_org_from_request(request)
    if request.method == "POST":
        name = request.POST.get("project_name")
        if name:
            Project.objects.get_or_create(organization=org, name=name)
    return redirect("my_projects")

@login_required(login_url='login')
def datas(request):
    """
    Landing page for 'New Estimate'.

    - Clears current selection (as you had before)
    - Reads ?work_type=original/repair from URL and stores in session
    - Defaults to 'original' if nothing selected
    """
    # Always start fresh when entering Datas
    request.session["fetched_items"] = []
    request.session["current_project_name"] = None
    request.session["qty_map"] = {}
    request.session["unit_map"] = {}
    request.session["work_name"] = ""
    request.session["grand_total"] = ""
    request.session["selected_backend_id"] = None  # Clear any previous backend selection
    request.session["current_saved_work_id"] = None  # Clear any resumed work so new estimate doesn't show "Update Work"
    # Clear uploaded custom items
    request.session["uploaded_items"] = []
    request.session["uploaded_file_id"] = None
    request.session["uploaded_item_blocks"] = {}
    request.session["uploaded_sheet_name"] = ""

    mode = request.GET.get("work_type")

    if mode in ("original", "repair"):
        request.session["work_type"] = mode

    # If nothing in session yet, default to original
    if "work_type" not in request.session:
        request.session["work_type"] = "original"

    return render(
        request,
        "core/datas.html",
        {"work_type": request.session["work_type"]},
    )



# -----------------------
# STEP 1: SELECT PROJECT
# -----------------------
@login_required(login_url='login')
def select_project(request):
    org = get_org_from_request(request)
    projects = Project.objects.for_org(org)

    use_id = request.GET.get("use")
    if use_id:
        request.session["selected_project_id"] = use_id
        request.session["fetched_items"] = []
        return redirect("choose_category")

    if request.method == "POST":
        project_name = request.POST.get("project_name")
        if project_name:
            project, created = Project.objects.get_or_create(organization=org, name=project_name)
            request.session["selected_project_id"] = project.id
            request.session["fetched_items"] = []
            return redirect("choose_category")

    return render(request, "core/select_project.html", {"projects": projects})


# -----------------------
# STEP 2: SELECT CATEGORY
# -----------------------
@login_required(login_url='login')
def choose_category(request):
    return render(request, "core/choose_category.html")


# -----------------------
# STEP 3: GROUPS PAGE (redirects to items)
# (imports consolidated at top)
# -----------------------
@login_required(login_url='login')
def datas_groups(request, category):
    # NEW: remember work_type in session if passed in URL
    work_type = (request.GET.get("work_type") or "").lower()
    if work_type in ("original", "repair"):
        request.session["work_type"] = work_type

    # NEW: Store selected backend_id in session for consistent use throughout the flow
    backend_id = request.GET.get("backend_id")
    if backend_id:
        try:
            request.session["selected_backend_id"] = int(backend_id)
        except (ValueError, TypeError):
            request.session["selected_backend_id"] = None
    
    # Get backend_id from session for loading
    selected_backend_id = request.session.get("selected_backend_id")
    
    try:
        items_list, groups_map, _, ws_data, filepath = load_backend(
            category, settings.BASE_DIR, 
            backend_id=selected_backend_id,
            module_code='new_estimate',
            user=request.user if request.user.is_authenticated else None
        )
    except FileNotFoundError as e:
        logger.info(f"No backend data available for category {category} - showing Coming Soon")
        # Check if the other category has backends available
        other_category = 'electrical' if category == 'civil' else 'civil'
        other_available = False
        try:
            from core.utils_excel import get_available_backends_for_module
            other_backends = get_available_backends_for_module('new_estimate', other_category)
            other_available = len(other_backends) > 0
        except:
            pass
        return render(request, "core/coming_soon.html", {
            "category": category,
            "module_name": "New Estimate",
            "other_category": other_category,
            "other_category_available": other_available,
        })
    except Exception as e:
        logger.error(f"Error loading backend for {category}: {e}")
        return render(request, "core/error.html", {
            "error_title": "Loading Error",
            "error_message": f"Could not load backend data: {str(e)}",
        })
    
    groups = sorted(groups_map.keys(), key=lambda s: s.lower())

    if not groups:
        return render(request, "core/groups.html", {
            "category": category,
            "groups": [],
            "error": "No groups found in backend Excel.",
        })

    default_group = request.GET.get("group") or groups[0]
    return redirect("datas_items", category=category, group=default_group)



# -----------------------
# STEP 4: ITEMS IN GROUP (3-panel UI)
# -----------------------
@login_required(login_url='login')
def datas_items(request, category, group):
    from core.utils_excel import get_available_backends_for_module
    
    # Check if backend_id is passed in URL (for switching backends on this page)
    url_backend_id = request.GET.get("backend_id")
    if url_backend_id:
        try:
            bid = int(url_backend_id)
            request.session["selected_backend_id"] = bid
            # Persist choice to DB so it survives re-login/redeployment
            if request.user.is_authenticated:
                try:
                    from subscriptions.models import ModuleBackend
                    from accounts.models import UserBackendPreference
                    _backend = ModuleBackend.objects.filter(pk=bid, is_active=True).first()
                    if _backend:
                        UserBackendPreference.set_user_backend(request.user, _backend)
                except Exception:
                    pass
        except (ValueError, TypeError):
            pass

    # Get backend_id from session for consistent loading throughout the flow
    selected_backend_id = request.session.get("selected_backend_id")

    # Initialize from user's saved preference if session has no selection
    if selected_backend_id is None and request.user.is_authenticated:
        try:
            from accounts.models import UserBackendPreference
            _pref_backend = UserBackendPreference.get_user_backend(request.user, 'new_estimate', category)
            if _pref_backend:
                selected_backend_id = _pref_backend.pk
                request.session["selected_backend_id"] = selected_backend_id
        except Exception:
            pass

    # Get available backends for the dropdown
    try:
        available_backends = get_available_backends_for_module('new_estimate', category)
    except Exception as e:
        logger.error(f"Error getting available backends: {e}")
        available_backends = []
    
    try:
        items_list, groups_map, backend_units_map, ws_data, filepath = load_backend(
            category, settings.BASE_DIR,
            backend_id=selected_backend_id,
            module_code='new_estimate',
            user=request.user if request.user.is_authenticated else None
        )
    except FileNotFoundError as e:
        logger.info(f"No backend data available for category {category} - showing Coming Soon")
        # Check if the other category has backends available
        other_category = 'electrical' if category == 'civil' else 'civil'
        other_available = False
        try:
            other_backends = get_available_backends_for_module('new_estimate', other_category)
            other_available = len(other_backends) > 0
        except:
            pass
        return render(request, "core/coming_soon.html", {
            "category": category,
            "module_name": "New Estimate",
            "other_category": other_category,
            "other_category_available": other_available,
        })
    except Exception as e:
        logger.error(f"Error loading backend for {category}: {e}")
        return render(request, "core/error.html", {
            "error_title": "Loading Error",
            "error_message": f"Could not load backend data: {str(e)}",
        })

    groups = sorted(groups_map.keys(), key=lambda s: s.lower())

    group_items = groups_map.get(group, [])
    detected_names = {i["name"] for i in items_list}
    display_items = [name for name in group_items if name in detected_names]

    wb_vals = load_workbook(filepath, data_only=True)
    ws_vals = wb_vals["Master Datas"]

    item_rates = {}
    item_descs = {}
    for info in items_list:
        name = info["name"]
        start_row = info["start_row"]
        end_row = info["end_row"]
        rate = None
        for r in range(end_row, start_row - 1, -1):
            val = ws_vals.cell(row=r, column=10).value  # column J
            if val not in (None, ""):
                rate = val
                break
        item_rates[name] = rate
        # Scan rows below header for the best description text
        # Check columns B(2) and D(4) in rows start_row+1 through end_row
        desc_candidates = []
        for r in range(start_row + 1, min(end_row + 1, start_row + 4)):
            for col in (4, 2):  # column D first, then B
                cell_val = ws_vals.cell(row=r, column=col).value
                if cell_val:
                    txt = str(cell_val).strip()
                    # Skip if it's just a number, rate value, or unit
                    if txt and len(txt) > 5 and not txt.replace('.', '').replace(',', '').isdigit():
                        desc_candidates.append(txt)
        if desc_candidates:
            # Pick the longest candidate as the most detailed description
            item_descs[name] = max(desc_candidates, key=len)
        else:
            item_descs[name] = name

    item_to_group = {}
    for grp_name, item_list_in_grp in groups_map.items():
        for nm in item_list_in_grp:
            item_to_group.setdefault(nm, grp_name)

    # Merge uploaded item data from session (uploaded items aren't in backend)
    uploaded_items_in_session = set(request.session.get('uploaded_items', []))
    session_saved_rates = request.session.get('item_rates', {})
    session_saved_units = request.session.get('item_units', {})
    session_saved_descs = request.session.get('item_descs', {})
    for uname in uploaded_items_in_session:
        if uname not in item_rates:
            item_rates[uname] = session_saved_rates.get(uname)
        if uname not in item_descs:
            item_descs[uname] = session_saved_descs.get(uname, uname)

    def units_for(name):
        # Uploaded items: use saved session unit
        if name in uploaded_items_in_session:
            u = session_saved_units.get(name, 'Nos')
            return (u, u)
        # First check backend_units_map (Column D from backend Excel)
        backend_unit = backend_units_map.get(name, "")
        if backend_unit:
            return (backend_unit, backend_unit)  # Use same for both plural and singular display
        # Fall back to group-based defaults
        grp_name = item_to_group.get(name, "")
        if grp_name in ("Piping", "Wiring & Cables"):
            return ("Mtrs", "Mtr")
        elif grp_name == "Points":
            return ("Pts", "Pt")
        else:
            return ("Nos", "No")

    # Build item subtypes map: items with ":" are subtypes
    # Group subtypes by their parent name (part before ":")
    # Support both " : " (with spaces) and ":" (without spaces)
    import re as _re
    _colon_re = _re.compile(r'\s*:\s*')

    def _has_colon(name):
        """Check if item name contains a colon separator (with or without spaces)."""
        return bool(_colon_re.search(name))

    def _split_parent(name):
        """Extract parent name from a colon-separated item name."""
        return _colon_re.split(name, 1)[0].strip()

    item_subtypes = {}  # parent_name -> [list of full subtype names]
    parent_items = set()  # items that have subtypes
    
    for name in display_items:
        if _has_colon(name):
            # This is a subtype - extract parent name
            parent_name = _split_parent(name)
            if parent_name not in item_subtypes:
                item_subtypes[parent_name] = []
            item_subtypes[parent_name].append(name)
            parent_items.add(parent_name)
    
    items_info = []
    seen_parents = set()
    for name in display_items:
        if _has_colon(name):
            # This is a subtype - check if we already added the parent
            parent_name = _split_parent(name)
            if parent_name not in seen_parents:
                # Add the parent item with subtypes info - serialize subtypes as JSON string
                subtypes_list = item_subtypes.get(parent_name, [])
                items_info.append({
                    "name": parent_name,
                    "rate": None,  # Parent doesn't have its own rate
                    "has_subtypes": True,
                    "subtypes": json.dumps(subtypes_list),
                    "subtypes_count": len(subtypes_list),
                })
                seen_parents.add(parent_name)
        else:
            # Regular item without subtypes
            items_info.append({
                "name": name,
                "rate": item_rates.get(name),
                "has_subtypes": False,
                "subtypes": "[]",
                "subtypes_count": 0,
            })


    fetched = request.session.get("fetched_items", [])

    # Normalize fetched_items: may contain dicts (from workslip session) or strings
    # Extract item name strings for the estimate view
    fetched_names = []
    for item in fetched:
        if isinstance(item, dict):
            fetched_names.append(item.get('item_name') or item.get('display_name') or item.get('name') or str(item))
        else:
            fetched_names.append(item)
    fetched = fetched_names

    qty_map = request.session.get("qty_map", {}) or {}
    unit_map = request.session.get("unit_map", {}) or {}
    work_name = request.session.get("work_name", "") or ""
    grand_total = request.session.get("grand_total", "") or ""

    estimate_rows = []
    # Build a serialisable rates map for session persistence so that
    # workslip generation can reuse the *exact* rates the user sees here.
    session_item_rates = {}
    session_item_units = {}
    for idx, name in enumerate(fetched, start=1):
        default_plural, singular = units_for(name)
        # Priority: 1) user-entered unit from UI, 2) backend_units_map default
        custom_unit = unit_map.get(name, "")
        display_unit = custom_unit if custom_unit else default_plural
        raw_rate = item_rates.get(name)
        estimate_rows.append({
            "sl": idx,
            "name": name,
            "rate": raw_rate,
            "unit": display_unit,
            "default_unit": default_plural,
            "qty": qty_map.get(name, ""),
        })
        # Store numeric rate for session (JSON-safe)
        try:
            session_item_rates[name] = float(raw_rate) if raw_rate is not None else 0.0
        except (ValueError, TypeError):
            session_item_rates[name] = 0.0
        session_item_units[name] = display_unit

    # Persist rates, units & descriptions so saved-works → workslip gets exact values
    request.session["item_rates"] = session_item_rates
    request.session["item_units"] = session_item_units
    request.session["item_descs"] = item_descs
    request.session.modified = True

    work_type = request.session.get("work_type", "original") or "original"
    excess_tp_percent = request.session.get("excess_tp_percent", "") or ""
    ls_special_name = request.session.get("ls_special_name", "") or ""
    ls_special_amount = request.session.get("ls_special_amount", "") or ""
    deduct_old_material = request.session.get("deduct_old_material", "") or ""
    
    return render(request, "core/items.html", {
        "category": category,
        "group": group,
        "groups": groups,
        "items_info": items_info,
        "fetched": fetched,
        "estimate_rows": estimate_rows,
        "work_name": work_name,
        "grand_total": grand_total,
        "work_type": work_type,
        "excess_tp_percent": excess_tp_percent,
        "ls_special_name": ls_special_name,
        "ls_special_amount": ls_special_amount,
        "deduct_old_material": deduct_old_material,
        "fetched_json": json.dumps(fetched),
        "available_backends": available_backends,
        "selected_backend_id": selected_backend_id,
        "item_descs_json": json.dumps(item_descs),
        "uploaded_items_json": json.dumps(list(uploaded_items_in_session)),
    })


# -----------------------
# STEP 5: FETCH / UN-FETCH ITEM (toggle)
# -----------------------
@login_required(login_url='login')
def fetch_item(request, category, group, item):
    fetched = request.session.get("fetched_items", []) or []

    if item in fetched:
        fetched.remove(item)
    else:
        fetched.append(item)

    request.session["fetched_items"] = fetched

    work_name = request.GET.get("work_name")
    if work_name is not None:
        request.session["work_name"] = work_name

    return redirect("datas_items", category=category, group=group)


# -----------------------
# AJAX TOGGLE ITEM (no page reload)
# -----------------------
@login_required(login_url='login')
def ajax_toggle_item(request, category):
    """
    AJAX endpoint to toggle an item in fetched list without page reload.
    POST with JSON: { "item": "item_name", "action": "add" or "remove", "work_name": "..." }
    Returns JSON: { "status": "ok", "fetched": [...], "action_taken": "added" or "removed", "item_info": {...} }
    """
    if request.method != "POST":
        return JsonResponse({"status": "error", "message": "POST required"}, status=405)
    
    try:
        if request.content_type and 'application/json' in request.content_type:
            data = json.loads(request.body.decode("utf-8") or "{}")
        else:
            data = request.POST
        
        item = data.get("item", "").strip()
        action = data.get("action", "toggle")  # "add", "remove", or "toggle"
        work_name = data.get("work_name")
        
        if not item:
            return JsonResponse({"status": "error", "message": "No item specified"}, status=400)
        
        fetched = request.session.get("fetched_items", []) or []
        action_taken = None
        
        if action == "add":
            if item not in fetched:
                fetched.append(item)
                action_taken = "added"
            else:
                action_taken = "already_exists"
        elif action == "remove":
            if item in fetched:
                fetched.remove(item)
                action_taken = "removed"
            else:
                action_taken = "not_found"
        else:  # toggle
            if item in fetched:
                fetched.remove(item)
                action_taken = "removed"
            else:
                fetched.append(item)
                action_taken = "added"
        
        request.session["fetched_items"] = fetched

        # If a removed item was an uploaded custom item, clean it from uploaded tracking
        if action_taken == "removed":
            uploaded = request.session.get("uploaded_items", [])
            if item in uploaded:
                uploaded.remove(item)
                request.session["uploaded_items"] = uploaded
                # Clean up associated data
                for key in ("item_rates", "item_units", "item_descs"):
                    d = request.session.get(key, {})
                    if isinstance(d, dict):
                        d.pop(item, None)
                uploaded_blocks = request.session.get("uploaded_item_blocks", {})
                if isinstance(uploaded_blocks, dict):
                    uploaded_blocks.pop(item, None)
                    request.session["uploaded_item_blocks"] = uploaded_blocks
                # If no uploaded items left, clear the upload file reference
                if not uploaded:
                    request.session["uploaded_file_id"] = None
                    request.session["uploaded_sheet_name"] = ""
                request.session.modified = True

        if work_name is not None:
            request.session["work_name"] = work_name
        
        # Get item info (rate, unit) for newly added items
        item_info = None
        if action_taken == "added":
            try:
                # Use same backend as the items page - get from session
                selected_backend_id = request.session.get("selected_backend_id")
                items_list, groups_map, units_map, ws_data, filepath = load_backend(
                    category, settings.BASE_DIR,
                    backend_id=selected_backend_id,
                    module_code='new_estimate',
                    user=request.user
                )
                
                # Get rate
                wb_vals = load_workbook(filepath, data_only=True)
                ws_vals = wb_vals["Master Datas"]
                
                item_rate = None
                for info in items_list:
                    if info["name"] == item:
                        start_row = info["start_row"]
                        end_row = info["end_row"]
                        for r in range(end_row, start_row - 1, -1):
                            val = ws_vals.cell(row=r, column=10).value  # column J
                            if val not in (None, ""):
                                item_rate = val
                                break
                        break
                
                # Get unit with smart fallback (same logic as datas_items)
                # Priority: 1) units_map from backend, 2) group-based defaults
                unit = units_map.get(item, "")
                if not unit:
                    # Find item's group for fallback
                    item_group = ""
                    for grp_name, grp_items in groups_map.items():
                        if item in grp_items:
                            item_group = grp_name
                            break
                    # Group-based defaults
                    if item_group in ("Piping", "Wiring & Cables", "Run of Mains", "Sheathed Cables", "U.G Cabling"):
                        unit = "Mtrs"
                    elif item_group == "Points":
                        unit = "Pts"
                    else:
                        unit = "Nos"
                
                item_info = {
                    "name": item,
                    "rate": item_rate,
                    "unit": unit
                }
                
                wb_vals.close()
            except Exception as e:
                # If we can't get item info, just return without it
                item_info = {"name": item, "rate": None, "unit": "Nos"}
        
        return JsonResponse({
            "status": "ok",
            "fetched": fetched,
            "action_taken": action_taken,
            "item": item,
            "item_info": item_info
        })
        
    except Exception as e:
        return JsonResponse({"status": "error", "message": str(e)}, status=400)


# -----------------------
# AJAX REORDER ITEMS
# -----------------------
@login_required(login_url='login')
def ajax_reorder_items(request, category):
    """
    AJAX endpoint to reorder fetched items list.
    POST with JSON: { "items": ["item1", "item2", ...] }
    Returns JSON: { "status": "ok", "fetched": [...] }
    """
    if request.method != "POST":
        return JsonResponse({"status": "error", "message": "POST required"}, status=405)
    
    try:
        if request.content_type and 'application/json' in request.content_type:
            data = json.loads(request.body.decode("utf-8") or "{}")
        else:
            return JsonResponse({"status": "error", "message": "JSON required"}, status=400)
        
        new_order = data.get("items", [])
        
        if not isinstance(new_order, list):
            return JsonResponse({"status": "error", "message": "items must be a list"}, status=400)
        
        # Validate: new_order should contain the same items as current fetched
        current_fetched = set(request.session.get("fetched_items", []) or [])
        new_order_set = set(new_order)
        
        # Only reorder if sets match (no items added/removed via this endpoint)
        if current_fetched == new_order_set:
            request.session["fetched_items"] = new_order
        else:
            # Allow partial reorder - use the intersection
            valid_items = [item for item in new_order if item in current_fetched]
            request.session["fetched_items"] = valid_items
        
        return JsonResponse({
            "status": "ok",
            "fetched": request.session["fetched_items"]
        })
        
    except Exception as e:
        return JsonResponse({"status": "error", "message": str(e)}, status=400)


# -----------------------
# AJAX UPLOAD CUSTOM ITEMS
# -----------------------
@login_required(login_url='login')
@require_POST
def ajax_upload_custom_items(request, category):
    """
    AJAX endpoint to upload an Excel file containing custom item blocks.
    Parses item blocks (yellow bg + red text) and adds them to the session
    alongside any backend-selected items.
    Returns JSON: { status, items: [{name, rate, unit, desc}], count, warnings }
    """
    from ..models import Upload
    from ..utils_excel import _extract_items_from_sheet, _determine_unit_from_heading

    try:
        uploaded_file = request.FILES.get('custom_items_file')
        if not uploaded_file:
            return JsonResponse({"status": "error", "message": "No file uploaded"}, status=400)

        if uploaded_file.size == 0:
            return JsonResponse({"status": "error", "message": "Uploaded file is empty"}, status=400)

        if uploaded_file.size > 10 * 1024 * 1024:
            return JsonResponse({"status": "error", "message": "File too large (max 10MB)"}, status=400)

        if not uploaded_file.name.endswith('.xlsx'):
            return JsonResponse({"status": "error", "message": "Only .xlsx files are supported"}, status=400)

        org = get_org_from_request(request)

        # Load workbook twice: formulas and values
        uploaded_file.seek(0)
        try:
            wb_formulas = load_workbook(uploaded_file, data_only=False)
        except Exception as e:
            return JsonResponse({"status": "error", "message": f"Failed to read Excel file: {e}"}, status=400)

        uploaded_file.seek(0)
        try:
            wb_values = load_workbook(uploaded_file, data_only=True)
        except Exception as e:
            return JsonResponse({"status": "error", "message": f"Failed to process Excel file: {e}"}, status=400)

        # Read units from Groups sheet if available
        upload_units_map = {}
        try:
            if "Groups" in wb_formulas.sheetnames:
                from ..utils_excel import read_groups
                _, upload_units_map = read_groups(wb_formulas["Groups"])
        except Exception:
            pass

        # Find item blocks across all sheets
        all_items = []          # [{name, rate, unit, desc}]
        all_item_blocks = {}    # {name: [start_row, end_row]}
        used_sheet_name = ""

        for sheet_name in wb_formulas.sheetnames:
            ws_src = wb_formulas[sheet_name]
            fetched_names, item_blocks = _extract_items_from_sheet(ws_src)
            if not fetched_names:
                continue

            if not used_sheet_name:
                used_sheet_name = sheet_name

            ws_vals = wb_values[sheet_name]

            for item_name in fetched_names:
                src_min, src_max = item_blocks[item_name]

                # Rate: last non-empty value in column J
                rate = None
                for r in range(src_max, src_min - 1, -1):
                    v = ws_vals.cell(row=r, column=10).value
                    if v not in (None, ""):
                        try:
                            rate = float(v)
                        except (ValueError, TypeError):
                            rate = None
                        break

                unit = _determine_unit_from_heading(item_name, upload_units_map)
                desc = str(ws_src.cell(row=src_min + 2, column=4).value or "").strip()

                all_items.append({
                    "name": item_name,
                    "rate": rate,
                    "unit": unit,
                    "desc": desc or item_name,
                })
                all_item_blocks[item_name] = [src_min, src_max]

        if not all_items:
            return JsonResponse({
                "status": "error",
                "message": "No item blocks found. Ensure item headers have yellow background and red text."
            }, status=400)

        # Remove old uploaded items from session if re-uploading
        old_uploaded = set(request.session.get('uploaded_items', []))
        if old_uploaded:
            fetched = request.session.get('fetched_items', [])
            request.session['fetched_items'] = [n for n in fetched if n not in old_uploaded]
            for n in old_uploaded:
                request.session.get('item_rates', {}).pop(n, None)
                request.session.get('item_units', {}).pop(n, None)
                request.session.get('item_descs', {}).pop(n, None)
                request.session.get('qty_map', {}).pop(n, None)

        # Detect name collisions with current backend items
        current_backend_items = set(request.session.get('fetched_items', []))
        warnings = []
        new_items = []
        for item in all_items:
            if item["name"] in current_backend_items:
                warnings.append(f"'{item['name']}' already selected from backend, skipped")
            else:
                new_items.append(item)

        if not new_items:
            return JsonResponse({
                "status": "error",
                "message": "All uploaded items already exist in your selection.",
                "warnings": warnings,
            }, status=400)

        # Save uploaded file via Upload model for persistence
        uploaded_file.seek(0)
        upload_obj = Upload.objects.create(
            organization=org,
            user=request.user,
            file=uploaded_file,
            filename=uploaded_file.name,
            file_size=uploaded_file.size,
            status='completed',
        )

        # Merge into session
        fetched = request.session.get('fetched_items', []) or []
        item_rates = request.session.get('item_rates', {}) or {}
        item_units = request.session.get('item_units', {}) or {}
        item_descs = request.session.get('item_descs', {}) or {}

        uploaded_names = []
        for item in new_items:
            name = item["name"]
            fetched.append(name)
            item_rates[name] = item["rate"]
            item_units[name] = item["unit"]
            item_descs[name] = item["desc"]
            uploaded_names.append(name)

        request.session['fetched_items'] = fetched
        request.session['item_rates'] = item_rates
        request.session['item_units'] = item_units
        request.session['item_descs'] = item_descs

        # Only keep blocks for items that were actually added (not skipped)
        filtered_blocks = {n: all_item_blocks[n] for n in uploaded_names if n in all_item_blocks}

        request.session['uploaded_items'] = uploaded_names
        request.session['uploaded_file_id'] = upload_obj.id
        request.session['uploaded_item_blocks'] = filtered_blocks
        request.session['uploaded_sheet_name'] = used_sheet_name
        request.session.modified = True

        return JsonResponse({
            "status": "ok",
            "items": [i for i in new_items],
            "count": len(new_items),
            "warnings": warnings,
        })

    except Exception as e:
        return JsonResponse({"status": "error", "message": str(e)}, status=500)


# -----------------------
# STEP 6: OUTPUT PANEL
# -----------------------
@login_required(login_url='login')
def output_panel(request, category):
    fetched = request.session.get("fetched_items", [])
    return render(request, "core/output.html", {
        "category": category,
        "items": fetched
    })


# local imports consolidated at top


@login_required(login_url='login')
def download_output(request, category):
    """
    Async Excel generation endpoint.
    
    POST with:
      - qty_map: JSON string of quantities
      - unit_map: JSON string of custom units per item
      - work_name: Name of work
      - work_type: "original" or "repair"
    
    Returns JSON with job_id and status_url for polling.
    Actual Excel generation happens asynchronously via Celery task.
    """
    fetched = request.session.get("fetched_items", [])
    if not fetched:
        return JsonResponse({"error": "No items selected"}, status=400)

    # Parse input
    item_qtys = {}
    item_units = {}
    work_name = ""
    grand_total = None
    
    # Initialize optional parameters with defaults (before POST check)
    excess_tp_enabled = False
    excess_tp_percent = None
    ls_special_enabled = False
    ls_special_name = None
    ls_special_amount = None
    deduct_old_material = None

    if request.method == "POST":
        qty_map_str = request.POST.get("qty_map", "")
        if qty_map_str:
            try:
                raw = json.loads(qty_map_str)
                if isinstance(raw, dict):
                    for k, v in raw.items():
                        try:
                            item_qtys[str(k)] = float(v)
                        except Exception:
                            continue
            except Exception:
                pass

        # Parse unit_map from POST
        unit_map_str = request.POST.get("unit_map", "")
        if unit_map_str:
            try:
                raw_units = json.loads(unit_map_str)
                if isinstance(raw_units, dict):
                    for k, v in raw_units.items():
                        item_units[str(k)] = str(v).strip()
            except Exception:
                pass

        work_name = (request.POST.get("work_name") or "").strip()
        
        # Parse grand_total from POST
        grand_total_str = request.POST.get("grand_total", "").strip()
        if grand_total_str:
            try:
                grand_total = float(grand_total_str)
            except ValueError:
                grand_total = None
        
        # Parse additional options
        excess_tp_enabled = request.POST.get("excess_tp_enabled", "").strip().lower() == 'true'
        if excess_tp_enabled:
            excess_tp_str = request.POST.get("excess_tp_percent", "").strip()
            if excess_tp_str:
                try:
                    excess_tp_percent = float(excess_tp_str)
                except ValueError:
                    excess_tp_percent = None
        
        ls_special_enabled = request.POST.get("ls_special_enabled", "").strip().lower() == 'true'
        if ls_special_enabled:
            ls_special_name = request.POST.get("ls_special_name", "").strip() or None
            ls_special_amount_str = request.POST.get("ls_special_amount", "").strip()
            if ls_special_amount_str:
                try:
                    ls_special_amount = float(ls_special_amount_str)
                except ValueError:
                    ls_special_amount = None
        
        # Parse Deduct Old Material Cost (for repair work)
        deduct_old_material_str = request.POST.get("deduct_old_material", "").strip()
        if deduct_old_material_str:
            try:
                deduct_old_material = float(deduct_old_material_str)
            except ValueError:
                deduct_old_material = None

    work_type = (request.POST.get("work_type")
                 or request.session.get("work_type")
                 or "original").lower()
    request.session["work_type"] = work_type
    
    # Get selected backend ID from session (for multi-backend support)
    selected_backend_id = request.session.get("selected_backend_id")

    # For development: Run synchronously without Celery
    # This bypasses the async task queue when Redis/Celery isn't available
    from django.conf import settings
    
    if getattr(settings, 'CELERY_TASK_ALWAYS_EAGER', True):
        # Synchronous mode - call task directly without queue
        try:
            org = get_org_from_request(request)
            
            job = Job.objects.create(
                organization=org,
                user=request.user,
                job_type='generate_output_excel',
                status='queued',
                current_step="Processing...",
            )
            
            # Store inputs in job result
            job.result = {
                'fetched_items': fetched,
                'qty_map': item_qtys,
                'unit_map': item_units,
                'work_name': work_name,
                'work_type': work_type,
                'grand_total': grand_total,
                'excess_tp_enabled': excess_tp_enabled,
                'excess_tp_percent': excess_tp_percent,
                'ls_special_enabled': ls_special_enabled,
                'ls_special_name': ls_special_name,
                'ls_special_amount': ls_special_amount,
                'deduct_old_material': deduct_old_material,
                'backend_id': selected_backend_id,
                # Uploaded custom items data for dual-source download
                'uploaded_items': request.session.get('uploaded_items', []),
                'uploaded_file_id': request.session.get('uploaded_file_id'),
                'uploaded_item_blocks': request.session.get('uploaded_item_blocks', {}),
                'uploaded_sheet_name': request.session.get('uploaded_sheet_name', ''),
                'item_descs': request.session.get('item_descs', {}),
                'item_units_saved': request.session.get('item_units', {}),
            }
            job.save()
            
            # Call task function directly (synchronous, no Celery)
            from core.tasks import generate_output_excel
            generate_output_excel.apply(args=(
                job.id,
                category,
                json.dumps(item_qtys),
                json.dumps(item_units),
                work_name,
                work_type,
                grand_total,
                excess_tp_percent,
                ls_special_name,
                ls_special_amount,
                deduct_old_material,
                selected_backend_id,
            )).get()

            # Refresh job to get updated result
            job.refresh_from_db()

            # Return JSON with job_id and status_url so JS can poll/download
            return JsonResponse({
                'job_id': job.id,
                'status_url': reverse('job_status', args=[job.id]),
                'message': job.current_step or 'Processing complete',
            })
                
        except Exception as e:
            logger.error(f"Failed to generate output Excel: {e}")
            return JsonResponse({"error": str(e)}, status=500)
    
    # Async mode with Celery (production)
    try:
        org = get_org_from_request(request)
        
        job = Job.objects.create(
            organization=org,
            user=request.user,
            job_type='generate_output_excel',
            status='queued',
            current_step="Queued for processing",
        )
        
        # Store inputs in job result temporarily
        job.result = {
            'fetched_items': fetched,
            'qty_map': item_qtys,
            'unit_map': item_units,
            'work_name': work_name,
            'work_type': work_type,
            'grand_total': grand_total,
            'excess_tp_enabled': excess_tp_enabled,
            'excess_tp_percent': excess_tp_percent,
            'ls_special_enabled': ls_special_enabled,
            'ls_special_name': ls_special_name,
            'ls_special_amount': ls_special_amount,
            'deduct_old_material': deduct_old_material,
            'backend_id': selected_backend_id,
            # Uploaded custom items data for dual-source download
            'uploaded_items': request.session.get('uploaded_items', []),
            'uploaded_file_id': request.session.get('uploaded_file_id'),
            'uploaded_item_blocks': request.session.get('uploaded_item_blocks', {}),
            'uploaded_sheet_name': request.session.get('uploaded_sheet_name', ''),
            'item_descs': request.session.get('item_descs', {}),
            'item_units_saved': request.session.get('item_units', {}),
        }
        job.save()

        # Enqueue async task
        from core.tasks import generate_output_excel
        task = generate_output_excel.delay(
            job.id,
            category,
            json.dumps(item_qtys),
            json.dumps(item_units),
            work_name,
            work_type,
            grand_total,
            excess_tp_percent,
            ls_special_name,
            ls_special_amount,
            deduct_old_material,
            selected_backend_id,
        )
        
        job.celery_task_id = task.id
        job.save()
        
        return JsonResponse({
            'job_id': job.id,
            'status_url': reverse('job_status', args=[job.id]),
            'message': f'Generating {category} output estimate. Please wait...'
        })
        
    except Exception as e:
        logger.error(f"Failed to enqueue output Excel task: {e}")
        return JsonResponse({"error": str(e)}, status=500)






# -----------------------
# CLEAR OUTPUT
# -----------------------
@login_required(login_url='login')
def clear_output(request, category):
    request.session["fetched_items"] = []
    request.session["qty_map"] = {}
    request.session["unit_map"] = {}
    request.session["work_name"] = ""
    request.session["grand_total"] = ""
    # Clear uploaded custom items
    request.session["uploaded_items"] = []
    request.session["uploaded_file_id"] = None
    request.session["uploaded_item_blocks"] = {}
    request.session["uploaded_sheet_name"] = ""

    group = request.GET.get("group")
    if group:
        return redirect("datas_items", category=category, group=group)

    return redirect("datas_groups", category=category)


# -----------------------
# SAVE QTY MAP (AJAX endpoint for preserving quantities during navigation)
# -----------------------
@login_required(login_url='login')
def save_qty_map(request, category):
    """
    AJAX endpoint to save quantity map, unit map, and grand total to session.
    Called before navigation to preserve entered quantities and units.
    """
    if request.method == "POST":
        try:
            qty_map_str = request.POST.get("qty_map", "")
            unit_map_str = request.POST.get("unit_map", "")
            grand_total_str = request.POST.get("grand_total", "")
            work_name = request.POST.get("work_name", "")
            excess_tp_percent = request.POST.get("excess_tp_percent", "")
            ls_special_name = request.POST.get("ls_special_name", "")
            ls_special_amount = request.POST.get("ls_special_amount", "")
            deduct_old_material = request.POST.get("deduct_old_material", "")
            work_type = request.POST.get("work_type", "")
            
            if qty_map_str:
                try:
                    qty_map = json.loads(qty_map_str)
                    if isinstance(qty_map, dict):
                        request.session["qty_map"] = qty_map
                except json.JSONDecodeError:
                    pass
            
            if unit_map_str:
                try:
                    unit_map = json.loads(unit_map_str)
                    if isinstance(unit_map, dict):
                        request.session["unit_map"] = unit_map
                except json.JSONDecodeError:
                    pass
            
            if grand_total_str:
                request.session["grand_total"] = grand_total_str
            
            if work_name:
                request.session["work_name"] = work_name
            
            if excess_tp_percent:
                request.session["excess_tp_percent"] = excess_tp_percent
            
            if ls_special_name:
                request.session["ls_special_name"] = ls_special_name
            
            if ls_special_amount:
                request.session["ls_special_amount"] = ls_special_amount
            
            if deduct_old_material:
                request.session["deduct_old_material"] = deduct_old_material
            
            if work_type:
                request.session["work_type"] = work_type
            
            # Ensure session is flushed to DB before the response is sent
            request.session.modified = True
            
            return JsonResponse({"status": "ok"})
        except Exception as e:
            return JsonResponse({"status": "error", "message": str(e)}, status=400)
    
    return JsonResponse({"status": "error", "message": "POST required"}, status=405)


@login_required(login_url='login')
@org_required
def save_project(request, category):
    org = get_org_from_request(request)
    fetched = request.session.get("fetched_items", []) or []
    if not fetched:
        return redirect("datas_groups", category=category)

    if request.method == "POST":
        qty_map = {}
        qty_map_str = request.POST.get("qty_map", "")
        if qty_map_str:
            try:
                raw = json.loads(qty_map_str)
                if isinstance(raw, dict):
                    for k, v in raw.items():
                        try:
                            num = float(v)
                        except (TypeError, ValueError):
                            continue
                        qty_map[str(k)] = num
            except json.JSONDecodeError:
                qty_map = {}

        work_name = (request.POST.get("work_name_hidden") or "").strip()

        request.session["qty_map"] = qty_map
        request.session["work_name"] = work_name

        project_name = request.POST.get("project_name") or request.session.get("current_project_name")

        if project_name:
            project, created = Project.objects.get_or_create(organization=org, name=project_name)
            project.category = category

            data = {
                "items": fetched,
                "qty_map": qty_map,
                "work_name": work_name,
            }
            project.set_items(data)
            project.save()

            request.session["current_project_name"] = project.name

            return redirect("my_projects")

    return redirect("datas_groups", category=category)


@org_required
def load_project(request, project_id):
    org = get_org_from_request(request)
    project = get_object_or_404(Project, id=project_id, organization=org)
    stored = project.get_items()

    if isinstance(stored, dict):
        fetched = stored.get("items", []) or []
        qty_map = stored.get("qty_map", {}) or {}
        work_name = stored.get("work_name", "") or ""
    else:
        fetched = stored or []
        qty_map = {}
        work_name = ""

    request.session["fetched_items"] = fetched
    request.session["qty_map"] = qty_map
    request.session["work_name"] = work_name
    request.session["current_project_name"] = project.name

    return redirect("datas_groups", category=project.category)


@login_required(login_url='login')
def download_estimate(request, category):
    """
    Async estimate Excel generation endpoint.
    
    Uses session['fetched_items'] to generate estimate-only workbook asynchronously.
    Returns JSON with job_id and status_url for polling.
    """
    fetched = request.session.get("fetched_items", [])
    if not fetched:
        return JsonResponse({"error": "No items selected"}, status=400)
    
    # Get selected backend ID from session (for multi-backend support)
    selected_backend_id = request.session.get("selected_backend_id")

    try:
        org = get_org_from_request(request)
        
        job = Job.objects.create(
            organization=org,
            user=request.user,
            job_type='generate_estimate_excel',
            status='queued',
            current_step="Queued for processing",
        )
        
        # Store fetched items and backend_id in result temporarily
        job.result = {
            'fetched_items': fetched,
            'backend_id': selected_backend_id,
        }
        job.save()
        
        # Enqueue async task
        from core.tasks import generate_estimate_excel
        task = generate_estimate_excel.delay(
            job.id,
            category,
            json.dumps(fetched),
            selected_backend_id,
        )
        
        job.celery_task_id = task.id
        job.save()
        
        return JsonResponse({
            'job_id': job.id,
            'status_url': reverse('job_status', args=[job.id]),
            'message': f'Generating {category} estimate. Please wait...'
        })
        
    except Exception as e:
        logger.error(f"Failed to enqueue estimate Excel task: {e}")
        return JsonResponse({"error": str(e)}, status=500)



def build_fetched_details(category, base_dir, fetched_names):
    """
    Build a list of dicts for the UI side panel:
    [
      {"slno": 1, "name": "Item 1", "rate": 123.45, "rate_display": "123.45"},
      ...
    ]
    Rates are read from the backend Excel (last non-empty J cell of each block).
    """
    if not fetched_names:
        return []

    try:
        items_list, groups_map, _, ws_data, filepath = load_backend(category, base_dir)
    except Exception:
        return [
            {"slno": idx, "name": name, "rate": "", "rate_display": ""}
            for idx, name in enumerate(fetched_names, start=1)
        ]

    block_map = {it["name"]: (it["start_row"], it["end_row"]) for it in items_list}

    try:
        wb_vals = load_workbook(filepath, data_only=True)
        ws_vals = wb_vals["Master Datas"]
    except Exception:
        wb_vals = None
        ws_vals = None

    details = []
    for idx, name in enumerate(fetched_names, start=1):
        rate_value = ""
        rate_display = ""
        block = block_map.get(name)

        if block and ws_vals is not None:
            start_row, end_row = block
            for r in range(end_row, start_row - 1, -1):
                val = ws_vals.cell(row=r, column=10).value
                if val not in (None, ""):
                    rate_display = val
                    try:
                        rate_value = float(val)
                    except Exception:
                        rate_value = ""
                    break

        details.append(
            {
                "slno": idx,
                "name": name,
                "rate": rate_value,
                "rate_display": rate_display,
            }
        )

    return details


@login_required(login_url='login')
@require_POST
def toggle_item(request, category, group):
    """
    AJAX endpoint: add/remove an item from session['fetched_items']
    based on checkbox state, and return updated side panel HTML.
    """
    item = request.POST.get("item", "").strip()
    checked = request.POST.get("checked") == "true"

    fetched = request.session.get("fetched_items", [])

    if checked:
        if item and item not in fetched:
            fetched.append(item)
    else:
        fetched = [x for x in fetched if x != item]

    request.session["fetched_items"] = fetched

    html = render_to_string(
        "core/sidebar_output.html",
        {"category": category, "fetched": fetched},
        request=request,
    )

    return JsonResponse({"html": html})


@login_required(login_url='login')
def new_project(request):
    """
    Clear current selection and start from scratch.
    """
    request.session["fetched_items"] = []
    request.session["qty_map"] = {}
    request.session["work_name"] = ""
    request.session["current_project_name"] = None
    return redirect("datas")


@org_required
def delete_project(request, project_id):
    org = get_org_from_request(request)
    project = get_object_or_404(Project, id=project_id, organization=org)
    project.delete()
    return redirect("my_projects")


# ==============================================================================
# SPECIFICATION REPORT & FORWARDING LETTER (relocated from estimate_views.py)
# ==============================================================================

@org_required
@login_required
def download_specification_report(request, estimate_id):
    """Generate and download specification report as Word document"""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    org = request.organization
    estimate = get_object_or_404(Estimate, id=estimate_id, organization=org)

    try:
        # Create Word document
        doc = Document()

        # Title
        title = doc.add_heading('Specification report accompanying the estimate :-', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in title.runs:
            run.font.size = Pt(12)
            run.font.bold = True

        # Introduction paragraph with work name
        estimate_data = estimate.estimate_data or {}
        work_name = estimate_data.get('name_of_work', 'the work')

        intro_text = f'The estimate is prepared for the work {work_name}'
        intro_para = doc.add_paragraph(intro_text)
        intro_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in intro_para.runs:
            run.font.size = Pt(11)

        # Estimate amount
        total_amount = estimate.total_amount or 0
        amount_para = doc.add_paragraph()
        amount_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        amount_run = amount_para.add_run(f'Est.Amount: Rs. {total_amount:,.2f}')
        amount_run.font.size = Pt(11)
        amount_run.font.bold = True
        amount_run.font.color.rgb = RGBColor(255, 0, 0)

        doc.add_paragraph()

        # Body of letter placeholder
        body_label = doc.add_paragraph('{{BODY_OF_LETTER}}')
        body_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in body_label.runs:
            run.font.size = Pt(11)
            run.font.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()

        doc.add_paragraph('Hence, this estimate has been prepared accordingly.')

        doc.add_paragraph()

        # Extract items with quantities and units as bullet points
        estimate_items = estimate_data.get('ws_estimate_rows', estimate_data.get('items', []))

        for item in estimate_items:
            item_description = item.get('desc', item.get('description', item.get('display_name', '')))
            quantity = item.get('qty_est', item.get('qty', item.get('quantity', '')))
            unit = item.get('unit', '')

            if quantity and unit:
                qty_str = str(quantity).strip()
                if '.' in qty_str and qty_str.endswith('.0'):
                    qty_str = qty_str.replace('.0', '')
                bullet_text = f'{item_description}  -  {qty_str} {unit}'
            elif quantity:
                qty_str = str(quantity).strip()
                if '.' in qty_str and qty_str.endswith('.0'):
                    qty_str = qty_str.replace('.0', '')
                bullet_text = f'{item_description}  -  {qty_str}'
            else:
                bullet_text = item_description

            bullet_para = doc.add_paragraph(bullet_text, style='List Bullet')
            for run in bullet_para.runs:
                run.font.size = Pt(11)

        doc.add_paragraph()

        # Calculate financial year
        from datetime import datetime
        today = datetime.now()
        if today.month >= 4:
            fy_start = today.year
            fy_end = (today.year + 1) % 100
        else:
            fy_start = today.year - 1
            fy_end = today.year % 100
        financial_year = f"{fy_start}-{fy_end:02d}"

        footer_text = (f'The rates proposed in the estimate are as per SQR {financial_year} and Approved rates. L.S. Provision is made in the '
                      'estimate towards GST at 18%, QC amount at 1%, Labour Cess at 1% and NAC amount at 0.1% as per actual and LS Provision Towards, unforeseen items & rounding off also proposed in the estimate.')
        footer_para = doc.add_paragraph(footer_text)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in footer_para.runs:
            run.font.size = Pt(10)

        doc.add_paragraph()

        funds_text = ('FUNDS: The estimate requires Administrative sanction and also fixes up the agency with provision of funds '
                     'under relevant head of account for taking up the work from the Government. Telangana State Hyderabad')
        funds_para = doc.add_paragraph(funds_text)
        funds_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in funds_para.runs:
            run.font.size = Pt(10)
            run.font.bold = True

        filename = 'Spec_Report.docx'

        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        doc.save(response)
        return response

    except Exception as e:
        logger.error(f'Error generating specification report: {str(e)}', exc_info=True)
        from django.contrib import messages
        messages.error(request, f'Error generating report: {str(e)}')
        return redirect('view_estimate', estimate_id=estimate_id)


@login_required(login_url='login')
def download_specification_report_live(request, category):
    """
    Generate specification report from live estimate items (New Estimate module).
    Receives items as JSON from the frontend.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if request.method != 'POST':
        return redirect('datas_groups', category=category)

    try:
        items_json = request.POST.get('items', '[]')
        work_name = request.POST.get('work_name', '{{NAME_OF_WORK}}')
        total_amount = request.POST.get('total_amount', '0.00')

        items = json.loads(items_json)

        if not items:
            from django.contrib import messages
            messages.error(request, 'No items with quantities to generate specification report')
            return redirect('datas_groups', category=category)

        doc = Document()

        title = doc.add_heading('Specification report accompanying the estimate :-', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in title.runs:
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.underline = True

        intro_para = doc.add_paragraph()
        intro_para.add_run(f'The estimate is prepared for the work {work_name}')

        doc.add_paragraph()

        amount_para = doc.add_paragraph()
        amount_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        amount_run = amount_para.add_run(f'Est.Amount: Rs. {total_amount}')
        amount_run.font.bold = True
        amount_run.font.underline = True

        doc.add_paragraph()

        body_para = doc.add_paragraph('{{BODY_OF_LETTER}}')
        for run in body_para.runs:
            run.font.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()

        doc.add_paragraph('Hence, this estimate has been prepared accordingly.')

        doc.add_paragraph()

        for item in items:
            desc = item.get('desc', '')
            qty = item.get('qty', '')
            unit = item.get('unit', '')

            if qty:
                try:
                    qty_float = float(qty)
                    if qty_float == int(qty_float):
                        qty = str(int(qty_float))
                    else:
                        qty = str(qty_float)
                except:
                    pass

            if qty and unit:
                bullet_text = f'{desc}  -  {qty} {unit}'
            elif qty:
                bullet_text = f'{desc}  -  {qty}'
            else:
                bullet_text = desc

            bullet_para = doc.add_paragraph(bullet_text, style='List Bullet')
            for run in bullet_para.runs:
                run.font.size = Pt(11)
                run.font.bold = True

        doc.add_paragraph()

        from datetime import datetime
        today = datetime.now()
        if today.month >= 4:
            fy_start = today.year
            fy_end = (today.year + 1) % 100
        else:
            fy_start = today.year - 1
            fy_end = today.year % 100
        financial_year = f"{fy_start}-{fy_end:02d}"

        footer_text = (f'The rates proposed in the estimate are as per SQR {financial_year} and Approved rates. L.S. Provision is made in the '
                      'estimate towards GST at 18%, QC amount at 1%, Labour Cess at 1% and NAC amount at 0.1% as per actual '
                      'and LS Provision Towards, unforeseen items & rounding off also proposed in the estimate.')
        footer_para = doc.add_paragraph(footer_text)
        for run in footer_para.runs:
            run.font.size = Pt(10)

        doc.add_paragraph()

        funds_para = doc.add_paragraph()
        funds_run = funds_para.add_run('FUNDS: ')
        funds_run.font.bold = True
        funds_run.font.underline = True
        funds_para.add_run('The estimate requires Administrative sanction and also fixes up the agency with provision of funds '
                          'under relevant head of account for taking up the work from the Government, Telangana State Hyderabad')

        filename = 'Spec_Report.docx'

        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        doc.save(response)
        return response

    except Exception as e:
        logger.error(f'Error generating specification report: {str(e)}', exc_info=True)
        from django.contrib import messages
        messages.error(request, f'Error generating specification report: {str(e)}')
        return redirect('datas_groups', category=category)


@login_required(login_url='login')
def download_forwarding_letter_live(request, category):
    """
    Generate forwarding letter from live estimate items (New Estimate module).
    Receives items as JSON from the frontend.
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    if request.method != 'POST':
        return redirect('datas_groups', category=category)

    try:
        items_json = request.POST.get('items', '[]')
        work_name = request.POST.get('work_name', '{{NAME_OF_WORK}}')
        total_amount = request.POST.get('total_amount', '0.00')

        items = json.loads(items_json)

        if not items:
            from django.contrib import messages
            messages.error(request, 'No items with quantities to generate forwarding letter')
            return redirect('datas_groups', category=category)

        try:
            grand_total = float(total_amount.replace(',', '').replace('Rs.', '').replace('\u20b9', '').strip())
        except:
            grand_total = 0.0

        current_date = _get_current_date_formatted()
        financial_year = _get_current_financial_year()
        today = timezone.now().date()

        letter_settings = _get_letter_settings(request.user)

        doc = Document()

        placeholder_color = RGBColor(169, 169, 169)

        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Header - Government/Organization name
        header1 = doc.add_paragraph()
        header1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if letter_settings and letter_settings.government_name:
            run1 = header1.add_run(letter_settings.government_name.upper())
            run1.font.bold = True
            run1.font.size = Pt(14)
        else:
            run1 = header1.add_run('[GOVERNMENT / ORGANIZATION NAME]')
            run1.font.bold = True
            run1.font.size = Pt(14)
            run1.font.color.rgb = placeholder_color
            run1.font.italic = True

        # Header - Department name
        header2 = doc.add_paragraph()
        header2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if letter_settings and letter_settings.department_name:
            run2 = header2.add_run(letter_settings.department_name.upper())
            run2.font.bold = True
            run2.font.size = Pt(13)
        else:
            run2 = header2.add_run('[DEPARTMENT NAME]')
            run2.font.bold = True
            run2.font.size = Pt(13)
            run2.font.color.rgb = placeholder_color
            run2.font.italic = True

        doc.add_paragraph()

        # From/To section in a table
        from_to_table = doc.add_table(rows=1, cols=2)
        from_to_table.autofit = True

        from_cell = from_to_table.cell(0, 0)
        from_para = from_cell.paragraphs[0]
        from_label = from_para.add_run('From: -\n')
        from_label.font.bold = True

        if letter_settings and letter_settings.officer_name:
            name_qual = letter_settings.officer_name
            if letter_settings.officer_qualification:
                name_qual += f", {letter_settings.officer_qualification}"
            from_run1 = from_para.add_run(f'{name_qual},\n')
        else:
            from_run1 = from_para.add_run('[Officer Name, Qualification],\n')
            from_run1.font.color.rgb = placeholder_color
            from_run1.font.italic = True

        if letter_settings and letter_settings.officer_designation:
            from_run2 = from_para.add_run(f'{letter_settings.officer_designation},\n')
        else:
            from_run2 = from_para.add_run('[Designation],\n')
            from_run2.font.color.rgb = placeholder_color
            from_run2.font.italic = True

        if letter_settings and (letter_settings.sub_division or letter_settings.office_address):
            sub_addr = letter_settings.sub_division
            if letter_settings.office_address:
                sub_addr += f", {letter_settings.office_address}" if sub_addr else letter_settings.office_address
            from_run3 = from_para.add_run(f'{sub_addr}.')
        else:
            from_run3 = from_para.add_run('[Sub Division, Office Address].')
            from_run3.font.color.rgb = placeholder_color
            from_run3.font.italic = True

        # To section
        to_cell = from_to_table.cell(0, 1)
        to_para = to_cell.paragraphs[0]
        to_label = to_para.add_run('To,\n')
        to_label.font.bold = True

        if letter_settings and letter_settings.recipient_designation:
            to_run1 = to_para.add_run(f'{letter_settings.recipient_designation},\n')
        else:
            to_run1 = to_para.add_run('[Officer Designation],\n')
            to_run1.font.color.rgb = placeholder_color
            to_run1.font.italic = True

        if letter_settings and letter_settings.recipient_division:
            to_run2 = to_para.add_run(f'{letter_settings.recipient_division},\n')
        else:
            to_run2 = to_para.add_run('[Division Name],\n')
            to_run2.font.color.rgb = placeholder_color
            to_run2.font.italic = True

        if letter_settings and letter_settings.recipient_address:
            to_run3 = to_para.add_run(f'{letter_settings.recipient_address}.')
        else:
            to_run3 = to_para.add_run('[Address].')
            to_run3.font.color.rgb = placeholder_color
            to_run3.font.italic = True

        doc.add_paragraph()

        # Letter number and date
        lr_para = doc.add_paragraph()
        lr_run = lr_para.add_run('Lr No. ')
        lr_run.font.bold = True
        lr_run.font.underline = True
        if letter_settings and letter_settings.office_code:
            lr_code = lr_para.add_run(letter_settings.office_code)
            lr_code.font.underline = True
            lr_code.font.bold = True
        else:
            lr_placeholder = lr_para.add_run('[Office Code]')
            lr_placeholder.font.color.rgb = placeholder_color
            lr_placeholder.font.italic = True
            lr_placeholder.font.underline = True
            lr_placeholder.font.bold = True
        lr_fy = lr_para.add_run(f'/{financial_year}/          ')
        lr_fy.font.bold = True
        lr_fy.font.underline = True
        lr_date = lr_para.add_run(f'\t\t\t\t\tDate:-    - {today.strftime("%m")} - {today.year}.')
        lr_date.font.bold = True

        doc.add_paragraph()

        sir_para = doc.add_paragraph()
        sir_para.add_run('Sir,')

        doc.add_paragraph()

        # Subject
        subject_para = doc.add_paragraph()
        subj_run = subject_para.add_run('Sub:-')
        subj_run.font.underline = True
        subject_para.add_run('\t')
        subj_work = subject_para.add_run(f'{work_name} ')
        subj_work.font.bold = True
        subject_para.add_run(f'for the year {financial_year}.  -  Submission  -  Request for obtaining administrative sanction  -  Regarding.')

        doc.add_paragraph()

        # Reference
        ref_para = doc.add_paragraph()
        ref_run = ref_para.add_run('Ref:-')
        ref_run.font.underline = True
        ref_para.add_run('\tMemo No.')
        ref_placeholder = ref_para.add_run('[Reference Number]')
        ref_placeholder.font.color.rgb = placeholder_color
        ref_placeholder.font.italic = True
        ref_placeholder.font.underline = True
        ref_para.add_run(f'/{financial_year} Dt.')
        ref_date_placeholder = ref_para.add_run('[DD.MM.YYYY]')
        ref_date_placeholder.font.color.rgb = placeholder_color
        ref_date_placeholder.font.italic = True
        ref_date_placeholder.font.underline = True

        doc.add_paragraph()

        stars_para = doc.add_paragraph()
        stars_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        stars_para.add_run('**.**')

        doc.add_paragraph()

        body_para = doc.add_paragraph()
        body_para.add_run('With reference to the subject cited, I submit here ')
        with_run = body_para.add_run('with  1')
        with_run.font.underline = True
        body_para.add_run(' No. estimate for the following work for the amount specified.')

        doc.add_paragraph()

        # Create table for estimate
        table = doc.add_table(rows=2, cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for cell in table.columns[0].cells:
            cell.width = Inches(0.5)
        for cell in table.columns[1].cells:
            cell.width = Inches(4.5)
        for cell in table.columns[2].cells:
            cell.width = Inches(1.5)

        header_cells = table.rows[0].cells
        header_cells[0].text = 'Sl.\nNo'
        header_cells[1].text = 'Name of work'
        header_cells[2].text = 'Amount'

        for cell in header_cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.bold = True

        row_cells = table.rows[1].cells
        row_cells[0].text = '1'
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[1].text = work_name

        formatted_amount = _format_indian_number(grand_total)
        row_cells[2].text = f"Rs.{formatted_amount}"
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph()

        spec_para = doc.add_paragraph()
        spec_para.add_run("Specification report accompanying the estimate explains the necessity and provisions made therein in detail.")

        doc.add_paragraph()

        request_para = doc.add_paragraph()
        request_para.add_run('I request the ')
        if letter_settings and letter_settings.superior_designation:
            req_run = request_para.add_run(letter_settings.superior_designation)
        else:
            req_placeholder = request_para.add_run('[Superior Officer Designation]')
            req_placeholder.font.color.rgb = placeholder_color
            req_placeholder.font.italic = True
        request_para.add_run(' to kindly arrange to obtain administrative sanction for the above estimate and arrange to finalize the agency at the earliest for taking up the work.')

        doc.add_paragraph()

        enc_para = doc.add_paragraph()
        enc_para.add_run('Enclosure: -')
        doc.add_paragraph('Estimate  - 1 No.')

        doc.add_paragraph()
        doc.add_paragraph()

        sign_para = doc.add_paragraph()
        sign_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sign_para.add_run('Yours faithfully,')

        doc.add_paragraph()
        doc.add_paragraph()

        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if letter_settings and letter_settings.officer_designation:
            run_title = title_para.add_run(f'{letter_settings.officer_designation}\n')
            run_title.font.bold = True
        else:
            run_title = title_para.add_run('[Officer Designation]\n')
            run_title.font.bold = True
            run_title.font.color.rgb = placeholder_color
            run_title.font.italic = True

        if letter_settings and letter_settings.sub_division:
            sub_div_run = title_para.add_run(f'{letter_settings.sub_division},\n')
        else:
            sub_div_run = title_para.add_run('[Sub Division Name],\n')
            sub_div_run.font.color.rgb = placeholder_color
            sub_div_run.font.italic = True

        if letter_settings and letter_settings.office_address:
            addr_run = title_para.add_run(f'{letter_settings.office_address}.')
        else:
            addr_run = title_para.add_run('[Office Address].')
            addr_run.font.color.rgb = placeholder_color
            addr_run.font.italic = True

        doc.add_paragraph()

        copy_para = doc.add_paragraph()
        copy_para.add_run('Copy to the ')
        if letter_settings and (letter_settings.copy_to_designation or letter_settings.copy_to_section):
            copy_text = letter_settings.copy_to_designation or ''
            if letter_settings.copy_to_section:
                copy_text += f", {letter_settings.copy_to_section}" if copy_text else letter_settings.copy_to_section
            copy_run = copy_para.add_run(copy_text)
        else:
            copy_placeholder = copy_para.add_run('[Officer Designation, Section Name]')
            copy_placeholder.font.color.rgb = placeholder_color
            copy_placeholder.font.italic = True
        copy_para.add_run(' for information.')

        filename = 'Fwd_Letter.docx'

        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        doc.save(response)
        return response

    except Exception as e:
        logger.error(f'Error generating forwarding letter: {str(e)}', exc_info=True)
        from django.contrib import messages
        messages.error(request, f'Error generating forwarding letter: {str(e)}')
        return redirect('datas_groups', category=category)

