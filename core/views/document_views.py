# Auto-generated from core/views.py split
import json
import os
import re
import logging
from copy import copy

import inflect
from docx import Document
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, Border, Side, PatternFill
from django.utils import timezone
from django.urls import reverse
from django.core.files.uploadedfile import InMemoryUploadedFile
from django.contrib.auth.decorators import login_required

from django.conf import settings
from django.http import HttpResponse, JsonResponse, HttpResponseNotAllowed
from django.shortcuts import render, redirect, get_object_or_404
from django.template.loader import render_to_string
from django.views.decorators.http import require_POST
from django.utils.crypto import get_random_string

import io
from io import BytesIO
from difflib import SequenceMatcher

from ..models import Project, SelfFormattedTemplate, Estimate, Organization, Membership, Upload, Job, OutputFile, LetterSettings
from ..decorators import org_required, role_required

logger = logging.getLogger(__name__)
from ..tasks import process_excel_upload, generate_bill_pdf, generate_workslip_pdf, generate_bill_document_task
from ..utils_excel import load_backend, copy_block_with_styles_and_formulas, build_temp_day_rates

p_engine = inflect.engine()
BILL_TEMPLATES_DIR = os.path.join(settings.BASE_DIR, "core", "templates", "core", "bill_templates")
_inflect_engine = inflect.engine()

from .utils import (_apply_print_settings, _format_indian_number,
    _number_to_words_rupees, _get_current_financial_year, _get_current_date_formatted,
    _get_letter_settings, get_org_from_request, create_job_for_excel)
from .bill_parsing import (_extract_header_data_fuzzy_from_wb,
    _extract_header_data_from_sheet, _extract_total_amount_from_bill_wb,
    _detect_bill_format, find_estimate_sheet_and_header_row,
    parse_estimate_items)
from .amount_utils import (_extract_total_amount_for_action,
    _build_mb_details_string, _resolve_cc_header,
    _fill_excel_template, _fill_docx_template)

@org_required
def bill_document(request):
    """
    Generate:
      - LS Form (Part / Final) -> Excel (multiple sheets if multiple bills)
      - Covering Letter (Word) -> Multiple pages if multiple bills
      - Movement Slip (Word) -> Multiple pages if multiple bills

    directly from an uploaded Bill Excel + MB details.
    Returns the file directly for download (synchronous).

    POST fields:
      - bill_file     : uploaded Excel (Bill / Estimate / Nth bill)
      - doc_kind      : 'ls_part', 'ls_final', 'covering', 'movement'
      - action        : same as in bill() (estimate_first_part, firstpart_nth_part, ...)
      - nth_number    : Nth number (for N>=2 bills)
      - mb_measure_no, mb_measure_p_from, mb_measure_p_to
      - mb_abstract_no, mb_abstract_p_from, mb_abstract_p_to
    """
    if request.method != "POST":
        return HttpResponseNotAllowed(["POST"], "Use POST to generate documents.")

    # 1) Get uploaded Bill Excel
    uploaded = request.FILES.get("bill_file")
    if not uploaded:
        return HttpResponse(
            "Please upload the Bill Excel file for document generation.",
            status=400,
        )

    # 2) Which document?
    doc_kind = (request.POST.get("doc_kind") or "").strip()
    if not doc_kind:
        return HttpResponse("Missing 'doc_kind' in request.", status=400)

    # Action + Nth number (for CC header)
    action = (request.POST.get("action") or "").strip()
    nth_number_str = (request.POST.get("nth_number") or "").strip()

    # For LS forms: if action is empty, derive action from doc_kind for CC header resolution
    # This happens when user clicks LS form buttons directly (which set doc_kind but not action)
    if not action and doc_kind in ("ls_part", "ls_final"):
        # Default to first bill if nth_number not specified
        if not nth_number_str or nth_number_str == "1":
            action = "estimate_first_part" if doc_kind == "ls_part" else "estimate_first_final"
        else:
            action = "nth_nth_part" if doc_kind == "ls_part" else "nth_nth_final"

    # 3) MB details from form
    mb_measure_no = (request.POST.get("mb_measure_no") or "").strip()
    mb_measure_p_from = (request.POST.get("mb_measure_p_from") or "").strip()
    mb_measure_p_to = (request.POST.get("mb_measure_p_to") or "").strip()
    mb_abs_no = (request.POST.get("mb_abstract_no") or "").strip()
    mb_abs_p_from = (request.POST.get("mb_abstract_p_from") or "").strip()
    mb_abs_p_to = (request.POST.get("mb_abstract_p_to") or "").strip()

    # 4) Open uploaded Excel
    try:
        wb_in = load_workbook(uploaded, data_only=True)
    except Exception as e:
        return HttpResponse(f"Error reading uploaded Bill Excel: {e}", status=400)
    
    # Also open with formulas for fallback calculation
    try:
        uploaded.seek(0)  # Reset file pointer
        wb_formulas = load_workbook(uploaded, data_only=False)
    except:
        wb_formulas = None

    # Find all bill sheets (sheets that look like bills/estimates)
    bill_sheets = [ws for ws in wb_in.worksheets if ws.title.startswith("Bill")]
    if not bill_sheets:
        # Fallback: try to find sheets with estimate/bill-like structure
        bill_sheets = []
        for ws in wb_in.worksheets:
            # Check if it has typical bill structure
            if ws.max_row > 5:
                bill_sheets.append(ws)
        if not bill_sheets:
            bill_sheets = wb_in.worksheets  # use all sheets
    
    has_multiple_bills = len(bill_sheets) > 1

    # Try to extract MB details from the uploaded bill file itself
    file_header = _extract_header_data_fuzzy_from_wb(wb_in)
    file_mb_details = (file_header.get("mb_details") or "").strip()

    # If user has entered MB details manually, use those; otherwise use file-extracted MB
    user_entered_mb = any([
        mb_measure_no, mb_measure_p_from, mb_measure_p_to,
        mb_abs_no, mb_abs_p_from, mb_abs_p_to,
    ])

    if user_entered_mb:
        # User manually entered MB fields — build string from those
        mb_details_str = _build_mb_details_string(
            mb_measure_no,
            mb_measure_p_from,
            mb_measure_p_to,
            mb_abs_no,
            mb_abs_p_from,
            mb_abs_p_to,
        )
    elif file_mb_details:
        # Auto-extracted from the uploaded bill file
        mb_details_str = file_mb_details
    else:
        # No MB details available at all — build empty string
        mb_details_str = _build_mb_details_string(
            mb_measure_no,
            mb_measure_p_from,
            mb_measure_p_to,
            mb_abs_no,
            mb_abs_p_from,
            mb_abs_p_to,
        )

    cc_header = _resolve_cc_header(action, nth_number_str=nth_number_str)
    # Use CC Header from uploaded file if available
    file_cc_header = (file_header.get("cc_header") or "").strip()
    if file_cc_header:
        cc_header = file_cc_header
    
    # Current month + year
    now = timezone.now()
    mm_yyyy = f"{now.month:02d}.{now.year}"

    # Helper function to extract total from a single sheet
    def _extract_total_from_sheet(ws, ws_formulas=None):
        """
        Extract the total amount from a bill sheet.

        Strategy:
        1. Scan header rows (1-15) for columns containing "amount" in the header.
           Use the rightmost such column as the amount column.
        2. Scan for "Grand Total" or "Total" row (prefer "Grand Total").
        3. Read the amount value from the identified column in that row.
        4. Fall back to formula evaluation if cell is empty.
        """
        total = 0.0
        max_scan = min(ws.max_row or 0, 200)
        max_col = min(ws.max_column or 0, 20)

        # Step 1: Find the rightmost column with "amount" in header rows 1-15
        amount_col = None
        for r in range(1, min(max_scan, 16)):
            for c in range(1, max_col + 1):
                hdr = str(ws.cell(row=r, column=c).value or "").strip().lower()
                if "amount" in hdr:
                    # Always keep the rightmost match
                    if amount_col is None or c > amount_col:
                        amount_col = c

        # Fallback: if no "amount" header found, use columns 8, 9, 10 (legacy)
        fallback_cols = [8, 9, 10] if amount_col is None else [amount_col]

        def _try_get_amount(row, col):
            """Try to read a numeric amount from a cell, with formula fallback."""
            amt_val = ws.cell(row=row, column=col).value

            # Formula fallback if cell is empty/zero
            if (amt_val is None or amt_val == 0 or amt_val == '') and ws_formulas:
                try:
                    formula_cell = ws_formulas.cell(row=row, column=col)
                    formula_val = formula_cell.value
                    if isinstance(formula_val, str) and formula_val.startswith('='):
                        match = re.match(r'=([A-Z]+)(\d+)([\+\-])([A-Z]+)(\d+)', formula_val)
                        if match:
                            col1, row1, op, col2, row2 = match.groups()
                            def col_to_num(c_letter):
                                result = 0
                                for ch in c_letter:
                                    result = result * 26 + (ord(ch) - ord('A') + 1)
                                return result
                            val1 = ws.cell(row=int(row1), column=col_to_num(col1)).value
                            val2 = ws.cell(row=int(row2), column=col_to_num(col2)).value
                            if val1 and val2:
                                if op == '+':
                                    amt_val = float(val1) + float(val2)
                                elif op == '-':
                                    amt_val = float(val1) - float(val2)
                except Exception:
                    pass

            try:
                num_val = float(amt_val) if amt_val else 0
                if num_val != 0:
                    return num_val
            except (TypeError, ValueError):
                pass
            return 0.0

        # Step 2: Scan for "Grand Total" first (preferred), then "Total"
        grand_total_row = None
        total_row = None

        for r in range(1, max_scan + 1):
            for check_col in range(1, min(max_col + 1, 8)):
                cell_val = str(ws.cell(row=r, column=check_col).value or "").strip().lower()
                if "grand total" in cell_val:
                    grand_total_row = r
                elif cell_val == "total":
                    total_row = r

        # Prefer Grand Total over Total
        target_row = grand_total_row or total_row

        if target_row:
            for amt_col in fallback_cols:
                val = _try_get_amount(target_row, amt_col)
                if val != 0:
                    return val

        return total

    # 10) LS FORMS (EXCEL) - Multiple sheets support
    if doc_kind in ("ls_part", "ls_final"):
        if doc_kind == "ls_part":
            template_name = "LS_Form_Part.xlsx"
            download_name = "LS_Form.xlsx"
        else:
            template_name = "LS_Form_Final.xlsx"
            download_name = "LS_Form.xlsx"

        template_path = os.path.join(BILL_TEMPLATES_DIR, template_name)
        if not os.path.exists(template_path):
            return HttpResponse(f"LS template not found: {template_name}", status=404)

        # If multiple bill sheets, create one LS sheet per bill
        if has_multiple_bills:
            wb_out = Workbook()
            # Remove default sheet
            if wb_out.worksheets:
                wb_out.remove(wb_out.worksheets[0])
            
            for sheet_idx, bill_ws in enumerate(bill_sheets, start=1):
                # Extract per-sheet header data
                sheet_header = _extract_header_data_from_sheet(bill_ws)
                sheet_name_of_work = sheet_header.get("name_of_work", "") or ""
                sheet_agreement_ref = sheet_header.get("agreement", "") or ""
                sheet_agency_name = sheet_header.get("agency", "") or ""
                
                # Per-sheet MB: use user-entered MB if provided, else try sheet-level MB
                sheet_mb = mb_details_str
                if not user_entered_mb:
                    sheet_mb_from_file = (sheet_header.get("mb_details") or "").strip()
                    if sheet_mb_from_file:
                        sheet_mb = sheet_mb_from_file
                
                # Get formulas sheet if available
                ws_formulas = None
                if wb_formulas:
                    try:
                        ws_formulas = wb_formulas[bill_ws.title]
                    except:
                        pass
                
                # Calculate total for this sheet
                sheet_total = _extract_total_from_sheet(bill_ws, ws_formulas)
                try:
                    sheet_total_str = f"{float(sheet_total):,.2f}"
                except:
                    sheet_total_str = str(sheet_total)
                sheet_amount_words = _number_to_words_rupees(sheet_total)
                
                # Build context for this sheet
                sheet_ctx = {
                    "NAME_OF_WORK": sheet_name_of_work,
                    "NAME_OF_AGENCY": sheet_agency_name,
                    "AGENCY_NAME": sheet_agency_name,
                    "REF_OF_AGREEMENT": sheet_agreement_ref,
                    "AGREEMENT_REF": sheet_agreement_ref,
                    "MB_DETAILS": sheet_mb,
                    "CC_HEADER": cc_header,
                    "AMOUNT": sheet_total_str,
                    "TOTAL_AMOUNT": sheet_total_str,
                    "AMOUNT_IN_WORDS": sheet_amount_words,
                }
                
                # Load fresh template for this sheet
                try:
                    template_wb = load_workbook(template_path)
                except Exception as e:
                    return HttpResponse(f"Error loading template: {e}", status=500)
                
                # Fill placeholders in template
                for ws in template_wb.worksheets:
                    for row in ws.iter_rows():
                        for cell in row:
                            if isinstance(cell.value, str):
                                text = cell.value
                                for key, val in sheet_ctx.items():
                                    placeholder = "{{" + key + "}}"
                                    if placeholder in text:
                                        text = text.replace(placeholder, str(val or ""))
                                cell.value = text
                    
                    # Copy this sheet to output workbook
                    new_sheet_name = f"{bill_ws.title}_LS" if len(bill_ws.title) <= 25 else f"LS_{sheet_idx}"
                    new_ws = wb_out.create_sheet(title=new_sheet_name)
                    
                    # Copy cell values, styles, and merged cells
                    for row in ws.iter_rows():
                        for cell in row:
                            new_cell = new_ws[cell.coordinate]
                            new_cell.value = cell.value
                            if cell.has_style:
                                new_cell.font = copy(cell.font)
                                new_cell.border = copy(cell.border)
                                new_cell.fill = copy(cell.fill)
                                new_cell.number_format = copy(cell.number_format)
                                new_cell.protection = copy(cell.protection)
                                new_cell.alignment = copy(cell.alignment)
                    
                    # Copy merged cells
                    for merged_range in ws.merged_cells.ranges:
                        new_ws.merge_cells(str(merged_range))
                    
                    # Copy column widths
                    for col_letter, col_dim in ws.column_dimensions.items():
                        new_ws.column_dimensions[col_letter].width = col_dim.width
                    
                    # Copy row heights
                    for row_num, row_dim in ws.row_dimensions.items():
                        new_ws.row_dimensions[row_num].height = row_dim.height
            
            download_name = "LS_Forms.xlsx"
        else:
            # Single sheet - use original logic
            header = _extract_header_data_fuzzy_from_wb(wb_in)
            total_amount = _extract_total_from_sheet(bill_sheets[0], wb_formulas[bill_sheets[0].title] if wb_formulas else None)
            try:
                total_amount_str = f"{float(total_amount):,.2f}"
            except:
                total_amount_str = str(total_amount)
            
            ctx = {
                "NAME_OF_WORK": header.get("name_of_work", ""),
                "NAME_OF_AGENCY": header.get("agency", ""),
                "AGENCY_NAME": header.get("agency", ""),
                "REF_OF_AGREEMENT": header.get("agreement", ""),
                "AGREEMENT_REF": header.get("agreement", ""),
                "MB_DETAILS": mb_details_str,
                "CC_HEADER": cc_header,
                "AMOUNT": total_amount_str,
                "TOTAL_AMOUNT": total_amount_str,
                "AMOUNT_IN_WORDS": _number_to_words_rupees(total_amount),
            }
            
            try:
                wb_out = _fill_excel_template(template_name, ctx)
            except FileNotFoundError:
                return HttpResponse(f"LS template not found: {template_name}", status=404)
            except Exception as e:
                return HttpResponse(f"Error filling LS template: {e}", status=500)

        _apply_print_settings(wb_out)
        resp = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        resp["Content-Disposition"] = f'attachment; filename="{download_name}"'
        wb_out.save(resp)
        return resp

    # 11) COVERING LETTER / MOVEMENT SLIP (WORD) - Multiple pages support
    if doc_kind in ("covering", "movement"):
        from core.template_views import get_user_template
        
        if doc_kind == "covering":
            template_type = "covering_letter"
            base_download_name = "Cover_Letter"
            template_type_display = "Covering Letter"
        else:
            template_type = "movement_slip"
            base_download_name = "Movement_Slip"
            template_type_display = "Movement Slip"
        
        # Check if user has uploaded their template
        user_template = get_user_template(request.user, template_type)
        
        if not user_template:
            # No user template - show styled popup-like page with link to upload
            error_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <title>Template Required</title>
                <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet">
                <link href="https://cdn.jsdelivr.net/npm/bootstrap-icons@1.10.0/font/bootstrap-icons.css" rel="stylesheet">
                <style>
                    body {{ background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); min-height: 100vh; display: flex; align-items: center; justify-content: center; }}
                    .card {{ max-width: 500px; border-radius: 15px; box-shadow: 0 10px 40px rgba(0,0,0,0.2); }}
                    .icon-warning {{ font-size: 4rem; color: #ffc107; }}
                </style>
            </head>
            <body>
                <div class="card text-center p-5">
                    <i class="bi bi-exclamation-triangle-fill icon-warning mb-3"></i>
                    <h3 class="mb-3">{template_type_display} Template Required</h3>
                    <p class="text-muted mb-4">
                        You haven't uploaded a <strong>{template_type_display}</strong> template yet.<br>
                        Please upload your own Word template (.docx) with your officer names and formatting.
                    </p>
                    <div class="d-grid gap-2">
                        <a href="/templates/" class="btn btn-primary btn-lg">
                            <i class="bi bi-upload me-2"></i>Upload Template
                        </a>
                        <a href="/bill/" class="btn btn-outline-secondary">
                            <i class="bi bi-arrow-left me-2"></i>Back to Bill Generator
                        </a>
                    </div>
                </div>
            </body>
            </html>
            """
            return HttpResponse(error_html, status=200)
        
        # Read template bytes from DB (survives redeployments)
        template_bytes = user_template.get_file_bytes()
        if not template_bytes:
            return HttpResponse("Template file not found. Please re-upload your template.", status=404)

        def replace_in_paragraphs(paragraphs, placeholder_map, mm_yyyy_val):
            for p in paragraphs:
                if not p.runs:
                    continue
                for run in p.runs:
                    if not run.text:
                        continue
                    text = run.text
                    changed = False
                    for ph, val in placeholder_map.items():
                        if ph in text:
                            text = text.replace(ph, val or "")
                            changed = True
                    if "dd.mm.yyyy" in text:
                        text = text.replace("dd.mm.yyyy", f"  .{mm_yyyy_val}")  # two spaces before dot
                        changed = True
                    if changed:
                        run.text = text

        # If multiple bill sheets, create combined document with page breaks
        if has_multiple_bills:
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            from copy import deepcopy
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn as _qn

            # Start with the template and clear its content
            combined_doc = Document(io.BytesIO(template_bytes))

            # Clear the template content first
            for element in list(combined_doc.element.body):
                combined_doc.element.body.remove(element)

            for sheet_idx, bill_ws in enumerate(bill_sheets, start=1):
                # Extract per-sheet header data
                sheet_header = _extract_header_data_from_sheet(bill_ws)
                sheet_name_of_work = sheet_header.get("name_of_work", "") or ""
                sheet_agreement_ref = sheet_header.get("agreement", "") or ""
                sheet_agency_name = sheet_header.get("agency", "") or ""

                # Per-sheet MB: use user-entered MB if provided, else try sheet-level MB
                sheet_mb = mb_details_str
                if not user_entered_mb:
                    sheet_mb_from_file = (sheet_header.get("mb_details") or "").strip()
                    if sheet_mb_from_file:
                        sheet_mb = sheet_mb_from_file

                # Get formulas sheet if available
                ws_formulas = None
                if wb_formulas:
                    try:
                        ws_formulas = wb_formulas[bill_ws.title]
                    except:
                        pass

                # Calculate total for this sheet
                sheet_total = _extract_total_from_sheet(bill_ws, ws_formulas)
                try:
                    sheet_total_str = f"{float(sheet_total):,.2f}"
                except:
                    sheet_total_str = str(sheet_total)
                sheet_amount_words = _number_to_words_rupees(sheet_total)

                # Build placeholder map for this sheet
                sheet_placeholder_map = {
                    "{{NAME_OF_WORK}}": sheet_name_of_work,
                    "{{AGENCY_NAME}}": sheet_agency_name,
                    "{{NAME_OF_AGENCY}}": sheet_agency_name,
                    "{{AGREEMENT_REF}}": sheet_agreement_ref,
                    "{{REF_OF_AGREEMENT}}": sheet_agreement_ref,
                    "{{CC_HEADER}}": cc_header,
                    "{{MB_DETAILS}}": sheet_mb,
                    "{{AMOUNT}}": sheet_total_str,
                    "{{TOTAL_AMOUNT}}": sheet_total_str,
                    "{{AMOUNT_IN_WORDS}}": sheet_amount_words,
                }

                # Load fresh template for this sheet
                sheet_doc = Document(io.BytesIO(template_bytes))

                # Replace placeholders in body
                replace_in_paragraphs(sheet_doc.paragraphs, sheet_placeholder_map, mm_yyyy)

                # Replace in tables
                for table in sheet_doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            replace_in_paragraphs(cell.paragraphs, sheet_placeholder_map, mm_yyyy)

                # Add page break at the END of this document (except for the last one)
                # This ensures each estimate starts on a new page without blank pages
                if sheet_idx < len(bill_sheets):
                    sheet_doc.add_page_break()

                # Deep-copy elements before appending to avoid shared references
                elements_to_add = [deepcopy(el) for el in sheet_doc.element.body]

                # For sheets after the first, restart any Word auto-numbering
                # so reference numbers don't continue from the previous letter
                if sheet_idx > 1:
                    try:
                        numbering_part = combined_doc.part.numbering_part._element
                        # Collect unique numId values used in these elements
                        used_numIds = set()
                        numId_elems = []
                        for el in elements_to_add:
                            for nid_el in el.iter(_qn('w:numId')):
                                val = nid_el.get(_qn('w:val'))
                                if val and val != '0':
                                    used_numIds.add(int(val))
                                    numId_elems.append(nid_el)

                        if used_numIds:
                            # Find max existing numId
                            max_num_id = 0
                            for num_elem in numbering_part.findall(_qn('w:num')):
                                nid = int(num_elem.get(_qn('w:numId'), '0'))
                                if nid > max_num_id:
                                    max_num_id = nid

                            # Create new numbering instances that restart at 1
                            id_map = {}
                            for old_id in used_numIds:
                                original_num = None
                                for num_elem in numbering_part.findall(_qn('w:num')):
                                    if int(num_elem.get(_qn('w:numId'), '0')) == old_id:
                                        original_num = num_elem
                                        break
                                if original_num is None:
                                    continue

                                max_num_id += 1
                                new_num = deepcopy(original_num)
                                new_num.set(_qn('w:numId'), str(max_num_id))

                                # Remove existing overrides
                                for ov in list(new_num.findall(_qn('w:lvlOverride'))):
                                    new_num.remove(ov)

                                # Add restart override for level 0
                                lvl_override = OxmlElement('w:lvlOverride')
                                lvl_override.set(_qn('w:ilvl'), '0')
                                start_ov = OxmlElement('w:startOverride')
                                start_ov.set(_qn('w:val'), '1')
                                lvl_override.append(start_ov)
                                new_num.append(lvl_override)

                                numbering_part.append(new_num)
                                id_map[old_id] = max_num_id

                            # Update numId references in elements to use new numbering
                            for nid_el in numId_elems:
                                old_val = int(nid_el.get(_qn('w:val'), '0'))
                                if old_val in id_map:
                                    nid_el.set(_qn('w:val'), str(id_map[old_val]))
                    except Exception:
                        pass  # If numbering manipulation fails, proceed anyway

                # Append elements to combined document
                for element in elements_to_add:
                    combined_doc.element.body.append(element)
            
            buf = io.BytesIO()
            combined_doc.save(buf)
            buf.seek(0)
            
            download_name = f"{base_download_name}s.docx"
            response = HttpResponse(
                buf.getvalue(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            response["Content-Disposition"] = f'attachment; filename="{download_name}"'
            return response
        
        else:
            # Single sheet - original logic
            header = _extract_header_data_fuzzy_from_wb(wb_in)
            total_amount = _extract_total_from_sheet(bill_sheets[0], wb_formulas[bill_sheets[0].title] if wb_formulas else None)
            try:
                total_amount_str = f"{float(total_amount):,.2f}"
            except:
                total_amount_str = str(total_amount)
            
            placeholder_map = {
                "{{NAME_OF_WORK}}": header.get("name_of_work", ""),
                "{{AGENCY_NAME}}": header.get("agency", ""),
                "{{NAME_OF_AGENCY}}": header.get("agency", ""),
                "{{AGREEMENT_REF}}": header.get("agreement", ""),
                "{{REF_OF_AGREEMENT}}": header.get("agreement", ""),
                "{{CC_HEADER}}": cc_header,
                "{{MB_DETAILS}}": mb_details_str,
                "{{AMOUNT}}": total_amount_str,
                "{{TOTAL_AMOUNT}}": total_amount_str,
                "{{AMOUNT_IN_WORDS}}": _number_to_words_rupees(total_amount),
            }

            doc = Document(io.BytesIO(template_bytes))
            replace_in_paragraphs(doc.paragraphs, placeholder_map, mm_yyyy)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        replace_in_paragraphs(cell.paragraphs, placeholder_map, mm_yyyy)

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)

            download_name = f"{base_download_name}.docx"
            response = HttpResponse(
                buf.getvalue(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            response["Content-Disposition"] = f'attachment; filename="{download_name}"'
            return response

    return HttpResponse(f"Unknown doc_kind: {doc_kind}", status=400)


@org_required
def self_formatted_document(request):
    """
    Generate self-formatted documents asynchronously.
    
    Now uses async job processing instead of in-request generation.
    
    User uploads:
      - bill_file      : Excel (your Bill / Estimate / Nth Bill)
      - template_file  : DOCX or XLSX (user's own format with {{PLACEHOLDERS}})
      - action         : same as bill() (estimate_first_part, workslip_first_final, ...)
      - nth_number     : N value for Nth bills (optional)
      - mb_measure_no, mb_measure_p_from, mb_measure_p_to
      - mb_abstract_no, mb_abstract_p_from, mb_abstract_p_to
    """
    if request.method != "POST":
        return HttpResponseNotAllowed(["POST"], "Use POST to generate self-formatted documents.")

    org = get_org_from_request(request)

    # ---------- Required uploads ----------
    bill_file = request.FILES.get("bill_file")
    template_file = request.FILES.get("template_file")

    if not bill_file:
        return JsonResponse({"error": "Please upload the Bill Excel file (bill_file)."}, status=400)
    if not template_file:
        return JsonResponse({"error": "Please upload the template file (template_file)."}, status=400)

    # ---------- Meta info ----------
    action = (request.POST.get("action") or "").strip()
    nth_number_str = (request.POST.get("nth_number") or "").strip()

    # MB details
    mb_measure_no = (request.POST.get("mb_measure_no") or "").strip()
    mb_measure_p_from = (request.POST.get("mb_measure_p_from") or "").strip()
    mb_measure_p_to = (request.POST.get("mb_measure_p_to") or "").strip()
    mb_abs_no = (request.POST.get("mb_abstract_no") or "").strip()
    mb_abs_p_from = (request.POST.get("mb_abstract_p_from") or "").strip()
    mb_abs_p_to = (request.POST.get("mb_abstract_p_to") or "").strip()

    # Create Upload for bill file
    try:
        upload = Upload.objects.create(
            organization=org,
            user=request.user,
            filename=bill_file.name,
            file_size=bill_file.size,
            status='processing'
        )
    except Exception as e:
        return JsonResponse({"error": f"Failed to save upload: {e}"}, status=400)

    # Enqueue async job
    metadata = {
        'action': action,
        'nth_number': nth_number_str,
        'mb_measure_no': mb_measure_no,
        'mb_measure_p_from': mb_measure_p_from,
        'mb_measure_p_to': mb_measure_p_to,
        'mb_abs_no': mb_abs_no,
        'mb_abs_p_from': mb_abs_p_from,
        'mb_abs_p_to': mb_abs_p_to,
        'template_filename': template_file.name,
        'upload_id': upload.id,
    }

    try:
        job, task = create_job_for_excel(
            request,
            upload=upload,
            job_type='generate_self_formatted_document',
            metadata=metadata
        )
        
        return JsonResponse({
            'job_id': job.id,
            'status_url': reverse('job_status', args=[job.id]),
            'message': 'Generating self-formatted document. You will be notified when ready.'
        })
    except Exception as e:
        return JsonResponse({"error": f"Failed to enqueue task: {e}"}, status=500)


# -----------------------
# SIMPLE PAGES & PROJECTS
# -----------------------
