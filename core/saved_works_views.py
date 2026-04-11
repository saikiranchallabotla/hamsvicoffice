# core/saved_works_views.py
"""
Views for the Saved Works feature.
Allows users to save their work-in-progress and resume later.
"""

import json
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.views.decorators.http import require_POST, require_GET
from django.http import JsonResponse
from django.contrib import messages
from django.db.models import Q
from django.utils import timezone
from django.urls import reverse

from .models import SavedWork, WorkFolder, Organization, Membership
from .decorators import org_required


# ==============================================================================
# SUBSCRIPTION ACCESS CONTROL FOR SAVED WORKS
# ==============================================================================

# Mapping from work_type to module_code for subscription checks
WORK_TYPE_TO_MODULE = {
    'new_estimate': 'new_estimate',
    'temporary_works': 'temp_works',
    'amc': 'amc',
    'workslip': 'workslip',
    'bill': 'bill',
}


def check_saved_work_access(user, saved_work):
    """
    Check if user has active subscription to access a saved work.
    
    Args:
        user: Django User object
        saved_work: SavedWork instance
        
    Returns:
        dict: {ok: bool, reason: str, module_code: str}
    """
    work_type = saved_work.work_type
    module_code = WORK_TYPE_TO_MODULE.get(work_type)
    
    if not module_code:
        # Unknown work type - allow access (fallback)
        return {'ok': True, 'reason': 'No module restriction', 'module_code': None}
    
    try:
        from subscriptions.services import SubscriptionService
        result = SubscriptionService.check_access(user, module_code)
        return {
            'ok': result.get('ok', False),
            'reason': result.get('reason', 'Access denied'),
            'module_code': module_code,
            'data': result.get('data', {})
        }
    except Exception as e:
        # If subscription service fails, deny access for safety
        import logging
        logging.error(f"Subscription check failed for saved work: {e}")
        return {
            'ok': False, 
            'reason': 'Unable to verify subscription. Please try again.',
            'module_code': module_code
        }


def get_org_from_request(request):
    """Get organization from request, creating default if needed."""
    if hasattr(request, 'organization') and request.organization:
        return request.organization
    
    if request.user.is_authenticated:
        from django.utils.text import slugify
        
        membership = Membership.objects.filter(user=request.user).select_related('organization').first()
        
        if membership:
            request.organization = membership.organization
            return membership.organization
        
        org_name = f"{request.user.username}'s Organization"
        org_slug = slugify(org_name)[:255]
        
        base_slug = org_slug
        counter = 1
        while Organization.objects.filter(slug=org_slug).exists():
            org_slug = f"{base_slug}-{counter}"
            counter += 1
        
        org, created = Organization.objects.get_or_create(
            name=org_name,
            defaults={
                'slug': org_slug,
                'owner': request.user,
                'is_active': True
            }
        )
        
        Membership.objects.get_or_create(
            user=request.user,
            organization=org,
            defaults={'role': 'owner'}
        )
        
        request.organization = org
        return org
    
    from django.http import Http404
    raise Http404("Please login to continue.")


# ==============================================================================
# HELPER FUNCTIONS
# ==============================================================================

def build_folder_tree_flat(all_folders):
    """
    Build a flat list of folders in DFS order with depth info for indented <select>.
    Returns: [{'id': ..., 'name': ..., 'depth': 0, 'color': ..., 'parent_id': ...}, ...]
    """
    folders_by_parent = {}
    for f in all_folders:
        folders_by_parent.setdefault(f.parent_id, []).append(f)

    result = []
    def walk(parent_id, depth):
        for f in sorted(folders_by_parent.get(parent_id, []), key=lambda x: x.name):
            result.append({
                'id': f.id,
                'name': f.name,
                'depth': depth,
                'color': f.color,
                'parent_id': f.parent_id,
            })
            walk(f.id, depth + 1)
    walk(None, 0)
    return result


def get_descendant_folder_ids(folder_id, all_folders):
    """Return set of all descendant folder IDs for the given folder_id."""
    children_map = {}
    for f in all_folders:
        children_map.setdefault(f.parent_id, []).append(f.id)

    result = set()
    stack = list(children_map.get(folder_id, []))
    while stack:
        fid = stack.pop()
        result.add(fid)
        stack.extend(children_map.get(fid, []))
    return result


# ==============================================================================
# SAVED WORKS LIST & MANAGEMENT
# ==============================================================================

@login_required(login_url='login')
def saved_works_list(request):
    """List all saved works with folder structure."""
    org = get_org_from_request(request)
    user = request.user
    
    # Get filter parameters
    work_type_filter = request.GET.get('type', 'all')
    folder_id = request.GET.get('folder')
    search_query = request.GET.get('q', '').strip()
    status_filter = request.GET.get('status', 'all')
    
    # Get current folder if viewing inside a folder
    current_folder = None
    if folder_id and folder_id != 'unfiled':
        current_folder = get_object_or_404(WorkFolder, id=folder_id, organization=org, user=user)
    
    # Get folders to display (root folders or subfolders of current folder)
    if current_folder:
        folders = WorkFolder.objects.filter(organization=org, user=user, parent=current_folder)
    else:
        folders = WorkFolder.objects.filter(organization=org, user=user, parent__isnull=True)
    
    # Get saved works
    works = SavedWork.objects.filter(organization=org, user=user)
    
    # Fix orphan workslips/bills: link them to their parent estimate
    # This handles works created before parent-linking was implemented
    orphan_children = works.filter(
        parent__isnull=True,
        work_type__in=['workslip', 'bill'],
    )
    for oc in orphan_children:
        wd = oc.work_data or {}
        if oc.work_type == 'workslip':
            # For workslips, link to the source estimate
            src_id = wd.get('ws_source_estimate_id')
            if src_id:
                try:
                    parent_work = SavedWork.objects.get(id=int(src_id), organization=org, user=user)
                    # Walk up to root estimate
                    root = parent_work
                    while root and root.work_type not in ('new_estimate', 'temporary_works', 'amc') and root.parent:
                        root = root.parent
                    if root and root.work_type in ('new_estimate', 'temporary_works', 'amc'):
                        oc.parent = root
                        oc.save(update_fields=['parent'])
                except SavedWork.DoesNotExist:
                    pass
        elif oc.work_type == 'bill':
            # For bills, try to link to the source workslip first, then estimate
            bill_parent_id = wd.get('bill_parent_work_id') or wd.get('bill_source_work_id')
            if bill_parent_id:
                try:
                    bill_parent = SavedWork.objects.get(id=int(bill_parent_id), organization=org, user=user)
                    oc.parent = bill_parent
                    oc.save(update_fields=['parent'])
                except SavedWork.DoesNotExist:
                    # Fallback: try ws_source_estimate_id
                    src_id = wd.get('ws_source_estimate_id')
                    if src_id:
                        try:
                            root = SavedWork.objects.get(id=int(src_id), organization=org, user=user)
                            oc.parent = root
                            oc.save(update_fields=['parent'])
                        except SavedWork.DoesNotExist:
                            pass

    # Re-fetch works after fixing orphans
    works = SavedWork.objects.filter(organization=org, user=user)

    # Hide child works (workslips/bills with a parent) from the list.
    # They are accessible through the parent estimate's detail page.
    works = works.filter(parent__isnull=True)
    
    # Apply filters
    if work_type_filter != 'all':
        works = works.filter(work_type=work_type_filter)
    
    if folder_id:
        if folder_id == 'unfiled':
            works = works.filter(folder__isnull=True)
        else:
            works = works.filter(folder_id=folder_id)
    else:
        # At root, show only unfiled works
        works = works.filter(folder__isnull=True)
    
    if status_filter != 'all':
        works = works.filter(status=status_filter)
    
    if search_query:
        works = works.filter(Q(name__icontains=search_query) | Q(notes__icontains=search_query))
    
    # Build breadcrumb path for nested folders
    breadcrumb_path = []
    if current_folder:
        folder = current_folder
        while folder:
            breadcrumb_path.insert(0, folder)
            folder = folder.parent
    
    # Get all folders for the dropdown/tree
    all_folders = WorkFolder.objects.filter(organization=org, user=user)
    all_folders_tree = build_folder_tree_flat(all_folders)
    
    # Check module access for each work type
    # Mapping: work_type -> module_code
    work_type_to_module = {
        'new_estimate': 'new_estimate',
        'temporary_works': 'temp_works',
        'amc': 'amc',
        'workslip': 'workslip',
        'bill': 'bill',
    }

    module_access = {}
    try:
        from subscriptions.services import SubscriptionService
        for work_type, module_code in work_type_to_module.items():
            result = SubscriptionService.check_access(user, module_code)
            module_access[work_type] = result.get('ok', False)
        # Also check workslip access for the generate workslip button
        workslip_result = SubscriptionService.check_access(user, 'workslip')
        module_access['can_generate_workslip'] = workslip_result.get('ok', False)
        # Check bill access
        bill_result = SubscriptionService.check_access(user, 'bill')
        module_access['can_generate_bill'] = bill_result.get('ok', False)
    except Exception:
        # If subscription service fails, deny access (secure fallback)
        for work_type in work_type_to_module.keys():
            module_access[work_type] = False
        module_access['can_generate_workslip'] = False
        module_access['can_generate_bill'] = False

    # Evaluate queryset to list so we can attach workflow chain data
    works_list = list(works)

    # Build workflow chain data for ALL works in the current view.
    # Attach .wf_chain to each work object so the template can render
    # the full E/W/B navigation for every row.
    _chain_cache = {}  # Cache chains by root estimate id

    def _build_chain_for_estimate(est):
        """Build and cache the E/W/B chain for an estimate."""
        if est is None:
            return None
        if est.id in _chain_cache:
            return _chain_cache[est.id]
        # Direct workslip children
        ws_list = list(
            SavedWork.objects.filter(
                organization=org, user=user,
                work_type='workslip', parent=est,
            ).order_by('workslip_number')
        )
        # Also find workslips chained to other workslips (W2→W1, W3→W2)
        # and re-parent them to the root estimate
        ws_ids = {ws.id for ws in ws_list}
        depth = 0
        search_ids = set(ws_ids)
        while depth < 10:
            chained = list(
                SavedWork.objects.filter(
                    organization=org, user=user,
                    work_type='workslip', parent_id__in=search_ids,
                ).exclude(id__in=ws_ids)
            )
            if not chained:
                break
            new_ids = set()
            for cw in chained:
                cw.parent = est
                cw.save(update_fields=['parent'])
                ws_list.append(cw)
                ws_ids.add(cw.id)
                new_ids.add(cw.id)
            search_ids = new_ids
            depth += 1
        ws_list.sort(key=lambda w: w.workslip_number or 0)

        bill_list = list(
            SavedWork.objects.filter(
                organization=org, user=user, work_type='bill',
            ).filter(
                Q(parent=est) | Q(parent_id__in=ws_ids)
            ).order_by('bill_number')
        )
        chain = {
            'estimate': est,
            'workslips': ws_list,
            'bills': bill_list,
        }
        _chain_cache[est.id] = chain
        return chain

    for work in works_list:
        # Build wf_chain for E/W/B navigation
        if work.work_type == 'new_estimate':
            work.wf_chain = _build_chain_for_estimate(work)
        elif work.work_type in ('workslip', 'bill'):
            root_est = None
            current = work.parent
            while current:
                if current.work_type == 'new_estimate':
                    root_est = current
                    break
                current = current.parent
            work.wf_chain = _build_chain_for_estimate(root_est)
        else:
            work.wf_chain = None

        # Also build workslip children data for dynamic W1, W2, W3... buttons
        if work.work_type in ('new_estimate', 'temporary_works', 'amc'):
            try:
                all_ws = work.get_all_workslips() if hasattr(work, 'get_all_workslips') else []
                work.workslip_children = all_ws
                work.next_ws_number = (max(ws.workslip_number for ws in all_ws) + 1) if all_ws else 1
                work.last_ws = all_ws[-1] if all_ws else None
                # Attach bill children from wf_chain
                work.bill_children = work.wf_chain.get('bills', []) if work.wf_chain else []
                work.next_bill_number = (max(b.bill_number for b in work.bill_children) + 1) if work.bill_children else 1
                # Attach subsequent_bills to each bill for edit-confirmation popups
                all_bills = work.bill_children
                for bill in all_bills:
                    bill.subsequent_bills = [b for b in all_bills if b.bill_number > bill.bill_number]
                # Attach bill_children_list to each workslip for edit-confirmation popups
                for ws in all_ws:
                    ws.bill_children_list = list(ws.children.filter(work_type='bill').order_by('bill_number'))
            except Exception:
                work.workslip_children = []
                work.next_ws_number = 1
                work.last_ws = None
                work.bill_children = []
                work.next_bill_number = 1

        # Attach subscription access flags for template gating
        work.module_code = work_type_to_module.get(work.work_type, work.work_type)
        work.has_estimate_access = module_access.get(work.work_type, False)

        # Attach item_count for display
        wd = work.work_data or {}
        if work.work_type == 'new_estimate':
            work.item_count = len(wd.get('fetched_items', []))
        elif work.work_type == 'workslip':
            work.item_count = len(wd.get('ws_estimate_rows', []))
        elif work.work_type == 'bill':
            work.item_count = len(wd.get('bill_ws_rows', wd.get('ws_estimate_rows', [])))
        else:
            work.item_count = 0

        # Build sibling workslips for workslip cards
        if work.work_type == 'workslip':
            # Attach bill children for edit-confirmation popup
            work.bill_children_list = list(work.children.filter(work_type='bill').order_by('bill_number'))
            try:
                root = work.get_root_estimate() if hasattr(work, 'get_root_estimate') else None
                if root:
                    all_ws = root.get_all_workslips() if hasattr(root, 'get_all_workslips') else []
                    work.sibling_workslips = all_ws
                    work.root_estimate = root
                else:
                    work.sibling_workslips = []
                    work.root_estimate = None
            except Exception:
                work.sibling_workslips = []
                work.root_estimate = None

        # Attach subsequent bills for standalone bill cards
        if work.work_type == 'bill':
            try:
                root = work.get_root_estimate() if hasattr(work, 'get_root_estimate') else None
                if root:
                    all_bills = root.get_all_bills() if hasattr(root, 'get_all_bills') else []
                    work.subsequent_bills = [b for b in all_bills if b.bill_number > work.bill_number]
                else:
                    work.subsequent_bills = []
            except Exception:
                work.subsequent_bills = []
    
    context = {
        'works': works_list,
        'folders': folders,
        'all_folders': all_folders,
        'all_folders_tree': all_folders_tree,
        'current_folder': current_folder,
        'breadcrumb_path': breadcrumb_path,
        'work_type_filter': work_type_filter,
        'status_filter': status_filter,
        'search_query': search_query,
        'work_type_choices': SavedWork.WORK_TYPE_CHOICES,
        'status_choices': SavedWork.STATUS_CHOICES,
        'module_access': module_access,
    }

    return render(request, 'core/saved_works/list.html', context)


@login_required(login_url='login')
@require_POST
def create_folder(request):
    """Create a new folder for organizing saved works."""
    org = get_org_from_request(request)
    user = request.user
    
    # Handle both form data and JSON
    if request.content_type and 'application/json' in request.content_type:
        import json
        data = json.loads(request.body.decode('utf-8'))
        name = data.get('name', '').strip()
        parent_id = data.get('parent_id')
        color = data.get('color', '#6366f1')
        description = data.get('description', '').strip()
    else:
        name = request.POST.get('name', '').strip()
        parent_id = request.POST.get('parent_id')
        color = request.POST.get('color', '#6366f1')
        description = request.POST.get('description', '').strip()
    
    if not name:
        return JsonResponse({'success': False, 'error': 'Folder name is required.'})
    
    parent = None
    if parent_id:
        parent = get_object_or_404(WorkFolder, id=parent_id, organization=org, user=user)
    
    # Check for duplicate name in same parent
    if WorkFolder.objects.filter(organization=org, user=user, name=name, parent=parent).exists():
        return JsonResponse({'success': False, 'error': 'A folder with this name already exists.'})
    
    folder = WorkFolder.objects.create(
        organization=org,
        user=user,
        name=name,
        parent=parent,
        color=color,
        description=description,
    )
    
    return JsonResponse({
        'success': True,
        'folder_id': folder.id,
        'folder_name': folder.name,
        'message': f'Folder "{name}" created successfully!'
    })


@login_required(login_url='login')
@require_POST
def rename_folder(request, folder_id):
    """Rename an existing folder."""
    org = get_org_from_request(request)
    user = request.user
    
    folder = get_object_or_404(WorkFolder, id=folder_id, organization=org, user=user)
    
    new_name = request.POST.get('name', '').strip()
    if not new_name:
        return JsonResponse({'success': False, 'error': 'Folder name is required.'})
    
    # Check for duplicate name in same parent
    if WorkFolder.objects.filter(
        organization=org, user=user, name=new_name, parent=folder.parent
    ).exclude(id=folder.id).exists():
        return JsonResponse({'success': False, 'error': 'A folder with this name already exists.'})
    
    folder.name = new_name
    folder.save()
    
    return JsonResponse({
        'success': True,
        'message': f'Folder renamed to "{new_name}"!'
    })


@login_required(login_url='login')
@require_POST
def delete_folder(request, folder_id):
    """Delete a folder and all its contents permanently."""
    from django.db import transaction
    
    org = get_org_from_request(request)
    user = request.user
    
    folder = get_object_or_404(WorkFolder, id=folder_id, organization=org, user=user)
    
    permanent = request.POST.get('permanent', 'false').lower() == 'true'
    
    folder_name = folder.name
    
    # Use transaction to ensure atomic deletion
    with transaction.atomic():
        if permanent:
            # Permanently delete all works and subfolders recursively
            def delete_folder_contents(f):
                # Delete all works in this folder
                f.saved_works.all().delete()
                # Recursively delete child folders
                for child in f.children.all():
                    delete_folder_contents(child)
                    child.delete()
            
            delete_folder_contents(folder)
        else:
            # Move saved works to parent folder or root
            folder.saved_works.update(folder=folder.parent)
            folder.children.update(parent=folder.parent)
        
        folder.delete()
    
    return JsonResponse({
        'success': True,
        'message': f'Folder "{folder_name}" deleted successfully!'
    })


@login_required(login_url='login')
@require_POST
def move_folder_to(request, folder_id):
    """Move a folder into another folder (or to root)."""
    org = get_org_from_request(request)
    user = request.user

    folder = get_object_or_404(WorkFolder, id=folder_id, organization=org, user=user)
    target_id = request.POST.get('folder_id')

    if target_id and target_id != 'none':
        target_folder = get_object_or_404(WorkFolder, id=target_id, organization=org, user=user)

        # Prevent circular move: can't move into self or any descendant
        all_user_folders = WorkFolder.objects.filter(organization=org, user=user)
        descendants = get_descendant_folder_ids(folder.id, all_user_folders)
        if target_folder.id == folder.id or target_folder.id in descendants:
            return JsonResponse({
                'success': False,
                'error': 'Cannot move a folder into itself or one of its subfolders.'
            })

        folder.parent = target_folder
    else:
        folder.parent = None

    # Check for duplicate name in target location
    if WorkFolder.objects.filter(
        organization=org, user=user, name=folder.name, parent=folder.parent
    ).exclude(id=folder.id).exists():
        return JsonResponse({
            'success': False,
            'error': f'A folder named "{folder.name}" already exists in the destination.'
        })

    folder.save()

    return JsonResponse({
        'success': True,
        'message': f'Folder "{folder.name}" moved successfully!'
    })


# ==============================================================================
# SAVE WORK FUNCTIONALITY
# ==============================================================================

@login_required(login_url='login')
@require_POST
def save_work(request):
    """
    Save current work-in-progress.
    Called from any module (New Estimate, Workslip, Temporary Works, AMC).
    """
    org = get_org_from_request(request)
    user = request.user
    
    # Get work details from POST
    work_name = request.POST.get('work_name', '').strip()
    work_type = request.POST.get('work_type', '')
    folder_id = request.POST.get('folder_id')
    notes = request.POST.get('notes', '').strip()
    category = request.POST.get('category', 'electrical')
    work_id = request.POST.get('work_id')  # For updating existing saved work
    
    if not work_name:
        return JsonResponse({'success': False, 'error': 'Work name is required.'})
    
    if work_type not in dict(SavedWork.WORK_TYPE_CHOICES):
        return JsonResponse({'success': False, 'error': 'Invalid work type.'})
    
    # Get folder if specified
    folder = None
    if folder_id:
        folder = get_object_or_404(WorkFolder, id=folder_id, organization=org, user=user)
    
    # Collect work data from session based on work type
    work_data = collect_work_data(request, work_type)
    
    # Calculate progress based on work data
    progress_percent = calculate_progress(work_data, work_type)
    last_step = get_last_step(request, work_type)
    
    if work_id:
        # Check if existing saved work has a different work_type
        # If so, create a new record instead of overwriting (to preserve estimate when saving workslip)
        existing_work = get_object_or_404(SavedWork, id=work_id, organization=org, user=user)
        
        if existing_work.work_type != work_type:
            # Different work type - create a new record, don't overwrite
            # This prevents overwriting estimate data when saving from workslip
            
            # For workslips, set the workslip_number and parent
            workslip_number = 1
            parent = None
            if work_type == 'workslip':
                workslip_number = request.session.get('ws_target_workslip', 1) or 1
                # Link workslip to its source estimate
                source_est_id = request.session.get('ws_source_estimate_id')
                if source_est_id:
                    try:
                        parent = SavedWork.objects.get(id=source_est_id, organization=org, user=user)
                    except SavedWork.DoesNotExist:
                        pass
                if not parent:
                    parent = existing_work  # fallback: link to the work we came from

            # For bills, set bill_number and parent from session
            bill_number = 1
            if work_type == 'bill':
                bill_number = request.session.get('bill_target_number', 1) or 1
                parent_id = request.session.get('bill_parent_work_id')
                if parent_id:
                    try:
                        parent = SavedWork.objects.get(id=parent_id, organization=org, user=user)
                    except SavedWork.DoesNotExist:
                        pass

            # Auto-inherit parent's folder if no folder specified
            child_folder = folder
            if not child_folder and parent and parent.folder:
                child_folder = parent.folder

            saved_work = SavedWork.objects.create(
                organization=org,
                user=user,
                folder=child_folder,
                parent=parent,
                name=work_name,
                work_type=work_type,
                work_data=work_data,
                category=category,
                notes=notes,
                progress_percent=progress_percent,
                last_step=last_step,
                workslip_number=workslip_number,
                bill_number=bill_number,
            )
            message = f'Work "{work_name}" saved successfully!'
        else:
            # Same work type - update existing
            saved_work = existing_work
            saved_work.name = work_name
            saved_work.folder = folder
            saved_work.notes = notes
            saved_work.work_data = work_data
            saved_work.progress_percent = progress_percent
            saved_work.last_step = last_step

            # Update workslip_number if this is a workslip
            if work_type == 'workslip':
                saved_work.workslip_number = request.session.get('ws_target_workslip', 1) or 1

            # Update bill_number if this is a bill
            if work_type == 'bill':
                saved_work.bill_number = request.session.get('bill_target_number', 1) or 1

            saved_work.save()
            message = f'Work "{work_name}" updated successfully!'

            # ── Propagate estimate qty changes to child workslips ──
            if work_type == 'new_estimate':
                _propagate_estimate_to_children(saved_work, work_data)

            # ── Back-propagate metadata from workslip to parent estimate ──
            if work_type == 'workslip' and saved_work.parent_id:
                _backpropagate_metadata_to_estimate(saved_work, work_data)
    else:
        # Create new saved work

        # For workslips, set the workslip_number and parent
        workslip_number = 1
        parent = None
        if work_type == 'workslip':
            workslip_number = request.session.get('ws_target_workslip', 1) or 1
            # Link workslip to its source estimate
            source_est_id = request.session.get('ws_source_estimate_id')
            if source_est_id:
                try:
                    parent = SavedWork.objects.get(id=source_est_id, organization=org, user=user)
                except SavedWork.DoesNotExist:
                    pass

        # For bills, set bill_number and parent from session
        bill_number = 1
        if work_type == 'bill':
            bill_number = request.session.get('bill_target_number', 1) or 1
            parent_id = request.session.get('bill_parent_work_id')
            if parent_id:
                try:
                    parent = SavedWork.objects.get(id=parent_id, organization=org, user=user)
                except SavedWork.DoesNotExist:
                    pass

        # Auto-inherit parent's folder if no folder specified
        child_folder = folder
        if not child_folder and parent and parent.folder:
            child_folder = parent.folder

        saved_work = SavedWork.objects.create(
            organization=org,
            user=user,
            folder=child_folder,
            parent=parent,
            name=work_name,
            work_type=work_type,
            work_data=work_data,
            category=category,
            notes=notes,
            progress_percent=progress_percent,
            last_step=last_step,
            workslip_number=workslip_number,
            bill_number=bill_number,
        )
        message = f'Work "{work_name}" saved successfully!'
    
    # Store saved work ID in session for quick access
    request.session['current_saved_work_id'] = saved_work.id
    
    return JsonResponse({
        'success': True,
        'work_id': saved_work.id,
        'message': message
    })


# ==============================================================================
# PROPAGATE ESTIMATE CHANGES TO CHILD WORKSLIPS / BILLS
# ==============================================================================

def _propagate_estimate_to_children(estimate_work, estimate_data):
    """
    When an estimate is re-saved, update qty_est in every child workslip's
    ws_estimate_rows so the workslip always reflects the latest estimate
    quantities, rates, and items.
    
    Also recalculates ws_estimate_grand_total for each child.
    """
    import logging
    logger = logging.getLogger(__name__)

    qty_map = estimate_data.get('qty_map', {})
    item_rates = estimate_data.get('item_rates', {})
    item_units = estimate_data.get('item_units', {})
    grand_total = estimate_data.get('grand_total', '')
    work_name = estimate_data.get('work_name', '')

    # Parse grand_total to float for children
    try:
        grand_total_float = float(grand_total) if grand_total else 0.0
    except (ValueError, TypeError):
        grand_total_float = 0.0

    # Find all descendant workslips — direct children of this estimate
    # AND chained workslips (Workslip-2 → parent=Workslip-1, etc.)
    # We walk the tree breadth-first so every workslip in the chain gets
    # its ws_estimate_rows updated.
    queue = list(SavedWork.objects.filter(parent=estimate_work, work_type='workslip'))
    visited = set()

    updated_count = 0
    while queue:
        child = queue.pop(0)
        if child.id in visited:
            continue
        visited.add(child.id)

        child_data = child.work_data or {}
        ws_rows = child_data.get('ws_estimate_rows', [])
        if not ws_rows:
            # Still enqueue grandchildren even if this child has no rows
            grandchildren = list(SavedWork.objects.filter(parent=child, work_type='workslip'))
            queue.extend(grandchildren)
            continue

        changed = False
        for row in ws_rows:
            item_name = row.get('item_name', '')
            if not item_name:
                continue

            # Update qty_est from estimate's qty_map
            new_qty = qty_map.get(item_name, 0)
            try:
                new_qty = float(new_qty) if new_qty else 0.0
            except (ValueError, TypeError):
                new_qty = 0.0

            old_qty = row.get('qty_est', 0)
            try:
                old_qty = float(old_qty) if old_qty else 0.0
            except (ValueError, TypeError):
                old_qty = 0.0

            if new_qty != old_qty:
                row['qty_est'] = new_qty
                changed = True

            # Update rate if available from estimate
            if item_name in item_rates:
                new_rate = float(item_rates[item_name] or 0)
                old_rate = float(row.get('rate', 0) or 0)
                if new_rate != old_rate:
                    row['rate'] = new_rate
                    changed = True

            # Update unit if available from estimate
            if item_name in item_units:
                new_unit = str(item_units[item_name])
                if new_unit != row.get('unit', ''):
                    row['unit'] = new_unit
                    changed = True

        if changed:
            child_data['ws_estimate_rows'] = ws_rows
            # Update grand total
            child_data['ws_estimate_grand_total'] = grand_total_float
            # Update work name if the estimate's work name changed
            if work_name:
                child_data['ws_work_name'] = work_name
            # Update metadata
            ws_meta = child_data.get('ws_metadata', {})
            if ws_meta:
                ws_meta['estimate_amount'] = str(grand_total_float) if grand_total_float else ''
                if work_name:
                    ws_meta['work_name'] = work_name
                ws_meta['grand_total'] = grand_total_float
                child_data['ws_metadata'] = ws_meta

            child.work_data = child_data
            child.save()
            updated_count += 1
            logger.info(f"[PROPAGATE] Updated child workslip '{child.name}' (ID={child.id}) with new estimate quantities")

        # Enqueue any grandchildren (next workslips in the chain)
        grandchildren = list(SavedWork.objects.filter(parent=child, work_type='workslip'))
        queue.extend(grandchildren)

    if updated_count:
        logger.info(f"[PROPAGATE] Propagated estimate changes to {updated_count} child workslip(s)")


def _backpropagate_metadata_to_estimate(workslip_work, workslip_data):
    """
    When a workslip is saved with metadata (admin_sanction, tech_sanction,
    agreement, agency_name), propagate those values back to the parent
    estimate's estimate_metadata field so they persist for the entire work.
    """
    import logging
    logger = logging.getLogger(__name__)

    ws_meta = workslip_data.get('ws_metadata', {})
    if not ws_meta:
        return

    # Walk up to root estimate
    parent = workslip_work.parent
    root_estimate = None
    while parent:
        if parent.work_type in ('new_estimate', 'temporary_works', 'amc'):
            root_estimate = parent
            break
        parent = parent.parent

    if not root_estimate:
        return

    est_data = root_estimate.work_data or {}
    est_meta = est_data.get('estimate_metadata', {})

    changed = False
    for field in ('admin_sanction', 'tech_sanction', 'agreement', 'agency_name'):
        new_val = ws_meta.get(field, '')
        old_val = est_meta.get(field, '')
        if new_val and new_val != old_val:
            est_meta[field] = new_val
            changed = True

    if changed:
        est_data['estimate_metadata'] = est_meta
        root_estimate.work_data = est_data
        root_estimate.save(update_fields=['work_data'])
        logger.info(f"[BACKPROP] Updated estimate '{root_estimate.name}' (ID={root_estimate.id}) metadata from workslip")


def collect_work_data(request, work_type):
    """Collect all relevant session data for a given work type.
    
    For new_estimate: prefer direct POST data (qty_map_json, etc.) over session
    so we don't depend on a prior save_qty_map request having flushed to DB.
    """
    import json as _json
    work_data = {}
    
    if work_type == 'new_estimate':
        # ---- qty_map: prefer direct POST JSON over session ----
        qty_map = request.session.get('qty_map', {})
        qty_map_json = request.POST.get('qty_map_json', '')
        if qty_map_json:
            try:
                parsed = _json.loads(qty_map_json)
                if isinstance(parsed, dict):
                    qty_map = parsed
                    # Also update session so subsequent reads are consistent
                    request.session['qty_map'] = qty_map
            except (ValueError, TypeError):
                pass

        # ---- unit_map: prefer direct POST JSON over session ----
        unit_map = request.session.get('unit_map', {})
        unit_map_json = request.POST.get('unit_map_json', '')
        if unit_map_json:
            try:
                parsed = _json.loads(unit_map_json)
                if isinstance(parsed, dict):
                    unit_map = parsed
                    request.session['unit_map'] = unit_map
            except (ValueError, TypeError):
                pass

        # ---- scalar fields: prefer POST *_value suffix over session ----
        def _post_or_session(post_key, session_key, default=''):
            val = request.POST.get(post_key, '')
            if val:
                request.session[session_key] = val
                return val
            return request.session.get(session_key, default)

        # ---- fetched_items: prefer direct POST JSON over session ----
        fetched_items = request.session.get('fetched_items', [])
        fetched_items_json = request.POST.get('fetched_items_json', '')
        if fetched_items_json:
            try:
                parsed = _json.loads(fetched_items_json)
                if isinstance(parsed, list) and parsed:
                    fetched_items = parsed
                    request.session['fetched_items'] = fetched_items
            except (ValueError, TypeError):
                pass

        work_data = {
            'fetched_items': fetched_items,
            'current_project_name': request.session.get('current_project_name'),
            'work_type': _post_or_session('work_type_value', 'work_type', 'original'),
            'qty_map': qty_map,
            'unit_map': unit_map,
            'work_name': _post_or_session('work_name_value', 'work_name', ''),
            'grand_total': _post_or_session('grand_total_value', 'grand_total', ''),
            'excess_tp_percent': _post_or_session('excess_tp_percent_value', 'excess_tp_percent', ''),
            'ls_special_name': _post_or_session('ls_special_name_value', 'ls_special_name', ''),
            'ls_special_amount': _post_or_session('ls_special_amount_value', 'ls_special_amount', ''),
            'deduct_old_material': _post_or_session('deduct_old_material_value', 'deduct_old_material', ''),
            'last_group': request.POST.get('group', ''),
            'selected_backend_id': request.session.get('selected_backend_id'),
            'item_rates': request.session.get('item_rates', {}),
            'item_units': request.session.get('item_units', {}),
            'item_descs': request.session.get('item_descs', {}),
            'estimate_source': request.session.get('estimate_source', ''),
            # Uploaded custom items data
            'uploaded_items': request.session.get('uploaded_items', []),
            'uploaded_file_id': request.session.get('uploaded_file_id'),
            'uploaded_item_blocks': request.session.get('uploaded_item_blocks', {}),
            'uploaded_sheet_name': request.session.get('uploaded_sheet_name', ''),
            # Estimate-level metadata: header fields unique to each work
            'estimate_metadata': request.session.get('estimate_metadata', {
                'admin_sanction': '',
                'tech_sanction': '',
                'agreement': '',
                'agency_name': '',
            }),
        }
        request.session.modified = True
    
    elif work_type == 'workslip':
        work_data = {
            'ws_estimate_rows': request.session.get('ws_estimate_rows', []),
            'ws_exec_map': request.session.get('ws_exec_map', {}),
            'ws_tp_percent': request.session.get('ws_tp_percent', 0.0),
            'ws_tp_type': request.session.get('ws_tp_type', 'Excess'),
            'ws_supp_items': request.session.get('ws_supp_items', []),
            'ws_estimate_grand_total': request.session.get('ws_estimate_grand_total', 0.0),
            'ws_work_name': request.session.get('ws_work_name', ''),
            'ws_deduct_old_material': request.session.get('ws_deduct_old_material', 0.0),
            'ws_lc_percent': request.session.get('ws_lc_percent', 0.0),
            'ws_qc_percent': request.session.get('ws_qc_percent', 0.0),
            'ws_nac_percent': request.session.get('ws_nac_percent', 0.0),
            'ws_current_phase': request.session.get('ws_current_phase', 1),
            'ws_target_workslip': request.session.get('ws_target_workslip', 1),
            'ws_previous_phases': request.session.get('ws_previous_phases', []),
            'ws_previous_ae_data': request.session.get('ws_previous_ae_data', []),
            'ws_previous_supp_items': request.session.get('ws_previous_supp_items', []),
            'ws_metadata': request.session.get('ws_metadata', {}),
            'selected_backend_id': request.session.get('ws_selected_backend_id'),
            'ws_source_estimate_id': request.session.get('ws_source_estimate_id'),
            'ws_work_mode': request.session.get('ws_work_mode', 'original'),
        }
    
    elif work_type == 'temporary_works':
        work_data = {
            'temp_entries': request.session.get('temp_entries', []),
            'temp_work_name': request.session.get('temp_work_name', ''),
            'temp_grand_total': request.session.get('temp_grand_total', ''),
            'temp_selected_backend_id': request.session.get('temp_selected_backend_id'),
            'temp_category': request.session.get('temp_category', 'electrical'),
            'last_group': request.POST.get('group', ''),
        }
    
    elif work_type == 'amc':
        work_data = {
            'amc_fetched_items': request.session.get('amc_fetched_items', []),
            'amc_qty_map': request.session.get('amc_qty_map', {}),
            'amc_category': request.session.get('amc_category', 'electrical'),
            'amc_work_name': request.session.get('amc_work_name', ''),
            'amc_grand_total': request.session.get('amc_grand_total', ''),
            'amc_selected_backend_id': request.session.get('amc_selected_backend_id'),
            'amc_work_type': request.session.get('amc_work_type', 'original'),
            'last_group': request.POST.get('group', ''),
        }

    elif work_type == 'bill':
        work_data = {
            'bill_source_work_id': request.session.get('bill_source_work_id'),
            'bill_source_work_type': request.session.get('bill_source_work_type', ''),
            'bill_source_work_name': request.session.get('bill_source_work_name', ''),
            'bill_from_workslip': request.session.get('bill_from_workslip', False),
            'bill_ws_rows': request.session.get('bill_ws_rows', []),
            'bill_ws_exec_map': request.session.get('bill_ws_exec_map', {}),
            'bill_ws_tp_percent': request.session.get('bill_ws_tp_percent', 0),
            'bill_ws_tp_type': request.session.get('bill_ws_tp_type', 'Excess'),
            'bill_target_number': request.session.get('bill_target_number', 1),
            'bill_type': request.session.get('bill_type', ''),
            'source_workslip_id': request.session.get('bill_source_work_id'),
            'bill_ws_metadata': request.session.get('bill_ws_metadata', {}),
            'bill_parent_work_id': request.session.get('bill_parent_work_id'),
            'ws_source_estimate_id': request.session.get('ws_source_estimate_id'),
        }

    return work_data


def calculate_progress(work_data, work_type):
    """Calculate progress percentage based on work data completeness."""
    if work_type == 'new_estimate':
        items = work_data.get('fetched_items', [])
        if not items:
            return 0
        qty_map = work_data.get('qty_map', {})
        if qty_map:
            return min(80, 20 + len(items) * 5)
        return min(50, len(items) * 5)
    
    elif work_type == 'workslip':
        rows = work_data.get('ws_estimate_rows', [])
        if not rows:
            return 0
        exec_map = work_data.get('ws_exec_map', {})
        if exec_map:
            return min(90, 30 + len(rows) * 3)
        return min(50, len(rows) * 3)
    
    elif work_type == 'temporary_works':
        items = work_data.get('temp_entries', [])
        if not items:
            return 0
        return min(80, len(items) * 10)
    
    elif work_type == 'bill':
        # Bill progress is determined by existence of bill data
        if work_data.get('bill_ws_rows'):
            return 50
        return 10

    elif work_type == 'amc':
        items = work_data.get('amc_fetched_items', [])
        if not items:
            return 0
        qty_map = work_data.get('amc_qty_map', {})
        if qty_map:
            return min(80, 20 + len(items) * 5)
        return min(50, len(items) * 5)
    
    return 0


def get_last_step(request, work_type):
    """Get description of last step user was on."""
    if work_type == 'new_estimate':
        items = request.session.get('fetched_items', [])
        if items:
            return f"Selected {len(items)} items"
        return "Category selection"
    
    elif work_type == 'workslip':
        rows = request.session.get('ws_estimate_rows', [])
        if rows:
            return f"Uploaded estimate with {len(rows)} items"
        return "Initial upload"
    
    elif work_type == 'temporary_works':
        items = request.session.get('temp_entries', [])
        if items:
            return f"Added {len(items)} temporary items"
        return "Category selection"
    
    elif work_type == 'amc':
        items = request.session.get('amc_fetched_items', [])
        if items:
            return f"Selected {len(items)} AMC items"
        return "Category selection"

    elif work_type == 'bill':
        bill_num = request.session.get('bill_target_number', 1)
        return f"Bill-{bill_num} generation"

    return "Started"


# ==============================================================================
# RESUME WORK FUNCTIONALITY
# ==============================================================================

@login_required(login_url='login')
def resume_saved_work(request, work_id):
    """Resume a saved work - restores session state and redirects to appropriate module."""
    org = get_org_from_request(request)
    user = request.user
    
    saved_work = get_object_or_404(SavedWork, id=work_id, organization=org, user=user)
    
    # Check subscription access BEFORE allowing resume
    access_result = check_saved_work_access(user, saved_work)
    if not access_result['ok']:
        module_code = access_result.get('module_code')
        messages.warning(
            request, 
            f'You need an active subscription to access this saved work. {access_result["reason"]}'
        )
        # Redirect to module subscription page if module_code exists
        if module_code:
            return redirect('module_access', module_code=module_code)
        return redirect('saved_works_list')
    
    # Restore session state based on work type
    restore_work_data(request, saved_work)
    
    # Store current saved work ID in session
    request.session['current_saved_work_id'] = saved_work.id
    request.session['current_saved_work_name'] = saved_work.name
    
    # Update last accessed
    saved_work.save()  # Updates updated_at timestamp
    
    # Redirect to appropriate module
    redirect_url = get_module_url(saved_work)

    return redirect(redirect_url)


def restore_work_data(request, saved_work):
    """Restore session state from saved work data."""
    import logging
    logger = logging.getLogger(__name__)
    
    work_data = saved_work.work_data
    work_type = saved_work.work_type
    
    logger.info(f"[RESTORE DEBUG] Restoring work_type={work_type}, work_data keys={work_data.keys() if work_data else 'None'}")
    
    if work_type == 'new_estimate':
        request.session['fetched_items'] = work_data.get('fetched_items', [])
        request.session['current_project_name'] = work_data.get('current_project_name')
        request.session['work_type'] = work_data.get('work_type', 'original')
        request.session['qty_map'] = work_data.get('qty_map', {})
        request.session['unit_map'] = work_data.get('unit_map', {})
        request.session['work_name'] = work_data.get('work_name', '')
        request.session['grand_total'] = work_data.get('grand_total', '')
        request.session['excess_tp_percent'] = work_data.get('excess_tp_percent', '')
        request.session['ls_special_name'] = work_data.get('ls_special_name', '')
        request.session['ls_special_amount'] = work_data.get('ls_special_amount', '')
        request.session['deduct_old_material'] = work_data.get('deduct_old_material', '')
        # Restore backend_id so modules use the correct backend
        if work_data.get('selected_backend_id'):
            request.session['selected_backend_id'] = work_data['selected_backend_id']
        # Restore saved item rates & units so workslip generation uses exact values
        if work_data.get('item_rates'):
            request.session['item_rates'] = work_data['item_rates']
        if work_data.get('item_units'):
            request.session['item_units'] = work_data['item_units']
        if work_data.get('item_descs'):
            request.session['item_descs'] = work_data['item_descs']
        # Restore estimate-level metadata (admin sanction, tech sanction, etc.)
        request.session['estimate_metadata'] = work_data.get('estimate_metadata', {})
        # Restore estimate source (uploaded vs datas)
        request.session['estimate_source'] = work_data.get('estimate_source', '')
        # Restore uploaded custom items data
        request.session['uploaded_items'] = work_data.get('uploaded_items', [])
        request.session['uploaded_file_id'] = work_data.get('uploaded_file_id')
        request.session['uploaded_item_blocks'] = work_data.get('uploaded_item_blocks', {})
        request.session['uploaded_sheet_name'] = work_data.get('uploaded_sheet_name', '')
        # Force session save
        request.session.modified = True
    
    elif work_type == 'workslip':
        ws_estimate_rows = work_data.get('ws_estimate_rows', [])
        logger.info(f"[RESTORE DEBUG] ws_estimate_rows count={len(ws_estimate_rows)}")
        if ws_estimate_rows:
            logger.info(f"[RESTORE DEBUG] First row: {ws_estimate_rows[0]}")
        
        request.session['ws_estimate_rows'] = ws_estimate_rows
        request.session['ws_exec_map'] = work_data.get('ws_exec_map', {})
        request.session['ws_tp_percent'] = work_data.get('ws_tp_percent', 0.0)
        request.session['ws_tp_type'] = work_data.get('ws_tp_type', 'Excess')
        request.session['ws_supp_items'] = work_data.get('ws_supp_items', [])
        request.session['ws_estimate_grand_total'] = work_data.get('ws_estimate_grand_total', 0.0)
        request.session['ws_work_name'] = work_data.get('ws_work_name', '')
        request.session['ws_deduct_old_material'] = work_data.get('ws_deduct_old_material', 0.0)
        request.session['ws_current_phase'] = work_data.get('ws_current_phase', 1)
        request.session['ws_target_workslip'] = work_data.get('ws_target_workslip', 1)
        request.session['ws_previous_phases'] = work_data.get('ws_previous_phases', [])
        request.session['ws_previous_ae_data'] = work_data.get('ws_previous_ae_data', [])
        request.session['ws_previous_supp_items'] = work_data.get('ws_previous_supp_items', [])
        request.session['ws_metadata'] = work_data.get('ws_metadata', {})
        
        # Restore backend_id so the workslip module uses the correct backend
        if work_data.get('selected_backend_id'):
            request.session['ws_selected_backend_id'] = work_data['selected_backend_id']
        request.session['ws_source_estimate_id'] = work_data.get('ws_source_estimate_id')
        # Restore work_mode (original/repair) — fall back to parent estimate's work_type
        ws_work_mode = work_data.get('ws_work_mode')
        if not ws_work_mode and saved_work.parent and saved_work.parent.work_data:
            ws_work_mode = saved_work.parent.work_data.get('work_type', 'original')
        request.session['ws_work_mode'] = ws_work_mode or 'original'
        request.session['ws_lc_percent'] = work_data.get('ws_lc_percent', 0.0)
        request.session['ws_qc_percent'] = work_data.get('ws_qc_percent', 0.0)
        request.session['ws_nac_percent'] = work_data.get('ws_nac_percent', 0.0)
        
        # Force session save
        request.session.modified = True
        logger.info(f"[RESTORE DEBUG] Session saved. ws_estimate_rows in session: {len(request.session.get('ws_estimate_rows', []))}")
    
    elif work_type == 'temporary_works':
        request.session['temp_entries'] = work_data.get('temp_entries', [])
        request.session['temp_work_name'] = work_data.get('temp_work_name', '')
        request.session['temp_grand_total'] = work_data.get('temp_grand_total', '')
        request.session['temp_selected_backend_id'] = work_data.get('temp_selected_backend_id')
        request.session['temp_category'] = work_data.get('temp_category', 'electrical')
        request.session.modified = True
    
    elif work_type == 'amc':
        request.session['amc_fetched_items'] = work_data.get('amc_fetched_items', [])
        request.session['amc_qty_map'] = work_data.get('amc_qty_map', {})
        request.session['amc_category'] = work_data.get('amc_category', 'electrical')
        request.session['amc_work_name'] = work_data.get('amc_work_name', '')
        request.session['amc_grand_total'] = work_data.get('amc_grand_total', '')
        request.session['amc_selected_backend_id'] = work_data.get('amc_selected_backend_id')
        request.session['amc_work_type'] = work_data.get('amc_work_type', 'original')
        request.session.modified = True

    elif work_type == 'bill':
        # Restore bill session data for the existing bill engine
        request.session['bill_source_work_id'] = work_data.get('bill_source_work_id')
        request.session['bill_source_work_type'] = work_data.get('bill_source_work_type', '')
        request.session['bill_source_work_name'] = work_data.get('bill_source_work_name', '')
        request.session['bill_from_workslip'] = work_data.get('bill_from_workslip', False)
        request.session['bill_ws_rows'] = work_data.get('bill_ws_rows', [])
        request.session['bill_ws_exec_map'] = work_data.get('bill_ws_exec_map', {})
        request.session['bill_ws_tp_percent'] = work_data.get('bill_ws_tp_percent', 0)
        request.session['bill_ws_tp_type'] = work_data.get('bill_ws_tp_type', 'Excess')
        request.session['bill_target_number'] = work_data.get('bill_target_number', 1)
        request.session['bill_type'] = work_data.get('bill_type', '')
        request.session['bill_ws_metadata'] = work_data.get('bill_ws_metadata', {})
        request.session.modified = True


def get_module_url(saved_work):
    """Get the URL to redirect to for resuming work."""
    from django.urls import reverse
    
    work_type = saved_work.work_type
    category = saved_work.category or 'electrical'
    work_data = saved_work.work_data or {}
    
    if work_type == 'new_estimate':
        # All new_estimate saved works go to the new module
        # (old uploaded-estimate works are also compatible with the new module)
        last_group = work_data.get('last_group', '')
        if last_group:
            return reverse('datas_items', kwargs={'category': category, 'group': last_group})
        return reverse('datas_groups', kwargs={'category': category})
    
    elif work_type == 'workslip':
        return reverse('workslip_main') + '?preserve=1'
    
    elif work_type == 'temporary_works':
        last_group = work_data.get('last_group', '')
        if last_group:
            return reverse('temp_items', kwargs={'category': category, 'group': last_group})
        return reverse('temp_groups', kwargs={'category': category})
    
    elif work_type == 'amc':
        last_group = work_data.get('last_group', '')
        if last_group:
            return reverse('amc_items', kwargs={'category': category, 'group': last_group})
        return reverse('amc_groups', kwargs={'category': category})

    elif work_type == 'bill':
        # Redirect to the bill_entry page which has all the functionality:
        # MB details, dates, items with quantities, save, and download.
        # Find the source workslip to use as the bill_entry context.
        source_workslip_id = None
        # Check parent chain: bill -> workslip
        if saved_work.parent_id:
            parent = saved_work.parent
            if parent and parent.work_type == 'workslip':
                source_workslip_id = parent.id
            elif parent and parent.work_type == 'new_estimate':
                # Bill was generated directly from estimate - find the workslip child
                # that matches the bill number
                ws = parent.children.filter(
                    work_type='workslip',
                    workslip_number=saved_work.bill_number or 1,
                ).first()
                if ws:
                    source_workslip_id = ws.id
        # Also check work_data for source workslip id
        if not source_workslip_id:
            source_workslip_id = work_data.get('bill_parent_work_id') or work_data.get('source_workslip_id') or work_data.get('bill_source_work_id')
        
        if source_workslip_id:
            bill_num = saved_work.bill_number or 1
            url = reverse('bill_entry', kwargs={'work_id': int(source_workslip_id)})
            # Pass bill_number so the page shows the correct bill, not the workslip's number
            return f"{url}?bill_number={bill_num}"
        # Fallback to old bill module page if no source workslip found
        return reverse('bill') + '?from_saved=1'

    return reverse('dashboard')


# ==============================================================================
# SAVED WORK DETAILS & ACTIONS
# ==============================================================================

@login_required(login_url='login')
@require_POST
def update_saved_work(request, work_id):
    """Update saved work metadata (name, folder, notes)."""
    org = get_org_from_request(request)
    user = request.user
    
    saved_work = get_object_or_404(SavedWork, id=work_id, organization=org, user=user)
    
    work_name = request.POST.get('name', '').strip()
    folder_id = request.POST.get('folder_id')
    notes = request.POST.get('notes', '').strip()
    status = request.POST.get('status')
    
    if work_name:
        saved_work.name = work_name
    
    if folder_id:
        if folder_id == 'none':
            saved_work.folder = None
        else:
            saved_work.folder = get_object_or_404(WorkFolder, id=folder_id, organization=org, user=user)
    
    saved_work.notes = notes
    
    if status in dict(SavedWork.STATUS_CHOICES):
        saved_work.status = status
    
    saved_work.save()
    
    return JsonResponse({
        'success': True,
        'message': f'Work "{saved_work.name}" updated successfully!'
    })


@login_required(login_url='login')
@require_POST
def delete_saved_work(request, work_id):
    """Delete a saved work."""
    org = get_org_from_request(request)
    user = request.user
    
    saved_work = get_object_or_404(SavedWork, id=work_id, organization=org, user=user)
    work_name = saved_work.name
    saved_work.delete()
    
    return JsonResponse({
        'success': True,
        'message': f'Work "{work_name}" deleted successfully!'
    })


@login_required(login_url='login')
@require_POST
def delete_children(request, work_id):
    """Delete all child works (workslips + bills) of a saved work."""
    org = get_org_from_request(request)
    user = request.user
    saved_work = get_object_or_404(SavedWork, id=work_id, organization=org, user=user)

    # Collect all descendants (BFS)
    to_delete = []
    queue = list(saved_work.children.all())
    while queue:
        child = queue.pop(0)
        to_delete.append(child)
        queue.extend(list(child.children.all()))

    count = len(to_delete)
    for child in to_delete:
        child.delete()

    return JsonResponse({
        'success': True,
        'message': f'Deleted {count} child work(s).'
    })


@login_required(login_url='login')
@require_POST
def delete_subsequent_bills(request, work_id):
    """Delete all bills with a higher bill_number in the same workflow."""
    org = get_org_from_request(request)
    user = request.user
    bill = get_object_or_404(SavedWork, id=work_id, organization=org, user=user, work_type='bill')

    # Find root estimate to locate all bills in the workflow
    root = bill.get_root_estimate() if hasattr(bill, 'get_root_estimate') else None
    if root:
        all_bills = root.get_all_bills() if hasattr(root, 'get_all_bills') else []
    else:
        # Fallback: find sibling bills under same parent
        all_bills = list(SavedWork.objects.filter(
            parent=bill.parent, work_type='bill', organization=org, user=user
        ))

    to_delete = [b for b in all_bills if b.bill_number > bill.bill_number]
    count = len(to_delete)
    for b in to_delete:
        b.delete()

    return JsonResponse({
        'success': True,
        'message': f'Deleted {count} subsequent bill(s).'
    })


@login_required(login_url='login')
@require_POST
def move_to_folder(request, work_id):
    """Move a saved work to a different folder."""
    org = get_org_from_request(request)
    user = request.user
    
    saved_work = get_object_or_404(SavedWork, id=work_id, organization=org, user=user)
    folder_id = request.POST.get('folder_id')
    
    if folder_id == 'none' or not folder_id:
        saved_work.folder = None
    else:
        saved_work.folder = get_object_or_404(WorkFolder, id=folder_id, organization=org, user=user)
    
    saved_work.save()
    
    return JsonResponse({
        'success': True,
        'message': f'Work moved successfully!'
    })


@login_required(login_url='login')
@require_POST
def duplicate_saved_work(request, work_id):
    """Duplicate a saved work."""
    org = get_org_from_request(request)
    user = request.user
    
    saved_work = get_object_or_404(SavedWork, id=work_id, organization=org, user=user)
    
    new_work = SavedWork.objects.create(
        organization=org,
        user=user,
        folder=saved_work.folder,
        name=f"{saved_work.name} (Copy)",
        work_type=saved_work.work_type,
        work_data=saved_work.work_data.copy() if saved_work.work_data else {},
        category=saved_work.category,
        notes=saved_work.notes,
        progress_percent=saved_work.progress_percent,
        last_step=saved_work.last_step,
    )
    
    return JsonResponse({
        'success': True,
        'work_id': new_work.id,
        'message': f'Work duplicated as "{new_work.name}"!'
    })


@login_required(login_url='login')
@require_POST
def copy_to_folder(request, work_id):
    """Copy a saved work to a specific folder."""
    org = get_org_from_request(request)
    user = request.user

    saved_work = get_object_or_404(SavedWork, id=work_id, organization=org, user=user)
    folder_id = request.POST.get('folder_id')

    target_folder = None
    if folder_id and folder_id != 'none':
        target_folder = get_object_or_404(WorkFolder, id=folder_id, organization=org, user=user)

    new_work = SavedWork.objects.create(
        organization=org,
        user=user,
        folder=target_folder,
        name=f"{saved_work.name} (Copy)",
        work_type=saved_work.work_type,
        work_data=saved_work.work_data.copy() if saved_work.work_data else {},
        category=saved_work.category,
        notes=saved_work.notes,
        progress_percent=saved_work.progress_percent,
        last_step=saved_work.last_step,
    )

    return JsonResponse({
        'success': True,
        'work_id': new_work.id,
        'message': f'Work copied as "{new_work.name}"!'
    })


@login_required(login_url='login')
@require_POST
def batch_action(request):
    """Batch move, copy, or delete multiple works."""
    org = get_org_from_request(request)
    user = request.user

    try:
        data = json.loads(request.body.decode('utf-8'))
    except (json.JSONDecodeError, UnicodeDecodeError):
        return JsonResponse({'success': False, 'error': 'Invalid request data.'})

    action = data.get('action')
    work_ids = data.get('work_ids', [])
    folder_id = data.get('folder_id')

    if not work_ids or not action:
        return JsonResponse({'success': False, 'error': 'No works or action specified.'})

    works = SavedWork.objects.filter(id__in=work_ids, organization=org, user=user)
    count = works.count()

    if count == 0:
        return JsonResponse({'success': False, 'error': 'No matching works found.'})

    target_folder = None
    if folder_id and folder_id != 'none':
        target_folder = get_object_or_404(WorkFolder, id=folder_id, organization=org, user=user)

    if action == 'move':
        works.update(folder=target_folder)
        return JsonResponse({'success': True, 'message': f'Moved {count} work(s) successfully!'})

    elif action == 'copy':
        for work in works:
            SavedWork.objects.create(
                organization=org,
                user=user,
                folder=target_folder,
                name=f"{work.name} (Copy)",
                work_type=work.work_type,
                work_data=work.work_data.copy() if work.work_data else {},
                category=work.category,
                notes=work.notes,
                progress_percent=work.progress_percent,
                last_step=work.last_step,
            )
        return JsonResponse({'success': True, 'message': f'Copied {count} work(s) successfully!'})

    elif action == 'delete':
        works.delete()
        return JsonResponse({'success': True, 'message': f'Deleted {count} work(s) successfully!'})

    return JsonResponse({'success': False, 'error': 'Invalid action.'})


# ==============================================================================
# SAVE MODAL DATA (for frontend save work modal)
# ==============================================================================

@login_required(login_url='login')
def check_work_name(request):
    """Check if a work with the same name already exists in the given folder."""
    org = get_org_from_request(request)
    user = request.user
    name = request.GET.get('name', '').strip()
    folder_id = request.GET.get('folder_id', '') or None
    exclude_id = request.GET.get('exclude_id', '') or None

    if not name:
        return JsonResponse({'exists': False})

    qs = SavedWork.objects.filter(organization=org, user=user, name__iexact=name)
    if folder_id:
        qs = qs.filter(folder_id=folder_id)
    else:
        qs = qs.filter(folder__isnull=True)
    if exclude_id:
        qs = qs.exclude(id=exclude_id)

    return JsonResponse({'exists': qs.exists()})


@login_required(login_url='login')
def get_save_work_modal_data(request):
    """Get data needed for save work modal (folders list, current work info)."""
    org = get_org_from_request(request)
    user = request.user
    
    folders = WorkFolder.objects.filter(organization=org, user=user).values('id', 'name', 'color')
    
    current_work_id = request.session.get('current_saved_work_id')
    current_work = None
    if current_work_id:
        try:
            work = SavedWork.objects.get(id=current_work_id, organization=org, user=user)
            current_work = {
                'id': work.id,
                'name': work.name,
                'folder_id': work.folder_id,
                'notes': work.notes,
            }
        except SavedWork.DoesNotExist:
            pass
    
    return JsonResponse({
        'folders': list(folders),
        'current_work': current_work,
    })


# ==============================================================================
# WORKFLOW CHAIN: ESTIMATE â†’ WORKSLIP â†’ BILL
# ==============================================================================

def load_item_rates_from_backend(category, item_names, backend_id=None, user=None, module_code=None):
    """
    Load item rates and descriptions from the backend Excel file for given item names.
    Uses the user's selected backend (ModuleBackend) if available, falling back to default.
    Returns a dict: {item_name: {'rate': value, 'unit': 'Nos/Mtrs/Pts', 'group': 'group_name', 'desc': 'description'}}
    """
    from django.conf import settings
    from openpyxl import load_workbook
    import os
    import logging
    logger = logging.getLogger(__name__)
    
    logger.info(f"[LOAD_RATES DEBUG] Loading rates for category={category}, items={item_names}, backend_id={backend_id}")
    
    try:
        from core.utils_excel import load_backend
        
        # Use load_backend which handles backend_id, user preferences, and fallbacks
        items_list, groups_map, _units_map, ws_data, filepath = load_backend(
            category, settings.BASE_DIR,
            backend_id=backend_id,
            module_code=module_code or 'new_estimate',
            user=user
        )
        
        logger.info(f"[LOAD_RATES DEBUG] Backend filepath: {filepath}, found {len(items_list)} items")
        
        if not filepath or not os.path.exists(filepath):
            logger.warning(f"[LOAD_RATES DEBUG] Backend file not found at {filepath}!")
            return {name: {'rate': 0, 'unit': 'Nos', 'group': '', 'desc': name} for name in item_names}
        
        # Open workbook with data_only=True to get calculated formula values
        wb_data = load_workbook(filepath, data_only=True)
        ws_data_only = wb_data["Master Datas"]
        # Build item to group mapping
        item_to_group = {}
        for grp_name, item_list_in_grp in groups_map.items():
            for nm in item_list_in_grp:
                item_to_group[nm] = grp_name
        
        # Get rates and descriptions for requested items
        result = {}
        for info in items_list:
            name = info["name"]
            if name not in item_names:
                continue
            
            start_row = info["start_row"]
            end_row = info["end_row"]
            rate = 0
            
            # Find rate from bottom up (last value in column J)
            # Try data_only workbook first (has calculated values), then formula workbook
            for r in range(end_row, start_row - 1, -1):
                val = ws_data_only.cell(row=r, column=10).value  # column J (data_only)
                if val not in (None, ""):
                    try:
                        rate = float(val)
                    except (ValueError, TypeError):
                        rate = 0
                    break
            
            # If rate is still 0, try the formula workbook (ws_data)
            # Some Excel files may only have literal values, not formula results
            if rate == 0 and ws_data is not None:
                for r in range(end_row, start_row - 1, -1):
                    val = ws_data.cell(row=r, column=10).value
                    if val not in (None, ""):
                        try:
                            rate = float(val)
                        except (ValueError, TypeError):
                            rate = 0
                        break
            
            # Get description from row start_row + 2, column D (4)
            desc = name  # default to item name
            if ws_data is not None:
                desc_cell = ws_data.cell(row=start_row + 2, column=4).value
                if desc_cell and str(desc_cell).strip():
                    desc = str(desc_cell).strip()
            
            # Determine unit from units_map (Column D of Groups sheet) — authoritative source
            unit = _units_map.get(name, "")
            if not unit:
                # Fallback: infer from group name
                grp_name = item_to_group.get(name, "")
                if grp_name in ("Piping", "Wiring & Cables", "Wiring and cables"):
                    unit = "Mtrs"
                elif grp_name == "Points":
                    unit = "Pts"
                else:
                    unit = "Nos"
            else:
                grp_name = item_to_group.get(name, "")
            
            result[name] = {'rate': rate, 'unit': unit, 'group': grp_name, 'desc': desc}
            logger.info(f"[LOAD_RATES DEBUG] Found rate for '{name}': rate={rate}, unit={unit}, desc={desc[:50] if desc else 'None'}")
        
        wb_data.close()
        
        # Fill in any missing items with defaults
        for name in item_names:
            if name not in result:
                logger.warning(f"[LOAD_RATES DEBUG] Item '{name}' NOT FOUND in backend!")
                result[name] = {'rate': 0, 'unit': 'Nos', 'group': '', 'desc': name}
        
        return result
        
    except Exception as e:
        logger.error(f"[LOAD_RATES DEBUG] Error loading item rates: {e}")
        import traceback
        traceback.print_exc()
        return {name: {'rate': 0, 'unit': 'Nos', 'group': '', 'desc': name} for name in item_names}


def load_prefix_map(category, backend_id=None, user=None, module_code='new_estimate'):
    """
    Load {item_name: prefix} mapping from the backend Excel's Groups sheet.
    Returns empty dict on error or if no prefixes found.
    """
    import os
    import logging
    from django.conf import settings
    logger = logging.getLogger(__name__)

    try:
        from openpyxl import load_workbook as _load_wb
        from core.utils_excel import load_backend

        _items_list, _groups_map, _units_map, _ws_data, filepath = load_backend(
            category, settings.BASE_DIR,
            backend_id=backend_id,
            user=user,
            module_code=module_code,
        )
        if not filepath or not os.path.exists(filepath):
            return {}

        backend_wb = _load_wb(filepath, data_only=False)
        ws_groups = backend_wb["Groups"]
        header_row_g = None
        col_item_g = None
        col_prefix_g = None
        for r in range(1, ws_groups.max_row + 1):
            for c in range(1, ws_groups.max_column + 1):
                val = str(ws_groups.cell(row=r, column=c).value or "").strip().lower()
                if val == "item name":
                    header_row_g = r
                    col_item_g = c
                elif val == "prefix":
                    col_prefix_g = c
            if header_row_g:
                break
        if not (header_row_g and col_item_g and col_prefix_g):
            return {}
        prefix_map = {}
        for r in range(header_row_g + 1, ws_groups.max_row + 1):
            nm = ws_groups.cell(r, col_item_g).value
            px = ws_groups.cell(r, col_prefix_g).value
            if nm and px not in (None, ""):
                prefix_map[str(nm).strip()] = str(px).strip()
        return prefix_map
    except Exception:
        logger.debug("[LOAD_PREFIX_MAP] Error loading prefix map", exc_info=True)
        return {}


def apply_prefix_to_desc(desc, item_name, prefix_map):
    """Apply prefix from prefix_map to a description string. Returns modified desc."""
    if not prefix_map:
        return desc
    prefix = prefix_map.get(item_name, '')
    if prefix:
        return f"{prefix} {desc}" if desc else prefix
    return desc


@login_required(login_url='login')
def generate_workslip_from_saved(request, work_id):
    """
    Generate a workslip from a saved estimate.
    Loads the estimate data into workslip session and redirects to workslip page.
    """
    import logging
    logger = logging.getLogger(__name__)
    
    org = get_org_from_request(request)
    user = request.user
    
    saved_work = get_object_or_404(SavedWork, id=work_id, organization=org, user=user)
    
    # Check workslip subscription access BEFORE allowing generation
    try:
        from subscriptions.services import SubscriptionService
        result = SubscriptionService.check_access(user, 'workslip')
        if not result.get('ok', False):
            messages.warning(request, 'You need an active Workslip subscription to generate workslips.')
            return redirect('module_access', module_code='workslip')
    except Exception:
        messages.error(request, 'Unable to verify subscription. Please try again.')
        return redirect('saved_works_list')
    
    # Verify this is an estimate, temporary_works, or amc
    if saved_work.work_type not in ['new_estimate', 'temporary_works', 'amc']:
        messages.error(request, 'Only estimates, temporary works, or AMC can be used to generate workslips.')
        return redirect('saved_works_list')

    # Check if estimate is finalized (completed) before allowing workslip generation
    # Auto-finalize if user confirmed via the list page popup
    if saved_work.status != 'completed':
        saved_work.status = 'completed'
        saved_work.save(update_fields=['status'])
    
    # Load estimate data into workslip session
    work_data = saved_work.work_data or {}
    
    # Get item names and quantities from saved data
    # fetched_items can be either a list of names (strings) or a list of dicts
    fetched_items = work_data.get('fetched_items', [])
    qty_map = work_data.get('qty_map', {})
    category = saved_work.category or 'electrical'
    
    # Get saved backend_id from estimate work_data for correct rate lookup
    saved_backend_id = work_data.get('selected_backend_id')
    
    logger.info(f"[GEN_WORKSLIP DEBUG] fetched_items={fetched_items}")
    logger.info(f"[GEN_WORKSLIP DEBUG] qty_map={qty_map}")
    logger.info(f"[GEN_WORKSLIP DEBUG] category={category}, backend_id={saved_backend_id}")
    
    # Determine module_code for backend loading
    if saved_work.work_type == 'amc':
        ws_module_code = 'amc'
    elif saved_work.work_type == 'temporary_works':
        ws_module_code = 'temp_works'
    else:
        ws_module_code = 'new_estimate'
    
    # Saved rates / units / descriptions from the estimate session (exact values the user saw)
    saved_item_rates = work_data.get('item_rates', {})
    saved_item_units = work_data.get('item_units', {})
    saved_item_descs = work_data.get('item_descs', {})
    
    logger.info(f"[GEN_WORKSLIP DEBUG] saved_item_rates keys={list(saved_item_rates.keys())[:10]}")
    
    # Check if fetched_items is a list of strings (item names) or dicts
    if fetched_items and isinstance(fetched_items[0], str):
        item_names = fetched_items
        
        # Fetch descriptions from backend for ALL items (desc always comes from backend row+2)
        item_info_map = {}
        if item_names:
            logger.info(f"[GEN_WORKSLIP DEBUG] Fetching descriptions from backend for {len(item_names)} items")
            item_info_map = load_item_rates_from_backend(
                category, item_names,
                backend_id=saved_backend_id,
                user=request.user,
                module_code=ws_module_code
            )
        
        # Convert to workslip format - match the format from estimate upload
        ws_estimate_rows = []
        for idx, item_name in enumerate(item_names):
            qty = qty_map.get(item_name, 0)
            try:
                qty = float(qty) if qty else 0.0
            except (ValueError, TypeError):
                qty = 0.0
            
            # Get backend info (always has description from row+2)
            info = item_info_map.get(item_name, {'rate': 0, 'unit': 'Nos', 'desc': item_name})
            backend_desc = str(info.get('desc', item_name) or item_name)
            # If backend returned header name as desc, prefer saved item_descs
            if backend_desc == item_name and item_name in saved_item_descs:
                backend_desc = saved_item_descs[item_name]
            
            # Priority for rate: 1) saved rates from estimate, 2) backend re-fetch
            if item_name in saved_item_rates and saved_item_rates[item_name]:
                rate = float(saved_item_rates[item_name])
                unit = str(saved_item_units.get(item_name, info.get('unit', 'Nos')))
            else:
                rate = float(info.get('rate', 0) or 0)
                unit = str(info.get('unit', 'Nos'))
            
            logger.info(f"[GEN_WORKSLIP DEBUG] Item '{item_name}': qty={qty}, rate={rate}, desc={backend_desc[:50] if backend_desc else 'N/A'}")
            
            ws_estimate_rows.append({
                'key': f"saved_{idx}",
                'item_name': str(item_name),       # Item header name for UI
                'display_name': str(item_name),    # Item header name for UI
                'item_desc': backend_desc,         # Backend row+2 description for downloads
                'desc': backend_desc,              # Backend row+2 description for downloads
                'unit': unit,
                'qty_est': qty,
                'rate': rate,
            })
    else:
        # It's already a list of dicts with full item info (from estimate upload)
        ws_estimate_rows = []
        for idx, item in enumerate(fetched_items):
            if isinstance(item, dict):
                item_id = item.get('id') or item.get('item_name') or item.get('name') or str(idx)
                qty = qty_map.get(str(item_id), item.get('qty', item.get('qty_est', 0)))
                try:
                    qty = float(qty) if qty else 0.0
                except (ValueError, TypeError):
                    qty = 0.0
                rate = float(item.get('rate', 0)) if item.get('rate') else 0.0
                
                # UI display: use display_name (yellow header)
                ui_name = str(item.get('display_name') or item.get('item_name') or item.get('name', ''))
                # Download description: use item_desc (row+2 content) preferentially
                download_desc = str(item.get('item_desc') or item.get('desc') or item.get('description') or '')
                # If no desc stored in item dict, check saved_item_descs map
                if (not download_desc or download_desc == ui_name) and ui_name in saved_item_descs:
                    download_desc = saved_item_descs[ui_name]
                if not download_desc:
                    download_desc = ui_name
                
                ws_estimate_rows.append({
                    'key': f"saved_{idx}",
                    'item_name': ui_name,            # Item header name for UI
                    'display_name': ui_name,         # Item header name for UI
                    'item_desc': download_desc,      # Backend row+2 description for downloads
                    'desc': download_desc,           # Backend row+2 description for downloads
                    'unit': str(item.get('unit', 'Nos')),
                    'qty_est': qty,
                    'rate': rate,
                })
    
    # Convert grand_total to float (it might be stored as string)
    grand_total = work_data.get('grand_total', 0)
    logger.info(f"[GEN_WORKSLIP DEBUG] Raw grand_total from work_data: '{grand_total}' (type: {type(grand_total).__name__})")
    try:
        grand_total = float(grand_total) if grand_total else 0.0
    except (ValueError, TypeError):
        grand_total = 0.0
    logger.info(f"[GEN_WORKSLIP DEBUG] Parsed grand_total: {grand_total}")
    
    # Get the work_name from the saved estimate data (entered in estimate preview)
    # This is the "Name of the Work" for Excel, NOT the project/file name (saved_work.name)
    estimate_work_name = work_data.get('work_name', '') or ''
    
    # Set workslip session data - ensure all values are JSON serializable
    request.session['ws_estimate_rows'] = ws_estimate_rows
    request.session['ws_exec_map'] = {}  # Start fresh execution quantities
    request.session['ws_tp_percent'] = 0.0  # TP will be entered via UI for Workslip-1
    request.session['ws_tp_type'] = 'Excess'
    request.session['ws_supp_items'] = []
    request.session['ws_estimate_grand_total'] = grand_total
    request.session['ws_work_name'] = str(estimate_work_name)
    request.session['ws_source_estimate_id'] = int(saved_work.id)
    request.session['ws_current_phase'] = 1
    request.session['ws_target_workslip'] = 1
    request.session['ws_previous_phases'] = []
    request.session['ws_previous_supp_items'] = []
    
    # For Workslip-1: Set initial metadata from estimate
    # work_name and grand_total come from estimate, TP will be entered via UI
    # Inherit header fields from estimate_metadata if they were saved earlier
    est_meta = work_data.get('estimate_metadata', {})
    request.session['ws_metadata'] = {
        'work_name': str(estimate_work_name),
        'estimate_amount': str(grand_total) if grand_total else '',
        'admin_sanction': est_meta.get('admin_sanction', ''),
        'tech_sanction': est_meta.get('tech_sanction', ''),
        'agreement': est_meta.get('agreement', ''),
        'agency_name': est_meta.get('agency_name', ''),
        'tp_percent': 0.0,
        'tp_type': 'Excess',
        'grand_total': grand_total,
    }
    
    # Set the backend_id in session so the workslip module uses the correct backend
    if saved_backend_id:
        request.session['ws_selected_backend_id'] = saved_backend_id
    
    # Set work type/mode/category for workslip module
    if saved_work.work_type == 'amc':
        request.session['ws_work_type'] = 'amc'
    elif saved_work.work_type == 'temporary_works':
        request.session['ws_work_type'] = 'tempworks'
    else:
        request.session['ws_work_type'] = 'new_estimate'
    request.session['ws_category'] = category

    # Propagate original/repair work mode from estimate to workslip session
    # This ensures prefixes from the backend Groups sheet are applied in Excel output
    estimate_work_mode = work_data.get('work_type', 'original')
    request.session['ws_work_mode'] = estimate_work_mode

    # ── Pre-create the SavedWork record for Workslip-1 so
    #    quickSaveWorkslip() finds it and auto-updates without asking
    #    for a name. Use get_or_create to prevent duplicates. ──
    new_ws_name = f"{saved_work.name} - W1"

    new_ws, ws_created = SavedWork.objects.get_or_create(
        organization=org,
        user=user,
        parent=saved_work,
        work_type='workslip',
        workslip_number=1,
        defaults={
            'folder': saved_work.folder,
            'name': new_ws_name,
            'work_data': {},  # will be filled on first quickSave
            'category': category,
            'notes': '',
            'progress_percent': 0,
            'last_step': 'workslip',
        },
    )

    request.session['current_saved_work_id'] = new_ws.id
    request.session['current_saved_work_name'] = new_ws_name
    request.session.modified = True

    return redirect(reverse('workslip_main') + '?preserve=1')


@login_required(login_url='login')
def generate_next_workslip_from_saved(request, work_id):
    """
    Generate the next workslip from a saved workslip.
    For example, if work_id is Workslip-1, this generates Workslip-2.
    Loads previous workslip data as phase data and redirects to workslip page.
    """
    import logging
    logger = logging.getLogger(__name__)
    
    org = get_org_from_request(request)
    user = request.user
    
    saved_work = get_object_or_404(SavedWork, id=work_id, organization=org, user=user)
    
    # Check workslip subscription access BEFORE allowing generation
    try:
        from subscriptions.services import SubscriptionService
        result = SubscriptionService.check_access(user, 'workslip')
        if not result.get('ok', False):
            messages.warning(request, 'You need an active Workslip subscription to generate workslips.')
            return redirect('module_access', module_code='workslip')
    except Exception:
        messages.error(request, 'Unable to verify subscription. Please try again.')
        return redirect('saved_works_list')
    
    # Verify this is a workslip
    if saved_work.work_type != 'workslip':
        messages.error(request, 'Only saved workslips can generate next workslips.')
        return redirect('saved_works_list')

    # Check if root estimate is finalized before allowing next workslip generation
    root_est = saved_work.parent
    while root_est and root_est.work_type != 'new_estimate':
        root_est = root_est.parent
    if root_est and root_est.status != 'completed':
        messages.warning(request, 'Please finalize the estimate before generating workslips.')
        return redirect('saved_work_detail', work_id=root_est.id)

    # Load workslip data
    work_data = saved_work.work_data or {}
    
    # Get the current workslip number and calculate next
    current_workslip_number = saved_work.workslip_number or 1
    next_workslip_number = current_workslip_number + 1
    
    logger.info(f"[NEXT_WORKSLIP] Generating Workslip-{next_workslip_number} from saved Workslip-{current_workslip_number} (ID: {work_id})")
    
    # Get estimate rows (base items) from saved workslip
    ws_estimate_rows = work_data.get('ws_estimate_rows', [])
    
    # Get previous phases from saved data (might already have phases from earlier workslips)
    existing_previous_phases = work_data.get('ws_previous_phases', [])
    
    # Get the execution map from this workslip - this becomes a new phase
    current_exec_map = work_data.get('ws_exec_map', {})
    
    # Get supplemental items from this workslip
    current_supp_items = work_data.get('ws_supp_items', [])
    
    # Get previous supplemental items
    existing_prev_supp_items = work_data.get('ws_previous_supp_items', [])
    
    # Build the new phase from current workslip's execution data
    new_phase_map = {}
    for key, qty in current_exec_map.items():
        if qty and str(qty).strip():
            try:
                new_phase_map[key] = float(qty)
            except (ValueError, TypeError):
                pass
    
    # Combine all previous phases plus the current workslip as a new phase
    all_previous_phases = list(existing_previous_phases) + [new_phase_map]
    
    # Build supplemental items from current workslip to add to previous supp items
    new_supp_items = []
    if current_supp_items and ws_estimate_rows:
        # Load rates from backend for supplemental items
        category = saved_work.category or 'electrical'
        saved_backend_id = work_data.get('selected_backend_id')
        supp_rates = load_item_rates_from_backend(
            category, current_supp_items,
            backend_id=saved_backend_id,
            user=request.user,
            module_code='new_estimate'
        )
        
        for supp_name in current_supp_items:
            supp_key = f"supp:{supp_name}"
            supp_qty = current_exec_map.get(supp_key, 0)
            try:
                supp_qty = float(supp_qty) if supp_qty else 0.0
            except:
                supp_qty = 0.0
            if supp_qty > 0:
                # Get rate from backend lookup
                supp_info = supp_rates.get(supp_name, {})
                supp_rate = float(supp_info.get('rate', 0) or 0)
                supp_unit = str(supp_info.get('unit', 'Nos') or 'Nos')
                supp_desc = str(supp_info.get('desc', supp_name) or supp_name)
                
                new_supp_items.append({
                    "name": supp_name,
                    "qty": supp_qty,
                    "phase": current_workslip_number,
                    "supp_section": current_workslip_number,
                    "desc": supp_desc,
                    "unit": supp_unit,
                    "rate": supp_rate,
                    "amount": supp_qty * supp_rate,
                })
    
    # Combine previous supplemental items
    all_previous_supp_items = list(existing_prev_supp_items) + new_supp_items
    
    # Set workslip session data for the next workslip
    request.session['ws_estimate_rows'] = ws_estimate_rows
    request.session['ws_exec_map'] = {}  # Start fresh execution quantities
    request.session['ws_previous_phases'] = all_previous_phases
    request.session['ws_previous_supp_items'] = all_previous_supp_items
    request.session['ws_current_phase'] = next_workslip_number
    request.session['ws_target_workslip'] = next_workslip_number
    request.session['ws_tp_percent'] = work_data.get('ws_tp_percent', 0.0)
    request.session['ws_tp_type'] = work_data.get('ws_tp_type', 'Excess')
    request.session['ws_supp_items'] = []  # Start fresh supplemental items for new workslip
    request.session['ws_estimate_grand_total'] = work_data.get('ws_estimate_grand_total', 0)
    request.session['ws_work_name'] = work_data.get('ws_work_name', '')
    request.session['ws_deduct_old_material'] = work_data.get('ws_deduct_old_material', 0)
    request.session['ws_lc_percent'] = work_data.get('ws_lc_percent', 0)
    request.session['ws_qc_percent'] = work_data.get('ws_qc_percent', 0)
    request.session['ws_nac_percent'] = work_data.get('ws_nac_percent', 0)

    # Propagate work_mode (original/repair) from previous workslip
    # Fall back to parent estimate's work_type if ws_work_mode is not in work_data
    ws_work_mode = work_data.get('ws_work_mode')
    if not ws_work_mode:
        # Walk up to root estimate to get work_type
        root = saved_work.parent
        while root:
            if root.work_type == 'new_estimate' and root.work_data:
                ws_work_mode = root.work_data.get('work_type', 'original')
                break
            root = root.parent
    request.session['ws_work_mode'] = ws_work_mode or 'original'

    # Carry over metadata from previous workslip (Name of work, Agency, Sanctions, Agreement, etc.)
    prev_metadata = work_data.get('ws_metadata', {})
    # Ensure all metadata fields are preserved
    metadata = {
        'work_name': prev_metadata.get('work_name', '') or work_data.get('ws_work_name', ''),
        'estimate_amount': prev_metadata.get('estimate_amount', ''),
        'admin_sanction': prev_metadata.get('admin_sanction', ''),
        'tech_sanction': prev_metadata.get('tech_sanction', ''),
        'agreement': prev_metadata.get('agreement', ''),
        'agency_name': prev_metadata.get('agency_name', ''),
        'tp_percent': work_data.get('ws_tp_percent', 0.0),
        'tp_type': work_data.get('ws_tp_type', 'Excess'),
        'grand_total': work_data.get('ws_estimate_grand_total', 0),
        'deduct_old_material': work_data.get('ws_deduct_old_material', 0),
        'lc_percent': work_data.get('ws_lc_percent', 0),
        'qc_percent': work_data.get('ws_qc_percent', 0),
        'nac_percent': work_data.get('ws_nac_percent', 0),
    }
    request.session['ws_metadata'] = metadata
    
    # Set the backend_id in session so the workslip module uses the correct backend
    saved_backend_id_session = work_data.get('selected_backend_id')
    if saved_backend_id_session:
        request.session['ws_selected_backend_id'] = saved_backend_id_session
    
    # Set parent work info for saving
    # Carry forward the source estimate ID so new workslip links to root estimate
    source_est_id = work_data.get('ws_source_estimate_id')
    if source_est_id:
        request.session['ws_source_estimate_id'] = source_est_id
    elif saved_work.parent_id:
        request.session['ws_source_estimate_id'] = saved_work.parent_id
    request.session['ws_parent_work_id'] = saved_work.id

    # ── Pre-create the SavedWork record for the new workslip so
    #    quickSaveWorkslip() finds it and auto-updates without asking
    #    for a name.  Derive name from parent workslip / root estimate. ─
    # Walk up to root estimate — ALL workslips should be direct children
    # of the root estimate so queries like parent=root_estimate find them.
    root_estimate = saved_work.parent
    while root_estimate:
        if root_estimate.work_type == 'new_estimate':
            break
        root_estimate = root_estimate.parent
    # Fallback: if no root estimate found, parent to the current workslip
    parent_for_new_ws = root_estimate if root_estimate else saved_work

    base_name = root_estimate.name if root_estimate else saved_work.name

    new_ws_name = f"{base_name} - W{next_workslip_number}"

    # Use get_or_create to prevent duplicate workslip records on repeat visits
    new_ws, ws_created = SavedWork.objects.get_or_create(
        organization=org,
        user=user,
        parent=parent_for_new_ws,
        work_type='workslip',
        workslip_number=next_workslip_number,
        defaults={
            'folder': saved_work.folder,
            'name': new_ws_name,
            'work_data': {},  # will be filled on first quickSave
            'category': saved_work.category or 'electrical',
            'notes': '',
            'progress_percent': 0,
            'last_step': 'workslip',
        },
    )

    request.session['current_saved_work_id'] = new_ws.id
    request.session['current_saved_work_name'] = new_ws_name
    request.session.modified = True

    return redirect(reverse('workslip_main') + '?preserve=1')


@login_required(login_url='login')
def bill_choice(request, work_id):
    """
    Show a choice page asking which workslip to generate a bill from.
    work_id is the root estimate. Lists all workslips under it so the user
    can pick "Bill N from Workslip N".
    If there is only one workslip, redirect straight to generate_bill_from_saved.
    """
    org = get_org_from_request(request)
    user = request.user
    saved_work = get_object_or_404(SavedWork, id=work_id, organization=org, user=user)

    # Check bill subscription access BEFORE showing bill choices
    try:
        from subscriptions.services import SubscriptionService
        result = SubscriptionService.check_access(user, 'bill')
        if not result.get('ok', False):
            messages.warning(request, 'You need an active Bill subscription to generate bills.')
            return redirect('module_access', module_code='bill')
    except Exception:
        messages.error(request, 'Unable to verify subscription. Please try again.')
        return redirect('saved_works_list')

    # Only estimates should reach this page
    if saved_work.work_type != 'new_estimate':
        messages.error(request, 'Bill choice is only available for estimates.')
        return redirect('saved_works_list')

    # Gather all workslips under this estimate
    all_workslips = list(
        SavedWork.objects.filter(
            organization=org, user=user,
            work_type='workslip',
            parent=saved_work,
        ).order_by('workslip_number')
    )

    if not all_workslips:
        messages.info(request, 'No workslips found. Generate a workslip first.')
        return redirect('saved_works_list')

    # Gather existing bills
    workslip_ids = [ws.id for ws in all_workslips]
    all_bills = list(
        SavedWork.objects.filter(
            Q(parent=saved_work, work_type='bill') |
            Q(parent_id__in=workslip_ids, work_type='bill')
        ).filter(organization=org, user=user).order_by('bill_number')
    )

    # Determine the next bill number
    existing_bill_numbers = [b.bill_number for b in all_bills if b.bill_number]
    next_bill_number = (max(existing_bill_numbers) + 1) if existing_bill_numbers else 1

    # Find the workslip that corresponds to the next bill number
    # Bill N is generated from Workslip N
    target_ws = None
    for ws in all_workslips:
        if ws.workslip_number == next_bill_number:
            target_ws = ws
            break

    # If only one workslip or we found the target workslip, go straight to bill entry
    if len(all_workslips) == 1:
        from django.urls import reverse
        url = reverse('bill_entry', args=[all_workslips[0].id]) + f'?bill_number={next_bill_number}'
        return redirect(url)

    if target_ws:
        from django.urls import reverse
        url = reverse('bill_entry', args=[target_ws.id]) + f'?bill_number={next_bill_number}'
        return redirect(url)

    # Map: which workslip already has a bill generated?
    bills_by_ws = {}
    for bill in all_bills:
        bn = bill.bill_number or 0
        bills_by_ws[bn] = bill

    return render(request, 'core/saved_works/bill_choice.html', {
        'work': saved_work,
        'all_workslips': all_workslips,
        'all_bills': all_bills,
        'bills_by_ws': bills_by_ws,
    })


# ==============================================================================
# BILL QUANTITY ENTRY PAGE (Enter quantities per billing period)
# ==============================================================================

@login_required(login_url='login')
def bill_entry(request, work_id):
    """
    Quantity entry UI for bill generation.
    work_id is a workslip SavedWork.id.

    Shows a table of all items with:
    - Previous bills' quantities per column (read-only, for reference)
    - Input field for the current bill's quantities
    Saves quantities as a 'bill' SavedWork draft.
    After saving, user proceeds to bill_generate to download the Excel.
    """
    org = get_org_from_request(request)
    user = request.user
    workslip = get_object_or_404(SavedWork, id=work_id, organization=org, user=user)

    if workslip.work_type != 'workslip':
        messages.error(request, 'Bill entry requires a workslip.')
        return redirect('saved_works_list')

    work_data = workslip.work_data or {}
    ws_rows = work_data.get('ws_estimate_rows', [])
    tp_percent = float(work_data.get('ws_tp_percent', 0) or 0)
    tp_type = work_data.get('ws_tp_type', 'Excess')
    ws_metadata = work_data.get('ws_metadata', {})

    if not ws_rows:
        messages.error(request, 'Workslip has no items. Cannot create bill entry.')
        return redirect('saved_work_detail', work_id=work_id)

    # Gather all previously generated bills for this workslip — used for previous quantities display
    # and for determining the next bill number.  We exclude only 'draft' status so that bills
    # created via the old session-based flow (status='in_progress') are still counted.
    existing_bills = list(
        SavedWork.objects.filter(
            organization=org, user=user,
            work_type='bill',
            parent=workslip,
        ).exclude(status='draft').order_by('bill_number')
    )

    # Determine next bill number based only on completed bills
    if existing_bills:
        max_bill_num = max(b.bill_number or 0 for b in existing_bills)
        next_bill_number = max_bill_num + 1
    else:
        next_bill_number = 1

    # Check if there's already a draft or saved (in_progress) bill for next bill
    draft_bill = SavedWork.objects.filter(
        organization=org, user=user,
        work_type='bill',
        parent=workslip,
        bill_number=next_bill_number,
        status__in=['draft', 'in_progress'],
    ).first()

    draft_exec_map = {}
    if draft_bill:
        draft_data = draft_bill.work_data or {}
        draft_exec_map = draft_data.get('bill_ws_exec_map', {}) or {}
        if not draft_exec_map:
            draft_exec_map = draft_data.get('ws_exec_map', {}) or {}

    # Build per-item data: previous bills' quantities + current draft quantities
    items = []
    for idx, row in enumerate(ws_rows):
        key = row.get('key', f'saved_{idx}')
        desc = row.get('item_desc') or row.get('desc') or row.get('display_name') or row.get('item_name', '')
        unit = row.get('unit', 'Nos')
        rate = round(float(row.get('rate', 0) or 0), 2)
        qty_est = float(row.get('qty_est', 0) or 0)

        # Previous bills' quantities per bill
        prev_qtys = []
        for bill in existing_bills:
            b_data = bill.work_data or {}
            b_exec = b_data.get('bill_ws_exec_map', {})
            q = b_exec.get(key, 0)
            try:
                q = float(q) if q else 0.0
            except (ValueError, TypeError):
                q = 0.0
            prev_qtys.append({'bill_number': bill.bill_number, 'qty': q})

        # Current draft quantity (pre-filled)
        draft_qty = draft_exec_map.get(key, '')
        try:
            draft_qty = float(draft_qty) if draft_qty != '' else ''
        except (ValueError, TypeError):
            draft_qty = ''

        items.append({
            'key': key,
            'desc': desc,
            'unit': unit,
            'rate': rate,
            'qty_est': qty_est,
            'prev_qtys': prev_qtys,
            'draft_qty': draft_qty,
        })

    if request.method == 'POST':
        # Save entered quantities
        new_exec_map = {}
        for idx, row in enumerate(ws_rows):
            key = row.get('key', f'saved_{idx}')
            val = request.POST.get(f'qty_{key}', '').strip()
            if val:
                try:
                    new_exec_map[key] = float(val)
                except (ValueError, TypeError):
                    pass

        bill_save_data = {
            'bill_ws_rows': ws_rows,
            'bill_ws_exec_map': new_exec_map,
            'bill_ws_tp_percent': tp_percent,
            'bill_ws_tp_type': tp_type,
            'bill_ws_metadata': {
                'name_of_work': ws_metadata.get('work_name', ''),
                'estimate_amount': ws_metadata.get('estimate_amount', ''),
                'admin_sanction': ws_metadata.get('admin_sanction', ''),
                'tech_sanction': ws_metadata.get('tech_sanction', ''),
                'agreement': ws_metadata.get('agreement', ''),
                'agency': ws_metadata.get('agency_name', ''),
            },
            'ws_source_estimate_id': workslip.parent_id,
        }

        if draft_bill:
            draft_bill.work_data = bill_save_data
            draft_bill.bill_number = next_bill_number
            draft_bill.save(update_fields=['work_data', 'bill_number'])
        else:
            base_name = workslip.parent.name if workslip.parent else workslip.name
            draft_bill = SavedWork.objects.create(
                organization=org,
                user=user,
                folder=workslip.folder,
                parent=workslip,
                name=f'{base_name} - B{next_bill_number} (Draft)',
                work_type='bill',
                work_data=bill_save_data,
                category=workslip.category or 'electrical',
                last_step='bill_entry',
                bill_number=next_bill_number,
                status='draft',
            )

        messages.success(request, f'Bill {next_bill_number} quantities saved. Now generate the bill.')
        return redirect('bill_generate', work_id=work_id)

    # Build list of bill numbers for table header
    bill_numbers = [b.bill_number for b in existing_bills]

    from core.views import ordinal_word
    context = {
        'workslip': workslip,
        'work': workslip.parent,
        'next_bill_number': next_bill_number,
        'bill_ord': ordinal_word(next_bill_number),
        'items': items,
        'bill_numbers': bill_numbers,
        'existing_bills': existing_bills,
        'draft_bill': draft_bill,
        'tp_percent': tp_percent,
        'tp_type': tp_type,
    }
    return render(request, 'core/saved_works/bill_entry.html', context)


# ==============================================================================
# DEDICATED BILL GENERATION PAGE (from Saved WorkSlip)
# ==============================================================================

@login_required(login_url='login')
def bill_generate(request, work_id):
    """
    Dedicated bill generation page from a saved workslip.
    Generates bills, LS forms, covering letters, and movement slips
    directly from workslip data — no file upload needed.

    work_id is a workslip SavedWork.id.
    Bill number matches workslip_number (B1 from W1, B2 from W2, etc.).
    Bill 1 uses first-bill 8-column format.
    Bill 2+ uses nth-bill 11-column format with previous-bill deductions.
    """
    from core.views import (
        create_first_bill_sheet, _populate_nth_bill_sheet, _apply_print_settings,
        ordinal_word, _format_date_to_ddmmyyyy, _number_to_words_rupees,
        _build_mb_details_string, _fill_excel_template, BILL_TEMPLATES_DIR,
    )
    from core.template_views import get_user_template
    from openpyxl import Workbook
    from django.http import HttpResponse
    import io
    import os

    org = get_org_from_request(request)
    user = request.user
    workslip = get_object_or_404(SavedWork, id=work_id, organization=org, user=user)

    if workslip.work_type != 'workslip':
        messages.error(request, 'Bill generation requires a workslip.')
        return redirect('saved_works_list')

    work_data = workslip.work_data or {}
    ws_rows = work_data.get('ws_estimate_rows', [])
    tp_percent = float(work_data.get('ws_tp_percent', 0) or 0)
    tp_type = work_data.get('ws_tp_type', 'Excess')

    # Check for a saved bill (entered via bill_entry page as 'draft', or saved
    # via save_work() as 'in_progress') — use its quantities if present.
    draft_bill_record = SavedWork.objects.filter(
        organization=org, user=user,
        work_type='bill',
        parent=workslip,
        status__in=['draft', 'in_progress'],
    ).order_by('-bill_number').first()

    # Completed/finalized bills: used for "Deduct Previous Measurements" in Bill 2+.
    # Search across ALL workslips under the same root estimate, not just the current workslip.
    # This ensures B2 (parented to W2) can find B1 (parented to W1) for deduction.
    all_bill_parent_ids = [workslip.id]  # at minimum, include current workslip
    if workslip.parent:
        # Find all sibling workslips under the same root estimate
        sibling_ws_ids = list(
            SavedWork.objects.filter(
                organization=org, user=user,
                work_type='workslip',
                parent=workslip.parent,
            ).values_list('id', flat=True)
        )
        all_bill_parent_ids = sibling_ws_ids + [workslip.parent_id]  # include estimate too
    completed_bills_qs = SavedWork.objects.filter(
        organization=org, user=user,
        work_type='bill',
        parent_id__in=all_bill_parent_ids,
    ).exclude(status='draft')
    if draft_bill_record:
        completed_bills_qs = completed_bills_qs.exclude(id=draft_bill_record.id)
    completed_bills = list(completed_bills_qs.order_by('bill_number'))

    if draft_bill_record:
        # Use quantities from the saved/draft bill
        bill_number = draft_bill_record.bill_number
        # bill_entry saves under 'bill_ws_exec_map'; save_work() also uses that key
        draft_data = draft_bill_record.work_data or {}
        ws_exec_map = draft_data.get('bill_ws_exec_map', {}) or {}
        # If bill_ws_exec_map is empty, also check the session-based keys from save_work()
        if not ws_exec_map:
            ws_exec_map = draft_data.get('ws_exec_map', {}) or {}
    else:
        # Fallback: use workslip's own exec_map (Bill 1 default)
        max_completed = max((b.bill_number or 0 for b in completed_bills), default=0)
        bill_number = max_completed + 1 if max_completed > 0 else (workslip.workslip_number or 1)
        ws_exec_map = work_data.get('ws_exec_map', {}) or {}

    # Build metadata for bill header
    ws_metadata = work_data.get('ws_metadata', {})
    header_data = {
        'name_of_work': ws_metadata.get('work_name', '') or work_data.get('ws_work_name', ''),
        'estimate_amount': ws_metadata.get('estimate_amount', '') or str(work_data.get('ws_estimate_grand_total', '')),
        'admin_sanction': ws_metadata.get('admin_sanction', ''),
        'tech_sanction': ws_metadata.get('tech_sanction', ''),
        'agreement': ws_metadata.get('agreement', ''),
        'agency': ws_metadata.get('agency_name', '') or ws_metadata.get('agency', ''),
        'cc_header': ws_metadata.get('cc_header', ''),
    }

    # Build items from workslip exec data
    # First, load full descriptions from backend for all item names
    # (stored item_desc may be just the short name if backend lookup failed during workslip creation)
    all_item_names = [row.get('item_name', '') for row in ws_rows if row.get('item_name')]
    backend_descs = {}
    if all_item_names:
        category = workslip.category or 'electrical'
        saved_backend_id = work_data.get('selected_backend_id')
        try:
            backend_descs = load_item_rates_from_backend(
                category, all_item_names,
                backend_id=saved_backend_id,
                user=request.user,
                module_code='new_estimate',
            )
        except Exception:
            backend_descs = {}

    # Load prefix mapping for repair works
    # When work_mode is "repair", prefixes from the backend Groups sheet
    # are prepended to item descriptions in the bill Excel output.
    work_mode = work_data.get('ws_work_mode', 'original')
    # Also check the bill draft record for work_mode (set by bill_entry_save)
    if draft_bill_record and draft_bill_record.work_data:
        work_mode = draft_bill_record.work_data.get('work_mode', work_mode)
    is_repair = (work_mode == 'repair')

    item_to_prefix = {}
    if is_repair:
        category = workslip.category or 'electrical'
        saved_backend_id = work_data.get('selected_backend_id')
        item_to_prefix = load_prefix_map(category, backend_id=saved_backend_id, user=request.user)

    items = []
    total_amount = 0.0
    for idx, row in enumerate(ws_rows):
        # Key derivation MUST match bill_entry view: key or item_name or item_{idx}
        key = row.get('key') or row.get('item_name') or f'item_{idx}'
        exec_qty = ws_exec_map.get(key, 0)
        # Also try fallback keys in case of older data
        if not exec_qty:
            for alt_key in [f'saved_{idx}', row.get('display_name', ''), row.get('desc', '')]:
                if alt_key and alt_key in ws_exec_map:
                    exec_qty = ws_exec_map[alt_key]
                    break
        try:
            exec_qty = float(exec_qty) if exec_qty else 0.0
        except (ValueError, TypeError):
            exec_qty = 0.0
        if exec_qty <= 0 and not ws_exec_map:
            try:
                exec_qty = float(row.get('qty_est', 0) or 0)
            except (ValueError, TypeError):
                exec_qty = 0.0
        if exec_qty <= 0:
            continue
        rate = round(float(row.get('rate', 0) or 0), 2)
        if rate == 0:
            continue

        # For downloads: Prioritize backend description (row+2 content) over item_name
        # For UI display: Use item_name (yellow header)
        item_name = row.get('item_name', '')
        
        # 1) First try backend description (authoritative source - row+2 of item block)
        backend_info = backend_descs.get(item_name, {})
        backend_desc = backend_info.get('desc', '')
        
        # 2) Then try stored item_desc
        stored_item_desc = row.get('item_desc', '')
        
        # 3) Then try stored desc
        stored_desc = row.get('desc', '')
        
        # Priority order: backend_desc > item_desc > desc > item_name
        # Use first non-empty value that's different from item_name
        desc = item_name  # fallback
        for candidate in [backend_desc, stored_item_desc, stored_desc]:
            if candidate and candidate.strip() and candidate != item_name:
                desc = candidate.strip()
                break

        # Apply repair prefix if work_mode is "repair"
        if is_repair and item_to_prefix:
            prefix = item_to_prefix.get(item_name, '')
            if prefix:
                desc = f"{prefix} {desc}" if desc else prefix

        unit = row.get('unit', 'Nos')
        is_ae = str(desc).lower().startswith('ae')
        amount = exec_qty * rate
        total_amount += amount
        items.append({
            'sl': len(items) + 1,
            'qty': exec_qty,
            'unit': unit,
            'desc': desc,
            'rate': rate,
            'is_ae': is_ae,
            'amount': round(amount, 2),
            'key': key,
        })

    # Include supplemental items from the workslip
    # Previous workslip supplemental items (have rate/unit stored)
    prev_supp_items = work_data.get('ws_previous_supp_items', [])
    seen_supp_keys = set()
    for supp in prev_supp_items:
        supp_name = supp.get('name', '')
        section = supp.get('supp_section', supp.get('phase', 1))
        supp_key = f"prev_supp:{section}:{supp_name}"
        if supp_key in seen_supp_keys:
            continue
        seen_supp_keys.add(supp_key)
        exec_qty = ws_exec_map.get(supp_key, 0)
        try:
            exec_qty = float(exec_qty) if exec_qty else 0.0
        except (ValueError, TypeError):
            exec_qty = 0.0
        if exec_qty <= 0:
            continue
        rate = round(float(supp.get('rate', 0) or 0), 2)
        if rate == 0:
            continue
        amount = exec_qty * rate
        total_amount += amount
        # Use description from row+2 (stored in 'desc'), fallback to item name
        supp_desc = supp.get('desc', supp_name) or supp_name
        # Apply repair prefix
        if is_repair and item_to_prefix:
            prefix = item_to_prefix.get(supp_name, '')
            if prefix:
                supp_desc = f"{prefix} {supp_desc}" if supp_desc else prefix
        items.append({
            'sl': len(items) + 1,
            'qty': exec_qty,
            'unit': supp.get('unit', 'Nos') or 'Nos',
            'desc': supp_desc,
            'rate': rate,
            'is_ae': False,
            'amount': round(amount, 2),
            'key': supp_key,
        })

    # Current workslip supplemental items (load rates from backend)
    current_supp_names = work_data.get('ws_supp_items', [])
    if current_supp_names:
        category = workslip.category or 'electrical'
        saved_backend_id = work_data.get('selected_backend_id')
        supp_rates = load_item_rates_from_backend(
            category, current_supp_names,
            backend_id=saved_backend_id,
            user=request.user,
            module_code='new_estimate',
        )
        for supp_name in current_supp_names:
            supp_key = f"supp:{supp_name}"
            exec_qty = ws_exec_map.get(supp_key, 0)
            try:
                exec_qty = float(exec_qty) if exec_qty else 0.0
            except (ValueError, TypeError):
                exec_qty = 0.0
            if exec_qty <= 0:
                continue
            supp_info = supp_rates.get(supp_name, {})
            rate = round(float(supp_info.get('rate', 0) or 0), 2)
            if rate == 0:
                continue
            amount = exec_qty * rate
            total_amount += amount
            # Use description from backend row+2, fallback to item name
            supp_desc = supp_info.get('desc', supp_name) or supp_name
            # Apply repair prefix
            if is_repair and item_to_prefix:
                prefix = item_to_prefix.get(supp_name, '')
                if prefix:
                    supp_desc = f"{prefix} {supp_desc}" if supp_desc else prefix
            items.append({
                'sl': len(items) + 1,
                'qty': exec_qty,
                'unit': supp_info.get('unit', 'Nos') or 'Nos',
                'desc': supp_desc,
                'rate': rate,
                'is_ae': False,
                'amount': round(amount, 2),
                'key': supp_key,
            })

    # ── AE (Excess over Estimate) splitting ──
    # When workslip qty > estimate qty, split the item into:
    #   - Base row: qty capped at estimate qty
    #   - AE row: excess qty (workslip qty - estimate qty)
    # Get estimate quantities from parent estimate
    ae_qty_map = {}
    parent_estimate = workslip.parent
    if parent_estimate and parent_estimate.work_data:
        ae_qty_map = parent_estimate.work_data.get('qty_map', {})

    # Also check ws_exec_map for AE keys saved by bill_entry (keys like "ae:...")
    ae_exec_keys = {k: v for k, v in ws_exec_map.items() if k.startswith('ae:')}

    split_items = []
    ae_counter = 1
    for item in items:
        item_key = item['key']
        item_name_for_lookup = item['desc']

        # Get est_qty from ws_rows first (authoritative source), then fallback to ae_qty_map
        est_qty = 0.0
        for row in ws_rows:
            rkey = row.get('key') or row.get('item_name') or ''
            if rkey == item_key:
                est_qty = float(row.get('qty', row.get('qty_est', 0)) or 0)
                break
        if est_qty <= 0:
            est_raw = ae_qty_map.get(item_key, ae_qty_map.get(item_name_for_lookup, 0))
            try:
                est_qty = float(est_raw) if est_raw else 0.0
            except (ValueError, TypeError):
                est_qty = 0.0

        exec_qty = item['qty']
        ae_key = f"ae:{item_key}"
        ae_saved_qty = 0
        try:
            ae_saved_qty = float(ae_exec_keys.get(ae_key, 0) or 0)
        except (ValueError, TypeError):
            ae_saved_qty = 0

        if est_qty > 0 and (exec_qty > est_qty or ae_saved_qty > 0):
            # Cap main item at estimate qty
            base_qty = min(exec_qty, est_qty)
            excess_qty = ae_saved_qty if ae_saved_qty > 0 else (exec_qty - est_qty)
            if excess_qty < 0:
                excess_qty = 0

            item['qty'] = base_qty
            item['amount'] = round(base_qty * item['rate'], 2)
            split_items.append(item)

            if excess_qty > 0:
                ae_amount = round(excess_qty * item['rate'], 2)
                split_items.append({
                    'sl': None,
                    'qty': excess_qty,
                    'unit': item['unit'],
                    'desc': item['desc'],
                    'rate': item['rate'],
                    'is_ae': True,
                    'amount': ae_amount,
                    'key': ae_key,
                    'ae_number': ae_counter,
                })
                ae_counter += 1
        else:
            split_items.append(item)

    # Re-number serial numbers and recalculate total
    items = split_items
    slno = 1
    total_amount = 0.0
    for item in items:
        if not item.get('is_ae'):
            item['sl'] = slno
            slno += 1
        total_amount += item['amount']

    # For Bill 2+, get previous bill's "Total Till Date" quantities for deduction.
    # Bill N's deduction = Bill (N-1)'s Total Till Date (which is already cumulative).
    # We only use the immediately previous bill, NOT the sum of all previous bills.
    prev_qty_map = {}
    if bill_number > 1:
        # Find the immediately previous bill (bill_number - 1)
        prev_bill_for_deduct = None
        for b in completed_bills:
            if (b.bill_number or 0) == bill_number - 1:
                prev_bill_for_deduct = b
                break
        if prev_bill_for_deduct:
            pb_data = prev_bill_for_deduct.work_data or {}
            pb_exec = pb_data.get('bill_ws_exec_map', {}) or {}
            # Also check bill_exec_map (used by bill_entry save)
            if not pb_exec:
                pb_exec = pb_data.get('bill_exec_map', {}) or {}
            pb_rows = pb_data.get('bill_ws_rows', ws_rows)  # fallback to workslip rows
            seen_pb_keys = set()
            for pidx, prow in enumerate(pb_rows):
                # Key derivation must match bill_entry: key or item_name or item_{idx}
                pkey = prow.get('key') or prow.get('item_name') or f'item_{pidx}'
                seen_pb_keys.add(pkey)
                pqty = pb_exec.get(pkey, 0)
                # Also try fallback keys for older data
                if not pqty:
                    for alt_key in [f'saved_{pidx}', prow.get('display_name', ''), prow.get('desc', '')]:
                        if alt_key and alt_key in pb_exec:
                            pqty = pb_exec[alt_key]
                            break
                try:
                    pqty = round(float(pqty), 2) if pqty else 0.0
                except (ValueError, TypeError):
                    pqty = 0.0
                if pqty > 0:
                    prev_qty_map[pkey] = {
                        'qty': pqty,
                        'rate': round(float(prow.get('rate', 0) or 0), 2),
                    }

            # Also check pb_exec for supplemental keys NOT found in pb_rows
            for exec_key, exec_val in pb_exec.items():
                if exec_key in seen_pb_keys:
                    continue
                try:
                    pqty = round(float(exec_val), 2) if exec_val else 0.0
                except (ValueError, TypeError):
                    pqty = 0.0
                if pqty > 0:
                    prev_qty_map[exec_key] = {
                        'qty': pqty,
                        'rate': 0.0,
                    }

    # User document templates
    covering_template = get_user_template(user, 'covering_letter')
    movement_template = get_user_template(user, 'movement_slip')

    # Existing completed bill record for the current bill_number
    existing_bill = next((b for b in completed_bills if b.bill_number == bill_number), None)
    if existing_bill is None:
        existing_bill = draft_bill_record  # show draft as "existing" so UI shows context

    # ── POST: Generate bill or document ──
    if request.method == 'POST':
        action_type = request.POST.get('action_type', '').strip()

        mb_measure_no = str(request.POST.get('mb_measure_no') or '').strip()
        mb_measure_p_from = str(request.POST.get('mb_measure_p_from') or '').strip()
        mb_measure_p_to = str(request.POST.get('mb_measure_p_to') or '').strip()
        mb_abs_no = str(request.POST.get('mb_abstract_no') or '').strip()
        mb_abs_p_from = str(request.POST.get('mb_abstract_p_from') or '').strip()
        mb_abs_p_to = str(request.POST.get('mb_abstract_p_to') or '').strip()
        doi = _format_date_to_ddmmyyyy(request.POST.get('doi') or '')
        doc_date = _format_date_to_ddmmyyyy(request.POST.get('doc') or '')
        domr = _format_date_to_ddmmyyyy(request.POST.get('domr') or '')
        dobr = _format_date_to_ddmmyyyy(request.POST.get('dobr') or '')

        if not items:
            return JsonResponse({'error': 'No executed items found (all quantities are zero).'}, status=400)

        # ── BILL GENERATION ──
        if action_type in ('bill_part', 'bill_final'):
            is_final = action_type == 'bill_final'
            ord_text = ordinal_word(bill_number)

            if bill_number == 1:
                title_text = f'CC First & {"Final" if is_final else "Part"} Bill'
                wb_out = Workbook()
                create_first_bill_sheet(
                    wb_out, sheet_name='Bill',
                    items=items, header_data=header_data, title_text=title_text,
                    tp_percent=tp_percent, tp_type=tp_type,
                    mb_measure_no=mb_measure_no, mb_measure_p_from=mb_measure_p_from,
                    mb_measure_p_to=mb_measure_p_to,
                    mb_abs_no=mb_abs_no, mb_abs_p_from=mb_abs_p_from,
                    mb_abs_p_to=mb_abs_p_to,
                    doi=doi, doc=doc_date, domr=domr, dobr=dobr,
                )
            else:
                title_text = f'CC {ord_text} & {"Final" if is_final else "Part"} Bill'
                nth_items = []
                for item in items:
                    prev = prev_qty_map.get(item['key'], {})
                    prev_qty = prev.get('qty', 0.0)
                    prev_amount = round(prev_qty * item['rate'], 2)
                    # item['qty'] is the "Total Measurements Till Date" entered by the user
                    # (NOT "since last"). Column C = Till Date, Column I = C - G (auto-calculated).
                    nth_items.append({
                        'desc': item['desc'],
                        'unit': item['unit'],
                        'rate': item['rate'],
                        'prev_qty': prev_qty,
                        'prev_amount': prev_amount,
                        'qty_till_date': item['qty'],
                        'is_ae': item.get('is_ae', False),
                        'ae_number': item.get('ae_number', ''),
                    })

                wb_out = Workbook()
                ws_out = wb_out.active
                ws_out.title = 'Bill'
                _populate_nth_bill_sheet(
                    ws_out, items=nth_items, header_data=header_data, title_text=title_text,
                    tp_percent=tp_percent, tp_type=tp_type,
                    mb_measure_no=mb_measure_no, mb_measure_p_from=mb_measure_p_from,
                    mb_measure_p_to=mb_measure_p_to,
                    mb_abs_no=mb_abs_no, mb_abs_p_from=mb_abs_p_from,
                    mb_abs_p_to=mb_abs_p_to,
                    doi=doi, doc=doc_date, domr=domr, dobr=dobr,
                )
                # Fill in Quantity Till Date (column C) from cumulative data
                data_start = 12
                for i, nit in enumerate(nth_items):
                    ws_out.cell(row=data_start + i, column=3, value=float(nit['qty_till_date'] or 0))

            _apply_print_settings(wb_out)
            resp = HttpResponse(
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            fname = 'Bill.xlsx'
            resp['Content-Disposition'] = f'attachment; filename="{fname}"'
            wb_out.save(resp)

            # Determine bill_type based on action and bill number
            if is_final:
                saved_bill_type = 'first_final' if bill_number == 1 else 'nth_final'
            else:
                saved_bill_type = 'first_part' if bill_number == 1 else 'nth_part'

            # Auto-save the bill record — promote draft → completed, or create new
            # Build complete rows list including supplemental items so the next bill
            # can find ALL items for deduction
            complete_bill_rows = list(ws_rows)  # main estimate rows
            seen_keys = set(
                (row.get('key') or row.get('item_name') or f'item_{i}')
                for i, row in enumerate(ws_rows)
            )
            # Add supplemental items that were included in the bill
            for item in items:
                if item['key'] not in seen_keys:
                    seen_keys.add(item['key'])
                    complete_bill_rows.append({
                        'key': item['key'],
                        'item_name': item['desc'],
                        'display_name': item['desc'],
                        'unit': item['unit'],
                        'rate': item['rate'],
                        'label': 'Supplemental',
                    })
            bill_save_data = {
                'bill_ws_rows': complete_bill_rows,
                'bill_ws_exec_map': ws_exec_map,
                'bill_ws_tp_percent': tp_percent,
                'bill_ws_tp_type': tp_type,
                'bill_ws_metadata': header_data,
                'ws_source_estimate_id': workslip.parent_id,
                'source_workslip_id': workslip.id,  # explicit link for generate_next_bill_from_saved
            }
            # If draft bill exists for this bill_number, promote it to completed
            target_bill = SavedWork.objects.filter(
                organization=org, user=user, work_type='bill',
                parent=workslip, bill_number=bill_number,
            ).first()
            if target_bill:
                target_bill.work_data = bill_save_data
                target_bill.status = 'completed'
                target_bill.bill_type = saved_bill_type
                base_name = workslip.parent.name if workslip.parent else workslip.name
                target_bill.name = f'{base_name} - B{bill_number}'
                target_bill.save(update_fields=['work_data', 'status', 'bill_type', 'name'])
            else:
                base_name = workslip.parent.name if workslip.parent else workslip.name
                SavedWork.objects.create(
                    organization=org, user=user,
                    folder=workslip.folder,
                    parent=workslip,
                    name=f'{base_name} - B{bill_number}',
                    work_type='bill',
                    work_data=bill_save_data,
                    category=workslip.category or 'electrical',
                    last_step='bill',
                    bill_number=bill_number,
                    bill_type=saved_bill_type,
                    status='completed',
                )

            return resp

        # ── LS FORM ──
        if action_type in ('ls_part', 'ls_final'):
            tp_adj = total_amount * tp_percent / 100
            grand_total = (total_amount - tp_adj) if tp_type == 'Less' else (total_amount + tp_adj)
            total_str = f'{grand_total:,.2f}'
            amount_words = _number_to_words_rupees(grand_total)
            ord_text = ordinal_word(bill_number)

            if bill_number == 1:
                cc_header = 'CC First & Part Bill' if action_type == 'ls_part' else 'CC First & Final Bill'
            else:
                cc_header = f'CC {ord_text} & {"Part" if action_type == "ls_part" else "Final"} Bill'

            mb_details_str = _build_mb_details_string(
                mb_measure_no, mb_measure_p_from, mb_measure_p_to,
                mb_abs_no, mb_abs_p_from, mb_abs_p_to,
            )

            template_name = 'LS_Form_Part.xlsx' if action_type == 'ls_part' else 'LS_Form_Final.xlsx'
            template_path = os.path.join(BILL_TEMPLATES_DIR, template_name)
            if not os.path.exists(template_path):
                return HttpResponse(f'LS template not found: {template_name}', status=404)

            ctx = {
                'NAME_OF_WORK': header_data.get('name_of_work', ''),
                'NAME_OF_AGENCY': header_data.get('agency', ''),
                'AGENCY_NAME': header_data.get('agency', ''),
                'REF_OF_AGREEMENT': header_data.get('agreement', ''),
                'AGREEMENT_REF': header_data.get('agreement', ''),
                'MB_DETAILS': mb_details_str,
                'CC_HEADER': cc_header,
                'AMOUNT': total_str,
                'TOTAL_AMOUNT': total_str,
                'AMOUNT_IN_WORDS': amount_words,
            }
            wb_out = _fill_excel_template(template_name, ctx)
            _apply_print_settings(wb_out)

            resp = HttpResponse(
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            dl_name = 'LS_Form.xlsx'
            resp['Content-Disposition'] = f'attachment; filename="{dl_name}"'
            wb_out.save(resp)
            return resp

        # ── COVERING LETTER / MOVEMENT SLIP ──
        if action_type in ('covering', 'covering_part', 'covering_final', 'movement'):
            template_type = 'covering_letter' if action_type.startswith('covering') else 'movement_slip'
            user_tmpl = get_user_template(user, template_type)
            if not user_tmpl:
                return HttpResponse(
                    'Template not uploaded. Please upload your template from the Bill Generator page first.',
                    status=400,
                )
            template_bytes = user_tmpl.get_file_bytes()
            if not template_bytes:
                return HttpResponse('Template file not found. Please re-upload.', status=404)

            tp_adj = total_amount * tp_percent / 100
            grand_total = (total_amount - tp_adj) if tp_type == 'Less' else (total_amount + tp_adj)
            total_str = f'{grand_total:,.2f}'
            amount_words = _number_to_words_rupees(grand_total)
            ord_text = ordinal_word(bill_number)

            # Determine bill type based on action_type (Part or Final)
            is_final = action_type == 'covering_final'
            bill_type_text = 'Final' if is_final else 'Part'
            if bill_number == 1:
                cc_header = f'CC First & {bill_type_text} Bill'
            else:
                cc_header = f'CC {ord_text} & {bill_type_text} Bill'

            # Use CC Header from saved work data if available
            file_cc_header = (header_data.get('cc_header') or '').strip()
            if file_cc_header:
                cc_header = file_cc_header

            mb_details_str = _build_mb_details_string(
                mb_measure_no, mb_measure_p_from, mb_measure_p_to,
                mb_abs_no, mb_abs_p_from, mb_abs_p_to,
            )

            now = timezone.now()
            mm_yyyy = f'{now.month:02d}.{now.year}'

            placeholder_map = {
                '{{NAME_OF_WORK}}': header_data.get('name_of_work', ''),
                '{{AGENCY_NAME}}': header_data.get('agency', ''),
                '{{NAME_OF_AGENCY}}': header_data.get('agency', ''),
                '{{AGREEMENT_REF}}': header_data.get('agreement', ''),
                '{{REF_OF_AGREEMENT}}': header_data.get('agreement', ''),
                '{{CC_HEADER}}': cc_header,
                '{{MB_DETAILS}}': mb_details_str,
                '{{AMOUNT}}': total_str,
                '{{TOTAL_AMOUNT}}': total_str,
                '{{AMOUNT_IN_WORDS}}': amount_words,
            }

            from docx import Document
            word_doc = Document(io.BytesIO(template_bytes))

            def _replace_runs(paragraphs):
                for p in paragraphs:
                    if not p.runs:
                        continue
                    for run in p.runs:
                        if not run.text:
                            continue
                        text = run.text
                        changed = False
                        for ph, val in placeholder_map.items():
                            if ph in text:
                                text = text.replace(ph, val or '')
                                changed = True
                        if 'dd.mm.yyyy' in text:
                            text = text.replace('dd.mm.yyyy', f'  .{mm_yyyy}')
                            changed = True
                        if changed:
                            run.text = text

            _replace_runs(word_doc.paragraphs)
            for table in word_doc.tables:
                for trow in table.rows:
                    for cell in trow.cells:
                        _replace_runs(cell.paragraphs)

            buf = io.BytesIO()
            word_doc.save(buf)
            buf.seek(0)

            dl_name = 'Cover_Letter.docx' if action_type.startswith('covering') else 'Movement_Slip.docx'
            resp = HttpResponse(
                buf.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            )
            resp['Content-Disposition'] = f'attachment; filename="{dl_name}"'
            return resp

        return HttpResponse(f'Unknown action: {action_type}', status=400)

    # ── GET: Render the dedicated bill generation page ──
    context = {
        'workslip': workslip,
        'work': workslip.parent,
        'bill_number': bill_number,
        'bill_ord': ordinal_word(bill_number),
        'items': items,
        'item_count': len(items),
        'total_amount': round(total_amount, 2),
        'tp_percent': tp_percent,
        'tp_type': tp_type,
        'header_data': header_data,
        'prev_qty_map': json.dumps({k: v for k, v in prev_qty_map.items()}),
        'has_previous': bool(prev_qty_map),
        'existing_bill': existing_bill,
        'draft_bill': draft_bill_record,
        'has_draft_quantities': bool(draft_bill_record),
        'covering_letter_template': covering_template,
        'movement_slip_template': movement_template,
        'completed_bills': completed_bills,
    }
    return render(request, 'core/saved_works/bill_generate.html', context)


@login_required(login_url='login')
def generate_bill_from_saved(request, work_id):
    """
    Generate a bill from a saved workslip or estimate.
    Saved Works acts as navigation layer only - passes data to existing bill engine.
    Redirects to bill page with appropriate data loaded.
    """
    import logging
    logger = logging.getLogger(__name__)

    org = get_org_from_request(request)
    user = request.user

    # Check bill subscription access BEFORE allowing generation
    try:
        from subscriptions.services import SubscriptionService
        result = SubscriptionService.check_access(user, 'bill')
        if not result.get('ok', False):
            messages.warning(request, 'You need an active Bill subscription to generate bills.')
            return redirect('module_access', module_code='bill')
    except Exception:
        messages.error(request, 'Unable to verify subscription. Please try again.')
        return redirect('saved_works_list')

    try:
        saved_work = get_object_or_404(SavedWork, id=work_id, organization=org, user=user)
    except Exception:
        messages.error(request, 'Saved work not found or you do not have access.')
        return redirect('saved_works_list')

    # Verify this is a workslip or estimate
    if saved_work.work_type not in ['workslip', 'new_estimate']:
        messages.error(request, 'Only workslips or estimates can be used to generate bills.')
        return redirect('saved_works_list')

    work_data = saved_work.work_data or {}

    # Validate: Estimate must have data
    if saved_work.work_type == 'new_estimate':
        fetched_items = work_data.get('fetched_items', [])
        if not fetched_items:
            messages.error(request, 'Estimate has no items. Cannot generate bill from empty estimate.')
            return redirect('saved_work_detail', work_id=work_id)

    # Validate: Workslip must have estimate rows
    if saved_work.work_type == 'workslip':
        ws_rows = work_data.get('ws_estimate_rows', [])
        if not ws_rows:
            messages.error(request, 'WorkSlip has no data. Cannot generate bill from empty workslip.')
            return redirect('saved_work_detail', work_id=work_id)

    # Store source info in session for bill page
    request.session['bill_source_work_id'] = saved_work.id
    request.session['bill_source_work_type'] = saved_work.work_type
    request.session['bill_source_work_name'] = saved_work.name

    if saved_work.work_type == 'workslip':
        # B(N) from W(N): bill_number matches the source workslip_number
        bill_number = saved_work.workslip_number or 1
        request.session['bill_target_number'] = bill_number
        request.session['bill_sequence_number'] = bill_number
        # Store parent workslip ID so save_work can link the bill to this workslip
        request.session['bill_parent_work_id'] = saved_work.id

        # Load workslip data for bill generation
        request.session['bill_from_workslip'] = True
        request.session['bill_ws_rows'] = work_data.get('ws_estimate_rows', [])
        request.session['bill_ws_exec_map'] = work_data.get('ws_exec_map', {})
        request.session['bill_ws_tp_percent'] = work_data.get('ws_tp_percent', 0)
        request.session['bill_ws_tp_type'] = work_data.get('ws_tp_type', 'Excess')
        # Pass metadata for bill header (name of work, estimate amount, sanctions, agency)
        ws_metadata = work_data.get('ws_metadata', {})
        request.session['bill_ws_metadata'] = {
            'name_of_work': ws_metadata.get('work_name', '') or work_data.get('ws_work_name', ''),
            'estimate_amount': ws_metadata.get('estimate_amount', '') or str(work_data.get('ws_estimate_grand_total', '')),
            'admin_sanction': ws_metadata.get('admin_sanction', ''),
            'tech_sanction': ws_metadata.get('tech_sanction', ''),
            'agreement': ws_metadata.get('agreement', ''),
            'agency': ws_metadata.get('agency_name', ''),
        }
        # Pass supplemental items if any
        request.session['bill_ws_supp_items'] = work_data.get('ws_supp_items', [])
        # Propagate work_mode (original/repair) so bill Excel gets prefix
        request.session['work_type'] = work_data.get('ws_work_mode', 'original')

        logger.info(f"[GEN_BILL] Generating Bill-{bill_number} from WorkSlip-{saved_work.workslip_number} '{saved_work.name}' (ID: {work_id})")
    else:
        # Bill from estimate (fallback)
        request.session['bill_target_number'] = 1
        request.session['bill_sequence_number'] = 1
        request.session['bill_parent_work_id'] = saved_work.id

        request.session['bill_from_workslip'] = False
        request.session['bill_estimate_items'] = work_data.get('fetched_items', [])
        request.session['bill_qty_map'] = work_data.get('qty_map', {})
        # Propagate work_mode (original/repair) so bill Excel gets prefix
        request.session['work_type'] = work_data.get('work_type', 'original')

        logger.info(f"[GEN_BILL] Generating Bill-1 from estimate '{saved_work.name}' (ID: {work_id})")

    # For workslip source: redirect directly to bill_entry so the user enters
    # quantities for this billing period via the new bill_entry → bill_generate flow.
    # bill_entry handles draft creation and proper status tracking.
    if saved_work.work_type == 'workslip':
        request.session.modified = True
        return redirect(reverse('bill_entry', kwargs={'work_id': saved_work.id}))

    # For estimate source (legacy fallback): redirect to old bill view with session data.
    bill_number = request.session.get('bill_target_number', 1)
    request.session.modified = True
    return redirect(reverse('bill') + '?from_saved=1')


@login_required(login_url='login')
def saved_work_detail(request, work_id):
    """
    View detailed information about a saved work including workflow chain.
    Provides E, W1-W3, B1-B3 button context for the Saved Works navigation layer.
    """
    org = get_org_from_request(request)
    user = request.user

    saved_work = get_object_or_404(SavedWork, id=work_id, organization=org, user=user)

    # Find the root estimate for this workflow chain
    root_estimate = None
    if saved_work.work_type == 'new_estimate':
        root_estimate = saved_work
    else:
        # Walk up parent chain to find root estimate
        current = saved_work.parent
        while current:
            if current.work_type == 'new_estimate':
                root_estimate = current
                break
            current = current.parent

        # Fallback for orphan workslips/bills (parent not set):
        # Try to find the estimate by ws_source_estimate_id stored in work_data
        if not root_estimate:
            work_data = saved_work.work_data or {}
            source_est_id = work_data.get('ws_source_estimate_id') or work_data.get('bill_source_work_id')
            if source_est_id:
                try:
                    candidate = SavedWork.objects.get(id=source_est_id, organization=org, user=user)
                    if candidate.work_type == 'new_estimate':
                        root_estimate = candidate
                        # Fix the parent relationship while we're at it
                        saved_work.parent = root_estimate
                        saved_work.save(update_fields=['parent'])
                    else:
                        # Maybe the source is a workslip; walk up from there
                        cur = candidate.parent
                        while cur:
                            if cur.work_type == 'new_estimate':
                                root_estimate = cur
                                break
                            cur = cur.parent
                except SavedWork.DoesNotExist:
                    pass

        # Last fallback: find estimate with same name in same org
        if not root_estimate:
            root_estimate = SavedWork.objects.filter(
                organization=org, user=user,
                work_type='new_estimate',
                name=saved_work.name,
            ).first()

    # Build workflow navigation: E, W1-W3, B1-B3 buttons
    workslips = []
    bills = []
    if root_estimate:
        # Find direct child workslips of root estimate
        workslips = list(
            SavedWork.objects.filter(
                organization=org, user=user,
                work_type='workslip',
                parent=root_estimate,
            ).order_by('workslip_number')
        )

        # Also find workslips chained to other workslips (W2→W1, W3→W2)
        # These were created before the fix that makes all workslips direct children
        existing_ws_ids = {ws.id for ws in workslips}
        chained_ws = SavedWork.objects.filter(
            organization=org, user=user,
            work_type='workslip',
            parent_id__in=existing_ws_ids,
        )
        for cw in chained_ws:
            if cw.id not in existing_ws_ids:
                # Fix: re-parent to root estimate
                cw.parent = root_estimate
                cw.save(update_fields=['parent'])
                workslips.append(cw)
                existing_ws_ids.add(cw.id)
        # Recursively find deeper chains (W3→W2→W1)
        depth = 0
        while depth < 10:
            deeper = SavedWork.objects.filter(
                organization=org, user=user,
                work_type='workslip',
                parent_id__in=existing_ws_ids,
            ).exclude(id__in=existing_ws_ids)
            if not deeper.exists():
                break
            for dw in deeper:
                dw.parent = root_estimate
                dw.save(update_fields=['parent'])
                workslips.append(dw)
                existing_ws_ids.add(dw.id)
            depth += 1

        # Also find orphan workslips (parent not set) that reference this estimate
        # via ws_source_estimate_id in their work_data — fix their parent and include them
        existing_ws_ids = {ws.id for ws in workslips}
        orphan_workslips = SavedWork.objects.filter(
            organization=org, user=user,
            work_type='workslip',
            parent__isnull=True,
        )
        for ow in orphan_workslips:
            wd = ow.work_data or {}
            src_id = wd.get('ws_source_estimate_id')
            if src_id and int(src_id) == root_estimate.id and ow.id not in existing_ws_ids:
                # Fix the parent relationship
                ow.parent = root_estimate
                ow.save(update_fields=['parent'])
                workslips.append(ow)
                existing_ws_ids.add(ow.id)
        # Re-sort workslips by workslip_number after adding orphans
        workslips.sort(key=lambda w: w.workslip_number or 0)

        # Bills: children of root_estimate or children of any of its workslips
        workslip_ids = [ws.id for ws in workslips]
        bills = list(
            SavedWork.objects.filter(
                Q(parent=root_estimate, work_type='bill') |
                Q(parent_id__in=workslip_ids, work_type='bill')
            ).filter(organization=org, user=user).order_by('bill_number')
        )

        # Also find orphan bills that reference this estimate or its workslips
        existing_bill_ids = {b.id for b in bills}
        orphan_bills = SavedWork.objects.filter(
            organization=org, user=user,
            work_type='bill',
            parent__isnull=True,
        )
        for ob in orphan_bills:
            bd = ob.work_data or {}
            src_id = bd.get('bill_source_work_id')
            if src_id:
                src_id = int(src_id)
                if (src_id == root_estimate.id or src_id in existing_ws_ids) and ob.id not in existing_bill_ids:
                    ob.parent = root_estimate
                    ob.save(update_fields=['parent'])
                    bills.append(ob)
                    existing_bill_ids.add(ob.id)
        bills.sort(key=lambda b: b.bill_number or 0)

    # Handle orphan workslips (saved via upload module with no parent estimate)
    is_orphan_workslip = False
    if not root_estimate and saved_work.work_type == 'workslip':
        is_orphan_workslip = True
        workslips = [saved_work]
        bills = list(
            SavedWork.objects.filter(
                organization=org, user=user,
                work_type='bill',
                parent=saved_work,
            ).order_by('bill_number')
        )

    # Handle bills belonging to orphan workslips
    # (when viewing a bill whose parent workslip has no parent estimate)
    if not root_estimate and saved_work.work_type == 'bill':
        parent_workslip = saved_work.parent
        if parent_workslip and parent_workslip.work_type == 'workslip':
            is_orphan_workslip = True
            workslips = [parent_workslip]
            bills = list(
                SavedWork.objects.filter(
                    organization=org, user=user,
                    work_type='bill',
                    parent=parent_workslip,
                ).order_by('bill_number')
            )

    # Get workflow chain (parents)
    parent_chain = []
    current = saved_work.parent
    while current:
        parent_chain.insert(0, current)
        current = current.parent

    children = saved_work.children.all()

    # Check subscription access for this work
    access_result = check_saved_work_access(user, saved_work)

    # Module access checks
    module_access = {}
    try:
        from subscriptions.services import SubscriptionService
        for wt, mc in WORK_TYPE_TO_MODULE.items():
            result = SubscriptionService.check_access(user, mc)
            module_access[wt] = result.get('ok', False)
    except Exception:
        for wt in WORK_TYPE_TO_MODULE:
            module_access[wt] = False

    # ===========================================================
    # BILL PREVIEW: Build preview rows for workslip/bill types
    # Shows data as it would appear in the generated Excel bill
    # ===========================================================
    bill_preview_rows = []
    bill_preview_number = 0
    bill_preview_total = 0.0
    previous_bill_rows = []  # For Bill 2+ to show deductions

    if saved_work.work_type == 'workslip':
        work_data = saved_work.work_data or {}
        ws_rows = work_data.get('ws_estimate_rows', [])
        ws_exec = work_data.get('ws_exec_map', {}) or {}
        bill_preview_number = saved_work.workslip_number or 1

        for idx, row in enumerate(ws_rows):
            key = row.get('key', f'saved_{idx}')
            exec_qty = ws_exec.get(key, 0)
            try:
                exec_qty = float(exec_qty) if exec_qty else 0.0
            except (ValueError, TypeError):
                exec_qty = 0.0
            if exec_qty <= 0:
                continue
            rate = round(float(row.get('rate', 0) or 0), 2)
            amount = round(exec_qty * rate, 2)
            bill_preview_total += amount
            _raw_name = row.get('desc') or row.get('display_name') or row.get('item_name', '')
            _raw_desc = row.get('desc', '')
            bill_preview_rows.append({
                'sl': len(bill_preview_rows) + 1,
                'name': _raw_name,
                'desc': _raw_desc,
                'unit': row.get('unit', 'Nos'),
                'qty': exec_qty,
                'rate': rate,
                'amount': amount,
                'key': key,
            })

        # For Bill 2+, find the previous bill's data for deductions
        if bill_preview_number > 1:
            prev_bill = None
            if root_estimate:
                prev_bill = SavedWork.objects.filter(
                    organization=org, user=user, work_type='bill',
                    bill_number=bill_preview_number - 1,
                    parent=root_estimate,
                ).first()
            if not prev_bill:
                # Check if the previous bill is a child of a workslip
                for ws in workslips:
                    prev_bill = SavedWork.objects.filter(
                        organization=org, user=user, work_type='bill',
                        bill_number=bill_preview_number - 1,
                        parent=ws,
                    ).first()
                    if prev_bill:
                        break

            if prev_bill:
                prev_data = prev_bill.work_data or {}
                prev_ws_rows = prev_data.get('ws_estimate_rows', [])
                prev_exec = prev_data.get('ws_exec_map', prev_data.get('bill_ws_exec_map', {})) or {}
                for pidx, prow in enumerate(prev_ws_rows):
                    pkey = prow.get('key', f'saved_{pidx}')
                    pqty = prev_exec.get(pkey, 0)
                    try:
                        pqty = float(pqty) if pqty else 0.0
                    except (ValueError, TypeError):
                        pqty = 0.0
                    previous_bill_rows.append({
                        'key': pkey,
                        'name': prow.get('display_name') or prow.get('item_name', ''),
                        'qty': pqty,
                    })

    elif saved_work.work_type == 'bill':
        work_data = saved_work.work_data or {}
        ws_rows = work_data.get('bill_ws_rows', work_data.get('ws_estimate_rows', []))
        ws_exec = work_data.get('bill_ws_exec_map', work_data.get('ws_exec_map', {})) or {}
        bill_preview_number = saved_work.bill_number or 1

        for idx, row in enumerate(ws_rows):
            key = row.get('key', f'saved_{idx}')
            exec_qty = ws_exec.get(key, 0)
            try:
                exec_qty = float(exec_qty) if exec_qty else 0.0
            except (ValueError, TypeError):
                exec_qty = 0.0
            if exec_qty <= 0:
                continue
            rate = round(float(row.get('rate', 0) or 0), 2)
            amount = round(exec_qty * rate, 2)
            bill_preview_total += amount
            _raw_name = row.get('desc') or row.get('display_name') or row.get('item_name', '')
            _raw_desc = row.get('desc', '')
            bill_preview_rows.append({
                'sl': len(bill_preview_rows) + 1,
                'name': _raw_name,
                'desc': _raw_desc,
                'unit': row.get('unit', 'Nos'),
                'qty': exec_qty,
                'rate': rate,
                'amount': amount,
                'key': key,
            })

    # Determine last workslip for the W button prompt
    last_workslip = workslips[-1] if workslips else None

    context = {
        'work': saved_work,
        'root_estimate': root_estimate,
        'is_orphan_workslip': is_orphan_workslip,
        'workslips': workslips,
        'bills': bills,
        'last_workslip': last_workslip,
        'parent_chain': parent_chain,
        'children': children,
        'can_generate_workslip': saved_work.can_generate_workslip(),
        'can_generate_bill': saved_work.can_generate_bill(),
        'has_subscription_access': access_result['ok'],
        'subscription_reason': access_result.get('reason', ''),
        'module_code': access_result.get('module_code'),
        'module_access': module_access,
        'bill_preview_rows': bill_preview_rows,
        'bill_preview_number': bill_preview_number,
        'bill_preview_total': bill_preview_total,
        'previous_bill_rows': json.dumps(previous_bill_rows),
    }

    return render(request, 'core/saved_works/detail.html', context)


# ==============================================================================
# BILL GENERATION FROM SAVED WORKS (B1, B2, B3 FLOW)
# ==============================================================================

@login_required(login_url='login')
def generate_next_bill_from_saved(request, work_id):
    """
    Generate the next bill from a saved bill.
    For example, if work_id is Bill-1, this generates Bill-2.
    Saved Works acts as navigation layer only — passes IDs to existing bill engine.
    """
    import logging
    logger = logging.getLogger(__name__)

    org = get_org_from_request(request)
    user = request.user

    saved_work = get_object_or_404(SavedWork, id=work_id, organization=org, user=user)

    # Validate: must be a bill
    if saved_work.work_type != 'bill':
        return JsonResponse({
            'success': False,
            'error': 'Only saved bills can generate next bills.'
        }, status=400)

    if saved_work.status != 'completed':
        return JsonResponse({
            'success': False,
            'error': 'Bill must be completed before generating the next bill.'
        }, status=400)

    current_bill_number = saved_work.bill_number or 1
    next_bill_number = current_bill_number + 1

    logger.info(
        f"[NEXT_BILL] Generating Bill-{next_bill_number} from saved Bill-{current_bill_number} (ID: {work_id})"
    )

    work_data = saved_work.work_data or {}

    # Find the source workslip for this bill chain — walk up parent to find the workslip
    source_workslip_id = work_data.get('source_workslip_id')
    if not source_workslip_id and saved_work.parent:
        # If this bill's parent is a workslip, use that
        if saved_work.parent.work_type == 'workslip':
            source_workslip_id = saved_work.parent.id
        elif saved_work.parent.work_type in ('new_estimate', 'temporary_works', 'amc'):
            # Bill parented to an estimate — try to find the newest completed workslip
            newest_ws = SavedWork.objects.filter(
                organization=org, user=user,
                work_type='workslip',
                parent=saved_work.parent,
                status='completed',
            ).order_by('-workslip_number').first()
            if newest_ws:
                source_workslip_id = newest_ws.id

    # Redirect to bill_entry for the next workslip or pass bill_number override
    if source_workslip_id:
        # Try to find the correct workslip for the next bill number
        # (e.g., Bill 2 should come from W2 if it exists)
        source_ws = SavedWork.objects.filter(id=source_workslip_id).first()
        if source_ws:
            root_estimate = source_ws.parent if source_ws.work_type == 'workslip' else source_ws
            if root_estimate:
                next_ws = SavedWork.objects.filter(
                    organization=org, user=user,
                    work_type='workslip',
                    parent=root_estimate,
                    workslip_number=next_bill_number,
                ).first()
                if next_ws:
                    # W(N+1) exists, use it directly
                    messages.success(request, f'Enter quantities for Bill-{next_bill_number}.')
                    return redirect(f"{reverse('bill_entry', kwargs={'work_id': next_ws.id})}")
        
        # No matching workslip found - use the same workslip with bill_number override
        messages.success(request, f'Enter quantities for Bill-{next_bill_number}.')
        return redirect(f"{reverse('bill_entry', kwargs={'work_id': int(source_workslip_id)})}?bill_number={next_bill_number}")

    # Fallback if no workslip found
    messages.error(request, 'Could not find source workslip for this bill chain.')
    return redirect('saved_works_list')


@login_required(login_url='login')
@require_POST
def saved_work_action(request, work_id):
    """
    Handle AJAX actions from the Saved Works detail page.
    This is the navigation layer that routes to existing module APIs.
    Actions: update_workslip, start_next_workslip, update_bill, start_next_bill
    """
    org = get_org_from_request(request)
    user = request.user

    saved_work = get_object_or_404(SavedWork, id=work_id, organization=org, user=user)

    action = request.POST.get('action', '')

    if action == 'update_estimate':
        # Resume existing estimate — same as resume_saved_work
        return JsonResponse({
            'success': True,
            'redirect_url': reverse('resume_saved_work', kwargs={'work_id': work_id}),
            'action': 'update_estimate',
        })

    elif action == 'update_workslip':
        # Resume existing workslip
        return JsonResponse({
            'success': True,
            'redirect_url': reverse('resume_saved_work', kwargs={'work_id': work_id}),
            'action': 'update_workslip',
        })

    elif action == 'start_next_workslip':
        # Validate: workslip must be completed
        if saved_work.work_type != 'workslip':
            return JsonResponse({
                'success': False,
                'error': 'Only completed workslips can generate next workslips.'
            }, status=400)
        return JsonResponse({
            'success': True,
            'redirect_url': reverse('generate_next_workslip_from_saved', kwargs={'work_id': work_id}),
            'action': 'start_next_workslip',
        })

    elif action == 'update_bill':
        # Resume existing bill
        return JsonResponse({
            'success': True,
            'redirect_url': reverse('resume_saved_work', kwargs={'work_id': work_id}),
            'action': 'update_bill',
        })

    elif action == 'start_next_bill':
        # Generate next bill from existing bill
        if saved_work.work_type != 'bill':
            return JsonResponse({
                'success': False,
                'error': 'Only completed bills can generate next bills.'
            }, status=400)
        return JsonResponse({
            'success': True,
            'redirect_url': reverse('generate_next_bill_from_saved', kwargs={'work_id': work_id}),
            'action': 'start_next_bill',
        })

    elif action == 'generate_first_bill':
        # Generate Bill-1 from a workslip — go through bill_entry so quantities
        # are entered per billing period and the record gets status='completed' correctly.
        if saved_work.work_type != 'workslip':
            return JsonResponse({
                'success': False,
                'error': 'Only workslips can generate bills.'
            }, status=400)
        return JsonResponse({
            'success': True,
            'redirect_url': reverse('bill_entry', kwargs={'work_id': work_id}),
            'action': 'generate_first_bill',
        })

    return JsonResponse({
        'success': False,
        'error': f'Unknown action: {action}'
    }, status=400)


@login_required(login_url='login')
@require_POST
def save_with_parent(request):
    """
    Save current work with a parent reference (for workflow chain).
    Called when saving a workslip generated from an estimate, or bill from workslip.
    """
    org = get_org_from_request(request)
    user = request.user
    
    work_name = request.POST.get('work_name', '').strip()
    work_type = request.POST.get('work_type', '')
    parent_id = request.POST.get('parent_id')
    folder_id = request.POST.get('folder_id')
    notes = request.POST.get('notes', '').strip()
    category = request.POST.get('category', 'electrical')
    
    if not work_name:
        return JsonResponse({'success': False, 'error': 'Work name is required.'})
    
    if work_type not in dict(SavedWork.WORK_TYPE_CHOICES):
        return JsonResponse({'success': False, 'error': 'Invalid work type.'})
    
    # Get parent if specified
    parent = None
    if parent_id:
        parent = get_object_or_404(SavedWork, id=parent_id, organization=org, user=user)
    
    # Get folder if specified
    folder = None
    if folder_id:
        folder = get_object_or_404(WorkFolder, id=folder_id, organization=org, user=user)
    
    # Collect work data
    work_data = collect_work_data(request, work_type)
    progress_percent = calculate_progress(work_data, work_type)
    last_step = get_last_step(request, work_type)
    
    # Auto-inherit parent's folder if no folder specified
    if not folder and parent and parent.folder:
        folder = parent.folder

    # Determine workslip_number / bill_number from session
    workslip_number = 1
    bill_number = 1
    if work_type == 'workslip':
        workslip_number = request.session.get('ws_target_workslip', 1) or 1
    elif work_type == 'bill':
        bill_number = request.session.get('bill_target_number', 1) or 1

    # Create saved work with parent reference
    saved_work = SavedWork.objects.create(
        organization=org,
        user=user,
        folder=folder,
        parent=parent,
        name=work_name,
        work_type=work_type,
        work_data=work_data,
        category=category,
        notes=notes,
        progress_percent=progress_percent,
        last_step=last_step,
        workslip_number=workslip_number,
        bill_number=bill_number,
    )
    
    # Store saved work ID in session
    request.session['current_saved_work_id'] = saved_work.id
    
    return JsonResponse({
        'success': True,
        'work_id': saved_work.id,
        'message': f'Work "{work_name}" saved successfully!'
    })
