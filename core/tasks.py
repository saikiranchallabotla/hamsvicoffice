# core/tasks.py
"""
Celery tasks for asynchronous Excel processing.

These tasks run in the background using Celery + Redis.
All Excel parsing, generation, and data processing happens here,
not in HTTP request handlers.
"""

import json
import logging
import traceback
from datetime import datetime
from celery import shared_task
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage
from core.models import Job, Upload, OutputFile, Organization


def normalize_text(text):
    """
    Normalize text by replacing special unicode characters with standard ASCII equivalents.
    Fixes encoding issues like em-dash (—) appearing as â€" in output.
    """
    if text is None:
        return ""
    text = str(text)
    # Replace various unicode dashes with standard hyphen
    text = text.replace('—', '-')  # em dash
    text = text.replace('–', '-')  # en dash
    text = text.replace('−', '-')  # minus sign
    text = text.replace('‐', '-')  # hyphen
    text = text.replace('‑', '-')  # non-breaking hyphen
    text = text.replace('‒', '-')  # figure dash
    # Replace other common problematic characters
    text = text.replace(''', "'")  # left single quote
    text = text.replace(''', "'")  # right single quote
    text = text.replace('"', '"')  # left double quote
    text = text.replace('"', '"')  # right double quote
    text = text.replace('…', '...')  # ellipsis
    text = text.replace('\u00a0', ' ')  # non-breaking space
    return text


# TODO: Implement these functions in utils_excel.py
# For now, use stubs to allow server to run
def read_excel_file(file_path):
    """Stub: Read and parse Excel file. To be implemented."""
    import openpyxl
    wb = openpyxl.load_workbook(file_path, data_only=True)
    return {'sheets': [ws.title for ws in wb.worksheets], 'status': 'parsed'}

def generate_bill(data, template=None):
    """Stub: Generate bill from data. To be implemented."""
    return {'status': 'generated', 'type': 'bill'}

def generate_workslip(data, template=None):
    """Stub: Generate workslip from data. To be implemented."""
    return {'status': 'generated', 'type': 'workslip'}


logger = logging.getLogger(__name__)


@shared_task(bind=True, max_retries=3)
def process_excel_upload(self, upload_id):
    """
    Process an uploaded Excel file.
    
    Args:
        upload_id: ID of the Upload object
    
    Returns:
        JSON result with status and output_file_id
    
    Raises:
        Retries up to 3 times on failure
    """
    try:
        upload = Upload.objects.get(id=upload_id)
        job = upload.jobs.first()  # Get the first job associated with this upload
        
        if not job:
            logger.error(f"No job found for upload {upload_id}")
            return {'status': 'error', 'error': 'No job found for upload'}
        
        # Update job: mark as running
        job.status = 'running'
        job.current_step = "Reading Excel file..."
        job.started_at = datetime.now()
        job.progress = 10
        job.save()
        
        logger.info(f"Starting Excel processing for upload {upload_id}")
        
        # Read Excel file
        try:
            file_path = upload.file.name
            data = read_excel_file(file_path)
            job.progress = 30
            job.current_step = "Parsing data..."
            job.save()
        except Exception as e:
            job.status = 'failed'
            job.error_message = f"Failed to read Excel file: {str(e)}"
            job.error_log.append({
                'timestamp': datetime.now().isoformat(),
                'step': 'read_excel',
                'error': str(e),
                'traceback': traceback.format_exc(),
            })
            job.save()
            upload.status = 'failed'
            upload.save()
            logger.error(f"Excel read failed for upload {upload_id}: {e}")
            return {
                'status': 'failed',
                'error': 'Failed to read Excel file',
            }
        
        # Store parsed data in job result
        job.result = data
        job.progress = 50
        job.current_step = "Data stored successfully"
        job.save()
        
        # Create OutputFile for the raw data
        output_file = OutputFile.objects.create(
            job=job,
            user=job.user,
            organization=job.organization,
            filename=f"parsed_data_{job.id}.json",
            file_type="json",
            file_size=len(json.dumps(data)),
        )
        
        # Store JSON content
        json_content = ContentFile(
            json.dumps(data, indent=2).encode('utf-8'),
            name=output_file.filename
        )
        output_file.file = default_storage.save(
            f"outputs/{job.organization.slug}/{output_file.filename}",
            json_content
        )
        output_file.save()
        
        # Mark upload as completed
        upload.status = 'completed'
        upload.save()
        
        # Mark job as completed
        job.status = 'completed'
        job.progress = 100
        job.current_step = "Processing complete"
        job.completed_at = datetime.now()
        job.save()
        
        logger.info(f"Excel processing completed for upload {upload_id}")
        
        return {
            'status': 'success',
            'output_file_id': output_file.id,
            'result': data,
        }
    
    except Upload.DoesNotExist:
        logger.error(f"Upload {upload_id} not found")
        return {
            'status': 'error',
            'error': f'Upload {upload_id} not found',
        }
    
    except Exception as e:
        logger.error(f"Unexpected error in process_excel_upload: {e}\n{traceback.format_exc()}")
        
        # Retry with exponential backoff
        try:
            upload = Upload.objects.get(id=upload_id)
            job = upload.jobs.first()
            if job:
                job.error_log.append({
                    'timestamp': datetime.now().isoformat(),
                    'attempt': self.request.retries,
                    'error': str(e),
                    'traceback': traceback.format_exc(),
                })
                job.save()
        except:
            pass
        
        # Retry after 60, 120, 240 seconds
        raise self.retry(exc=e, countdown=60 * (2 ** self.request.retries))


@shared_task(bind=True)
def generate_bill_pdf(self, job_id, project_id):
    """
    Generate a bill PDF from job data.
    
    Args:
        job_id: ID of the Job object
        project_id: ID of the Project object
    
    Returns:
        JSON result with output_file_id
    """
    try:
        job = Job.objects.get(id=job_id)
        
        job.status = 'running'
        job.current_step = "Generating bill PDF..."
        job.progress = 50
        job.save()
        
        logger.info(f"Starting bill generation for job {job_id}")
        
        # Generate bill using existing utility
        bill_content = generate_bill(
            project_id=project_id,
            data=job.result,
            job_id=job_id,
        )
        
        # Create OutputFile for PDF
        output_file = OutputFile.objects.create(
            job=job,
            user=job.user,
            organization=job.organization,
            filename=f"bill_{job_id}.pdf",
            file_type="pdf",
            file_size=len(bill_content),
        )
        
        # Store PDF
        pdf_file = ContentFile(bill_content, name=output_file.filename)
        output_file.file = default_storage.save(
            f"outputs/{job.organization.slug}/{output_file.filename}",
            pdf_file
        )
        output_file.save()
        
        # Update job
        job.status = 'completed'
        job.progress = 100
        job.current_step = "Bill generated"
        job.completed_at = datetime.now()
        job.save()
        
        logger.info(f"Bill generation completed for job {job_id}")
        
        return {
            'status': 'success',
            'output_file_id': output_file.id,
        }
    
    except Exception as e:
        logger.error(f"Bill generation failed: {e}\n{traceback.format_exc()}")
        
        try:
            job = Job.objects.get(id=job_id)
            job.status = 'failed'
            job.error_message = f"Failed to generate bill: {str(e)}"
            job.error_log.append({
                'timestamp': datetime.now().isoformat(),
                'step': 'generate_bill',
                'error': str(e),
                'traceback': traceback.format_exc(),
            })
            job.save()
        except:
            pass
        
        return {
            'status': 'failed',
            'error': str(e),
        }


@shared_task(bind=True)
def generate_workslip_pdf(self, job_id, project_id):
    """
    Generate a workslip PDF from job data.
    
    Args:
        job_id: ID of the Job object
        project_id: ID of the Project object
    
    Returns:
        JSON result with output_file_id
    """
    try:
        job = Job.objects.get(id=job_id)
        
        job.status = 'running'
        job.current_step = "Generating workslip PDF..."
        job.progress = 50
        job.save()
        
        logger.info(f"Starting workslip generation for job {job_id}")
        
        # Generate workslip using existing utility
        workslip_content = generate_workslip(
            project_id=project_id,
            data=job.result,
            job_id=job_id,
        )
        
        # Create OutputFile for PDF
        output_file = OutputFile.objects.create(
            job=job,
            user=job.user,
            organization=job.organization,
            filename=f"workslip_{job_id}.pdf",
            file_type="pdf",
            file_size=len(workslip_content),
        )
        
        # Store PDF
        pdf_file = ContentFile(workslip_content, name=output_file.filename)
        output_file.file = default_storage.save(
            f"outputs/{job.organization.slug}/{output_file.filename}",
            pdf_file
        )
        output_file.save()
        
        # Update job
        job.status = 'completed'
        job.progress = 100
        job.current_step = "Workslip generated"
        job.completed_at = datetime.now()
        job.save()
        
        logger.info(f"Workslip generation completed for job {job_id}")
        
        return {
            'status': 'success',
            'output_file_id': output_file.id,
        }
    
    except Exception as e:
        logger.error(f"Workslip generation failed: {e}\n{traceback.format_exc()}")
        
        try:
            job = Job.objects.get(id=job_id)
            job.status = 'failed'
            job.error_message = f"Failed to generate workslip: {str(e)}"
            job.error_log.append({
                'timestamp': datetime.now().isoformat(),
                'step': 'generate_workslip',
                'error': str(e),
                'traceback': traceback.format_exc(),
            })
            job.save()
        except:
            pass
        
        return {
            'status': 'failed',
            'error': str(e),
        }


@shared_task(bind=True, max_retries=2)
def generate_output_excel(self, job_id, category, qty_map_json, unit_map_json, work_name, work_type, grand_total=None, excess_tp_percent=None, ls_special_name=None, ls_special_amount=None, deduct_old_material=None, backend_id=None):
    """
    Generate Output + Estimate Excel workbook asynchronously.
    
    This task handles the heavy lifting for download_output() view.
    
    Args:
        job_id: Job.id to track progress
        category: Category name for backend loading
        qty_map_json: JSON string of item quantities
        unit_map_json: JSON string of custom units per item
        work_name: Name of the work
        work_type: "original" or "repair"
        grand_total: Manually entered grand total amount (optional)
        excess_tp_percent: Percentage for Excess T.P (optional)
        ls_special_name: Name for L.S Provision special item (optional)
        ls_special_amount: Amount for L.S Provision special item (optional)
        deduct_old_material: Amount to deduct for old material cost (repair work, optional)
        backend_id: ID of the ModuleBackend to use (optional, for multi-backend support)
    
    Returns:
        JSON with status and output_file_id
    """
    try:
        from django.conf import settings
        from django.utils import timezone
        from openpyxl import Workbook, load_workbook
        from openpyxl.styles import Alignment, Font, Border, Side
        from io import BytesIO
        from core.utils_excel import load_backend, copy_block_with_styles_and_formulas
        
        job = Job.objects.get(id=job_id)
        job.status = 'running'
        job.started_at = timezone.now()
        job.progress = 5
        job.current_step = "Loading backend data..."
        job.save()
        
        # Parse input
        qty_map = json.loads(qty_map_json) if qty_map_json else {}
        unit_map = json.loads(unit_map_json) if unit_map_json else {}
        is_repair = (work_type == "repair")
        
        # Load backend
        job.progress = 15
        job.current_step = "Loading items and groups..."
        job.save()
        
        # Get backend_id from job result if not passed directly (for backward compatibility)
        if not backend_id and job.result:
            backend_id = job.result.get('backend_id')
        
        # Determine module code - use 'amc' for AMC module, 'new_estimate' for New Estimate
        module_code = 'new_estimate'
        if job.result and job.result.get('module') == 'amc':
            module_code = 'amc'
        
        items_list, groups_map, backend_units_map, ws_src, filepath = load_backend(
            category, settings.BASE_DIR,
            backend_id=backend_id,
            module_code=module_code
        )
        name_to_info = {it["name"]: it for it in items_list}
        
        # Map items to groups
        item_to_group = {}
        for grp, items_in_grp in groups_map.items():
            for nm in items_in_grp:
                item_to_group.setdefault(nm, grp)
        
        # Load prefix mapping
        job.progress = 20
        job.current_step = "Loading prefix data..."
        job.save()
        
        item_to_prefix = {}
        try:
            backend_wb = load_workbook(filepath, data_only=False)
            ws_groups = backend_wb["Groups"]
            
            header_row = None
            col_item = None
            col_prefix = None
            
            for r in range(1, ws_groups.max_row + 1):
                for c in range(1, ws_groups.max_column + 1):
                    val = str(ws_groups.cell(row=r, column=c).value or "").strip().lower()
                    if val == "item name":
                        header_row = r
                        col_item = c
                    elif val == "prefix":
                        col_prefix = c
                if header_row:
                    break
            
            if header_row and col_item and col_prefix:
                for r in range(header_row + 1, ws_groups.max_row + 1):
                    nm = ws_groups.cell(r, col_item).value
                    px = ws_groups.cell(r, col_prefix).value
                    if nm and px not in (None, ""):
                        item_to_prefix[str(nm).strip()] = str(px).strip()
        except Exception:
            pass  # Continue without prefixes
        
        # Assuming fetched_items is in job metadata or passed somehow
        # For now, we get it from the session-based workflow
        # In future, this should be stored in Job.result as input data
        fetched = job.result.get('fetched_items', []) if job.result else []
        
        # Log items found vs missing for debugging
        missing_items = [name for name in fetched if name not in name_to_info]
        if missing_items:
            logger.warning(f"Job {job_id}: {len(missing_items)} items not found in backend: {missing_items[:5]}{'...' if len(missing_items) > 5 else ''}")
        logger.info(f"Job {job_id}: Processing {len(fetched)} fetched items, {len(fetched) - len(missing_items)} found in backend")
        
        job.progress = 30
        job.current_step = "Building Output sheet..."
        job.save()
        
        # Create workbook
        wb = Workbook()
        ws_out = wb.active
        ws_out.title = "Output"

        thin = Side(border_style="thin", color="000000")
        border_all = Border(top=thin, left=thin, right=thin, bottom=thin)

        # Add "Name of Work" header at top of Output sheet
        ws_out.merge_cells("A1:J1")
        c1 = ws_out["A1"]
        c1.value = f"Name of the work : {normalize_text(work_name)}" if work_name else "Name of the work : "
        c1.font = Font(bold=True, size=11)
        c1.alignment = Alignment(horizontal="left", vertical="center")
        for col in range(1, 11):
            ws_out.cell(row=1, column=col).border = border_all

        cursor = 3  # start item blocks after header + blank row
        rate_pos = {}
        data_serial = 1

        # Build Output sheet
        for idx, name in enumerate(fetched):
            if idx % max(1, len(fetched) // 5) == 0:
                job.progress = 30 + int((idx / len(fetched)) * 40)
                job.save()
            
            info = name_to_info.get(name)
            if not info:
                continue
            
            src_min = info["start_row"]
            src_max = info["end_row"]
            
            rate_src_row = None
            for r in range(src_max, src_min, -1):
                v = ws_src.cell(row=r, column=10).value
                if v not in (None, ""):
                    rate_src_row = r
                    break
            
            dst_start = cursor
            
            copy_block_with_styles_and_formulas(
                ws_src=ws_src,
                ws_dst=ws_out,
                src_min_row=src_min,
                src_max_row=src_max,
                dst_start_row=dst_start,
                col_start=1,
                col_end=10
            )
            
            ws_out.cell(row=dst_start, column=1).value = f"Data {data_serial}"
            data_serial += 1
            
            if rate_src_row:
                rate_pos[name] = dst_start + (rate_src_row - src_min)
            
            if is_repair:
                prefix = item_to_prefix.get(name, "")
                if prefix:
                    desc_cell = ws_out.cell(row=dst_start + 2, column=4)
                    base = desc_cell.value
                    base_str = str(base).strip() if base not in (None, "") else ""
                    desc_cell.value = f"{prefix} {base_str}" if base_str else prefix
            
            cursor += (src_max - src_min + 1)
        
        job.progress = 70
        job.current_step = "Building Estimate sheet..."
        job.save()
        
        # Create Estimate sheet
        ws_est = wb.create_sheet("Estimate")
        
        ws_est.merge_cells("A1:H1")
        c1 = ws_est["A1"]
        c1.value = "ESTIMATE"
        c1.font = Font(bold=True, size=14)
        c1.alignment = Alignment(horizontal="center", vertical="center")
        
        ws_est.merge_cells("A2:H2")
        c2 = ws_est["A2"]
        c2.value = f"Name of the work : {normalize_text(work_name)}" if work_name else "Name of the work : "
        c2.font = Font(bold=True, size=11)
        c2.alignment = Alignment(horizontal="left", vertical="center")
        
        for row in (1, 2):
            for col in range(1, 9):
                ws_est.cell(row=row, column=col).border = border_all
        
        headers = ["Sl.No", "Quantity (Unit)", "", "Item Description", "Rate", "Per Unit", "", "Amount"]
        for col, text in enumerate(headers, start=1):
            cell = ws_est.cell(row=3, column=col, value=text)
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = border_all
        
        ws_est.merge_cells("B3:C3")
        ws_est.merge_cells("F3:G3")
        
        ws_est.column_dimensions["A"].width = 6
        ws_est.column_dimensions["B"].width = 10
        ws_est.column_dimensions["C"].width = 10
        ws_est.column_dimensions["D"].width = 38
        ws_est.column_dimensions["E"].width = 10
        ws_est.column_dimensions["F"].width = 8
        ws_est.column_dimensions["G"].width = 10
        ws_est.column_dimensions["H"].width = 15
        
        # Fill estimate rows
        row_est = 4
        slno = 1
        
        for name in fetched:
            info = name_to_info.get(name)
            if not info:
                continue
            
            start = info["start_row"]
            base_desc = ws_src.cell(row=start + 2, column=4).value or ""
            base_desc_str = normalize_text(base_desc).strip()
            
            prefix = item_to_prefix.get(name, "") if is_repair else ""
            if prefix:
                desc = f"{prefix} {base_desc_str}" if base_desc_str else prefix
            else:
                desc = base_desc_str
            
            rr = rate_pos.get(name)
            rate_formula = f"=Output!J{rr}" if rr else ""
            
            # Priority for unit: 1) user-entered unit_map, 2) backend_units_map (Column D), 3) group-based default
            custom_unit = unit_map.get(name, "").strip()
            backend_unit = backend_units_map.get(name, "").strip() if backend_units_map else ""
            
            if custom_unit:
                # Use user-entered custom unit - derive singular form
                unit_plural = custom_unit
                unit_lower = custom_unit.lower()
            elif backend_unit:
                # Use unit from backend Column D - derive singular form
                unit_plural = backend_unit
                unit_lower = backend_unit.lower()
            else:
                unit_lower = ""
            
            if unit_lower:
                # Derive singular form from the unit
                if unit_lower in ("nos", "no"):
                    unit_singular = "No"
                elif unit_lower in ("mtrs", "mtr"):
                    unit_singular = "Mtr"
                elif unit_lower in ("pts", "pt"):
                    unit_singular = "Pt"
                elif unit_lower == "cum":
                    unit_singular = "Cum"
                elif unit_lower == "sqm":
                    unit_singular = "Sqm"
                elif unit_lower == "rmt":
                    unit_singular = "Rmt"
                elif unit_lower == "kg":
                    unit_singular = "Kg"
                elif unit_lower == "l":
                    unit_singular = "L"
                elif unit_lower == "set":
                    unit_singular = "Set"
                elif unit_lower == "lot":
                    unit_singular = "Lot"
                elif unit_lower == "ls":
                    unit_singular = "LS"
                else:
                    unit_singular = unit_plural
            else:
                # Use default unit based on group
                grp = (item_to_group.get(name, "") or "").lower()
                if grp in ("piping", "wiring & cables", "wiring and cables"):
                    unit_plural, unit_singular = "Mtrs", "Mtr"
                elif grp == "points":
                    unit_plural, unit_singular = "Pts", "Pt"
                else:
                    unit_plural, unit_singular = "Nos", "No"
            
            qty_val = qty_map.get(name)
            
            ws_est.cell(row=row_est, column=1, value=slno).alignment = Alignment(horizontal="center", vertical="top")
            ws_est.cell(row=row_est, column=1).border = border_all
            
            ws_est.cell(row=row_est, column=2, value=qty_val).alignment = Alignment(horizontal="center", vertical="top")
            ws_est.cell(row=row_est, column=2).border = border_all
            
            ws_est.cell(row=row_est, column=3, value=unit_plural).alignment = Alignment(horizontal="center", vertical="top")
            ws_est.cell(row=row_est, column=3).border = border_all
            
            ws_est.cell(row=row_est, column=4, value=desc).alignment = Alignment(horizontal="justify", vertical="top", wrap_text=True)
            ws_est.cell(row=row_est, column=4).border = border_all
            
            ws_est.cell(row=row_est, column=5, value=rate_formula).alignment = Alignment(horizontal="center", vertical="top")
            ws_est.cell(row=row_est, column=5).border = border_all
            
            ws_est.cell(row=row_est, column=6, value=1).alignment = Alignment(horizontal="center", vertical="top")
            ws_est.cell(row=row_est, column=6).border = border_all
            
            ws_est.cell(row=row_est, column=7, value=unit_singular).alignment = Alignment(horizontal="center", vertical="top")
            ws_est.cell(row=row_est, column=7).border = border_all
            
            ws_est.cell(row=row_est, column=8, value=f"=B{row_est}*E{row_est}").alignment = Alignment(horizontal="center", vertical="top")
            ws_est.cell(row=row_est, column=8).border = border_all
            
            row_est += 1
            slno += 1
        
        # Track current row for dynamic row placement
        current_row = row_est
        
        # Add Deduct Old Material Cost row if provided (for repair work, ABOVE ECV)
        deduct_row = None
        if deduct_old_material is not None and deduct_old_material > 0:
            deduct_row = current_row
            ws_est.cell(row=deduct_row, column=4, value="Deduct Old Material Cost")
            ws_est.cell(row=deduct_row, column=8, value=-deduct_old_material)  # Negative value for deduction
            current_row += 1
        
        # Add ECV row (finalized amount after deduction)
        ecv_row = current_row
        ws_est.cell(row=ecv_row, column=4, value="ECV")
        # ECV = SUM of items - Deduct Old Material (if any)
        if deduct_row:
            ws_est.cell(row=ecv_row, column=8, value=f"=SUM(H4:H{deduct_row-1})+H{deduct_row}")
        else:
            ws_est.cell(row=ecv_row, column=8, value=f"=SUM(H4:H{ecv_row-1})")
        
        # Add Excess T.P row if enabled (after ECV)
        excess_tp_row = None
        if excess_tp_percent is not None and excess_tp_percent > 0:
            current_row += 1
            excess_tp_row = current_row
            # Calculate Excess T.P based on finalized ECV
            ws_est.cell(row=excess_tp_row, column=4, value=f"Add Excess T.P @ {excess_tp_percent} %")
            ws_est.cell(row=excess_tp_row, column=8, value=f"=H{ecv_row}*{excess_tp_percent/100}")
        
        # LC, QC, NAC rows - calculations based on finalized ECV only
        lc_row = current_row + 1
        qc_row = current_row + 2
        nac_row = current_row + 3
        sub_row = current_row + 4
        gst_row = current_row + 5
        ls_row = current_row + 6
        
        # L.S Provision towards Special Items row (if enabled, comes before Grand Total)
        ls_special_row = None
        if ls_special_name and ls_special_amount is not None and ls_special_amount > 0:
            ls_special_row = current_row + 7
            gt_row = current_row + 8
        else:
            gt_row = current_row + 7
        
        # LC, QC, NAC are calculated based on ECV only (not including Excess T.P)
        ws_est.cell(row=lc_row, column=4, value="Add LC @ 1 %")
        ws_est.cell(row=lc_row, column=8, value=f"=H{ecv_row}*0.01")
        
        ws_est.cell(row=qc_row, column=4, value="Add QC @ 1 %")
        ws_est.cell(row=qc_row, column=8, value=f"=H{ecv_row}*0.01")
        
        ws_est.cell(row=nac_row, column=4, value="Add NAC @ 0.1 %")
        ws_est.cell(row=nac_row, column=8, value=f"=H{ecv_row}*0.001")
        
        ws_est.cell(row=sub_row, column=4, value="Sub Total")
        # Sub Total = ECV + Excess T.P (if any) + LC + QC + NAC
        sub_total_parts = [f"H{ecv_row}"]
        if excess_tp_row:
            sub_total_parts.append(f"H{excess_tp_row}")
        sub_total_parts.extend([f"H{lc_row}", f"H{qc_row}", f"H{nac_row}"])
        ws_est.cell(row=sub_row, column=8, value=f"={'+'.join(sub_total_parts)}")
        
        ws_est.cell(row=gst_row, column=4, value="Add GST@18 %")
        ws_est.cell(row=gst_row, column=8, value=f"=H{sub_row}*0.18")
        
        ws_est.cell(row=ls_row, column=4, value="L.S Provision towards unforeseen items")
        # Calculate L.S as difference between Grand Total and (Sub Total + GST + Special Items if any)
        if ls_special_row:
            ws_est.cell(row=ls_row, column=8, value=f"=H{gt_row}-H{gst_row}-H{sub_row}-H{ls_special_row}")
        else:
            ws_est.cell(row=ls_row, column=8, value=f"=H{gt_row}-H{gst_row}-H{sub_row}")
        
        # Add L.S Provision towards Special Items row (if enabled)
        if ls_special_row:
            ws_est.cell(row=ls_special_row, column=4, value=f"L.S Provision towards {ls_special_name}")
            ws_est.cell(row=ls_special_row, column=8, value=ls_special_amount)
        
        ws_est.cell(row=gt_row, column=4, value="Grand Total")
        # Set Grand Total value if provided by user, otherwise leave empty
        if grand_total is not None and grand_total > 0:
            ws_est.cell(row=gt_row, column=8, value=grand_total)
        
        for r in range(ecv_row, gt_row + 1):
            for c in range(1, 9):
                ws_est.cell(row=r, column=c).border = border_all
                if c == 4:
                    ws_est.cell(row=r, column=c).alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
                else:
                    ws_est.cell(row=r, column=c).alignment = Alignment(horizontal="center", vertical="top")
            ws_est.cell(row=r, column=4).font = Font(bold=True)
            ws_est.cell(row=r, column=8).font = Font(bold=True)
        
        job.progress = 85
        job.current_step = "Saving Excel file..."
        job.save()
        
        # Reorder sheets: Estimate first, then Output (ItemBlocks)
        if "Estimate" in wb.sheetnames:
            est_idx = wb.sheetnames.index("Estimate")
            if est_idx > 0:
                wb.move_sheet("Estimate", offset=-est_idx)

        # Apply print settings: Portrait, A4, fit columns, Times New Roman
        from core.views import _apply_print_settings
        _apply_print_settings(wb)

        # Save to file
        output_buffer = BytesIO()
        wb.save(output_buffer)
        output_buffer.seek(0)
        
        # Create OutputFile record
        output_file = OutputFile.objects.create(
            job=job,
            user=job.user,
            organization=job.organization,
            filename=f"{category}_output_estimate.xlsx",
            file_type="excel",
            file_size=len(output_buffer.getvalue()),
        )
        
        # Save file to storage
        output_file.file.save(
            f"{job.id}_{category}_output_estimate.xlsx",
            ContentFile(output_buffer.getvalue()),
        )
        
        job.progress = 100
        job.status = 'completed'
        job.completed_at = timezone.now()
        job.current_step = "Complete"
        job.result = {
            'output_file_id': output_file.id,
            'filename': output_file.filename,
        }
        job.save()
        
        logger.info(f"Generated output Excel for job {job_id}")
        return {'status': 'completed', 'output_file_id': output_file.id}
        
    except Exception as e:
        logger.error(f"Failed to generate output Excel: {e}")
        try:
            job = Job.objects.get(id=job_id)
            job.status = 'failed'
            job.error_message = str(e)
            job.completed_at = timezone.now()
            job.error_log.append({
                'timestamp': timezone.now().isoformat(),
                'step': 'generate_output_excel',
                'error': str(e),
                'traceback': traceback.format_exc(),
            })
            job.save()
        except:
            pass
        
        return {'status': 'failed', 'error': str(e)}


@shared_task(bind=True, max_retries=2)
def generate_estimate_excel(self, job_id, category, fetched_items_json, backend_id=None):
    """
    Generate Estimate-only Excel workbook asynchronously.
    
    This task handles the lighter workbook generation for download_estimate() view.
    
    Args:
        job_id: Job.id to track progress
        category: Category name for backend loading
        fetched_items_json: JSON string of fetched item names
        backend_id: ID of the ModuleBackend to use (optional, for multi-backend support)
    
    Returns:
        JSON with status and output_file_id
    """
    try:
        from django.conf import settings
        from django.utils import timezone
        from io import BytesIO
        from core.utils_excel import load_backend, build_estimate_wb
        
        job = Job.objects.get(id=job_id)
        job.status = 'running'
        job.started_at = timezone.now()
        job.progress = 10
        job.current_step = "Loading items..."
        job.save()
        
        fetched = json.loads(fetched_items_json) if fetched_items_json else []
        
        if not fetched:
            raise ValueError("No items to generate estimate from")
        
        job.progress = 30
        job.current_step = "Building estimate workbook..."
        job.save()
        
        # Get backend_id from job result if not passed directly (for backward compatibility)
        if not backend_id and job.result:
            backend_id = job.result.get('backend_id')
        
        # Determine module code - use 'amc' for AMC module, 'new_estimate' for New Estimate
        module_code = 'new_estimate'
        if job.result and job.result.get('module') == 'amc':
            module_code = 'amc'
        
        items_list, groups_map, _, ws_src, _ = load_backend(
            category, settings.BASE_DIR,
            backend_id=backend_id,
            module_code=module_code
        )
        name_to_block = {it["name"]: it for it in items_list}
        blocks = [name_to_block[n] for n in fetched if n in name_to_block]
        
        est_wb = build_estimate_wb(ws_src, blocks)
        
        job.progress = 80
        job.current_step = "Saving Excel file..."
        job.save()

        from core.views import _apply_print_settings
        _apply_print_settings(est_wb)

        # Save to file
        output_buffer = BytesIO()
        est_wb.save(output_buffer)
        output_buffer.seek(0)
        
        # Create OutputFile record
        output_file = OutputFile.objects.create(
            job=job,
            user=job.user,
            organization=job.organization,
            filename=f"{category}_estimate.xlsx",
            file_type="excel",
            file_size=len(output_buffer.getvalue()),
        )
        
        # Save file to storage
        output_file.file.save(
            f"{job.id}_{category}_estimate.xlsx",
            ContentFile(output_buffer.getvalue()),
        )
        
        job.progress = 100
        job.status = 'completed'
        job.completed_at = timezone.now()
        job.current_step = "Complete"
        job.result = {
            'output_file_id': output_file.id,
            'filename': output_file.filename,
        }
        job.save()
        
        logger.info(f"Generated estimate Excel for job {job_id}")
        return {'status': 'completed', 'output_file_id': output_file.id}
        
    except Exception as e:
        logger.error(f"Failed to generate estimate Excel: {e}")
        try:
            job = Job.objects.get(id=job_id)
            job.status = 'failed'
            job.error_message = str(e)
            job.completed_at = timezone.now()
            job.error_log.append({
                'timestamp': timezone.now().isoformat(),
                'step': 'generate_estimate_excel',
                'error': str(e),
                'traceback': traceback.format_exc(),
            })
            job.save()
        except:
            pass
        
        return {'status': 'failed', 'error': str(e)}


@shared_task
def cleanup_old_files(days=30):
    """
    Cleanup old output files (optional maintenance task).
    
    Args:
        days: Delete files older than this many days
    """
    from datetime import timedelta
    from django.utils import timezone
    
    cutoff_date = timezone.now() - timedelta(days=days)
    deleted_count, _ = OutputFile.objects.filter(
        created_at__lt=cutoff_date,
        download_count=0,  # Only delete unused files
    ).delete()
    
    logger.info(f"Cleaned up {deleted_count} old output files")
    return {'deleted_count': deleted_count}


@shared_task(bind=True, max_retries=3)
def generate_bill_document_task(self, job_id):
    """
    Generate bill documents (LS Forms, Covering Letter, Movement Slip).
    
    Reads metadata from job.result['metadata'] which contains:
      - doc_kind: 'ls_part', 'ls_final', 'covering', 'movement'
      - action, nth_number, mb details, etc.
      
    Args:
        job_id: ID of the Job object
        
    Returns:
        Dict with status and output_file_id
    """
    import os
    import io
    import re
    from io import BytesIO
    from copy import copy
    from django.conf import settings
    from django.utils import timezone
    from openpyxl import Workbook, load_workbook
    from openpyxl.styles import Alignment, Font, Border, Side, PatternFill
    from docx import Document
    
    BILL_TEMPLATES_DIR = os.path.join(settings.BASE_DIR, "core", "templates", "core", "bill_templates")
    
    try:
        job = Job.objects.get(id=job_id)
        upload = job.upload
        
        if not upload or not upload.file:
            raise ValueError("No upload file associated with this job")
        
        # Update job status
        job.status = 'running'
        job.current_step = "Reading bill file..."
        job.started_at = timezone.now()
        job.progress = 10
        job.save()
        
        # Get metadata from job
        metadata = job.result.get('metadata', {}) if job.result else {}
        doc_kind = metadata.get('doc_kind', '')
        action = metadata.get('action', '')
        nth_number_str = metadata.get('nth_number', '')
        mb_measure_no = metadata.get('mb_measure_no', '')
        mb_measure_p_from = metadata.get('mb_measure_p_from', '')
        mb_measure_p_to = metadata.get('mb_measure_p_to', '')
        mb_abs_no = metadata.get('mb_abs_no', '')
        mb_abs_p_from = metadata.get('mb_abs_p_from', '')
        mb_abs_p_to = metadata.get('mb_abs_p_to', '')
        
        logger.info(f"Generating {doc_kind} document for job {job_id}")
        
        # Load the uploaded Excel file
        try:
            wb_in = load_workbook(upload.file.path, data_only=True)
        except Exception as e:
            raise ValueError(f"Failed to read Excel file: {e}")
        
        job.progress = 20
        job.current_step = "Extracting header data..."
        job.save()
        
        # Helper functions (inline to avoid circular imports)
        def _number_to_words_rupees(n):
            """Convert number to rupees in words."""
            try:
                import inflect
                p = inflect.engine()
                
                n = float(n)
                rupees = int(n)
                paise = int(round((n - rupees) * 100))
                
                if rupees == 0 and paise == 0:
                    return "Zero Rupees Only"
                
                words_parts = []
                if rupees > 0:
                    rupee_words = p.number_to_words(rupees, andword='').title()
                    words_parts.append(f"Rupees {rupee_words}")
                
                if paise > 0:
                    paise_words = p.number_to_words(paise, andword='').title()
                    words_parts.append(f"Paise {paise_words}")
                
                return " and ".join(words_parts) + " Only"
            except Exception:
                return f"Rupees {n} Only"
        
        def _extract_header_data_fuzzy_from_wb(wb):
            """Extract header data from workbook."""
            import re as _re
            header = {
                "name_of_work": "",
                "estimate_amount": "",
                "admin_sanction": "",
                "tech_sanction": "",
                "agreement": "",
                "agency": "",
                "mb_details": "",
            }
            
            def clean_value(val):
                if not val:
                    return ""
                s = str(val).strip()
                if ":" in s:
                    s = s.split(":", 1)[1].strip()
                return s
            
            def _has_mb(low, tokens):
                if _re.search(r'm\.?b\.?\s*no', low):
                    return True
                if "measurement" in tokens and "book" in tokens:
                    return True
                if "mb" in tokens and ("details" in tokens or "no" in tokens or "nos" in tokens):
                    return True
                return False
            
            for ws in wb.worksheets:
                max_row = min(ws.max_row, 40)
                max_col = min(ws.max_column, 20)
                for r in range(1, max_row + 1):
                    for c in range(1, max_col + 1):
                        raw = ws.cell(row=r, column=c).value
                        if raw is None:
                            continue
                        s_full = str(raw).strip()
                        low = s_full.lower()
                        if not low:
                            continue
                        tokens = set(low.replace(":", " ").replace(".", " ").split())
                        
                        if not header["name_of_work"] and "name" in tokens and "work" in tokens:
                            header["name_of_work"] = clean_value(s_full)
                        if not header["agreement"] and ("agreement" in low or "agrmt" in low or "agt" in tokens):
                            header["agreement"] = clean_value(s_full)
                        if not header["agency"] and ("agency" in tokens or "contractor" in tokens or "firm" in tokens):
                            header["agency"] = clean_value(s_full)
                        if not header["mb_details"] and _has_mb(low, tokens):
                            header["mb_details"] = clean_value(s_full)
            
            return header
        
        def _build_mb_details_string(mb_no, mb_from, mb_to, abs_no, abs_from, abs_to):
            return (
                f"MB.No. {mb_no} P.No. {mb_from} to {mb_to} (Measurements) "
                f"& MB.No. {abs_no} P.No. {abs_from} to {abs_to} (Abstract)"
            )
        
        def _resolve_cc_header(action, nth_str=None):
            def ordinal_word(n):
                mapping = {1: "First", 2: "Second", 3: "Third", 4: "Fourth", 5: "Fifth",
                           6: "Sixth", 7: "Seventh", 8: "Eighth", 9: "Ninth", 10: "Tenth"}
                if n in mapping:
                    return mapping[n]
                if 10 < n % 100 < 14:
                    suffix = "th"
                else:
                    suffix = {1: "st", 2: "nd", 3: "rd"}.get(n % 10, "th")
                return f"{n}{suffix}"
            
            if action in ("estimate_first_part", "workslip_first_part"):
                return "CC First & Part Bill"
            if action in ("estimate_first_final", "workslip_first_final"):
                return "CC First & Final Bill"
            
            try:
                n_val = int(nth_str or "2")
            except:
                n_val = 2
            if n_val < 2:
                n_val = 2
            ord_word = ordinal_word(n_val)
            
            if action in ("firstpart_nth_part", "nth_nth_part"):
                return f"CC {ord_word} & Part Bill"
            if action in ("firstpart_2nd_final", "nth_nth_final"):
                return f"CC {ord_word} & Final Bill"
            
            return "CC Bill"
        
        def _extract_total_from_bill(wb):
            """Extract total amount from bill workbook."""
            total = 0.0
            for ws in wb.worksheets:
                max_scan = min(ws.max_row, 200)
                for r in range(1, max_scan + 1):
                    for check_col in [3, 4, 5]:
                        cell_val = str(ws.cell(row=r, column=check_col).value or "").strip().lower()
                        if cell_val == "total":
                            for amt_col in [8, 9, 10]:
                                amt_val = ws.cell(row=r, column=amt_col).value
                                try:
                                    num_val = float(amt_val) if amt_val else 0
                                    if num_val != 0:
                                        total += num_val
                                        break
                                except:
                                    continue
                            break
            return total
        
        # Extract data
        header = _extract_header_data_fuzzy_from_wb(wb_in)
        name_of_work = header.get("name_of_work", "")
        agreement_ref = header.get("agreement", "")
        agency_name = header.get("agency", "")
        file_mb_details = (header.get("mb_details") or "").strip()
        
        # Use user-entered MB if provided, else use file-extracted MB
        user_entered_mb = any([
            mb_measure_no, mb_measure_p_from, mb_measure_p_to,
            mb_abs_no, mb_abs_p_from, mb_abs_p_to,
        ])
        if user_entered_mb:
            mb_details_str = _build_mb_details_string(
                mb_measure_no, mb_measure_p_from, mb_measure_p_to,
                mb_abs_no, mb_abs_p_from, mb_abs_p_to
            )
        elif file_mb_details:
            mb_details_str = file_mb_details
        else:
            mb_details_str = _build_mb_details_string(
                mb_measure_no, mb_measure_p_from, mb_measure_p_to,
                mb_abs_no, mb_abs_p_from, mb_abs_p_to
            )
        cc_header = _resolve_cc_header(action, nth_number_str)
        total_amount = _extract_total_from_bill(wb_in)
        total_amount_str = f"{float(total_amount):,.2f}"
        amount_in_words = _number_to_words_rupees(total_amount)
        
        job.progress = 40
        job.current_step = f"Generating {doc_kind} document..."
        job.save()
        
        # Context for templates
        ctx = {
            "NAME_OF_WORK": name_of_work,
            "NAME_OF_AGENCY": agency_name,
            "AGENCY_NAME": agency_name,
            "REF_OF_AGREEMENT": agreement_ref,
            "AGREEMENT_REF": agreement_ref,
            "MB_DETAILS": mb_details_str,
            "CC_HEADER": cc_header,
            "AMOUNT": total_amount_str,
            "TOTAL_AMOUNT": total_amount_str,
            "AMOUNT_IN_WORDS": amount_in_words,
        }
        
        now = timezone.now()
        mm_yyyy = f"{now.month:02d}.{now.year}"
        
        output_buffer = BytesIO()
        
        # Generate document based on doc_kind
        if doc_kind in ("ls_part", "ls_final"):
            template_name = "LS_Form_Part.xlsx" if doc_kind == "ls_part" else "LS_Form_Final.xlsx"
            template_path = os.path.join(BILL_TEMPLATES_DIR, template_name)
            
            if not os.path.exists(template_path):
                raise FileNotFoundError(f"Template not found: {template_name}")
            
            wb_out = load_workbook(template_path)
            
            # Replace placeholders
            for ws in wb_out.worksheets:
                for row in ws.iter_rows():
                    for cell in row:
                        if isinstance(cell.value, str):
                            text = cell.value
                            for key, val in ctx.items():
                                placeholder = "{{" + key + "}}"
                                if placeholder in text:
                                    text = text.replace(placeholder, str(val or ""))
                            cell.value = text
            
            wb_out.save(output_buffer)
            file_ext = "xlsx"
            content_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            download_name = f"LS_Form_{'Part' if doc_kind == 'ls_part' else 'Final'}.xlsx"
            
        elif doc_kind in ("covering", "movement"):
            template_name = "covering_letter.docx" if doc_kind == "covering" else "Movement_Slip.docx"
            template_path = os.path.join(BILL_TEMPLATES_DIR, template_name)
            
            if not os.path.exists(template_path):
                raise FileNotFoundError(f"Template not found: {template_name}")
            
            doc = Document(template_path)
            
            placeholder_map = {
                "{{NAME_OF_WORK}}": name_of_work,
                "{{AGENCY_NAME}}": agency_name,
                "{{NAME_OF_AGENCY}}": agency_name,
                "{{AGREEMENT_REF}}": agreement_ref,
                "{{REF_OF_AGREEMENT}}": agreement_ref,
                "{{CC_HEADER}}": cc_header,
                "{{MB_DETAILS}}": mb_details_str,
                "{{AMOUNT}}": total_amount_str,
                "{{TOTAL_AMOUNT}}": total_amount_str,
                "{{AMOUNT_IN_WORDS}}": amount_in_words,
            }
            
            def replace_in_paragraphs(paragraphs):
                for p in paragraphs:
                    if not p.runs:
                        continue
                    for run in p.runs:
                        if not run.text:
                            continue
                        text = run.text
                        changed = False
                        for ph, val in placeholder_map.items():
                            if ph in text:
                                text = text.replace(ph, val or "")
                                changed = True
                        if "dd.mm.yyyy" in text:
                            text = text.replace("dd.mm.yyyy", f"dd.{mm_yyyy}")
                            changed = True
                        if changed:
                            run.text = text
            
            replace_in_paragraphs(doc.paragraphs)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        replace_in_paragraphs(cell.paragraphs)
            
            doc.save(output_buffer)
            file_ext = "docx"
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            download_name = f"{'Covering_Letter' if doc_kind == 'covering' else 'Movement_Slip'}.docx"
        else:
            raise ValueError(f"Unknown doc_kind: {doc_kind}")
        
        output_buffer.seek(0)
        
        job.progress = 80
        job.current_step = "Saving output file..."
        job.save()
        
        # Create OutputFile
        output_file = OutputFile.objects.create(
            job=job,
            user=job.user,
            organization=job.organization,
            filename=download_name,
            file_type=file_ext,
            file_size=len(output_buffer.getvalue()),
        )
        
        output_file.file.save(
            f"{job.id}_{download_name}",
            ContentFile(output_buffer.getvalue()),
        )
        
        # Mark job as completed
        job.status = 'completed'
        job.progress = 100
        job.current_step = "Complete"
        job.completed_at = timezone.now()
        job.result = {
            'output_file_id': output_file.id,
            'filename': download_name,
            'download_url': f"/api/output/{output_file.id}/download/",
        }
        job.save()
        
        logger.info(f"Generated {doc_kind} document for job {job_id}")
        return {'status': 'completed', 'output_file_id': output_file.id}
        
    except Exception as e:
        logger.error(f"Failed to generate bill document for job {job_id}: {e}\n{traceback.format_exc()}")
        try:
            job = Job.objects.get(id=job_id)
            job.status = 'failed'
            job.error_message = str(e)
            job.completed_at = timezone.now()
            job.error_log.append({
                'timestamp': timezone.now().isoformat(),
                'step': 'generate_bill_document',
                'error': str(e),
                'traceback': traceback.format_exc(),
            })
            job.save()
        except:
            pass
        
        return {'status': 'failed', 'error': str(e)}
