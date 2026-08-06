"""
Tests for filling a saved (reusable) self-formatted format from several source
files at once: one sheet per work for Excel templates, one page per work for
Word templates, or a ZIP of separate documents.
"""

import io
import zipfile

import pytest
from docx import Document
from openpyxl import Workbook, load_workbook

from core.views.self_formatted_views import (
    _build_multi_docx,
    _extract_works_with_lines,
    _fill_template_multi,
    _safe_sheet_title,
)


# ---------------------------------------------------------------- fixtures --

def _xlsx_template_bytes(placeholder_rows, sheet_titles=("Bill",)):
    """Workbook template with {{PLACEHOLDER}} markers on each sheet."""
    wb = Workbook()
    wb.remove(wb.active)
    for title in sheet_titles:
        ws = wb.create_sheet(title=title)
        for r, (label, placeholder) in enumerate(placeholder_rows, start=1):
            ws.cell(row=r, column=1, value=label)
            ws.cell(row=r, column=2, value=placeholder)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _docx_template_bytes(lines):
    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _works(*names):
    return [
        (name, {"{{NAME_OF_WORK}}": name, "{{AMOUNT}}": f"{i + 1}000"})
        for i, name in enumerate(names)
    ]


def _source_workbook(sheets):
    """Source workbook: {sheet_title: work_name} — one work per sheet."""
    wb = Workbook()
    wb.remove(wb.active)
    for title, work_name in sheets.items():
        ws = wb.create_sheet(title=title)
        ws.cell(row=1, column=1, value="Name of the work")
        ws.cell(row=1, column=2, value=work_name)
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


# ------------------------------------------------------------ excel output --

def test_excel_template_gets_one_sheet_per_work():
    template = _xlsx_template_bytes(
        [("Name of the work", "{{NAME_OF_WORK}}"), ("Amount", "{{AMOUNT}}")]
    )
    works = _works("Road Repair", "Drain Work", "Street Lighting")

    resp = _fill_template_multi(template, "bill.xlsx", works)

    assert resp.status_code == 200
    wb = load_workbook(io.BytesIO(resp.content))
    assert len(wb.worksheets) == 3
    assert wb.sheetnames == ["Road Repair", "Drain Work", "Street Lighting"]
    for ws, (name, _map) in zip(wb.worksheets, works):
        assert ws.cell(row=1, column=2).value == name
    assert wb["Drain Work"].cell(row=2, column=2).value == "2000"


def test_multi_sheet_excel_template_repeats_the_whole_set_per_work():
    template = _xlsx_template_bytes(
        [("Name of the work", "{{NAME_OF_WORK}}")],
        sheet_titles=("Summary", "Detail"),
    )

    resp = _fill_template_multi(template, "bill.xlsx", _works("Work A", "Work B"))

    wb = load_workbook(io.BytesIO(resp.content))
    assert len(wb.worksheets) == 4
    assert wb.sheetnames == ["Summary (1)", "Detail (1)", "Summary (2)", "Detail (2)"]
    assert wb["Summary (1)"].cell(row=1, column=2).value == "Work A"
    assert wb["Detail (2)"].cell(row=1, column=2).value == "Work B"


def test_single_work_keeps_the_original_single_file_output():
    template = _xlsx_template_bytes([("Name of the work", "{{NAME_OF_WORK}}")])

    resp = _fill_template_multi(template, "bill.xlsx", _works("Only Work"))

    assert 'filename="Filled_bill.xlsx"' in resp["Content-Disposition"]
    wb = load_workbook(io.BytesIO(resp.content))
    assert len(wb.worksheets) == 1
    assert wb.sheetnames == ["Bill"]  # template's own sheet name is kept
    assert wb.worksheets[0].cell(row=1, column=2).value == "Only Work"


def test_sheet_titles_are_sanitised_and_unique():
    used = set()
    assert _safe_sheet_title("A/B:C", used) == "A-B-C"
    assert _safe_sheet_title("A/B:C", used) == "A-B-C (2)"
    assert len(_safe_sheet_title("x" * 60, used)) == 31


# ------------------------------------------------------------- word output --

def test_word_template_gets_one_page_per_work():
    template = _docx_template_bytes(["Work: {{NAME_OF_WORK}}", "Amount: {{AMOUNT}}"])
    works = _works("Road Repair", "Drain Work")

    resp = _fill_template_multi(template, "letter.docx", works)

    doc = Document(io.BytesIO(resp.content))
    texts = [p.text for p in doc.paragraphs]
    assert "Work: Road Repair" in texts
    assert "Work: Drain Work" in texts
    assert "{{NAME_OF_WORK}}" not in "".join(texts)
    # A page break separates the two works
    xml = doc.element.body.xml
    assert xml.count('w:type="page"') == 1


def test_merged_docx_keeps_a_single_body_section():
    template = _docx_template_bytes(["Work: {{NAME_OF_WORK}}"])

    data = _build_multi_docx(template, _works("A", "B", "C"))

    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        doc_xml = zf.read("word/document.xml").decode("utf-8")
    assert doc_xml.count("<w:body") == 1
    assert doc_xml.count("<w:sectPr") == 1
    assert doc_xml.index("<w:sectPr") > doc_xml.rindex("Work: C")


def test_merged_docx_repeats_tables_from_the_template():
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Name of work"
    table.cell(0, 1).text = "{{NAME_OF_WORK}}"
    buf = io.BytesIO()
    doc.save(buf)

    resp = _fill_template_multi(buf.getvalue(), "letter.docx", _works("Work A", "Work B"))

    out = Document(io.BytesIO(resp.content))
    assert len(out.tables) == 2
    assert out.tables[0].cell(0, 1).text == "Work A"
    assert out.tables[1].cell(0, 1).text == "Work B"


def test_excel_template_without_markers_fills_each_work_by_label():
    # No {{PLACEHOLDER}} markers: values are matched by the label text instead
    template = _xlsx_template_bytes([("Name of the work :", None)])

    resp = _fill_template_multi(template, "bill.xlsx", _works("Road Repair", "Drain Work"))

    wb = load_workbook(io.BytesIO(resp.content))
    assert wb["Road Repair"].cell(row=1, column=2).value == "Road Repair"
    assert wb["Drain Work"].cell(row=1, column=2).value == "Drain Work"


# -------------------------------------------------------------- zip output --

def test_separate_files_mode_returns_a_zip_per_work():
    template = _xlsx_template_bytes([("Name of the work", "{{NAME_OF_WORK}}")])

    resp = _fill_template_multi(
        template, "bill.xlsx", _works("Road Repair", "Drain Work"), separate_files=True
    )

    assert resp["Content-Type"] == "application/zip"
    with zipfile.ZipFile(io.BytesIO(resp.content)) as zf:
        names = zf.namelist()
        assert names == ["01_Road Repair.xlsx", "02_Drain Work.xlsx"]
        wb = load_workbook(io.BytesIO(zf.read(names[1])))
        assert wb.worksheets[0].cell(row=1, column=2).value == "Drain Work"


def test_unsupported_template_type_is_rejected():
    resp = _fill_template_multi(b"data", "template.pdf", _works("A", "B"))
    assert resp.status_code == 400


# ------------------------------------------------------- source extraction --

def test_each_sheet_of_a_source_workbook_is_its_own_work():
    from django.core.files.uploadedfile import SimpleUploadedFile

    buf = _source_workbook({"Bill 1": "Road Repair", "Bill 2": "Drain Work"})
    upload = SimpleUploadedFile("bills.xlsx", buf.getvalue())

    works = _extract_works_with_lines(upload)

    assert len(works) == 2
    names = [labels.get("name_of_work") for _src, labels, _lines in works]
    assert names == ["Road Repair", "Drain Work"]
    assert all(lines for _src, _labels, lines in works)


# ------------------------------------------------------------- view level ---

def _saved_format(name, template_filename, template_bytes):
    """A saved reusable format owned by a fresh user/org."""
    from django.contrib.auth.models import User
    from django.core.files.uploadedfile import SimpleUploadedFile

    from core.models import Membership, Organization, SelfFormattedTemplate

    username = f"sfuser-{name}".lower().replace(" ", "-")
    user = User.objects.create_user(username, password="x")
    # A signal already gives every new user an organization + membership
    membership = Membership.objects.filter(user=user).select_related("organization").first()
    if membership:
        org = membership.organization
    else:
        org = Organization.objects.create(name=f"Org {name}", slug=f"{username}-org", owner=user)
        Membership.objects.create(user=user, organization=org, role="owner")

    fmt = SelfFormattedTemplate.objects.create(
        organization=org,
        user=user,
        name=name,
        template_file=SimpleUploadedFile(template_filename, template_bytes),
        custom_placeholders="",
    )
    return fmt, user, org


@pytest.mark.django_db
def test_use_format_view_fills_one_sheet_per_uploaded_file(rf):
    """The view is called directly so the subscription middleware, which these
    endpoints sit behind in production, is not part of what is asserted."""
    from django.core.files.uploadedfile import SimpleUploadedFile

    from core.views.self_formatted_form_views import self_formatted_use_format

    fmt, user, org = _saved_format(
        "Bill Format",
        "bill.xlsx",
        _xlsx_template_bytes([("Name of the work", "{{NAME_OF_WORK}}")]),
    )

    request = rf.post(
        f"/self-formatted/use/{fmt.pk}/",
        {
            "source_files": [
                SimpleUploadedFile("a.xlsx", _source_workbook({"S1": "Road Repair"}).getvalue()),
                SimpleUploadedFile("b.xlsx", _source_workbook({"S1": "Drain Work"}).getvalue()),
            ],
            "split_sheets": "1",
            "output_mode": "combined",
        },
    )
    request.user = user
    request.organization = org

    resp = self_formatted_use_format(request, fmt.pk)

    assert resp.status_code == 200
    assert resp["X-Works-Generated"] == "2"
    wb = load_workbook(io.BytesIO(resp.content))
    assert len(wb.worksheets) == 2
    values = [ws.cell(row=1, column=2).value for ws in wb.worksheets]
    assert values == ["Road Repair", "Drain Work"]


@pytest.mark.django_db
def test_use_format_view_fills_one_page_per_uploaded_file(rf):
    """Word format + three uploaded files -> one document, one filled page each."""
    from django.core.files.uploadedfile import SimpleUploadedFile

    from core.views.self_formatted_form_views import self_formatted_use_format

    fmt, user, org = _saved_format(
        "Letter Format",
        "letter.docx",
        _docx_template_bytes(["Name of work: {{NAME_OF_WORK}}", "Agency: {{AGENCY_NAME}}"]),
    )

    names = ["Road Repair", "Drain Work", "Street Lighting"]
    request = rf.post(
        f"/self-formatted/use/{fmt.pk}/",
        {
            "source_files": [
                SimpleUploadedFile(f"{n}.xlsx", _source_workbook({"S1": n}).getvalue())
                for n in names
            ],
            "split_sheets": "1",
            "output_mode": "combined",
        },
    )
    request.user = user
    request.organization = org

    resp = self_formatted_use_format(request, fmt.pk)

    assert resp.status_code == 200
    assert resp["X-Works-Generated"] == "3"
    assert resp["Content-Type"].endswith("wordprocessingml.document")

    doc = Document(io.BytesIO(resp.content))
    texts = [p.text for p in doc.paragraphs]
    for n in names:
        assert f"Name of work: {n}" in texts
    # Three pages: a break between each consecutive pair
    assert doc.element.body.xml.count('w:type="page"') == 2


@pytest.mark.django_db
def test_use_format_view_can_return_separate_files_as_zip(rf):
    from django.core.files.uploadedfile import SimpleUploadedFile

    from core.views.self_formatted_form_views import self_formatted_use_format

    fmt, user, org = _saved_format(
        "Zip Format",
        "letter.docx",
        _docx_template_bytes(["Name of work: {{NAME_OF_WORK}}"]),
    )

    request = rf.post(
        f"/self-formatted/use/{fmt.pk}/",
        {
            "source_files": [
                SimpleUploadedFile("a.xlsx", _source_workbook({"S1": "Road Repair"}).getvalue()),
                SimpleUploadedFile("b.xlsx", _source_workbook({"S1": "Drain Work"}).getvalue()),
            ],
            "output_mode": "separate",
        },
    )
    request.user = user
    request.organization = org

    resp = self_formatted_use_format(request, fmt.pk)

    assert resp["Content-Type"] == "application/zip"
    with zipfile.ZipFile(io.BytesIO(resp.content)) as zf:
        assert len(zf.namelist()) == 2
        first = Document(io.BytesIO(zf.read(zf.namelist()[0])))
        assert "Name of work: Road Repair" in [p.text for p in first.paragraphs]
