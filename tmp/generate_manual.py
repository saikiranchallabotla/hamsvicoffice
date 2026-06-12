"""Generate Hamsvic Office User Manual as a Word document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date
import os

OUT_PATH = os.path.join(os.path.dirname(__file__), "Hamsvic_User_Manual.docx")

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def add_heading(text, level=1, color=(31, 73, 125)):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*color)
    return h


def add_para(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    return p


def add_numbered(text):
    return doc.add_paragraph(text, style='List Number')


def add_callout(title, body, fill="EAF1FB"):
    """A shaded one-cell table used as a tip/note box."""
    t = doc.add_table(rows=1, cols=1)
    t.autofit = True
    cell = t.rows[0].cells[0]
    shade_cell(cell, fill)
    p1 = cell.paragraphs[0]
    r = p1.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 73, 125)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    return t


def page_break():
    doc.add_page_break()


def kv_table(rows, col1="Field / Button", col2="What it does"):
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = col1
    hdr[1].text = col2
    for c in hdr:
        shade_cell(c, "1F497D")
        for run in c.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    for name, desc in rows:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = desc
    doc.add_paragraph()


# ======================== COVER ========================
for _ in range(4):
    doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run("HAMSVIC OFFICE")
r.bold = True
r.font.size = Pt(36)
r.font.color.rgb = RGBColor(31, 73, 125)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("User Manual")
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = RGBColor(80, 80, 80)

tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tag.add_run("How every module works — Estimate, Workslip, Bill, AMC, Temporary Works, Self-Formatted Forms")
r.italic = True
r.font.size = Pt(13)

for _ in range(8):
    doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run(f"Website: hamsvic.com\nVersion 1.0  |  {date.today().strftime('%B %Y')}")
r.font.size = Pt(12)

page_break()

# ======================== TOC ========================
add_heading("Table of Contents", level=1)
toc = [
    "1. Introduction",
    "2. Getting Started (Login & Dashboard)",
    "3. Navigation Overview",
    "4. The 3-Panel Item Picker (used by Estimate, Workslip, AMC, Temp Works)",
    "5. Drag-and-Drop Reordering",
    "6. Estimate Module",
    "7. Workslip Module",
    "8. Bill Module",
    "9. AMC Module",
    "10. Temporary Works Module",
    "11. Self-Formatted Forms Module",
    "12. My Saved Works (Folders, Projects, E → W → B Flow)",
    "13. My Custom Items",
    "14. Document Templates & Letter Settings",
    "15. Exports & Downloads",
    "16. Account, Notifications, Payments",
    "17. Tips, Shortcuts, Best Practices",
    "18. FAQ & Support",
    "Appendix A. Supplement Datas Sheet Format (for Workslip uploads)",
]
for t in toc:
    p = doc.add_paragraph(t)
    p.paragraph_format.left_indent = Inches(0.3)
page_break()

# ======================== 1. INTRODUCTION ========================
add_heading("1. Introduction", level=1)
add_para(
    "Hamsvic Office is a web platform for preparing project documentation — Estimates, "
    "Workslips, Bills, AMC contracts, Temporary Works estimates, and custom Self-"
    "Formatted forms. The system stores your work, lets you build documents from a "
    "Schedule of Rates (SOR) catalogue, and exports professional Excel and PDF outputs."
)
add_para(
    "The work usually flows in this order: you start with an Estimate (E), generate one "
    "or more Workslips (W1, W2, W3 …) from it, and raise Bills (B1, B2, B3 …) against "
    "those Workslips. Saved Works keeps the whole chain organised."
)
add_callout("Key idea",
    "Every project follows the chain  E  →  W1, W2, …  →  B1, B2, …  "
    "Each link is generated from the previous one, so numbers and items stay consistent.")

page_break()

# ======================== 2. GETTING STARTED ========================
add_heading("2. Getting Started", level=1)
add_heading("2.1 Logging In", level=2)
add_numbered("Open https://hamsvic.com in your browser.")
add_numbered("Enter your registered email and password and click Login.")
add_numbered("You land on the Dashboard.")

add_heading("2.2 The Dashboard", level=2)
add_para("Dashboard tiles give one-click access to every module:", bold=True)
for b in [
    "Estimate, Workslip, Bill, AMC, Temporary Works, Self-Formatted Forms tiles.",
    "Stats cards: total Estimates, Workslips and Bills you have created.",
    "Sidebar links: My Saved Works, My Custom Items, Settings, Payment History, Help Center.",
    "Bell icon (top-right) for announcements; account avatar for profile and logout.",
]:
    add_bullet(b)

page_break()

# ======================== 3. NAVIGATION ========================
add_heading("3. Navigation Overview", level=1)
add_heading("3.1 Sidebar Menu", level=2)
kv_table([
    ("Dashboard", "Home page with module tiles and stats."),
    ("My Saved Works", "Every Estimate, Workslip and Bill, organised in folders."),
    ("My Custom Items", "Items you create yourself, in addition to the standard SOR catalogue."),
    ("Settings", "Profile, password, letter settings and document templates."),
    ("Payment History", "List of past subscription payments."),
    ("Help Center", "Support documentation and contact details."),
], col1="Menu Item")

add_heading("3.2 Top Bar", level=2)
for b in [
    "Page title with breadcrumbs showing where you are.",
    "Bell icon — opens announcements dropdown.",
    "Account avatar — opens profile menu and Logout.",
]:
    add_bullet(b)

page_break()

# ======================== 4. 3-PANEL ITEM PICKER ========================
add_heading("4. The 3-Panel Item Picker", level=1)
add_para(
    "Estimate, Workslip, AMC and Temporary Works all share the same item-selection "
    "screen. Learn it once and you can use it everywhere."
)

add_heading("4.1 Layout — three columns", level=2)
kv_table([
    ("Left panel — Groups",
     "Top-level categories of items (e.g. Wiring, Switchgear, Masonry). Has a "
     "‘Search groups…’ box at the top. Click a group to load its items in the middle."),
    ("Middle panel — Items",
     "Items that belong to the selected group. Has a ‘Search items…’ box. Each item "
     "has a checkbox; some items have subtypes and show a ‘X types’ badge — clicking "
     "those opens a small modal where you choose the subtype."),
    ("Right panel — Selected Items table",
     "The items you have added. Each row has Item name, Unit dropdown, Quantity input, "
     "(in Repair) Deduct Old Material input, and a red delete button."),
], col1="Panel")

add_heading("4.2 How to add an item", level=2)
add_numbered("Click a group in the left panel. The middle panel updates with that group’s items.")
add_numbered("Type in ‘Search items…’ if the group is long. Results filter as you type.")
add_numbered("Tick the checkbox next to the item. It instantly appears as a row in the right panel with quantity 0.")
add_numbered("If the item has subtypes, a small modal opens — click the subtype you want.")
add_numbered("Type the quantity in the Qty field of the right-panel row. The amount and grand total update live.")

add_heading("4.3 How to remove an item", level=2)
for b in [
    "Click the red delete (X) button at the end of the row in the right panel.",
    "Or untick the checkbox in the middle panel — the row disappears.",
    "Use the Clear All button (red) at the bottom of the page to remove every item at once.",
]:
    add_bullet(b)

add_heading("4.4 Searching", level=2)
for b in [
    "Left ‘Search groups…’ — filters the groups list.",
    "Middle ‘Search items…’ — filters items within the selected group, by name / code / description.",
    "Clearing a search box restores the full list.",
]:
    add_bullet(b)

add_heading("4.5 Quantity, Units and Live Totals", level=2)
for b in [
    "The Unit column shows the SOR unit; some items have a dropdown if multiple units are allowed.",
    "Qty is a number input — you can also use the up/down arrows.",
    "On every keystroke the Amount column for that row and the Grand Total below update.",
    "In Repair work type a ‘Deduct Old Material’ field appears — type the deduction amount per item.",
]:
    add_bullet(b)

add_callout("Tip",
    "Use the search boxes — they are the fastest way to find items in a long catalogue. "
    "You don’t need to scroll through every group manually.")

page_break()

# ======================== 5. DRAG AND DROP ========================
add_heading("5. Drag-and-Drop Reordering", level=1)
add_para(
    "The right-panel table (the Selected Items list) is fully reorderable by drag-and-"
    "drop. The order you set here is the order that appears in the downloaded Excel, PDF "
    "and printed bill."
)

add_heading("5.1 How to drag a row", level=2)
add_numbered("Move your mouse over the row you want to move.")
add_numbered("Press and hold the LEFT mouse button anywhere on the row.")
add_numbered("Drag the row up or down. The row you are dragging becomes semi-transparent. The row you are hovering over highlights as the drop target.")
add_numbered("Release the mouse button when the row is in the correct position. The list reshuffles instantly.")
add_numbered("The new order is saved automatically — no extra Save click is needed for ordering.")

add_heading("5.2 Where reordering works", level=2)
for b in [
    "Estimate — items.html selected items table.",
    "Workslip — workslip.html selected items table.",
    "Temporary Works — temp_items.html selected items table.",
    "AMC — amc_items.html selected items table.",
]:
    add_bullet(b)

add_callout("Note",
    "On a touchscreen, press and hold a row for about half a second, then drag with "
    "your finger. The behaviour is the same as the mouse.")

page_break()

# ======================== 6. ESTIMATE ========================
add_heading("6. Estimate Module", level=1)
add_para(
    "An Estimate is the cost-budget document for a project, built from the SOR "
    "catalogue. It is usually the FIRST document you create — every Workslip and Bill "
    "later flows from this Estimate."
)

add_heading("6.1 Step 1 — Choose Work Type", level=2)
add_para("From the Dashboard, click the Estimate tile. You see two cards:", bold=True)
for b in [
    "Original Work (blue, lightning-bolt icon) — for new construction or fresh installation.",
    "Repair Work (gray, tools icon) — for maintenance or repair jobs. Items get an automatic ‘Repair to …’ description prefix where applicable, and a ‘Deduct Old Material’ column appears.",
]:
    add_bullet(b)

add_heading("6.2 Step 2 — Choose Category", level=2)
for b in [
    "Electrical (amber icon).",
    "Civil (green icon).",
    "The system loads the correct backend Excel (Master Datas / Groups) for that category.",
]:
    add_bullet(b)

add_heading("6.3 Step 3 — Build the Estimate (Items page)", level=2)
add_para(
    "You arrive at the 3-panel item picker (see Chapter 4). Pick groups → tick items → "
    "type quantities → drag to reorder. The grand total updates live at the bottom."
)

add_heading("6.4 Action buttons at the bottom of the Items page", level=2)
kv_table([
    ("Download Estimate (green)",
     "Generates the formatted Excel estimate with item list, Quantity × Rate, totals."),
    ("Specification Report (green gradient)",
     "Generates a Word/PDF document of technical specifications for each selected item."),
    ("Forwarding Letter (blue gradient)",
     "Generates a covering letter using your Letter Settings (Chapter 14)."),
    ("Clear All (red)",
     "Removes every selected item. A confirmation appears before clearing."),
    ("Save Work (purple)",
     "Opens a modal: type a name for the Estimate, choose a folder, and click Save. "
     "The Estimate is then visible in My Saved Works."),
], col1="Button")

add_callout("Tip",
    "Always click Save Work before downloading. Saved Estimates can be reopened, edited, "
    "and used to generate Workslips and Bills later.")

page_break()

# ======================== 7. WORKSLIP ========================
add_heading("7. Workslip Module", level=1)
add_para(
    "A Workslip is the on-site work record. Each Estimate can have multiple Workslips "
    "(numbered Workslip-1, Workslip-2 … up to Workslip-10)."
)

add_heading("7.1 Two ways to start a Workslip", level=2)
for b in [
    "From a saved Estimate in My Saved Works → click the +W1 button on the workflow bar (E → W → B). This is the recommended path — it links the Workslip to the Estimate.",
    "From the Dashboard → click the Workslip tile (used when starting fresh).",
]:
    add_bullet(b)

add_heading("7.2 The Workslip page", level=2)
add_para("Same 3-panel item picker as Estimate, plus these workslip-specific controls:", bold=True)
kv_table([
    ("Target Workslip dropdown",
     "Choose which workslip number you are filling: ‘Workslip-1 (First Workslip)’, "
     "‘Workslip-2’, etc. The system shows whether each is empty, in-progress or completed."),
    ("Upload Workslip Data (file input)",
     "Optional. Upload an .xlsx / .xlsm / .xls of an existing workslip — the system "
     "parses items and quantities and pre-fills the table. Click ‘Upload Workslip Data’ "
     "(green button) to apply."),
    ("Add Supplemental Items",
     "Below the items table. Type an amount and click ‘Add Supplemental Items’ (green) "
     "to add an extra row. Each supplemental row has its own delete (X) button."),
    ("Download Workslip (green)",
     "Exports the Workslip Excel — formatted for site use."),
    ("Save Work",
     "Saves the Workslip under its parent Estimate in My Saved Works."),
], col1="Control")

add_heading("7.3 Generating the next Workslip", level=2)
add_para(
    "Once a Workslip is completed, open its parent Estimate in My Saved Works. The "
    "workflow bar shows the next +W button (e.g. +W2). Click it to start the next "
    "Workslip — items from the parent Estimate are pre-loaded so you only adjust quantities."
)

add_heading("7.4 Workslip 2+ behaviour", level=2)
for b in [
    "Supplemental items appear by name only (concise format).",
    "Repair prefix is applied automatically to descriptions where applicable.",
    "Quantities can override defaults from the parent Estimate.",
]:
    add_bullet(b)

page_break()

# ======================== 8. BILL ========================
add_heading("8. Bill Module", level=1)
add_para(
    "Bills are raised against Workslips. A single Workslip can have several Bills "
    "(B1, B2, B3 …). The Bill page is a structured form, not the 3-panel item picker — "
    "items are inherited from the parent Workslip and you only fill in measurement "
    "and quantity-till-date data."
)

add_heading("8.1 Starting a Bill", level=2)
add_numbered("Open the parent Workslip in My Saved Works.")
add_numbered("On the workflow bar (E → W → B), click +B1 (or +BN for the next bill).")
add_numbered("The Bill Entry page opens.")

add_heading("8.2 Header", level=2)
for b in [
    "Title shows ‘Bill N’ with a small badge ‘BN’.",
    "Subtitle shows the lineage, e.g. ‘First Bill from Workslip-1’ or ‘Bill 2 (linked to B1 & W1)’.",
]:
    add_bullet(b)

add_heading("8.3 Form Sections", level=2)

add_para("Measurement Book Details (blue header):", bold=True)
for b in [
    "M.B. No., From Page, To Page — for measurements.",
    "M.B. No., From Page, To Page — for abstract.",
]:
    add_bullet(b)

add_para("Important Dates (purple header):", bold=True)
for b in [
    "Date of Inspection.",
    "Date of Completion.",
    "Date of MR (Measurement Recording).",
]:
    add_bullet(b)

add_para("Quantities & Amounts table:", bold=True)
kv_table([
    ("Item / Unit / Rate", "Pre-filled from the parent Workslip — read-only."),
    ("W Qty", "Workslip quantity — read-only reference."),
    ("Deduct Previous", "Quantity already billed in Bill (N-1). Read-only — auto-calculated."),
    ("Total Till Date",
     "Type the cumulative quantity completed up to this Bill. The Since Last and "
     "Amount columns update automatically."),
    ("Since Last", "Calculated: Total Till Date − Deduct Previous."),
    ("Amount (₹)", "Calculated: Since Last × Rate."),
], col1="Column")

add_para("Summary box at the bottom:", bold=True)
for b in [
    "Items count.",
    "Deduct (Previous) — total amount already billed.",
    "Net Bill Amount — what is payable in this bill.",
]:
    add_bullet(b)

add_heading("8.4 Action buttons", level=2)
kv_table([
    ("Save Bill (green)", "Saves the Bill against the parent Workslip."),
    ("Download Bill (red)",
     "Opens a modal with two options: ‘N & Part Bill’ (partial) or ‘N & Final Bill’ (final). Pick one to download."),
    ("Download L.S Form (blue)",
     "Lump-Sum form — modal with ‘L.S Form (Part)’ and ‘L.S Form (Final)’ choices."),
    ("Covering Letter (amber)",
     "Modal with ‘Covering Letter (Part)’ and ‘Covering Letter (Final)’ — uses your "
     "uploaded Covering Letter template."),
    ("Back (outline)", "Returns to the Saved Works detail page."),
], col1="Button")

add_callout("Bill types — Part vs Final",
    "Part Bill is a partial / interim bill raised before completion. Final Bill is the "
    "closing bill once the work is done. The system automatically labels and numbers "
    "each accordingly.")

page_break()

# ======================== 9. AMC ========================
add_heading("9. AMC Module", level=1)
add_para(
    "AMC = Annual Maintenance Contract. The page is functionally identical to the "
    "Estimate items page — same 3-panel layout, same search, same drag-to-reorder."
)

add_heading("9.1 Workflow", level=2)
add_numbered("From the Dashboard, click the AMC tile.")
add_numbered("Choose Original or Repair work type.")
add_numbered("Choose AMC category (AMC Electrical, AMC Civil, …).")
add_numbered("Pick groups → tick items → type the annual quantity → drag to reorder.")
add_numbered("Click Save Work and download the AMC Excel / Specification Report / Forwarding Letter.")

add_heading("9.2 What is different from a normal Estimate", level=2)
for b in [
    "The header reads ‘AMC {category} → {group}’ rather than the plain category name.",
    "Items are sourced from the AMC backend rate book (different rates and items than the standard SOR).",
    "Specification Reports include maintenance schedules where applicable.",
]:
    add_bullet(b)

page_break()

# ======================== 10. TEMP WORKS ========================
add_heading("10. Temporary Works Module", level=1)
add_para(
    "For short-term setups: temporary lighting, temporary power, scaffolding, hoardings. "
    "Costing combines a Quantity AND a Days input per item (day-rate based)."
)

add_heading("10.1 Workflow", level=2)
add_numbered("From the Dashboard, click Temporary Works.")
add_numbered("Pick the category (Temporary Electrical or Temporary Civil).")
add_numbered("Use the 3-panel item picker — same as Estimate.")
add_numbered("For each selected item, fill BOTH inputs in the right panel: Qty AND Days.")
add_numbered("Optionally override the Grand Total field at the bottom; tax breakdown recalculates.")
add_numbered("Save Work and download.")

add_heading("10.2 What the Days field does", level=2)
for b in [
    "‘Days’ defaults to 1.",
    "Day rates are stored per item — the system picks the correct rate for the entered number of days.",
    "Amount = Qty × Rate (for that number of days).",
]:
    add_bullet(b)

add_callout("Example",
    "If you select ‘Temporary Generator 25 kVA’ with Qty=2 and Days=15, the system "
    "uses the 15-day rate for that item and multiplies by 2 generators.")

page_break()

# ======================== 11. SELF-FORMATTED ========================
add_heading("11. Self-Formatted Forms Module", level=1)
add_para(
    "Self-Formatted Forms is a placeholder-substitution tool. You upload a Source File "
    "(the data) and a Template File (a Word/Excel with {{PLACEHOLDERS}}). The system "
    "fills the placeholders and gives you the merged document."
)

add_heading("11.1 Page layout — two cards", level=2)
add_para("Left card — Quick Generate (one-off output)", bold=True)
kv_table([
    ("Source File (file input)",
     "Excel / Word / PDF / image / TXT / CSV containing the data to extract. Even "
     "blurred scans are supported."),
    ("Template File (file input, required)",
     "Word or Excel file containing your placeholders, like {{NAME}} or {{AMOUNT}}."),
    ("Custom Placeholders (optional)",
     "Each row maps a placeholder name (e.g. ‘DIVISION_NAME’) to a label to find in "
     "the source (e.g. ‘Name of Division’). Click + to add more rows."),
    ("Preview (eye icon)", "Shows a preview of the merged document before downloading."),
    ("Generate (download icon)", "Builds the merged file and downloads it."),
    ("Placeholders (lightbulb)", "Shows the list of placeholders detected in your template."),
    ("How it works? (question)", "Opens an in-app guide with examples."),
], col1="Field / Button")

add_para("Right card — Save Reusable Format", bold=True)
for b in [
    "Format name and (optional) description.",
    "Upload the Format template file.",
    "Same placeholder mapping as the Quick Generate card.",
    "Save Format — stores it permanently so you can reuse it across projects.",
]:
    add_bullet(b)

add_heading("11.2 When to use it", level=2)
for b in [
    "Letterheads, agreements, indents — anything where the layout is fixed and only data changes.",
    "Recurring monthly returns — save the format once, just change the source data each month.",
    "Forms that have to match a department’s exact wording — paste the wording into the template, mark the variables with {{PLACEHOLDERS}}.",
]:
    add_bullet(b)

page_break()

# ======================== 12. SAVED WORKS ========================
add_heading("12. My Saved Works", level=1)
add_para(
    "My Saved Works is your file manager for everything you have created. It looks "
    "and behaves like Windows Explorer."
)

add_heading("12.1 Page layout", level=2)
kv_table([
    ("Left sidebar — Folders",
     "Your folder tree. Each folder shows its name, colour and item-count badge. "
     "Click to expand, click a folder name to open it."),
    ("Top toolbar",
     "Breadcrumb of your current folder path, plus a New Folder button (blue). When "
     "you tick one or more works, batch buttons appear: Move, Copy, Delete, Clear."),
    ("Status filters",
     "Quick links to All / In Progress / Completed / Archived works."),
    ("Main grid",
     "Folder cards (folder icon) and Work cards (estimate / workslip / bill icons). "
     "Each card shows name, type badge and status badge."),
], col1="Area")

add_heading("12.2 Folder operations", level=2)
for b in [
    "New Folder (toolbar) — creates a folder under the current location.",
    "Right-click a folder card — context menu: Rename, Delete, Move, Change Colour.",
    "Double-click a folder — opens it.",
]:
    add_bullet(b)

add_heading("12.3 Work operations (right-click a work card)", level=2)
for b in [
    "Edit / Resume — re-opens the work in its module so you can continue.",
    "Download — direct download of the work’s primary export.",
    "Move — pick another folder.",
    "Delete — removes the work (with confirmation; you can choose whether children are deleted too).",
]:
    add_bullet(b)

add_heading("12.4 The detail page — E → W → B workflow bar", level=2)
add_para(
    "Click any saved work to open its detail page. At the top is the workflow bar:"
)
kv_table([
    ("E (blue)", "The Estimate. Shows a checkmark if completed."),
    ("W1, W2, W3 … (green)",
     "Each existing Workslip. Click to open. Each shows a status badge."),
    ("+W1, +WN (blue)",
     "Generate the next Workslip. +W1 appears when there are no workslips yet; +WN "
     "appears once the previous workslip is completed."),
    ("B1, B2, B3 … (red)", "Each existing Bill — click to open."),
    ("+B1, +BN (red)",
     "Generate the next Bill. +B1 appears when at least one workslip exists; +BN "
     "appears after the previous bill is saved."),
    ("Locked (gray, lock icon)",
     "You don’t have a subscription for that module. Contact support to enable."),
], col1="Button")

add_callout("Why use this bar",
    "The +W and +B buttons keep the parent–child links clean. Numbers stay correct, "
    "items inherit from the parent, and Saved Works can show the full project history.")

add_heading("12.5 Batch actions", level=2)
for b in [
    "Tick the checkboxes on multiple work cards.",
    "Move (folder-symlink icon) — pick a destination folder for all selected.",
    "Copy (files icon) — duplicates them into another folder.",
    "Delete (red trash) — deletes all selected.",
    "Clear (X) — clears the current selection.",
]:
    add_bullet(b)

page_break()

# ======================== 13. CUSTOM ITEMS ========================
add_heading("13. My Custom Items", level=1)
add_para(
    "Use this when an item you bill regularly is not in the standard SOR catalogue. "
    "Custom items appear in the item picker for the relevant category alongside the "
    "standard items."
)
for b in [
    "Add new — code, description, unit, default rate, default quantity.",
    "Edit — change any field.",
    "Delete — removes it; existing saved works keep their data.",
    "Custom items are private to your account.",
]:
    add_bullet(b)

page_break()

# ======================== 14. TEMPLATES & LETTER SETTINGS ========================
add_heading("14. Document Templates & Letter Settings", level=1)

add_heading("14.1 Letter Settings", level=2)
add_para("Settings → Letter Settings — configure once, use everywhere:", bold=True)
for b in [
    "Company letterhead — name, address, phone, email, GST, logo.",
    "Signatory names and titles (used on Bills and Workslips).",
    "Footer text and standard terms & conditions.",
]:
    add_bullet(b)

add_heading("14.2 User Document Templates", level=2)
add_para("Settings → Templates — upload your own .docx templates:", bold=True)
for b in [
    "Covering Letter — attached when you click ‘Covering Letter (Part / Final)’ on a Bill.",
    "Movement Slip — material movement certificate.",
]:
    add_bullet(b)
add_para("For each template you can: View, Download, Replace, Delete, Set Active.", italic=True)

page_break()

# ======================== 15. EXPORTS ========================
add_heading("15. Exports & Downloads", level=1)
add_para(
    "Large files are generated in the background. A progress indicator appears while "
    "the file is being built; the download starts automatically when it is ready."
)
kv_table([
    ("Estimate Excel", "Item list, Qty × Rate, totals, ready for billing."),
    ("Workslip Excel", "Workslip-format Excel for site use."),
    ("Bill PDF (Part / Final)", "Formatted bill with measurement details and signatures."),
    ("L.S Form (Part / Final)", "Lump-Sum form variant of the bill."),
    ("Specification Report", "Item-wise technical specifications and standards."),
    ("Forwarding Letter", "Covering letter using your Letter Settings."),
    ("Covering Letter (Part / Final)", "Cover letter using your uploaded template."),
    ("AMC Excel / PDF", "AMC contract document with maintenance schedule."),
    ("Temporary Works Excel", "Day-rate based temporary works estimate."),
    ("Self-Formatted Output", "Your template merged with the source data."),
], col1="Document")

page_break()

# ======================== 16. ACCOUNT ========================
add_heading("16. Account, Notifications, Payments", level=1)

add_heading("16.1 Profile & Password", level=2)
for b in [
    "Settings → edit name, email, phone.",
    "Change password (old + new + confirm).",
    "View subscription / plan details and renewal date.",
    "Session management — see active devices and log out from others.",
]:
    add_bullet(b)

add_heading("16.2 Notifications", level=2)
for b in [
    "Bell icon (top-right) shows unread announcements count.",
    "Click the bell to read or dismiss; click View All for the full history.",
    "Toast messages appear briefly after Save / Delete / Download actions.",
]:
    add_bullet(b)

add_heading("16.3 Payment History", level=2)
for b in [
    "Date, plan/module, amount, payment method.",
    "Invoice download link where available.",
    "Subscription status (Active / Expired / Trial).",
]:
    add_bullet(b)

page_break()

# ======================== 17. TIPS ========================
add_heading("17. Tips, Shortcuts and Best Practices", level=1)
for b in [
    "Always Save Work BEFORE downloading — you can re-open and edit only saved works.",
    "Configure Letter Settings on day one — it propagates to every Forwarding Letter and Bill cover letter.",
    "Upload your Covering Letter and Movement Slip templates before raising bills.",
    "Use folders in My Saved Works to keep one site’s documents together.",
    "Use the +W / +B buttons on the workflow bar to generate next workslips/bills — it preserves parent–child links and numbering.",
    "Drag-to-reorder is your friend — set the printed order before downloading.",
    "Use the search boxes in the item picker; don’t scroll long lists manually.",
    "For Repair projects, always pick the Repair work type at Step 1 so prefixes apply automatically.",
    "Use My Custom Items for parts you bill repeatedly that are not in the SOR.",
    "Clear All only when you really mean it — there is no undo for Clear All.",
]:
    add_bullet(b)

page_break()

# ======================== 18. FAQ ========================
add_heading("18. FAQ & Support", level=1)
faqs = [
    ("Can I edit a saved Estimate later?",
     "Yes. Open it from My Saved Works and click Edit / Resume."),
    ("How do I add an item that is not in the SOR?",
     "Use My Custom Items. The item will appear in the picker for the chosen category."),
    ("How are Workslips numbered?",
     "Automatically. The first under an Estimate is Workslip-1; subsequent are Workslip-2, Workslip-3, … up to Workslip-10."),
    ("Can I raise more than one Bill against a single Workslip?",
     "Yes. Use Part Bills for partial billing and Final Bill for closing."),
    ("What happens if I delete an Estimate that has Workslips and Bills?",
     "You will be asked whether to delete the children. Confirm only if you also want them removed."),
    ("Where do I change the company name on my bills?",
     "Settings → Letter Settings. The change applies to every new document."),
    ("My large download is taking time — is something wrong?",
     "Large files generate in the background. A progress bar shows status; the file downloads automatically when ready."),
    ("How do I rearrange items in the printed Excel?",
     "Drag rows in the right-panel selected items table. The printed order matches the order on screen."),
    ("Can I reuse a Self-Formatted form I built last month?",
     "Yes. Save it once using the right card on the Self-Formatted page. Next time, pick the saved format and supply only the new source data."),
]
for q, a in faqs:
    add_para(q, bold=True)
    add_para(a)
    doc.add_paragraph()

add_heading("Support", level=2)
for b in [
    "Website: https://hamsvic.com",
    "Email: support@hamsvic.com",
    "In-app: Help Center (sidebar)",
]:
    add_bullet(b)

page_break()

# ======================== APPENDIX A — Supplement Datas Format ========================
add_heading("Appendix A. Supplement Datas Sheet Format", level=1)
add_para(
    "Hamsvic-generated Excel files use specific sheet names that other parts of the "
    "system can recognise:"
)
kv_table([
    ("Estimate output (Estimate / AMC modules)",
     "Two sheets, in this order: ‘Estimate’ first (the formatted estimate table) "
     "and ‘Datas’ second (the raw item blocks, previously named ‘Output’)."),
    ("Workslip output",
     "Two sheets, in this order: ‘WorkSlip’ first (the quantity table); then "
     "‘Supplement Datas N’ (the supplemental item headings, where N is the workslip "
     "number being generated). For Workslip-1 the second sheet is ‘Supplement Datas 1’; "
     "for Workslip-2 it is ‘Supplement Datas 2’; and so on. If a workslip has no "
     "supplemental items, the second sheet is omitted entirely."),
], col1="Output")

add_para(
    "When you upload a previous Workslip Excel into the Workslip module, Hamsvic looks "
    "inside the file for either a ‘Supplement Datas N’ sheet (current format) or a "
    "legacy ‘ItemBlocks’ / ‘Items Blocks’ sheet (older files) so it can pick up your "
    "supplemental item names. If the Excel was generated by Hamsvic itself, the format "
    "is already correct. If you are preparing the file manually (or editing one from "
    "another source), it MUST follow the rules below — otherwise the upload will not "
    "detect your items and supplemental rows will be missing in the output."
)

add_callout(
    "What is the Supplement Datas sheet?",
    "It is a second sheet inside the workslip workbook that lists each supplemental "
    "item heading and the rows that belong to it. The main WorkSlip sheet holds the "
    "quantities; the Supplement Datas sheet holds the item titles in a colour format "
    "that the parser can recognise.")

add_heading("A.1 Sheet name", level=2)
add_para(
    "The parser recognises two naming patterns (case is ignored):"
)
for b in [
    "Current format — the title contains both ‘supplement’ and ‘data’. Examples: "
    "‘Supplement Datas 1’, ‘Supplement Datas 2’, ‘Supplemental Data 3’.",
    "Legacy format — the title contains both ‘item’ and ‘block’. Examples: "
    "‘ItemBlocks’, ‘Items Blocks’, ‘Item Block’.",
]:
    add_bullet(b)
add_para(
    "Names without one of the two keyword pairs (e.g. just ‘Supplement’ or just "
    "‘Datas’) are NOT recognised.",
    italic=True,
)

add_heading("A.2 The colour rule (most important)", level=2)
add_para(
    "Each item heading is identified by ONE specific cell formatting: solid yellow "
    "background AND red font colour, in the SAME cell, in any of columns A to J of "
    "the heading row. The exact colour codes are:"
)
kv_table([
    ("Cell fill (background)",
     "Solid pattern, foreground colour ending in HEX FFFF00 (pure yellow)."),
    ("Font (text) colour",
     "Colour ending in HEX FF0000 (pure red)."),
    ("Cell value",
     "Any non-empty text. This text becomes the item / heading name."),
    ("Columns scanned",
     "A through J only (columns 1 to 10). Cells beyond column J are ignored."),
], col1="Property")
add_callout(
    "How to apply this in Excel",
    "Select the heading cell → Home tab → Fill Color → choose the standard pure "
    "Yellow swatch. Then with the same cell selected → Font Color → choose the "
    "standard pure Red swatch. Type the item name in that cell. That is all the "
    "parser needs to detect a block heading.",
    fill="FFF7E0")

add_heading("A.3 What a ‘block’ is", level=2)
for b in [
    "A heading row is any row that has at least one yellow+red cell (in A..J).",
    "The block belonging to that heading runs from the heading row down to the row immediately ABOVE the next heading row.",
    "If there is no next heading, the block runs to the last used row of the sheet.",
    "Inside the block, you can have any rows you like (description, sub-items, rates) — only the heading row needs the yellow+red colour.",
]:
    add_bullet(b)

add_heading("A.4 Optional ‘SUPPLEMENTAL’ divider", level=2)
add_para(
    "If your sheet contains BOTH normal estimate item headings and supplemental item "
    "headings, you can separate them with a divider row whose text contains the word "
    "‘SUPPLEMENTAL’ (any case). Headings AFTER this divider are treated as supplemental; "
    "headings before it are treated as estimate items."
)
add_para(
    "If there is NO ‘SUPPLEMENTAL’ divider, the system assumes every heading in the "
    "sheet is supplemental. This is the format Hamsvic generates by default — the "
    "‘Supplement Datas N’ sheet contains only supplemental blocks and no divider.",
    italic=True,
)

add_heading("A.5 Row 1 — Name of Work header (Hamsvic-generated files)", level=2)
add_para(
    "Files generated by Hamsvic put a ‘Name of Work’ banner in row 1 (cells A1 to J1 "
    "are merged). This row is decorative — it is NOT a heading because the parser "
    "ignores rows that do not have the yellow+red colour combination. You may keep, "
    "remove, or change it without affecting parsing."
)

add_heading("A.6 Visual layout (example for ‘Supplement Datas 2’)", level=2)
example_table = doc.add_table(rows=1, cols=2)
example_table.style = "Light Grid Accent 1"
hdr = example_table.rows[0].cells
hdr[0].text = "Row"
hdr[1].text = "Cell A content & formatting"
for c in hdr:
    shade_cell(c, "1F497D")
    for run in c.paragraphs[0].runs:
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
example_rows = [
    ("1", "Name of Work: …    (merged A1:J1, decorative)"),
    ("2", "Supply and laying of cable    ←  YELLOW fill + RED font = heading"),
    ("3", "Description / sub-items / rate row (any formatting)"),
    ("4", "More body rows"),
    ("5", "Erection of distribution board    ←  YELLOW fill + RED font = next heading"),
    ("6", "Description / sub-items / rate row"),
    ("7", "SUPPLEMENTAL ITEMS    (optional divider; plain text, any formatting)"),
    ("8", "Cable trench excavation    ←  YELLOW fill + RED font = supplemental heading"),
    ("9", "Body rows for that supplemental item"),
]
for row_num, content in example_rows:
    row = example_table.add_row().cells
    row[0].text = row_num
    row[1].text = content
doc.add_paragraph()

add_heading("A.7 Common mistakes that cause the upload to miss items", level=2)
for b in [
    "Heading cell is yellow but font is BLACK — not detected (red font is required).",
    "Heading text is red but the cell has NO fill — not detected (yellow background is required).",
    "Yellow + red applied in column K or beyond — not detected (only A..J are scanned).",
    "Yellow fill chosen via a theme that does not resolve to FFFF00 — be safe and pick the standard Yellow swatch.",
    "Sheet name does not match either keyword pair — the title must contain ‘supplement’ + ‘data’ OR ‘item’ + ‘block’.",
    "The heading cell is empty — a heading must have non-empty text.",
    "Merged cells where only the merged anchor has the colour — apply yellow+red to the anchor cell BEFORE merging, and ensure the visible value sits in the anchor.",
]:
    add_bullet(b)

add_heading("A.8 The companion WorkSlip sheet", level=2)
add_para(
    "The Workslip workbook also contains a WorkSlip sheet (the data sheet). Although "
    "it is not the Supplement Datas sheet itself, the parser needs a few things in it:"
)
for b in [
    "A header row that contains a ‘Sl’ (Serial) cell in column A, somewhere in rows 1–15. By default this is row 8.",
    "A ‘Description’ column.",
    "An ‘Est Qty’ (estimate quantity) and ‘Est Rate’ (estimate rate) column.",
    "One or more ‘Execution Qty’ and ‘Execution Amount’ columns (for Phase / Workslip stages).",
    "Footer rows containing TP / Grand Total / Deduct Old Material / LC / QC / NAC, with the percentage written as ‘@ N %’ where applicable.",
]:
    add_bullet(b)

add_callout(
    "Quickest way to get a valid template",
    "Download any Excel produced by Hamsvic, open it, replace the headings and bodies "
    "with your own content, and keep the existing yellow+red formatting on the heading "
    "rows. Save and re-upload. This guarantees the format is exactly what the parser "
    "expects.")

doc.add_paragraph()
end = doc.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = end.add_run("— End of Manual —")
r.italic = True
r.font.color.rgb = RGBColor(120, 120, 120)

doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")
